


#CrossGuard AI
#A production Streamlit application for bidirectional cross-domain data filtering
#(Low -> High and High -> Low transfer workflows). The filter engine is fully local
#Heuristic, and an optional organization-supplied custom term list — with NO outbound
#network calls and NO third-party LLM dependency. Includes a configurable synthetic
#data generator for validation/testing and multi-format export (PDF, Word, TXT, JSON).

#Developed by Randy Singh | Kalsnet (KNet) Consulting

#DEPLOYMENT SCOPE: This application is a sensitive-data redaction and pre-screening
#utility for business data, Controlled Unclassified Information (CUI) workflows,
#vendor/cloud handoffs, and internal Low<->High / High<->Low data movement between
#environments. Because the filter engine runs entirely locally with no external API
#calls, it does not have the network-disclosure problem that ruled out earlier LLM-
#based designs. It is still NOT, by itself, an accredited Cross-Domain Solution (CDS)
#under DoDI 8540.01 / NCDSMO Raise-the-Bar — accreditation also requires hardware-
#enforced guard components, formal NCDSMO Lab-Based Security Assessment, and CDTAB/AO
#sign-off, none of which a general-purpose application server provides on its own.
#Do not route actual classified national-security data through this tool without that
#accreditation. See ACCREDITATION_GUIDE.md and the in-app DoD Policy & Approval tab
#for the real-world pathway.


import io
import json
import logging
import os
import random
import re
import uuid
import datetime as dt

import pandas as pd
import streamlit as st

# =========================================================================
# LOGGING (production audit trail — never logs API keys or raw record content)
# =========================================================================
LOG_DIR = os.environ.get("CROSSGUARD_LOG_DIR", "./logs")
os.makedirs(LOG_DIR, exist_ok=True)
logger = logging.getLogger("crossguard")
if not logger.handlers:
    logger.setLevel(logging.INFO)
    _fh = logging.FileHandler(os.path.join(LOG_DIR, "crossguard_audit.log"))
    _fh.setFormatter(logging.Formatter("%(asctime)s | %(levelname)s | %(message)s"))
    logger.addHandler(_fh)

# =========================================================================
# ENVIRONMENT / DEPLOYMENT CONFIG
# No LLM API keys required — the filter engine is fully local.
# =========================================================================
MAX_UPLOAD_MB = int(os.environ.get("CROSSGUARD_MAX_UPLOAD_MB", "25"))
DEPLOY_ENV = os.environ.get("CROSSGUARD_ENV", "production")  # "production" | "staging" | "dev"

# =========================================================================
# PAGE CONFIG
# =========================================================================
st.set_page_config(
    page_title="CrossGuard AI | Cross-Domain Data Filter",
    page_icon="",
    layout="wide",
    initial_sidebar_state="expanded",
)

# =========================================================================
# GLOBAL STYLE
# =========================================================================
st.markdown(
    """
    <style>
    .block-container {padding-top: 1.5rem;}
    .title-line-1 {
        text-align: center;
        color: #0047AB;
        font-weight: 800;
        font-size: 2.2rem;
        margin-bottom: 0;
        letter-spacing: 0.5px;
    }
    .title-line-2 {
        text-align: center;
        color: #0047AB;
        font-weight: 700;
        font-size: 1.05rem;
        margin-top: 0;
    }
    .subtle-divider {
        border-top: 2px solid #0047AB;
        opacity: 0.25;
        margin: 0.6rem 0 1.2rem 0;
    }
    .status-pill {
        display:inline-block; padding:3px 10px; border-radius:12px;
        font-size:0.78rem; font-weight:700; color:white;
    }
    .pill-clean {background:#1a8a3e;}
    .pill-flag {background:#c0392b;}
    .pill-info {background:#0047AB;}
    div[data-testid="stMetricValue"] {color:#0047AB;}
    </style>
    """,
    unsafe_allow_html=True,
)

# =========================================================================
# TITLE BAR
# =========================================================================
_env_badge_color = {"production": "#1a8a3e", "staging": "#b8860b", "dev": "#6e6e73"}.get(DEPLOY_ENV, "#1a8a3e")
st.markdown(
    f"""
    <div class="title-line-1"> CrossGuard AI — Cross-Domain Data Filter</div>
    <div class="title-line-2">Developed by Randy Singh from Kalsnet (KNet) Consulting</div>
    <div style="text-align:center;margin-top:4px;">
        <span style="background:{_env_badge_color};color:white;border-radius:10px;padding:2px 10px;
        font-size:0.7rem;font-weight:700;letter-spacing:0.5px;">{DEPLOY_ENV.upper()} ENVIRONMENT</span>
    </div>
    <div class="subtle-divider"></div>
    """,
    unsafe_allow_html=True,
)
with st.expander(" Deployment Scope & Compliance Notice", expanded=False):
    st.markdown(
        "**What this system is:** an internal, fully local sensitive-data redaction and "
        "pre-screening utility for business data, Controlled Unclassified Information (CUI) "
        "workflows, and Low↔High / High↔Low data transfers between environments (e.g., dev → prod, "
        "on-prem → cloud, internal → vendor). The filter engine runs entirely on this server with "
        "**no outbound network calls and no third-party LLM dependency** — data never leaves this "
        "deployment during filtering."
    )
    st.markdown(
        "**What this system is not:** an accredited Cross-Domain Solution (CDS) under "
        "DoDI 8540.01 / NCDSMO Raise-the-Bar. Removing the external-API dependency closes the "
        "network-disclosure gap, but accreditation also requires hardware-enforced guard "
        "components, an NCDSMO Lab-Based Security Assessment of the filter logic itself, and "
        "CDTAB/AO sign-off — none of which a general-purpose application server provides on its "
        "own. **Do not route actual classified or national-security data through this tool** "
        "without that accreditation. See the *DoD Policy & Approval* tab and `ACCREDITATION_GUIDE.md` "
        "for when full accreditation is required and how to pursue it."
    )

# =========================================================================
# CONSTANTS
# =========================================================================
CLASSIFICATION_LEVELS = ["UNCLASSIFIED", "CONFIDENTIAL", "SECRET", "TOP SECRET"]
CLASSIFICATION_MARKERS = ["TOP SECRET", "SECRET", "CONFIDENTIAL", "RESTRICTED", "NOFORN", "FOUO", "SCI"]

FIRST_NAMES = ["James", "Maria", "Wei", "Fatima", "Liam", "Sara", "Carlos", "Aisha",
               "Noah", "Elena", "Raj", "Olivia", "Hassan", "Mei", "Lucas", "Amara"]
LAST_NAMES = ["Smith", "Johnson", "Chen", "Khan", "Garcia", "Singh", "Müller", "Rossi",
              "Tanaka", "Brown", "Patel", "Kim", "Nguyen", "Ibrahim", "Lopez", "Walker"]
CITIES = ["Ottawa", "Toronto", "Calgary", "Vancouver", "Halifax", "Winnipeg", "Quebec City", "Edmonton"]
DEPARTMENTS = ["Defence Logistics", "Cyber Operations", "Finance & Audit", "Diplomatic Affairs",
               "Border Security", "Health Intelligence", "Procurement", "Critical Infrastructure"]

NOTE_TEMPLATES = [
    "Subject {name} (SSN {ssn}) was reviewed under file ref {ref}. Contact: {email}, {phone}.",
    "Routine status update for {name}. Address on file: {address}, {city}. Marked {marker}.",
    "Procurement note: vendor contacted via {email}. Internal classification: {marker}.",
    "Personnel record {name} updated. Phone {phone}. No additional remarks.",
    "Incident summary regarding {name} ({ssn}). Distribution limited - {marker}.",
    "General correspondence with {email} regarding logistics in {city}. Unremarkable.",
    "Audit trail entry for case {ref}. No personal data referenced. Status: cleared.",
    "Financial reconciliation note. Card on file ending pattern: {cc}. Marked {marker}.",
]

PII_PATTERNS = {
    "SSN": r"\b\d{3}-\d{2}-\d{4}\b",
    "EMAIL": r"\b[\w.\-]+@[\w.\-]+\.\w+\b",
    "PHONE": r"\b\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b",
    "CREDIT_CARD": r"\b(?:\d[ -]?){13,16}\b",
}

DEFAULT_SKIP_PHRASES = {
    "Main St", "Elm Ave", "King St", "Oak Blvd", "New York", "United States",
    "Quebec City",
}

# =========================================================================
# DoD / DEFENSE CDS POLICY & APPROVAL FRAMEWORK (reference content)
# Public, unclassified policy overview only — informational, not a substitute
# for guidance from an Authorizing Official (AO), CDSE, or NCDSMO.
# =========================================================================
CDS_GOVERNANCE = [
    ("DoDI 8540.01 — Cross Domain (CD) Policy",
     "The core DoD instruction governing how Cross Domain Solutions (CDS) are evaluated, authorized, "
     "and deployed. It defines the roles of the Authorizing Official (AO), the Cross Domain Support "
     "Element (CDSE), and the Cross Domain Technical Advisory Board (CDTAB), and requires every CDS to "
     "receive a Cross Domain Solution Authorization (CDSA) in addition to (not instead of) the standard "
     "DoD Risk Management Framework (RMF) process."),
    ("National Cross Domain Strategy and Management Office (NCDSMO)",
     "Stood up in February 2019 under the National Security Agency's national-security-systems authority "
     "(per National Security Directive 42), replacing the former UCDSMO. NCDSMO sets technical security "
     "requirements for CDS used on U.S. national security systems across DoD and the Intelligence "
     "Community, and maintains the approved CDS baseline list."),
    ("Raise the Bar (RTB)",
     "NCDSMO's CDS security framework, first published in 2018, covering design, development, assessment, "
     "and deployment standards for CDS filter components. National Security Memorandum 8 (NSM-8, January "
     "2022) requires agencies operating a CDS on a national security system to report progress toward "
     "RTB compliance."),
    ("CNSSI No. 1253 + CDS Overlay (Appendix E / Attachment 3)",
     "Applies the NIST SP 800-53 security control catalog to national security systems and adds a "
     "CDS-specific overlay describing which confidentiality/integrity/availability controls apply to "
     "Access, Transfer, and Multilevel CDS types."),
    ("DISN Connection Process Guide (DISA)",
     "Once a CDS is authorized, this guide (implementing DoDI 8010.01 and CJCSI 6211.02D) governs how it "
     "is actually connected to DoD networks and registered in DoD's connection-tracking repositories."),
]

CDS_TYPES = [
    "**Access** — lets an authorized user view/work with data from multiple domains without the data itself moving.",
    "**Transfer** — moves data between domains (the use case this app simulates).",
    "**Multilevel (MLS)** — stores data at multiple levels and serves it back according to each user's clearance.",
]

APPROVAL_STEPS = [
    ("Step 0 — Engage the CDSE",
     "Mission owner engages the Cross Domain Support Element early to scope the requirement and check "
     "whether an existing NCDSMO baseline-listed CDS already meets it, before any new build is considered."),
    ("RMF Step 1 — Categorize",
     "The information system and the data types crossing the boundary are categorized for confidentiality, "
     "integrity, and availability impact."),
    ("RMF Step 2 — Select Controls",
     "Security controls are selected from NIST SP 800-53 plus the CNSSI 1253 CDS Overlay, tailored to "
     "whether the CDS is Access, Transfer, or Multilevel."),
    ("RMF Step 3 — Implement",
     "Controls — including the content/data filters — are implemented and documented."),
    ("RMF Step 4 — Assess",
     "An independent Security Control Assessor evaluates the system; for CDS this typically includes an "
     "NCDSMO Lab-Based Security Assessment of the filter logic against Raise-the-Bar requirements."),
    ("RMF Step 5 — Authorize (CDSA)",
     "The Authorizing Official issues a Cross Domain Solution Authorization, informed by CDTAB review — "
     "a separate, additional gate beyond a standard RMF Authorization to Operate."),
    ("Ongoing — Connect, Monitor, Reauthorize",
     "Once authorized, the CDS is registered and connected per the DISN Connection Process Guide, then "
     "subject to continuous monitoring and periodic reassessment."),
]

DIRECTION_NOTES = {
    "Low → High (Upload)": {
        "risk": (
            "Primary risk is introducing malicious or malformed content into a higher-assurance domain "
            "(malware, exploits, format-confusion attacks) rather than disclosure, since the destination "
            "already protects the data at an equal or higher level."
        ),
        "filters": [
            "Default-deny data-type / format allow-listing — only fully-specified, explicitly permitted formats pass.",
            "Deep content/structure validation confirming files truly match their declared format.",
            "Multi-engine malware and static content inspection.",
            "Size/rate limiting and full transaction logging.",
            "Basic dirty-word / sensitive-pattern screening as a safety net, though not the primary control.",
        ],
        "approval_note": (
            "Because the consequence of a filter failure (malware ingress) is generally lower-severity than "
            "a spillage event, Low→High transfer CDS accreditation focuses heavily on integrity and assurance "
            "testing of the content-validation and anti-malware filter chain."
        ),
    },
    "High → Low (Downgrade/Export)": {
        "risk": (
            "Primary risk is unauthorized disclosure ('spillage') of higher-classification or sensitive data "
            "into a domain with a different, larger, or less-cleared population of users. This is the harder "
            "accreditation case."
        ),
        "filters": [
            "Default-deny data-type allow-listing, plus stricter normalization (flatten active content, strip metadata/embedded objects/macros).",
            "Classification/marking inspection — automated detection of banner and portion markings and classification keywords (TOP SECRET / SECRET / CONFIDENTIAL / NOFORN / FOUO, etc.).",
            "PII and sensitive-data detection and redaction, since the destination's user population usually differs.",
            "Releasability and dissemination-control checks (REL TO, NOFORN, ORCON) against the destination domain's authorized recipients.",
            "Mandatory human review / two-person integrity before release for most Raise-the-Bar-conformant downgrade solutions, particularly for free-text or non-deterministic content.",
            "Comprehensive, tamper-evident audit logging of every release decision.",
        ],
        "approval_note": (
            "Given the higher consequence of a filter failure, downgrade/export CDS accreditation requires "
            "more rigorous independent testing (NCDSMO Lab-Based Security Assessment) of the filter logic "
            "before a CDSA is granted, typically with a higher overall assurance level than Low→High transfer."
        ),
    },
}


def render_direction_policy_note(direction):
    """Inline, direction-aware DoD policy/filter-requirement reminder shown next to the direction selector."""
    info = DIRECTION_NOTES.get(direction)
    if not info:
        return
    with st.expander(f" DoD filter & policy requirements — {direction}"):
        st.markdown(f"**Primary risk:** {info['risk']}")
        st.markdown("**Required filter behaviors (Raise-the-Bar style):**")
        for f in info["filters"]:
            st.markdown(f"- {f}")
        st.markdown(f"**Accreditation note:** {info['approval_note']}")

# =========================================================================
# SESSION STATE DEFAULTS
# =========================================================================
DEFAULTS = {
    "synthetic_df": None,
    "synthetic_filtered_df": None,
    "synthetic_findings": [],
    "real_raw": None,
    "real_kind": None,        # "table" or "text"
    "real_filtered": None,
    "real_findings": [],
    "audit_log": [],
}
for k, v in DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = v

if "session_id" not in st.session_state:
    st.session_state.session_id = str(uuid.uuid4())[:8]


def log_audit(action, direction, records, findings_count, engine):
    entry = {
        "timestamp": dt.datetime.now().isoformat(timespec="seconds"),
        "session_id": st.session_state.session_id,
        "action": action,
        "direction": direction,
        "records_processed": records,
        "findings_redacted": findings_count,
        "engine": engine,
    }
    st.session_state.audit_log.append(entry)
    # Persistent, server-side audit trail (content-free — no PII/record data is ever logged).
    logger.info(json.dumps(entry))


# =========================================================================
# SIDEBAR — FILTER ENGINE CONFIGURATION (fully local, no LLM dependency)
# =========================================================================
with st.sidebar:
    st.header(" Filter Engine Configuration")
    st.success("Local deterministic + heuristic engine — no outbound network calls, no third-party LLM dependency.")

    enable_name_heuristic = st.checkbox(
        "Enable name-pattern detection heuristic",
        value=True,
        help="In addition to exact-match PII patterns (SSN/email/phone/card) and classification "
             "markers, scan for likely person names — both an exact match against a configurable "
             "name list and a generic Capitalized-Word-Pair pattern. The generic pattern can "
             "produce false positives (e.g., place names); tune the skip-list below as needed.",
    )

    custom_terms_raw = st.text_area(
        "Custom sensitive terms to redact (one per line, optional)",
        value="",
        height=90,
        placeholder="PROJECT-CHIMERA\nINTERNAL-ONLY\n...",
        help="Organization-specific dirty-word list — project codenames, internal program names, "
             "or any other term that should always be redacted regardless of the built-in patterns.",
    )
    custom_terms = [t.strip() for t in custom_terms_raw.splitlines() if t.strip()]

    skip_phrases_raw = st.text_area(
        "Name-heuristic skip list (one per line)",
        value="\n".join(sorted(DEFAULT_SKIP_PHRASES)),
        height=90,
        help="Two-word capitalized phrases that should NOT be treated as a possible name by the "
             "generic heuristic (e.g., place names, street types).",
    )
    skip_phrases = {t.strip() for t in skip_phrases_raw.splitlines() if t.strip()}

    with st.expander(" How the filter engine works"):
        st.markdown(
            "1. **Exact-pattern pass** — regex match on SSNs, emails, phone numbers, "
            "card-like numbers, and classification banner words (TOP SECRET / SECRET / "
            "CONFIDENTIAL / RESTRICTED / NOFORN / FOUO / SCI).\n"
            "2. **Name-pattern pass** (optional) — known-name-list match plus a generic "
            "Capitalized-Word-Pair heuristic, filtered against the skip list.\n"
            "3. **Custom term pass** — any organization-supplied terms above, matched "
            "case-insensitively.\n\n"
            "All three passes run locally inside this process. Nothing is sent over the network "
            "during filtering."
        )

    st.markdown("---")
    st.caption(f"CrossGuard AI v1.0 · {DEPLOY_ENV} build")


# =========================================================================
# RULE-BASED FILTER ENGINE (fully local — PII, classification markers,
# name-pattern heuristic, and custom organizational terms)
# =========================================================================
_NAME_LIST_PATTERN = (
    r"\b(?:" + "|".join(re.escape(n) for n in FIRST_NAMES) + r")\s+"
    r"(?:" + "|".join(re.escape(n) for n in LAST_NAMES) + r")\b"
)
_GENERIC_NAME_PATTERN = r"\b[A-Z][a-z]{1,15}\s[A-Z][a-z]{1,15}\b"


def rule_based_redact(text, enable_name_heuristic=True, custom_terms=None, skip_phrases=None):
    if text is None:
        return text, []
    text = str(text)
    findings = []
    redacted = text
    custom_terms = custom_terms or []
    skip_phrases = skip_phrases or set()

    for label, pattern in PII_PATTERNS.items():
        matches = re.findall(pattern, redacted)
        if matches:
            findings.append({"type": label, "count": len(matches)})
            redacted = re.sub(pattern, f"[REDACTED-{label}]", redacted)

    for marker in CLASSIFICATION_MARKERS:
        count = len(re.findall(re.escape(marker), redacted, flags=re.IGNORECASE))
        if count:
            findings.append({"type": f"CLASSIFICATION:{marker}", "count": count})
            redacted = re.sub(re.escape(marker), "[CLASSIFICATION-REDACTED]", redacted, flags=re.IGNORECASE)

    if enable_name_heuristic:
        known_matches = re.findall(_NAME_LIST_PATTERN, redacted)
        if known_matches:
            findings.append({"type": "NAME (known-list match)", "count": len(known_matches)})
            redacted = re.sub(_NAME_LIST_PATTERN, "[REDACTED-NAME]", redacted)

        generic_matches = [m for m in re.findall(_GENERIC_NAME_PATTERN, redacted) if m not in skip_phrases]
        if generic_matches:
            findings.append({"type": "NAME (heuristic pattern)", "count": len(generic_matches)})
            for m in set(generic_matches):
                redacted = redacted.replace(m, "[REDACTED-NAME?]")

    for term in custom_terms:
        if not term:
            continue
        count = len(re.findall(re.escape(term), redacted, flags=re.IGNORECASE))
        if count:
            findings.append({"type": f"CUSTOM-TERM:{term}", "count": count})
            redacted = re.sub(re.escape(term), "[REDACTED-CUSTOM]", redacted, flags=re.IGNORECASE)

    return redacted, findings


def redact_dataframe(df, columns_to_scan=None, enable_name_heuristic=True, custom_terms=None, skip_phrases=None):
    """Run the local filter engine over every cell of the given (or all string) columns."""
    df = df.copy()
    all_findings = []
    cols = columns_to_scan or [c for c in df.columns if df[c].dtype == object]
    for col in cols:
        new_vals = []
        for val in df[col]:
            red, finds = rule_based_redact(val, enable_name_heuristic, custom_terms, skip_phrases)
            new_vals.append(red)
            all_findings.extend(finds)
        df[col] = new_vals
    return df, all_findings


# =========================================================================
# SYNTHETIC DATA GENERATOR
# =========================================================================
def generate_synthetic_dataset(n, direction):
    rows = []
    for i in range(n):
        name = f"{random.choice(FIRST_NAMES)} {random.choice(LAST_NAMES)}"
        ssn = f"{random.randint(100,999)}-{random.randint(10,99)}-{random.randint(1000,9999)}"
        email = f"{name.split()[0].lower()}.{name.split()[1].lower()}@example.com"
        phone = f"({random.randint(200,999)}) {random.randint(200,999)}-{random.randint(1000,9999)}"
        address = f"{random.randint(10,9999)} {random.choice(['Main St','Elm Ave','King St','Oak Blvd'])}"
        city = random.choice(CITIES)
        cc = " ".join(str(random.randint(1000, 9999)) for _ in range(4))
        marker = random.choice(CLASSIFICATION_MARKERS)
        ref = f"REF-{random.randint(10000,99999)}"
        note = random.choice(NOTE_TEMPLATES).format(
            name=name, ssn=ssn, email=email, phone=phone, address=address,
            city=city, cc=cc, marker=marker, ref=ref,
        )
        rows.append(
            {
                "record_id": i + 1,
                "full_name": name,
                "ssn": ssn,
                "email": email,
                "phone": phone,
                "address": f"{address}, {city}",
                "department": random.choice(DEPARTMENTS),
                "classification": random.choice(CLASSIFICATION_LEVELS),
                "notes": note,
                "transfer_direction": direction,
            }
        )
    return pd.DataFrame(rows)


# =========================================================================
# EXPORT HELPERS
# =========================================================================
def _safe_latin1(text):
    return str(text).encode("latin-1", "replace").decode("latin-1")


def df_to_json_bytes(df):
    return df.to_json(orient="records", indent=2).encode("utf-8")


def df_to_text_bytes(df, title):
    buf = io.StringIO()
    buf.write(f"{title}\n")
    buf.write("=" * len(title) + "\n\n")
    buf.write(df.to_string(index=False))
    return buf.getvalue().encode("utf-8")


def df_to_docx_bytes(df, title):
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Generated: {dt.datetime.now().isoformat(timespec='seconds')}")
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            cells[i].text = str(row[col])
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue()


def df_to_pdf_bytes(df, title):
    from fpdf import FPDF

    pdf = FPDF(orientation="L", unit="mm", format="A4")
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 10, _safe_latin1(title), ln=True)
    pdf.set_font("Helvetica", "", 7)

    n_cols = max(len(df.columns), 1)
    col_width = 270 / n_cols

    pdf.set_font("Helvetica", "B", 7)
    for col in df.columns:
        pdf.cell(col_width, 7, _safe_latin1(str(col))[:22], border=1)
    pdf.ln()
    pdf.set_font("Helvetica", "", 7)
    for _, row in df.iterrows():
        for col in df.columns:
            pdf.cell(col_width, 7, _safe_latin1(str(row[col]))[:28], border=1)
        pdf.ln()
    out = pdf.output(dest="S")
    if isinstance(out, str):
        out = out.encode("latin-1", "replace")
    return bytes(out)


def text_to_docx_bytes(text, title):
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Generated: {dt.datetime.now().isoformat(timespec='seconds')}")
    for line in text.splitlines():
        doc.add_paragraph(line)
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue()


def text_to_pdf_bytes(text, title):
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 10, _safe_latin1(title), ln=True)
    pdf.set_font("Helvetica", "", 10)
    for line in text.splitlines() or [""]:
        pdf.multi_cell(0, 6, _safe_latin1(line))
    out = pdf.output(dest="S")
    if isinstance(out, str):
        out = out.encode("latin-1", "replace")
    return bytes(out)


def export_block(data, base_filename, title, key_prefix):
    """Render the 4 download buttons (PDF / DOCX / TXT / JSON) for a DataFrame or raw text."""
    is_df = isinstance(data, pd.DataFrame)
    c1, c2, c3, c4 = st.columns(4)

    with c1:
        try:
            payload = df_to_pdf_bytes(data, title) if is_df else text_to_pdf_bytes(data, title)
            st.download_button("⬇️ PDF", payload, file_name=f"{base_filename}.pdf",
                                mime="application/pdf", key=f"{key_prefix}_pdf")
        except Exception as e:
            st.error(f"PDF export error: {e}")

    with c2:
        try:
            payload = df_to_docx_bytes(data, title) if is_df else text_to_docx_bytes(data, title)
            st.download_button("⬇️ Word", payload, file_name=f"{base_filename}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"{key_prefix}_docx")
        except Exception as e:
            st.error(f"Word export error: {e}")

    with c3:
        try:
            payload = df_to_text_bytes(data, title) if is_df else data.encode("utf-8")
            st.download_button("⬇️ Text", payload, file_name=f"{base_filename}.txt",
                                mime="text/plain", key=f"{key_prefix}_txt")
        except Exception as e:
            st.error(f"Text export error: {e}")

    with c4:
        try:
            payload = df_to_json_bytes(data) if is_df else json.dumps({"text": data}, indent=2).encode("utf-8")
            st.download_button("⬇️ JSON", payload, file_name=f"{base_filename}.json",
                                mime="application/json", key=f"{key_prefix}_json")
        except Exception as e:
            st.error(f"JSON export error: {e}")


# =========================================================================
# MAIN TABS
# =========================================================================
tab_synth, tab_real, tab_audit, tab_policy = st.tabs(
    [" Synthetic Data Demo", " Real Data Upload", " Audit Log", " DoD Policy & Approval"]
)

# -------------------------------------------------------------------------
# TAB 1: SYNTHETIC DATA DEMO
# -------------------------------------------------------------------------
with tab_synth:
    st.subheader("Synthetic Data Generator & Filter Demonstration")
    st.write(
        "Generate fake-but-realistic records (with embedded PII and classification markers) "
        "to demonstrate the filter without touching any real data."
    )

    col_a, col_b, col_c = st.columns([3, 2, 1])
    with col_a:
        n_records = st.slider("Number of synthetic records", min_value=0, max_value=300, value=50, step=1)
    with col_b:
        synth_direction = st.selectbox(
            "Transfer direction", ["Low → High (Upload)", "High → Low (Downgrade/Export)"], key="synth_dir"
        )
    with col_c:
        st.write("")
        st.write("")
        gen_clicked = st.button(" Generate", use_container_width=True)

    render_direction_policy_note(synth_direction)

    reset_col, _ = st.columns([1, 5])
    with reset_col:
        if st.button("🗑️ Reset Data", use_container_width=True):
            st.session_state.synthetic_df = None
            st.session_state.synthetic_filtered_df = None
            st.session_state.synthetic_findings = []
            st.rerun()

    if gen_clicked:
        if n_records == 0:
            st.session_state.synthetic_df = pd.DataFrame()
            st.session_state.synthetic_filtered_df = None
            st.info("0 records selected — nothing to generate.")
        else:
            st.session_state.synthetic_df = generate_synthetic_dataset(n_records, synth_direction)
            st.session_state.synthetic_filtered_df = None
            st.success(f"Generated {n_records} synthetic records.")

    if st.session_state.synthetic_df is not None and not st.session_state.synthetic_df.empty:
        st.markdown("**Raw synthetic data (unfiltered):**")
        st.dataframe(st.session_state.synthetic_df, use_container_width=True, height=260)

        run_filter = st.button(" Run Cross-Domain Filter", type="primary")
        if run_filter:
            df = st.session_state.synthetic_df
            with st.spinner("Running local filter engine (PII + classification + name heuristic + custom terms)..."):
                filtered_df, findings = redact_dataframe(
                    df,
                    columns_to_scan=["notes", "address", "email", "ssn", "phone", "full_name"],
                    enable_name_heuristic=enable_name_heuristic,
                    custom_terms=custom_terms,
                    skip_phrases=skip_phrases,
                )

            st.session_state.synthetic_filtered_df = filtered_df
            st.session_state.synthetic_findings = findings
            log_audit("Synthetic data filter", synth_direction, len(df), len(findings), "Local rule-based + heuristic")
            st.success("Filter complete.")

    if st.session_state.synthetic_filtered_df is not None:
        st.markdown("**Filtered output:**")
        n_flags = len(st.session_state.synthetic_findings)
        pill = "pill-flag" if n_flags else "pill-clean"
        label = f"{n_flags} item(s) redacted" if n_flags else "No sensitive data detected"
        st.markdown(f'<span class="status-pill {pill}">{label}</span>', unsafe_allow_html=True)
        st.dataframe(st.session_state.synthetic_filtered_df, use_container_width=True, height=260)

        if st.session_state.synthetic_findings:
            with st.expander(" Findings detail"):
                st.json(st.session_state.synthetic_findings)

        st.markdown("**Download filtered data:**")
        export_block(
            st.session_state.synthetic_filtered_df,
            "crossguard_synthetic_filtered",
            "CrossGuard AI — Synthetic Filtered Data",
            "synth",
        )

# -------------------------------------------------------------------------
# TAB 2: REAL DATA UPLOAD
# -------------------------------------------------------------------------
with tab_real:
    st.subheader("Real Data Upload & Filter")
    st.warning(
        "IMPORTANT-NOTE: This system is approved for business data, Controlled Unclassified Information (CUI), and "
        "other sensitive-but-unclassified records. It is **not** an accredited Cross-Domain Solution — "
        "do not upload actual classified national-security information here. See *DoD Policy & Approval* "
        "and `ACCREDITATION_GUIDE.md` if your use case involves classified data or a national security system."
    )

    real_direction = st.selectbox(
        "Transfer direction", ["Low → High (Upload)", "High → Low (Downgrade/Export)"], key="real_dir"
    )
    render_direction_policy_note(real_direction)

    uploaded = st.file_uploader(
        "Upload a file to filter", type=["csv", "json", "txt", "xlsx"], key="real_uploader"
    )

    reset_col2, _ = st.columns([1, 5])
    with reset_col2:
        if st.button(" Reset Real Data", use_container_width=True):
            st.session_state.real_raw = None
            st.session_state.real_kind = None
            st.session_state.real_filtered = None
            st.session_state.real_findings = []
            st.rerun()

    if uploaded is not None:
        size_mb = uploaded.size / (1024 * 1024)
        if size_mb > MAX_UPLOAD_MB:
            st.error(
                f"File is {size_mb:.1f} MB, which exceeds the {MAX_UPLOAD_MB} MB limit for this "
                f"deployment (CROSSGUARD_MAX_UPLOAD_MB). Ask an administrator to raise the limit if needed."
            )
            logger.warning(f"session={st.session_state.session_id} upload_rejected_oversize size_mb={size_mb:.1f} filename_redacted=true")
        else:
            try:
                if uploaded.name.endswith(".csv"):
                    st.session_state.real_raw = pd.read_csv(uploaded)
                    st.session_state.real_kind = "table"
                elif uploaded.name.endswith(".xlsx"):
                    st.session_state.real_raw = pd.read_excel(uploaded)
                    st.session_state.real_kind = "table"
                elif uploaded.name.endswith(".json"):
                    content = json.load(uploaded)
                    st.session_state.real_raw = pd.json_normalize(content) if isinstance(content, list) else pd.json_normalize([content])
                    st.session_state.real_kind = "table"
                else:  # .txt
                    st.session_state.real_raw = uploaded.read().decode("utf-8", errors="replace")
                    st.session_state.real_kind = "text"
                st.session_state.real_filtered = None
                st.success(f"Loaded '{uploaded.name}' ({size_mb:.2f} MB).")
                logger.info(f"session={st.session_state.session_id} upload_loaded size_mb={size_mb:.2f} kind={st.session_state.real_kind}")
            except Exception as e:
                st.error(f"Could not read file: {e}")
                logger.error(f"session={st.session_state.session_id} upload_parse_error error={e}")

    if st.session_state.real_raw is not None:
        st.markdown("**Raw uploaded data (unfiltered):**")
        if st.session_state.real_kind == "table":
            st.dataframe(st.session_state.real_raw, use_container_width=True, height=260)
        else:
            st.text_area("Raw text", st.session_state.real_raw, height=200, disabled=True)

        if st.button(" Run Cross-Domain Filter", type="primary", key="real_filter_btn"):
            if st.session_state.real_kind == "table":
                df = st.session_state.real_raw
                with st.spinner("Running local filter engine..."):
                    filtered, findings = redact_dataframe(
                        df,
                        enable_name_heuristic=enable_name_heuristic,
                        custom_terms=custom_terms,
                        skip_phrases=skip_phrases,
                    )

                st.session_state.real_filtered = filtered
                st.session_state.real_findings = findings
                log_audit("Real data filter (table)", real_direction, len(df), len(findings), "Local rule-based + heuristic")

            else:
                text = st.session_state.real_raw
                with st.spinner("Running local filter engine..."):
                    final_text, findings = rule_based_redact(
                        text,
                        enable_name_heuristic=enable_name_heuristic,
                        custom_terms=custom_terms,
                        skip_phrases=skip_phrases,
                    )

                st.session_state.real_filtered = final_text
                st.session_state.real_findings = findings
                log_audit("Real data filter (text)", real_direction, 1, len(findings), "Local rule-based + heuristic")

            st.success("Filter complete.")

    if st.session_state.real_filtered is not None:
        st.markdown("**Filtered output:**")
        n_flags = len(st.session_state.real_findings)
        pill = "pill-flag" if n_flags else "pill-clean"
        label = f"{n_flags} item(s) redacted" if n_flags else "No sensitive data detected"
        st.markdown(f'<span class="status-pill {pill}">{label}</span>', unsafe_allow_html=True)

        if isinstance(st.session_state.real_filtered, pd.DataFrame):
            st.dataframe(st.session_state.real_filtered, use_container_width=True, height=260)
        else:
            st.text_area("Filtered text", st.session_state.real_filtered, height=200, disabled=True)

        if st.session_state.real_findings:
            with st.expander(" Findings detail"):
                st.json(st.session_state.real_findings)

        st.markdown("**Download filtered data:**")
        export_block(
            st.session_state.real_filtered,
            "crossguard_real_filtered",
            "CrossGuard AI — Real Data Filtered Output",
            "real",
        )

# -------------------------------------------------------------------------
# TAB 3: AUDIT LOG
# -------------------------------------------------------------------------
with tab_audit:
    st.subheader("Filter Run Audit Log")
    if st.session_state.audit_log:
        audit_df = pd.DataFrame(st.session_state.audit_log)
        st.dataframe(audit_df, use_container_width=True, height=300)
        st.download_button(
            "⬇️ Download Audit Log (CSV)",
            audit_df.to_csv(index=False).encode("utf-8"),
            file_name="crossguard_audit_log.csv",
            mime="text/csv",
        )
        if st.button(" Clear Audit Log"):
            st.session_state.audit_log = []
            st.rerun()
    else:
        st.info("No filter runs logged yet. Run a filter in the Synthetic or Real Data tabs to populate this log.")

# -------------------------------------------------------------------------
# TAB 4: DoD POLICY & APPROVAL FRAMEWORK
# -------------------------------------------------------------------------
with tab_policy:
    st.subheader("DoD / Defense Cross-Domain Solution (CDS) Policy & Approval Framework")
    st.caption(
        "Reference overview of the public, unclassified U.S. DoD / national-security-system policy "
        "landscape this kind of filter sits within. This tab is informational — it does not grant, "
        "imply, or substitute for an accreditation decision. See `ACCREDITATION_GUIDE.md` for the full "
        "pathway and engagement contacts."
    )

    with st.expander(" Do you need full CDS accreditation for your use case? (start here)", expanded=True):
        st.markdown(
            "**Full DoDI 8540.01 / NCDSMO accreditation is required if *any* of these are true:**\n"
            "- The system moves data between two different classification levels, or between domains "
            "with different releasability/dissemination controls, on a national security system (NSS).\n"
            "- The system is operated by, or on behalf of, DoD or an Intelligence Community element and "
            "touches classified information at any point.\n"
            "- The receiving or sending domain is governed by CNSS policy.\n\n"
            "**Accreditation is *not* the applicable framework, and this application's current "
            "architecture may fit, if instead:**\n"
            "- The data is business-sensitive, proprietary, or Controlled Unclassified Information (CUI) "
            "moving between unclassified environments (e.g., dev→prod, on-prem→cloud, internal→vendor) — "
            "governed instead by NIST SP 800-171 / 800-53 and your organization's data-handling policy.\n"
            "- You are pre-screening or masking data *before* it would ever reach an accredited boundary, "
            "with a human/AO sign-off step still in front of the actual domain transfer.\n\n"
            "If you are unsure which bucket you're in, treat the data as requiring accreditation until "
            "your AO or CDSE confirms otherwise — do not make that call unilaterally."
        )

    st.markdown("### Governing policy & authorities")
    for name, desc in CDS_GOVERNANCE:
        with st.expander(name):
            st.write(desc)

    st.markdown("### CDS types")
    for t in CDS_TYPES:
        st.markdown(f"- {t}")

    st.markdown("### Approval / authorization process (DoDI 8540.01 + DoD RMF)")
    st.caption("A CDS does not receive a standard RMF Authorization to Operate alone — it must also earn a Cross Domain Solution Authorization (CDSA).")
    steps_df = pd.DataFrame(APPROVAL_STEPS, columns=["Step", "What happens"])
    st.table(steps_df)

    with st.expander(" What goes into a CDSA package (typical artifacts)"):
        st.markdown(
            "- System Security Plan (SSP) describing the CDS architecture and data flows\n"
            "- Security control implementation evidence (NIST SP 800-53 + CNSSI 1253 CDS Overlay)\n"
            "- Test & Evaluation Master Plan (TEMP) and test results\n"
            "- Raise-the-Bar conformance evidence for the filter/guard components used\n"
            "- NCDSMO Lab-Based Security Assessment results\n"
            "- Concept of Operations (CONOPS) for the transfer direction(s) in use\n"
            "- Risk assessment and residual-risk acceptance by the AO\n"
            "- CDTAB recommendation memo\n\n"
            "Exact requirements vary by Component and CDS type — confirm the current checklist with your CDSE."
        )

    st.markdown("### Filter requirement policy by transfer direction")
    pcol1, pcol2 = st.columns(2)
    for pcol, direction in zip([pcol1, pcol2], DIRECTION_NOTES.keys()):
        with pcol:
            st.markdown(f"#### {direction}")
            info = DIRECTION_NOTES[direction]
            st.markdown(f"**Primary risk:** {info['risk']}")
            st.markdown("**Required filter behaviors:**")
            for f in info["filters"]:
                st.markdown(f"- {f}")
            st.info(info["approval_note"])

    with st.expander(" Where CrossGuard AI fits — and where it doesn't", expanded=False):
        st.markdown(
            "This application's filter engine runs entirely locally with no outbound network "
            "calls, which removes the specific disclosure risk an LLM-API-based design would "
            "carry. That's real, but **not sufficient for CDS accreditation on its own.**\n\n"
            "**What it is today:** a pre-screening / business-data tool, appropriate for CUI and "
            "business-sensitive data, or as a pre-screening layer that still hands off to a "
            "human/AO decision in front of an actual accredited boundary.\n\n"
            "**What it would take to become part of an accredited CDS:**\n"
            "- The filter would need to run *inside* the accredited boundary, on hardware/software "
            "itself assessed via NCDSMO Lab-Based Security Assessment — not a general-purpose server.\n"
            "- The filter logic would need independent assessment of completeness and false-negative "
            "rate; heuristic pattern matching (like the name-detection heuristic here) is unlikely to "
            "clear the bar for classified-spillage prevention on its own.\n"
            "- The full CDSA package (SSP, TEMP, CONOPS, risk acceptance, CDTAB review) would still "
            "be required.\n"
            "- Depending on CDS type, hardware-enforced guards/diodes may be required in addition to "
            "software filtering.\n\n"
            "See `ACCREDITATION_GUIDE.md` for the complete writeup."
        )

    st.markdown("### Who to engage")
    st.markdown(
        "- **Your Component's Cross Domain Support Element (CDSE)** — first stop for scoping and for "
        "checking the NCDSMO baseline list.\n"
        "- **Authorizing Official (AO)** — owns the final risk-acceptance and CDSA decision.\n"
        "- **NCDSMO** — sets and assesses against Raise-the-Bar technical requirements; documentation is "
        "distributed via Intelink-U to cleared personnel.\n"
        "- **DISA** (via the DISN Connection Process Guide) — handles network connection once authorized."
    )

    st.markdown("---")
    st.caption(
        "Sources (public/unclassified): DoDI 8540.01 Cross Domain Policy; NCDSMO Raise the Bar program; "
        "CNSSI No. 1253 + CDS Overlay; DISA DISN Connection Process Guide. This tab is a reference summary "
        "and does not constitute, and cannot substitute for, an accreditation decision."
    )

st.markdown("---")
st.caption(f"CrossGuard AI v1.0 · {DEPLOY_ENV} · session {st.session_state.session_id} · Randy Singh · Kalsnet (KNet) Consulting")
