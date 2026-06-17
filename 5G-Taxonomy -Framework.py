# 5G Taxonomy & Framework
import streamlit as st
import pandas as pd
import numpy as np
import random
import json
import io
import datetime
import textwrap
import base64
import hashlib
import zlib
import html
import urllib.request
import urllib.error
import streamlit.components.v1 as components
from PIL import Image as PILImage
# ReportLab for PDF export
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    PageBreak, HRFlowable, Image as RLImage, Preformatted
)
from reportlab.lib.enums import TA_LEFT, TA_CENTER
# python-docx for Word export
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# =============================================================================
# BRAND / THEME CONSTANTS
# =============================================================================
BRAND_NAME = "Kalsnet (KNet) Consulting Group"
AUTHOR_NAME = "Randy Singh"
FRAMEWORK_NAME = "5G Taxonomy & Reference Framework (5G-TRF)"
FRAMEWORK_VERSION = "v2.1 — 2026 Edition"

NAVY = "#0B2E5E"
BLUE = "#1357C7"
LIGHT_BLUE = "#EAF1FC"
ACCENT = "#0E8F8F"
SLATE = "#3A4456"
GOLD = "#C8932B"

st.set_page_config(
    page_title="5G Taxonomy & Framework | Kalsnet (KNet) Consulting Group",
    page_icon="📡",
    layout="wide",
    initial_sidebar_state="expanded",
)

# =============================================================================
# GLOBAL CSS
# =============================================================================
st.markdown(f"""
<style>
    .main .block-container {{
        padding-top: 1.2rem;
        max-width: 1200px;
    }}
    .knet-header-wrap {{
        background: linear-gradient(135deg, {LIGHT_BLUE} 0%, #FFFFFF 65%);
        border: 1px solid #D7E3F7;
        border-radius: 14px;
        padding: 28px 32px 22px 32px;
        margin-bottom: 18px;
        box-shadow: 0 2px 14px rgba(11,46,94,0.06);
    }}
    .knet-title {{
        color: {NAVY};
        font-weight: 800;
        font-size: 2.3rem;
        letter-spacing: -0.5px;
        margin-bottom: 0.15rem;
        line-height: 1.15;
    }}
    .knet-subtitle {{
        color: {BLUE};
        font-weight: 700;
        font-size: 1.15rem;
        margin-bottom: 0.6rem;
    }}
    .knet-meta {{
        color: {SLATE};
        font-size: 0.92rem;
        font-weight: 500;
    }}
    .knet-meta b {{ color: {NAVY}; }}
    .knet-pill {{
        display:inline-block;
        background:{NAVY};
        color:white;
        padding:3px 12px;
        border-radius:20px;
        font-size:0.78rem;
        font-weight:600;
        margin-right:8px;
        letter-spacing:0.3px;
    }}
    .knet-pill-accent {{
        display:inline-block;
        background:{ACCENT};
        color:white;
        padding:3px 12px;
        border-radius:20px;
        font-size:0.78rem;
        font-weight:600;
        margin-right:8px;
        letter-spacing:0.3px;
    }}
    .section-card {{
        background:#FFFFFF;
        border:1px solid #E3E8F0;
        border-radius:12px;
        padding:22px 26px;
        margin-bottom:16px;
        box-shadow: 0 1px 6px rgba(20,30,60,0.04);
    }}
    .section-heading {{
        color:{NAVY};
        font-weight:750;
        font-size:1.35rem;
        border-left:5px solid {BLUE};
        padding-left:12px;
        margin-bottom:10px;
    }}
    .field-term {{
        color:{NAVY};
        font-weight:700;
    }}
    .layer-badge {{
        background:{LIGHT_BLUE};
        color:{NAVY};
        border:1px solid #C9DAF3;
        border-radius:8px;
        padding:10px 14px;
        font-weight:600;
        font-size:0.95rem;
        margin-bottom:6px;
    }}
    .formula-box {{
        background:#0B2E5E0D;
        border-left:4px solid {GOLD};
        padding:10px 16px;
        font-family: "Courier New", monospace;
        font-size:0.92rem;
        border-radius:6px;
        margin:8px 0;
        color:{SLATE};
    }}
    .footer-note {{
        text-align:center;
        color:#8A93A6;
        font-size:0.82rem;
        margin-top:30px;
        padding-top:14px;
        border-top:1px solid #E3E8F0;
    }}
    div[data-testid="stMetricValue"] {{
        color:{NAVY};
    }}
    .stTabs [data-baseweb="tab"] {{
        font-weight:600;
    }}
</style>
""", unsafe_allow_html=True)

# =============================================================================
# HEADER
# =============================================================================
st.markdown(f"""
<div class="knet-header-wrap">
    <div class="knet-title">{FRAMEWORK_NAME}</div>
    <div class="knet-subtitle">A Structured Taxonomy &amp; Reference Framework for 5G Networks</div>
    <div class="knet-meta">
        <span class="knet-pill">Developed by {AUTHOR_NAME}</span>
        <span class="knet-pill-accent">{BRAND_NAME}</span>
        <br><br>
        <b>Framework Version:</b> {FRAMEWORK_VERSION} &nbsp;|&nbsp;
        <b>Domain:</b> Telecommunications — 5G / Next-Gen Networks &nbsp;|&nbsp;
        <b>Publisher:</b> {BRAND_NAME}
    </div>
</div>
""", unsafe_allow_html=True)
st.caption("Professional reference tool for classifying, explaining, and exporting 5G network architecture concepts, sample data, and structural relationships.")

# =============================================================================
# CORE TAXONOMY DATA MODEL
# =============================================================================
TAXONOMY = {
    "L1": {
        "layer": "Layer 1 — Spectrum & Physical Resource Domain",
        "description": "Defines the radio frequency resources 5G operates on and how raw spectrum is classified.",
        "categories": [
            {"name": "Low-Band (Sub-1 GHz)", "detail": "Frequencies below 1 GHz (e.g., 600–850 MHz). Offers the widest coverage and best building penetration but lower peak throughput. Used for broad rural and indoor coverage."},
            {"name": "Mid-Band (1–6 GHz / FR1)", "detail": "Often called 'C-Band' (e.g., 3.3–4.2 GHz). Balances coverage and capacity — the workhorse band for most nationwide 5G deployments."},
            {"name": "High-Band / mmWave (FR2, 24–100 GHz)", "detail": "Millimeter wave spectrum delivering multi-Gbps peak speeds over very short range; highly susceptible to obstruction and rain fade. Used for dense urban hotspots and fixed wireless access."},
            {"name": "Shared & Unlicensed Spectrum", "detail": "CBRS (3.5 GHz in the US) and unlicensed bands enabling private 5G networks and neutral-host deployments without exclusive licensing."},
        ]
    },
    "L2": {
        "layer": "Layer 2 — Radio Access Network (RAN) Domain",
        "description": "Covers how user devices connect to the network over the air interface.",
        "categories": [
            {"name": "gNodeB (gNB)", "detail": "The 5G base station — the radio access node that terminates the air interface (Uu) toward the User Equipment (UE) and connects to the 5G Core."},
            {"name": "Massive MIMO & Beamforming", "detail": "Antenna arrays with dozens/hundreds of elements that steer focused signal beams toward users, multiplying capacity and range without new spectrum."},
            {"name": "Open RAN (O-RAN)", "detail": "A disaggregated RAN architecture splitting baseband (CU/DU) and radio (RU) functions across open, interoperable interfaces, reducing vendor lock-in."},
            {"name": "Non-Standalone (NSA) vs Standalone (SA)", "detail": "NSA anchors 5G radio to an existing 4G LTE core for control signaling; SA uses a fully native 5G Core end-to-end, unlocking low latency and slicing."},
        ]
    },
    "L3": {
        "layer": "Layer 3 — Core Network Domain (5GC)",
        "description": "The service-based, cloud-native core that manages sessions, mobility, policy, and data routing.",
        "categories": [
            {"name": "AMF — Access & Mobility Management Function", "detail": "Handles registration, connection, reachability, and mobility management for the UE; the first point of contact in the 5G Core."},
            {"name": "SMF — Session Management Function", "detail": "Establishes, modifies, and releases PDU (Packet Data Unit) sessions, allocating IP addresses and selecting the UPF."},
            {"name": "UPF — User Plane Function", "detail": "Routes and forwards actual user data traffic; can be deployed at the network edge to minimize latency."},
            {"name": "Service-Based Architecture (SBA)", "detail": "5GC functions expose and consume APIs over HTTP/2, enabling modular, independently scalable microservices instead of monolithic network elements."},
        ]
    },
    "L4": {
        "layer": "Layer 4 — Network Slicing & Virtualization Domain",
        "description": "How a single physical network is logically partitioned into multiple virtual, purpose-built networks.",
        "categories": [
            {"name": "eMBB Slice (Enhanced Mobile Broadband)", "detail": "Optimized for high data throughput — video streaming, AR/VR, broadband-like consumer experiences."},
            {"name": "URLLC Slice (Ultra-Reliable Low-Latency Communication)", "detail": "Optimized for sub-5ms latency and 99.999%+ reliability — autonomous vehicles, remote surgery, industrial automation."},
            {"name": "mMTC Slice (Massive Machine-Type Communication)", "detail": "Optimized for connection density — supporting up to ~1 million devices per km² for IoT sensors and smart-city deployments."},
            {"name": "Network Functions Virtualization (NFV)", "detail": "Decouples network functions from dedicated hardware, running them as software instances on commodity infrastructure (NFVI)."},
        ]
    },
    "L5": {
        "layer": "Layer 5 — Edge, Orchestration & Automation Domain",
        "description": "How compute is distributed and how the network is automated end-to-end.",
        "categories": [
            {"name": "Multi-Access Edge Computing (MEC)", "detail": "Places compute and storage resources physically near the radio edge, cutting round-trip latency for time-critical applications."},
            {"name": "MANO (Management & Orchestration)", "detail": "ETSI NFV-MANO framework coordinating the lifecycle of virtual network functions (VNFs) — onboarding, scaling, healing."},
            {"name": "Closed-Loop Automation / AI-RAN", "detail": "Self-optimizing networks (SON) using AI/ML to detect anomalies and auto-remediate without human intervention."},
            {"name": "Network Slicing Orchestration", "detail": "Coordinates end-to-end slice lifecycle across RAN, transport, and core to meet a tenant's committed SLA."},
        ]
    },
    "L6": {
        "layer": "Layer 6 — Security & Trust Domain",
        "description": "Mechanisms that protect 5G's distributed, virtualized, multi-tenant architecture.",
        "categories": [
            {"name": "SUPI / SUCI Concealment", "detail": "The Subscription Permanent Identifier is never sent in the clear; it is encrypted as a SUCI (Subscription Concealed Identifier) to prevent IMSI-catcher attacks."},
            {"name": "Zero Trust Network Segmentation", "detail": "Each network function and slice is treated as untrusted by default, requiring mutual authentication (typically via TLS and OAuth2) for every API call."},
            {"name": "Security Edge Protection Proxy (SEPP)", "detail": "Filters and protects signaling traffic exchanged between operators' 5G Cores at network boundaries (the N32 interface)."},
            {"name": "Lawful Intercept & Privacy Compliance", "detail": "Regulatory-mandated capabilities for authorized monitoring, balanced against subscriber privacy protections (e.g., GDPR-aligned data handling)."},
        ]
    },
    "L7": {
        "layer": "Layer 7 — Service & Use-Case Domain",
        "description": "How the technical capabilities of 5G translate into deployable services and verticals.",
        "categories": [
            {"name": "Fixed Wireless Access (FWA)", "detail": "5G-based home/business broadband replacing fiber or cable in markets where wired infrastructure is costly to deploy."},
            {"name": "Private 5G / Enterprise Networks", "detail": "Dedicated 5G networks (often on CBRS or licensed spectrum) deployed for a single enterprise — factories, ports, campuses."},
            {"name": "Connected Vehicles & V2X", "detail": "Vehicle-to-Everything communication enabling cars to talk to infrastructure, other vehicles, and pedestrians using URLLC slices."},
            {"name": "Smart Manufacturing / Industry 4.0", "detail": "Combines mMTC sensor density with URLLC control loops to enable wireless robotics and predictive maintenance on the factory floor."},
        ]
    },
}

FIELD_GLOSSARY = [
    {"field": "record_id", "meaning": "Unique synthetic identifier for the simulated 5G session/device record."},
    {"field": "timestamp", "meaning": "Simulated UTC capture time of the record, generated over a rolling 30-day window."},
    {"field": "cell_id", "meaning": "Identifier of the simulated gNodeB/cell sector serving the connection."},
    {"field": "spectrum_band", "meaning": "Frequency layer used: Low-Band, Mid-Band, or mmWave (see Layer 1 of the taxonomy)."},
    {"field": "ran_mode", "meaning": "Whether the connection runs Non-Standalone (NSA, anchored to 4G) or Standalone (SA, native 5G Core)."},
    {"field": "slice_type", "meaning": "Network slice category serving the session: eMBB, URLLC, or mMTC (see Layer 4)."},
    {"field": "device_category", "meaning": "Simulated UE type, e.g., Smartphone, IoT Sensor, FWA Gateway, AR/VR Headset, Connected Vehicle."},
    {"field": "use_case", "meaning": "Business/service use-case the record represents (see Layer 7), e.g., Video Streaming, Industrial Automation."},
    {"field": "signal_rsrp_dbm", "meaning": "Reference Signal Received Power in dBm — simulated radio signal strength (typical range -140 to -44 dBm; closer to 0 is stronger)."},
    {"field": "latency_ms", "meaning": "Simulated end-to-end latency in milliseconds for the session."},
    {"field": "throughput_mbps", "meaning": "Simulated achieved data throughput in megabits per second."},
    {"field": "reliability_pct", "meaning": "Simulated packet/session success rate, expressed as a percentage (relevant to URLLC SLA targets of 99.999%)."},
    {"field": "connected_devices", "meaning": "Number of devices simulated as concurrently attached to this cell/slice (relevant to mMTC density)."},
    {"field": "qos_5qi", "meaning": "Simulated 5G QoS Identifier (5QI) — a standard value mapping to a pre-defined QoS profile (priority, latency budget, error rate)."},
    {"field": "security_profile", "meaning": "Simulated security posture applied to the session, e.g., SUCI-Concealed, Zero-Trust-Segmented, SEPP-Filtered."},
    {"field": "edge_node", "meaning": "Simulated Multi-Access Edge Computing (MEC) node ID handling the session, if applicable."},
]

FORMULAS = [
    {"name": "Shannon Capacity (Theoretical Channel Capacity)", "formula": "C = B × log2(1 + SNR)",
     "explain": "C = max data rate (bps), B = channel bandwidth (Hz), SNR = signal-to-noise ratio. Sets the theoretical ceiling on throughput for a given spectrum allocation — the basis for why wider mid-band/mmWave channels enable higher peak rates."},
    {"name": "Spectral Efficiency", "formula": "η = Throughput (bps) / Bandwidth (Hz)",
     "explain": "Measures how efficiently bits are packed into each Hz of spectrum (bps/Hz). Massive MIMO and advanced beamforming push η higher without consuming more spectrum."},
    {"name": "Latency Budget (URLLC)", "formula": "T_total = T_air + T_proc + T_transport + T_core",
     "explain": "End-to-end latency is the sum of air-interface, processing, transport, and core latencies. URLLC slices minimize T_transport and T_core via MEC placement to hit sub-5ms targets."},
    {"name": "Network Slice Capacity Allocation", "formula": "Σ(slice_i resource share) ≤ 100% of RAN/Core resource pool",
     "explain": "Each slice (eMBB, URLLC, mMTC) is guaranteed a resource share by the orchestrator; the sum of committed shares across co-existing slices cannot exceed total available capacity."},
    {"name": "Connection Density (mMTC)", "formula": "Density = Devices / km²",
     "explain": "IMT-2020 target for mMTC is up to 1,000,000 devices per km², driving the design of lightweight signaling and power-efficient protocols."},
    {"name": "Reliability Target (URLLC)", "formula": "Reliability = (Successful Packets / Total Packets) × 100%",
     "explain": "IMT-2020 URLLC target is 99.999% (the 'five nines'), meaning no more than 1 failure in 100,000 transmissions."},
]

QOS_5QI_TABLE = [
    {"5QI": 1, "Resource Type": "GBR", "Priority": "Highest", "Typical Use": "Conversational Voice (VoNR)"},
    {"5QI": 2, "Resource Type": "GBR", "Priority": "High", "Typical Use": "Conversational Video"},
    {"5QI": 80, "Resource Type": "Non-GBR", "Priority": "Very High", "Typical Use": "Low-Latency eMBB / V2X"},
    {"5QI": 82, "Resource Type": "Delay-Critical GBR", "Priority": "Highest", "Typical Use": "URLLC — Discrete Automation"},
    {"5QI": 86, "Resource Type": "Delay-Critical GBR", "Priority": "Very High", "Typical Use": "V2X Messages"},
    {"5QI": 9, "Resource Type": "Non-GBR", "Priority": "Lowest", "Typical Use": "Best-Effort Data / Browsing"},
]

GLOSSARY_TERMS = [
    {"term": "5G NR", "def": "5G New Radio — the global standard air interface defined by 3GPP for 5G."},
    {"term": "3GPP", "def": "3rd Generation Partnership Project — the standards body defining 5G specifications across releases (e.g., Release 15, 16, 17, 18)."},
    {"term": "gNB", "def": "5G base station (gNodeB) terminating the radio interface to user devices."},
    {"term": "UE", "def": "User Equipment — any device (phone, IoT sensor, router) that connects to the 5G network."},
    {"term": "5GC", "def": "5G Core — the cloud-native, service-based core network defined for 5G Standalone deployments."},
    {"term": "AMF", "def": "Access and Mobility Management Function — manages registration and mobility in the 5G Core."},
    {"term": "SMF", "def": "Session Management Function — manages PDU session establishment and IP allocation."},
    {"term": "UPF", "def": "User Plane Function — forwards user data traffic; deployable at the edge for low latency."},
    {"term": "eMBB", "def": "Enhanced Mobile Broadband — the 5G use-case category optimized for high throughput."},
    {"term": "URLLC", "def": "Ultra-Reliable Low-Latency Communication — the 5G use-case category optimized for latency and reliability."},
    {"term": "mMTC", "def": "Massive Machine-Type Communication — the 5G use-case category optimized for device density (IoT)."},
    {"term": "Network Slicing", "def": "Logically partitioning one physical network into multiple virtual networks, each tuned to a specific service's needs."},
    {"term": "MEC", "def": "Multi-Access Edge Computing — placing compute resources near the radio edge to reduce latency."},
    {"term": "Massive MIMO", "def": "Multiple-Input Multiple-Output antenna technology using many antenna elements to multiply capacity via beamforming."},
    {"term": "O-RAN", "def": "Open RAN — a disaggregated, standardized, multi-vendor approach to building radio access networks."},
    {"term": "NSA / SA", "def": "Non-Standalone (5G radio + 4G core) vs Standalone (fully native 5G Core) deployment modes."},
    {"term": "SBA", "def": "Service-Based Architecture — 5G Core design where network functions communicate via APIs (HTTP/2)."},
    {"term": "NFV", "def": "Network Functions Virtualization — running network functions as software on commodity hardware instead of dedicated appliances."},
    {"term": "MANO", "def": "Management and Orchestration — the ETSI framework for automating the lifecycle of virtualized network functions."},
    {"term": "SUCI / SUPI", "def": "Subscription Concealed/Permanent Identifier — 5G's privacy mechanism encrypting subscriber identity over the air."},
    {"term": "SEPP", "def": "Security Edge Protection Proxy — protects signaling exchanged between operators' 5G Cores."},
    {"term": "5QI", "def": "5G QoS Identifier — a standardized value representing a pre-configured Quality of Service profile."},
    {"term": "FWA", "def": "Fixed Wireless Access — using 5G as a wireless substitute for fixed-line home/business broadband."},
    {"term": "V2X", "def": "Vehicle-to-Everything — communication between vehicles and infrastructure, other vehicles, networks, and pedestrians."},
    {"term": "CBRS", "def": "Citizens Broadband Radio Service — shared 3.5 GHz US spectrum enabling private 5G without exclusive licensing."},
]

# =============================================================================
# SESSION STATE INIT
# =============================================================================
if "synthetic_df" not in st.session_state:
    st.session_state.synthetic_df = None
if "record_count" not in st.session_state:
    st.session_state.record_count = 100
if "data_seed" not in st.session_state:
    st.session_state.data_seed = 42

def reset_all_data():
    """Clears generated synthetic data and resets controls to defaults."""
    st.session_state.synthetic_df = None
    st.session_state.record_count = 100
    st.session_state.data_seed = 42

# =============================================================================
# SYNTHETIC DATA GENERATOR
# =============================================================================
SPECTRUM_BANDS = ["Low-Band (Sub-1GHz)", "Mid-Band (C-Band 3.3-4.2GHz)", "mmWave (FR2 24-100GHz)"]
RAN_MODES = ["NSA", "SA"]
SLICE_TYPES = ["eMBB", "URLLC", "mMTC"]
DEVICE_CATS = ["Smartphone", "IoT Sensor", "FWA Gateway", "AR/VR Headset", "Connected Vehicle", "Industrial Robot Arm"]
USE_CASES = [
    "Video Streaming", "Cloud Gaming", "Remote Surgery", "Industrial Automation",
    "Smart Metering", "Fixed Wireless Broadband", "Autonomous Driving (V2X)",
    "Smart City Sensors", "AR/VR Collaboration", "Asset Tracking"
]
SECURITY_PROFILES = ["SUCI-Concealed", "Zero-Trust-Segmented", "SEPP-Filtered", "Standard-TLS"]

def generate_synthetic_5g_data(n_records: int, seed: int = 42) -> pd.DataFrame:
    """Generates a synthetic 5G record dataset for demonstration purposes only.
    No real network, subscriber, or operator data is used or represented."""
    rng = np.random.default_rng(seed)
    py_rng = random.Random(seed)
    rows = []
    start_time = datetime.datetime(2026, 5, 18, 0, 0, 0)
    for i in range(n_records):
        slice_type = py_rng.choice(SLICE_TYPES)
        band = py_rng.choice(SPECTRUM_BANDS)
        ran_mode = py_rng.choices(RAN_MODES, weights=[0.4, 0.6])[0]
        if slice_type == "URLLC":
            latency = round(float(rng.uniform(0.5, 5.0)), 2)
            throughput = round(float(rng.uniform(5, 100)), 2)
            reliability = round(float(rng.uniform(99.990, 99.9999)), 4)
            qos = py_rng.choice([82, 86, 1])
            device = py_rng.choice(["Connected Vehicle", "Industrial Robot Arm"])
            use_case = py_rng.choice(["Remote Surgery", "Industrial Automation", "Autonomous Driving (V2X)"])
        elif slice_type == "mMTC":
            latency = round(float(rng.uniform(20, 200)), 2)
            throughput = round(float(rng.uniform(0.01, 2)), 3)
            reliability = round(float(rng.uniform(97.0, 99.9)), 3)
            qos = py_rng.choice([9, 80])
            device = "IoT Sensor"
            use_case = py_rng.choice(["Smart Metering", "Smart City Sensors", "Asset Tracking"])
        else:
            latency = round(float(rng.uniform(8, 40)), 2)
            throughput = round(float(rng.uniform(50, 2000)), 2)
            reliability = round(float(rng.uniform(98.5, 99.95)), 3)
            qos = py_rng.choice([2, 9, 80])
            device = py_rng.choice(["Smartphone", "FWA Gateway", "AR/VR Headset"])
            use_case = py_rng.choice(["Video Streaming", "Cloud Gaming", "Fixed Wireless Broadband", "AR/VR Collaboration"])
        rsrp = round(float(rng.uniform(-120, -60)), 1)
        connected_devices = int(rng.integers(1, 50000)) if slice_type == "mMTC" else int(rng.integers(1, 300))
        rows.append({
            "record_id": f"5GREC-{i+1:05d}",
            "timestamp": (start_time + datetime.timedelta(minutes=int(rng.integers(0, 43200)))).strftime("%Y-%m-%d %H:%M:%S"),
            "cell_id": f"gNB-{py_rng.randint(1000,1999)}-S{py_rng.randint(1,3)}",
            "spectrum_band": band,
            "ran_mode": ran_mode,
            "slice_type": slice_type,
            "device_category": device,
            "use_case": use_case,
            "signal_rsrp_dbm": rsrp,
            "latency_ms": latency,
            "throughput_mbps": throughput,
            "reliability_pct": reliability,
            "connected_devices": connected_devices,
            "qos_5qi": qos,
            "security_profile": py_rng.choice(SECURITY_PROFILES),
            "edge_node": f"MEC-{py_rng.randint(1,40):02d}" if py_rng.random() > 0.25 else "N/A",
        })
    return pd.DataFrame(rows)

# =============================================================================
# DIAGRAM REGISTRY (single source of truth for both the live tab and exports)
# =============================================================================
DIAGRAMS = [
    {
        "id": "taxonomy_tree",
        "title": "1. Taxonomy Structure Diagram",
        "description": "A top-down view of how the 7 layers branch into their core categories.",
        "ui_height": 560,
        "code": """flowchart TB
    ROOT["5G Taxonomy and Reference Framework"]
    L1["Layer 1: Spectrum and Physical Resource"]
    L2["Layer 2: Radio Access Network RAN"]
    L3["Layer 3: Core Network 5GC"]
    L4["Layer 4: Network Slicing and Virtualization"]
    L5["Layer 5: Edge Orchestration and Automation"]
    L6["Layer 6: Security and Trust"]
    L7["Layer 7: Service and Use Case"]
    ROOT --> L1
    ROOT --> L2
    ROOT --> L3
    ROOT --> L4
    ROOT --> L5
    ROOT --> L6
    ROOT --> L7
    L1 --> L1a["Low Band"]
    L1 --> L1b["Mid Band"]
    L1 --> L1c["mmWave"]
    L2 --> L2a["gNodeB"]
    L2 --> L2b["Massive MIMO"]
    L2 --> L2c["Open RAN"]
    L3 --> L3a["AMF"]
    L3 --> L3b["SMF"]
    L3 --> L3c["UPF"]
    L4 --> L4a["eMBB Slice"]
    L4 --> L4b["URLLC Slice"]
    L4 --> L4c["mMTC Slice"]
    L7 --> L7a["FWA"]
    L7 --> L7b["Private 5G"]
    L7 --> L7c["V2X"]""",
    },
    {
        "id": "connection_flow",
        "title": "2. Connection Establishment Flow (Sequence Diagram)",
        "description": "Shows the simplified message flow when a device (UE) attaches to the network and establishes a data session, from radio connection through to the data network.",
        "ui_height": 500,
        "code": """sequenceDiagram
    participant UE as User Equipment
    participant gNB as gNodeB Radio
    participant AMF as AMF Mobility
    participant SMF as SMF Session
    participant UPF as UPF Data Plane
    participant DN as Data Network
    UE->>gNB: RRC Connection Request
    gNB->>AMF: Registration Request N2
    AMF->>AMF: Authenticate via SUCI to SUPI
    AMF->>SMF: Create Session Request
    SMF->>UPF: Configure Data Path N4
    UPF->>DN: Establish PDU Session
    DN-->>UPF: Data Response
    UPF-->>gNB: User Plane Data N3
    gNB-->>UE: Radio Bearer Data""",
    },
    {
        "id": "slicing_flow",
        "title": "3. Network Slicing Flow",
        "description": "Illustrates how one shared physical infrastructure is logically divided by the orchestration layer into the three core slice types, each serving distinct use-cases.",
        "ui_height": 460,
        "code": """flowchart LR
    PHY["Shared Physical Infrastructure RAN, Transport, Core"]
    ORCH["Orchestration and MANO Layer"]
    PHY --> ORCH
    ORCH --> S1["eMBB Slice: High Throughput"]
    ORCH --> S2["URLLC Slice: Low Latency"]
    ORCH --> S3["mMTC Slice: High Density"]
    S1 --> U1["Video, Cloud Gaming, AR VR"]
    S2 --> U2["Autonomous Vehicles, Remote Surgery"]
    S3 --> U3["IoT Sensors, Smart Metering"]""",
    },
    {
        "id": "security_flow",
        "title": "4. Identity & Security Flow",
        "description": "Shows how subscriber identity is protected (SUCI/SUPI), how internal functions authenticate under zero trust, and how inter-operator signaling is filtered through the SEPP.",
        "ui_height": 460,
        "code": """flowchart TB
    UE["User Equipment"] -->|"SUCI Encrypted ID"| GNB["gNodeB"]
    GNB --> AMF["AMF: Decrypt SUCI to SUPI"]
    AMF --> AUSF["AUSF: Authentication"]
    AUSF --> UDM["UDM: Subscriber Profile"]
    AMF --> NF1["Internal Network Functions Zero Trust mTLS"]
    NF1 --> SEPP["SEPP: Security Edge Protection Proxy"]
    SEPP --> PARTNER["Partner Operator 5G Core via N32"]""",
    },
]

# =============================================================================
# MERMAID RENDER HELPER (live, in-app rendering)
# =============================================================================
def render_mermaid(mermaid_code: str, height: int = 480, diagram_id: str = None):
    """Renders a Mermaid.js diagram inside the Streamlit app via an HTML component.

    Fixes vs. the previous version:
      - Uses the UMD build (pinned version) instead of an ESM import — simpler and
        more reliably loaded inside Streamlit's sandboxed component iframe.
      - Sets startOnLoad=False and calls mermaid.run() exactly once, removing the
        startOnLoad/run() race condition that could leave diagrams blank.
      - Forces the rendered SVG to scale responsively (max-width:100%; height:auto)
        so wide/tall diagrams are never clipped inside the panel.
      - On any failure (e.g., no internet access to the CDN), shows a clear,
        readable fallback with the raw Mermaid source instead of a blank box.
    """
    code = mermaid_code.strip()
    uid = diagram_id or hashlib.md5(code.encode("utf-8")).hexdigest()[:10]
    cid = f"mermaid-{uid}"
    escaped_code = html.escape(code)

    html_code = f"""
    <div style="background:white; border:1px solid #E3E8F0; border-radius:10px; padding:14px; overflow:auto;">
      <div id="{cid}" class="mermaid" style="display:flex; justify-content:center;">
{code}
      </div>
      <div id="fallback-{cid}" style="display:none;">
        <p style="color:#B23A3A; font-weight:600; font-family:sans-serif; font-size:0.85rem;">
          ⚠️ This diagram could not render visually (e.g., no internet access to load Mermaid.js).
          Raw diagram source is shown below — paste it into
          <a href="https://mermaid.live" target="_blank" rel="noopener">mermaid.live</a> to view it graphically.
        </p>
        <pre style="white-space:pre-wrap; font-family:'Courier New',monospace; font-size:0.78rem; background:#F7F8FA; border:1px solid #E3E8F0; border-radius:6px; padding:10px; color:#3A4456;">{escaped_code}</pre>
      </div>
    </div>
    <style>
      #{cid} svg {{ max-width: 100% !important; height: auto !important; }}
    </style>
    <script src="https://cdn.jsdelivr.net/npm/mermaid@10.9.1/dist/mermaid.min.js"></script>
    <script>
      (function() {{
        function showFallback() {{
          var fb = document.getElementById("fallback-{cid}");
          if (fb) {{ fb.style.display = "block"; }}
        }}
        try {{
          if (typeof mermaid === "undefined") {{
            showFallback();
          }} else {{
            mermaid.initialize({{
              startOnLoad: false,
              theme: "base",
              themeVariables: {{
                primaryColor: "#EAF1FC",
                primaryTextColor: "#0B2E5E",
                primaryBorderColor: "#1357C7",
                lineColor: "#1357C7",
                secondaryColor: "#0E8F8F",
                tertiaryColor: "#FFFFFF"
              }}
            }});
            mermaid.run({{ querySelector: "#{cid}" }}).catch(showFallback);
          }}
        }} catch (e) {{
          showFallback();
        }}
      }})();
    </script>
    """
    components.html(html_code, height=height, scrolling=True)

# =============================================================================
# MERMAID -> STATIC IMAGE HELPER (for PDF / DOCX exports)
# =============================================================================
@st.cache_data(show_spinner=False, ttl=86400)
def fetch_diagram_image(mermaid_code: str, width: int = 1000):
    """Renders Mermaid source to PNG bytes for embedding in static exports (PDF/DOCX).

    Tries the public mermaid.ink rendering service first, then kroki.io as a
    fallback. Returns None if neither is reachable (e.g., no internet access at
    export time) — callers must fall back to including the raw Mermaid source as
    text so the diagram's content is never silently dropped from an export.
    """
    code = mermaid_code.strip()

    # 1) mermaid.ink
    try:
        encoded = base64.urlsafe_b64encode(code.encode("utf-8")).decode("ascii")
        url = f"https://mermaid.ink/img/{encoded}?type=png&width={width}&bgColor=FFFFFF"
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=12) as resp:
            data = resp.read()
            if data:
                return data
    except Exception:
        pass

    # 2) kroki.io fallback
    try:
        compressed = zlib.compress(code.encode("utf-8"), 9)
        encoded = base64.urlsafe_b64encode(compressed).decode("ascii")
        url = f"https://kroki.io/mermaid/png/{encoded}"
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=12) as resp:
            data = resp.read()
            if data:
                return data
    except Exception:
        pass

    return None

# =============================================================================
# SIDEBAR
# =============================================================================
with st.sidebar:
    st.markdown(f"### 📡 {BRAND_NAME}")
    st.markdown(f"**{FRAMEWORK_NAME}**")
    st.caption(f"Developed by {AUTHOR_NAME}")
    st.divider()
    st.markdown("#### 🎛️ Sample Data Controls")
    record_count = st.slider(
        "Number of synthetic 5G records to generate",
        min_value=0, max_value=300,
        value=st.session_state.record_count,
        step=10,
        help="Generates fully synthetic, illustrative 5G session records. No real subscriber or network data is used."
    )
    st.session_state.record_count = record_count
    seed_val = st.number_input(
        "Random seed (for reproducibility)",
        min_value=0, max_value=99999,
        value=st.session_state.data_seed,
        step=1
    )
    st.session_state.data_seed = int(seed_val)
    col_gen, col_reset = st.columns(2)
    with col_gen:
        generate_clicked = st.button("🔄 Generate", use_container_width=True, type="primary")
    with col_reset:
        reset_clicked = st.button("🗑️ Reset Data", use_container_width=True)
    if generate_clicked:
        if record_count == 0:
            st.session_state.synthetic_df = pd.DataFrame()
        else:
            st.session_state.synthetic_df = generate_synthetic_5g_data(record_count, st.session_state.data_seed)
        st.toast(f"Generated {record_count} synthetic 5G records.", icon="✅")
    if reset_clicked:
        reset_all_data()
        st.toast("Sample data and controls have been reset.", icon="🗑️")
        st.rerun()
    st.divider()
    st.markdown("#### 📤 Export Center")
    st.caption("Use the **Export & Reports** tab to download the full framework in PDF, Word, Text, or JSON format.")
    st.divider()
    st.markdown(
        f"<div style='font-size:0.78rem; color:#8A93A6;'>© {datetime.date.today().year} {BRAND_NAME}. "
        f"Synthetic data is illustrative only and does not represent any real operator, subscriber, or device.</div>",
        unsafe_allow_html=True
    )

# =============================================================================
# MAIN NAVIGATION TABS
# =============================================================================
tab_overview, tab_taxonomy, tab_diagrams, tab_data, tab_glossary, tab_export = st.tabs([
    "📘 Overview & Concepts",
    "🧩 Taxonomy Framework",
    "🔀 Flow & Architecture Diagrams",
    "📊 Sample 5G Data",
    "📖 Glossary & Formulas",
    "📤 Export & Reports",
])

# =============================================================================
# TAB 1: OVERVIEW & CONCEPTS
# =============================================================================
with tab_overview:
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<div class="section-heading">What is a "5G Taxonomy & Framework"?</div>', unsafe_allow_html=True)
    st.markdown(f"""
A **taxonomy** is a structured classification system — it organizes a complex domain into clear categories,
sub-categories, and relationships so that anyone can locate where a specific concept, technology, or component
belongs. A **framework** goes one step further: it explains *how those categories interact*, defining the flows,
dependencies, and rules that connect them into a working system.

Put together, the **{FRAMEWORK_NAME}** organizes the entire 5G ecosystem — from raw radio spectrum up to
business-facing services — into **7 structural layers** (a "5G stack"), each broken into concrete categories.
This gives engineers, analysts, and business stakeholders a shared map of 5G: a common vocabulary to discuss
architecture decisions, a reference for compliance and vendor evaluation, and a teaching tool for newcomers to
the technology.

**Why does 5G need this?** 5G is not a single technology — it is a *system of systems* spanning spectrum
policy, radio engineering, cloud-native software, virtualization, security, and vertical-specific services
(manufacturing, automotive, healthcare). Without a taxonomy, conversations about "5G" become ambiguous — a
spectrum engineer and a slicing-orchestration architect may use the same words to mean different things. This
framework removes that ambiguity by assigning every concept a precise place in the structure.
""")
    st.markdown('</div>', unsafe_allow_html=True)

    col1, col2, col3 = st.columns(3)
    with col1:
        st.markdown('<div class="layer-badge">🗂️ 7 Structural Layers</div>', unsafe_allow_html=True)
        st.caption("From physical spectrum to business services")
    with col2:
        st.markdown('<div class="layer-badge">🧱 28 Core Categories</div>', unsafe_allow_html=True)
        st.caption("Concrete technologies & functions per layer")
    with col3:
        st.markdown('<div class="layer-badge">📐 6 Quantitative Formulas</div>', unsafe_allow_html=True)
        st.caption("The math underlying 5G performance claims")

    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<div class="section-heading">How the Framework is Organized</div>', unsafe_allow_html=True)
    st.markdown("""
The framework reads **bottom-up**, mirroring how a 5G connection is actually built — starting from raw spectrum
and ending at the business service the customer experiences:
""")
    layer_summary_rows = []
    for key, layer in TAXONOMY.items():
        layer_summary_rows.append({
            "Layer": layer["layer"],
            "Purpose": layer["description"],
            "# Categories": len(layer["categories"])
        })
    st.dataframe(pd.DataFrame(layer_summary_rows), use_container_width=True, hide_index=True)
    st.markdown('</div>', unsafe_allow_html=True)

    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<div class="section-heading">How to Use This Tool</div>', unsafe_allow_html=True)
    st.markdown("""
1. **Taxonomy Framework tab** — drill into each of the 7 layers and their categories with plain-English explanations.
2. **Flow & Architecture Diagrams tab** — view Mermaid-based diagrams showing how a connection is established, how slicing works, and how security is enforced.
3. **Sample 5G Data tab** — generate between 0 and 300 synthetic 5G session records (no real data), with every field explained, plus summary charts.
4. **Glossary & Formulas tab** — a quick-reference dictionary of every acronym and the key quantitative formulas behind 5G performance targets.
5. **Export & Reports tab** — download the complete framework (including the diagrams and any generated sample data) as **PDF, Word (DOCX), Text, or JSON**.
""")
    st.markdown('</div>', unsafe_allow_html=True)

# =============================================================================
# TAB 2: TAXONOMY FRAMEWORK (drill-down)
# =============================================================================
with tab_taxonomy:
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<div class="section-heading">The 7-Layer 5G Taxonomy</div>', unsafe_allow_html=True)
    st.markdown(
        "Each layer below represents a distinct structural concern in 5G. Expand a layer to see its "
        "categories and a plain-English explanation of each — written so both technical and non-technical "
        "readers can follow."
    )
    st.markdown('</div>', unsafe_allow_html=True)

    layer_icons = {"L1": "📶", "L2": "📡", "L3": "🧠", "L4": "🧬", "L5": "⚙️", "L6": "🔐", "L7": "🏭"}
    for key, layer in TAXONOMY.items():
        icon = layer_icons.get(key, "▶️")
        with st.expander(f"{icon}  **{layer['layer']}**", expanded=(key == "L1")):
            st.markdown(f"*{layer['description']}*")
            st.write("")
            for cat in layer["categories"]:
                st.markdown(f"<span class='field-term'>{cat['name']}</span>", unsafe_allow_html=True)
                st.markdown(f"<div style='margin-bottom:10px; color:#3A4456;'>{cat['detail']}</div>", unsafe_allow_html=True)

    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<div class="section-heading">Layer Composition at a Glance</div>', unsafe_allow_html=True)
    layer_chart_df = pd.DataFrame([
        {"Layer": layer["layer"].split("—")[0].strip(), "Categories": len(layer["categories"])}
        for layer in TAXONOMY.values()
    ])
    st.bar_chart(layer_chart_df.set_index("Layer"), color=BLUE, height=320)
    st.caption("Number of defined categories per taxonomy layer (illustrative count of items in this reference framework).")
    st.markdown('</div>', unsafe_allow_html=True)

# =============================================================================
# TAB 3: FLOW & ARCHITECTURE DIAGRAMS
# =============================================================================
with tab_diagrams:
    for d in DIAGRAMS:
        st.markdown('<div class="section-card">', unsafe_allow_html=True)
        st.markdown(f'<div class="section-heading">{d["title"]}</div>', unsafe_allow_html=True)
        st.markdown(d["description"])
        st.markdown('</div>', unsafe_allow_html=True)
        render_mermaid(d["code"], height=d["ui_height"], diagram_id=d["id"])

    st.info(
        "💡 Diagrams render live using **Mermaid.js**. If a diagram can't render visually in your environment "
        "(e.g., no internet access), a readable fallback with the raw Mermaid source appears in its place — "
        "and every diagram (as an image when possible, otherwise as Mermaid source) is included automatically "
        "in every exported report: PDF, Word, Text, and JSON.",
        icon="💡"
    )

# =============================================================================
# TAB 4: SAMPLE 5G DATA
# =============================================================================
with tab_data:
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<div class="section-heading">Synthetic 5G Record Dataset</div>', unsafe_allow_html=True)
    st.markdown(f"""
Use the **sidebar controls** (0–300 records) to generate a fully **synthetic** dataset of illustrative 5G
session records, modeled on the taxonomy's structural concepts (spectrum band, RAN mode, slice type, QoS, etc.).
This data does **not** represent any real network, operator, or subscriber — it exists purely to demonstrate
how the taxonomy's fields map to real-world telemetry.

Click **🔄 Generate** in the sidebar to create records, or **🗑️ Reset Data** to clear everything back to a
blank state.
""")
    st.markdown('</div>', unsafe_allow_html=True)

    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<div class="section-heading">Field-by-Field Explanation</div>', unsafe_allow_html=True)
    st.markdown("Each column in the dataset corresponds to a concept from the taxonomy above:")
    glossary_df = pd.DataFrame(FIELD_GLOSSARY).rename(columns={"field": "Field", "meaning": "Explanation"})
    st.dataframe(glossary_df, use_container_width=True, hide_index=True, height=420)
    st.markdown('</div>', unsafe_allow_html=True)

    df = st.session_state.synthetic_df
    if df is None:
        st.warning("No sample data generated yet. Use the sidebar to set a record count (0–300) and click **Generate**.", icon="ℹ️")
    elif df.empty:
        st.warning("Record count is set to 0 — no rows to display. Increase the slider and click **Generate**.", icon="ℹ️")
    else:
        st.markdown('<div class="section-card">', unsafe_allow_html=True)
        st.markdown('<div class="section-heading">Generated Records</div>', unsafe_allow_html=True)
        m1, m2, m3, m4 = st.columns(4)
        m1.metric("Total Records", len(df))
        m2.metric("Avg Latency (ms)", f"{df['latency_ms'].mean():.2f}")
        m3.metric("Avg Throughput (Mbps)", f"{df['throughput_mbps'].mean():.1f}")
        m4.metric("Avg Reliability (%)", f"{df['reliability_pct'].mean():.3f}")
        st.dataframe(df, use_container_width=True, hide_index=True, height=380)
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown('<div class="section-card">', unsafe_allow_html=True)
        st.markdown('<div class="section-heading">Visual Summary</div>', unsafe_allow_html=True)
        c1, c2 = st.columns(2)
        with c1:
            st.markdown("**Records by Slice Type**")
            st.bar_chart(df["slice_type"].value_counts(), color=BLUE, height=280)
        with c2:
            st.markdown("**Records by Spectrum Band**")
            st.bar_chart(df["spectrum_band"].value_counts(), color=ACCENT, height=280)
        c3, c4 = st.columns(2)
        with c3:
            st.markdown("**Avg Throughput by Slice Type (Mbps)**")
            st.bar_chart(df.groupby("slice_type")["throughput_mbps"].mean(), color=GOLD, height=280)
        with c4:
            st.markdown("**Avg Latency by Slice Type (ms)**")
            st.bar_chart(df.groupby("slice_type")["latency_ms"].mean(), color=NAVY, height=280)
        st.markdown("**Latency vs. Throughput (scatter, sampled up to 300 pts)**")
        scatter_df = df[["latency_ms", "throughput_mbps", "slice_type"]].copy()
        st.scatter_chart(scatter_df, x="latency_ms", y="throughput_mbps", color="slice_type", height=320)
        st.markdown('</div>', unsafe_allow_html=True)

# =============================================================================
# TAB 5: GLOSSARY & FORMULAS
# =============================================================================
with tab_glossary:
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<div class="section-heading">Glossary of Key Terms</div>', unsafe_allow_html=True)
    st.markdown("A quick-reference dictionary for every acronym and concept used throughout the framework.")
    gdf = pd.DataFrame(GLOSSARY_TERMS).rename(columns={"term": "Term", "def": "Definition"})
    search_term = st.text_input("🔎 Search glossary", placeholder="e.g., URLLC, SUCI, slicing...")
    if search_term:
        mask = gdf["Term"].str.contains(search_term, case=False) | gdf["Definition"].str.contains(search_term, case=False)
        st.dataframe(gdf[mask], use_container_width=True, hide_index=True, height=400)
    else:
        st.dataframe(gdf, use_container_width=True, hide_index=True, height=400)
    st.markdown('</div>', unsafe_allow_html=True)

    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<div class="section-heading">Key Formulas Behind 5G Performance</div>', unsafe_allow_html=True)
    st.markdown("These are the quantitative foundations referenced by the taxonomy's performance-related fields.")
    for f in FORMULAS:
        st.markdown(f"<span class='field-term'>{f['name']}</span>", unsafe_allow_html=True)
        st.markdown(f"<div class='formula-box'>{f['formula']}</div>", unsafe_allow_html=True)
        st.markdown(f"<div style='margin-bottom:14px; color:#3A4456;'>{f['explain']}</div>", unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<div class="section-heading">Illustrative 5QI (QoS Identifier) Reference Table</div>', unsafe_allow_html=True)
    st.markdown(
        "A simplified, illustrative subset of standardized 5QI values used to map a session to a "
        "pre-configured Quality of Service profile (priority, resource type, typical use)."
    )
    st.dataframe(pd.DataFrame(QOS_5QI_TABLE), use_container_width=True, hide_index=True)
    st.caption("Note: This is a simplified educational reference, not the complete 3GPP TS 23.501 5QI table.")
    st.markdown('</div>', unsafe_allow_html=True)

# =============================================================================
# EXPORT BUILDERS
# =============================================================================
def build_export_payload():
    """Assembles the complete framework + diagrams + sample data into one structured
    dict, used as the canonical source for JSON, TXT, DOCX, and PDF exports."""
    df = st.session_state.synthetic_df
    sample_records = df.to_dict(orient="records") if isinstance(df, pd.DataFrame) and not df.empty else []
    payload = {
        "meta": {
            "framework_name": FRAMEWORK_NAME,
            "framework_version": FRAMEWORK_VERSION,
            "developed_by": AUTHOR_NAME,
            "organization": BRAND_NAME,
            "generated_at_utc": datetime.datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S UTC"),
            "sample_record_count": len(sample_records),
            "disclaimer": "All sample 5G records are synthetically generated for illustrative purposes only and do not represent real network, operator, or subscriber data."
        },
        "taxonomy": TAXONOMY,
        "diagrams": [
            {
                "id": d["id"],
                "title": d["title"],
                "description": d["description"],
                "mermaid_code": d["code"],
            }
            for d in DIAGRAMS
        ],
        "field_glossary": FIELD_GLOSSARY,
        "formulas": FORMULAS,
        "qos_5qi_table": QOS_5QI_TABLE,
        "glossary_terms": GLOSSARY_TERMS,
        "sample_5g_records": sample_records,
    }
    return payload

# ---------- JSON ----------
def export_json(payload: dict) -> bytes:
    return json.dumps(payload, indent=2, default=str).encode("utf-8")

# ---------- TXT ----------
def export_txt(payload: dict) -> bytes:
    lines = []
    meta = payload["meta"]
    width = 88
    def hr(char="="):
        lines.append(char * width)
    hr()
    lines.append(meta["framework_name"].center(width))
    lines.append(f"Developed by {meta['developed_by']} | {meta['organization']}".center(width))
    hr()
    lines.append(f"Version: {meta['framework_version']}")
    lines.append(f"Generated: {meta['generated_at_utc']}")
    lines.append(f"Sample Records Included: {meta['sample_record_count']}")
    lines.append("")
    lines.append("DISCLAIMER: " + meta["disclaimer"])
    lines.append("")
    hr("-")
    lines.append("SECTION 1: WHAT IS A 5G TAXONOMY & FRAMEWORK?".center(width))
    hr("-")
    lines.append(textwrap.fill(
        "A taxonomy organizes a complex domain into clear categories and relationships. A framework explains "
        "how those categories interact end-to-end. This document organizes the 5G ecosystem into 7 structural "
        "layers, from raw spectrum up to business-facing services, giving a shared reference vocabulary for "
        "engineers, analysts, and business stakeholders.", width=width
    ))
    lines.append("")
    hr("-")
    lines.append("SECTION 2: TAXONOMY — 7 LAYERS".center(width))
    hr("-")
    for key, layer in payload["taxonomy"].items():
        lines.append(f"\n{layer['layer']}")
        lines.append(textwrap.fill(layer["description"], width=width))
        for cat in layer["categories"]:
            lines.append(f"  - {cat['name']}: {textwrap.shorten(cat['detail'], width=width-6, placeholder='...')}")
    lines.append("")
    hr("-")
    lines.append("SECTION 3: FLOW & ARCHITECTURE DIAGRAMS (MERMAID SOURCE)".center(width))
    hr("-")
    lines.append(textwrap.fill(
        "The diagrams below are defined in Mermaid syntax. Paste any block into https://mermaid.live to "
        "render it visually, or view it rendered live in the app's Flow & Architecture Diagrams tab.",
        width=width
    ))
    for d in payload["diagrams"]:
        lines.append(f"\n{d['title']}")
        lines.append(textwrap.fill(d["description"], width=width))
        lines.append("-" * 40)
        lines.append(d["mermaid_code"])
        lines.append("-" * 40)
    lines.append("")
    hr("-")
    lines.append("SECTION 4: SAMPLE 5G RECORD FIELD GLOSSARY".center(width))
    hr("-")
    for f in payload["field_glossary"]:
        lines.append(f"  - {f['field']}: {f['meaning']}")
    lines.append("")
    hr("-")
    lines.append("SECTION 5: KEY FORMULAS".center(width))
    hr("-")
    for f in payload["formulas"]:
        lines.append(f"\n{f['name']}")
        lines.append(f"  Formula: {f['formula']}")
        lines.append(textwrap.fill(f["explain"], width=width, initial_indent="  ", subsequent_indent="  "))
    lines.append("")
    hr("-")
    lines.append("SECTION 6: GLOSSARY OF TERMS".center(width))
    hr("-")
    for g in payload["glossary_terms"]:
        lines.append(f"  - {g['term']}: {g['def']}")
    lines.append("")
    hr("-")
    lines.append(f"SECTION 7: SAMPLE 5G RECORDS (n={len(payload['sample_5g_records'])})".center(width))
    hr("-")
    if payload["sample_5g_records"]:
        cols = list(payload["sample_5g_records"][0].keys())
        lines.append(" | ".join(cols))
        for rec in payload["sample_5g_records"]:
            lines.append(" | ".join(str(rec[c]) for c in cols))
    else:
        lines.append("No sample records were generated for this export.")
    lines.append("")
    hr()
    lines.append(f"END OF REPORT — {meta['organization']}".center(width))
    hr()
    return "\n".join(lines).encode("utf-8")

# ---------- DOCX ----------
def _set_cell_background(cell, hex_color):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)

def _add_docx_code_block(doc, code_text, font_size=8):
    """Adds a shaded, monospace 'code block' to the docx (used as a fallback when a
    diagram image couldn't be generated, so the Mermaid source is still legible)."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    _set_cell_background(cell, "F4F6FA")
    cell.paragraphs[0].text = ""
    first = True
    for line in code_text.split("\n"):
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        run = p.add_run(line if line.strip() else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor(0x3A, 0x44, 0x56)
    return tbl

def export_docx(payload: dict) -> bytes:
    doc = DocxDocument()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    meta = payload["meta"]
    navy = RGBColor(0x0B, 0x2E, 0x5E)
    blue = RGBColor(0x13, 0x57, 0xC7)
    slate = RGBColor(0x3A, 0x44, 0x56)

    # ---- Title block ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(meta["framework_name"])
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = navy

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = subtitle.add_run("A Structured Taxonomy & Reference Framework for 5G Networks")
    r2.bold = True
    r2.font.size = Pt(13)
    r2.font.color.rgb = blue

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta_p.add_run(
        f"Developed by {meta['developed_by']}  |  {meta['organization']}\n"
        f"Version {meta['framework_version']}  |  Generated {meta['generated_at_utc']}"
    )
    mr.font.size = Pt(10)
    mr.font.color.rgb = slate
    mr.italic = True

    doc.add_paragraph()
    disc = doc.add_paragraph()
    dr = disc.add_run(f"Disclaimer: {meta['disclaimer']}")
    dr.font.size = Pt(8.5)
    dr.italic = True
    dr.font.color.rgb = slate
    doc.add_page_break()

    # ---- Section 1: What is ----
    h1 = doc.add_heading("1. What is a 5G Taxonomy & Framework?", level=1)
    for run in h1.runs:
        run.font.color.rgb = navy
    doc.add_paragraph(
        "A taxonomy is a structured classification system that organizes a complex domain into clear "
        "categories, sub-categories, and relationships. A framework explains how those categories interact — "
        "the flows, dependencies, and rules that connect them into a working system."
    )
    doc.add_paragraph(
        f"The {meta['framework_name']} organizes the 5G ecosystem into 7 structural layers, from raw radio "
        "spectrum up to business-facing services, giving engineers, analysts, and business stakeholders a "
        "shared map of 5G architecture, terminology, and performance concepts."
    )

    # ---- Section 2: Taxonomy ----
    h2 = doc.add_heading("2. Taxonomy — 7 Structural Layers", level=1)
    for run in h2.runs:
        run.font.color.rgb = navy
    for key, layer in payload["taxonomy"].items():
        hl = doc.add_heading(layer["layer"], level=2)
        for run in hl.runs:
            run.font.color.rgb = blue
        p = doc.add_paragraph()
        pr = p.add_run(layer["description"])
        pr.italic = True
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Category"
        hdr[1].text = "Explanation"
        for c in hdr:
            for p2 in c.paragraphs:
                for r2 in p2.runs:
                    r2.bold = True
            _set_cell_background(c, "0B2E5E")
            for p2 in c.paragraphs:
                for r2 in p2.runs:
                    r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for cat in layer["categories"]:
            row = table.add_row().cells
            row[0].text = cat["name"]
            row[1].text = cat["detail"]
        doc.add_paragraph()
    doc.add_page_break()

    # ---- Section 3: Flow & Architecture Diagrams ----
    h3d = doc.add_heading("3. Flow & Architecture Diagrams", level=1)
    for run in h3d.runs:
        run.font.color.rgb = navy
    doc.add_paragraph(
        "These diagrams illustrate the taxonomy's structure and the key operational flows — connection "
        "establishment, network slicing, and identity/security — described throughout this framework."
    )
    for d in payload["diagrams"]:
        hd = doc.add_heading(d["title"], level=2)
        for run in hd.runs:
            run.font.color.rgb = blue
        dp = doc.add_paragraph()
        dpr = dp.add_run(d["description"])
        dpr.italic = True

        img_bytes = fetch_diagram_image(d["mermaid_code"])
        inserted = False
        if img_bytes:
            try:
                doc.add_picture(io.BytesIO(img_bytes), width=Inches(6.2))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                inserted = True
            except Exception:
                inserted = False
        if not inserted:
            note = doc.add_paragraph()
            nr = note.add_run(
                "Diagram image could not be generated automatically (no internet access to the rendering "
                "service at export time). The Mermaid source is provided below — paste it into "
                "https://mermaid.live to view it graphically."
            )
            nr.italic = True
            nr.font.size = Pt(8.5)
            nr.font.color.rgb = slate
            _add_docx_code_block(doc, d["mermaid_code"])
        doc.add_paragraph()
    doc.add_page_break()

    # ---- Section 4: Field glossary ----
    h4 = doc.add_heading("4. Sample 5G Record — Field-by-Field Glossary", level=1)
    for run in h4.runs:
        run.font.color.rgb = navy
    doc.add_paragraph("Each field below corresponds to a structural concept defined in the taxonomy.")
    ft = doc.add_table(rows=1, cols=2)
    ft.style = "Light Grid Accent 1"
    fhdr = ft.rows[0].cells
    fhdr[0].text = "Field"
    fhdr[1].text = "Explanation"
    for c in fhdr:
        _set_cell_background(c, "0B2E5E")
        for p2 in c.paragraphs:
            for r2 in p2.runs:
                r2.bold = True
                r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for f in payload["field_glossary"]:
        row = ft.add_row().cells
        row[0].text = f["field"]
        row[1].text = f["meaning"]
    doc.add_page_break()

    # ---- Section 5: Formulas ----
    h5 = doc.add_heading("5. Key Formulas Behind 5G Performance", level=1)
    for run in h5.runs:
        run.font.color.rgb = navy
    for f in payload["formulas"]:
        hf = doc.add_heading(f["name"], level=2)
        for run in hf.runs:
            run.font.color.rgb = blue
        fp = doc.add_paragraph()
        fr = fp.add_run(f["formula"])
        fr.font.name = "Courier New"
        fr.bold = True
        doc.add_paragraph(f["explain"])
    doc.add_page_break()

    # ---- Section 6: Glossary ----
    h6 = doc.add_heading("6. Glossary of Terms", level=1)
    for run in h6.runs:
        run.font.color.rgb = navy
    gt = doc.add_table(rows=1, cols=2)
    gt.style = "Light Grid Accent 1"
    ghdr = gt.rows[0].cells
    ghdr[0].text = "Term"
    ghdr[1].text = "Definition"
    for c in ghdr:
        _set_cell_background(c, "0B2E5E")
        for p2 in c.paragraphs:
            for r2 in p2.runs:
                r2.bold = True
                r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for g in payload["glossary_terms"]:
        row = gt.add_row().cells
        row[0].text = g["term"]
        row[1].text = g["def"]
    doc.add_page_break()

    # ---- Section 7: Sample records ----
    h7 = doc.add_heading(f"7. Sample 5G Records (n={len(payload['sample_5g_records'])})", level=1)
    for run in h7.runs:
        run.font.color.rgb = navy
    records = payload["sample_5g_records"]
    if records:
        cols = list(records[0].keys())
        max_rows = min(len(records), 150)
        rt = doc.add_table(rows=1, cols=len(cols))
        rt.style = "Light Grid Accent 1"
        rhdr = rt.rows[0].cells
        for i, c in enumerate(cols):
            rhdr[i].text = c
            _set_cell_background(rhdr[i], "0B2E5E")
            for p2 in rhdr[i].paragraphs:
                for r2 in p2.runs:
                    r2.bold = True
                    r2.font.size = Pt(7.5)
                    r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for rec in records[:max_rows]:
            row = rt.add_row().cells
            for i, c in enumerate(cols):
                row[i].text = str(rec[c])
                for p2 in row[i].paragraphs:
                    for r2 in p2.runs:
                        r2.font.size = Pt(7.5)
        if len(records) > max_rows:
            doc.add_paragraph(f"(Showing first {max_rows} of {len(records)} records. Full dataset available via JSON/Text export.)")
    else:
        doc.add_paragraph("No sample records were generated for this export.")

    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr2 = footer_p.add_run(f"{meta['organization']} — {meta['framework_name']} {meta['framework_version']}")
    fr2.font.size = Pt(8)
    fr2.font.color.rgb = slate
    fr2.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# ---------- PDF ----------
def _pdf_diagram_flowable(png_bytes, max_width_inch=6.3, max_height_inch=8.5):
    """Builds a properly-scaled ReportLab Image flowable from PNG bytes, preserving
    aspect ratio so diagrams never get stretched or clipped on the page."""
    if not png_bytes:
        return None
    try:
        img_buf = io.BytesIO(png_bytes)
        pil_img = PILImage.open(img_buf)
        w, h = pil_img.size
        if w <= 0 or h <= 0:
            return None
        max_w = max_width_inch * inch
        disp_w = max_w
        disp_h = h * (max_w / w)
        max_h = max_height_inch * inch
        if disp_h > max_h:
            disp_h = max_h
            disp_w = w * (max_h / h)
        img_buf.seek(0)
        return RLImage(img_buf, width=disp_w, height=disp_h)
    except Exception:
        return None

def export_pdf(payload: dict) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=LETTER,
        topMargin=0.7 * inch, bottomMargin=0.7 * inch,
        leftMargin=0.7 * inch, rightMargin=0.7 * inch,
    )
    styles = getSampleStyleSheet()
    navy = colors.HexColor("#0B2E5E")
    blue = colors.HexColor("#1357C7")
    slate = colors.HexColor("#3A4456")
    gold = colors.HexColor("#C8932B")
    light_blue = colors.HexColor("#EAF1FC")

    style_title = ParagraphStyle(
        "TitleBlue", parent=styles["Title"], textColor=navy,
        fontSize=22, leading=26, alignment=TA_CENTER, spaceAfter=4,
    )
    style_subtitle = ParagraphStyle(
        "Subtitle", parent=styles["Normal"], textColor=blue,
        fontSize=13, leading=16, alignment=TA_CENTER, spaceAfter=10, fontName="Helvetica-Bold",
    )
    style_meta = ParagraphStyle(
        "Meta", parent=styles["Normal"], textColor=slate,
        fontSize=9.5, leading=13, alignment=TA_CENTER, spaceAfter=4, fontName="Helvetica-Oblique",
    )
    style_h1 = ParagraphStyle(
        "H1", parent=styles["Heading1"], textColor=navy, fontSize=15, spaceBefore=14, spaceAfter=8,
    )
    style_h2 = ParagraphStyle(
        "H2", parent=styles["Heading2"], textColor=blue, fontSize=12, spaceBefore=10, spaceAfter=5,
    )
    style_body = ParagraphStyle(
        "Body", parent=styles["Normal"], textColor=slate, fontSize=9.5, leading=13.5, spaceAfter=6,
    )
    style_formula = ParagraphStyle(
        "Formula", parent=styles["Normal"], textColor=colors.HexColor("#7A5A12"),
        fontName="Courier-Bold", fontSize=10, backColor=colors.HexColor("#FBF3E1"),
        borderPadding=6, spaceAfter=6, spaceBefore=2,
    )
    style_code = ParagraphStyle(
        "DiagramCode", parent=styles["Normal"], textColor=slate,
        fontName="Courier", fontSize=7.3, leading=9.2,
        backColor=colors.HexColor("#F4F6FA"), borderPadding=6, spaceAfter=8, spaceBefore=2,
    )
    style_disclaimer = ParagraphStyle(
        "Disc", parent=styles["Normal"], textColor=slate, fontSize=7.8, fontName="Helvetica-Oblique",
        alignment=TA_CENTER,
    )

    meta = payload["meta"]
    elements = []

    # ---- Cover ----
    elements.append(Spacer(1, 0.6 * inch))
    elements.append(Paragraph(meta["framework_name"], style_title))
    elements.append(Paragraph("A Structured Taxonomy &amp; Reference Framework for 5G Networks", style_subtitle))
    elements.append(HRFlowable(width="60%", thickness=1.2, color=blue, spaceBefore=4, spaceAfter=10, hAlign="CENTER"))
    elements.append(Paragraph(f"Developed by <b>{meta['developed_by']}</b>", style_meta))
    elements.append(Paragraph(f"{meta['organization']}", style_meta))
    elements.append(Spacer(1, 0.15 * inch))
    elements.append(Paragraph(f"Version: {meta['framework_version']} &nbsp;|&nbsp; Generated: {meta['generated_at_utc']}", style_meta))
    elements.append(Paragraph(f"Sample Records Included: {meta['sample_record_count']}", style_meta))
    elements.append(Spacer(1, 0.4 * inch))
    elements.append(Paragraph(f"Disclaimer: {meta['disclaimer']}", style_disclaimer))
    elements.append(PageBreak())

    # ---- Section 1 ----
    elements.append(Paragraph("1. What is a 5G Taxonomy &amp; Framework?", style_h1))
    elements.append(Paragraph(
        "A taxonomy is a structured classification system that organizes a complex domain into clear "
        "categories, sub-categories, and relationships. A framework explains how those categories interact — "
        "the flows, dependencies, and rules that connect them into a working system.", style_body))
    elements.append(Paragraph(
        f"The {meta['framework_name']} organizes the 5G ecosystem into 7 structural layers, from raw radio "
        "spectrum up to business-facing services, giving engineers, analysts, and business stakeholders a "
        "shared map of 5G architecture, terminology, and performance concepts.", style_body))

    # ---- Section 2: Taxonomy ----
    elements.append(Paragraph("2. Taxonomy — 7 Structural Layers", style_h1))
    for key, layer in payload["taxonomy"].items():
        elements.append(Paragraph(layer["layer"], style_h2))
        elements.append(Paragraph(f"<i>{layer['description']}</i>", style_body))
        table_data = [["Category", "Explanation"]]
        for cat in layer["categories"]:
            table_data.append([Paragraph(cat["name"], style_body), Paragraph(cat["detail"], style_body)])
        t = Table(table_data, colWidths=[1.8 * inch, 4.5 * inch])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), navy),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8.5),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#D7E3F7")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, light_blue]),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        elements.append(t)
        elements.append(Spacer(1, 0.12 * inch))
    elements.append(PageBreak())

    # ---- Section 3: Flow & Architecture Diagrams ----
    elements.append(Paragraph("3. Flow & Architecture Diagrams", style_h1))
    elements.append(Paragraph(
        "These diagrams illustrate the taxonomy's structure and the key operational flows — connection "
        "establishment, network slicing, and identity/security — described throughout this framework.",
        style_body))
    for d in payload["diagrams"]:
        elements.append(Paragraph(d["title"], style_h2))
        elements.append(Paragraph(f"<i>{d['description']}</i>", style_body))
        png_bytes = fetch_diagram_image(d["mermaid_code"])
        flow_img = _pdf_diagram_flowable(png_bytes)
        if flow_img is not None:
            elements.append(flow_img)
        else:
            elements.append(Paragraph(
                "Diagram image could not be generated automatically (no internet access to the rendering "
                "service at export time). The Mermaid source is provided below — paste it into "
                "<a href='https://mermaid.live'>mermaid.live</a> to view it graphically.", style_body))
            elements.append(Preformatted(d["mermaid_code"], style_code))
        elements.append(Spacer(1, 0.2 * inch))
    elements.append(PageBreak())

    # ---- Section 4: Field glossary ----
    elements.append(Paragraph("4. Sample 5G Record — Field-by-Field Glossary", style_h1))
    elements.append(Paragraph("Each field below corresponds to a structural concept defined in the taxonomy.", style_body))
    fg_data = [["Field", "Explanation"]]
    for f in payload["field_glossary"]:
        fg_data.append([Paragraph(f["field"], style_body), Paragraph(f["meaning"], style_body)])
    tf = Table(fg_data, colWidths=[1.6 * inch, 4.7 * inch])
    tf.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), navy),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8.5),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#D7E3F7")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, light_blue]),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    elements.append(tf)
    elements.append(PageBreak())

    # ---- Section 5: Formulas ----
    elements.append(Paragraph("5. Key Formulas Behind 5G Performance", style_h1))
    for f in payload["formulas"]:
        elements.append(Paragraph(f["name"], style_h2))
        elements.append(Paragraph(f["formula"], style_formula))
        elements.append(Paragraph(f["explain"], style_body))
    elements.append(PageBreak())

    # ---- Section 6: Glossary ----
    elements.append(Paragraph("6. Glossary of Terms", style_h1))
    gl_data = [["Term", "Definition"]]
    for g in payload["glossary_terms"]:
        gl_data.append([Paragraph(g["term"], style_body), Paragraph(g["def"], style_body)])
    tg = Table(gl_data, colWidths=[1.4 * inch, 4.9 * inch])
    tg.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), navy),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8.5),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#D7E3F7")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, light_blue]),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    elements.append(tg)
    elements.append(PageBreak())

    # ---- Section 7: Sample records ----
    records = payload["sample_5g_records"]
    elements.append(Paragraph(f"7. Sample 5G Records (n={len(records)})", style_h1))
    if records:
        max_rows = min(len(records), 80)
        cols = list(records[0].keys())
        short_cols = ["record_id", "spectrum_band", "ran_mode", "slice_type",
                      "latency_ms", "throughput_mbps", "reliability_pct", "qos_5qi"]
        short_cols = [c for c in short_cols if c in cols]
        rec_data = [short_cols]
        style_tiny = ParagraphStyle("Tiny", parent=style_body, fontSize=7, leading=9)
        for rec in records[:max_rows]:
            rec_data.append([Paragraph(str(rec[c]), style_tiny) for c in short_cols])
        tr = Table(rec_data, repeatRows=1)
        tr.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), navy),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 7),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#D7E3F7")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, light_blue]),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 2),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
        ]))
        elements.append(tr)
        if len(records) > max_rows:
            elements.append(Spacer(1, 0.1 * inch))
            elements.append(Paragraph(
                f"Showing first {max_rows} of {len(records)} records (subset of columns). "
                "Full dataset with all columns is available via the JSON and Text exports.", style_body))
    else:
        elements.append(Paragraph("No sample records were generated for this export.", style_body))

    elements.append(Spacer(1, 0.3 * inch))
    elements.append(HRFlowable(width="100%", thickness=0.6, color=colors.HexColor("#D7E3F7")))
    elements.append(Spacer(1, 0.08 * inch))
    elements.append(Paragraph(f"{meta['organization']} — {meta['framework_name']} {meta['framework_version']}", style_disclaimer))

    doc.build(elements)
    return buf.getvalue()

# =============================================================================
# TAB 6: EXPORT & REPORTS
# =============================================================================
with tab_export:
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<div class="section-heading">Export the Complete Framework</div>', unsafe_allow_html=True)
    st.markdown(f"""
Download the entire **{FRAMEWORK_NAME}** — including the overview explanation, all 7 taxonomy layers,
the 4 flow & architecture diagrams, field-by-field glossary, formulas, terminology dictionary, and any
currently generated sample 5G records — in your preferred format. Every export carries the
**{BRAND_NAME}** title block and attribution to **{AUTHOR_NAME}**.

Diagrams are embedded as images when an internet connection to the rendering service is available at
export time; otherwise the Mermaid source is included as readable text so no diagram content is ever lost.
""")
    df_preview = st.session_state.synthetic_df
    n_preview = 0 if df_preview is None else len(df_preview)
    if n_preview == 0:
        st.warning(
            "No sample data is currently generated — exports will still include the full taxonomy, diagrams, "
            "glossary, and formulas, but the sample-records section will be empty. Generate data from the "
            "sidebar first if you'd like it included.",
            icon="ℹ️"
        )
    else:
        st.success(f"{n_preview} synthetic 5G record(s) currently in memory will be included in your export.", icon="✅")
    st.markdown('</div>', unsafe_allow_html=True)

    payload = build_export_payload()

    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<div class="section-heading">Choose a Format</div>', unsafe_allow_html=True)
    colA, colB, colC, colD = st.columns(4)
    with colA:
        st.markdown("**📕 PDF Report**")
        st.caption("Polished, print-ready document with title page, diagrams, tables, and formulas.")
        try:
            pdf_bytes = export_pdf(payload)
            st.download_button(
                "Download PDF", data=pdf_bytes,
                file_name="5G_Taxonomy_Framework_KNet.pdf",
                mime="application/pdf", use_container_width=True
            )
        except Exception as e:
            st.error(f"PDF generation error: {e}")
    with colB:
        st.markdown("**📄 Word (DOCX)**")
        st.caption("Fully editable Word document with styled headings, diagrams, and tables.")
        try:
            docx_bytes = export_docx(payload)
            st.download_button(
                "Download Word", data=docx_bytes,
                file_name="5G_Taxonomy_Framework_KNet.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        except Exception as e:
            st.error(f"DOCX generation error: {e}")
    with colC:
        st.markdown("**📝 Plain Text**")
        st.caption("Lightweight, universally readable .txt summary, including diagram source code.")
        try:
            txt_bytes = export_txt(payload)
            st.download_button(
                "Download Text", data=txt_bytes,
                file_name="5G_Taxonomy_Framework_KNet.txt",
                mime="text/plain", use_container_width=True
            )
        except Exception as e:
            st.error(f"Text generation error: {e}")
    with colD:
        st.markdown("**🗂️ JSON**")
        st.caption("Fully structured, machine-readable export for integration/reuse.")
        try:
            json_bytes = export_json(payload)
            st.download_button(
                "Download JSON", data=json_bytes,
                file_name="5G_Taxonomy_Framework_KNet.json",
                mime="application/json", use_container_width=True
            )
        except Exception as e:
            st.error(f"JSON generation error: {e}")
    st.markdown('</div>', unsafe_allow_html=True)

    with st.expander("🔍 Preview export payload structure (JSON)"):
        preview_payload = dict(payload)
        preview_payload["sample_5g_records"] = preview_payload["sample_5g_records"][:5]
        st.json(preview_payload)
        st.caption("Preview truncated to first 5 sample records for readability — full export contains all generated records.")

# =============================================================================
# FOOTER
# =============================================================================
st.markdown(f"""
<div class="footer-note">
    {FRAMEWORK_NAME} {FRAMEWORK_VERSION} &nbsp;•&nbsp; Developed by {AUTHOR_NAME} &nbsp;•&nbsp; {BRAND_NAME}<br>
    Synthetic sample data is illustrative only and does not represent any real network, operator, or subscriber.
</div>
""", unsafe_allow_html=True)
