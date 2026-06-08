


# AI-Drone-Inspection-Risk-Assessment-Platform.py

import streamlit as st
import pandas as pd
import numpy as np
import random
import json
from datetime import datetime
import plotly.express as px

from io import BytesIO
from docx import Document

from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer
)

from reportlab.lib.styles import getSampleStyleSheet


# -------------------------
# PAGE CONFIG
# -------------------------
st.set_page_config(
    page_title="KNet SkyGuard AI",
    layout="wide"
)

# -------------------------
# GLOBAL HEADER STYLE FUNCTION
# -------------------------
def blue_header(text, size=34):
    st.markdown(
        f"""
        <h1 style='
            color:#0056D2;
            font-weight:bold;
            font-size:{size}px;
            text-align:left;
        '>
        {text}
        </h1>
        """,
        unsafe_allow_html=True
    )


# -------------------------
# MAIN HEADER
# -------------------------
st.markdown("""
<div style='text-align:center;'>

<h1 style='color:#0056D2;font-weight:bold;font-size:42px;'>
KNet SkyGuard AI™
</h1>

<h2 style='color:#0056D2;font-weight:bold;'>
AI Drone Inspection & Risk Assessment Platform
</h2>

<h3 style='color:#0056D2;font-weight:bold;'>
Developed by Randy Singh – Kalsnet (KNet) Consulting Group
</h3>

</div>
""", unsafe_allow_html=True)


# -------------------------
# SESSION STATE
# -------------------------
if "df" not in st.session_state:
    st.session_state.df = pd.DataFrame()


# -------------------------
# SIDEBAR
# -------------------------
page = st.sidebar.selectbox(
    "Navigation",
    [
        "Dashboard",
        "Upload Data",
        "Synthetic Data",
        "Analytics",
        "Export Results",
        "Field Explanations",
        "Formula Explanations"
    ]
)


# -------------------------
# RISK FORMULA
# -------------------------
def calculate_risk(severity, weather, battery, confidence):
    score = (
        severity * 0.40 +
        weather * 0.15 +
        battery * 0.15 +
        confidence * 0.30
    )
    return round(score, 2)


# -------------------------
# REPORT GENERATORS
# -------------------------
def create_word_report(df):
    doc = Document()

    doc.add_heading("Kalsnet (KNet) Analysis Report", level=1)
    doc.add_paragraph(f"Generated: {datetime.now()}")
    doc.add_paragraph(f"Total Inspections: {len(df)}")

    if not df.empty:
        doc.add_paragraph(f"Average Risk Score: {round(df['Risk Score'].mean(),2)}")
        doc.add_paragraph("\nInspection Data:\n")
        doc.add_paragraph(df.to_string())

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def create_pdf_report(df):
    output = BytesIO()
    pdf = SimpleDocTemplate(output)
    styles = getSampleStyleSheet()
    content = []

    content.append(
        Paragraph(
            "<font color='blue'><b>Kalsnet (KNet) Analysis Report</b></font>",
            styles["Title"]
        )
    )

    content.append(Spacer(1, 12))
    content.append(Paragraph(f"Generated: {datetime.now()}", styles["Normal"]))
    content.append(Spacer(1, 12))
    content.append(Paragraph(f"Total Inspections: {len(df)}", styles["Normal"]))

    if not df.empty:
        content.append(
            Paragraph(
                f"Average Risk Score: {round(df['Risk Score'].mean(),2)}",
                styles["Normal"]
            )
        )

    pdf.build(content)
    output.seek(0)
    return output


# -------------------------
# DASHBOARD
# -------------------------
if page == "Dashboard":

    blue_header("Executive Dashboard", 30)

    df = st.session_state.df

    c1, c2, c3, c4 = st.columns(4)

    c1.metric("Total Inspections", len(df))

    if not df.empty:
        avg_risk = round(df["Risk Score"].mean(), 2)
        high_risk = len(df[df["Risk Score"] > 60])
    else:
        avg_risk = 0
        high_risk = 0

    c2.metric("Average Risk Score", avg_risk)
    c3.metric("High Risk Assets", high_risk)
    c4.metric("Compliance Score", "93%")

    st.markdown("---")
    st.write("This dashboard provides enterprise-level drone inspection intelligence.")


# -------------------------
# UPLOAD DATA
# -------------------------
elif page == "Upload Data":

    blue_header("Upload Real Data", 30)

    uploaded = st.file_uploader(
        "Upload CSV, Excel or JSON",
        type=["csv", "xlsx", "json"]
    )

    if uploaded:

        if uploaded.name.endswith(".csv"):
            df = pd.read_csv(uploaded)
        elif uploaded.name.endswith(".xlsx"):
            df = pd.read_excel(uploaded)
        elif uploaded.name.endswith(".json"):
            df = pd.read_json(uploaded)

        st.session_state.df = df

        st.success("Data uploaded successfully!")
        st.dataframe(df, use_container_width=True)


# -------------------------
# SYNTHETIC DATA
# -------------------------
elif page == "Synthetic Data":

    blue_header("Synthetic Data Generator", 30)

    n = st.slider("Number of Records", 0, 250, 50)

    col1, col2 = st.columns(2)

    with col1:
        if st.button("Generate Records"):

            rows = []

            for i in range(n):

                severity = random.randint(1, 100)
                weather = random.randint(1, 100)
                battery = random.randint(1, 100)
                confidence = random.randint(1, 100)

                risk = calculate_risk(severity, weather, battery, confidence)

                rows.append({
                    "Inspection ID": f"INS-{1000+i}",
                    "Drone ID": f"DR-{random.randint(1,50)}",
                    "Asset ID": f"AST-{random.randint(100,999)}",
                    "Latitude": round(random.uniform(36, 39), 6),
                    "Longitude": round(random.uniform(-79, -75), 6),
                    "Altitude": random.randint(50, 400),
                    "Weather": weather,
                    "Battery": battery,
                    "Severity": severity,
                    "Confidence": confidence,
                    "Risk Score": risk,
                    "Timestamp": str(datetime.now())
                })

            st.session_state.df = pd.DataFrame(rows)

    with col2:
        if st.button("Reset Records"):
            st.session_state.df = pd.DataFrame()
            st.success("All records reset.")

    st.dataframe(st.session_state.df, use_container_width=True)


# -------------------------
# ANALYTICS
# -------------------------
elif page == "Analytics":

    blue_header("Inspection Analytics", 30)

    df = st.session_state.df

    if not df.empty:

        fig = px.histogram(df, x="Risk Score", title="Risk Distribution")
        st.plotly_chart(fig, use_container_width=True)

        fig2 = px.scatter(df, x="Severity", y="Risk Score", title="Severity vs Risk")
        st.plotly_chart(fig2, use_container_width=True)


# -------------------------
# EXPORT RESULTS
# -------------------------
elif page == "Export Results":

    blue_header("Export Results", 30)

    df = st.session_state.df

    if df.empty:
        st.warning("No data available for export.")

    else:
        st.success(f"{len(df)} records ready for export.")

        pdf_file = create_pdf_report(df)
        st.download_button(
            "📄 Download PDF Report",
            data=pdf_file,
            file_name="KNet_Analysis_Report.pdf",
            mime="application/pdf"
        )

        word_file = create_word_report(df)
        st.download_button(
            "📝 Download Word Report",
            data=word_file,
            file_name="KNet_Analysis_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        st.download_button(
            "📊 Download JSON",
            data=df.to_json(orient="records", indent=2),
            file_name="KNet_Analysis_Report.json",
            mime="application/json"
        )

        st.dataframe(df.head(25), use_container_width=True)


# -------------------------
# FIELD EXPLANATIONS
# -------------------------
elif page == "Field Explanations":

    blue_header("Field Explanations", 30)

    st.table(pd.DataFrame({
        "Field": [
            "Inspection ID", "Drone ID", "Asset ID",
            "Latitude", "Longitude", "Altitude",
            "Weather", "Battery", "Severity",
            "Confidence", "Risk Score"
        ],
        "Purpose": [
            "Unique inspection", "Drone identity", "Asset tracked",
            "GPS location", "GPS location", "Flight height",
            "Weather impact", "Battery status", "Damage severity",
            "AI confidence", "Overall risk"
        ]
    }))


# -------------------------
# FORMULA EXPLANATION
# -------------------------
elif page == "Formula Explanations":

    blue_header("Formula Explanation", 30)

    st.latex(r"""
    RiskScore =
    (Severity \times 0.40)
    +
    (Weather \times 0.15)
    +
    (Battery \times 0.15)
    +
    (Confidence \times 0.30)
    """)

    st.write("""
    **Severity (40%)** contributes most to risk.

    **Weather (15%)** affects inspection reliability.

    **Battery (15%)** affects drone stability.

    **Confidence (30%)** represents AI certainty.
    """)