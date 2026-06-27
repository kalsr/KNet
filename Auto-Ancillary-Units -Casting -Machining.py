# AI for Auto Ancillary Units — Casting & Machining
# Single-file Streamlit application.
# Developed by Randy Singh | Kalsnet (KNet) Consulting Group
# Modules:
#   1. Vision QC — Casting Defect Detection
#   2. Predictive Maintenance — Machine Health
#   3. Tool Wear Prediction
#   4. Demand & Inventory Forecast
#   5. Energy Optimization
#   6. Traceability / Root Cause
#   7. ROI Summary
# Each module has:
#   - Generate Synthetic Data button
#   - Upload Real Data (CSV/Excel) — real data takes priority over synthetic when present
#   - Charts / tables
#   - Formula / methodology explanation
#   - Export to CSV / JSON / TXT / Word / PDF
#   - Reset Data button
import io
import logging
import os
import datetime as dt
from datetime import datetime, timedelta
from typing import Dict, Optional

import numpy as np
import pandas as pd
import streamlit as st

try:
    from fpdf import FPDF
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# ----------------------------------------------------------------------------
# Logging
# ----------------------------------------------------------------------------
LOG_DIR = os.environ.get("AAU_LOG_DIR", "./logs")
os.makedirs(LOG_DIR, exist_ok=True)
logger = logging.getLogger("auto_ancillary_units")
if not logger.handlers:
    logger.setLevel(logging.INFO)
    _fh = logging.FileHandler(os.path.join(LOG_DIR, "auto_ancillary_units.log"))
    _fh.setFormatter(logging.Formatter("%(asctime)s | %(levelname)s | %(message)s"))
    logger.addHandler(_fh)

# ----------------------------------------------------------------------------
# Page config
# ----------------------------------------------------------------------------
st.set_page_config(
    page_title="AI for Auto Ancillary Units",
    page_icon="⚙️",
    layout="wide",
)

RNG_SEED = 42

# ----------------------------------------------------------------------------
# Styling — bold blue title bar (matches ShopFloor Sentinel / ChainIntel Pro)
# ----------------------------------------------------------------------------
st.markdown(
    """
    <style>
    .main-title {
        text-align: center;
        color: #0047AB;
        font-weight: 900;
        font-size: 2.6rem;
        margin-bottom: 0.3rem;
        letter-spacing: 0.6px;
        font-family: 'Arial Black', sans-serif;
    }
    .subtitle-line {
        text-align: center;
        color: #0047AB;
        font-weight: 800;
        font-size: 1.15rem;
        margin-top: 0;
        margin-bottom: 1.1rem;
        font-family: 'Arial', sans-serif;
    }
    .section-header {
        color: #0047AB;
        font-weight: 800;
        font-size: 1.6rem;
        margin-top: 0.5rem;
        margin-bottom: 0.4rem;
    }
    .formula-box {
        background: #fff9e6;
        border: 2px solid #ffd700;
        padding: 1rem;
        border-radius: 6px;
        font-size: 0.92rem;
        margin: 0.8rem 0;
    }
    .info-box {
        background: #e8f4f8;
        border-left: 4px solid #0099cc;
        padding: 0.8rem;
        border-radius: 4px;
        font-size: 0.85rem;
        margin: 0.6rem 0;
    }
    .divider-line {
        border-top: 2px solid #0047AB;
        margin: 1rem 0;
        opacity: 0.3;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# Top title bar — always visible above the module navigation, both lines bold blue
st.markdown('<div class="main-title">⚙️ AI for Auto Ancillary Units — Casting &amp; Machining</div>', unsafe_allow_html=True)
st.markdown('<div class="subtitle-line">Developed by Randy Singh, Kalsnet (KNet) Consulting Group</div>', unsafe_allow_html=True)
st.markdown('<div class="divider-line"></div>', unsafe_allow_html=True)

# ----------------------------------------------------------------------------
# Export helpers (CSV / JSON / TXT / Word / PDF)
# ----------------------------------------------------------------------------
# FPDF's built-in core fonts only support Latin-1. Emoji and special symbols
# (⚙️🔍🛠️🔧📦⚡🧩💰▶️🔄⬇️🟢🟡🔴⚠️, smart quotes, em-dashes, ₹) crash pdf.cell()/
# multi_cell() with a Unicode error unless sanitized first.
_PDF_CHAR_MAP = {
    "⚙️": "", "🔍": "", "🛠️": "", "🔧": "", "📦": "", "⚡": "", "🧩": "", "💰": "",
    "▶️": "", "🔄": "[RESET] ", "⬇️": "", "🏠": "", "📊": "", "📋": "", "📄": "", "📝": "",
    "🟢": "[OK] ", "🟡": "[MONITOR] ", "🔴": "[ALERT] ", "⚠️": "[WARNING] ", "⚠": "[WARNING] ",
    "✅": "[OK] ", "❌": "[X] ",
    "₹": "Rs. ", "–": "-", "—": "-",
    "’": "'", "‘": "'", "“": '"', "”": '"', "…": "...",
}


def _pdf_safe(value) -> str:
    """Make text safe for FPDF's Latin-1-only core fonts."""
    text = str(value)
    for original, replacement in _PDF_CHAR_MAP.items():
        text = text.replace(original, replacement)
    return text.encode("latin-1", "ignore").decode("latin-1").strip()


def export_to_pdf(df: pd.DataFrame, title: str, formulas: Optional[Dict] = None) -> Optional[bytes]:
    """Export dataframe (+ optional formula explanations) to PDF."""
    if not HAS_PDF:
        return None
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", "B", 16)
        pdf.cell(0, 10, _pdf_safe(title), ln=True, align="C")

        pdf.set_font("Arial", "B", 10)
        pdf.cell(0, 5, f"Generated: {dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", ln=True)
        pdf.cell(0, 5, "AI for Auto Ancillary Units | Kalsnet (KNet) Consulting Group", ln=True)

        pdf.set_font("Arial", "", 8)
        pdf.ln(5)

        col_width = 190 / max(1, len(df.columns))
        row_height = 7
        page_bottom = pdf.h - pdf.b_margin

        def draw_header():
            pdf.set_font("Arial", "B", 8)
            for col in df.columns:
                pdf.cell(col_width, row_height, _pdf_safe(col)[:20], border=1, align="C")
            pdf.ln()
            pdf.set_font("Arial", "", 8)

        draw_header()

        for _, row in df.head(60).iterrows():
            if pdf.get_y() + row_height > page_bottom:
                pdf.add_page()
                draw_header()
            for val in row:
                pdf.cell(col_width, row_height, _pdf_safe(val)[:20], border=1, align="L")
            pdf.ln()

        if formulas:
            pdf.add_page()
            pdf.set_font("Arial", "B", 12)
            pdf.cell(0, 10, "Formulas & Explanations", ln=True)
            pdf.set_font("Arial", "", 9)
            for name, text in formulas.items():
                pdf.cell(0, 5, _pdf_safe(name) + ":", ln=True)
                pdf.multi_cell(0, 4, _pdf_safe(text)[:600])
                pdf.ln(2)

        raw_output = pdf.output(dest="S")
        if isinstance(raw_output, str):
            return raw_output.encode("latin-1")
        return bytes(raw_output)
    except Exception as e:
        st.error(f"PDF export error: {e}")
        logger.exception("PDF export failed")
        return None


def export_to_docx(df: pd.DataFrame, title: str, formulas: Optional[Dict] = None) -> Optional[bytes]:
    """Export dataframe (+ optional formula explanations) to a Word document."""
    if not HAS_DOCX:
        return None
    try:
        doc = Document()

        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 71, 171)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        meta_para = doc.add_paragraph()
        meta_run = meta_para.add_run(f"Generated: {dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        meta_run.font.size = Pt(10)
        meta_run = meta_para.add_run("AI for Auto Ancillary Units | Developed by Randy Singh | Kalsnet (KNet) Consulting Group")
        meta_run.font.size = Pt(10)
        meta_run.italic = True

        doc.add_paragraph()

        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = "Light Grid Accent 1"
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr_cells[i].text = str(col)

        for _, row in df.head(60).iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val)[:100]

        if formulas:
            doc.add_page_break()
            heading_para = doc.add_paragraph()
            heading_run = heading_para.add_run("Formulas & Explanations")
            heading_run.font.size = Pt(14)
            heading_run.font.bold = True
            heading_run.font.color.rgb = RGBColor(0, 71, 171)

            for name, text in formulas.items():
                doc.add_paragraph(name, style="Heading 3")
                doc.add_paragraph(text[:600])

        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        return docx_bytes.getvalue()
    except Exception as e:
        st.error(f"Word export error: {e}")
        logger.exception("Word export failed")
        return None


def export_to_text(df: pd.DataFrame, title: str, formulas: Optional[Dict] = None) -> str:
    """Export to plain text format."""
    text = f"\n{'=' * 80}\n{title}\n{'=' * 80}\n\n"
    text += f"Generated: {dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
    text += "AI for Auto Ancillary Units | Developed by Randy Singh | Kalsnet (KNet) Consulting Group\n\n"
    text += df.to_string()

    if formulas:
        text += f"\n\n{'=' * 80}\nFormulas & Explanations\n{'=' * 80}\n\n"
        for name, val in formulas.items():
            text += f"\n{name}:\n{val}\n"
    return text


def load_uploaded_table(uploaded_file) -> Optional[pd.DataFrame]:
    """Read an uploaded CSV/Excel file into a DataFrame with a friendly error on failure."""
    try:
        if uploaded_file.name.lower().endswith(".csv"):
            return pd.read_csv(uploaded_file)
        return pd.read_excel(uploaded_file)
    except Exception as e:
        st.error(f"Could not read file: {e}")
        logger.exception("Failed to read uploaded file")
        return None


def full_export_buttons(df: pd.DataFrame, title: str, formulas: Optional[Dict], key_prefix: str, file_stub: str):
    """Render a row of CSV / JSON / TXT / Word / PDF export buttons for a dataframe."""
    st.markdown("**💾 Export Results:**")
    col1, col2, col3, col4, col5 = st.columns(5)
    with col1:
        st.download_button("📊 CSV", df.to_csv(index=False).encode(), f"{file_stub}.csv", "text/csv",
                            key=f"csv_{key_prefix}", use_container_width=True)
    with col2:
        st.download_button("📋 JSON", df.to_json(orient="records", indent=2).encode(), f"{file_stub}.json",
                            "application/json", key=f"json_{key_prefix}", use_container_width=True)
    with col3:
        txt = export_to_text(df, title, formulas)
        st.download_button("📄 TXT", txt.encode(), f"{file_stub}.txt", "text/plain",
                            key=f"txt_{key_prefix}", use_container_width=True)
    with col4:
        if HAS_DOCX:
            docx_bytes = export_to_docx(df, title, formulas)
            if docx_bytes:
                st.download_button("📝 WORD", docx_bytes, f"{file_stub}.docx",
                                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key=f"docx_{key_prefix}", use_container_width=True)
        else:
            st.caption("Word export unavailable")
    with col5:
        if HAS_PDF:
            pdf_bytes = export_to_pdf(df, title, formulas)
            if pdf_bytes:
                st.download_button("🔴 PDF", pdf_bytes, f"{file_stub}.pdf", "application/pdf",
                                    key=f"pdf_{key_prefix}", use_container_width=True)
        else:
            st.caption("PDF export unavailable")


# ----------------------------------------------------------------------------
# Generic helpers
# ----------------------------------------------------------------------------
def df_download_button(df: pd.DataFrame, label: str, filename: str, key: str):
    """Render a CSV export button for a dataframe (kept for backward compatibility)."""
    buf = io.StringIO()
    df.to_csv(buf, index=False)
    st.download_button(
        label=label,
        data=buf.getvalue(),
        file_name=filename,
        mime="text/csv",
        key=key,
        use_container_width=True,
    )


def reset_button(state_key: str, label: str = "🔄 Reset Data", key: str = None, also_clear=None):
    """Clear synthetic (and optionally real-upload) state and rerun to regenerate a fresh batch."""
    if st.button(label, key=key or f"reset_{state_key}", use_container_width=True):
        if state_key in st.session_state:
            del st.session_state[state_key]
        if also_clear:
            for k in also_clear:
                if k in st.session_state:
                    del st.session_state[k]
        st.rerun()


def upload_real_data(label: str, real_key: str, expected_cols: str, upload_widget_key: str, clear_key: str):
    """Render an upload widget + clear button for real-world data, storing it in session_state[real_key]."""
    st.markdown(f"**📁 OR Upload Real Data — {label}**")
    st.caption(f"Expected columns (case-insensitive recommended): {expected_cols}")
    col_u1, col_u2 = st.columns([4, 1])
    with col_u1:
        uploaded = st.file_uploader("Upload CSV/Excel file", type=["csv", "xlsx"], key=upload_widget_key,
                                     label_visibility="collapsed")
    with col_u2:
        if st.session_state.get(real_key) is not None:
            if st.button("❌ Clear Upload", key=clear_key, use_container_width=True):
                st.session_state[real_key] = None
                st.rerun()

    if uploaded is not None:
        real_df = load_uploaded_table(uploaded)
        if real_df is not None:
            st.session_state[real_key] = real_df
            st.success(f"✅ Loaded {len(real_df)} real records — now overriding synthetic data below.")


def pick_active_df(real_key: str, synth_key: str):
    """Real uploaded data takes priority over synthetic data when both are present."""
    real_df = st.session_state.get(real_key)
    if real_df is not None:
        return real_df, True
    return st.session_state.get(synth_key), False


def section_title(emoji: str, title: str, subtitle: str):
    st.markdown(f"## {emoji} {title}")
    st.caption(subtitle)
    st.divider()


def kpi_row(items):
    cols = st.columns(len(items))
    for c, (label, value, delta) in zip(cols, items):
        c.metric(label, value, delta)


def formula_box(markdown_text: str):
    st.markdown('<div class="formula-box">', unsafe_allow_html=True)
    st.markdown(markdown_text)
    st.markdown("</div>", unsafe_allow_html=True)


# ----------------------------------------------------------------------------
# Synthetic data generators
# ----------------------------------------------------------------------------
def gen_casting_defects(n_days=30, seed=RNG_SEED):
    rng = np.random.default_rng(seed)
    dates = pd.date_range(end=datetime.today(), periods=n_days)
    defect_types = ["Blowhole", "Shrinkage", "Cold Shut", "Gas Porosity", "Surface Crack"]
    rows = []
    base_rate = 0.035
    for i, d in enumerate(dates):
        produced = rng.integers(800, 1400)
        trend = base_rate * (1 - 0.5 * i / n_days)
        rejected = int(produced * max(0.01, rng.normal(trend, 0.006)))
        rows.append({
            "Date": d.date(),
            "Produced": produced,
            "Rejected": rejected,
            "Rejection_%": round(100 * rejected / produced, 2),
            "Top_Defect": rng.choice(defect_types, p=[0.30, 0.25, 0.15, 0.20, 0.10]),
            "Melt_Temp_C": round(rng.normal(720, 8), 1),
            "Pour_Temp_C": round(rng.normal(700, 6), 1),
            "Mold_Temp_C": round(rng.normal(180, 10), 1),
            "Humidity_%": round(rng.normal(55, 7), 1),
        })
    return pd.DataFrame(rows)


def gen_machine_health(n_hours=240, seed=RNG_SEED):
    rng = np.random.default_rng(seed + 1)
    machines = ["VMC-01", "VMC-02", "HMC-01", "Press-01", "Press-02"]
    rows = []
    start = datetime.today() - timedelta(hours=n_hours)
    for m in machines:
        degradation_start = rng.integers(int(n_hours * 0.6), n_hours)
        for h in range(n_hours):
            t = start + timedelta(hours=h)
            failing = h > degradation_start
            vib = rng.normal(2.5, 0.3) + (3.0 * (h - degradation_start) / n_hours if failing else 0)
            temp = rng.normal(55, 3) + (15 * (h - degradation_start) / n_hours if failing else 0)
            current = rng.normal(12, 1) + (4 * (h - degradation_start) / n_hours if failing else 0)
            risk = min(99, max(1, (vib / 6 + temp / 90 + current / 18) * 33))
            rows.append({
                "Timestamp": t,
                "Machine": m,
                "Vibration_mm_s": round(vib, 2),
                "Temperature_C": round(temp, 1),
                "Current_A": round(current, 2),
                "Failure_Risk_%": round(risk, 1),
                "Predicted_Days_to_Failure": round(max(0.5, 21 - risk / 5), 1),
            })
    return pd.DataFrame(rows)


def gen_tool_wear(n_records=200, seed=RNG_SEED):
    rng = np.random.default_rng(seed + 2)
    tools = ["Insert-CNMG", "Insert-WNMG", "Drill-12mm", "Boring-Bar", "End-Mill-10mm"]
    rows = []
    for i in range(n_records):
        tool = rng.choice(tools)
        cycles = rng.integers(10, 500)
        wear = min(100, round(cycles / rng.uniform(4, 7) + rng.normal(0, 4), 1))
        wear = max(0, wear)
        spindle_load = round(rng.normal(60, 10) + wear * 0.2, 1)
        vib = round(rng.normal(1.8, 0.4) + wear * 0.02, 2)
        recommended_change = "Change Now" if wear > 80 else ("Monitor" if wear > 55 else "OK")
        rows.append({
            "Tool_ID": f"{tool}-{i % 30:02d}",
            "Tool_Type": tool,
            "Cycles_Run": cycles,
            "Wear_%": wear,
            "Spindle_Load_%": spindle_load,
            "Vibration_mm_s": vib,
            "Recommendation": recommended_change,
        })
    return pd.DataFrame(rows)


def gen_demand_forecast(n_weeks=26, seed=RNG_SEED):
    rng = np.random.default_rng(seed + 3)
    parts = ["Bracket-A1", "Housing-B2", "Gear-C3", "Hub-D4"]
    rows = []
    weeks = pd.date_range(end=datetime.today(), periods=n_weeks, freq="W")
    for p in parts:
        base = rng.integers(2000, 6000)
        season_amp = base * 0.15
        for i, w in enumerate(weeks):
            actual = None if i >= n_weeks - 4 else int(
                base + season_amp * np.sin(i / 4) + rng.normal(0, base * 0.05)
            )
            forecast = int(base + season_amp * np.sin(i / 4) + rng.normal(0, base * 0.02))
            rows.append({
                "Week": w.date(),
                "Part": p,
                "Actual_Demand": actual,
                "AI_Forecast": max(0, forecast),
                "Safety_Stock": int(forecast * 0.12),
                "Reorder_Point": int(forecast * 0.35),
            })
    return pd.DataFrame(rows)


def gen_energy(n_days=30, seed=RNG_SEED):
    rng = np.random.default_rng(seed + 4)
    dates = pd.date_range(end=datetime.today(), periods=n_days)
    rows = []
    for d in dates:
        baseline_kwh = rng.normal(4200, 200)
        optimized_kwh = baseline_kwh * rng.uniform(0.82, 0.92)
        rows.append({
            "Date": d.date(),
            "Baseline_kWh": round(baseline_kwh, 0),
            "AI_Optimized_kWh": round(optimized_kwh, 0),
            "Savings_%": round(100 * (1 - optimized_kwh / baseline_kwh), 1),
            "Peak_Load_Shifted_kWh": round(rng.uniform(150, 500), 0),
        })
    return pd.DataFrame(rows)


def gen_traceability(n_events=15, seed=RNG_SEED):
    rng = np.random.default_rng(seed + 5)
    machines = ["VMC-01", "VMC-02", "HMC-01"]
    operators = ["Op-A", "Op-B", "Op-C", "Op-D"]
    rows = []
    for i in range(n_events):
        rows.append({
            "Defect_ID": f"DEF-{1000 + i}",
            # FIX: rng.integers() returns numpy.int64, which Python's
            # timedelta() constructor rejects ("unsupported type for
            # timedelta days component: numpy.int64") — this was the
            # exact TypeError that crashed this module. Casting to a
            # native int fixes it.
            "Date": (datetime.today() - timedelta(days=int(rng.integers(0, 30)))).date(),
            "Melt_Lot": f"ML-{rng.integers(100, 999)}",
            "Machine": rng.choice(machines),
            "Operator": rng.choice(operators),
            "Tool_Used": f"Insert-{rng.integers(1, 9)}",
            "Root_Cause_AI": rng.choice([
                "Die temperature deviation",
                "Raw material batch variance",
                "Tool wear beyond threshold",
                "Coolant pressure drop",
                "Operator parameter override",
            ]),
            "Manual_RCA_Time_Hrs": round(rng.uniform(8, 48), 1),
            "AI_RCA_Time_Min": round(rng.uniform(0.3, 3), 1),
        })
    return pd.DataFrame(rows)


# ----------------------------------------------------------------------------
# Sidebar
# ----------------------------------------------------------------------------
st.sidebar.title("⚙️ AI Ops Console")
st.sidebar.markdown(
    "Demo dashboard for **casting + machining auto ancillary units** — "
    "scrap reduction, downtime prediction, tool cost, demand forecasting, "
    "energy savings, and root-cause tracing."
)
st.sidebar.divider()
module = st.sidebar.radio(
    "Select Module",
    [
        "🏠 Overview",
        "🔍 Vision QC — Casting Defects",
        "🛠️ Predictive Maintenance",
        "🔧 Tool Wear Prediction",
        "📦 Demand & Inventory Forecast",
        "⚡ Energy Optimization",
        "🧩 Traceability & Root Cause",
        "💰 ROI Summary",
    ],
)
st.sidebar.divider()
st.sidebar.caption("Synthetic data is for demo purposes. Upload real shop-floor data in any module to override it.")

# ----------------------------------------------------------------------------
# OVERVIEW
# ----------------------------------------------------------------------------
if module == "🏠 Overview":
    st.markdown(
        """
This app demonstrates where AI creates measurable value on the shop floor of a
**casting + machining auto ancillary unit**: cutting scrap, predicting downtime,
optimizing tool cost, sharpening demand forecasts, trimming energy spend, and
collapsing root-cause analysis from days to minutes.

Use the sidebar to open each module. Every module lets you:
- **Generate synthetic data** to simulate a live shop-floor feed
- **Upload your own real data** (CSV/Excel) — it automatically takes priority over synthetic data
- **Visualize** the key metrics AI would surface, with the underlying formula explained
- **Export** results to CSV, JSON, TXT, Word, or PDF
- **Reset** to clear the simulated/uploaded data and start fresh
        """
    )
    st.divider()
    kpi_row([
        ("Casting Rejection", "3.5% → 1.8%", "-48%"),
        ("Machining Tool Cost/Part", "↓ 12–18%", None),
        ("Unplanned Downtime", "↓ 25%", None),
        ("Inventory Reduction", "15–25%", None),
    ])
    st.divider()
    st.markdown("### Module Map")
    c1, c2, c3 = st.columns(3)
    with c1:
        st.markdown("**🔍 Vision QC**\nCatch blowholes, shrinkage, cold shuts, porosity, cracks.")
        st.markdown("**🛠️ Predictive Maintenance**\nFlag bearing/spindle/die failures 2–4 weeks early.")
    with c2:
        st.markdown("**🔧 Tool Wear**\nReplace inserts on condition, not fixed schedule.")
        st.markdown("**📦 Demand Forecast**\nForecast FG/RM better than Excel for JIT OEM supply.")
    with c3:
        st.markdown("**⚡ Energy Optimization**\nShift furnace/machining load to cut bills 8–20%.")
        st.markdown("**🧩 Root Cause**\nCorrelate batch, machine, tool, operator in seconds.")

# ----------------------------------------------------------------------------
# 1. VISION QC
# ----------------------------------------------------------------------------
elif module == "🔍 Vision QC — Casting Defects":
    section_title("🔍", "Vision QC — Casting Defect Detection",
                  "Computer-vision style defect tracking: blowholes, shrinkage, cold shuts, porosity, cracks.")
    c1, c2, c3 = st.columns([2.5, 1, 1])
    with c1:
        n_days = st.slider("Days of shop-floor data to simulate", 7, 90, 30)
    with c2:
        st.write("")
        gen_clicked = st.button("▶️ Generate Synthetic Data", key="gen_casting", use_container_width=True)
    with c3:
        st.write("")
        reset_button("casting_df", key="reset_casting", also_clear=["real_casting_df"])

    if gen_clicked:
        try:
            st.session_state["casting_df"] = gen_casting_defects(n_days)
        except Exception as e:
            st.error(f"Could not generate data: {e}")
            logger.exception("gen_casting_defects failed")

    if "casting_df" not in st.session_state:
        try:
            st.session_state["casting_df"] = gen_casting_defects(n_days)
        except Exception as e:
            st.error(f"Could not generate initial data: {e}")
            logger.exception("gen_casting_defects initial failed")

    upload_real_data("Casting QC Data", "real_casting_df",
                      "Date, Produced, Rejected, Rejection_%, Top_Defect, Melt_Temp_C, Pour_Temp_C, Mold_Temp_C, Humidity_%",
                      "upload_casting", "clear_casting")

    df, is_real = pick_active_df("real_casting_df", "casting_df")

    if df is not None and len(df) > 0:
        if is_real:
            st.info("📁 Showing uploaded real data (synthetic data is hidden while real data is active).")

        has_cols = lambda *cols: all(c in df.columns for c in cols)

        kpi_items = []
        if has_cols("Rejection_%"):
            kpi_items.append(("Avg Rejection %", f"{df['Rejection_%'].mean():.2f}%", None))
        if has_cols("Produced"):
            kpi_items.append(("Total Produced", f"{df['Produced'].sum():,.0f}", None))
        if has_cols("Rejected"):
            kpi_items.append(("Total Rejected", f"{df['Rejected'].sum():,.0f}", None))
        if has_cols("Top_Defect"):
            kpi_items.append(("Most Common Defect", df["Top_Defect"].mode()[0], None))
        if kpi_items:
            kpi_row(kpi_items)

        if has_cols("Date", "Rejection_%"):
            st.markdown("#### Rejection % Trend (AI flags drift before it becomes a batch problem)")
            st.line_chart(df.set_index("Date")[["Rejection_%"]])

        cc1, cc2 = st.columns(2)
        with cc1:
            if has_cols("Top_Defect"):
                st.markdown("#### Defect Type Distribution")
                st.bar_chart(df["Top_Defect"].value_counts())
        with cc2:
            if has_cols("Pour_Temp_C", "Rejection_%"):
                st.markdown("#### Process Parameters vs Rejection")
                st.scatter_chart(df, x="Pour_Temp_C", y="Rejection_%")

        formula_box("""
        **📐 Rejection % Formula**

        **Rejection % = (Rejected_Units / Produced_Units) × 100**

        AI watches this trend day over day rather than waiting for a monthly quality report —
        a rising trend gets flagged *before* it becomes a full batch write-off. The synthetic
        generator also models a downward drift over time, simulating the improvement typically
        seen after deploying vision-based defect detection on the line.

        **Defect drivers tracked:** Melt_Temp_C, Pour_Temp_C, Mold_Temp_C, Humidity_% — these are
        the process parameters most correlated with blowholes, shrinkage, cold shuts, porosity,
        and surface cracks in casting.

        **ROI:** Typical deployments cut rejection rate by ~45-50% (3.5% → 1.8%), directly reducing
        scrap and rework cost.
        """)

        st.markdown("#### Raw Data")
        st.dataframe(df, use_container_width=True, height=280)

        formulas_dict = {"Rejection % Formula": "Rejection % = (Rejected / Produced) x 100"}
        full_export_buttons(df, "Vision QC - Casting Defects", formulas_dict, "casting", "casting_defects")
    else:
        st.info("No data yet — click **Generate Synthetic Data** or upload a real QC file above.")

# ----------------------------------------------------------------------------
# 2. PREDICTIVE MAINTENANCE
# ----------------------------------------------------------------------------
elif module == "🛠️ Predictive Maintenance":
    section_title("🛠️", "Predictive Maintenance — Machine Health",
                  "Vibration, temperature, current sensors → failure-risk scoring for VMC/HMC/Press machines.")
    c1, c2, c3 = st.columns([2.5, 1, 1])
    with c1:
        n_hours = st.slider("Hours of sensor history to simulate", 48, 720, 240, step=24)
    with c2:
        st.write("")
        gen_clicked = st.button("▶️ Generate Synthetic Data", key="gen_health", use_container_width=True)
    with c3:
        st.write("")
        reset_button("health_df", key="reset_health", also_clear=["real_health_df"])

    if gen_clicked:
        try:
            st.session_state["health_df"] = gen_machine_health(n_hours)
        except Exception as e:
            st.error(f"Could not generate data: {e}")
            logger.exception("gen_machine_health failed")

    if "health_df" not in st.session_state:
        try:
            st.session_state["health_df"] = gen_machine_health(n_hours)
        except Exception as e:
            st.error(f"Could not generate initial data: {e}")
            logger.exception("gen_machine_health initial failed")

    upload_real_data("Machine Health Sensor Data", "real_health_df",
                      "Timestamp, Machine, Vibration_mm_s, Temperature_C, Current_A, Failure_Risk_%, Predicted_Days_to_Failure",
                      "upload_health", "clear_health")

    df, is_real = pick_active_df("real_health_df", "health_df")

    if df is not None and len(df) > 0 and "Machine" in df.columns:
        if is_real:
            st.info("📁 Showing uploaded real data (synthetic data is hidden while real data is active).")

        machine_sel = st.selectbox("Select machine", sorted(df["Machine"].unique()))
        mdf = df[df["Machine"] == machine_sel]
        if "Timestamp" in mdf.columns:
            mdf = mdf.sort_values("Timestamp")
        latest = mdf.iloc[-1]

        kpi_items = []
        if "Failure_Risk_%" in mdf.columns:
            kpi_items.append(("Current Failure Risk", f"{latest['Failure_Risk_%']:.0f}%", None))
        if "Predicted_Days_to_Failure" in mdf.columns:
            kpi_items.append(("Predicted Days to Failure", f"{latest['Predicted_Days_to_Failure']:.1f} d", None))
        if "Vibration_mm_s" in mdf.columns:
            kpi_items.append(("Vibration", f"{latest['Vibration_mm_s']} mm/s", None))
        if "Temperature_C" in mdf.columns:
            kpi_items.append(("Temperature", f"{latest['Temperature_C']} °C", None))
        if kpi_items:
            kpi_row(kpi_items)

        if "Timestamp" in mdf.columns and "Failure_Risk_%" in mdf.columns:
            st.markdown(f"#### Failure Risk Trend — {machine_sel}")
            st.line_chart(mdf.set_index("Timestamp")[["Failure_Risk_%"]])

        cc1, cc2 = st.columns(2)
        with cc1:
            if "Timestamp" in mdf.columns and {"Vibration_mm_s", "Current_A"}.issubset(mdf.columns):
                st.markdown("#### Vibration & Current")
                st.line_chart(mdf.set_index("Timestamp")[["Vibration_mm_s", "Current_A"]])
        with cc2:
            if "Timestamp" in mdf.columns and "Temperature_C" in mdf.columns:
                st.markdown("#### Temperature")
                st.line_chart(mdf.set_index("Timestamp")[["Temperature_C"]])

        if "Failure_Risk_%" in mdf.columns:
            risk_val = latest["Failure_Risk_%"]
            if risk_val > 65:
                st.error(f"⚠️ {machine_sel} is showing elevated failure risk — schedule maintenance within "
                         f"{latest.get('Predicted_Days_to_Failure', 'a few')} days.")
            elif risk_val > 40:
                st.warning(f"🟡 {machine_sel} risk trending up — monitor closely.")
            else:
                st.success(f"🟢 {machine_sel} operating within normal parameters.")

        formula_box("""
        **📐 Failure Risk Score Formula**

        **Failure_Risk = min(99, max(1, (Vibration/6 + Temperature/90 + Current/18) × 33))**

        This is a simple weighted composite of three sensor channels, each normalized against a
        rough "danger ceiling" (6 mm/s vibration, 90°C temperature, 18A current) before being
        averaged and scaled to a 1-99% risk score. As a machine drifts toward failure, all three
        readings tend to climb together, which the formula picks up as a rising risk score.

        **Action thresholds:** >65% → schedule maintenance now · 40-65% → monitor closely · <40% → normal.

        **ROI:** Predictive maintenance catches degrading bearings/spindles/dies 2-4 weeks before
        failure, cutting unplanned downtime by ~25%.
        """)

        st.markdown("#### Raw Data")
        st.dataframe(df, use_container_width=True, height=280)

        formulas_dict = {
            "Failure Risk Formula": "Failure_Risk = min(99, max(1, (Vibration/6 + Temperature/90 + Current/18) x 33))",
        }
        full_export_buttons(df, "Predictive Maintenance - Machine Health", formulas_dict, "health", "machine_health")
    else:
        st.info("No data yet — click **Generate Synthetic Data** or upload a real sensor file above.")

# ----------------------------------------------------------------------------
# 3. TOOL WEAR
# ----------------------------------------------------------------------------
elif module == "🔧 Tool Wear Prediction":
    section_title("🔧", "Tool Wear Prediction",
                  "Spindle load, vibration, and cycle count → condition-based insert/tool replacement.")
    c1, c2, c3 = st.columns([2.5, 1, 1])
    with c1:
        n_records = st.slider("Number of tools to simulate", 50, 500, 200, step=50)
    with c2:
        st.write("")
        gen_clicked = st.button("▶️ Generate Synthetic Data", key="gen_tool", use_container_width=True)
    with c3:
        st.write("")
        reset_button("tool_df", key="reset_tool", also_clear=["real_tool_df"])

    if gen_clicked:
        try:
            st.session_state["tool_df"] = gen_tool_wear(n_records)
        except Exception as e:
            st.error(f"Could not generate data: {e}")
            logger.exception("gen_tool_wear failed")

    if "tool_df" not in st.session_state:
        try:
            st.session_state["tool_df"] = gen_tool_wear(n_records)
        except Exception as e:
            st.error(f"Could not generate initial data: {e}")
            logger.exception("gen_tool_wear initial failed")

    upload_real_data("Tool Wear Data", "real_tool_df",
                      "Tool_ID, Tool_Type, Cycles_Run, Wear_%, Spindle_Load_%, Vibration_mm_s, Recommendation",
                      "upload_tool", "clear_tool")

    df, is_real = pick_active_df("real_tool_df", "tool_df")

    if df is not None and len(df) > 0:
        if is_real:
            st.info("📁 Showing uploaded real data (synthetic data is hidden while real data is active).")

        has_cols = lambda *cols: all(c in df.columns for c in cols)

        kpi_items = []
        if "Wear_%" in df.columns:
            kpi_items.append(("Avg Wear %", f"{df['Wear_%'].mean():.1f}%", None))
        if "Recommendation" in df.columns:
            kpi_items.append(("Tools to Change Now", int((df["Recommendation"] == "Change Now").sum()), None))
            kpi_items.append(("Tools to Monitor", int((df["Recommendation"] == "Monitor").sum()), None))
            kpi_items.append(("Tools OK", int((df["Recommendation"] == "OK").sum()), None))
        if kpi_items:
            kpi_row(kpi_items)

        cc1, cc2 = st.columns(2)
        with cc1:
            if has_cols("Tool_Type", "Wear_%"):
                st.markdown("#### Wear % by Tool Type")
                st.bar_chart(df.groupby("Tool_Type")["Wear_%"].mean())
        with cc2:
            if has_cols("Wear_%", "Spindle_Load_%"):
                st.markdown("#### Spindle Load vs Wear")
                st.scatter_chart(df, x="Wear_%", y="Spindle_Load_%")

        formula_box("""
        **📐 Tool Wear & Recommendation Logic**

        **Wear % ≈ (Cycles_Run / Expected_Tool_Life) + sensor noise**, capped to 0-100%.

        **Recommendation thresholds:**
        - Wear ≤ 55% → **OK**
        - 55% < Wear ≤ 80% → **Monitor**
        - Wear > 80% → **Change Now**

        Spindle load and vibration both rise alongside wear (a worn edge cuts less efficiently),
        so AI cross-checks wear estimates against these two signals rather than relying on a fixed
        cycle-count schedule alone.

        **ROI:** Condition-based replacement (instead of fixed-interval changes) typically cuts
        machining tool cost per part by 12-18% while avoiding tool-failure scrap.
        """)

        if "Recommendation" in df.columns:
            st.markdown("#### Tools Needing Attention")
            st.dataframe(
                df[df["Recommendation"] != "OK"].sort_values("Wear_%", ascending=False)
                if "Wear_%" in df.columns else df[df["Recommendation"] != "OK"],
                use_container_width=True, height=240,
            )

        st.markdown("#### Full Tool Data")
        st.dataframe(df, use_container_width=True, height=280)

        formulas_dict = {
            "Wear % Logic": "Wear % ~ Cycles_Run / Expected_Tool_Life + sensor noise (0-100%)",
            "Recommendation Thresholds": "<=55% OK | 55-80% Monitor | >80% Change Now",
        }
        full_export_buttons(df, "Tool Wear Prediction", formulas_dict, "tool", "tool_wear")
    else:
        st.info("No data yet — click **Generate Synthetic Data** or upload a real tool-wear file above.")

# ----------------------------------------------------------------------------
# 4. DEMAND & INVENTORY FORECAST
# ----------------------------------------------------------------------------
elif module == "📦 Demand & Inventory Forecast":
    section_title("📦", "Demand & Inventory Forecast",
                  "AI forecast vs actual demand by part, with safety stock and reorder points.")
    c1, c2, c3 = st.columns([2.5, 1, 1])
    with c1:
        n_weeks = st.slider("Weeks of history to simulate", 12, 52, 26)
    with c2:
        st.write("")
        gen_clicked = st.button("▶️ Generate Synthetic Data", key="gen_demand", use_container_width=True)
    with c3:
        st.write("")
        reset_button("demand_df", key="reset_demand", also_clear=["real_demand_df"])

    if gen_clicked:
        try:
            st.session_state["demand_df"] = gen_demand_forecast(n_weeks)
        except Exception as e:
            st.error(f"Could not generate data: {e}")
            logger.exception("gen_demand_forecast failed")

    if "demand_df" not in st.session_state:
        try:
            st.session_state["demand_df"] = gen_demand_forecast(n_weeks)
        except Exception as e:
            st.error(f"Could not generate initial data: {e}")
            logger.exception("gen_demand_forecast initial failed")

    upload_real_data("Demand/Inventory Data", "real_demand_df",
                      "Week, Part, Actual_Demand, AI_Forecast, Safety_Stock, Reorder_Point",
                      "upload_demand", "clear_demand")

    df, is_real = pick_active_df("real_demand_df", "demand_df")

    if df is not None and len(df) > 0 and "Part" in df.columns:
        if is_real:
            st.info("📁 Showing uploaded real data (synthetic data is hidden while real data is active).")

        part_sel = st.selectbox("Select part", sorted(df["Part"].unique()))
        pdf_ = df[df["Part"] == part_sel]
        if "Week" in pdf_.columns:
            pdf_ = pdf_.sort_values("Week")

        kpi_items = []
        if {"Actual_Demand", "AI_Forecast"}.issubset(pdf_.columns):
            valid = pdf_.dropna(subset=["Actual_Demand"])
            if len(valid) > 0 and (valid["Actual_Demand"] != 0).any():
                mape = (abs(valid["Actual_Demand"] - valid["AI_Forecast"]) / valid["Actual_Demand"]).mean() * 100
                kpi_items.append(("Forecast Accuracy (MAPE)", f"{100 - mape:.1f}%", None))
                kpi_items.append(("Avg Weekly Demand", f"{valid['Actual_Demand'].mean():.0f}", None))
        if "Safety_Stock" in pdf_.columns and len(pdf_) > 0:
            kpi_items.append(("Latest Safety Stock", int(pdf_.iloc[-1]["Safety_Stock"]), None))
        if "Reorder_Point" in pdf_.columns and len(pdf_) > 0:
            kpi_items.append(("Latest Reorder Point", int(pdf_.iloc[-1]["Reorder_Point"]), None))
        if kpi_items:
            kpi_row(kpi_items)

        if "Week" in pdf_.columns and {"Actual_Demand", "AI_Forecast"}.issubset(pdf_.columns):
            st.markdown(f"#### Actual vs AI Forecast — {part_sel}")
            st.line_chart(pdf_.set_index("Week")[["Actual_Demand", "AI_Forecast"]])

        if "Week" in pdf_.columns and {"Safety_Stock", "Reorder_Point"}.issubset(pdf_.columns):
            st.markdown("#### Safety Stock & Reorder Point")
            st.line_chart(pdf_.set_index("Week")[["Safety_Stock", "Reorder_Point"]])

        formula_box("""
        **📐 Forecast Accuracy, Safety Stock & Reorder Point Formulas**

        **MAPE (Mean Absolute Percentage Error) = avg(|Actual − Forecast| / Actual) × 100**
        Forecast Accuracy is reported as **100% − MAPE**.

        **Safety_Stock ≈ AI_Forecast × 12%** — buffer held to absorb demand variability and
        lead-time uncertainty.

        **Reorder_Point ≈ AI_Forecast × 35%** — inventory level at which a new purchase order
        should be triggered.

        **ROI:** Sharper forecasts typically reduce raw-material/finished-goods inventory by
        15-25% while still protecting JIT OEM supply commitments.
        """)

        st.markdown("#### Raw Data")
        st.dataframe(df, use_container_width=True, height=280)

        formulas_dict = {
            "MAPE / Accuracy": "Accuracy % = 100 - avg(|Actual - Forecast| / Actual) x 100",
            "Safety Stock": "Safety_Stock ~ AI_Forecast x 12%",
            "Reorder Point": "Reorder_Point ~ AI_Forecast x 35%",
        }
        full_export_buttons(df, "Demand & Inventory Forecast", formulas_dict, "demand", "demand_forecast")
    else:
        st.info("No data yet — click **Generate Synthetic Data** or upload a real demand file above.")

# ----------------------------------------------------------------------------
# 5. ENERGY OPTIMIZATION
# ----------------------------------------------------------------------------
elif module == "⚡ Energy Optimization":
    section_title("⚡", "Energy Optimization",
                  "Baseline vs AI-optimized furnace/machining load, with peak-shift savings.")
    c1, c2, c3 = st.columns([2.5, 1, 1])
    with c1:
        n_days = st.slider("Days of energy data to simulate", 7, 90, 30, key="energy_days")
    with c2:
        st.write("")
        gen_clicked = st.button("▶️ Generate Synthetic Data", key="gen_energy", use_container_width=True)
    with c3:
        st.write("")
        reset_button("energy_df", key="reset_energy", also_clear=["real_energy_df"])

    if gen_clicked:
        try:
            st.session_state["energy_df"] = gen_energy(n_days)
        except Exception as e:
            st.error(f"Could not generate data: {e}")
            logger.exception("gen_energy failed")

    if "energy_df" not in st.session_state:
        try:
            st.session_state["energy_df"] = gen_energy(n_days)
        except Exception as e:
            st.error(f"Could not generate initial data: {e}")
            logger.exception("gen_energy initial failed")

    upload_real_data("Energy Consumption Data", "real_energy_df",
                      "Date, Baseline_kWh, AI_Optimized_kWh, Savings_%, Peak_Load_Shifted_kWh",
                      "upload_energy", "clear_energy")

    df, is_real = pick_active_df("real_energy_df", "energy_df")

    if df is not None and len(df) > 0:
        if is_real:
            st.info("📁 Showing uploaded real data (synthetic data is hidden while real data is active).")

        kpi_items = []
        if {"Baseline_kWh", "AI_Optimized_kWh"}.issubset(df.columns):
            total_baseline = df["Baseline_kWh"].sum()
            total_optimized = df["AI_Optimized_kWh"].sum()
            savings_pct = 100 * (1 - total_optimized / total_baseline) if total_baseline else 0
            kpi_items.append(("Total Baseline kWh", f"{total_baseline:,.0f}", None))
            kpi_items.append(("Total AI-Optimized kWh", f"{total_optimized:,.0f}", None))
            kpi_items.append(("Avg Savings %", f"{savings_pct:.1f}%", None))
        if "Peak_Load_Shifted_kWh" in df.columns:
            kpi_items.append(("Peak Load Shifted (avg/day)", f"{df['Peak_Load_Shifted_kWh'].mean():.0f} kWh", None))
        if kpi_items:
            kpi_row(kpi_items)

        if "Date" in df.columns and {"Baseline_kWh", "AI_Optimized_kWh"}.issubset(df.columns):
            st.markdown("#### Baseline vs AI-Optimized Consumption")
            st.line_chart(df.set_index("Date")[["Baseline_kWh", "AI_Optimized_kWh"]])

        if "Date" in df.columns and "Savings_%" in df.columns:
            st.markdown("#### Daily Savings %")
            st.bar_chart(df.set_index("Date")[["Savings_%"]])

        formula_box("""
        **📐 Energy Savings Formula**

        **Savings % = (1 − AI_Optimized_kWh / Baseline_kWh) × 100**

        AI shifts furnace pre-heat cycles, machining schedules, and compressed-air usage away from
        peak-tariff windows and trims idle-running load, without changing production volume.
        **Peak_Load_Shifted_kWh** tracks how much load was moved out of the most expensive tariff
        band each day.

        **ROI:** Typical deployments cut the energy bill by 8-20% depending on tariff structure and
        existing equipment efficiency.
        """)

        st.markdown("#### Raw Data")
        st.dataframe(df, use_container_width=True, height=280)

        formulas_dict = {"Energy Savings Formula": "Savings % = (1 - AI_Optimized_kWh / Baseline_kWh) x 100"}
        full_export_buttons(df, "Energy Optimization", formulas_dict, "energy", "energy_optimization")
    else:
        st.info("No data yet — click **Generate Synthetic Data** or upload a real energy file above.")

# ----------------------------------------------------------------------------
# 6. TRACEABILITY / ROOT CAUSE
# ----------------------------------------------------------------------------
elif module == "🧩 Traceability & Root Cause":
    section_title("🧩", "Traceability & Root Cause Analysis",
                  "AI correlates melt lot, machine, tool, and operator to find root cause in seconds.")
    c1, c2, c3 = st.columns([2.5, 1, 1])
    with c1:
        n_events = st.slider("Number of defect events to simulate", 5, 50, 15)
    with c2:
        st.write("")
        gen_clicked = st.button("▶️ Generate Synthetic Data", key="gen_trace", use_container_width=True)
    with c3:
        st.write("")
        reset_button("trace_df", key="reset_trace", also_clear=["real_trace_df"])

    if gen_clicked:
        try:
            st.session_state["trace_df"] = gen_traceability(n_events)
        except Exception as e:
            st.error(f"Could not generate data: {e}")
            logger.exception("gen_traceability failed")

    if "trace_df" not in st.session_state:
        try:
            st.session_state["trace_df"] = gen_traceability(n_events)
        except Exception as e:
            st.error(f"Could not generate initial data: {e}")
            logger.exception("gen_traceability initial failed")

    upload_real_data("Traceability / RCA Data", "real_trace_df",
                      "Defect_ID, Date, Melt_Lot, Machine, Operator, Tool_Used, Root_Cause_AI, Manual_RCA_Time_Hrs, AI_RCA_Time_Min",
                      "upload_trace", "clear_trace")

    df, is_real = pick_active_df("real_trace_df", "trace_df")

    if df is not None and len(df) > 0:
        if is_real:
            st.info("📁 Showing uploaded real data (synthetic data is hidden while real data is active).")

        kpi_items = []
        if {"Manual_RCA_Time_Hrs", "AI_RCA_Time_Min"}.issubset(df.columns) and df["AI_RCA_Time_Min"].mean() > 0:
            kpi_items.append(("Avg Manual RCA Time", f"{df['Manual_RCA_Time_Hrs'].mean():.1f} hrs", None))
            kpi_items.append(("Avg AI RCA Time", f"{df['AI_RCA_Time_Min'].mean():.1f} min", None))
            speedup = (df["Manual_RCA_Time_Hrs"].mean() * 60 / df["AI_RCA_Time_Min"].mean())
            kpi_items.append(("Speed-up Factor", f"{speedup:.0f}x", None))
        if "Root_Cause_AI" in df.columns:
            kpi_items.append(("Top Root Cause", df["Root_Cause_AI"].mode()[0], None))
        if kpi_items:
            kpi_row(kpi_items)

        if "Root_Cause_AI" in df.columns:
            st.markdown("#### Root Cause Frequency")
            st.bar_chart(df["Root_Cause_AI"].value_counts())

        if {"Defect_ID", "Manual_RCA_Time_Hrs", "AI_RCA_Time_Min"}.issubset(df.columns):
            st.markdown("#### Manual vs AI RCA Time (per event)")
            comp = df[["Defect_ID", "Manual_RCA_Time_Hrs"]].copy()
            comp["AI_RCA_Time_Hrs"] = df["AI_RCA_Time_Min"] / 60
            st.bar_chart(comp.set_index("Defect_ID")[["Manual_RCA_Time_Hrs", "AI_RCA_Time_Hrs"]])

        formula_box("""
        **📐 Root Cause Speed-up Formula**

        **Speed-up Factor = (Avg_Manual_RCA_Time_Hrs × 60) / Avg_AI_RCA_Time_Min**

        AI cross-references melt lot, machine, tool, and operator against every historical defect
        event to surface the most statistically likely root cause in seconds, instead of an
        engineer manually pulling logs across multiple systems over hours or days.

        **ROI:** Root-cause analysis time typically collapses from ~2 days (manual) to ~30 seconds
        (AI-assisted) — a >99% reduction, freeing engineering time for prevention instead of
        firefighting.
        """)

        st.markdown("#### Event Log")
        st.dataframe(df, use_container_width=True, height=300)

        formulas_dict = {"RCA Speed-up Formula": "Speed-up = (Avg_Manual_RCA_Hrs x 60) / Avg_AI_RCA_Min"}
        full_export_buttons(df, "Traceability & Root Cause", formulas_dict, "trace", "traceability_rca")
    else:
        st.info("No data yet — click **Generate Synthetic Data** or upload a real RCA log above.")

# ----------------------------------------------------------------------------
# 7. ROI SUMMARY
# ----------------------------------------------------------------------------
elif module == "💰 ROI Summary":
    section_title("💰", "ROI Summary",
                  "Typical results reported by Indian auto ancillary units after AI deployment.")
    roi_df = pd.DataFrame([
        {"Area": "Casting Rejection Rate", "Before": "3.5%", "After": "1.8%", "Improvement": "-48%"},
        {"Area": "Machining Tool Cost / Part", "Before": "Baseline", "After": "↓ 12–18%", "Improvement": "12-18%"},
        {"Area": "Unplanned Downtime", "Before": "Baseline", "After": "↓ 25%", "Improvement": "25%"},
        {"Area": "Inventory / Working Capital", "Before": "Baseline", "After": "↓ 15–25%", "Improvement": "15-25%"},
        {"Area": "Energy Bill", "Before": "Baseline", "After": "↓ 8–20%", "Improvement": "8-20%"},
        {"Area": "Customer Complaints (Vision QC)", "Before": "Baseline", "After": "↓ 50–80%", "Improvement": "50-80%"},
        {"Area": "Root Cause Analysis Time", "Before": "~2 days", "After": "~30 sec", "Improvement": ">99%"},
    ])
    st.dataframe(roi_df, use_container_width=True, height=300)

    formula_box("""
    **📐 How These Numbers Are Derived**

    Each row reflects the typical *before vs after* delta reported across the six modules in this
    app once deployed on a live line: Vision QC drives the rejection-rate and complaint-reduction
    rows, Predictive Maintenance drives downtime, Tool Wear drives tool cost, Demand Forecast
    drives inventory/working capital, Energy Optimization drives the energy-bill row, and
    Traceability drives RCA time. Treat these as directional industry benchmarks, not a
    site-specific guarantee — actual results depend on baseline process maturity and the quality
    of sensor/data coverage at rollout.
    """)

    st.markdown("#### Suggested Rollout Sequence (Low Budget First)")
    st.markdown(
        """
1. **Vision QC pilot** — 2–3 cameras + off-the-shelf vision software on highest-rejection line. Fastest ROI, easiest operator buy-in.
2. **Sensor retrofit** — vibration + current sensors on top 3 bottleneck machines (~₹50k/machine).
3. **Use existing data** — CMM logs, rejection logs, tool-change sheets feed the first predictive models.
4. **Scale to scheduling, demand forecasting, and energy** once shop floor trusts the QC/maintenance pilots.
        """
    )

    formulas_dict = {"Methodology": "Before/after deltas aggregated from the Vision QC, Predictive Maintenance, Tool Wear, Demand Forecast, Energy, and Traceability modules."}
    full_export_buttons(roi_df, "ROI Summary", formulas_dict, "roi", "roi_summary")

st.sidebar.divider()
st.sidebar.caption("Demo build · All figures synthetic / illustrative unless real data is uploaded · Not a substitute for a site-specific audit.")
