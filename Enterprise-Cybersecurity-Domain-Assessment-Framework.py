#CyberGuard Pro — Enterprise Cybersecurity Domain Assessment Framework
#100+ Security Domains aligned with NIST CSF 2.0 & CIS Controls v8
#Developed by Randy Singh from Kalsnet (KNet) Consulting
import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
from datetime import datetime
import json
from io import BytesIO
# ── Optional export dependencies (graceful degradation) ────────────────────
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import (SimpleDocTemplate, Table, TableStyle,
                                    Paragraph, Spacer, HRFlowable, PageBreak)
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors as rl_colors
    from reportlab.lib.units import inch
    REPORTLAB_OK = True
except ImportError:
    REPORTLAB_OK = False
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_OK = True
except ImportError:
    DOCX_OK = False
# ════════════════════════════════════════════════════════════════════════════
# PAGE CONFIG
# ════════════════════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="CyberGuard Pro — Cybersecurity Assessment Framework",
    page_icon="🛡️",
    layout="wide",
    initial_sidebar_state="expanded",
)
# ════════════════════════════════════════════════════════════════════════════
# GLOBAL CSS
# ════════════════════════════════════════════════════════════════════════════
st.markdown("""
<style>
.app-title{font-size:2.2rem;font-weight:900;color:#fff;text-align:center;letter-spacing:.5px;margin:0}
.app-subtitle{font-size:1rem;font-weight:700;color:#90CAF9;text-align:center;margin:3px 0 0}
.header-wrap{background:linear-gradient(100deg,#002B7F 0%,#1565C0 55%,#0288D1 100%);
  border-radius:14px;padding:24px 30px 20px;margin-bottom:18px;
  box-shadow:0 6px 24px rgba(0,43,127,.3)}
.kpi-card{border-radius:12px;padding:16px 18px;color:#fff;text-align:center;
  box-shadow:0 4px 16px rgba(0,0,0,.14)}
.kpi-val{font-size:1.9rem;font-weight:800;line-height:1.1}
.kpi-lbl{font-size:.8rem;opacity:.9;margin-top:2px}
.maturity-0{background:#C62828;color:#fff;padding:2px 10px;border-radius:12px;font-size:.75rem;font-weight:700}
.maturity-1{background:#E65100;color:#fff;padding:2px 10px;border-radius:12px;font-size:.75rem;font-weight:700}
.maturity-2{background:#F9A825;color:#fff;padding:2px 10px;border-radius:12px;font-size:.75rem;font-weight:700}
.maturity-3{background:#7CB342;color:#fff;padding:2px 10px;border-radius:12px;font-size:.75rem;font-weight:700}
.maturity-4{background:#2E7D32;color:#fff;padding:2px 10px;border-radius:12px;font-size:.75rem;font-weight:700}
.maturity-5{background:#1565C0;color:#fff;padding:2px 10px;border-radius:12px;font-size:.75rem;font-weight:700}
.stTabs [data-baseweb="tab"]{font-weight:600}
.domain-card{background:#fff;border:1px solid #D6E4F0;border-radius:10px;padding:12px 16px;margin-bottom:8px}
.gap-critical{border-left:4px solid #C62828}
.gap-high{border-left:4px solid #E65100}
.gap-medium{border-left:4px solid #F9A825}
.gap-low{border-left:4px solid #2E7D32}
.info-box{background:#EEF4FF;border:1px solid #BBD6F5;border-radius:10px;padding:14px 18px;margin-bottom:10px}
.formula-box{background:#FFF8E1;border:1px solid #F0D98C;border-radius:8px;padding:10px 14px;margin:6px 0;font-family:monospace;font-size:.85rem}
</style>
""", unsafe_allow_html=True)
# ════════════════════════════════════════════════════════════════════════════
# FRAMEWORK DATA — 108 CYBERSECURITY DOMAINS
# NIST CSF 2.0 Functions: GOVERN, IDENTIFY, PROTECT, DETECT, RESPOND, RECOVER
# Each domain mapped to NIST CSF 2.0 category + CIS Controls v8
# ════════════════════════════════════════════════════════════════════════════
FRAMEWORK = {
    "GOVERN (GV)": {
        "color": "#4A148C",
        "description": "Establish and monitor cybersecurity risk management strategy, expectations, and policy.",
        "domains": [
            ("GV-01", "Organizational Cybersecurity Strategy", "GV.OC", "CIS 17", "Enterprise security strategy aligned with business objectives and risk appetite."),
            ("GV-02", "Cybersecurity Governance Structure", "GV.OC", "CIS 17", "Board oversight, CISO reporting lines, security steering committees."),
            ("GV-03", "Security Policy Management", "GV.PO", "CIS 14", "Development, approval, communication, and lifecycle of security policies."),
            ("GV-04", "Roles, Responsibilities & Authorities", "GV.RR", "CIS 17", "RACI for security functions, accountability frameworks, job descriptions."),
            ("GV-05", "Risk Management Strategy", "GV.RM", "CIS 17", "Enterprise risk framework, risk appetite statements, risk tolerance thresholds."),
            ("GV-06", "Risk Assessment Program", "GV.RM", "CIS 17", "Periodic risk assessments, threat modeling, risk registers, heat maps."),
            ("GV-07", "Cyber Risk Quantification", "GV.RM", "CIS 17", "FAIR analysis, monetary risk quantification, risk-based prioritization."),
            ("GV-08", "Third-Party Risk Management", "GV.SC", "CIS 15", "Vendor security assessments, contract security clauses, ongoing monitoring."),
            ("GV-09", "Supply Chain Security", "GV.SC", "CIS 15", "Software supply chain integrity, SBOM management, hardware provenance."),
            ("GV-10", "Compliance Management", "GV.OC", "CIS 17", "Regulatory mapping (GDPR, HIPAA, PCI-DSS, SOX), audit management."),
            ("GV-11", "Security Metrics & Reporting", "GV.OV", "CIS 17", "KPIs/KRIs, executive dashboards, board reporting, trend analysis."),
            ("GV-12", "Security Budget & Investment", "GV.RR", "CIS 17", "Security spend optimization, ROI analysis, resource allocation."),
            ("GV-13", "Cyber Insurance Management", "GV.RM", "CIS 17", "Coverage assessment, policy requirements, claims readiness."),
            ("GV-14", "Legal & Regulatory Liaison", "GV.OC", "CIS 17", "Breach notification laws, legal hold, regulator relationships."),
            ("GV-15", "Security Culture Program", "GV.RR", "CIS 14", "Culture measurement, behavioral change, security champions network."),
            ("GV-16", "Audit & Assurance Management", "GV.OV", "CIS 17", "Internal/external audit coordination, findings remediation tracking."),
        ],
    },
    "IDENTIFY (ID)": {
        "color": "#1565C0",
        "description": "Understand the organization's assets, risks, and vulnerabilities.",
        "domains": [
            ("ID-01", "Hardware Asset Inventory", "ID.AM", "CIS 1", "Discovery and inventory of all hardware: servers, endpoints, IoT, OT devices."),
            ("ID-02", "Software Asset Inventory", "ID.AM", "CIS 2", "Authorized software catalog, license management, shadow IT detection."),
            ("ID-03", "Data Asset Inventory & Classification", "ID.AM", "CIS 3", "Data discovery, classification schemes, data flow mapping, crown jewels."),
            ("ID-04", "Cloud Asset Management", "ID.AM", "CIS 1,2", "Multi-cloud inventory, CSPM, cloud resource tagging, orphaned resources."),
            ("ID-05", "Identity Inventory", "ID.AM", "CIS 5", "Human and non-human identity catalog, service accounts, API keys."),
            ("ID-06", "Network Architecture Documentation", "ID.AM", "CIS 12", "Network diagrams, data flows, trust boundaries, segmentation maps."),
            ("ID-07", "Vulnerability Management Program", "ID.RA", "CIS 7", "Scanning cadence, risk-based prioritization (CVSS/EPSS), remediation SLAs."),
            ("ID-08", "Penetration Testing Program", "ID.RA", "CIS 18", "Internal/external pentests, red team exercises, purple teaming."),
            ("ID-09", "Threat Intelligence Program", "ID.RA", "CIS 17", "TI feeds, IOC management, threat actor tracking, intelligence sharing (ISACs)."),
            ("ID-10", "Attack Surface Management", "ID.RA", "CIS 1,7", "External attack surface discovery, digital footprint, exposed services."),
            ("ID-11", "Business Impact Analysis", "ID.RA", "CIS 11", "Critical process identification, RTO/RPO definition, dependency mapping."),
            ("ID-12", "Threat Modeling", "ID.RA", "CIS 16", "STRIDE/PASTA analysis, design-phase security review, abuse cases."),
            ("ID-13", "Security Baseline Assessment", "ID.RA", "CIS 4", "Configuration benchmarks (CIS/STIG), baseline drift detection."),
            ("ID-14", "Insider Threat Program", "ID.RA", "CIS 14", "Insider risk indicators, UEBA, privileged user monitoring."),
            ("ID-15", "Crown Jewels Analysis", "ID.AM", "CIS 3", "Mission-critical asset identification, high-value target protection."),
            ("ID-16", "OT/ICS Asset Discovery", "ID.AM", "CIS 1", "Industrial control system inventory, SCADA visibility, passive discovery."),
        ],
    },
    "PROTECT (PR)": {
        "color": "#2E7D32",
        "description": "Implement safeguards to secure assets and limit the impact of cybersecurity events.",
        "domains": [
            ("PR-01", "Identity & Access Management (IAM)", "PR.AA", "CIS 5,6", "Identity lifecycle, provisioning/deprovisioning, access certification."),
            ("PR-02", "Multi-Factor Authentication", "PR.AA", "CIS 6", "MFA coverage for all users, phishing-resistant methods (FIDO2, passkeys)."),
            ("PR-03", "Privileged Access Management (PAM)", "PR.AA", "CIS 5", "Vault-based credential management, session recording, JIT elevation."),
            ("PR-04", "Single Sign-On & Federation", "PR.AA", "CIS 6", "SSO deployment, SAML/OIDC federation, identity provider hardening."),
            ("PR-05", "Zero Trust Architecture", "PR.AA", "CIS 6,12,13", "Never trust always verify, micro-perimeters, continuous verification."),
            ("PR-06", "Endpoint Protection (EPP/EDR)", "PR.PS", "CIS 10", "Anti-malware, EDR deployment, endpoint hardening, device control."),
            ("PR-07", "Mobile Device Management", "PR.PS", "CIS 4", "MDM/MAM enrollment, BYOD policy, mobile threat defense."),
            ("PR-08", "Secure Configuration Management", "PR.PS", "CIS 4", "Hardened images, CIS benchmarks, configuration drift monitoring."),
            ("PR-09", "Patch Management", "PR.PS", "CIS 7", "Patch deployment SLAs, emergency patching, virtual patching."),
            ("PR-10", "Application Whitelisting", "PR.PS", "CIS 2", "Allowlisting execution control, script blocking, LOLBin protections."),
            ("PR-11", "Network Segmentation", "PR.IR", "CIS 12", "VLAN/micro-segmentation, east-west traffic control, DMZ architecture."),
            ("PR-12", "Firewall & Perimeter Security", "PR.IR", "CIS 13", "NGFW rules management, egress filtering, rule review lifecycle."),
            ("PR-13", "Secure Remote Access (VPN/ZTNA)", "PR.AA", "CIS 12", "VPN hardening, ZTNA migration, split-tunnel policy, posture checks."),
            ("PR-14", "Wireless Network Security", "PR.IR", "CIS 12", "WPA3-Enterprise, rogue AP detection, guest network isolation."),
            ("PR-15", "Email Security", "PR.PS", "CIS 9", "SPF/DKIM/DMARC, anti-phishing gateway, attachment sandboxing."),
            ("PR-16", "Web Security Gateway", "PR.PS", "CIS 9", "URL filtering, SSL inspection, DNS security, browser isolation."),
            ("PR-17", "Data Loss Prevention (DLP)", "PR.DS", "CIS 3", "Endpoint/network/cloud DLP, content inspection, exfiltration prevention."),
            ("PR-18", "Encryption at Rest", "PR.DS", "CIS 3", "Full-disk encryption, database TDE, file-level encryption, key rotation."),
            ("PR-19", "Encryption in Transit", "PR.DS", "CIS 3", "TLS 1.2+ enforcement, certificate management, mTLS for services."),
            ("PR-20", "Key & Certificate Management", "PR.DS", "CIS 3", "HSM/KMS deployment, PKI lifecycle, certificate expiry monitoring."),
            ("PR-21", "Database Security", "PR.DS", "CIS 3", "DB activity monitoring, least-privilege DB accounts, query auditing."),
            ("PR-22", "Cloud Security Posture (CSPM)", "PR.PS", "CIS 4", "Misconfiguration detection, IaC scanning, compliance guardrails."),
            ("PR-23", "Cloud Workload Protection (CWPP)", "PR.PS", "CIS 10", "Container/VM runtime protection, serverless security, drift detection."),
            ("PR-24", "Container & Kubernetes Security", "PR.PS", "CIS 4,16", "Image scanning, admission control, pod security standards, RBAC."),
            ("PR-25", "Secure SDLC / DevSecOps", "PR.PS", "CIS 16", "Shift-left security, SAST/DAST/SCA pipeline gates, security requirements."),
            ("PR-26", "API Security", "PR.PS", "CIS 16", "API gateway controls, OAuth scopes, rate limiting, schema validation."),
            ("PR-27", "Secrets Management", "PR.DS", "CIS 16", "Vault adoption, secret rotation, hardcoded secret scanning."),
            ("PR-28", "Security Awareness Training", "PR.AT", "CIS 14", "Role-based training, phishing simulations, measured behavior change."),
            ("PR-29", "Developer Security Training", "PR.AT", "CIS 14,16", "Secure coding training, OWASP Top-10, security champions."),
            ("PR-30", "Physical Security", "PR.IR", "CIS 12", "Datacenter access controls, badge systems, CCTV, environmental controls."),
            ("PR-31", "Media Protection & Sanitization", "PR.DS", "CIS 3", "Secure disposal, degaussing, data destruction certificates."),
            ("PR-32", "Backup Security & Immutability", "PR.DS", "CIS 11", "Immutable backups, 3-2-1 rule, air-gapped copies, backup encryption."),
            ("PR-33", "Browser Security", "PR.PS", "CIS 9", "Browser hardening, extension control, enterprise browser management."),
            ("PR-34", "OT/ICS Security Controls", "PR.IR", "CIS 12", "Purdue model zoning, unidirectional gateways, OT protocol security."),
            ("PR-35", "AI/ML Security", "PR.PS", "CIS 16", "Model security, prompt injection defense, AI supply chain, LLM guardrails."),
        ],
    },
    "DETECT (DE)": {
        "color": "#E65100",
        "description": "Identify the occurrence of cybersecurity events through monitoring and detection.",
        "domains": [
            ("DE-01", "Security Operations Center (SOC)", "DE.CM", "CIS 8,13", "24×7 monitoring capability, tiered analyst model, SOC metrics."),
            ("DE-02", "SIEM Platform Management", "DE.CM", "CIS 8", "Log aggregation, correlation rules, use case development, tuning."),
            ("DE-03", "Log Management & Retention", "DE.CM", "CIS 8", "Centralized logging, retention policy, log integrity, time sync."),
            ("DE-04", "Endpoint Detection & Response", "DE.CM", "CIS 10,13", "EDR telemetry, behavioral detection, automated containment."),
            ("DE-05", "Network Detection & Response (NDR)", "DE.CM", "CIS 13", "Traffic analysis, lateral movement detection, encrypted traffic analytics."),
            ("DE-06", "Extended Detection & Response (XDR)", "DE.CM", "CIS 13", "Cross-layer correlation: endpoint + network + cloud + identity."),
            ("DE-07", "User & Entity Behavior Analytics", "DE.AE", "CIS 8", "Baseline behavior modeling, anomaly detection, insider threat detection."),
            ("DE-08", "Threat Hunting Program", "DE.AE", "CIS 13", "Hypothesis-driven hunting, MITRE ATT&CK coverage, hunt playbooks."),
            ("DE-09", "Intrusion Detection/Prevention", "DE.CM", "CIS 13", "IDS/IPS signatures, custom rules, inline blocking, alert tuning."),
            ("DE-10", "File Integrity Monitoring", "DE.CM", "CIS 3", "Critical file change detection, registry monitoring, baseline compare."),
            ("DE-11", "Cloud Security Monitoring", "DE.CM", "CIS 8,13", "CloudTrail/Activity Log analysis, cloud-native detections, CIEM alerts."),
            ("DE-12", "Container Runtime Detection", "DE.CM", "CIS 13", "Container behavioral monitoring, syscall analysis, escape detection."),
            ("DE-13", "Deception Technology", "DE.AE", "CIS 13", "Honeypots, honey tokens, decoy credentials, canary files."),
            ("DE-14", "Dark Web Monitoring", "DE.AE", "CIS 17", "Credential leak monitoring, brand abuse, data breach chatter."),
            ("DE-15", "Alert Triage & Case Management", "DE.AE", "CIS 8", "Alert workflow, false positive management, MTTD/MTTA tracking."),
            ("DE-16", "Detection Engineering", "DE.AE", "CIS 8", "Detection-as-code, Sigma rules, ATT&CK mapping, detection coverage."),
            ("DE-17", "Vulnerability Exploitation Detection", "DE.CM", "CIS 7,13", "Exploit attempt detection, virtual patching validation, IOA monitoring."),
            ("DE-18", "OT/ICS Anomaly Detection", "DE.CM", "CIS 13", "Industrial protocol monitoring, process anomaly detection, OT baselining."),
        ],
    },
    "RESPOND (RS)": {
        "color": "#C62828",
        "description": "Take action regarding detected cybersecurity incidents to contain their impact.",
        "domains": [
            ("RS-01", "Incident Response Plan & Playbooks", "RS.MA", "CIS 17", "IR plan, scenario playbooks (ransomware, BEC, data breach), runbooks."),
            ("RS-02", "Incident Response Team (CSIRT)", "RS.MA", "CIS 17", "Team structure, on-call rotation, retainer agreements, skills matrix."),
            ("RS-03", "Incident Classification & Escalation", "RS.MA", "CIS 17", "Severity matrix, escalation paths, declaration criteria, SLAs."),
            ("RS-04", "Digital Forensics Capability", "RS.AN", "CIS 17", "Evidence acquisition, chain of custody, memory/disk forensics, tooling."),
            ("RS-05", "Malware Analysis", "RS.AN", "CIS 17", "Sandbox analysis, reverse engineering, IOC extraction, YARA rules."),
            ("RS-06", "Containment & Eradication", "RS.MI", "CIS 17", "Isolation procedures, credential reset at scale, persistence removal."),
            ("RS-07", "SOAR & Response Automation", "RS.MI", "CIS 17", "Automated playbooks, orchestration workflows, enrichment automation."),
            ("RS-08", "Crisis Communication Plan", "RS.CO", "CIS 17", "Internal comms, media handling, customer notification templates."),
            ("RS-09", "Breach Notification Compliance", "RS.CO", "CIS 17", "72-hour GDPR readiness, state breach laws, regulator notifications."),
            ("RS-10", "Ransomware Response Readiness", "RS.MI", "CIS 11,17", "Ransomware playbook, negotiation policy, decryption capability, OFAC."),
            ("RS-11", "Tabletop Exercises & Simulations", "RS.MA", "CIS 17", "Executive tabletops, technical drills, full simulation exercises."),
            ("RS-12", "Post-Incident Review", "RS.AN", "CIS 17", "Lessons learned, root cause analysis, corrective action tracking."),
            ("RS-13", "Threat Actor Engagement Policy", "RS.CO", "CIS 17", "Law enforcement liaison, attribution policy, ransom payment stance."),
            ("RS-14", "Business Email Compromise Response", "RS.MI", "CIS 9,17", "BEC playbook, wire fraud recovery, mailbox forensics."),
        ],
    },
    "RECOVER (RC)": {
        "color": "#00695C",
        "description": "Restore capabilities and services impaired by cybersecurity incidents.",
        "domains": [
            ("RC-01", "Business Continuity Planning", "RC.RP", "CIS 11", "BCP documentation, alternate sites, workaround procedures."),
            ("RC-02", "Disaster Recovery Program", "RC.RP", "CIS 11", "DR plans, failover procedures, DR site management, RTO/RPO validation."),
            ("RC-03", "Backup & Restoration Testing", "RC.RP", "CIS 11", "Regular restore testing, recovery time measurement, backup validation."),
            ("RC-04", "Cyber Recovery Vault", "RC.RP", "CIS 11", "Isolated recovery environment, clean room, golden images."),
            ("RC-05", "Recovery Prioritization", "RC.RP", "CIS 11", "Tiered recovery sequencing, dependency-aware restoration order."),
            ("RC-06", "Post-Incident System Hardening", "RC.IM", "CIS 4", "Rebuild standards, compromise-informed hardening, validation scans."),
            ("RC-07", "Recovery Communication", "RC.CO", "CIS 17", "Stakeholder updates, service restoration announcements, trust rebuilding."),
            ("RC-08", "Reputation & Trust Recovery", "RC.CO", "CIS 17", "Customer confidence programs, transparency reports, PR strategy."),
            ("RC-09", "Recovery Metrics & Improvement", "RC.IM", "CIS 17", "MTTR tracking, recovery exercise results, plan update cadence."),
        ],
    },
}
MATURITY_LEVELS = {
    0: ("Not Performed", "#C62828", "No capability exists; control is absent."),
    1: ("Initial / Ad-hoc", "#E65100", "Informal, reactive, undocumented practices."),
    2: ("Developing", "#F9A825", "Documented but inconsistently applied."),
    3: ("Defined", "#7CB342", "Standardized, consistently implemented across the org."),
    4: ("Managed", "#2E7D32", "Measured, monitored with metrics, actively managed."),
    5: ("Optimized", "#1565C0", "Continuously improved, automated, industry-leading."),
}
FUNCTION_LIST = list(FRAMEWORK.keys())
TOTAL_DOMAINS = sum(len(f["domains"]) for f in FRAMEWORK.values())
# ════════════════════════════════════════════════════════════════════════════
# GUIDANCE TEXT — explains each NIST CSF 2.0 function for the person scoring it
# ════════════════════════════════════════════════════════════════════════════
FUNCTION_GUIDANCE = {
    "GOVERN (GV)": (
        "**What it covers:** GOVERN is the foundation function — it sets the cybersecurity risk management "
        "strategy, expectations, and policy that everything else operates under. It answers: *who is accountable, "
        "what is our risk appetite, and how do we oversee third parties and the supply chain?*\n\n"
        "**How to score it:** Look for a documented and board-approved security strategy, a named risk owner (e.g. "
        "CISO) with clear reporting lines, current and enforced policies, a formal risk register, active third-party/"
        "supply-chain due diligence, and regular metrics reporting to leadership. Score low (0-1) if governance is "
        "informal or undocumented; score high (4-5) if it is measured, board-visible, and continuously improved."
    ),
    "IDENTIFY (ID)": (
        "**What it covers:** IDENTIFY builds the organization's understanding of its assets, data, systems, risks, "
        "and vulnerabilities — you cannot protect what you don't know you have.\n\n"
        "**How to score it:** Look for a current, automated hardware/software/cloud/identity asset inventory, data "
        "classification, an active vulnerability management and penetration testing program, threat intelligence "
        "consumption, and business impact analysis (RTO/RPO). Score low if inventories are manual or stale; score "
        "high if discovery is continuous/automated and feeds risk-based prioritization."
    ),
    "PROTECT (PR)": (
        "**What it covers:** PROTECT is the largest function (35 domains) — the preventive safeguards that reduce "
        "the likelihood or impact of a cybersecurity event: identity & access management, endpoint and network "
        "security, data protection/encryption, secure development, and security training.\n\n"
        "**How to score it:** Look for enforced MFA and least-privilege access, encryption at rest/in transit, "
        "endpoint protection (EDR) on all assets, network segmentation, secure SDLC gates, and measurable security "
        "awareness training. Score low if controls are partial or optional; score high if controls are enforced "
        "org-wide, automated, and regularly tested."
    ),
    "DETECT (DE)": (
        "**What it covers:** DETECT is the organization's ability to identify that a cybersecurity event is "
        "happening — monitoring, logging, and analysis across endpoint, network, cloud, and identity layers.\n\n"
        "**How to score it:** Look for 24×7 monitoring (SOC/SIEM), centralized log retention, EDR/NDR/XDR "
        "telemetry, tuned detection content mapped to MITRE ATT&CK, and tracked mean-time-to-detect (MTTD). Score "
        "low if detection is manual/reactive; score high if there is continuous automated monitoring with proactive "
        "threat hunting."
    ),
    "RESPOND (RS)": (
        "**What it covers:** RESPOND is the organization's capability to act once an incident is detected — "
        "containment, communication, forensics, and coordinated execution of the incident response plan.\n\n"
        "**How to score it:** Look for a documented and tested IR plan with scenario playbooks, a staffed CSIRT, "
        "forensics capability, crisis communication templates, breach notification procedures, and SOAR automation. "
        "Score low if response is ad-hoc with no rehearsed plan; score high if playbooks are automated, tested via "
        "tabletop exercises, and continuously improved from post-incident reviews."
    ),
    "RECOVER (RC)": (
        "**What it covers:** RECOVER is the organization's ability to restore systems, data, and services after an "
        "incident and rebuild stakeholder trust — business continuity, disaster recovery, and clean-room recovery.\n\n"
        "**How to score it:** Look for documented BCP/DR plans with validated RTO/RPO, regularly tested backup "
        "restoration, immutable/air-gapped recovery copies, a defined recovery sequencing/prioritization order, and "
        "post-incident hardening. Score low if recovery has never been tested; score high if recovery is rehearsed, "
        "measured (MTTR), and continuously improved."
    ),
}
# ════════════════════════════════════════════════════════════════════════════
# SESSION STATE
# ════════════════════════════════════════════════════════════════════════════
def _init_state():
    if "scores" not in st.session_state:
        st.session_state.scores = {}          # domain_id -> maturity 0-5
        for f in FRAMEWORK.values():
            for d in f["domains"]:
                st.session_state.scores[d[0]] = 0
    if "targets" not in st.session_state:
        st.session_state.targets = {}
        for f in FRAMEWORK.values():
            for d in f["domains"]:
                st.session_state.targets[d[0]] = 3
    if "org_name" not in st.session_state:
        st.session_state.org_name = "My Organization"
    if "assessor" not in st.session_state:
        st.session_state.assessor = ""
    if "notes" not in st.session_state:
        st.session_state.notes = {}
_init_state()
# ════════════════════════════════════════════════════════════════════════════
# HELPERS
# ════════════════════════════════════════════════════════════════════════════
def get_assessment_df() -> pd.DataFrame:
    rows = []
    for fname, f in FRAMEWORK.items():
        for d in f["domains"]:
            did, name, nist, cis, desc = d
            score = st.session_state.scores.get(did, 0)
            target = st.session_state.targets.get(did, 3)
            rows.append({
                "Domain ID": did, "Domain": name, "Function": fname,
                "NIST CSF 2.0": nist, "CIS Controls": cis,
                "Current Maturity": score,
                "Maturity Label": MATURITY_LEVELS[score][0],
                "Target Maturity": target,
                "Gap": max(0, target - score),
                "Description": desc,
                "Notes": st.session_state.notes.get(did, ""),
            })
    return pd.DataFrame(rows)
def gap_priority(gap: int) -> str:
    if gap >= 3: return "Critical"
    if gap == 2: return "High"
    if gap == 1: return "Medium"
    return "None"
def function_scores(df: pd.DataFrame) -> pd.DataFrame:
    return df.groupby("Function").agg(
        Avg_Current=("Current Maturity", "mean"),
        Avg_Target=("Target Maturity", "mean"),
        Domains=("Domain ID", "count"),
        Total_Gap=("Gap", "sum"),
    ).reset_index()
# ════════════════════════════════════════════════════════════════════════════
# EXPORT FUNCTIONS
# ════════════════════════════════════════════════════════════════════════════
def export_json(df: pd.DataFrame):
    payload = {
        "framework": "CyberGuard Pro — Cybersecurity Domain Assessment Framework",
        "aligned_standards": ["NIST CSF 2.0", "CIS Controls v8"],
        "developed_by": "Randy Singh — Kalsnet (KNet) Consulting",
        "organization": st.session_state.org_name,
        "assessor": st.session_state.assessor,
        "assessment_date": datetime.now().isoformat(),
        "total_domains": TOTAL_DOMAINS,
        "overall_maturity": round(df["Current Maturity"].mean(), 2),
        "maturity_scale": {str(k): v[0] for k, v in MATURITY_LEVELS.items()},
        "function_summary": function_scores(df).to_dict(orient="records"),
        "domains": df.to_dict(orient="records"),
    }
    return json.dumps(payload, indent=2).encode(), "cyberguard_assessment.json", "application/json"
def export_text(df: pd.DataFrame):
    ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    sep = "=" * 90
    lines = [
        sep,
        "  CYBERGUARD PRO — CYBERSECURITY DOMAIN ASSESSMENT REPORT",
        "  Aligned with NIST CSF 2.0 & CIS Controls v8",
        "  Developed by Randy Singh | Kalsnet (KNet) Consulting",
        sep,
        f"  Organization : {st.session_state.org_name}",
        f"  Assessor     : {st.session_state.assessor or 'N/A'}",
        f"  Date         : {ts}",
        f"  Domains      : {TOTAL_DOMAINS}",
        f"  Overall Maturity: {df['Current Maturity'].mean():.2f} / 5.00",
        sep, "",
        "FUNCTION SUMMARY", "-" * 60,
    ]
    for _, r in function_scores(df).iterrows():
        lines.append(f"  {r['Function']:<18} Current: {r['Avg_Current']:.2f}  Target: {r['Avg_Target']:.2f}  Gap Total: {int(r['Total_Gap'])}")
    lines += ["", "DOMAIN DETAIL", "-" * 60]
    for fname in FUNCTION_LIST:
        lines.append(f"\n### {fname} ###")
        fdf = df[df["Function"] == fname]
        for _, r in fdf.iterrows():
            lines.append(f"  [{r['Domain ID']}] {r['Domain']}")
            lines.append(f"      NIST: {r['NIST CSF 2.0']} | CIS: {r['CIS Controls']} | Current: {r['Current Maturity']} ({r['Maturity Label']}) | Target: {r['Target Maturity']} | Gap: {r['Gap']} ({gap_priority(r['Gap'])})")
            if r["Notes"]:
                lines.append(f"      Notes: {r['Notes']}")
    lines += ["", sep, "  END OF REPORT — © Kalsnet (KNet) Consulting", sep]
    return "\n".join(lines).encode(), "cyberguard_assessment.txt", "text/plain"
def export_pdf(df: pd.DataFrame):
    if not REPORTLAB_OK:
        return None, None, None
    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter,
                            leftMargin=.55*inch, rightMargin=.55*inch,
                            topMargin=.65*inch, bottomMargin=.65*inch)
    styles = getSampleStyleSheet()
    tS  = ParagraphStyle("T",  parent=styles["Title"],  textColor=rl_colors.HexColor("#002B7F"), fontSize=17, spaceAfter=2)
    sS  = ParagraphStyle("S",  parent=styles["Normal"], textColor=rl_colors.HexColor("#1565C0"), fontSize=10, fontName="Helvetica-Bold", spaceAfter=8)
    h2S = ParagraphStyle("H2", parent=styles["Heading2"], textColor=rl_colors.HexColor("#002B7F"), fontSize=13)
    nS  = styles["Normal"]
    story = [
        Paragraph("🛡️ CyberGuard Pro — Cybersecurity Domain Assessment Report", tS),
        Paragraph("Aligned with NIST CSF 2.0 &amp; CIS Controls v8  |  Developed by Randy Singh — Kalsnet (KNet) Consulting", sS),
        HRFlowable(width="100%", thickness=1.2, color=rl_colors.HexColor("#1565C0")),
        Spacer(1, 10),
        Paragraph(f"<b>Organization:</b> {st.session_state.org_name} &nbsp;&nbsp; <b>Assessor:</b> {st.session_state.assessor or 'N/A'} &nbsp;&nbsp; <b>Date:</b> {datetime.now():%Y-%m-%d %H:%M}", nS),
        Paragraph(f"<b>Domains Assessed:</b> {TOTAL_DOMAINS} &nbsp;&nbsp; <b>Overall Maturity:</b> {df['Current Maturity'].mean():.2f} / 5.00 &nbsp;&nbsp; <b>Total Gap Points:</b> {int(df['Gap'].sum())}", nS),
        Spacer(1, 14),
        Paragraph("Executive Summary — Function Maturity", h2S),
    ]
    # Function summary table
    fs = function_scores(df)
    fdata = [["Function", "Domains", "Avg Current", "Avg Target", "Total Gap"]]
    for _, r in fs.iterrows():
        fdata.append([r["Function"], str(r["Domains"]), f"{r['Avg_Current']:.2f}", f"{r['Avg_Target']:.2f}", str(int(r["Total_Gap"]))])
    ft = Table(fdata, colWidths=[2.2*inch, 0.9*inch, 1.1*inch, 1.05*inch, 1.0*inch], repeatRows=1)
    ft.setStyle(TableStyle([
        ("BACKGROUND", (0,0),(-1,0), rl_colors.HexColor("#002B7F")),
        ("TEXTCOLOR",  (0,0),(-1,0), rl_colors.white),
        ("FONTNAME",   (0,0),(-1,0), "Helvetica-Bold"),
        ("FONTSIZE",   (0,0),(-1,-1), 8.5),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),[rl_colors.white, rl_colors.HexColor("#EEF4FF")]),
        ("GRID", (0,0),(-1,-1), .4, rl_colors.grey),
        ("PADDING", (0,0),(-1,-1), 5),
    ]))
    story += [ft, Spacer(1, 14), Paragraph("Top Priority Gaps (Critical &amp; High)", h2S)]
    gaps = df[df["Gap"] >= 2].sort_values("Gap", ascending=False).head(20)
    if len(gaps):
        gdata = [["ID", "Domain", "Function", "Current", "Target", "Gap"]]
        for _, r in gaps.iterrows():
            gdata.append([r["Domain ID"], r["Domain"][:38], r["Function"].split(" ")[0], str(r["Current Maturity"]), str(r["Target Maturity"]), str(r["Gap"])])
        gt = Table(gdata, colWidths=[0.6*inch, 2.9*inch, 1.0*inch, 0.7*inch, 0.65*inch, 0.5*inch], repeatRows=1)
        gt.setStyle(TableStyle([
            ("BACKGROUND", (0,0),(-1,0), rl_colors.HexColor("#C62828")),
            ("TEXTCOLOR",  (0,0),(-1,0), rl_colors.white),
            ("FONTNAME",   (0,0),(-1,0), "Helvetica-Bold"),
            ("FONTSIZE",   (0,0),(-1,-1), 8),
            ("ROWBACKGROUNDS",(0,1),(-1,-1),[rl_colors.white, rl_colors.HexColor("#FFF5F5")]),
            ("GRID", (0,0),(-1,-1), .4, rl_colors.grey),
            ("PADDING", (0,0),(-1,-1), 4),
        ]))
        story.append(gt)
    else:
        story.append(Paragraph("No critical or high gaps identified.", nS))
    story += [PageBreak(), Paragraph("Full Domain Assessment Detail", h2S), Spacer(1, 6)]
    ddata = [["ID", "Domain", "NIST", "CIS", "Cur", "Tgt", "Gap"]]
    for _, r in df.iterrows():
        ddata.append([r["Domain ID"], r["Domain"][:42], r["NIST CSF 2.0"], r["CIS Controls"][:8],
                      str(r["Current Maturity"]), str(r["Target Maturity"]), str(r["Gap"])])
    dt = Table(ddata, colWidths=[0.55*inch, 3.0*inch, 0.7*inch, 0.75*inch, 0.45*inch, 0.45*inch, 0.45*inch], repeatRows=1)
    dt.setStyle(TableStyle([
        ("BACKGROUND", (0,0),(-1,0), rl_colors.HexColor("#002B7F")),
        ("TEXTCOLOR",  (0,0),(-1,0), rl_colors.white),
        ("FONTNAME",   (0,0),(-1,0), "Helvetica-Bold"),
        ("FONTSIZE",   (0,0),(-1,-1), 7),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),[rl_colors.white, rl_colors.HexColor("#EEF4FF")]),
        ("GRID", (0,0),(-1,-1), .3, rl_colors.grey),
        ("PADDING", (0,0),(-1,-1), 3),
    ]))
    story.append(dt)
    story += [Spacer(1, 12),
              Paragraph("<i>© 2025 Kalsnet (KNet) Consulting. This assessment is confidential and intended solely for the named organization.</i>", nS)]
    doc.build(story)
    buf.seek(0)
    return buf.read(), "cyberguard_assessment.pdf", "application/pdf"
def export_word(df: pd.DataFrame):
    if not DOCX_OK:
        return None, None, None
    doc = Document()
    h = doc.add_heading("CyberGuard Pro — Cybersecurity Domain Assessment Report", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs: run.font.color.rgb = RGBColor(0, 43, 127)
    s = doc.add_paragraph("Aligned with NIST CSF 2.0 & CIS Controls v8  |  Developed by Randy Singh — Kalsnet (KNet) Consulting")
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in s.runs:
        run.font.bold = True; run.font.size = Pt(10.5); run.font.color.rgb = RGBColor(21,101,192)
    doc.add_paragraph(f"Organization: {st.session_state.org_name}    Assessor: {st.session_state.assessor or 'N/A'}    Date: {datetime.now():%Y-%m-%d %H:%M}")
    doc.add_paragraph(f"Domains Assessed: {TOTAL_DOMAINS}    Overall Maturity: {df['Current Maturity'].mean():.2f} / 5.00    Total Gap Points: {int(df['Gap'].sum())}")
    doc.add_heading("Executive Summary — Function Maturity", 1)
    fs = function_scores(df)
    t = doc.add_table(rows=1, cols=5); t.style = "Table Grid"
    for i, htxt in enumerate(["Function","Domains","Avg Current","Avg Target","Total Gap"]):
        c = t.rows[0].cells[i]; c.text = htxt
        c.paragraphs[0].runs[0].font.bold = True
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0,43,127)
    for _, r in fs.iterrows():
        cells = t.add_row().cells
        cells[0].text = r["Function"]; cells[1].text = str(r["Domains"])
        cells[2].text = f"{r['Avg_Current']:.2f}"; cells[3].text = f"{r['Avg_Target']:.2f}"
        cells[4].text = str(int(r["Total_Gap"]))
    doc.add_heading("Priority Gap Analysis", 1)
    gaps = df[df["Gap"] >= 2].sort_values("Gap", ascending=False)
    if len(gaps):
        for _, r in gaps.head(25).iterrows():
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"[{r['Domain ID']}] {r['Domain']} — ")
            run.font.bold = True
            p.add_run(f"Current {r['Current Maturity']} → Target {r['Target Maturity']} (Gap {r['Gap']}, {gap_priority(r['Gap'])} priority). {r['Description']}")
    else:
        doc.add_paragraph("No critical or high gaps identified.")
    doc.add_heading("Full Domain Assessment", 1)
    for fname in FUNCTION_LIST:
        doc.add_heading(fname, 2)
        fdf = df[df["Function"] == fname]
        t2 = doc.add_table(rows=1, cols=6); t2.style = "Table Grid"
        for i, htxt in enumerate(["ID","Domain","NIST CSF","CIS","Current","Gap"]):
            c = t2.rows[0].cells[i]; c.text = htxt
            c.paragraphs[0].runs[0].font.bold = True
        for _, r in fdf.iterrows():
            cells = t2.add_row().cells
            cells[0].text = r["Domain ID"]; cells[1].text = r["Domain"]
            cells[2].text = r["NIST CSF 2.0"]; cells[3].text = r["CIS Controls"]
            cells[4].text = f"{r['Current Maturity']} ({r['Maturity Label']})"
            cells[5].text = str(r["Gap"])
    doc.add_paragraph()
    fin = doc.add_paragraph("© 2025 Kalsnet (KNet) Consulting. Confidential.")
    fin.alignment = WD_ALIGN_PARAGRAPH.CENTER
    buf = BytesIO(); doc.save(buf); buf.seek(0)
    return buf.read(), "cyberguard_assessment.docx", \
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
# ════════════════════════════════════════════════════════════════════════════
# HEADER
# ════════════════════════════════════════════════════════════════════════════
st.markdown("""
<div class="header-wrap">
  <p class="app-title">🛡️ CyberGuard Pro — Enterprise Cybersecurity Domain Assessment Framework</p>
  <p class="app-subtitle">Developed By Randy Singh from Kalsnet (KNet) Consulting  ·  NIST CSF 2.0  ·  CIS Controls v8  ·  108 Security Domains</p>
</div>
""", unsafe_allow_html=True)
with st.expander("ℹ️ About This Application — How It Works", expanded=True):
    st.markdown(f"""
CyberGuard Pro is a self-assessment tool that helps an organization measure its cybersecurity maturity across
**{TOTAL_DOMAINS} security domains**, organized into the **6 NIST Cybersecurity Framework (CSF) 2.0 functions** —
GOVERN, IDENTIFY, PROTECT, DETECT, RESPOND, RECOVER — and cross-mapped to **CIS Controls v8**. It works as follows:

1. **📋 Domain Assessment** — for every domain, you rate your **Current Maturity** (0–5) and set a **Target
   Maturity**. This is the only data-entry step; every other page is calculated automatically from these scores.
2. **📊 Executive Dashboard** — summarizes your overall posture with KPIs and charts, so leadership can see the
   big picture at a glance.
3. **📈 Gap Analysis** — automatically calculates the gap (Target − Current) for every domain and prioritizes it as
   Critical, High, or Medium so you know what to fix first.
4. **🗺️ Framework Explorer** — lets you browse and search all {TOTAL_DOMAINS} domains with their NIST/CIS mappings
   and descriptions.
5. **🎯 Roadmap Builder** — auto-generates a phased remediation plan (0–3, 3–9, 9–18 months) based on gap severity.
6. **📤 Export Reports** — produces PDF, Word, Text, or JSON reports of the full assessment for stakeholders or
   GRC tooling.

Use the **🎲 Demo Data** button in the sidebar to populate sample scores instantly, or **🔄 Reset All** to start over.
""")
# ════════════════════════════════════════════════════════════════════════════
# SIDEBAR
# ════════════════════════════════════════════════════════════════════════════
with st.sidebar:
    st.markdown("## 🛡️ CyberGuard Pro")
    st.divider()
    st.session_state.org_name = st.text_input("🏢 Organization Name", st.session_state.org_name)
    st.session_state.assessor = st.text_input("👤 Assessor Name", st.session_state.assessor)
    st.divider()
    page = st.radio("Navigation", [
        "📊 Executive Dashboard",
        "📋 Domain Assessment",
        "📈 Gap Analysis",
        "🗺️ Framework Explorer",
        "🎯 Roadmap Builder",
        "📤 Export Reports",
    ], label_visibility="collapsed")
    st.divider()
    st.markdown("**⚡ Quick Actions**")
    c1, c2 = st.columns(2)
    with c1:
        if st.button("🎲 Demo Data", use_container_width=True):
            rng = np.random.default_rng(42)
            for f in FRAMEWORK.values():
                for d in f["domains"]:
                    st.session_state.scores[d[0]] = int(rng.choice([0,1,1,2,2,2,3,3,3,4,4,5], 1)[0])
            st.rerun()
    with c2:
        if st.button("🔄 Reset All", use_container_width=True):
            for f in FRAMEWORK.values():
                for d in f["domains"]:
                    st.session_state.scores[d[0]] = 0
            st.rerun()
    st.divider()
    st.markdown("**📖 Maturity Scale**")
    for lvl, (name, color, _) in MATURITY_LEVELS.items():
        st.markdown(f'<span class="maturity-{lvl}">{lvl} — {name}</span>', unsafe_allow_html=True)
        st.markdown("<div style='height:4px'></div>", unsafe_allow_html=True)
df = get_assessment_df()
# ════════════════════════════════════════════════════════════════════════════
# PAGE: EXECUTIVE DASHBOARD
# ════════════════════════════════════════════════════════════════════════════
if page == "📊 Executive Dashboard":
    st.subheader("📊 Executive Dashboard")
    st.caption("A leadership-level summary of cybersecurity maturity across all domains, calculated live from the scores entered on the Domain Assessment page.")
    overall = df["Current Maturity"].mean()
    total_gap = int(df["Gap"].sum())
    critical = len(df[df["Gap"] >= 3])
    assessed = len(df[df["Current Maturity"] > 0])
    c1, c2, c3, c4, c5 = st.columns(5)
    kpis = [
        (c1, "🎯 Overall Maturity", f"{overall:.2f}/5", "#002B7F"),
        (c2, "📚 Total Domains", TOTAL_DOMAINS, "#1565C0"),
        (c3, "✅ Assessed", assessed, "#2E7D32"),
        (c4, "⚠️ Gap Points", total_gap, "#E65100"),
        (c5, "🚨 Critical Gaps", critical, "#C62828"),
    ]
    for col, lbl, val, color in kpis:
        col.markdown(
            f'<div class="kpi-card" style="background:linear-gradient(135deg,{color},{color}bb)">'
            f'<div class="kpi-lbl">{lbl}</div><div class="kpi-val">{val}</div></div>',
            unsafe_allow_html=True)
    with st.expander("ℹ️ How are these KPIs calculated?"):
        st.markdown("**🎯 Overall Maturity** — the average Current Maturity score across every domain in the framework.")
        st.markdown('<div class="formula-box">Overall Maturity = SUM(Current Maturity, all domains) ÷ Total Domains</div>', unsafe_allow_html=True)
        st.markdown(f"Right now: SUM of all Current Maturity scores ÷ {TOTAL_DOMAINS} domains = **{overall:.2f} / 5.00**")
        st.markdown("---")
        st.markdown("**📚 Total Domains** — the fixed number of security domains in the framework (does not change with scoring).")
        st.markdown('<div class="formula-box">Total Domains = COUNT(all domains across all 6 NIST CSF functions)</div>', unsafe_allow_html=True)
        st.markdown("---")
        st.markdown("**✅ Assessed** — the number of domains that have been given a Current Maturity score greater than 0 (i.e. touched by the assessor). A domain left at 0 is treated as 'not yet assessed / not performed'.")
        st.markdown('<div class="formula-box">Assessed = COUNT(domains WHERE Current Maturity > 0)</div>', unsafe_allow_html=True)
        st.markdown("---")
        st.markdown("**⚠️ Gap Points** — the sum of every domain's individual gap, where a gap is how far the Current Maturity is below the Target Maturity (gaps can never be negative — a domain that already exceeds its target contributes 0).")
        st.markdown('<div class="formula-box">Gap (per domain) = MAX(0, Target Maturity − Current Maturity)\nGap Points (total) = SUM(Gap, all domains)</div>', unsafe_allow_html=True)
        st.markdown("---")
        st.markdown("**🚨 Critical Gaps** — the count of domains whose individual Gap is 3 or more (i.e. Current Maturity is at least 3 levels below Target). See the Gap Analysis page for the full priority scale (Critical / High / Medium).")
        st.markdown('<div class="formula-box">Critical Gaps = COUNT(domains WHERE Gap >= 3)</div>', unsafe_allow_html=True)
    st.markdown("<br>", unsafe_allow_html=True)
    r1c1, r1c2 = st.columns([1.1, 1])
    with r1c1:
        # Radar chart per NIST function
        fs = function_scores(df)
        categories = [f.split(" ")[0] for f in fs["Function"]]
        fig = go.Figure()
        fig.add_trace(go.Scatterpolar(
            r=fs["Avg_Current"].tolist() + [fs["Avg_Current"].iloc[0]],
            theta=categories + [categories[0]],
            fill="toself", name="Current", line_color="#1565C0"))
        fig.add_trace(go.Scatterpolar(
            r=fs["Avg_Target"].tolist() + [fs["Avg_Target"].iloc[0]],
            theta=categories + [categories[0]],
            fill="toself", name="Target", line_color="#E65100", opacity=0.5))
        fig.update_layout(title="🕸️ NIST CSF 2.0 Function Maturity Radar",
                          polar=dict(radialaxis=dict(range=[0,5])), height=400)
        st.plotly_chart(fig, use_container_width=True)
        st.caption("**How to read this:** each spoke is one of the 6 NIST CSF functions. The blue shape plots the average Current Maturity of all domains in that function; the orange shape plots the average Target. The bigger the gap between the two shapes on a spoke, the further that function is from its goal.")
    with r1c2:
        fig = px.bar(fs, x="Function", y=["Avg_Current","Avg_Target"],
                     barmode="group", title="📊 Current vs Target by Function",
                     color_discrete_map={"Avg_Current":"#1565C0","Avg_Target":"#F9A825"})
        fig.update_layout(height=400, xaxis_tickangle=25, legend_title="")
        st.plotly_chart(fig, use_container_width=True)
        st.caption("**How to read this:** the same Current vs Target averages as the radar chart, shown side-by-side per function for easier numeric comparison. Avg Current/Target = mean of all domain scores within that function.")
    r2c1, r2c2 = st.columns(2)
    with r2c1:
        mat_dist = df["Maturity Label"].value_counts().reindex(
            [MATURITY_LEVELS[i][0] for i in range(6)], fill_value=0)
        fig = px.pie(values=mat_dist.values, names=mat_dist.index,
                     title="🎨 Maturity Distribution Across All Domains",
                     color=mat_dist.index,
                     color_discrete_map={MATURITY_LEVELS[i][0]: MATURITY_LEVELS[i][1] for i in range(6)})
        fig.update_layout(height=360)
        st.plotly_chart(fig, use_container_width=True)
        st.caption(f"**How to read this:** counts how many of the {TOTAL_DOMAINS} domains fall into each of the 6 maturity labels (0 Not Performed → 5 Optimized). A pie skewed toward red/orange slices means most domains are still immature.")
    with r2c2:
        heat = df.pivot_table(index="Function", values="Current Maturity", aggfunc="mean").reset_index()
        fig = px.bar(df.groupby("Function")["Gap"].sum().reset_index(),
                     x="Gap", y="Function", orientation="h",
                     title="🔥 Total Gap Points by Function",
                     color="Gap", color_continuous_scale="Reds")
        fig.update_layout(height=360, showlegend=False)
        st.plotly_chart(fig, use_container_width=True)
        st.caption("**How to read this:** sums the Gap (Target − Current, floored at 0) of every domain within each function. Longer/darker bars mean that function has the most remediation work outstanding — a good place to focus budget and staffing.")
    # Heatmap of all domains
    st.markdown("**🗺️ Full Domain Heatmap (all 108 domains)**")
    st.caption("**How to read this:** every single domain plotted as one dot, grouped by function (row) and colored by Current Maturity (red = 0/Not Performed, blue = 5/Optimized). Hover a dot to see the domain name, ID, and gap. This is the most granular view on the dashboard — use it to spot clusters of weak domains within a function.")
    heat_df = df.copy()
    heat_df["Row"] = heat_df.groupby("Function").cumcount()
    fig = px.scatter(heat_df, x="Row", y="Function", color="Current Maturity",
                     size=[14]*len(heat_df), hover_name="Domain",
                     hover_data={"Domain ID":True,"Current Maturity":True,"Gap":True,"Row":False},
                     color_continuous_scale=["#C62828","#E65100","#F9A825","#7CB342","#2E7D32","#1565C0"],
                     range_color=[0,5])
    fig.update_layout(height=320, xaxis_title="Domain Index", yaxis_title="")
    st.plotly_chart(fig, use_container_width=True)
    st.markdown("**📚 All Domains — Assessed Status (scrollable)**")
    st.caption(f"Every one of the {TOTAL_DOMAINS} domains with its current status. Scroll inside the table to see all rows — this is not limited to the first 10.")
    st.dataframe(
        df[["Domain ID","Domain","Function","NIST CSF 2.0","CIS Controls",
            "Current Maturity","Maturity Label","Target Maturity","Gap"]].reset_index(drop=True),
        use_container_width=True, height=420)
# ════════════════════════════════════════════════════════════════════════════
# PAGE: DOMAIN ASSESSMENT
# ════════════════════════════════════════════════════════════════════════════
elif page == "📋 Domain Assessment":
    st.subheader("📋 Domain Maturity Assessment")
    st.markdown("Rate each domain's **current maturity** (0–5) and set a **target**. Add notes for context.")
    st.info("**Scoring scale reminder:** 0 = Not Performed, 1 = Initial/Ad-hoc, 2 = Developing, 3 = Defined, 4 = Managed, 5 = Optimized. See the sidebar for the full definitions of each level.")
    fn_tabs = st.tabs([f.split(" ")[0] for f in FUNCTION_LIST])
    for tab, fname in zip(fn_tabs, FUNCTION_LIST):
        with tab:
            f = FRAMEWORK[fname]
            st.markdown(f"**{fname}** — {f['description']} ({len(f['domains'])} domains)")
            with st.expander(f"📘 Guidance — how to assess {fname} domains", expanded=False):
                st.markdown(FUNCTION_GUIDANCE[fname])
            st.divider()
            for d in f["domains"]:
                did, name, nist, cis, desc = d
                with st.expander(f"**{did}** — {name}   `NIST: {nist}` `CIS: {cis}`"):
                    st.caption(desc)
                    c1, c2 = st.columns(2)
                    with c1:
                        st.session_state.scores[did] = st.select_slider(
                            "Current Maturity", options=list(range(6)),
                            value=st.session_state.scores[did],
                            format_func=lambda x: f"{x} — {MATURITY_LEVELS[x][0]}",
                            key=f"cur_{did}")
                    with c2:
                        st.session_state.targets[did] = st.select_slider(
                            "Target Maturity", options=list(range(6)),
                            value=st.session_state.targets[did],
                            format_func=lambda x: f"{x} — {MATURITY_LEVELS[x][0]}",
                            key=f"tgt_{did}")
                    st.session_state.notes[did] = st.text_input(
                        "Notes / evidence", st.session_state.notes.get(did, ""), key=f"note_{did}")
# ════════════════════════════════════════════════════════════════════════════
# PAGE: GAP ANALYSIS
# ════════════════════════════════════════════════════════════════════════════
elif page == "📈 Gap Analysis":
    st.subheader("📈 Gap Analysis & Prioritization")
    st.markdown(
        "This page calculates the **Gap** for every domain — how far Current Maturity is below Target Maturity — "
        "and sorts that work into priority tiers so remediation effort goes where it matters most."
    )
    with st.expander("ℹ️ How Gap and Priority are calculated", expanded=True):
        st.markdown('<div class="formula-box">Gap = MAX(0, Target Maturity − Current Maturity)</div>', unsafe_allow_html=True)
        st.markdown(
            "A domain scored Current=1 with Target=4 has a Gap of 3. A domain already at or above its target has a "
            "Gap of 0 and drops out of this page entirely (there is nothing to remediate)."
        )
        st.markdown("**Priority tiers — how they differ:**")
        st.markdown(
            "- **🚨 Critical (Gap ≥ 3):** the domain is at least 3 maturity levels behind its target — typically "
            "close to non-existent capability against a moderate-to-high target. These pose the most material risk "
            "and should be addressed with executive sponsorship within 0–3 months.\n"
            "- **🔶 High (Gap = 2):** two levels behind target — a meaningful capability shortfall that should be "
            "budgeted and resourced in the near term (3–9 months).\n"
            "- **🔸 Medium (Gap = 1):** one level behind target — a smaller, often quick-to-close gap suited to "
            "longer-term continuous improvement (9–18 months).\n\n"
            "In short: the **size of the number** (3, 2, or 1) is simply how many maturity levels separate where "
            "the organization is today from where it wants to be — the larger the gap, the more urgent the fix."
        )
    gaps = df[df["Gap"] > 0].copy()
    gaps["Priority"] = gaps["Gap"].apply(gap_priority)
    c1, c2, c3 = st.columns(3)
    c1.metric("🚨 Critical Gaps (3+)", len(gaps[gaps["Gap"] >= 3]))
    c2.metric("🔶 High Gaps (2)", len(gaps[gaps["Gap"] == 2]))
    c3.metric("🔸 Medium Gaps (1)", len(gaps[gaps["Gap"] == 1]))
    if len(gaps):
        fc1, fc2 = st.columns(2)
        with fc1:
            fn_filter = st.multiselect("Filter by Function", FUNCTION_LIST, default=FUNCTION_LIST)
        with fc2:
            pr_filter = st.multiselect("Filter by Priority", ["Critical","High","Medium"], default=["Critical","High","Medium"])
        fgaps = gaps[gaps["Function"].isin(fn_filter) & gaps["Priority"].isin(pr_filter)] \
                    .sort_values(["Gap","Function"], ascending=[False, True])
        fig = px.bar(fgaps.head(25), x="Gap", y="Domain", orientation="h",
                     color="Priority", title="Top 25 Domain Gaps",
                     color_discrete_map={"Critical":"#C62828","High":"#E65100","Medium":"#F9A825"},
                     hover_data=["Domain ID","Function","NIST CSF 2.0","CIS Controls"])
        fig.update_layout(height=620, yaxis=dict(autorange="reversed"))
        st.plotly_chart(fig, use_container_width=True)
        st.caption("**How to read this:** the 25 domains with the largest gaps, longest bar first. Color shows priority tier (red=Critical, orange=High, yellow=Medium).")
        st.markdown("**📋 Gap Register**")
        st.dataframe(
            fgaps[["Domain ID","Domain","Function","NIST CSF 2.0","CIS Controls",
                   "Current Maturity","Target Maturity","Gap","Priority","Notes"]].reset_index(drop=True),
            use_container_width=True, height=380)
    else:
        st.success("🎉 No gaps — all domains meet or exceed their targets!")
# ════════════════════════════════════════════════════════════════════════════
# PAGE: FRAMEWORK EXPLORER
# ════════════════════════════════════════════════════════════════════════════
elif page == "🗺️ Framework Explorer":
    st.subheader(f"🗺️ Framework Explorer — All {TOTAL_DOMAINS} Domains")
    st.markdown(
        f"Browse or search every domain in the CyberGuard Pro framework. All {TOTAL_DOMAINS} domains are organized "
        "under the 6 NIST CSF 2.0 functions (GOVERN, IDENTIFY, PROTECT, DETECT, RESPOND, RECOVER) and each is "
        "cross-mapped to its NIST CSF 2.0 category and CIS Controls v8 control group, with a short description of "
        "what the domain covers. Use the sunburst chart below to visually navigate by function, or use the search "
        "box and per-function lists to find and review specific domains — colors reflect the Current Maturity score "
        "entered on the Domain Assessment page."
    )
    search = st.text_input("🔍 Search domains", placeholder="e.g. encryption, MFA, cloud, ransomware …")
    sunburst_df = df.copy()
    fig = px.sunburst(sunburst_df, path=["Function","Domain"], values=None,
                      color="Current Maturity",
                      color_continuous_scale=["#C62828","#F9A825","#2E7D32","#1565C0"],
                      range_color=[0,5], title="🌐 Framework Structure (click to zoom)")
    fig.update_layout(height=550)
    st.plotly_chart(fig, use_container_width=True)
    st.caption("**How to read this:** the inner ring is the 6 NIST CSF functions; the outer ring is every domain within that function. Click a slice to zoom in, click the center to zoom back out. Color reflects that domain's Current Maturity score.")
    fdf = df.copy()
    if search:
        mask = fdf.apply(lambda r: r.astype(str).str.contains(search, case=False).any(), axis=1)
        fdf = fdf[mask]
        st.markdown(f"**{len(fdf)} domains match '{search}'**")
    for fname in FUNCTION_LIST:
        sub = fdf[fdf["Function"] == fname]
        if not len(sub): continue
        color = FRAMEWORK[fname]["color"]
        st.markdown(f'<h4 style="color:{color};border-bottom:2px solid {color};padding-bottom:4px">{fname} ({len(sub)} domains)</h4>', unsafe_allow_html=True)
        for _, r in sub.iterrows():
            lvl = r["Current Maturity"]
            st.markdown(
                f'<div class="domain-card gap-{"critical" if r["Gap"]>=3 else "high" if r["Gap"]==2 else "medium" if r["Gap"]==1 else "low"}">'
                f'<b>{r["Domain ID"]} — {r["Domain"]}</b> &nbsp; '
                f'<span class="maturity-{lvl}">{lvl} · {r["Maturity Label"]}</span>'
                f'<br><small style="color:#5A7A9A">NIST CSF: {r["NIST CSF 2.0"]} &nbsp;|&nbsp; CIS: {r["CIS Controls"]} &nbsp;|&nbsp; {r["Description"]}</small></div>',
                unsafe_allow_html=True)
# ════════════════════════════════════════════════════════════════════════════
# PAGE: ROADMAP BUILDER
# ════════════════════════════════════════════════════════════════════════════
elif page == "🎯 Roadmap Builder":
    st.subheader("🎯 Remediation Roadmap Builder")
    st.markdown("Auto-generated remediation phases based on gap severity and NIST function priority.")
    with st.expander("ℹ️ How the roadmap is built", expanded=True):
        st.markdown(
            "Every domain with a Gap greater than 0 (Current Maturity below Target) is automatically bucketed into "
            "one of three phases based on its Gap value, matching the same priority tiers used on the Gap Analysis "
            "page:\n\n"
            "- **🚨 Phase 1 — Immediate (0–3 months):** domains with Gap ≥ 3 (Critical). These represent the "
            "largest shortfalls and are scheduled first, with executive sponsorship.\n"
            "- **🔶 Phase 2 — Near-Term (3–9 months):** domains with Gap = 2 (High). Scheduled to start once Phase "
            "1 is underway, budgeted in the current fiscal cycle.\n"
            "- **🔸 Phase 3 — Strategic (9–18 months):** domains with Gap = 1 (Medium). Longer-horizon continuous "
            "improvement items.\n\n"
            "The **🗓 Remediation Timeline** chart below turns those three phases into a Gantt-style view: each bar "
            "is one domain, its start date is the beginning of its phase window, and its length is a fixed "
            "planning duration per phase (3 months for Phase 1, 6 months for Phase 2, 9 months for Phase 3). Only "
            "the 30 highest-gap domains are plotted to keep the chart readable."
        )
    gaps = df[df["Gap"] > 0].copy()
    gaps["Priority"] = gaps["Gap"].apply(gap_priority)
    if len(gaps) == 0:
        st.success("🎉 No remediation needed — all domains meet their targets!")
    else:
        phase1 = gaps[gaps["Gap"] >= 3]
        phase2 = gaps[gaps["Gap"] == 2]
        phase3 = gaps[gaps["Gap"] == 1]
        phases = [
            ("🚨 Phase 1 — Immediate (0–3 months)", phase1, "#C62828",
             "Critical gaps posing material risk. Address urgently with executive sponsorship."),
            ("🔶 Phase 2 — Near-Term (3–9 months)", phase2, "#E65100",
             "High-priority improvements. Budget and resource in the current fiscal cycle."),
            ("🔸 Phase 3 — Strategic (9–18 months)", phase3, "#F9A825",
             "Continuous improvement items to reach target maturity across the framework."),
        ]
        # Gantt-style timeline
        timeline_rows = []
        for pname, pdf_, color, _ in phases:
            for _, r in pdf_.iterrows():
                start = {"🚨":0, "🔶":3, "🔸":9}[pname[:1]]
                dur   = {"🚨":3, "🔶":6, "🔸":9}[pname[:1]]
                timeline_rows.append({"Domain": f"{r['Domain ID']} {r['Domain'][:30]}",
                                      "Start": start, "Finish": start+dur, "Phase": pname.split(" — ")[0][2:].strip()})
        if timeline_rows:
            tdf = pd.DataFrame(timeline_rows).head(30)
            fig = px.timeline(tdf.assign(
                                Start=pd.Timestamp("2026-01-01") + pd.to_timedelta(tdf["Start"]*30, unit="D"),
                                Finish=pd.Timestamp("2026-01-01") + pd.to_timedelta(tdf["Finish"]*30, unit="D")),
                              x_start="Start", x_end="Finish", y="Domain", color="Phase",
                              title="🗓 Remediation Timeline (top 30 items)",
                              color_discrete_map={"Phase 1":"#C62828","Phase 2":"#E65100","Phase 3":"#F9A825"})
            fig.update_layout(height=600, yaxis=dict(autorange="reversed"))
            st.plotly_chart(fig, use_container_width=True)
            st.caption("**How to read this:** each horizontal bar is one domain's planned remediation window. Bar color shows which phase it belongs to; bar position/length shows the scheduled start and planning duration described above.")
        for pname, pdf_, color, pdesc in phases:
            st.markdown(f'<h4 style="color:{color}">{pname} — {len(pdf_)} domains</h4>', unsafe_allow_html=True)
            st.caption(pdesc)
            if len(pdf_):
                st.markdown(
                    "Table columns: **Current Maturity** (score today), **Target Maturity** (goal score), and "
                    "**Gap** (Target − Current) — the same values that placed this domain into this phase."
                )
                st.dataframe(pdf_[["Domain ID","Domain","Function","Current Maturity","Target Maturity","Gap"]]
                             .sort_values("Gap", ascending=False).reset_index(drop=True),
                             use_container_width=True, height=min(280, 45+len(pdf_)*36))
# ════════════════════════════════════════════════════════════════════════════
# PAGE: EXPORT
# ════════════════════════════════════════════════════════════════════════════
elif page == "📤 Export Reports":
    st.subheader("📤 Export Professional Assessment Reports")
    st.markdown(f"**Organization:** {st.session_state.org_name} &nbsp;·&nbsp; "
                f"**Overall Maturity:** {df['Current Maturity'].mean():.2f}/5.00 &nbsp;·&nbsp; "
                f"**{TOTAL_DOMAINS} domains** &nbsp;·&nbsp; **{int(df['Gap'].sum())} gap points**",
                unsafe_allow_html=True)
    st.divider()
    c1, c2, c3, c4 = st.columns(4)
    with c1:
        st.markdown("#### 📄 PDF Report")
        st.caption("Executive summary, priority gaps, full domain table with KNet branding.")
        if not REPORTLAB_OK:
            st.warning("`pip install reportlab`")
        else:
            data, fname, mime = export_pdf(df)
            if data:
                st.download_button("⬇️ Download PDF", data=data, file_name=fname,
                                   mime=mime, use_container_width=True, type="primary")
    with c2:
        st.markdown("#### 📝 Word Report")
        st.caption("Fully editable .docx with function tables and gap analysis bullets.")
        if not DOCX_OK:
            st.warning("`pip install python-docx`")
        else:
            data, fname, mime = export_word(df)
            if data:
                st.download_button("⬇️ Download Word", data=data, file_name=fname,
                                   mime=mime, use_container_width=True, type="primary")
    with c3:
        st.markdown("#### 📋 Text Report")
        st.caption("Plain-text structured report for tickets, email, or archival.")
        data, fname, mime = export_text(df)
        st.download_button("⬇️ Download Text", data=data, file_name=fname,
                           mime=mime, use_container_width=True, type="primary")
    with c4:
        st.markdown("#### 🔧 JSON Data")
        st.caption("Machine-readable full assessment for GRC tool import or APIs.")
        data, fname, mime = export_json(df)
        st.download_button("⬇️ Download JSON", data=data, file_name=fname,
                           mime=mime, use_container_width=True, type="primary")
    st.divider()
    st.markdown("**Preview — Function Summary**")
    st.dataframe(function_scores(df), use_container_width=True)
    st.markdown(f"**Preview — All {TOTAL_DOMAINS} Domains (scrollable)**")
    st.caption("Scroll inside the table below to view every domain — this preview is no longer limited to the first 10 rows.")
    st.dataframe(df[["Domain ID","Domain","Function","NIST CSF 2.0","CIS Controls",
                     "Current Maturity","Target Maturity","Gap"]].reset_index(drop=True),
                 use_container_width=True, height=420)
# ── Footer ───────────────────────────────────────────────────────────────────
st.markdown("---")
st.markdown(
    '<p style="text-align:center;color:#7A9ABB;font-size:.8rem">'
    '🛡️ CyberGuard Pro · 108 Security Domains · NIST CSF 2.0 · CIS Controls v8<br>'
    '© 2025 Kalsnet (KNet) Consulting — Developed by Randy Singh. All rights reserved.</p>',
    unsafe_allow_html=True)
