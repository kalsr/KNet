# Rational Decision Support System (RDSS)
# Developed by Randy Singh - Kalsnet (KNet) Consulting
# A Streamlit application applying Rational Choice Theory / Multi-Attribute
# Utility Theory (MAUT) and Expected Utility Theory to three decision domains:
# 1. Strategic Decision-Making
# 2. Resource Allocation
# 3. Threat Assessment
# Supports synthetic demo data generation, real data upload (CSV/XLSX),
# interactive visualizations, and export to PDF, Word, TXT, and CSV.

import io
import datetime as dt
import numpy as np
import pandas as pd
import streamlit as st
import matplotlib.pyplot as plt
import matplotlib.cm as cm
from scipy.optimize import linprog
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors as rl_colors
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image as RLImage
)

# --------------------------------------------------------------------------
# GLOBAL CONFIG / BRANDING
# --------------------------------------------------------------------------
BRAND_NAME = "Kalsnet (KNet) Consulting"
DEVELOPER = "Randy Singh"
NAVY = "#0A2F5C"
BLUE = "#0E4C92"
ACCENT_GOLD = "#C9A227"
LIGHT_BG = "#F4F7FB"
DEV_BLUE = "#3FA7FF"

st.set_page_config(
    page_title="Rational Decision Support System | KNet Consulting",
    layout="wide",
    initial_sidebar_state="expanded",
)

CUSTOM_CSS = f"""
<style>
    .stApp {{
        background-color: {LIGHT_BG};
    }}
    .knet-title-bar {{
        background: linear-gradient(90deg, {NAVY} 0%, {BLUE} 100%);
        padding: 22px 30px;
        border-radius: 6px;
        margin-bottom: 18px;
        box-shadow: 0 4px 10px rgba(0,0,0,0.25);
    }}
    .knet-title-bar h1 {{
        color: #FFFFFF;
        font-weight: 800;
        font-size: 32px;
        margin: 0;
        letter-spacing: 0.5px;
    }}
    .knet-title-bar p {{
        color: #D9E4F5;
        margin: 4px 0 0 0;
        font-size: 15px;
        font-weight: 500;
    }}
    .knet-title-bar .knet-dev-credit {{
        color: {DEV_BLUE};
        font-weight: 800;
        font-size: 14px;
        margin: 8px 0 0 0;
        letter-spacing: 0.3px;
    }}
    .knet-subheader {{
        background-color: #FFFFFF;
        border-left: 6px solid {BLUE};
        padding: 10px 16px;
        border-radius: 4px;
        margin: 10px 0 16px 0;
        box-shadow: 0 1px 4px rgba(0,0,0,0.08);
    }}
    .knet-footer {{
        text-align: center;
        color: #6b7280;
        font-size: 12.5px;
        padding: 18px 0 6px 0;
        border-top: 1px solid #d9dee5;
        margin-top: 30px;
    }}
    .knet-card {{
        background-color: #FFFFFF;
        border-radius: 8px;
        padding: 16px 20px;
        box-shadow: 0 1px 6px rgba(0,0,0,0.10);
        margin-bottom: 14px;
    }}
    .stButton>button, .stDownloadButton>button {{
        background-color: {BLUE};
        color: white;
        font-weight: 600;
        border-radius: 5px;
        border: none;
    }}
    .stButton>button:hover, .stDownloadButton>button:hover {{
        background-color: {NAVY};
        color: white;
    }}
    section[data-testid="stSidebar"] {{
        background-color: #101E33;
    }}
    section[data-testid="stSidebar"] * {{
        color: #E7ECF5 !important;
    }}
</style>
"""
st.markdown(CUSTOM_CSS, unsafe_allow_html=True)


def title_bar(subtitle: str):
    st.markdown(
        f"""
        <div class="knet-title-bar">
            <h1>RATIONAL DECISION SUPPORT SYSTEM</h1>
            <p>{subtitle}</p>
            <p class="knet-dev-credit">Developed by Randy Singh from Kalsnet (KNet) Consulting Team</p>
        </div>
        """,
        unsafe_allow_html=True,
    )


def footer():
    st.markdown(
        f"""
        <div class="knet-footer">
            (c) {dt.datetime.now().year} {BRAND_NAME} &nbsp;|&nbsp;
            Developed by <b>{DEVELOPER}</b> &nbsp;|&nbsp;
            Rational Decision Support System v1.0
        </div>
        """,
        unsafe_allow_html=True,
    )


# --------------------------------------------------------------------------
# EXPORT HELPERS
# --------------------------------------------------------------------------
def fig_to_png_buffer(fig) -> io.BytesIO:
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=160, bbox_inches="tight")
    buf.seek(0)
    return buf


def df_to_csv_bytes(df: pd.DataFrame) -> bytes:
    return df.to_csv(index=False).encode("utf-8")


def df_to_txt_bytes(title: str, subtitle: str, notes: str, df: pd.DataFrame) -> bytes:
    lines = []
    lines.append(title)
    lines.append("=" * len(title))
    lines.append(subtitle)
    lines.append(f"Generated: {dt.datetime.now():%Y-%m-%d %H:%M}")
    lines.append(f"Prepared by: {DEVELOPER} - {BRAND_NAME}")
    lines.append("")
    if notes:
        lines.append("SUMMARY / RATIONALE")
        lines.append("-" * 20)
        lines.append(notes)
        lines.append("")
    lines.append("DATA")
    lines.append("-" * 20)
    lines.append(df.to_string(index=False))
    return "\n".join(lines).encode("utf-8")


def build_word_report(title: str, subtitle: str, notes: str, df: pd.DataFrame,
                       chart_buf: io.BytesIO = None) -> io.BytesIO:
    doc = Document()
    heading = doc.add_heading(title, level=1)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x0A, 0x2F, 0x5C)
    p = doc.add_paragraph(subtitle)
    p.runs[0].italic = True
    meta = doc.add_paragraph()
    meta.add_run(f"Generated: {dt.datetime.now():%Y-%m-%d %H:%M}\n").font.size = Pt(10)
    meta.add_run(f"Prepared by: {DEVELOPER} - {BRAND_NAME}").font.size = Pt(10)
    if notes:
        doc.add_heading("Summary & Rationale", level=2)
        doc.add_paragraph(notes)
    doc.add_heading("Supporting Data", level=2)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = f"{val:.3f}" if isinstance(val, float) else str(val)
    if chart_buf is not None:
        doc.add_heading("Visualization", level=2)
        doc.add_picture(chart_buf, width=Inches(6.0))
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run(f"{BRAND_NAME} - Confidential Analysis")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out


def build_pdf_report(title: str, subtitle: str, notes: str, df: pd.DataFrame,
                      chart_buf: io.BytesIO = None) -> io.BytesIO:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter,
                             topMargin=0.6 * inch, bottomMargin=0.6 * inch)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "KNetTitle", parent=styles["Heading1"],
        textColor=rl_colors.HexColor(NAVY), fontSize=18, spaceAfter=6
    )
    sub_style = ParagraphStyle(
        "KNetSub", parent=styles["Normal"], textColor=rl_colors.HexColor("#444444"),
        fontSize=10, spaceAfter=2
    )
    section_style = ParagraphStyle(
        "KNetSection", parent=styles["Heading2"],
        textColor=rl_colors.HexColor(BLUE), fontSize=13, spaceBefore=10, spaceAfter=6
    )
    elements = [
        Paragraph(title, title_style),
        Paragraph(subtitle, sub_style),
        Paragraph(f"Generated: {dt.datetime.now():%Y-%m-%d %H:%M}", sub_style),
        Paragraph(f"Prepared by: {DEVELOPER} - {BRAND_NAME}", sub_style),
        Spacer(1, 10),
    ]
    if notes:
        elements.append(Paragraph("Summary &amp; Rationale", section_style))
        elements.append(Paragraph(notes, styles["Normal"]))
        elements.append(Spacer(1, 8))
    elements.append(Paragraph("Supporting Data", section_style))
    data = [list(df.columns)] + df.round(3).astype(str).values.tolist()
    tbl = Table(data, repeatRows=1)
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), rl_colors.HexColor(NAVY)),
        ("TEXTCOLOR", (0, 0), (-1, 0), rl_colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("GRID", (0, 0), (-1, -1), 0.5, rl_colors.HexColor("#cccccc")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [rl_colors.white, rl_colors.HexColor("#EEF3FA")]),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    elements.append(tbl)
    if chart_buf is not None:
        elements.append(Spacer(1, 12))
        elements.append(Paragraph("Visualization", section_style))
        chart_buf.seek(0)
        elements.append(RLImage(chart_buf, width=440, height=260))
    doc.build(elements)
    buf.seek(0)
    return buf


def render_export_bar(key_prefix: str, title: str, subtitle: str, notes: str,
                       df: pd.DataFrame, chart_buf: io.BytesIO = None):
    st.markdown("##### Export Results")
    c1, c2, c3, c4 = st.columns(4)
    with c1:
        st.download_button("CSV", df_to_csv_bytes(df),
                            file_name=f"{key_prefix}.csv", mime="text/csv",
                            use_container_width=True, key=f"{key_prefix}_csv")
    with c2:
        st.download_button("TXT", df_to_txt_bytes(title, subtitle, notes, df),
                            file_name=f"{key_prefix}.txt", mime="text/plain",
                            use_container_width=True, key=f"{key_prefix}_txt")
    with c3:
        word_buf = build_word_report(title, subtitle, notes, df, chart_buf)
        st.download_button("Word", word_buf,
                            file_name=f"{key_prefix}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True, key=f"{key_prefix}_docx")
    with c4:
        pdf_buf = build_pdf_report(title, subtitle, notes, df, chart_buf)
        st.download_button("PDF", pdf_buf,
                            file_name=f"{key_prefix}.pdf", mime="application/pdf",
                            use_container_width=True, key=f"{key_prefix}_pdf")


def blue_bar_chart(categories, values, ylabel, title):
    fig, ax = plt.subplots(figsize=(8, 4.2))
    norm_vals = (np.array(values) - min(values)) / (max(values) - min(values) + 1e-9)
    colors_list = cm.Blues(0.35 + 0.55 * norm_vals)
    bars = ax.bar(categories, values, color=colors_list, edgecolor=NAVY)
    ax.set_ylabel(ylabel)
    ax.set_title(title, color=NAVY, fontweight="bold")
    ax.spines[["top", "right"]].set_visible(False)
    plt.xticks(rotation=25, ha="right")
    for b, v in zip(bars, values):
        ax.text(b.get_x() + b.get_width() / 2, v, f"{v:.2f}",
                 ha="center", va="bottom", fontsize=8)
    fig.tight_layout()
    return fig


# --------------------------------------------------------------------------
# THEORY / FORMULA / SCHEMA EXPLAINER (educational panel shown atop each module)
# --------------------------------------------------------------------------
def render_explainer_section(icon: str, use_case_title: str, how_it_works_md: str,
                              formula_md: str, field_ref: pd.DataFrame,
                              schema_preview: pd.DataFrame, expanded: bool = True):
    """
    Renders a top-of-page educational block explaining:
      - How the use case works in plain language
      - The exact formula(s) applied
      - A field-by-field reference table (what each column means & why it matters)
      - A live 'schema scan' preview of sample data with those fields populated
    This is purely explanatory and does not alter any downstream calculation.
    """
    with st.expander(f"How {use_case_title} Works - Theory, Formulas and Field Guide",
                      expanded=expanded):
        st.markdown("###### Plain-Language Explanation")
        st.markdown(how_it_works_md)
        st.markdown("###### Formula(s) Applied")
        st.markdown(formula_md)
        st.markdown("###### Field Reference and Relevance")
        st.dataframe(field_ref, use_container_width=True, hide_index=True)
        st.markdown("###### Sample Data Schema Scan (illustrative preview)")
        st.caption("A sample of demo records showing every field described above, populated with example values.")
        st.dataframe(schema_preview, use_container_width=True, hide_index=True)


# --------------------------------------------------------------------------
# DATA UPLOAD HELPER
# --------------------------------------------------------------------------
def data_source_selector(demo_generator, required_cols, key):
    mode = st.radio(
        "Data Source", ["Synthetic Demo Data", "Upload Real Data (CSV / XLSX)"],
        horizontal=True, key=f"{key}_mode"
    )
    if mode == "Synthetic Demo Data":
        if st.button("Regenerate Demo Data", key=f"{key}_regen"):
            st.session_state[f"{key}_seed"] = np.random.randint(0, 100000)
        seed = st.session_state.get(f"{key}_seed", 42)
        df = demo_generator(seed)
        st.caption("Using synthetic demo data (deterministic seed - click Regenerate for a new sample).")
    else:
        up = st.file_uploader(f"Upload data (columns expected: {', '.join(required_cols)})",
                               type=["csv", "xlsx"], key=f"{key}_upload")
        if up is not None:
            df = pd.read_csv(up) if up.name.endswith("csv") else pd.read_excel(up)
            missing = [c for c in required_cols if c not in df.columns]
            if missing:
                st.error(f"Missing required columns: {missing}. Falling back to demo data.")
                df = demo_generator(42)
        else:
            st.info("No file uploaded yet - showing demo data as a preview.")
            df = demo_generator(42)
    return df


# --------------------------------------------------------------------------
# PAGE 1: STRATEGIC DECISION-MAKING  (Multi-Attribute Utility Theory)
# --------------------------------------------------------------------------
CRITERIA = ["Market_Growth", "Cost_Efficiency", "Risk_Level_Inv", "Innovation", "Feasibility"]


def gen_strategy_demo(seed):
    rng = np.random.default_rng(seed)
    names = [f"Strategy {c}" for c in "ABCDEF"]
    data = {"Alternative": names}
    for c in CRITERIA:
        data[c] = rng.integers(3, 10, size=len(names))
    return pd.DataFrame(data)


def page_strategic():
    title_bar("Strategic Decision-Making - Multi-Attribute Utility Theory (MAUT)")
    strat_field_ref = pd.DataFrame([
        {"Field": "Alternative", "Type": "Text (identifier)",
         "Description": "The name of the candidate strategy being evaluated (e.g. 'Strategy A').",
         "Relevance": "Row label only - not used in the score calculation itself."},
        {"Field": "Market_Growth", "Type": "Number (1-10)",
         "Description": "Estimated growth potential of the market this strategy targets.",
         "Relevance": "One of the weighted criteria summed into the Utility Score - higher is better."},
        {"Field": "Cost_Efficiency", "Type": "Number (1-10)",
         "Description": "How cost-efficient the strategy is relative to alternatives.",
         "Relevance": "Weighted criterion input - higher efficiency raises the Utility Score."},
        {"Field": "Risk_Level_Inv", "Type": "Number (1-10)",
         "Description": "Inverted risk rating: 10 = very low risk, 1 = very high risk.",
         "Relevance": "Weighted criterion input - stored inverted so 'higher is always better' holds for every column."},
        {"Field": "Innovation", "Type": "Number (1-10)",
         "Description": "Degree of innovation or competitive differentiation the strategy offers.",
         "Relevance": "Weighted criterion input into the Utility Score."},
        {"Field": "Feasibility", "Type": "Number (1-10)",
         "Description": "How practically achievable the strategy is with current capability/resources.",
         "Relevance": "Weighted criterion input into the Utility Score."},
        {"Field": "Rational_Utility_Score", "Type": "Computed (0-1)",
         "Description": "The weighted, normalized aggregate score produced by the model.",
         "Relevance": "The single number used to rank alternatives - higher = more rational choice."},
        {"Field": "Rank", "Type": "Computed (integer)",
         "Description": "Ordinal position after sorting by Rational_Utility_Score, descending.",
         "Relevance": "Rank 1 is the recommended, utility-maximizing strategy."},
    ])
    strat_schema_preview = gen_strategy_demo(42).assign(
        Rational_Utility_Score="(computed after weights are set)", Rank="(computed)"
    )
    render_explainer_section(
        icon="",
        use_case_title="Strategic Decision-Making",
        how_it_works_md=(
            "Rational Choice Theory holds that when a decision-maker faces several strategic "
            "alternatives, the rational course of action is to score each alternative against "
            "the criteria that matter, weight those criteria by relative importance, and choose "
            "the alternative with the highest combined **expected utility**.\n\n"
            "**Steps this module follows:**\n"
            "1. Collect raw scores (1-10) for each alternative across five criteria.\n"
            "2. **Normalize** each criterion to a common 0-1 scale (min-max normalization), so a "
            "criterion measured in different units or ranges doesn't unfairly dominate.\n"
            "3. **Weight** each normalized criterion using the sliders you set (weights auto-sum to 100%).\n"
            "4. **Sum** the weighted values into one Rational Utility Score per alternative.\n"
            "5. **Rank** alternatives from highest to lowest score - Rank 1 is the rational recommendation."
        ),
        formula_md=(
            "**Step 1 - Min-Max Normalization** (per criterion, across all alternatives):\n\n"
            "> Normalized_i = (x_i - min(x)) / (max(x) - min(x))\n\n"
            "**Step 2 - Weight normalization** (so weights always sum to 1):\n\n"
            "> weight_c = raw_slider_c / sum(all raw_sliders)\n\n"
            "**Step 3 - Rational Utility Score** (weighted sum of expected utility):\n\n"
            "> Utility(alternative) = sum over all criteria c of ( weight_c x Normalized_c )\n\n"
            "This is the standard **Multi-Attribute Utility Theory (MAUT)** additive model, "
            "and is mathematically equivalent to computing an expected utility across weighted "
            "attributes."
        ),
        field_ref=strat_field_ref,
        schema_preview=strat_schema_preview,
    )
    st.markdown(
        """<div class="knet-subheader">
        Rational Choice Theory holds that an actor facing several strategic
        alternatives ranks them by <b>expected utility</b> - a weighted
        combination of relevant criteria - and rationally selects the
        alternative with the highest aggregate utility score.
        </div>""",
        unsafe_allow_html=True,
    )
    df = data_source_selector(gen_strategy_demo, ["Alternative"] + CRITERIA, "strat")
    criteria_present = [c for c in CRITERIA if c in df.columns]
    st.markdown('<div class="knet-card">', unsafe_allow_html=True)
    st.markdown("**Assign Criterion Weights** (rational preference structure - auto-normalized to sum to 1)")
    cols = st.columns(len(criteria_present))
    raw_weights = {}
    defaults = [25, 20, 20, 20, 15]
    for i, c in enumerate(criteria_present):
        with cols[i]:
            raw_weights[c] = st.slider(c.replace("_", " "), 0, 100,
                                        defaults[i % len(defaults)], key=f"w_{c}")
    total_w = sum(raw_weights.values()) or 1
    weights = {c: w / total_w for c, w in raw_weights.items()}
    st.markdown("</div>", unsafe_allow_html=True)
    norm = df.copy()
    for c in criteria_present:
        col = df[c].astype(float)
        rng_ = (col.max() - col.min()) or 1
        norm[c] = (col - col.min()) / rng_
    norm["Rational_Utility_Score"] = sum(norm[c] * weights[c] for c in criteria_present)
    result = df.copy()
    result["Rational_Utility_Score"] = norm["Rational_Utility_Score"].round(3)
    result = result.sort_values("Rational_Utility_Score", ascending=False).reset_index(drop=True)
    result.insert(0, "Rank", result.index + 1)
    st.subheader("Ranked Alternatives")
    st.dataframe(result, use_container_width=True, hide_index=True)
    top = result.iloc[0]
    notes = (
        f"Based on the weighted-utility model, '{top['Alternative']}' is the rational choice "
        f"with the highest expected utility score of {top['Rational_Utility_Score']:.3f}. "
        f"Criterion weights applied: " +
        ", ".join(f"{c.replace('_',' ')}={weights[c]*100:.0f}%" for c in criteria_present) + "."
    )
    st.markdown('<div class="knet-card">', unsafe_allow_html=True)
    st.markdown(f"**Rational Recommendation:** {notes}")
    st.markdown("</div>", unsafe_allow_html=True)
    fig = blue_bar_chart(result["Alternative"], result["Rational_Utility_Score"],
                          "Utility Score", "Strategic Alternatives - Rational Utility Ranking")
    st.pyplot(fig)
    chart_buf = fig_to_png_buffer(fig)
    render_export_bar("strategic_decision", "Strategic Decision-Making Report",
                       "Multi-Attribute Utility Theory Analysis", notes, result, chart_buf)
    footer()


# --------------------------------------------------------------------------
# PAGE 2: RESOURCE ALLOCATION  (Constrained Utility Maximization)
# --------------------------------------------------------------------------
def gen_resource_demo(seed):
    rng = np.random.default_rng(seed)
    names = [f"Initiative {c}" for c in "PQRSTU"]
    n = len(names)
    utility = rng.uniform(0.4, 0.95, n).round(2)
    risk = rng.uniform(0.1, 0.7, n).round(2)
    min_alloc = rng.integers(5, 15, n)
    max_alloc = min_alloc + rng.integers(10, 40, n)
    return pd.DataFrame({
        "Initiative": names,
        "Expected_Utility": utility,
        "Risk_Factor": risk,
        "Min_Allocation": min_alloc,
        "Max_Allocation": max_alloc,
    })


def page_resource():
    title_bar("Resource Allocation - Constrained Rational Utility Maximization")
    res_field_ref = pd.DataFrame([
        {"Field": "Initiative", "Type": "Text (identifier)",
         "Description": "Name of the project, department, or initiative competing for funding.",
         "Relevance": "Row label only - identifies which allocation belongs to which initiative."},
        {"Field": "Expected_Utility", "Type": "Number (0-1)",
         "Description": "Estimated payoff/ROI per unit of resource invested in this initiative.",
         "Relevance": "The main driver the optimizer tries to maximize - higher utility attracts more funding."},
        {"Field": "Risk_Factor", "Type": "Number (0-1)",
         "Description": "Estimated riskiness of the initiative (probability/severity of it underdelivering).",
         "Relevance": "Subtracted from Expected_Utility (scaled by your Risk Aversion slider) to penalize risky bets."},
        {"Field": "Min_Allocation", "Type": "Number (units)",
         "Description": "The smallest resource commitment the initiative can realistically operate on.",
         "Relevance": "Hard lower bound the optimizer must respect for every initiative that receives funding."},
        {"Field": "Max_Allocation", "Type": "Number (units)",
         "Description": "The largest resource amount the initiative can effectively absorb/utilize.",
         "Relevance": "Hard upper bound - prevents over-funding a single initiative beyond its capacity."},
        {"Field": "Allocated", "Type": "Computed (units)",
         "Description": "The resource amount the optimizer assigns to this initiative.",
         "Relevance": "The output of the linear program - always between Min_Allocation and Max_Allocation."},
        {"Field": "% of Budget", "Type": "Computed (%)",
         "Description": "Allocated amount expressed as a share of the total allocated budget.",
         "Relevance": "Helps quickly compare relative funding priority across initiatives."},
    ])
    res_schema_preview = gen_resource_demo(42).assign(
        Allocated="(computed by optimizer)", **{"% of Budget": "(computed)"}
    )
    render_explainer_section(
        icon="",
        use_case_title="Resource Allocation",
        how_it_works_md=(
            "Under scarcity, a rational actor cannot fund every initiative fully - resources must "
            "be allocated to **maximize total expected benefit** while respecting real-world "
            "constraints (a fixed budget, and minimum/maximum funding levels each initiative can "
            "actually use).\n\n"
            "**Steps this module follows:**\n"
            "1. Take each initiative's Expected Utility and Risk Factor.\n"
            "2. Adjust utility downward for risk using your chosen Risk Aversion Factor.\n"
            "3. Solve a **linear program**: distribute the total budget across initiatives to "
            "maximize the sum of risk-adjusted utility, while keeping every initiative's "
            "allocation within its Min/Max bounds and the grand total within budget.\n"
            "4. Present the resulting allocation, ranked by amount funded."
        ),
        formula_md=(
            "**Step 1 - Risk-adjusted utility:**\n\n"
            "> Adjusted_Utility_i = Expected_Utility_i - (Risk_Aversion x Risk_Factor_i)\n\n"
            "**Step 2 - Optimization objective (maximize total adjusted utility):**\n\n"
            "> maximize sum of ( Adjusted_Utility_i x Allocation_i )\n\n"
            "**Step 3 - Constraints:**\n\n"
            "> sum of Allocation_i <= Total_Budget\n"
            "> Min_Allocation_i <= Allocation_i <= Max_Allocation_i  for every initiative i\n\n"
            "Solved using linear programming (`scipy.optimize.linprog`, HiGHS solver) - the same "
            "constrained-optimization approach used in classical rational-actor resource theory."
        ),
        field_ref=res_field_ref,
        schema_preview=res_schema_preview,
    )
    st.markdown(
        """<div class="knet-subheader">
        Under scarcity, rational actors allocate limited resources to maximize
        total expected utility, subject to constraints (budget ceiling,
        minimum commitments, maximum capacity). This module solves that
        allocation as a linear program.
        </div>""",
        unsafe_allow_html=True,
    )
    req_cols = ["Initiative", "Expected_Utility", "Risk_Factor", "Min_Allocation", "Max_Allocation"]
    df = data_source_selector(gen_resource_demo, req_cols, "res")
    st.markdown('<div class="knet-card">', unsafe_allow_html=True)
    total_budget = st.slider("Total Available Budget (units)", 20, 500,
                              int(df["Min_Allocation"].sum() + df["Max_Allocation"].sum()) // 3, step=5)
    risk_aversion = st.slider("Risk Aversion Factor (penalizes risky initiatives)", 0.0, 1.0, 0.3, 0.05)
    st.markdown("</div>", unsafe_allow_html=True)
    n = len(df)
    adj_utility = df["Expected_Utility"] - risk_aversion * df["Risk_Factor"]
    c = -adj_utility.values  # linprog minimizes, so negate to maximize
    A_ub = [np.ones(n)]
    b_ub = [total_budget]
    bounds = list(zip(df["Min_Allocation"].values, df["Max_Allocation"].values))
    feasible_min = df["Min_Allocation"].sum()
    if feasible_min > total_budget:
        st.error(f"Budget ({total_budget}) is below the sum of minimum commitments "
                  f"({feasible_min}). Increase budget or adjust minimums.")
        return
    res_lp = linprog(c, A_ub=A_ub, b_ub=b_ub, bounds=bounds, method="highs")
    result = df.copy()
    if res_lp.success:
        result["Allocated"] = np.round(res_lp.x, 2)
    else:
        st.warning("Optimizer could not find an optimal solution; showing proportional fallback.")
        weights_fallback = adj_utility / adj_utility.sum()
        result["Allocated"] = np.round(weights_fallback * total_budget, 2)
    result["% of Budget"] = (result["Allocated"] / result["Allocated"].sum() * 100).round(1)
    result = result.sort_values("Allocated", ascending=False).reset_index(drop=True)
    st.subheader("Rational Resource Allocation Plan")
    st.dataframe(result, use_container_width=True, hide_index=True)
    top = result.iloc[0]
    notes = (
        f"Given a total budget of {total_budget} units and a risk-aversion factor of {risk_aversion}, "
        f"the utility-maximizing allocation directs the largest share "
        f"({top['Allocated']:.1f} units, {top['% of Budget']:.1f}%) to '{top['Initiative']}'. "
        f"All allocations respect stated minimum/maximum constraints while maximizing aggregate "
        f"expected utility across initiatives."
    )
    st.markdown('<div class="knet-card">', unsafe_allow_html=True)
    st.markdown(f"**Rational Recommendation:** {notes}")
    st.markdown("</div>", unsafe_allow_html=True)
    fig, axes = plt.subplots(1, 2, figsize=(10, 4.2))
    axes[0].pie(result["Allocated"], labels=result["Initiative"], autopct="%1.0f%%",
                colors=cm.Blues(np.linspace(0.4, 0.85, len(result))),
                wedgeprops={"edgecolor": "white"})
    axes[0].set_title("Allocation Share", color=NAVY, fontweight="bold")
    x = np.arange(len(result))
    axes[1].bar(x - 0.2, result["Max_Allocation"], width=0.4, label="Max Requested", color="#B7C9E2")
    axes[1].bar(x + 0.2, result["Allocated"], width=0.4, label="Allocated", color=NAVY)
    axes[1].set_xticks(x)
    axes[1].set_xticklabels(result["Initiative"], rotation=25, ha="right")
    axes[1].set_title("Requested vs. Allocated", color=NAVY, fontweight="bold")
    axes[1].legend()
    axes[1].spines[["top", "right"]].set_visible(False)
    fig.tight_layout()
    st.pyplot(fig)
    chart_buf = fig_to_png_buffer(fig)
    render_export_bar("resource_allocation", "Resource Allocation Report",
                       "Constrained Utility-Maximization Analysis", notes, result, chart_buf)
    footer()


# --------------------------------------------------------------------------
# PAGE 3: THREAT ASSESSMENT  (Expected-Value Risk Scoring)
# --------------------------------------------------------------------------
def gen_threat_demo(seed):
    rng = np.random.default_rng(seed)
    names = ["Cyber Intrusion", "Supply Chain Disruption", "Insider Threat",
             "Geopolitical Instability", "Regulatory Change", "Market Volatility",
             "Physical Security Breach", "Data Privacy Violation"]
    likelihood = rng.integers(1, 6, len(names))
    impact = rng.integers(1, 6, len(names))
    return pd.DataFrame({"Threat": names, "Likelihood": likelihood, "Impact": impact})


def categorize_risk(score):
    if score <= 5:
        return "Low"
    elif score <= 12:
        return "Medium"
    elif score <= 19:
        return "High"
    return "Critical"


def page_threat():
    title_bar("Threat Assessment - Expected-Value Risk Scoring")
    threat_field_ref = pd.DataFrame([
        {"Field": "Threat", "Type": "Text (identifier)",
         "Description": "Name or short description of the threat scenario being assessed.",
         "Relevance": "Row label only - identifies which risk score/category belongs to which threat."},
        {"Field": "Likelihood", "Type": "Number (1-5)",
         "Description": "Estimated probability of the threat occurring (1 = rare, 5 = near-certain).",
         "Relevance": "One of the two factors multiplied together to compute Risk_Score."},
        {"Field": "Impact", "Type": "Number (1-5)",
         "Description": "Estimated severity of consequences if the threat materializes (1 = minor, 5 = severe).",
         "Relevance": "The other factor multiplied into Risk_Score - high-impact threats are prioritized even at moderate likelihood."},
        {"Field": "Risk_Score", "Type": "Computed (1-25)",
         "Description": "The expected-value risk score: Likelihood x Impact.",
         "Relevance": "The core rational prioritization metric - higher score means higher expected cost of ignoring the threat."},
        {"Field": "Risk_Category", "Type": "Computed (Low/Medium/High/Critical)",
         "Description": "A human-readable bucket derived from Risk_Score using fixed thresholds.",
         "Relevance": "Makes the numeric score actionable for triage and reporting (color-coded in the table)."},
        {"Field": "Priority_Rank", "Type": "Computed (integer)",
         "Description": "Ordinal position after sorting all threats by Risk_Score, descending.",
         "Relevance": "Rank 1 is the threat requiring the most urgent rational resource commitment."},
    ])
    threat_schema_preview = gen_threat_demo(42).assign(
        Risk_Score="(computed = Likelihood x Impact)", Risk_Category="(computed)"
    )
    render_explainer_section(
        icon="",
        use_case_title="Threat Assessment",
        how_it_works_md=(
            "Rational threat assessment treats risk as a classic **expected-value problem**: the "
            "expected cost of a threat is not just how bad it would be, but how bad it would be "
            "*multiplied by* how likely it is to happen. A rational actor with limited attention "
            "and countermeasure resources should therefore prioritize threats by this combined "
            "expected value, not by likelihood or impact alone.\n\n"
            "**Steps this module follows:**\n"
            "1. Take a Likelihood rating (1-5) and an Impact rating (1-5) for each threat.\n"
            "2. Multiply them to get a single **Risk Score** (expected value of the threat).\n"
            "3. Bucket the score into a category (Low/Medium/High/Critical) using fixed thresholds.\n"
            "4. Rank all threats by Risk Score, highest first, so mitigation effort follows "
            "expected cost rather than gut feeling."
        ),
        formula_md=(
            "**Risk Score (expected value):**\n\n"
            "> Risk_Score = Likelihood x Impact\n\n"
            "**Category thresholds** (max possible score is 5 x 5 = 25):\n\n"
            "> Risk_Score <= 5  ->  **Low**\n"
            "> 6 <= Risk_Score <= 12  ->  **Medium**\n"
            "> 13 <= Risk_Score <= 19  ->  **High**\n"
            "> Risk_Score >= 20  ->  **Critical**\n\n"
            "This mirrors the standard 5x5 risk-matrix approach used across security, "
            "operational, and geopolitical risk frameworks."
        ),
        field_ref=threat_field_ref,
        schema_preview=threat_schema_preview,
    )
    st.markdown(
        """<div class="knet-subheader">
        Rational threat assessment treats risk as an expected-value problem:
        <b>Risk Score = Likelihood x Impact</b>. Rational prioritization
        directs limited attention and countermeasure resources toward the
        threats with the highest expected cost.
        </div>""",
        unsafe_allow_html=True,
    )
    df = data_source_selector(gen_threat_demo, ["Threat", "Likelihood", "Impact"], "threat")
    result = df.copy()
    result["Risk_Score"] = result["Likelihood"] * result["Impact"]
    result["Risk_Category"] = result["Risk_Score"].apply(categorize_risk)
    result = result.sort_values("Risk_Score", ascending=False).reset_index(drop=True)
    result.insert(0, "Priority_Rank", result.index + 1)
    st.subheader("Prioritized Threat Register")

    def highlight_cat(val):
        colors_map = {"Low": "#D9F2D9", "Medium": "#FFF3C4",
                      "High": "#FFD9B3", "Critical": "#FFB3B3"}
        return f"background-color: {colors_map.get(val, '')}"

    # NOTE: pandas 2.1+ renamed Styler.applymap -> Styler.map, and pandas 3.x
    # removed applymap entirely. Use .map when available, falling back to
    # .applymap on older pandas so this keeps working either way.
    styler = result.style
    if hasattr(styler, "map"):
        styler = styler.map(highlight_cat, subset=["Risk_Category"])
    else:
        styler = styler.applymap(highlight_cat, subset=["Risk_Category"])
    st.dataframe(styler, use_container_width=True, hide_index=True)

    top = result.iloc[0]
    crit_count = (result["Risk_Category"] == "Critical").sum()
    high_count = (result["Risk_Category"] == "High").sum()
    notes = (
        f"The highest-priority threat is '{top['Threat']}' with a risk score of "
        f"{top['Risk_Score']} ({top['Risk_Category']}). The register currently shows "
        f"{crit_count} Critical and {high_count} High-severity threats requiring "
        f"immediate rational resource commitment to mitigation."
    )
    st.markdown('<div class="knet-card">', unsafe_allow_html=True)
    st.markdown(f"**Rational Recommendation:** {notes}")
    st.markdown("</div>", unsafe_allow_html=True)
    fig, axes = plt.subplots(1, 2, figsize=(10.5, 4.4))
    cat_colors = {"Low": "#4CAF50", "Medium": "#E8B93B", "High": "#E8792B", "Critical": "#C62828"}
    for cat, sub in result.groupby("Risk_Category"):
        axes[0].scatter(sub["Likelihood"], sub["Impact"],
                         s=sub["Risk_Score"] * 25, color=cat_colors[cat],
                         label=cat, alpha=0.85, edgecolor="white")
    for _, r in result.iterrows():
        axes[0].annotate(r["Threat"], (r["Likelihood"], r["Impact"]),
                          fontsize=6.5, xytext=(3, 3), textcoords="offset points")
    axes[0].set_xlabel("Likelihood")
    axes[0].set_ylabel("Impact")
    axes[0].set_title("Risk Matrix", color=NAVY, fontweight="bold")
    axes[0].set_xlim(0, 6)
    axes[0].set_ylim(0, 6)
    axes[0].legend(fontsize=7, loc="upper left")
    axes[0].spines[["top", "right"]].set_visible(False)
    bar_colors = [cat_colors[c] for c in result["Risk_Category"]]
    axes[1].barh(result["Threat"], result["Risk_Score"], color=bar_colors, edgecolor="white")
    axes[1].invert_yaxis()
    axes[1].set_xlabel("Risk Score")
    axes[1].set_title("Threats by Risk Score", color=NAVY, fontweight="bold")
    axes[1].spines[["top", "right"]].set_visible(False)
    fig.tight_layout()
    st.pyplot(fig)
    chart_buf = fig_to_png_buffer(fig)
    render_export_bar("threat_assessment", "Threat Assessment Report",
                       "Expected-Value Risk Scoring Analysis", notes, result, chart_buf)
    footer()


# --------------------------------------------------------------------------
# PAGE 0: HOME / ABOUT
# --------------------------------------------------------------------------
def page_home():
    title_bar("Home - Overview of Rational Decision Theory")
    st.markdown(
        """
        <div class="knet-card">
        <h4 style="color:#0A2F5C;">What is Rational Choice / Rational Decision Theory?</h4>
        <p>Rational Decision Theory posits that a decision-maker, when facing a set of
        alternatives under known or estimable criteria, will:</p>
        <ol>
          <li>Identify all feasible alternatives and relevant criteria.</li>
          <li>Assign weights or probabilities reflecting the relative importance
              or likelihood of each criterion.</li>
          <li>Compute the <b>expected utility</b> (or expected value) of each
              alternative as a function of those weights/probabilities.</li>
          <li>Rank alternatives and select the one that maximizes expected utility,
              subject to any binding constraints (e.g. budget, capacity).</li>
        </ol>
        <p>This application operationalizes that framework across three domains:</p>
        <ul>
          <li><b>Strategic Decision-Making</b> - Multi-Attribute Utility Theory (MAUT)
              ranks strategic alternatives by weighted criteria.</li>
          <li><b>Resource Allocation</b> - Constrained optimization allocates scarce
              resources to maximize aggregate expected utility.</li>
          <li><b>Threat Assessment</b> - Expected-value risk scoring
              (Likelihood x Impact) prioritizes threats rationally.</li>
        </ul>
        </div>
        """,
        unsafe_allow_html=True,
    )
    st.markdown(
        """
        <div class="knet-card">
        <h4 style="color:#0A2F5C;">Application Capabilities</h4>
        <ul>
          <li>Synthetic demo data generation for instant exploration</li>
          <li>Real data upload via CSV / XLSX</li>
          <li>Interactive charts (bar, pie, scatter risk matrix)</li>
          <li>Export to <b>PDF</b>, <b>Word (.docx)</b>, <b>TXT</b>, and <b>CSV</b></li>
          <li>Adjustable weights, budgets, and risk-aversion parameters</li>
        </ul>
        </div>
        """,
        unsafe_allow_html=True,
    )
    footer()


# --------------------------------------------------------------------------
# NAVIGATION
# --------------------------------------------------------------------------
def main():
    st.sidebar.markdown(
        f"""
        <div style="text-align:center; padding: 10px 0 18px 0;">
            <div style="font-weight:800; font-size:16px; letter-spacing:0.5px;">
                RATIONAL DSS
            </div>
            <div style="font-size:11px; color:#9fb3d1;">{BRAND_NAME}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )
    page = st.sidebar.radio(
        "Navigation",
        ["Home", "Strategic Decision-Making", "Resource Allocation", "Threat Assessment"],
        label_visibility="collapsed",
    )
    st.sidebar.markdown("---")
    st.sidebar.caption(f"Developed by {DEVELOPER}\n{BRAND_NAME}")
    if page == "Home":
        page_home()
    elif page == "Strategic Decision-Making":
        page_strategic()
    elif page == "Resource Allocation":
        page_resource()
    elif page == "Threat Assessment":
        page_threat()


if __name__ == "__main__":
    main()
