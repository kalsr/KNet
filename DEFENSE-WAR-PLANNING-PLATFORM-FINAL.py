

# DEFENSE-WAR-PLANNING-PLATFORM-Final.py

import streamlit as st
import pandas as pd
import numpy as np
import networkx as nx
import matplotlib.pyplot as plt
from reportlab.platypus import SimpleDocTemplate, Paragraph
from datetime import datetime

# Safe import for WORD (fix crash issue)
try:
    from docx import Document
    DOCX_AVAILABLE = True
except:
    DOCX_AVAILABLE = False

# Optional YOLO
try:
    from ultralytics import YOLO
    YOLO_AVAILABLE = True
except:
    YOLO_AVAILABLE = False

st.set_page_config(layout="wide")

# ---------------- STYLE ----------------
st.markdown("""
<style>
.title-bar {
    color:#0033cc;
    font-size:42px;
    font-weight:bold;
    text-align:center;
}
.subtitle {
    color:#0033cc;
    font-size:20px;
    font-weight:bold;
    text-align:center;
}
.section-header {
    color:#0047AB;
    font-size:24px;
    font-weight:bold;
}
</style>
""", unsafe_allow_html=True)

# ---------------- TITLE ----------------
st.markdown("<div class='title-bar'>Kalsnet (KNet) AI WAR COMMAND PLATFORM</div>", unsafe_allow_html=True)
st.markdown("<div class='subtitle'>Developed by Randy Singh – KNet Consulting Group</div>", unsafe_allow_html=True)

# ---------------- ROLE ----------------
role = st.sidebar.selectbox("Select Role", ["Commander", "Analyst"])

st.markdown("### 🔹 Role Description")

if role == "Commander":
    st.info("""
**Commander Role:**  
Responsible for mission success, strategic decisions, and operational execution.  
Uses AI outputs (COA, Threat Scores, Logistics) to take action.

👉 Focus: **Decision Making + Execution**
""")
else:
    st.info("""
**Analyst Role:**  
Responsible for analyzing data, identifying patterns, and generating intelligence.  
Supports commanders with insights.

👉 Focus: **Data + Intelligence**
""")

# ---------------- MENU ----------------
menu = st.sidebar.radio("Select Module", [
    "Battlefield Map",
    "COA Decision Engine",
    "Threat Detection",
    "Logistics Optimization"
])

# ---------------- EXPORT FUNCTION ----------------
def export_data(df, name="output"):

    st.download_button("Download CSV", df.to_csv(), f"{name}.csv")
    st.download_button("Download JSON", df.to_json(), f"{name}.json")

    # PDF
    pdf_file = f"{name}.pdf"
    doc = SimpleDocTemplate(pdf_file)
    doc.build([Paragraph(df.head().to_string())])
    with open(pdf_file, "rb") as f:
        st.download_button("Download PDF", f, file_name=pdf_file)

    # WORD (SAFE)
    if DOCX_AVAILABLE:
        word_file = f"{name}.docx"
        document = Document()
        document.add_heading(name.upper(), 0)
        document.add_paragraph(df.to_string())
        document.save(word_file)

        with open(word_file, "rb") as f:
            st.download_button("Download WORD", f, file_name=word_file)
    else:
        st.warning("WORD export not available (python-docx not installed)")

# ==========================================================
# ---------------- BATTLEFIELD MAP (FIXED) ----------------
# ==========================================================
if menu == "Battlefield Map":

    st.markdown("<div class='section-header'>Battlefield Map</div>", unsafe_allow_html=True)

    st.info("""
**WHAT THIS MODULE DOES:**

The Battlefield Map provides a **real-time geospatial visualization** of operational environments.

**KEY CAPABILITIES:**
- Displays command centers, threat zones, and operational regions
- Helps commanders understand battlefield positioning
- Allows analysts to monitor geographic threat distribution

**TOOLS USED:**
- Folium (Interactive Maps)
- Streamlit Integration

**REAL-WORLD USE:**
Used in military and intelligence systems for:
- Situational awareness
- Mission planning
- Threat tracking

👉 This acts as a **live operational dashboard map**
""")

    # SAFE MAP LOADING
    try:
        import folium
        from streamlit_folium import st_folium

        m = folium.Map(location=[20, 0], zoom_start=2)

        # Sample markers
        folium.Marker([28.6, 77.2], tooltip="Command Center").add_to(m)
        folium.Marker([34.5, 69.2], tooltip="Threat Zone").add_to(m)

        st_folium(m, width=1200, height=600)

    except:
        st.warning("Interactive map not available (folium not installed). Showing fallback view.")

        # Fallback visualization
        df_map = pd.DataFrame({
            "lat": [28.6, 34.5],
            "lon": [77.2, 69.2]
        })

        st.map(df_map)

# ==========================================================
# ---------------- COA ENGINE ----------------
# ==========================================================
elif menu == "COA Decision Engine":

    st.markdown("<div class='section-header'>COA Decision Engine</div>", unsafe_allow_html=True)

    st.info("""
**DATA FIELDS EXPLAINED:**
- Terrain → Type of battlefield
- Enemy → Enemy force size
- Logistics → Available resources
- Weather → Environmental impact
""")

    n = st.slider("Synthetic Scenarios", 0, 200, 50)

    if st.button("Generate COA Data"):
        df = pd.DataFrame({
            "Terrain": np.random.choice(["Urban","Desert","Mountain"], n),
            "Enemy": np.random.randint(50,500,n),
            "Logistics": np.random.randint(50,200,n),
            "Weather": np.random.rand(n)
        })
        st.session_state.coa = df

    df = st.session_state.get("coa", pd.DataFrame())
    st.write(df)

    if not df.empty:

        df["Success"] = (df["Logistics"]/(df["Enemy"]+1))*df["Weather"]

        st.markdown("**Success = (Logistics / Enemy) × Weather**")

        st.bar_chart(df["Success"])

        G = nx.DiGraph()
        for i in range(len(df)):
            G.add_edge("Start", f"COA_{i}", weight=df["Success"][i])

        fig, ax = plt.subplots()
        nx.draw(G, with_labels=True)
        st.pyplot(fig)

        export_data(df, "coa")

# ==========================================================
# ---------------- THREAT DETECTION ----------------
# ==========================================================
elif menu == "Threat Detection":

    st.markdown("<div class='section-header'>Threat Detection</div>", unsafe_allow_html=True)

    n = st.slider("Synthetic Threat Data", 0, 200, 50)

    if st.button("Generate Threat Data"):
        df = pd.DataFrame({
            "Speed": np.random.randint(10,100,n),
            "Distance": np.random.randint(1,50,n),
            "Signal": np.random.rand(n)
        })
        st.session_state.threat = df

    df = st.session_state.get("threat", pd.DataFrame())
    st.write(df)

    if not df.empty:

        df["ThreatScore"] = df["Speed"]/(df["Distance"]+1)

        st.bar_chart(df["ThreatScore"])

        G = nx.Graph()
        for i in range(len(df)):
            G.add_edge("Base", f"T{i}", weight=df["ThreatScore"][i])

        fig, ax = plt.subplots()
        nx.draw(G, with_labels=True)
        st.pyplot(fig)

        export_data(df, "threat")

# ==========================================================
# ---------------- LOGISTICS ----------------
# ==========================================================
elif menu == "Logistics Optimization":

    st.markdown("<div class='section-header'>Logistics Optimization</div>", unsafe_allow_html=True)

    n = st.slider("Synthetic Logistics Data", 0, 200, 50)

    if st.button("Generate Logistics Data"):
        df = pd.DataFrame({
            "From": np.random.choice(["BaseA","BaseB"], n),
            "To": np.random.choice(["Zone1","Zone2"], n),
            "Cost": np.random.randint(100,500,n),
            "Time": np.random.randint(1,20,n)
        })
        st.session_state.log = df

    df = st.session_state.get("log", pd.DataFrame())
    st.write(df)

    if not df.empty:

        df["Efficiency"] = df["Cost"]/df["Time"]

        st.bar_chart(df["Efficiency"])

        G = nx.Graph()
        for i in range(len(df)):
            G.add_edge(df["From"][i], df["To"][i], weight=df["Cost"][i])

        fig, ax = plt.subplots()
        nx.draw(G, with_labels=True)
        st.pyplot(fig)

        export_data(df, "logistics")

# ---------------- FOOTER ----------------
st.markdown("---")
st.caption(f"System Time: {datetime.now()}")