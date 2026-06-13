# scrm_streamlit_app.py – Enhanced Edition
# Developed by Randy Singh, Kalsnet (KNet) Consulting Group

import streamlit as st
import json
import io
from datetime import datetime
import pandas as pd

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
except ImportError:
    canvas = None

# ------------------------------------------------------------------------------------
# Core SCRM knowledge base (categories 1–7, richly expanded)
# ------------------------------------------------------------------------------------

SCRM_CATEGORIES = {
    1: {
        "name": "SCRM Stakeholders Tiers",
        "description": (
            "This category models organizational stakeholders across three tiers: "
            "executive leadership (Tier 1), mission/business management (Tier 2), and "
            "systems management (Tier 3). It aligns with NIST SP 800-161r1 multi-tier "
            "risk management for ICT supply chains. Each tier has distinct roles, "
            "responsibilities, and activities that collectively ensure a holistic "
            "approach to supply chain risk."
        ),
        "subcategories": {
            "Tier 1 – Organization Stakeholders": {
                "text": (
                    "Executive Leadership (CEO, CIO, COO, CFO, CISO, CTO, General Counsel, "
                    "Chief Risk Officer, etc.) constitute Tier 1. They provide risk executive "
                    "functions and ensure that ICT SCRM mitigation strategies are cost-effective, "
                    "efficient, and consistent with the strategic goals of the organization.\n\n"
                    "Typical Tier-1 responsibilities:\n"
                    "- Define corporate strategy, policy, goals, and objectives.\n"
                    "- Establish ICT SCRM policies based on external requirements (laws, "
                    "regulations, EO 14028, NIST CSF, CMMC, etc.).\n"
                    "- Set risk tolerance and appetite for ICT supply chain risks.\n"
                    "- Charter a dedicated ICT SCRM team and fund it appropriately.\n"
                    "- Integrate SCRM with enterprise risk management (ERM) and board reporting.\n"
                    "- Respond to supply-chain incidents at the enterprise level.\n\n"
                    "Recommendations:\n"
                    "★ Appoint a dedicated Supply Chain Risk Officer (SCRO) at the C-suite.\n"
                    "★ Mandate annual SCRM tabletop exercises for executive leadership.\n"
                    "★ Include supply-chain KRIs (Key Risk Indicators) in board dashboards.\n"
                    "★ Align SCRM investments to annual budget cycles with dedicated funding."
                ),
                "examples": [
                    "SolarWinds (2020): Executives were unaware of a compromised build pipeline "
                    "inserting malicious code (SUNBURST) into Orion updates. Better Tier-1 "
                    "oversight—mandatory supply-chain security reviews on build systems—could "
                    "have triggered earlier detection.",
                    "Apple M-series supply chain: Apple's board-level commitment to hardware "
                    "security drove investment in custom silicon, reducing reliance on third-party "
                    "chip vendors with opaque supply chains.",
                    "CISO defines acceptable risk thresholds for using offshore software "
                    "development vendors and publishes them as binding policy.",
                ],
                "mermaid": """
flowchart TD
    A[Board of Directors] --> B[C-Suite: CEO / CISO / CIO / CTO]
    B --> C[Define Enterprise SCRM Policy]
    C --> D[Set Risk Tolerance & Appetite]
    D --> E[Charter & Fund SCRM Team]
    E --> F[Integrate with Enterprise Risk Mgmt]
    F --> G[Board-level KRI Reporting]
    G --> H{Risk Acceptable?}
    H -- Yes --> I[Continue Operations]
    H -- No --> J[Invoke Risk Response]
    J --> C
""",
                "ascii_diagram": """
┌─────────────────────────────────────────────────┐
│              TIER 1 – EXECUTIVE                  │
│  CEO ─── CISO ─── CIO ─── CTO ─── CFO ─── GC   │
│                      │                           │
│              SCRM Policy & Governance            │
│          Risk Tolerance │ Budget Allocation      │
│          Supply Chain Risk Officer (SCRO)        │
└──────────────────────┬──────────────────────────┘
                       │ directs
              ┌────────▼────────┐
              │   TIER 2        │
              │ Mission/Business│
              └────────┬────────┘
                       │ directs
              ┌────────▼────────┐
              │   TIER 3        │
              │   IT Systems    │
              └─────────────────┘
"""
            },
            "Tier 1 – Organization Activities": {
                "text": (
                    "Tier-1 activities translate strategy into concrete policies and enterprise-wide "
                    "practices. These are the operational heartbeat of executive SCRM governance.\n\n"
                    "Key activities:\n"
                    "- Establish ICT SCRM policies and dedicated funding models.\n"
                    "- Map mission/business requirements (cost, schedule, performance, security, "
                    "privacy, quality, safety) to SCRM needs.\n"
                    "- Embed SCRM requirements into corporate processes: procurement, architecture, "
                    "HR (background checks), legal (contract review).\n"
                    "- Maintain a strategic supplier list and prohibited vendor registry.\n"
                    "- Oversee threat intelligence sharing (ISACs, government partnerships).\n\n"
                    "Recommendations:\n"
                    "★ Publish and annually review a Prohibited/High-Risk Vendor List.\n"
                    "★ Participate in sector ISACs for supply chain threat intelligence.\n"
                    "★ Require legal review of all ICT contracts for SCRM clauses.\n"
                    "★ Automate supplier risk scoring and surface in executive dashboards."
                ),
                "examples": [
                    "U.S. Federal Government: EO 14028 mandated enterprise-wide software "
                    "supply chain security, requiring agencies to obtain SBOMs and use only "
                    "approved software from CISA's catalog—a classic Tier-1 activity.",
                    "Enterprise procurement policy requires security clauses in all ICT "
                    "contracts, including SBOM delivery and incident-notification SLAs within 72 hours.",
                    "Corporate architecture board mandates supplier risk assessments during "
                    "all design reviews for systems handling sensitive data.",
                ],
                "mermaid": """
flowchart TD
    P[Corporate Strategy & Legal] --> Q[Draft SCRM Policy]
    Q --> R[Map Mission Requirements to SCRM]
    R --> S[Embed in Procurement Contracts]
    S --> T[Embed in Architecture Reviews]
    T --> U[Prohibited Vendor Registry]
    U --> V[Threat Intel Sharing - ISAC]
    V --> W[Annual Policy Review]
    W --> Q
""",
                "ascii_diagram": """
  Corporate Strategy
         │
         ▼
  ┌─────────────────────────────────────┐
  │  Tier-1 Activity Cycle              │
  │                                     │
  │  Draft Policy ──► Map Requirements  │
  │       ▲                │            │
  │       │         Embed in Contracts  │
  │  Annual Review         │            │
  │       ▲         Embed in Arch Reviews│
  │       │                │            │
  │  Threat Intel ◄── Prohibited Vendor │
  └─────────────────────────────────────┘
"""
            },
            "Tier 2 – Mission Stakeholders": {
                "text": (
                    "Business Management (program management, R&D, engineering, acquisitions, "
                    "cost accounting, quality, safety, security officers) own mission/business "
                    "processes at Tier 2.\n\n"
                    "They translate enterprise SCRM policy into program-level requirements and "
                    "manage trust relationships with system integrators, sub-tier suppliers, and "
                    "managed service providers.\n\n"
                    "Key responsibilities:\n"
                    "- Maintain a program-level supplier risk register.\n"
                    "- Define security and resilience criteria for supplier selection.\n"
                    "- Manage ongoing supplier performance reviews and audits.\n"
                    "- Coordinate incident response with affected suppliers.\n\n"
                    "Recommendations:\n"
                    "★ Implement Supplier Security Scorecards reviewed quarterly.\n"
                    "★ Use structured supplier questionnaires aligned to ISO 27036 / NIST SP 800-161.\n"
                    "★ Conduct on-site or virtual supplier audits for high-criticality components.\n"
                    "★ Build contractual rights to audit into every critical supplier agreement."
                ),
                "examples": [
                    "Target (2013 breach): The HVAC contractor (Fazio Mechanical) had weak "
                    "credential practices. A Tier-2 mission-level supplier security scorecard "
                    "and mandatory MFA for remote vendors could have prevented the breach.",
                    "Program manager for a critical logistics system defines supplier onboarding "
                    "criteria: ISO 27001 certification, annual pen tests, and right-to-audit clauses.",
                    "Engineering lead requires code escrow and secure development lifecycle "
                    "attestation for all third-party software components in safety-critical systems.",
                ],
                "mermaid": """
flowchart TD
    A1[Mission Owner / Program Manager] --> B1[Define Program SCRM Requirements]
    B1 --> C1[Supplier Selection Criteria]
    C1 --> D1[Onboard with Security Scorecard]
    D1 --> E1[Contractual SCRM Clauses]
    E1 --> F1[Quarterly Supplier Reviews]
    F1 --> G1{Issues Found?}
    G1 -- No --> F1
    G1 -- Yes --> H1[Remediation Plan / Escalate to Tier 1]
    H1 --> F1
""",
                "ascii_diagram": """
  ┌───────────────────────────────────────┐
  │         TIER 2 – MISSION LEVEL        │
  │                                       │
  │  Program Mgr ──► Define Requirements  │
  │       │                │              │
  │  Supplier              │              │
  │  Scorecard ◄── Select Suppliers       │
  │       │                               │
  │  Contracts ──► Audits ──► Reviews     │
  │       │                  │            │
  │  Escalate to Tier 1 ◄── Issues        │
  └───────────────────────────────────────┘
"""
            },
            "Tier 2 – Mission Activities": {
                "text": (
                    "Tier-2 activities focus on operationalizing SCRM within mission/business processes.\n\n"
                    "- Define risk response strategies for critical supply chain scenarios.\n"
                    "- Maintain supplier contingency and business continuity plans.\n"
                    "- Integrate SCRM into enterprise architecture and program governance.\n"
                    "- Establish SCRM checkpoints in change management workflows.\n"
                    "- Manage SBOM intake, review, and lifecycle tracking.\n\n"
                    "Recommendations:\n"
                    "★ Map all critical mission dependencies to their Nth-tier suppliers.\n"
                    "★ Maintain alternate supplier lists for all single-source components.\n"
                    "★ Integrate SBOM review into change advisory board (CAB) workflows.\n"
                    "★ Define supplier disruption playbooks for top-5 critical suppliers."
                ),
                "examples": [
                    "COVID-19 pandemic (2020–2022): Semiconductor shortage exposed single-source "
                    "dependencies. Companies with Tier-2 alternate supplier lists (e.g., Toyota's "
                    "Kanban buffer strategy) recovered 3x faster than those without.",
                    "Mission owner defines a contingency plan for logistics software supplier "
                    "outage: pre-approved alternate vendor, manual fallback procedures, and "
                    "a 4-hour RTO.",
                    "Business process team adds supplier risk checkpoints to change management "
                    "workflows, requiring SBOM delta review for any third-party library update.",
                ],
                "mermaid": """
flowchart TD
    M[Identify Mission-Critical Dependencies] --> N[Map Nth-Tier Supplier Chain]
    N --> O[Define SCRM Requirements per Dependency]
    O --> P[Build Alternate Supplier List]
    P --> Q[Integrate SBOM into Change Mgmt]
    Q --> R[Supplier Disruption Playbooks]
    R --> S[Annual BCP Test with Suppliers]
    S --> M
""",
                "ascii_diagram": """
  Mission Process
       │
       ▼
  Map Critical Dependencies ──► Nth-Tier Mapping
       │
       ▼
  Alternate Suppliers ──► Disruption Playbooks
       │
       ▼
  SBOM in Change Mgmt ──► Annual BCP Test
       │
       └──────────────► (loop back)
"""
            },
            "Tier 3 – Information Systems Stakeholders": {
                "text": (
                    "Systems Management (architects, developers, system owners, QA/QC, testers, "
                    "contracting personnel, maintenance and disposal teams) implement SCRM at the "
                    "system level.\n\n"
                    "They decide how to acquire, integrate, operate, and retire ICT components "
                    "in line with SCRM policies cascaded from Tiers 1 and 2.\n\n"
                    "Key responsibilities:\n"
                    "- Maintain a hardware and software asset inventory with supplier metadata.\n"
                    "- Apply software composition analysis (SCA) to identify vulnerable libraries.\n"
                    "- Enforce trusted source requirements for all components.\n"
                    "- Perform supply-chain-aware security testing (hardware integrity, firmware).\n\n"
                    "Recommendations:\n"
                    "★ Use SCA tools (e.g., Sonatype, Snyk, FOSSA) in every CI/CD pipeline.\n"
                    "★ Maintain a live SBOM for every production system.\n"
                    "★ Enforce 'golden image' builds with locked dependency versions.\n"
                    "★ Require tamper-evident packaging and chain-of-custody for hardware."
                ),
                "examples": [
                    "Log4Shell (2021): Organizations with live SBOMs identified vulnerable "
                    "systems in hours vs. weeks for those without. Tier-3 software composition "
                    "analysis was the decisive differentiator.",
                    "System architect specifies hardware only from vetted suppliers with secure "
                    "manufacturing attestations, anti-counterfeiting markings, and tamper-evident seals.",
                    "QA team adds supply-chain-focused test cases: firmware integrity checks, "
                    "signature verification, and hash comparison against vendor-published values.",
                ],
                "mermaid": """
flowchart TD
    S1[System Architect] --> T1[Define Trusted Component Sources]
    T1 --> U1[Apply SCRM Requirements in Design]
    U1 --> V1[SCA in CI/CD Pipeline]
    V1 --> W1[Hardware Integrity Testing]
    W1 --> X1[Maintain Live SBOM]
    X1 --> Y1[Runtime Monitoring]
    Y1 --> Z1[Secure Disposal & Sanitization]
""",
                "ascii_diagram": """
  ┌──────────────────────────────────────────┐
  │       TIER 3 – SYSTEM LEVEL              │
  │                                          │
  │  Architect ──► Trusted Sources           │
  │      │                │                  │
  │  Developers           │                  │
  │  (SCA in CI/CD) ◄─────┘                  │
  │      │                                   │
  │  QA / Testing ──► Firmware Checks        │
  │      │                                   │
  │  Live SBOM ──► Runtime Monitor           │
  │      │                                   │
  │  Secure Disposal ◄── End-of-Life         │
  └──────────────────────────────────────────┘
"""
            },
            "Tier 3 – Information Systems Activities": {
                "text": (
                    "Tier-3 activities integrate SCRM into every phase of the SDLC:\n\n"
                    "- Apply supply-chain security controls during requirements, design, "
                    "development, testing, deployment, operations, and disposal.\n"
                    "- Secure build pipelines: signed artifacts, reproducible builds, "
                    "hermetic build environments.\n"
                    "- Enforce runtime integrity: attestation, allow-lists, anomaly detection.\n"
                    "- Maintain disposal procedures that include supplier notification, "
                    "secure data wipe (NIST 800-88), and hardware destruction records.\n\n"
                    "Recommendations:\n"
                    "★ Adopt a 'Secure by Design' SDLC with supply-chain gates at each phase.\n"
                    "★ Implement signed commits, signed builds, and signed deployments (sigstore).\n"
                    "★ Automate SBOM generation and diff analysis on every release.\n"
                    "★ Define a supplier-aware asset retirement checklist to prevent data leakage."
                ),
                "examples": [
                    "3CX Supply Chain Attack (2023): A compromised upstream library "
                    "(trading app) injected malware into 3CX's build pipeline. Hermetic "
                    "builds and artifact signing would have flagged the unexpected dependency.",
                    "DevOps team enforces signed artifacts and provenance tracking in CI/CD, "
                    "rejecting any build that cannot produce a valid SLSA Level 3 provenance.",
                    "Operations team maintains a supplier-tagged asset inventory; during "
                    "Tier-1 recall of a vendor's NIC, all affected assets were located in "
                    "under 2 hours.",
                ],
                "mermaid": """
flowchart TD
    SDLC[Secure SDLC Entry] --> Req[Requirements: SCRM Criteria]
    Req --> Design[Design: Trusted Architecture]
    Design --> Dev[Dev: SCA + Signed Commits]
    Dev --> Test[Test: Supply-Chain Security Tests]
    Test --> Deploy[Deploy: Signed Artifacts + SLSA]
    Deploy --> Ops[Operations: Runtime Attestation]
    Ops --> Retire[Disposal: NIST 800-88 + Records]
    Retire --> Audit[Post-Retirement Audit]
""",
                "ascii_diagram": """
  SECURE SDLC SUPPLY-CHAIN GATES

  [Requirements] ──SCRM criteria──►
  [Design]       ──Trusted arch────►
  [Development]  ──SCA + Signing───►
  [Testing]      ──SC Test cases───►
  [Deployment]   ──Signed artifacts►
  [Operations]   ──Attestation─────►
  [Disposal]     ──NIST 800-88─────►  Audit Trail
"""
            },
        },
    },
    2: {
        "name": "Supply Chain Threat Agents",
        "description": (
            "Threat agents are entities—intentional or unintentional—capable of compromising "
            "the supply chain. NIST SP 800-161r1 categorizes them by capability, intent, and "
            "access. Understanding WHO poses a threat is the first step in targeted mitigation. "
            "Threat agents operate across geopolitical, organizational, and technical dimensions."
        ),
        "subcategories": {
            "External Adversaries": {
                "text": (
                    "Nation-states, organized crime syndicates, hacktivists, industrial spies, "
                    "and competitors may target ICT supply chains to:\n"
                    "- Insert malicious components or backdoors.\n"
                    "- Steal intellectual property or design files.\n"
                    "- Disrupt production or logistics.\n"
                    "- Establish persistent footholds via trusted supplier channels.\n\n"
                    "Capability levels range from script kiddies to advanced persistent threats "
                    "(APTs) with nation-state resources.\n\n"
                    "Recommendations:\n"
                    "★ Subscribe to CISA, NSA, and sector ISAC threat feeds for supplier IOCs.\n"
                    "★ Conduct hardware authenticity testing on high-risk components.\n"
                    "★ Require suppliers to disclose any government investigations or ownership.\n"
                    "★ Vet country-of-origin for components in national security systems."
                ),
                "examples": [
                    "SUNBURST/SolarWinds (2020 – Russian SVR/APT29): Nation-state inserted "
                    "malicious DLL into SolarWinds Orion build pipeline. ~18,000 organizations "
                    "received the backdoored update. Countermeasure: build integrity monitoring.",
                    "SuperMicro Hardware Implant Allegations (2018): Reports alleged tiny "
                    "chips inserted during manufacturing in China. Countermeasure: independent "
                    "hardware validation labs and country-of-origin procurement controls.",
                    "A criminal group tampers with shipping manifests to divert GPU shipments "
                    "to gray markets, then sells counterfeits into the legitimate channel.",
                ],
                "mermaid": """
flowchart TD
    A[External Adversary] --> B1[Nation-State APT]
    A --> B2[Organized Crime]
    A --> B3[Hacktivist]
    B1 --> C[Compromise Build Pipeline]
    B2 --> D[Counterfeit / Gray-Market Parts]
    B3 --> E[DDoS Supplier Infrastructure]
    C --> F[Malicious Update to Customers]
    D --> G[Defective Components in Production]
    E --> H[Supply Disruption]
""",
                "ascii_diagram": """
  THREAT AGENT TAXONOMY

  External Adversaries
  ├── Nation-State APT  ──► Build Pipeline Compromise
  │                     ──► Hardware Implants
  ├── Organized Crime   ──► Counterfeiting / Diversion
  │                     ──► Ransomware on Suppliers
  └── Hacktivists       ──► Supplier DDoS / Defacement

  Insider Threats
  ├── Malicious Employee──► IP Theft / Sabotage
  └── Negligent Staff   ──► Misconfiguration / Data Leak

  Natural / Accidental
  ├── Natural Disaster  ──► Fab / Logistics Disruption
  └── SW Bugs in Tools  ──► Unintentional Backdoors
"""
            },
            "Insiders and Partners": {
                "text": (
                    "Employees, contractors, partner staff, and managed service providers can "
                    "abuse access—intentionally or accidentally—to introduce supply-chain risk.\n\n"
                    "Types:\n"
                    "- Malicious insiders: disgruntled employees, bribed staff, planted operatives.\n"
                    "- Negligent insiders: accidental data exposure, misconfigurations, phishing victims.\n"
                    "- Compromised partners: third parties with elevated access that are themselves breached.\n\n"
                    "Recommendations:\n"
                    "★ Apply zero-trust principles to all third-party remote access.\n"
                    "★ Enforce MFA and privileged access management (PAM) for all suppliers.\n"
                    "★ Conduct background checks on personnel with access to critical systems.\n"
                    "★ Monitor for anomalous data exfiltration by contractors and partners.\n"
                    "★ Offboard supplier access immediately upon contract termination."
                ),
                "examples": [
                    "Target Breach (2013): HVAC contractor Fazio Mechanical's stolen credentials "
                    "allowed attackers to pivot to POS systems. Countermeasure: network segmentation "
                    "and MFA for all third-party remote access.",
                    "Tesla Insider (2020): Former employee exfiltrated gigabytes of proprietary "
                    "data to competitors. Countermeasure: DLP tools and anomaly-based user "
                    "behavior analytics (UBA).",
                    "A partner misconfigures a shared S3 bucket, exposing firmware source code "
                    "for a medical device to the public internet.",
                ],
                "mermaid": """
flowchart TD
    I[Insider / Partner Threat] --> J1[Malicious Insider]
    I --> J2[Negligent Insider]
    I --> J3[Compromised Partner]
    J1 --> K1[Exfiltrate IP / Sabotage Build]
    J2 --> K2[Misconfigure Access / Fall for Phishing]
    J3 --> K3[Pivot from Partner to Core Systems]
    K1 --> L[Detect via UBA / DLP]
    K2 --> L
    K3 --> L
    L --> M[Respond & Remediate]
""",
                "ascii_diagram": """
  INSIDER / PARTNER RISK MODEL

  ┌─────────────────┐      ┌──────────────────┐
  │ Malicious       │      │  Negligent        │
  │ Insider         │      │  Insider          │
  │ (IP Theft,      │      │  (Misconfiguration│
  │  Sabotage)      │      │   Phishing)       │
  └────────┬────────┘      └────────┬──────────┘
           │                        │
           └──────────┬─────────────┘
                      ▼
            ┌──────────────────┐
            │  Compromised     │
            │  Partner Access  │
            └────────┬─────────┘
                     ▼
            Core Systems at Risk
                     │
            ┌────────▼─────────┐
            │ Detection:        │
            │ UBA + DLP + PAM  │
            └──────────────────┘
"""
            },
        },
    },
    3: {
        "name": "Supply Chain Threat Considerations",
        "description": (
            "Threat considerations answer WHERE and HOW threats manifest across the supply "
            "chain lifecycle. A threat agent (WHO) uses a vector (HOW) to exploit a "
            "vulnerability at a specific lifecycle stage (WHERE). Understanding these "
            "dimensions allows targeted countermeasure design."
        ),
        "subcategories": {
            "Lifecycle Stages": {
                "text": (
                    "Threats can emerge at any stage of the ICT component lifecycle:\n\n"
                    "1. Design: IP theft, malicious specification changes, hardware Trojans inserted in EDA tools.\n"
                    "2. Manufacturing: counterfeit parts, unauthorized clones, tampered firmware in fab.\n"
                    "3. Integration: malicious code injection during system integration, insecure defaults.\n"
                    "4. Distribution / Logistics: diversion, tampering in transit, gray-market substitution.\n"
                    "5. Operation: exploitation of backdoors, unauthorized updates, persistent malware.\n"
                    "6. Maintenance: malicious patch insertion, unauthorized remote access.\n"
                    "7. Retirement / Disposal: data recovery from improperly wiped hardware, component resale.\n\n"
                    "Recommendations:\n"
                    "★ Apply lifecycle-stage-specific security controls and checkpoints.\n"
                    "★ Require supplier lifecycle traceability documentation for critical components.\n"
                    "★ Use hardware security modules (HSMs) and TPMs to anchor trust across stages.\n"
                    "★ Conduct post-disposal audits to confirm secure destruction."
                ),
                "examples": [
                    "Design stage: ITAR-controlled chip design files stolen via insecure CAD "
                    "collaboration platform; PCB layout modified to add hidden debug ports.",
                    "Manufacturing: Supermicro server motherboards allegedly had counterfeit "
                    "chips mixed into production batches at contract manufacturers.",
                    "Logistics: In 2010, counterfeit Cisco routers were seized by US authorities "
                    "after being diverted through unauthorized resellers.",
                    "Disposal: NSA documents revealed adversaries recovered data from "
                    "improperly disposed drives sold in secondary markets.",
                ],
                "mermaid": """
flowchart LR
    D[Design] -->|IP theft risk| M[Manufacturing]
    M -->|Counterfeit risk| I[Integration]
    I -->|Code injection risk| L[Logistics]
    L -->|Diversion / tampering| O[Operation]
    O -->|Backdoor exploitation| Maint[Maintenance]
    Maint -->|Malicious patch| R[Retirement]
    R -->|Data recovery risk| Audit[Audit & Record]
""",
                "ascii_diagram": """
  ICT COMPONENT LIFECYCLE THREAT MAP

  [Design]──────► [Manufacturing]──────► [Integration]
      │                 │                      │
  IP Theft          Counterfeit            Code Inject
      │                 │                      │
  [Logistics]────► [Operation]──────► [Maintenance]
      │                 │                      │
  Diversion         Backdoor             Malicious Patch
      │
  [Retirement/Disposal]
      │
  Data Recovery Risk
"""
            },
            "Attack Vectors": {
                "text": (
                    "Attack vectors are the pathways adversaries use to introduce threats:\n\n"
                    "1. Compromised build tooling: CI/CD pipeline, IDEs, compilers.\n"
                    "2. Insecure source repositories: public or private repos with weak access control.\n"
                    "3. Malicious open-source dependencies: typosquatting, dependency confusion, repo hijack.\n"
                    "4. Weak contractual controls: no SBOM requirement, no incident notification SLA.\n"
                    "5. Physical: tampered hardware in transit, counterfeit components.\n"
                    "6. Social engineering: phishing supplier staff to harvest credentials.\n\n"
                    "Recommendations:\n"
                    "★ Scan all open-source dependencies for known CVEs before use.\n"
                    "★ Pin dependency versions and verify checksums in lock files.\n"
                    "★ Harden CI/CD: least-privilege service accounts, signed pipelines.\n"
                    "★ Conduct supplier social engineering awareness training annually."
                ),
                "examples": [
                    "Dependency Confusion (2021 – Alex Birsan): Researcher uploaded packages "
                    "to PyPI/npm with same names as internal packages of major companies; "
                    "pip/npm preferred the public version, executing malicious code on build servers.",
                    "Codecov bash uploader compromise (2021): Attacker modified the uploader "
                    "script hosted on Codecov's CDN; CI pipelines downloaded and executed it, "
                    "exfiltrating environment variables (including secrets).",
                    "Third-party libraries pulled from PyPI without hash verification were "
                    "found to contain crypto-miner payloads in a financial firm's ML pipeline.",
                ],
                "mermaid": """
flowchart TD
    AV[Attack Vector Entry] --> T1[Compromised Build Tools]
    AV --> T2[Malicious OSS Dependency]
    AV --> T3[Insecure Repository]
    AV --> T4[Physical Tampering]
    AV --> T5[Social Engineering]
    T1 --> E1[Malicious Binary Delivered]
    T2 --> E2[Dependency Confusion / Typosquatting]
    T3 --> E3[Unauthorized Code Merge]
    T4 --> E4[Counterfeit Component in Production]
    T5 --> E5[Credential Theft → Lateral Move]
""",
                "ascii_diagram": """
  ATTACK VECTOR OVERVIEW

  ┌────────────────────────────────────────┐
  │           ATTACK SURFACE               │
  │                                        │
  │  OSS Repo ──► Dependency Confusion     │
  │  CI/CD    ──► Pipeline Compromise      │
  │  Physical ──► Counterfeit / Tamper     │
  │  Social   ──► Phishing Supplier Staff  │
  │  Contract ──► No SBOM / No SLA        │
  └────────────────────────────────────────┘
           │ all paths lead to ►
  ┌────────▼────────────────────────────────┐
  │     Compromised Product / System        │
  └─────────────────────────────────────────┘
"""
            },
        },
    },
    4: {
        "name": "Supply Chain Vulnerability Considerations",
        "description": (
            "Vulnerabilities are weaknesses—in processes, technology, or people—that enable "
            "threat agents to succeed. Unlike threats (external), vulnerabilities are internal "
            "gaps that organizations can directly address. NIST SP 800-161r1 categorizes them "
            "across organizational, process, and technical dimensions."
        ),
        "subcategories": {
            "Process Vulnerabilities": {
                "text": (
                    "Gaps in organizational processes that enable supply-chain attacks:\n\n"
                    "- No formal supplier risk assessment or onboarding checklist.\n"
                    "- Absence of SBOM requirements in contracts.\n"
                    "- Weak change control: security bypassed for urgent hotfixes.\n"
                    "- Lack of a supplier incident response plan.\n"
                    "- No periodic supplier re-assessment after initial onboarding.\n"
                    "- Missing offboarding process: ex-supplier retains access post-contract.\n\n"
                    "Recommendations:\n"
                    "★ Implement a formal Supplier Security Assessment Program (SSAP).\n"
                    "★ Require SBOM delivery at contract award and on each major release.\n"
                    "★ Add supply-chain security gate to every change advisory board (CAB).\n"
                    "★ Conduct annual re-assessments of all Tier-1 and Tier-2 suppliers.\n"
                    "★ Automate supplier access revocation upon contract end."
                ),
                "examples": [
                    "Change Healthcare (2024 ransomware): No MFA on a Citrix remote access "
                    "portal used by a third-party vendor—a process gap in vendor access management "
                    "policy. Impact: $1B+ losses; healthcare claims disrupted nationwide.",
                    "No formal supplier risk assessment before onboarding a new hardware vendor "
                    "led to procurement of counterfeit memory modules, causing data corruption "
                    "in a financial trading system.",
                    "Emergency change bypasses security review for a 'critical hotfix'; the patch "
                    "introduced a hard-coded credential that persisted for 18 months.",
                ],
                "mermaid": """
flowchart TD
    PV[Process Vulnerability] --> A1[No Supplier Assessment]
    PV --> A2[No SBOM Requirement]
    PV --> A3[Weak Change Control]
    PV --> A4[No Incident Response Plan]
    PV --> A5[No Re-assessment Cycle]
    A1 --> B1[Unknown Supplier Risk Profile]
    A2 --> B2[Opaque Software Composition]
    A3 --> B3[Unreviewed Code in Production]
    A4 --> B4[Slow Breach Containment]
    A5 --> B5[Drift in Supplier Posture]
""",
                "ascii_diagram": """
  PROCESS VULNERABILITY MAP

  No Supplier Assessment  ──► Unknown Risk
  No SBOM                 ──► Blind to Vulnerabilities
  Weak Change Control     ──► Unreviewed Code in Prod
  No IR Plan              ──► Slow Breach Response
  No Re-assessment        ──► Supplier Posture Drift
  No Offboarding          ──► Orphaned Access

  All converge to: EXPLOITABLE SUPPLY CHAIN GAP
"""
            },
            "Technical Vulnerabilities": {
                "text": (
                    "Technology-level weaknesses that adversaries can exploit:\n\n"
                    "- Unsigned firmware: devices accept firmware without signature validation.\n"
                    "- Insecure update channels: updates delivered over HTTP without TLS.\n"
                    "- Missing integrity checks: no hash verification on downloaded components.\n"
                    "- Hardcoded credentials in embedded systems.\n"
                    "- Lack of runtime attestation or trusted boot (no TPM).\n"
                    "- Outdated dependencies with known CVEs.\n\n"
                    "Recommendations:\n"
                    "★ Mandate TLS 1.3+ and certificate pinning for all update channels.\n"
                    "★ Implement UEFI Secure Boot and measured boot with TPM 2.0.\n"
                    "★ Integrate CVE scanning (Trivy, Grype) into every CI/CD build.\n"
                    "★ Eliminate all hardcoded credentials; use secrets management (Vault, AWS Secrets Manager).\n"
                    "★ Enforce firmware signing with hardware security modules (HSMs)."
                ),
                "examples": [
                    "ASUS Live Update (2019 – ShadowHammer): Attackers signed malicious firmware "
                    "with ASUS's own stolen certificate. Countermeasure: certificate rotation, "
                    "anomaly-based build signing detection, and binary diff monitoring.",
                    "Mirai Botnet (2016): IoT devices with hardcoded default credentials "
                    "(admin/admin) were mass-compromised. 600,000+ devices formed a DDoS botnet "
                    "that took down Dyn DNS and major websites.",
                    "Devices accepting unsigned firmware over HTTP in a healthcare network allowed "
                    "an attacker to push backdoored firmware to 200+ infusion pumps.",
                ],
                "mermaid": """
flowchart TD
    TV[Technical Vulnerability] --> B1[Unsigned Firmware]
    TV --> B2[Insecure Update Channel]
    TV --> B3[Missing Integrity Checks]
    TV --> B4[Hardcoded Credentials]
    TV --> B5[No Trusted Boot / TPM]
    B1 --> C1[Malicious Firmware Accepted]
    B2 --> C2[MITM Update Injection]
    B3 --> C3[Tampered Binary Executes]
    B4 --> C4[Mass Exploitation at Scale]
    B5 --> C5[Undetected Boot Compromise]
""",
                "ascii_diagram": """
  TECHNICAL VULNERABILITY RISK MATRIX

  Vulnerability            │ Exploitability │ Impact
  ─────────────────────────┼────────────────┼──────────
  Unsigned Firmware        │    HIGH        │  CRITICAL
  Insecure Update Channel  │    HIGH        │  HIGH
  Missing Integrity Checks │    MEDIUM      │  HIGH
  Hardcoded Credentials    │    HIGH        │  CRITICAL
  No Trusted Boot (TPM)    │    MEDIUM      │  HIGH
  Outdated Dependencies    │    HIGH        │  VARIABLE
"""
            },
        },
    },
    5: {
        "name": "Supply Chain Constraints",
        "description": (
            "Constraints are practical limits that shape what SCRM controls are feasible. "
            "They are not excuses to avoid SCRM—they are inputs to risk-based prioritization. "
            "Effective SCRM acknowledges constraints and designs compensating controls."
        ),
        "subcategories": {
            "Business Constraints": {
                "text": (
                    "Cost, schedule, existing contracts, and market dynamics can restrict:\n\n"
                    "- Depth of supplier assessments (time/cost prohibitive for every supplier).\n"
                    "- Ability to switch suppliers (long-lead-time parts, proprietary interfaces).\n"
                    "- Implementation of technical controls (budget, legacy systems).\n"
                    "- Staffing of dedicated SCRM teams.\n\n"
                    "Risk-based approach: Tier suppliers by criticality. Apply deep assessments "
                    "only to Tier-1 critical suppliers; lighter touch for commodity vendors.\n\n"
                    "Recommendations:\n"
                    "★ Build a supplier criticality matrix: rank by impact × replaceability.\n"
                    "★ Phase SCRM controls over 12–24 months aligned to budget cycles.\n"
                    "★ Use shared assessments (Shared Assessments SIG, CAIQ) to reduce cost.\n"
                    "★ Negotiate multi-year contracts to amortize security assessment costs."
                ),
                "examples": [
                    "A defense subcontractor has 400+ suppliers but only budget to deeply "
                    "assess 20/year. Criticality matrix prioritizes the 20 highest-impact "
                    "single-source suppliers for immediate attention.",
                    "Project deadline forces use of existing GPU supplier without full audit; "
                    "compensating control: mandatory hardware authenticity testing on receipt.",
                    "Budget limitations prevent HSM-based firmware signing on all IoT devices; "
                    "compensating control: enhanced runtime anomaly detection and IDPS.",
                ],
                "mermaid": """
flowchart TD
    BC[Business Constraint Identified] --> BM[Build Criticality Matrix]
    BM --> P1{High-Criticality Supplier?}
    P1 -- Yes --> D1[Deep Assessment + Audit Rights]
    P1 -- No --> D2[Shared Assessment / Self-Attestation]
    D1 --> M1[Contractual SCRM Clauses]
    D2 --> M2[Annual Questionnaire]
    M1 --> R1[Ongoing Monitoring]
    M2 --> R1
""",
                "ascii_diagram": """
  SUPPLIER CRITICALITY TRIAGE

  Impact HIGH + Hard to Replace = DEEP ASSESSMENT
  Impact HIGH + Easy to Replace = MODERATE + Alt Sourcing
  Impact LOW  + Hard to Replace = LIGHT + Stockpile
  Impact LOW  + Easy to Replace = QUESTIONNAIRE ONLY
"""
            },
            "Regulatory and Geopolitical Constraints": {
                "text": (
                    "Legal, regulatory, and geopolitical factors constrain sourcing and operations:\n\n"
                    "- Export controls (EAR/ITAR, EU Dual-Use): restrict technology transfer.\n"
                    "- Data residency laws (GDPR, CCPA, India DPDP): mandate local processing.\n"
                    "- Sanctions (OFAC, EU sanctions): prohibit transactions with listed entities.\n"
                    "- Country-of-origin restrictions (NDAA Section 889, Supply Chain Act): ban "
                    "specific vendors in federal procurement (Huawei, ZTE, Hikvision, etc.).\n"
                    "- Tariffs and trade policy: impact cost and availability of components.\n\n"
                    "Recommendations:\n"
                    "★ Screen all suppliers against OFAC SDN and NDAA Section 889 banned lists.\n"
                    "★ Maintain a legal register of applicable regulations updated quarterly.\n"
                    "★ Implement country-of-origin tracking in procurement systems.\n"
                    "★ Engage trade counsel before sourcing from geopolitically sensitive regions."
                ),
                "examples": [
                    "NDAA Section 889 (2019): U.S. federal agencies barred from procuring "
                    "Huawei, ZTE, Hikvision, Dahua, and Hytera equipment. Agencies scrambled "
                    "to audit and replace thousands of installed devices.",
                    "Sanctions on Russian entities post-2022: companies had to rapidly identify "
                    "and remove Russian-origin components and software from supply chains, "
                    "revealing opaque dependencies in legacy systems.",
                    "GDPR data residency requirements forced a U.S. SaaS company to redesign "
                    "its cloud architecture to keep EU customer data within EU borders, "
                    "eliminating certain low-cost Asian cloud providers from the supply chain.",
                ],
                "mermaid": """
flowchart TD
    RG[Regulatory / Geopolitical Input] --> L1[Export Controls - EAR/ITAR]
    RG --> L2[Sanctions - OFAC / EU]
    RG --> L3[NDAA Section 889 Bans]
    RG --> L4[Data Residency Laws]
    L1 --> A1[Restrict Technology Transfer]
    L2 --> A2[Prohibit Specific Transactions]
    L3 --> A3[Remove Banned Vendor Equipment]
    L4 --> A4[Redesign Cloud Architecture]
    A1 --> Comply[Compliance Action Plan]
    A2 --> Comply
    A3 --> Comply
    A4 --> Comply
""",
                "ascii_diagram": """
  REGULATORY CONSTRAINT LANDSCAPE

  ┌──────────────────────────────────────────────────┐
  │  EXPORT CONTROLS    │  SANCTIONS     │  TRADE     │
  │  EAR / ITAR         │  OFAC SDN List │  Tariffs   │
  ├──────────────────────────────────────────────────┤
  │  NDAA §889 Bans     │  Data Residency│  CMMC      │
  │  Huawei/ZTE/etc.    │  GDPR / CCPA  │  FedRAMP   │
  └──────────────────────────────────────────────────┘
  All constrain supplier choice, architecture, and operations
"""
            },
        },
    },
    6: {
        "name": "Supply Chain Vulnerabilities Mapped to Organizations",
        "description": (
            "This category systematically maps identified supply-chain vulnerabilities to "
            "specific organizational tiers, business functions, and information systems. "
            "The mapping clarifies ownership, prioritizes remediation, and enables targeted "
            "investment of limited resources."
        ),
        "subcategories": {
            "Tier Mapping": {
                "text": (
                    "Each vulnerability is assigned to the tier that owns it and is best "
                    "positioned to remediate it:\n\n"
                    "- Tier-1 vulnerabilities: missing SCRM policy, no board reporting, "
                    "undefined risk tolerance.\n"
                    "- Tier-2 vulnerabilities: no supplier risk register, missing contractual "
                    "SCRM clauses, no incident response plan with suppliers.\n"
                    "- Tier-3 vulnerabilities: unsigned firmware, unsigned builds, missing SBOM, "
                    "no runtime integrity monitoring.\n\n"
                    "Recommendations:\n"
                    "★ Build a Vulnerability-to-Tier mapping table for all identified gaps.\n"
                    "★ Assign a named owner and remediation timeline to each vulnerability.\n"
                    "★ Track remediation progress in a centralized SCRM risk register.\n"
                    "★ Report unmitigated Tier-1 gaps directly to executive leadership."
                ),
                "examples": [
                    "Tier-1: No enterprise SCRM policy → Owner: CISO → Remediation: Draft "
                    "and approve policy within 90 days.",
                    "Tier-2: Logistics system has no supplier incident response SLA → "
                    "Owner: Program Manager → Remediation: Add SLA clause at next contract renewal.",
                    "Tier-3: ERP system accepts third-party plugins without signature verification "
                    "→ Owner: System Owner → Remediation: Implement plugin signing within 60 days.",
                ],
                "mermaid": """
flowchart TD
    V[Identified Vulnerability] --> C1{Which Tier Owns It?}
    C1 -- Policy/Strategy --> T1[Tier 1: C-Suite / CISO]
    C1 -- Mission/Process --> T2[Tier 2: Program Manager]
    C1 -- System/Tech --> T3[Tier 3: System Owner]
    T1 --> R1[Policy / Governance Fix]
    T2 --> R2[Process / Contract Fix]
    T3 --> R3[Technical Control Fix]
    R1 --> Track[SCRM Risk Register]
    R2 --> Track
    R3 --> Track
    Track --> Review[Quarterly Review]
""",
                "ascii_diagram": """
  VULNERABILITY-TO-TIER MAPPING TABLE

  Vulnerability               │ Tier │ Owner      │ Action
  ────────────────────────────┼──────┼────────────┼──────────────
  No SCRM Policy              │  1   │ CISO       │ Draft Policy
  No Board Reporting          │  1   │ CRO        │ Add KRI Dash
  No Supplier Risk Register   │  2   │ Prog. Mgr  │ Create Register
  No Incident SLA in Contract │  2   │ Legal/PM   │ Amend Contract
  Unsigned Firmware           │  3   │ Sys. Owner │ Enable Signing
  No SBOM                     │  3   │ DevOps     │ Add to Pipeline
  Missing Runtime Monitoring  │  3   │ SecOps     │ Deploy SIEM
"""
            },
            "Organizational Heatmap": {
                "text": (
                    "Risk heatmaps visualize the concentration and severity of supply-chain "
                    "vulnerabilities across departments, systems, and suppliers.\n\n"
                    "Heatmap dimensions typically include:\n"
                    "- Likelihood of exploitation (1–5).\n"
                    "- Business impact if exploited (1–5).\n"
                    "- Current control effectiveness (reduces residual risk).\n"
                    "- Supplier criticality (weight by mission dependency).\n\n"
                    "Recommendations:\n"
                    "★ Generate quarterly heatmaps and share with Tier-1 and Tier-2 owners.\n"
                    "★ Use heatmaps to drive annual SCRM budget allocation decisions.\n"
                    "★ Include supplier heatmap in board risk report.\n"
                    "★ Automate heatmap generation from GRC platform data."
                ),
                "examples": [
                    "Heatmap reveals HIGH risk concentration in legacy ERP system relying on "
                    "a single-source supplier with no audit rights → triggers emergency "
                    "supplier diversification project.",
                    "Supplier heatmap identifies 2 vendors with 3+ security incidents in the "
                    "past year → placed on enhanced monitoring with 30-day remediation deadline.",
                    "Department heatmap shows R&D and supply chain ops have the highest "
                    "concentration of unmitigated Tier-2 vulnerabilities → drives targeted "
                    "training and process improvement investment.",
                ],
                "mermaid": """
flowchart TD
    HM[SCRM Risk Heatmap] --> D1[Department View]
    HM --> D2[System View]
    HM --> D3[Supplier View]
    D1 --> E1[Identify High-Risk Departments]
    D2 --> E2[Identify Vulnerable Systems]
    D3 --> E3[Flag High-Risk Suppliers]
    E1 --> Act[Targeted Remediation Plan]
    E2 --> Act
    E3 --> Act
    Act --> Budget[Budget Allocation]
    Budget --> Track[Risk Register Update]
""",
                "ascii_diagram": """
  SUPPLY CHAIN RISK HEATMAP (Illustrative)

          IMPACT
          5 │  ■  ■  ■  ■  ■
          4 │  □  ■  ■  ■  ■
          3 │  □  □  ■  ■  ■
          2 │  □  □  □  ■  ■
          1 │  □  □  □  □  ■
            └─────────────────
              1  2  3  4  5  LIKELIHOOD

  ■ = High/Critical Risk (immediate action)
  □ = Low/Medium Risk (monitor / schedule)
"""
            },
        },
    },
    7: {
        "name": "SCRM Plan Controls at Tiers 1, 2, & 3",
        "description": (
            "Controls are specific, implementable measures applied at each organizational tier "
            "to reduce supply-chain risk to acceptable levels. Controls should be traceable to "
            "identified vulnerabilities, proportionate to risk, and validated through testing. "
            "NIST SP 800-161r1 Appendix D provides a comprehensive control catalog."
        ),
        "subcategories": {
            "Tier 1 Controls": {
                "text": (
                    "Enterprise-level governance controls:\n\n"
                    "- C-SCRM Policy: formally approved, annually reviewed SCRM policy.\n"
                    "- C-SCRM Governance Body: dedicated SCRM office or committee with cross-functional representation.\n"
                    "- Standardized Supplier Risk Framework: tiered assessment criteria, scoring rubric.\n"
                    "- Prohibited Vendor Registry: maintained and enforced in procurement systems.\n"
                    "- Board-Level Reporting: quarterly supply-chain KRI dashboard.\n"
                    "- SCRM Budget Line: dedicated funding in annual IT security budget.\n"
                    "- Threat Intelligence Program: subscriptions to sector ISACs and government feeds.\n\n"
                    "Recommendations:\n"
                    "★ Align SCRM policy to NIST SP 800-161r1, ISO 28001, and CMMC as applicable.\n"
                    "★ Include SCRM in annual enterprise risk management (ERM) review.\n"
                    "★ Establish a SCRM maturity model and track progress annually.\n"
                    "★ Conduct enterprise-level SCRM tabletop exercises annually."
                ),
                "examples": [
                    "U.S. DoD: CMMC Level 2/3 requirements effectively impose Tier-1 SCRM "
                    "controls on defense industrial base contractors, including formal SCRM "
                    "policies, supplier assessments, and board accountability.",
                    "Quarterly board report includes supply-chain risk metrics: number of "
                    "high-risk suppliers, open findings, remediation rate, and incident count.",
                    "Enterprise SCRM office maintains a unified supplier risk register with "
                    "3,200 active suppliers scored on 45 security attributes.",
                ],
                "mermaid": """
flowchart TD
    C1[Tier 1 Controls] --> P1[C-SCRM Policy & Governance]
    C1 --> P2[Prohibited Vendor Registry]
    C1 --> P3[Supplier Risk Framework]
    C1 --> P4[Board KRI Dashboard]
    C1 --> P5[Threat Intelligence Program]
    C1 --> P6[Annual SCRM Budget Line]
    P1 --> Out1[Policy Published & Enforced]
    P2 --> Out2[Blocked in Procurement System]
    P3 --> Out3[Supplier Scored on Intake]
    P4 --> Out4[Quarterly Board Visibility]
    P5 --> Out5[Proactive Threat Awareness]
    P6 --> Out6[Funded Controls Implementation]
""",
                "ascii_diagram": """
  TIER 1 CONTROL FRAMEWORK

  ┌──────────────────────────────────────────────────┐
  │              ENTERPRISE SCRM GOVERNANCE          │
  │                                                  │
  │  Policy ──────────────────────────────────────► │
  │  Governance Body (SCRM Office) ───────────────► │
  │  Prohibited Vendor Registry ──────────────────► │
  │  Supplier Risk Framework ─────────────────────► │
  │  Board KRI Reporting ─────────────────────────► │
  │  Threat Intelligence ─────────────────────────► │
  │  Dedicated Budget ────────────────────────────► │
  └──────────────────────┬───────────────────────────┘
                         │ cascades to
                    TIER 2 & 3 CONTROLS
"""
            },
            "Tier 2 Controls": {
                "text": (
                    "Mission/process-level controls:\n\n"
                    "- Supplier Onboarding/Offboarding Procedures: structured intake and exit.\n"
                    "- Contractual SCRM Clauses: SBOM requirements, vulnerability disclosure, "
                    "incident notification (<72h), right-to-audit, security standards compliance.\n"
                    "- Mission-Specific Risk Registers: tracked at program level.\n"
                    "- Supplier Contingency Plans: alternate sources, buffer stock, fallback procedures.\n"
                    "- Supplier Performance Reviews: quarterly scorecards with security KPIs.\n"
                    "- SBOM Management: intake, review, and lifecycle tracking of SBOMs.\n\n"
                    "Recommendations:\n"
                    "★ Use standardized contract language (NIST SP 800-161r1 Appendix F).\n"
                    "★ Integrate SBOM review into change advisory board workflows.\n"
                    "★ Define supplier SLAs for vulnerability remediation (e.g., Critical: 15 days).\n"
                    "★ Conduct at least one supplier site visit or virtual audit per critical supplier annually."
                ),
                "examples": [
                    "Microsoft SSPA (Supplier Security and Privacy Assurance): Microsoft requires "
                    "all suppliers with access to customer data to complete annual security "
                    "assessments, including SBOM delivery and penetration testing results.",
                    "Program contracts for a critical infrastructure system require SBOM within "
                    "30 days of contract award, plus incident notification within 24 hours of discovery.",
                    "Mission risk register tracks 47 supply-chain scenarios with likelihood, "
                    "impact, current controls, residual risk, and named owners.",
                ],
                "mermaid": """
flowchart TD
    C2[Tier 2 Controls] --> Onb[Supplier Onboarding Process]
    C2 --> SLA[Contractual SCRM Clauses]
    C2 --> RR[Mission Risk Register]
    C2 --> CP[Contingency Plans]
    C2 --> Perf[Quarterly Supplier Scorecards]
    C2 --> SBOM[SBOM Management]
    Onb --> Gate1[Security Gate at Intake]
    SLA --> Gate2[Legal Enforcement Mechanism]
    RR --> Gate3[Named Owner per Risk]
    CP --> Gate4[Alternate Supplier Activated if Needed]
    Perf --> Gate5[Underperformer Escalation]
    SBOM --> Gate6[Vulnerability Identification]
""",
                "ascii_diagram": """
  TIER 2 CONTROL CHECKLIST

  ☑ Supplier Onboarding Security Gate
  ☑ Contractual SBOM Requirement
  ☑ Incident Notification SLA (< 72h)
  ☑ Right-to-Audit Clause
  ☑ Mission Risk Register (named owners)
  ☑ Alternate Supplier / Contingency Plan
  ☑ Quarterly Security Scorecard Review
  ☑ SBOM Delta Review on Each Release
  ☑ Offboarding / Access Revocation Procedure
"""
            },
            "Tier 3 Controls": {
                "text": (
                    "System-level technical controls:\n\n"
                    "- Secure SDLC: supply-chain security gates at each development phase.\n"
                    "- Integrity Verification: code signing, build artifact signing (SLSA), "
                    "hardware attestation (TPM).\n"
                    "- SBOM Generation: automated SBOM per build, stored with release artifacts.\n"
                    "- Software Composition Analysis (SCA): automated CVE scanning in CI/CD.\n"
                    "- Runtime Monitoring: SIEM, EDR, IDPS with supply-chain-aware rules.\n"
                    "- Secure Decommissioning: NIST 800-88 data sanitization, destruction records.\n"
                    "- Trusted Execution Environment: TEE/HSM for sensitive operations.\n\n"
                    "Recommendations:\n"
                    "★ Target SLSA Level 2+ for all internally developed software.\n"
                    "★ Enforce SBOM generation on every CI/CD build (CycloneDX or SPDX format).\n"
                    "★ Integrate EPSS scoring alongside CVSS for vulnerability prioritization.\n"
                    "★ Automate decommissioning workflow with supplier notification step."
                ),
                "examples": [
                    "Google SLSA Framework: Google's Supply-chain Levels for Software Artifacts "
                    "(SLSA) framework defines 4 levels of build integrity. Google's own production "
                    "systems target SLSA Level 4 (hermetic, reproducible, signed builds).",
                    "Build pipeline verifies SLSA provenance and signatures of all third-party "
                    "libraries; unsigned or unrecognized artifacts are blocked and alerted.",
                    "Decommissioning procedure for servers: NIST 800-88 Clear/Purge/Destroy "
                    "based on data sensitivity, plus supplier RMA notification and destruction certificate.",
                ],
                "mermaid": """
flowchart TD
    C3[Tier 3 Controls] --> SC[Secure SDLC Gates]
    C3 --> IV[Integrity Verification - SLSA + TPM]
    C3 --> SB[Automated SBOM Generation]
    C3 --> SCA[SCA in CI/CD - CVE Scanning]
    C3 --> RM[Runtime Monitoring - SIEM/EDR]
    C3 --> SD[Secure Decommissioning]
    SC --> O1[Block Non-Compliant Builds]
    IV --> O2[Reject Unsigned Artifacts]
    SB --> O3[Visibility into All Dependencies]
    SCA --> O4[Block Known-Vulnerable Libraries]
    RM --> O5[Alert on Anomalous Supplier Activity]
    SD --> O6[Destruction Records + Supplier Notify]
""",
                "ascii_diagram": """
  TIER 3 TECHNICAL CONTROL STACK

  ┌──────────────────────────────────────────────┐
  │  SECURE SDLC PIPELINE                        │
  │                                              │
  │  [Code Commit] → Signed Commit (GPG/SSH)     │
  │       ↓                                      │
  │  [Build]       → Hermetic + SLSA Provenance  │
  │       ↓                                      │
  │  [SCA Scan]    → CVE + EPSS Scoring          │
  │       ↓                                      │
  │  [SBOM Gen]    → CycloneDX / SPDX            │
  │       ↓                                      │
  │  [Sign Artifact] → HSM-backed Signature      │
  │       ↓                                      │
  │  [Deploy]      → Verified Signature + SBOM   │
  │       ↓                                      │
  │  [Runtime]     → Attestation + SIEM Rules    │
  │       ↓                                      │
  │  [Retire]      → NIST 800-88 + Records       │
  └──────────────────────────────────────────────┘
"""
            },
        },
    },
}

# ------------------------------------------------------------------------------------
# Helper functions for export
# ------------------------------------------------------------------------------------

def build_text_export(category_id, subcat_key):
    cat = SCRM_CATEGORIES[category_id]
    sub = cat["subcategories"][subcat_key]
    lines = []
    lines.append(f"SCRM Category {category_id}: {cat['name']}")
    lines.append(f"Subcategory: {subcat_key}")
    lines.append("=" * 60)
    lines.append("Description:")
    lines.append(cat["description"])
    lines.append("")
    lines.append("Details:")
    lines.append(sub["text"])
    lines.append("")
    lines.append("Real-world examples:")
    for ex in sub["examples"]:
        lines.append(f"- {ex}")
    lines.append("")
    lines.append("ASCII Diagram:")
    lines.append(sub.get("ascii_diagram", "N/A"))
    lines.append("")
    lines.append("Mermaid conceptual flow:")
    lines.append(sub["mermaid"])
    lines.append("")
    lines.append("Developed by Randy Singh, Kalsnet (KNet) Consulting Group")
    return "\n".join(lines)


def build_docx_export(category_id, subcat_key):
    if Document is None:
        return None
    cat = SCRM_CATEGORIES[category_id]
    sub = cat["subcategories"][subcat_key]
    doc = Document()
    doc.add_heading(f"SCRM Category {category_id}: {cat['name']}", level=1)
    doc.add_heading(f"Subcategory: {subcat_key}", level=2)
    doc.add_heading("Description", level=3)
    doc.add_paragraph(cat["description"])
    doc.add_heading("Details", level=3)
    doc.add_paragraph(sub["text"])
    doc.add_heading("Real-world Examples", level=3)
    for ex in sub["examples"]:
        doc.add_paragraph(f"• {ex}", style="List Bullet")
    doc.add_heading("ASCII Diagram", level=3)
    doc.add_paragraph(sub.get("ascii_diagram", "N/A"))
    doc.add_heading("Mermaid Conceptual Flow", level=3)
    doc.add_paragraph(sub["mermaid"])
    doc.add_paragraph("\nDeveloped by Randy Singh, Kalsnet (KNet) Consulting Group")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_pdf_export(category_id, subcat_key):
    if canvas is None:
        return None
    cat = SCRM_CATEGORIES[category_id]
    sub = cat["subcategories"][subcat_key]
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    width, height = letter
    y = height - 50

    def draw_line(text, y_pos, size=11):
        if y_pos < 60:
            c.showPage()
            y_pos = height - 50
        c.setFontSize(size)
        c.drawString(40, y_pos, str(text)[:110])
        return y_pos - (size + 4)

    y = draw_line(f"SCRM Category {category_id}: {cat['name']}", y, 14)
    y = draw_line(f"Subcategory: {subcat_key}", y, 12)
    y = draw_line("Developed by Randy Singh, Kalsnet (KNet) Consulting Group", y, 9)
    y = draw_line("", y)
    y = draw_line("Description:", y, 12)
    for line in cat["description"].split("\n"):
        y = draw_line(line, y)
    y = draw_line("", y)
    y = draw_line("Details:", y, 12)
    for line in sub["text"].split("\n"):
        y = draw_line(line, y)
    y = draw_line("", y)
    y = draw_line("Real-world examples:", y, 12)
    for ex in sub["examples"]:
        for line in ex.split("\n"):
            y = draw_line(f"  {line}", y)
    c.showPage()
    c.save()
    buf.seek(0)
    return buf


def build_json_export(category_id, subcat_key):
    cat = SCRM_CATEGORIES[category_id]
    sub = cat["subcategories"][subcat_key]
    payload = {
        "category_id": category_id,
        "category_name": cat["name"],
        "subcategory": subcat_key,
        "description": cat["description"],
        "details": sub["text"],
        "examples": sub["examples"],
        "ascii_diagram": sub.get("ascii_diagram", ""),
        "mermaid": sub["mermaid"],
        "author": "Randy Singh, Kalsnet (KNet) Consulting Group",
        "exported_at": datetime.utcnow().isoformat() + "Z",
    }
    return json.dumps(payload, indent=2)


def generate_synthetic_scrm_data(num_rows: int = 20):
    records = []
    tiers = ["Tier 1", "Tier 2", "Tier 3"]
    risk_levels = ["Low", "Medium", "High", "Critical"]
    vulnerabilities = [
        "Unsigned firmware", "Weak supplier vetting", "No SBOM",
        "Missing incident SLA", "No right-to-audit", "Hardcoded credentials",
        "Insecure update channel", "No runtime monitoring",
    ]
    threat_agents = [
        "External adversary (APT)", "Insider / Partner", "Organized Crime",
        "Nation-State", "Negligent Employee",
    ]
    for i in range(num_rows):
        records.append({
            "id": i + 1,
            "tier": tiers[i % len(tiers)],
            "asset": f"System-{(i % 5) + 1}",
            "supplier": f"Supplier-{(i % 7) + 1}",
            "risk_level": risk_levels[i % len(risk_levels)],
            "vulnerability": vulnerabilities[i % len(vulnerabilities)],
            "threat_agent": threat_agents[i % len(threat_agents)],
            "likelihood": (i % 5) + 1,
            "impact": ((i + 2) % 5) + 1,
        })
    df = pd.DataFrame(records)
    df["risk_score"] = df["likelihood"] * df["impact"]
    return df

# ------------------------------------------------------------------------------------
# Streamlit UI
# ------------------------------------------------------------------------------------

st.set_page_config(
    page_title="SCRM Framework Dashboard",
    layout="wide",
)

# Bold Blue Title + Developer Credit
st.markdown(
    """
    <h1 style='color:#1a5fb4; font-weight:900; margin-bottom:4px;'>
        Supply Chain Risk Management (SCRM) Framework Dashboard
    </h1>
    <p style='color:#555; font-size:14px; margin-top:0; margin-bottom:16px;'>
        Developed by <strong>Randy Singh</strong> &mdash;
        <em>Kalsnet (KNet) Consulting Group</em>
    </p>
    """,
    unsafe_allow_html=True,
)

st.markdown(
    "> **Executive Leadership** (CEO, CIO, COO, CFO, CISO, CTO, etc.) provide risk "
    "executive functions and help ensure that ICT SCRM mitigation strategies are "
    "cost-effective, efficient, and consistent with strategic goals. "
    "This dashboard covers all seven SCRM categories with detailed guidance, "
    "real-world examples, recommendations, and visual diagrams."
)

st.sidebar.header("Navigation")

category_ids = sorted(SCRM_CATEGORIES.keys())
category_labels = [f"{cid}. {SCRM_CATEGORIES[cid]['name']}" for cid in category_ids]
selected_label = st.sidebar.selectbox("Select SCRM Category", category_labels)
selected_category_id = int(selected_label.split(".")[0])
selected_category = SCRM_CATEGORIES[selected_category_id]

subcat_keys = list(selected_category["subcategories"].keys())
selected_subcat_key = st.sidebar.selectbox("Select Subcategory", subcat_keys)

st.sidebar.markdown("---")
st.sidebar.subheader("Exports & Synthetic Data")

# Main layout
col_main, col_side = st.columns([3, 2])

with col_main:
    st.markdown(
        f"<h2 style='color:#1a5fb4;'>{selected_category_id}. {selected_category['name']}</h2>",
        unsafe_allow_html=True,
    )
    st.subheader(selected_subcat_key)

    st.markdown("**Category description:**")
    st.info(selected_category["description"])

    st.markdown("**Detailed framework text:**")
    st.write(selected_category["subcategories"][selected_subcat_key]["text"])

    st.markdown("**Real-world examples:**")
    for ex in selected_category["subcategories"][selected_subcat_key]["examples"]:
        st.markdown(f"- {ex}")

    # ASCII / Text Diagram
    ascii_diag = selected_category["subcategories"][selected_subcat_key].get("ascii_diagram")
    if ascii_diag:
        st.markdown("**Structural / Text Diagram:**")
        st.code(ascii_diag, language="text")

    # Mermaid Diagram
    st.markdown("**Conceptual Flow (Mermaid diagram):**")
    st.markdown(
        f"```mermaid\n{selected_category['subcategories'][selected_subcat_key]['mermaid']}\n```"
    )

with col_side:
    st.subheader("Export current view")

    text_content = build_text_export(selected_category_id, selected_subcat_key)
    st.download_button(
        label="⬇ Download as TXT",
        data=text_content,
        file_name=f"scrm_category_{selected_category_id}_{selected_subcat_key}.txt",
        mime="text/plain",
    )

    json_content = build_json_export(selected_category_id, selected_subcat_key)
    st.download_button(
        label="⬇ Download as JSON",
        data=json_content,
        file_name=f"scrm_category_{selected_category_id}_{selected_subcat_key}.json",
        mime="application/json",
    )

    docx_buf = build_docx_export(selected_category_id, selected_subcat_key)
    if docx_buf is not None:
        st.download_button(
            label="⬇ Download as Word (DOCX)",
            data=docx_buf,
            file_name=f"scrm_category_{selected_category_id}_{selected_subcat_key}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    else:
        st.info("Install `python-docx` to enable Word export.")

    pdf_buf = build_pdf_export(selected_category_id, selected_subcat_key)
    if pdf_buf is not None:
        st.download_button(
            label="⬇ Download as PDF",
            data=pdf_buf,
            file_name=f"scrm_category_{selected_category_id}_{selected_subcat_key}.pdf",
            mime="application/pdf",
        )
    else:
        st.info("Install `reportlab` to enable PDF export.")

    st.markdown("---")
    st.subheader("Synthetic SCRM Data Generator")

    num_rows = st.slider("Number of synthetic records", min_value=5, max_value=100, value=20)
    if st.button("Generate synthetic SCRM dataset"):
        df = generate_synthetic_scrm_data(num_rows)
        st.dataframe(df, use_container_width=True)

        # Risk score bar chart
        st.markdown("**Risk Score by Supplier (synthetic):**")
        supplier_risk = df.groupby("supplier")["risk_score"].mean().sort_values(ascending=False)
        st.bar_chart(supplier_risk)

        st.markdown("**Risk Level Distribution:**")
        risk_dist = df["risk_level"].value_counts()
        st.bar_chart(risk_dist)

        csv_buf = df.to_csv(index=False).encode("utf-8")
        st.download_button(
            label="⬇ Download synthetic data (CSV)",
            data=csv_buf,
            file_name="synthetic_scrm_data.csv",
            mime="text/csv",
        )

        json_buf = df.to_json(orient="records", indent=2).encode("utf-8")
        st.download_button(
            label="⬇ Download synthetic data (JSON)",
            data=json_buf,
            file_name="synthetic_scrm_data.json",
            mime="application/json",
        )

st.markdown("---")
st.caption(
    "SCRM Framework Dashboard · Developed by Randy Singh, Kalsnet (KNet) Consulting Group · "
    "Based on NIST SP 800-161r1 multi-tier supply chain risk management concepts."
)
