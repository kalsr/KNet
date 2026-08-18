


# Taxonomy to Ontology Conversion Studio
# Developed by Randy Singh, Kalsnet (KNet) Consulting

# Single-file Streamlit application. Contains the full conversion engine,
# sample data, synthetic data generator, diagram rendering, export helpers,
# and the UI (tabs, sidebar, synthetic data bar, uploads, exports) all in
# one place.


import io
import random
import textwrap
from datetime import datetime

import pandas as pd
import networkx as nx
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import streamlit as st


# =============================================================================
# SECTION 1: DOMAIN DEFINITIONS
# =============================================================================

DOMAINS = {
    "ecommerce": "E-commerce Product Taxonomy",
    "organization": "Organizational Role Taxonomy",
    "biological": "Biological Species Taxonomy",
    "document": "Document / Subject Taxonomy",
}

# =============================================================================
# SECTION 2: SAMPLE (CURATED) TAXONOMIES
# Small, hand-picked examples shown in every tab before the user uploads
# anything or generates synthetic data.
# =============================================================================

def sample_taxonomy(domain: str) -> pd.DataFrame:
    """Return a small curated Parent/Child taxonomy DataFrame for a domain."""
    if domain == "ecommerce":
        edges = [
            ("Electronics", "Root"),
            ("Computers", "Electronics"),
            ("Laptops", "Computers"),
            ("Desktops", "Computers"),
            ("Mobile Phones", "Electronics"),
            ("Smartphones", "Mobile Phones"),
            ("Feature Phones", "Mobile Phones"),
            ("Clothing", "Root"),
            ("Men", "Clothing"),
            ("Shirts", "Men"),
            ("Pants", "Men"),
            ("Women", "Clothing"),
            ("Dresses", "Women"),
            ("Tops", "Women"),
        ]
    elif domain == "organization":
        edges = [
            ("VP Engineering", "CEO"),
            ("Engineering Manager", "VP Engineering"),
            ("Software Engineer", "Engineering Manager"),
            ("QA Engineer", "Engineering Manager"),
            ("VP Sales", "CEO"),
            ("Sales Manager", "VP Sales"),
            ("Sales Representative", "Sales Manager"),
            ("Account Executive", "Sales Manager"),
        ]
    elif domain == "biological":
        edges = [
            ("Chordata", "Animalia"),
            ("Mammalia", "Chordata"),
            ("Carnivora", "Mammalia"),
            ("Felidae", "Carnivora"),
            ("Panthera", "Felidae"),
            ("Aves", "Chordata"),
            ("Passeriformes", "Aves"),
            ("Corvidae", "Passeriformes"),
        ]
    elif domain == "document":
        edges = [
            ("Science", "Root"),
            ("Physics", "Science"),
            ("Quantum Mechanics", "Physics"),
            ("Classical Mechanics", "Physics"),
            ("Biology", "Science"),
            ("Genetics", "Biology"),
            ("Ecology", "Biology"),
        ]
    else:
        edges = [("Item A", "Root"), ("Item B", "Root")]
    return pd.DataFrame(edges, columns=["Child", "Parent"])


DOMAIN_EXPLANATION = {
    "ecommerce": textwrap.dedent("""
        Taxonomy (before): A simple hierarchy of product categories
        (Electronics, then Computers, then Laptops). It tells you where
        something sits, but nothing about how products relate to brands,
        prices, or each other.

        Ontology (after): Each category becomes a formal Class. We add
        object properties (manufacturedBy, compatibleWith, hasAccessory),
        data properties (hasPrice, hasWarrantyMonths), and concrete
        individuals (actual products) that instantiate the classes.
        This lets a machine answer questions like "show me all Laptops
        manufactured by Dell under 1000 dollars" - which a plain taxonomy
        cannot do.
    """).strip(),
    "organization": textwrap.dedent("""
        Taxonomy (before): An org chart hierarchy (CEO, then VP, then
        Manager, then Individual Contributor). It only encodes reporting
        levels.

        Ontology (after): Roles become Classes. We add object properties
        (reportsTo, manages, collaboratesWith), data properties
        (hasSalaryBand, hasMinExperienceYears), and individual employees
        who hold those roles. This supports reasoning such as "who
        ultimately reports to the CEO through the Engineering chain".
    """).strip(),
    "biological": textwrap.dedent("""
        Taxonomy (before): The classic Linnaean rank hierarchy
        (Kingdom, Phylum, Class, Order, Family, Genus). Purely structural.

        Ontology (after): Ranks become Classes, linked by subClassOf.
        We enrich with object properties (eats, livesIn, preyedOnBy) and
        data properties (hasConservationStatus, hasAvgLifespanYears), plus
        individual species as instances. This enables food-web and habitat
        reasoning that a taxonomy alone cannot express.
    """).strip(),
    "document": textwrap.dedent("""
        Taxonomy (before): A subject classification tree (Science, then
        Physics, then Quantum Mechanics), as used in library catalogs.

        Ontology (after): Topics become Classes. We add object properties
        (prerequisiteOf, citesConcept, relatedTo) and data properties
        (hasDifficultyLevel), plus individual documents and papers tagged
        to topics. This enables recommendation and prerequisite-chain
        reasoning across a content library.
    """).strip(),
}

# =============================================================================
# SECTION 3: DOMAIN SEMANTIC TEMPLATES
# Used to enrich ANY taxonomy - curated sample, synthetic, or user upload.
# =============================================================================

DOMAIN_TEMPLATES = {
    "ecommerce": {
        "object_properties": ["manufacturedBy", "compatibleWith", "hasAccessory"],
        "data_properties": ["hasPrice", "hasWarrantyMonths"],
        "individual_prefix": "Product",
        "individual_object_values": ["Acme Corp", "Global Traders Inc", "NovaTech"],
    },
    "organization": {
        "object_properties": ["reportsTo", "manages", "collaboratesWith"],
        "data_properties": ["hasSalaryBand", "hasMinExperienceYears"],
        "individual_prefix": "Employee",
        "individual_object_values": ["Team Alpha", "Team Beta", "Team Gamma"],
    },
    "biological": {
        "object_properties": ["eats", "livesIn", "preyedOnBy"],
        "data_properties": ["hasConservationStatus", "hasAvgLifespanYears"],
        "individual_prefix": "Species",
        "individual_object_values": ["Savanna", "Rainforest", "Tundra"],
    },
    "document": {
        "object_properties": ["prerequisiteOf", "citesConcept", "relatedTo"],
        "data_properties": ["hasDifficultyLevel", "hasPageCount"],
        "individual_prefix": "Document",
        "individual_object_values": ["Journal A", "Journal B", "Conference C"],
    },
    "generic": {
        "object_properties": ["relatedTo", "partOf"],
        "data_properties": ["hasCode", "hasLevel"],
        "individual_prefix": "Instance",
        "individual_object_values": ["GroupX", "GroupY", "GroupZ"],
    },
}

# =============================================================================
# SECTION 4: SYNTHETIC DATA GENERATOR
# Used by the in-app "Synthetic Data Bar" and to produce the standalone
# 300-record demo CSV.
# =============================================================================

_WORD_BANK = [
    "Alpha", "Beta", "Gamma", "Delta", "Nova", "Orion", "Vertex", "Apex",
    "Summit", "Horizon", "Fusion", "Pulse", "Quantum", "Crystal", "Element",
    "Matrix", "Vector", "Cascade", "Prime", "Nexus", "Zenith", "Cobalt",
    "Titan", "Echo", "Lumen", "Onyx", "Sable", "Ivory", "Coral", "Amber",
]

_TOP_LEVEL_BY_DOMAIN = {
    "ecommerce": ["Electronics", "Apparel", "Home and Garden", "Sports and Outdoors",
                  "Toys and Games", "Beauty and Health", "Automotive", "Groceries",
                  "Books and Media", "Office Supplies", "Pet Supplies", "Furniture"],
    "organization": ["Engineering", "Sales", "Marketing", "Finance", "Human Resources",
                      "Operations", "Customer Support", "Legal", "Product", "IT",
                      "Procurement", "R and D"],
    "biological": ["Mammalia", "Aves", "Reptilia", "Amphibia", "Pisces", "Insecta",
                   "Arachnida", "Mollusca", "Crustacea", "Annelida", "Cnidaria", "Echinodermata"],
    "document": ["Physics", "Biology", "Chemistry", "Mathematics", "Computer Science",
                 "Economics", "Psychology", "History", "Literature", "Engineering",
                 "Medicine", "Philosophy"],
}


def generate_synthetic_taxonomy(domain: str, n_records: int = 300, seed: int = 42) -> pd.DataFrame:
    """
    Generate a synthetic multi-level Parent/Child taxonomy with approximately
    n_records Child-Parent edges, themed for the given domain. Used both by
    the in-app Synthetic Data Bar and to create the standalone demo CSV.
    """
    rng = random.Random(seed)
    tops = _TOP_LEVEL_BY_DOMAIN.get(domain, _TOP_LEVEL_BY_DOMAIN["ecommerce"])

    rows = []
    for top in tops:
        rows.append((top, "Root"))

    remaining = n_records - len(rows)
    if remaining <= 0:
        return pd.DataFrame(rows[:n_records], columns=["Child", "Parent"])

    n_level2 = max(1, remaining // 6)
    n_level3 = remaining - n_level2

    level2_nodes = []
    made_l2 = 0
    while made_l2 < n_level2:
        top = tops[made_l2 % len(tops)]
        word = rng.choice(_WORD_BANK)
        name = f"{top} - {word} Group {made_l2 + 1}"
        rows.append((name, top))
        level2_nodes.append(name)
        made_l2 += 1

    made_l3 = 0
    while made_l3 < n_level3:
        parent = level2_nodes[made_l3 % len(level2_nodes)]
        word1 = rng.choice(_WORD_BANK)
        word2 = rng.choice(_WORD_BANK)
        name = f"{parent} - {word1} {word2} Item {made_l3 + 1}"
        rows.append((name, parent))
        made_l3 += 1

    df = pd.DataFrame(rows, columns=["Child", "Parent"])
    return df.head(n_records).reset_index(drop=True)


# =============================================================================
# SECTION 5: GENERIC CONVERSION ENGINE
# =============================================================================

def build_hierarchy_graph(df: pd.DataFrame) -> nx.DiGraph:
    """Build a directed graph Parent to Child from a taxonomy dataframe."""
    G = nx.DiGraph()
    for _, row in df.iterrows():
        child, parent = str(row["Child"]).strip(), str(row["Parent"]).strip()
        if not child:
            continue
        G.add_node(child)
        if parent:
            G.add_node(parent)
            G.add_edge(parent, child)
    return G


def convert_taxonomy_to_ontology(df: pd.DataFrame, domain: str = "generic", max_individuals: int = 12) -> dict:
    """
    Core conversion routine. Takes a Parent/Child taxonomy dataframe and a
    domain key, and returns a dictionary describing the resulting ontology:
      - classes: list of class names
      - subclass_relations: list of (child_class, parent_class)
      - object_properties: list of (subject, property, object) triples
      - data_properties: list of (subject, property, literal_value)
      - individuals: list of (individual_name, "type", class_name)
      - triples: unified list of (Subject, Predicate, Object, Type) for export
    """
    template = DOMAIN_TEMPLATES.get(domain, DOMAIN_TEMPLATES["generic"])
    G = build_hierarchy_graph(df)

    classes = sorted(set(G.nodes()) - {"Root"})
    subclass_relations = [(c, p) for p, c in G.edges() if p != "Root"]

    leaves = [n for n in classes if G.out_degree(n) == 0]
    rng = random.Random(7)
    obj_props = template["object_properties"]
    data_props = template["data_properties"]
    obj_values = template["individual_object_values"]
    prefix = template["individual_prefix"]

    object_properties = []
    data_properties = []
    individuals = []

    sample_leaves = leaves[:max_individuals] if leaves else classes[:max_individuals]
    for i, leaf in enumerate(sample_leaves):
        ind_name = f"{prefix}_{i + 1}"
        individuals.append((ind_name, "instanceOf", leaf))
        if obj_props:
            op = obj_props[i % len(obj_props)]
            ov = obj_values[i % len(obj_values)]
            object_properties.append((ind_name, op, ov))
        if data_props:
            dp = data_props[i % len(data_props)]
            dv = rng.randint(10, 999)
            data_properties.append((ind_name, dp, dv))

    class_level_object_properties = []
    if len(classes) >= 2 and obj_props:
        for i in range(min(3, len(classes) - 1)):
            class_level_object_properties.append((classes[i], obj_props[0], classes[i + 1]))

    triples = []
    for child, parent in subclass_relations:
        triples.append((child, "subClassOf", parent, "Class Hierarchy"))
    for s, p, o in class_level_object_properties:
        triples.append((s, p, o, "Object Property (Class-level)"))
    for s, p, o in object_properties:
        triples.append((s, p, o, "Object Property (Individual)"))
    for s, p, o in data_properties:
        triples.append((s, p, str(o), "Data Property"))
    for s, p, o in individuals:
        triples.append((s, p, o, "Individual"))

    return {
        "domain": domain,
        "classes": classes,
        "subclass_relations": subclass_relations,
        "object_properties": object_properties + class_level_object_properties,
        "data_properties": data_properties,
        "individuals": individuals,
        "triples": triples,
        "graph": G,
    }


def ontology_to_dataframe(ontology: dict) -> pd.DataFrame:
    return pd.DataFrame(ontology["triples"], columns=["Subject", "Predicate", "Object", "Type"])


# =============================================================================
# SECTION 6: DIAGRAMS AND GRAPHS
# Matplotlib and NetworkX only - no external binaries required.
# =============================================================================

def draw_hierarchy_tree(G: nx.DiGraph, title: str = "Taxonomy Hierarchy (before)"):
    fig, ax = plt.subplots(figsize=(8, 5))
    try:
        pos = nx.nx_agraph.graphviz_layout(G, prog="dot")
    except Exception:
        pos = nx.spring_layout(G, k=0.6, seed=3)
    nx.draw(
        G, pos, ax=ax, with_labels=True, node_color="#dbe9ff",
        edge_color="#1a4fb4", node_size=1200, font_size=7, arrows=True,
        linewidths=1, edgecolors="#1a4fb4",
    )
    ax.set_title(title, fontsize=12, color="#0b2f7a", fontweight="bold")
    fig.tight_layout()
    return fig


def draw_ontology_network(ontology: dict, title: str = "Ontology Network (after)", max_nodes: int = 40):
    G2 = nx.DiGraph()
    color_map = {
        "subClassOf": "#1a4fb4",
        "objprop": "#d97706",
        "dataprop": "#0f9d58",
        "individual": "#a52a2a",
    }
    count = 0
    for child, parent in ontology["subclass_relations"]:
        if count >= max_nodes:
            break
        G2.add_edge(parent, child, kind="subClassOf")
        count += 1
    for s, p, o in ontology["object_properties"][:max_nodes]:
        G2.add_edge(s, o, kind="objprop", label=p)
    for s, p, o in ontology["individuals"][:max_nodes]:
        G2.add_edge(o, s, kind="individual", label=p)

    fig, ax = plt.subplots(figsize=(9, 6))
    pos = nx.spring_layout(G2, k=0.7, seed=5)
    for kind, color in color_map.items():
        edges = [(u, v) for u, v, d in G2.edges(data=True) if d.get("kind") == kind]
        nx.draw_networkx_edges(G2, pos, edgelist=edges, ax=ax, edge_color=color, arrows=True, width=1.4)
    nx.draw_networkx_nodes(G2, pos, ax=ax, node_color="#f5f7ff", edgecolors="#1a4fb4", node_size=900)
    nx.draw_networkx_labels(G2, pos, ax=ax, font_size=7)

    legend_handles = [
        mpatches.Patch(color="#1a4fb4", label="subClassOf"),
        mpatches.Patch(color="#d97706", label="Object Property"),
        mpatches.Patch(color="#a52a2a", label="Individual (instanceOf)"),
    ]
    ax.legend(handles=legend_handles, loc="upper right", fontsize=7)
    ax.set_title(title, fontsize=12, color="#0b2f7a", fontweight="bold")
    ax.axis("off")
    fig.tight_layout()
    return fig


def draw_conversion_flow():
    """A simple boxes-and-arrows flow diagram of the conversion pipeline."""
    steps = [
        "1. Ingest\nTaxonomy\n(Parent/Child)",
        "2. Build\nClass Hierarchy\n(subClassOf)",
        "3. Apply Domain\nSemantic Template\n(Object/Data Props)",
        "4. Generate\nIndividuals\n(Instances)",
        "5. Emit\nOntology\n(Classes, Props, Rules)",
    ]
    fig, ax = plt.subplots(figsize=(11, 2.4))
    n = len(steps)
    box_w, box_h, gap = 1.8, 1.0, 0.55
    x = 0.2
    for i, step in enumerate(steps):
        rect = mpatches.FancyBboxPatch((x, 0.3), box_w, box_h,
                                        boxstyle="round,pad=0.05,rounding_size=0.08",
                                        linewidth=1.5, edgecolor="#1a4fb4",
                                        facecolor="#eaf1ff")
        ax.add_patch(rect)
        ax.text(x + box_w / 2, 0.3 + box_h / 2, step, ha="center", va="center",
                 fontsize=8, color="#0b2f7a", fontweight="bold")
        if i < n - 1:
            ax.annotate("", xy=(x + box_w + gap - 0.1, 0.3 + box_h / 2),
                        xytext=(x + box_w, 0.3 + box_h / 2),
                        arrowprops=dict(arrowstyle="->", color="#0b2f7a", lw=1.6))
        x += box_w + gap
    ax.set_xlim(0, x)
    ax.set_ylim(0, 1.6)
    ax.axis("off")
    fig.tight_layout()
    return fig


# =============================================================================
# SECTION 7: EXPORT HELPERS
# =============================================================================

def export_csv_bytes(df: pd.DataFrame) -> bytes:
    return df.to_csv(index=False).encode("utf-8")


def export_txt_bytes(ontology: dict, domain: str) -> bytes:
    lines = []
    lines.append(f"Ontology generated from taxonomy - domain: {domain}")
    lines.append(f"Generated: {datetime.now().isoformat(timespec='seconds')}")
    lines.append("=" * 70)
    lines.append("\nCLASSES\n")
    for c in ontology["classes"]:
        lines.append(f"Class: {c}")
    lines.append("\nSUBCLASS RELATIONS (Turtle-like)\n")
    for child, parent in ontology["subclass_relations"]:
        lines.append(f":{child} rdfs:subClassOf :{parent} .")
    lines.append("\nOBJECT PROPERTIES\n")
    for s, p, o in ontology["object_properties"]:
        lines.append(f":{s} :{p} :{o} .")
    lines.append("\nDATA PROPERTIES\n")
    for s, p, o in ontology["data_properties"]:
        lines.append(f':{s} :{p} "{o}" .')
    lines.append("\nINDIVIDUALS\n")
    for s, p, o in ontology["individuals"]:
        lines.append(f":{s} rdf:type :{o} .")
    return "\n".join(lines).encode("utf-8")


def export_docx_bytes(ontology: dict, domain: str, explanation: str = "") -> bytes:
    from docx import Document
    from docx.shared import RGBColor

    doc = Document()

    title = doc.add_heading("Taxonomy to Ontology Conversion Report", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x0B, 0x2F, 0x7A)

    doc.add_paragraph(f"Domain / Use case: {DOMAINS.get(domain, domain)}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("Developed by Randy Singh, Kalsnet (KNet) Consulting")

    if explanation:
        doc.add_heading("Explanation", level=1)
        doc.add_paragraph(explanation)

    doc.add_heading("Classes", level=1)
    doc.add_paragraph(", ".join(ontology["classes"]) or "None")

    doc.add_heading("Subclass Relations", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text = "Child Class", "Parent Class"
    for child, parent in ontology["subclass_relations"]:
        row = table.add_row().cells
        row[0].text, row[1].text = str(child), str(parent)

    doc.add_heading("Object Properties", level=1)
    table2 = doc.add_table(rows=1, cols=3)
    table2.style = "Light Grid Accent 1"
    hdr2 = table2.rows[0].cells
    hdr2[0].text, hdr2[1].text, hdr2[2].text = "Subject", "Property", "Object"
    for s, p, o in ontology["object_properties"]:
        row = table2.add_row().cells
        row[0].text, row[1].text, row[2].text = str(s), str(p), str(o)

    doc.add_heading("Data Properties", level=1)
    table3 = doc.add_table(rows=1, cols=3)
    table3.style = "Light Grid Accent 1"
    hdr3 = table3.rows[0].cells
    hdr3[0].text, hdr3[1].text, hdr3[2].text = "Subject", "Property", "Value"
    for s, p, o in ontology["data_properties"]:
        row = table3.add_row().cells
        row[0].text, row[1].text, row[2].text = str(s), str(p), str(o)

    doc.add_heading("Individuals", level=1)
    table4 = doc.add_table(rows=1, cols=2)
    table4.style = "Light Grid Accent 1"
    hdr4 = table4.rows[0].cells
    hdr4[0].text, hdr4[1].text = "Individual", "Type (Class)"
    for s, p, o in ontology["individuals"]:
        row = table4.add_row().cells
        row[0].text, row[1].text = str(s), str(o)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_pdf_bytes(ontology: dict, domain: str, explanation: str = "") -> bytes:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

    buf = io.BytesIO()
    docpdf = SimpleDocTemplate(buf, pagesize=letter, topMargin=40, bottomMargin=40)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("TitleBlue", parent=styles["Title"],
                                  textColor=colors.HexColor("#0B2F7A"))
    h_style = ParagraphStyle("HeadBlue", parent=styles["Heading2"],
                              textColor=colors.HexColor("#1A4FB4"))
    body = styles["BodyText"]

    elems = []
    elems.append(Paragraph("Taxonomy to Ontology Conversion Report", title_style))
    elems.append(Paragraph(f"Domain / Use case: {DOMAINS.get(domain, domain)}", body))
    elems.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", body))
    elems.append(Paragraph("Developed by Randy Singh, Kalsnet (KNet) Consulting", body))
    elems.append(Spacer(1, 12))

    if explanation:
        elems.append(Paragraph("Explanation", h_style))
        for para in explanation.split("\n\n"):
            elems.append(Paragraph(para, body))
        elems.append(Spacer(1, 10))

    def add_table(heading, data, colnames):
        elems.append(Paragraph(heading, h_style))
        rows = [colnames] + [[str(x) for x in r] for r in data]
        if len(rows) == 1:
            rows.append(["-"] * len(colnames))
        t = Table(rows, repeatRows=1)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1A4FB4")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTSIZE", (0, 0), (-1, -1), 7),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EAF1FF")]),
        ]))
        elems.append(t)
        elems.append(Spacer(1, 10))

    add_table("Subclass Relations", ontology["subclass_relations"], ["Child Class", "Parent Class"])
    add_table("Object Properties", ontology["object_properties"], ["Subject", "Property", "Object"])
    add_table("Data Properties", ontology["data_properties"], ["Subject", "Property", "Value"])
    add_table("Individuals", ontology["individuals"], ["Individual", "Type (Class)"])

    docpdf.build(elems)
    return buf.getvalue()


# =============================================================================
# SECTION 8: STREAMLIT UI
# =============================================================================

st.set_page_config(
    page_title="Taxonomy to Ontology Conversion Studio",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.markdown(
    """
    <style>
    .title-bar-app   { font-size: 40px; font-weight: 800; color: #0B2F7A; line-height: 1.15; margin-bottom: 0px;}
    .title-bar-dev   { font-size: 26px; font-weight: 800; color: #0B2F7A; line-height: 1.2; margin-top: 0px;}
    .subtle-caption  { color: #444; font-size: 15px; margin-top: 4px; }
    .stTabs [data-baseweb="tab"] { font-weight: 700; font-size: 16px; }
    .synth-bar { background-color:#EAF1FF; border:1px solid #1A4FB4; border-radius:8px; padding:10px 14px; margin-bottom:14px;}
    </style>
    <div class="title-bar-app">Taxonomy to Ontology Conversion Studio</div>
    <div class="title-bar-dev">Developed by Randy Singh, Kalsnet (KNet) Consulting</div>
    <div class="subtle-caption">Turn flat, hierarchical taxonomies into rich, machine-reasonable ontologies, with live examples, your own data, and instant exports.</div>
    <hr style="border: 1px solid #0B2F7A;">
    """,
    unsafe_allow_html=True,
)

with st.sidebar:
    st.markdown("### About this app")
    st.write(
        "This app demonstrates, across four use cases in one single "
        "application, how a taxonomy (a simple parent/child hierarchy) is "
        "transformed into an ontology (classes, relationships, properties, "
        "individuals, and rules)."
    )
    st.markdown("### What changes, exactly")
    st.markdown(
        "Taxonomy: is-a hierarchy only.\n\n"
        "Ontology adds: object properties (relationships between classes), "
        "data properties (attributes and values), individuals (real "
        "instances), and optionally axioms and rules."
    )
    st.markdown("### Every use case tab lets you")
    st.markdown(
        "1. See a curated example conversion\n"
        "2. Generate synthetic demo data (Synthetic Data Bar)\n"
        "3. Upload your own taxonomy CSV\n"
        "4. View hierarchy and ontology diagrams\n"
        "5. Export results (PDF, Word, TXT, CSV)"
    )
    st.markdown("---")
    st.caption("Randy Singh, Kalsnet (KNet) Consulting")


def render_use_case(domain_key: str, tab_label: str):
    domain_name = DOMAINS[domain_key]
    st.header(domain_name)
    st.markdown(DOMAIN_EXPLANATION[domain_key])

    st.markdown("#### 1. Conversion Pipeline")
    st.pyplot(draw_conversion_flow(), clear_figure=True)

    st.markdown("---")
    st.markdown("#### 2. Built-in Example")
    sample_df = sample_taxonomy(domain_key)
    col_a, col_b = st.columns([1, 1])
    with col_a:
        st.caption("Sample taxonomy (Parent/Child pairs)")
        st.dataframe(sample_df, use_container_width=True, height=260)
    with col_b:
        st.caption("Hierarchy diagram (before)")
        st.pyplot(draw_hierarchy_tree(build_hierarchy_graph(sample_df)), clear_figure=True)

    sample_onto = convert_taxonomy_to_ontology(sample_df, domain_key)
    st.caption("Resulting ontology network (after). Blue lines are subClassOf, orange lines are object properties, red lines are individuals.")
    st.pyplot(draw_ontology_network(sample_onto), clear_figure=True)

    with st.expander("View resulting ontology triples (sample example)"):
        st.dataframe(ontology_to_dataframe(sample_onto), use_container_width=True)

    _download_row(sample_onto, domain_key, "sample", DOMAIN_EXPLANATION[domain_key])

    st.markdown("---")
    st.markdown("#### 3. Synthetic Data Bar - generate demo data instantly")
    st.markdown('<div class="synth-bar">', unsafe_allow_html=True)
    c1, c2, c3 = st.columns([2, 1, 1])
    with c1:
        n_records = st.slider(
            f"Number of synthetic taxonomy records ({tab_label})",
            min_value=20, max_value=500, value=300, step=20,
            key=f"synth_slider_{domain_key}",
        )
    with c2:
        seed = st.number_input("Random seed", min_value=1, max_value=9999, value=42, key=f"seed_{domain_key}")
    with c3:
        generate = st.button("Generate synthetic data", key=f"gen_{domain_key}", use_container_width=True)
    st.markdown('</div>', unsafe_allow_html=True)

    if generate:
        st.session_state[f"synth_df_{domain_key}"] = generate_synthetic_taxonomy(
            domain_key, n_records=n_records, seed=int(seed)
        )

    if f"synth_df_{domain_key}" in st.session_state:
        synth_df = st.session_state[f"synth_df_{domain_key}"]
        st.success(f"Generated {len(synth_df)} synthetic taxonomy records for {domain_name}.")
        st.dataframe(synth_df, use_container_width=True, height=240)
        st.download_button(
            "Download synthetic taxonomy as CSV",
            data=export_csv_bytes(synth_df),
            file_name=f"synthetic_taxonomy_{domain_key}.csv",
            mime="text/csv",
            key=f"dl_synth_csv_{domain_key}",
        )
        synth_onto = convert_taxonomy_to_ontology(synth_df, domain_key, max_individuals=15)
        st.caption("Ontology network generated from synthetic data")
        st.pyplot(draw_ontology_network(synth_onto, title="Ontology from Synthetic Data"), clear_figure=True)
        _download_row(synth_onto, domain_key, "synthetic", DOMAIN_EXPLANATION[domain_key])

    st.markdown("---")
    st.markdown("#### 4. Upload Your Own Taxonomy")
    st.caption("CSV must contain two columns named Child and Parent. Top-level items can use Root as parent.")
    uploaded = st.file_uploader(
        f"Upload taxonomy CSV for {domain_name}", type=["csv"], key=f"upload_{domain_key}"
    )
    if uploaded is not None:
        try:
            user_df = pd.read_csv(uploaded)
            cols = {c.lower(): c for c in user_df.columns}
            if "child" not in cols or "parent" not in cols:
                st.error("CSV must have Child and Parent columns. Please check your file and re-upload.")
            else:
                user_df = user_df.rename(columns={cols["child"]: "Child", cols["parent"]: "Parent"})[["Child", "Parent"]]
                st.success(f"Loaded {len(user_df)} rows from your file.")
                col_c, col_d = st.columns([1, 1])
                with col_c:
                    st.caption("Your taxonomy")
                    st.dataframe(user_df, use_container_width=True, height=240)
                with col_d:
                    st.caption("Your hierarchy diagram")
                    st.pyplot(draw_hierarchy_tree(build_hierarchy_graph(user_df)), clear_figure=True)

                user_onto = convert_taxonomy_to_ontology(user_df, domain_key, max_individuals=15)
                st.caption("Your ontology network")
                st.pyplot(draw_ontology_network(user_onto, title="Your Ontology"), clear_figure=True)
                with st.expander("View resulting ontology triples (your data)"):
                    st.dataframe(ontology_to_dataframe(user_onto), use_container_width=True)
                _download_row(user_onto, domain_key, "your_data", DOMAIN_EXPLANATION[domain_key])
        except Exception as e:
            st.error(f"Could not process file: {e}")


def _download_row(ontology: dict, domain_key: str, tag: str, explanation: str):
    st.markdown("**Export this result:**")
    df_export = ontology_to_dataframe(ontology)
    c1, c2, c3, c4 = st.columns(4)
    with c1:
        st.download_button(
            "Download CSV", data=export_csv_bytes(df_export),
            file_name=f"ontology_{domain_key}_{tag}.csv", mime="text/csv",
            key=f"csv_{domain_key}_{tag}",
        )
    with c2:
        st.download_button(
            "Download TXT", data=export_txt_bytes(ontology, domain_key),
            file_name=f"ontology_{domain_key}_{tag}.txt", mime="text/plain",
            key=f"txt_{domain_key}_{tag}",
        )
    with c3:
        st.download_button(
            "Download Word (docx)", data=export_docx_bytes(ontology, domain_key, explanation),
            file_name=f"ontology_{domain_key}_{tag}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=f"docx_{domain_key}_{tag}",
        )
    with c4:
        st.download_button(
            "Download PDF", data=export_pdf_bytes(ontology, domain_key, explanation),
            file_name=f"ontology_{domain_key}_{tag}.pdf", mime="application/pdf",
            key=f"pdf_{domain_key}_{tag}",
        )


tab1, tab2, tab3, tab4 = st.tabs([
    "E-commerce Products",
    "Organization Roles",
    "Biological Species",
    "Document / Subject",
])

with tab1:
    render_use_case("ecommerce", "E-commerce")
with tab2:
    render_use_case("organization", "Organization")
with tab3:
    render_use_case("biological", "Biological")
with tab4:
    render_use_case("document", "Document")

st.markdown("---")
st.caption(
    "Taxonomy to Ontology Conversion Studio. Developed by Randy Singh, "
    "Kalsnet (KNet) Consulting. Built with Streamlit, NetworkX, python-docx "
    "and ReportLab."
)