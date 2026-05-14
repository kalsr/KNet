

# =========================================================
# KNet PolicyForge AI
# Developed by Randy Singh from Kalsnet (KNet) Consulting Group
# =========================================================
#
# REQUIREMENTS:
#
# pip install streamlit groq pandas reportlab python-docx
#
# RUN:
#
# streamlit run app.py
#
# =========================================================

import streamlit as st
from groq import Groq
import json
import pandas as pd
from io import BytesIO
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from datetime import datetime

# =========================================================
# PAGE CONFIG
# =========================================================

st.set_page_config(
    page_title="KNet PolicyForge AI",
    layout="wide",
    page_icon="🛡️"
)

# =========================================================
# CUSTOM CSS
# =========================================================

st.markdown("""
<style>

.main {
    background-color: #0E1117;
}

h1, h2, h3, h4 {
    color: #4FC3F7;
}

.stButton>button {
    background-color: #1F6FEB;
    color: white;
    border-radius: 10px;
    padding: 0.5rem 1rem;
    border: none;
}

.stTextInput>div>div>input {
    border-radius: 8px;
}

textarea {
    border-radius: 8px !important;
}

.policy-header {
    text-align:center;
    padding:20px;
    background-color:#111827;
    border-radius:15px;
    margin-bottom:20px;
}

.footer-box {
    background-color:#111827;
    padding:15px;
    border-radius:10px;
    margin-top:20px;
}

</style>
""", unsafe_allow_html=True)

# =========================================================
# HEADER
# =========================================================

st.markdown("""
<div class="policy-header">
<h1>🛡️ KNet PolicyForge AI</h1>
<h4>Developed by Randy Singh from Kalsnet (KNet) Consulting Group</h4>
</div>
""", unsafe_allow_html=True)

# =========================================================
# SIDEBAR
# =========================================================

st.sidebar.title("⚙️ Configuration")

groq_api_key = st.sidebar.text_input(
    "Groq API Key",
    type="password"
)

st.sidebar.markdown("""
### How to Get Groq API Key

1. Visit:
https://console.groq.com

2. Create a free account

3. Navigate to API Keys

4. Generate a new API key

5. Paste the key here
""")

model_name = st.sidebar.selectbox(
    "Select Groq Model",
    [
        "llama-3.3-70b-versatile",
        "mixtral-8x7b-32768",
        "gemma2-9b-it"
    ]
)

policy_area = st.sidebar.selectbox(
    "Policy Area",
    [
        "Cybersecurity Policy",
        "Acceptable Use Policy",
        "Remote Work Policy",
        "AI Governance Policy",
        "Incident Response Policy",
        "Disaster Recovery Policy",
        "Business Continuity Policy",
        "Data Retention Policy",
        "Privacy Policy",
        "HR Policy",
        "Zero Trust Policy",
        "Vendor Management Policy"
    ]
)

organization_type = st.sidebar.selectbox(
    "Organization Type",
    [
        "Government",
        "Enterprise",
        "Healthcare",
        "Financial",
        "Educational",
        "Small Business",
        "Technology Company"
    ]
)

detail_level = st.sidebar.selectbox(
    "Policy Detail Level",
    [
        "Basic",
        "Standard",
        "Advanced",
        "Enterprise Grade"
    ]
)

# =========================================================
# MAIN INPUTS
# =========================================================

col1, col2 = st.columns(2)

with col1:
    organization_name = st.text_input(
        "Organization Name",
        placeholder="Enter organization name"
    )

    industry = st.text_input(
        "Industry",
        placeholder="Healthcare, Finance, Government etc."
    )

    compliance_requirements = st.text_area(
        "Compliance Requirements",
        placeholder="NIST, ISO 27001, HIPAA, GDPR, PCI-DSS etc."
    )

with col2:
    policy_objective = st.text_area(
        "Policy Objective",
        placeholder="Describe the policy objective"
    )

    risk_level = st.selectbox(
        "Risk Level",
        [
            "Low",
            "Moderate",
            "High",
            "Critical"
        ]
    )

    additional_requirements = st.text_area(
        "Additional Requirements",
        placeholder="Any additional controls or requirements"
    )

# =========================================================
# PROMPT TEMPLATE
# =========================================================

def build_prompt():
    prompt = f"""
Generate a highly professional enterprise-grade policy document.

Policy Area:
{policy_area}

Organization Name:
{organization_name}

Organization Type:
{organization_type}

Industry:
{industry}

Compliance Requirements:
{compliance_requirements}

Policy Objective:
{policy_objective}

Risk Level:
{risk_level}

Detail Level:
{detail_level}

Additional Requirements:
{additional_requirements}

The policy must include:

1. Policy Title
2. Purpose
3. Scope
4. Definitions
5. Policy Statements
6. Roles and Responsibilities
7. Standards and Controls
8. Procedures
9. Exceptions
10. Compliance Requirements
11. Enforcement
12. Review Schedule
13. Approval Section
14. Version History

Write in formal professional enterprise policy format.
"""
    return prompt

# =========================================================
# GROQ GENERATION
# =========================================================

generated_policy = ""

if st.button("🚀 Generate Policy"):

    if not groq_api_key:
        st.error("Please enter Groq API Key.")
    else:

        try:

            client = Groq(api_key=groq_api_key)

            with st.spinner("Generating professional policy..."):

                response = client.chat.completions.create(
                    model=model_name,
                    messages=[
                        {
                            "role": "system",
                            "content": "You are an enterprise policy development expert."
                        },
                        {
                            "role": "user",
                            "content": build_prompt()
                        }
                    ],
                    temperature=0.3,
                    max_tokens=6000
                )

                generated_policy = response.choices[0].message.content

                st.session_state["generated_policy"] = generated_policy

        except Exception as e:
            st.error(f"Error: {e}")

# =========================================================
# DISPLAY GENERATED POLICY
# =========================================================

if "generated_policy" in st.session_state:

    st.subheader("📄 Generated Policy")

    edited_policy = st.text_area(
        "Policy Editor",
        value=st.session_state["generated_policy"],
        height=700
    )

    # =====================================================
    # EXPORT FUNCTIONS
    # =====================================================

    def generate_pdf(text):

        buffer = BytesIO()

        doc = SimpleDocTemplate(buffer)

        styles = getSampleStyleSheet()

        story = []

        for line in text.split("\n"):
            story.append(Paragraph(line, styles['BodyText']))
            story.append(Spacer(1, 6))

        doc.build(story)

        buffer.seek(0)

        return buffer

    def generate_docx(text):

        doc = Document()

        doc.add_heading("Generated Policy", level=1)

        doc.add_paragraph(text)

        buffer = BytesIO()

        doc.save(buffer)

        buffer.seek(0)

        return buffer

    def generate_json(text):

        policy_json = {
            "generated_date": str(datetime.now()),
            "policy_area": policy_area,
            "organization_name": organization_name,
            "content": text
        }

        return json.dumps(policy_json, indent=4)

    def generate_csv(text):

        rows = []

        for line in text.split("\n"):
            rows.append({"Policy_Content": line})

        df = pd.DataFrame(rows)

        return df.to_csv(index=False).encode('utf-8')

    # =====================================================
    # EXPORT BUTTONS
    # =====================================================

    st.subheader("⬇️ Export Policy")

    c1, c2, c3, c4 = st.columns(4)

    with c1:

        pdf_file = generate_pdf(edited_policy)

        st.download_button(
            label="📄 Download PDF",
            data=pdf_file,
            file_name="policy.pdf",
            mime="application/pdf"
        )

    with c2:

        docx_file = generate_docx(edited_policy)

        st.download_button(
            label="📝 Download DOCX",
            data=docx_file,
            file_name="policy.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    with c3:

        csv_file = generate_csv(edited_policy)

        st.download_button(
            label="📊 Download CSV",
            data=csv_file,
            file_name="policy.csv",
            mime="text/csv"
        )

    with c4:

        json_file = generate_json(edited_policy)

        st.download_button(
            label="🧾 Download JSON",
            data=json_file,
            file_name="policy.json",
            mime="application/json"
        )

# =========================================================
# FILE UPLOAD SECTION
# =========================================================

st.subheader("📤 Upload Existing Policy")

uploaded_file = st.file_uploader(
    "Upload TXT Policy File",
    type=["txt"]
)

if uploaded_file is not None:

    content = uploaded_file.read().decode("utf-8")

    st.text_area(
        "Uploaded Policy Content",
        value=content,
        height=300
    )

# =========================================================
# FOOTER
# =========================================================

st.markdown("""
<div class="footer-box">

<h4>About KNet PolicyForge AI</h4>

This enterprise-grade AI policy generation platform helps organizations
develop professional governance, cybersecurity, compliance,
risk management, HR, and operational policies using advanced AI models.

Features:
- AI-powered policy development
- Enterprise-grade formatting
- Multi-format exports
- Professional GUI
- Editable policy workspace
- Upload & review capability

Developed by Randy Singh from Kalsnet (KNet) Consulting Group

</div>
""", unsafe_allow_html=True)