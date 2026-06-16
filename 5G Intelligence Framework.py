


#════════════════════════════════════════════════════════════════╗
#║   KNet 5G Intelligence Framework  |  Developed by Randy Singh   ║
#║   Kalsnet (KNet) Consulting Group                                ║
#╚══════════════════════════════════════════════════════════════════╝

# Run:  streamlit run knet_5g_framework.py
# Requires: streamlit pandas numpy plotly reportlab python-docx
# pip install streamlit pandas numpy plotly reportlab python-docx


import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import json
import io
import random
from datetime import datetime

# ── PDF / DOCX exports ──────────────────────────────────────────────────────
try:
    from reportlab.lib.pagesizes import letter, landscape
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    Table, TableStyle, HRFlowable)
    from reportlab.lib.units import inch
    REPORTLAB_OK = True
except ImportError:
    REPORTLAB_OK = False

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

# ── Page config ──────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="KNet 5G Intelligence Framework",
    page_icon="📡",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── Custom CSS ───────────────────────────────────────────────────────────────
st.markdown("""
<style>
  @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700;900&display=swap');
  html, body, [class*="css"] { font-family: 'Inter', sans-serif; }

  .knet-hero {
    background: linear-gradient(135deg, #0A1628 0%, #0D2B6B 50%, #1565C0 100%);
    border-radius: 16px;
    padding: 36px 40px 28px 40px;
    margin-bottom: 28px;
    box-shadow: 0 8px 40px rgba(21,101,192,0.4);
  }
  .knet-title {
    font-size: 2.6rem;
    font-weight: 900;
    color: #1E90FF;
    letter-spacing: -0.5px;
    line-height: 1.15;
    margin: 0;
  }
  .knet-subtitle {
    font-size: 1.05rem;
    color: #90CAF9;
    margin-top: 6px;
  }
  .knet-badge {
    display: inline-block;
    background: rgba(30,144,255,0.18);
    border: 1px solid #1E90FF;
    color: #64B5F6;
    padding: 4px 14px;
    border-radius: 20px;
    font-size: 0.82rem;
    font-weight: 600;
    margin-top: 14px;
    letter-spacing: 0.5px;
  }
  .metric-card {
    background: linear-gradient(135deg,#0D2B6B,#1565C0);
    border-radius: 12px;
    padding: 18px 20px;
    text-align: center;
    box-shadow: 0 4px 16px rgba(21,101,192,0.3);
    border: 1px solid rgba(30,144,255,0.3);
  }
  .metric-val { font-size: 1.9rem; font-weight: 900; color: #1E90FF; }
  .metric-lbl { font-size: 0.78rem; color: #90CAF9; margin-top: 4px; }
  .section-hdr {
    font-size: 1.3rem; font-weight: 700; color: #1E90FF;
    border-left: 4px solid #1E90FF;
    padding-left: 12px; margin: 28px 0 14px 0;
  }
  .taxonomy-box {
    background: #0A1628;
    border: 1px solid #1565C0;
    border-radius: 10px;
    padding: 18px 22px;
    color: #CFE2FF;
    font-size: 0.9rem;
    line-height: 1.7;
  }
  .field-card {
    background: #0D2148;
    border-left: 3px solid #1E90FF;
    border-radius: 6px;
    padding: 12px 16px;
    margin-bottom: 10px;
    color: #B0C4DE;
    font-size: 0.875rem;
  }
  .field-name { color: #1E90FF; font-weight: 700; font-size: 0.95rem; }
  .formula-tag {
    background: rgba(30,144,255,0.12);
    color: #64B5F6;
    padding: 2px 8px;
    border-radius: 4px;
    font-family: monospace;
    font-size: 0.82rem;
  }
</style>
""", unsafe_allow_html=True)

# ════════════════════════════════════════════════════════════════════════════
# CONSTANTS & LOOKUPS
# ════════════════════════════════════════════════════════════════════════════
REGIONS   = ["North America","Europe","Asia-Pacific","Middle East","Latin America","Africa"]
BANDS     = ["Sub-6 GHz","mmWave (24-100 GHz)","FR1 (450 MHz–6 GHz)","FR2 (24.25–52.6 GHz)"]
DEPLOY    = ["Standalone (SA)","Non-Standalone (NSA)","Hybrid"]
USE_CASES = ["eMBB","URLLC","mMTC","FWA","Network Slicing","V2X","Private 5G"]
OPERATORS = ["Verizon","AT&T","T-Mobile","Ericsson Net","Nokia Net","Huawei Net","Samsung Net","Rakuten"]
VENDORS   = ["Ericsson","Nokia","Huawei","Samsung","ZTE","Mavenir","Rakuten Symphony"]
SLICES    = ["eMBB Slice","URLLC Slice","mMTC Slice","Custom Slice"]

FIELD_DOCS = {
    "Record_ID":                 ("Unique identifier per site record.",
                                  "KNet-5G-{NNNN} sequential"),
    "Region":                    ("Geographic deployment region.",
                                  "Categorical — 6 global regions"),
    "Operator":                  ("Mobile Network Operator name.",
                                  "Categorical — 8 major operators"),
    "Vendor":                    ("Radio Access Network (RAN) equipment vendor.",
                                  "Categorical — 7 leading vendors"),
    "Frequency_Band":            ("Spectrum band used for 5G transmission.",
                                  "Categorical — 4 3GPP-defined bands"),
    "Deployment_Mode":           ("SA = standalone 5G Core; NSA = LTE anchor; Hybrid = mixed.",
                                  "Categorical — 3 modes per 3GPP Rel-15/16"),
    "Use_Case":                  ("Primary 5G application type per ITU IMT-2020.",
                                  "Categorical — 7 use case categories"),
    "Network_Slice":             ("Logical partition of the 5G network (3GPP TS 23.501).",
                                  "Categorical — 4 slice types"),
    "Throughput_Mbps":           ("Peak downlink data rate in Megabits per second.",
                                  "Normal(μ=1000 if mmWave/FR2 else 400, σ=20%) clipped ≥50"),
    "Latency_ms":                ("End-to-end round-trip delay in milliseconds.",
                                  "Normal(μ=1 if URLLC else 10, σ=30%) clipped ≥0.5"),
    "Jitter_ms":                 ("Variation in packet delay (PDV).",
                                  "|Normal(0.5, 0.2)|"),
    "Packet_Loss_Pct":           ("Percentage of packets not delivered.",
                                  "|Normal(0.01, 0.005)|"),
    "SNR_dB":                    ("Signal-to-Noise Ratio in decibels.",
                                  "Uniform(5, 30) dB"),
    "Spectral_Efficiency_bpHz":  ("Bits per second per Hz of bandwidth.",
                                  "log₂(1 + 10^(SNR_dB/10))  [Shannon Capacity]"),
    "Device_Density_per_km2":    ("Connected devices per square kilometre.",
                                  "Uniform(1K–1M if mMTC, else 100–10K)"),
    "Coverage_Radius_km":        ("Cell radius for the deployment.",
                                  "Uniform(0.1–2 km if mmWave/FR2, else 0.5–10 km)"),
    "Energy_Efficiency_MbpsW":   ("Throughput per Watt — key 5G green KPI.",
                                  "Throughput_Mbps / Uniform(10, 50) W"),
    "Network_Availability_Pct":  ("Percentage uptime; targets '5-nines' (99.999%).",
                                  "Uniform(99.0, 99.999)"),
    "QoS_Score":                 ("Composite Quality-of-Service index (0–100).",
                                  "(Throughput/2000×40)+(max(0,10−Latency)/10×30)+\n"
                                  "((Avail−99)×20)+((1−PacketLoss)×10), clipped [0,100]"),
    "Active_Sites":              ("Number of active base station sites.",
                                  "Uniform Integer(10, 5000)"),
    "Year":                      ("Deployment or measurement year.",
                                  "Choice from [2023, 2024, 2025, 2026]"),
}

# ════════════════════════════════════════════════════════════════════════════
# DATA GENERATION
# ════════════════════════════════════════════════════════════════════════════
def generate_data(n: int = 300, seed: int = 42) -> pd.DataFrame:
    np.random.seed(seed)
    random.seed(seed)
    rows = []
    for i in range(1, n + 1):
        band   = random.choice(BANDS)
        uc     = random.choice(USE_CASES)
        hi_bw  = "mmWave" in band or "FR2" in band

        tput_mu = 1000 if hi_bw else 400
        throughput = max(50.0, round(float(np.random.normal(tput_mu, tput_mu * 0.2)), 2))

        lat_mu = 1.0 if uc == "URLLC" else 10.0
        latency = max(0.5, round(abs(float(np.random.normal(lat_mu, lat_mu * 0.3))), 2))

        jitter       = round(abs(float(np.random.normal(0.5, 0.2))), 2)
        packet_loss  = round(abs(float(np.random.normal(0.01, 0.005))), 4)
        snr_db       = round(float(np.random.uniform(5, 30)), 1)
        spectral_eff = round(np.log2(1 + 10 ** (snr_db / 10)), 2)
        dev_density  = int(np.random.uniform(1000, 1_000_000) if uc == "mMTC"
                          else np.random.uniform(100, 10_000))
        coverage     = round(float(np.random.uniform(0.1, 2) if hi_bw
                                   else np.random.uniform(0.5, 10)), 2)
        power_w      = float(np.random.uniform(10, 50))
        energy_eff   = round(throughput / power_w, 2)
        availability = round(float(np.random.uniform(99.0, 99.999)), 3)

        qos = (throughput / 2000 * 40
               + max(0, (10 - latency)) / 10 * 30
               + (availability - 99) * 20
               + (1 - packet_loss) * 10)
        qos = round(min(100.0, max(0.0, qos)), 2)

        rows.append({
            "Record_ID":               f"KNet-5G-{i:04d}",
            "Region":                  random.choice(REGIONS),
            "Operator":                random.choice(OPERATORS),
            "Vendor":                  random.choice(VENDORS),
            "Frequency_Band":          band,
            "Deployment_Mode":         random.choice(DEPLOY),
            "Use_Case":                uc,
            "Network_Slice":           random.choice(SLICES),
            "Throughput_Mbps":         throughput,
            "Latency_ms":              latency,
            "Jitter_ms":               jitter,
            "Packet_Loss_Pct":         packet_loss,
            "SNR_dB":                  snr_db,
            "Spectral_Efficiency_bpHz": spectral_eff,
            "Device_Density_per_km2":  dev_density,
            "Coverage_Radius_km":      coverage,
            "Energy_Efficiency_MbpsW": energy_eff,
            "Network_Availability_Pct": availability,
            "QoS_Score":               qos,
            "Active_Sites":            random.randint(10, 5000),
            "Year":                    random.choice([2023, 2024, 2025, 2026]),
        })
    return pd.DataFrame(rows)

# ════════════════════════════════════════════════════════════════════════════
# SESSION STATE
# ════════════════════════════════════════════════════════════════════════════
if "df" not in st.session_state or st.session_state.get("df") is None:
    st.session_state.df = generate_data(300)
if "n_records" not in st.session_state:
    st.session_state.n_records = 300

# ════════════════════════════════════════════════════════════════════════════
# SIDEBAR
# ════════════════════════════════════════════════════════════════════════════
with st.sidebar:
    st.markdown("## ⚙️ Controls")
    n_records = st.slider("Number of Records", 0, 300,
                          st.session_state.n_records, step=10)
    seed_val  = st.number_input("Random Seed", 0, 9999, 42)

    col1, col2 = st.columns(2)
    with col1:
        if st.button("🔄 Generate", use_container_width=True):
            st.session_state.df = generate_data(n_records, seed=seed_val)
            st.session_state.n_records = n_records
            st.rerun()
    with col2:
        if st.button("🗑️ Reset", use_container_width=True):
            st.session_state.df = generate_data(300)
            st.session_state.n_records = 300
            st.rerun()

    st.markdown("---")
    st.markdown("### 🔍 Filters")
    df_all = st.session_state.df
    sel_region  = st.multiselect("Region",    REGIONS,   default=REGIONS)
    sel_band    = st.multiselect("Freq Band", BANDS,     default=BANDS)
    sel_usecase = st.multiselect("Use Case",  USE_CASES, default=USE_CASES)

    st.markdown("---")
    st.markdown("### 📤 Export")
    export_fmt = st.selectbox("Format", ["PDF","Word (.docx)","Text (.txt)","JSON"])

df = st.session_state.df
if len(df) > 0:
    mask = (df["Region"].isin(sel_region) &
            df["Frequency_Band"].isin(sel_band) &
            df["Use_Case"].isin(sel_usecase))
    df_f = df[mask].copy()
else:
    df_f = df.copy()

# ════════════════════════════════════════════════════════════════════════════
# HERO HEADER
# ════════════════════════════════════════════════════════════════════════════
st.markdown("""
<div class="knet-hero">
  <p class="knet-title">📡 KNet 5G Intelligence Framework</p>
  <p class="knet-subtitle">Advanced 5G Network Analytics, Taxonomy & Diagnostics Platform</p>
  <span class="knet-badge">🏢 Developed by Randy Singh &nbsp;|&nbsp; Kalsnet (KNet) Consulting Group</span>
</div>
""", unsafe_allow_html=True)

# ════════════════════════════════════════════════════════════════════════════
# TABS
# ════════════════════════════════════════════════════════════════════════════
tabs = st.tabs([
    "🏠 Framework Overview",
    "📊 Data Explorer",
    "📈 Analytics",
    "🗺️ Diagrams & Flowcharts",
    "📖 Field Dictionary",
    "📤 Export",
])

# ─────────────────────────────────────────────────────────────────────────────
# TAB 1 — FRAMEWORK OVERVIEW
# ─────────────────────────────────────────────────────────────────────────────
with tabs[0]:
    st.markdown('<p class="section-hdr">5G Taxonomy</p>', unsafe_allow_html=True)
    st.markdown("""
<div class="taxonomy-box">
<b style="color:#1E90FF">5G (Fifth Generation)</b> is the latest global wireless standard defined by
the <b>ITU IMT-2020</b> specification and implemented via <b>3GPP Release 15/16/17/18</b>.
It delivers three foundational service categories:
<br><br>
<b style="color:#64B5F6">① eMBB — Enhanced Mobile Broadband</b><br>
Peak downlink ≥ 20 Gbps, uplink ≥ 10 Gbps. Powers AR/VR, 4K/8K streaming, ultra-HD conferencing.
<br><br>
<b style="color:#64B5F6">② URLLC — Ultra-Reliable Low-Latency Communications</b><br>
End-to-end latency ≤ 1 ms, reliability ≥ 99.9999%. Enables autonomous vehicles, remote surgery, industrial automation.
<br><br>
<b style="color:#64B5F6">③ mMTC — Massive Machine-Type Communications</b><br>
Supports ≥ 1 million devices/km². Powers smart cities, IoT sensors, agriculture, logistics.
<br><br>
<b style="color:#64B5F6">Key 5G Architecture Pillars:</b>
<ul>
  <li><b>Service-Based Architecture (SBA)</b> — microservice 5G Core (5GC) with NRF, AMF, SMF, UPF</li>
  <li><b>Network Slicing</b> — isolated logical networks over shared physical infrastructure</li>
  <li><b>Open RAN (O-RAN)</b> — disaggregated, vendor-neutral radio networks</li>
  <li><b>Edge Computing (MEC)</b> — compute at the network edge for sub-ms applications</li>
  <li><b>Spectrum</b> — Sub-6 GHz (coverage) + mmWave (capacity) + DSS (LTE coexistence)</li>
  <li><b>Beamforming / Massive MIMO</b> — spatial multiplexing with 64T64R+ antenna arrays</li>
</ul>
</div>
""", unsafe_allow_html=True)

    st.markdown('<p class="section-hdr">KNet Framework Philosophy</p>', unsafe_allow_html=True)
    st.markdown("""
<div class="taxonomy-box">
The <b style="color:#1E90FF">KNet 5G Intelligence Framework</b> by <b>Randy Singh</b> at
<b>Kalsnet (KNet) Consulting Group</b> is a structured methodology for assessing, deploying,
and optimising 5G networks. It is built on five pillars:
<ol>
  <li><b>Performance Benchmarking</b> — KPI baselines against ITU IMT-2020 targets</li>
  <li><b>Use-Case Alignment</b> — matching network slice type to application SLA</li>
  <li><b>Spectral Intelligence</b> — band selection strategy (FR1 vs FR2) per geography</li>
  <li><b>Energy Optimisation</b> — Green 5G: Mbps per Watt as a first-class KPI</li>
  <li><b>Vendor Agnostic Analytics</b> — O-RAN aligned, multi-vendor scoring</li>
</ol>
</div>
""", unsafe_allow_html=True)

    # KPI Summary Cards
    st.markdown('<p class="section-hdr">Live KPI Summary</p>', unsafe_allow_html=True)
    if len(df_f) > 0:
        c1,c2,c3,c4,c5,c6 = st.columns(6)
        kpi_data = [
            (c1, f"{len(df_f)}", "Total Records"),
            (c2, f"{df_f['Throughput_Mbps'].mean():.0f}", "Avg Throughput (Mbps)"),
            (c3, f"{df_f['Latency_ms'].mean():.2f}", "Avg Latency (ms)"),
            (c4, f"{df_f['QoS_Score'].mean():.1f}", "Avg QoS Score"),
            (c5, f"{df_f['Network_Availability_Pct'].mean():.3f}%", "Avg Availability"),
            (c6, f"{df_f['Energy_Efficiency_MbpsW'].mean():.1f}", "Avg Energy Eff."),
        ]
        for col, val, lbl in kpi_data:
            with col:
                st.markdown(f"""
                <div class="metric-card">
                  <div class="metric-val">{val}</div>
                  <div class="metric-lbl">{lbl}</div>
                </div>""", unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────────────────────
# TAB 2 — DATA EXPLORER
# ─────────────────────────────────────────────────────────────────────────────
with tabs[1]:
    st.markdown('<p class="section-hdr">Synthetic 5G Network Dataset</p>',
                unsafe_allow_html=True)
    st.info(f"Showing **{len(df_f)}** of **{len(df)}** records after filters.")
    st.dataframe(df_f, use_container_width=True, height=480)

    st.markdown('<p class="section-hdr">Descriptive Statistics</p>',
                unsafe_allow_html=True)
    num_cols = df_f.select_dtypes(include=np.number).columns.tolist()
    st.dataframe(df_f[num_cols].describe().T.style.format("{:.3f}"),
                 use_container_width=True)

# ─────────────────────────────────────────────────────────────────────────────
# TAB 3 — ANALYTICS
# ─────────────────────────────────────────────────────────────────────────────
with tabs[2]:
    if len(df_f) == 0:
        st.warning("No data to display. Adjust filters.")
    else:
        st.markdown('<p class="section-hdr">Throughput by Frequency Band</p>',
                    unsafe_allow_html=True)
        fig1 = px.box(df_f, x="Frequency_Band", y="Throughput_Mbps",
                      color="Frequency_Band",
                      color_discrete_sequence=px.colors.sequential.Blues_r,
                      template="plotly_dark")
        fig1.update_layout(showlegend=False, paper_bgcolor="#0A1628",
                           plot_bgcolor="#0D2148")
        st.plotly_chart(fig1, use_container_width=True)

        col_a, col_b = st.columns(2)
        with col_a:
            st.markdown('<p class="section-hdr">Use Case Distribution</p>',
                        unsafe_allow_html=True)
            uc_cnt = df_f["Use_Case"].value_counts().reset_index()
            uc_cnt.columns = ["Use_Case","Count"]
            fig2 = px.pie(uc_cnt, names="Use_Case", values="Count",
                          color_discrete_sequence=px.colors.sequential.Blues_r,
                          template="plotly_dark")
            fig2.update_layout(paper_bgcolor="#0A1628")
            st.plotly_chart(fig2, use_container_width=True)

        with col_b:
            st.markdown('<p class="section-hdr">QoS Score by Region</p>',
                        unsafe_allow_html=True)
            fig3 = px.bar(df_f.groupby("Region")["QoS_Score"].mean().reset_index(),
                          x="Region", y="QoS_Score",
                          color="QoS_Score",
                          color_continuous_scale="Blues",
                          template="plotly_dark")
            fig3.update_layout(paper_bgcolor="#0A1628", plot_bgcolor="#0D2148",
                               showlegend=False)
            st.plotly_chart(fig3, use_container_width=True)

        st.markdown('<p class="section-hdr">Throughput vs Latency (SNR coloured)</p>',
                    unsafe_allow_html=True)
        fig4 = px.scatter(df_f, x="Latency_ms", y="Throughput_Mbps",
                          color="SNR_dB", size="QoS_Score",
                          hover_data=["Record_ID","Region","Use_Case"],
                          color_continuous_scale="Blues",
                          template="plotly_dark")
        fig4.update_layout(paper_bgcolor="#0A1628", plot_bgcolor="#0D2148")
        st.plotly_chart(fig4, use_container_width=True)

        st.markdown('<p class="section-hdr">Energy Efficiency by Vendor</p>',
                    unsafe_allow_html=True)
        fig5 = px.bar(df_f.groupby("Vendor")["Energy_Efficiency_MbpsW"].mean()
                      .reset_index().sort_values("Energy_Efficiency_MbpsW", ascending=False),
                      x="Vendor", y="Energy_Efficiency_MbpsW",
                      color="Energy_Efficiency_MbpsW",
                      color_continuous_scale="Blues",
                      template="plotly_dark")
        fig5.update_layout(paper_bgcolor="#0A1628", plot_bgcolor="#0D2148",
                           showlegend=False)
        st.plotly_chart(fig5, use_container_width=True)

        st.markdown('<p class="section-hdr">Spectral Efficiency Distribution</p>',
                    unsafe_allow_html=True)
        fig6 = px.histogram(df_f, x="Spectral_Efficiency_bpHz",
                            color="Frequency_Band",
                            nbins=40,
                            color_discrete_sequence=px.colors.sequential.Blues_r,
                            template="plotly_dark", barmode="overlay", opacity=0.75)
        fig6.update_layout(paper_bgcolor="#0A1628", plot_bgcolor="#0D2148")
        st.plotly_chart(fig6, use_container_width=True)

        st.markdown('<p class="section-hdr">Heatmap: Avg Throughput (Region × Use Case)</p>',
                    unsafe_allow_html=True)
        pivot = df_f.pivot_table(values="Throughput_Mbps",
                                  index="Region", columns="Use_Case",
                                  aggfunc="mean").fillna(0)
        fig7 = px.imshow(pivot, color_continuous_scale="Blues",
                         text_auto=".0f", template="plotly_dark")
        fig7.update_layout(paper_bgcolor="#0A1628")
        st.plotly_chart(fig7, use_container_width=True)

# ─────────────────────────────────────────────────────────────────────────────
# TAB 4 — DIAGRAMS & FLOWCHARTS  (Mermaid via HTML)
# ─────────────────────────────────────────────────────────────────────────────
MERMAID_HTML = """
<!DOCTYPE html><html><head>
<script src="https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.min.js">
<style>
  body { background:#0A1628; color:#90CAF9; font-family:Inter,sans-serif; padding:20px; }
  h3 { color:#1E90FF; border-left:4px solid #1E90FF; padding-left:10px; margin-top:28px; }
  .mermaid { background:#0D2148; border-radius:10px; padding:20px; margin:12px 0; }
</style>
</head>
<body>
<script>mermaid.initialize({theme:'dark',themeVariables:{primaryColor:'#1E90FF',
  primaryTextColor:'#E3F2FD',primaryBorderColor:'#1565C0',
  lineColor:'#64B5F6',background:'#0D2148',nodeBorder:'#1E90FF'}});</script>

<h3>① 5G Architecture Layers</h3>
<div class="mermaid">
graph TB
  UE[📱 User Equipment]-->RAN[🗼 5G RAN\nNext-Gen NodeB gNB]
  RAN-->CN[⚙️ 5G Core\nService-Based Architecture]
  CN-->DN[🌐 Data Networks\nInternet / MEC / Slices]
  CN-->OAM[🛠️ OAM\nNetwork Management]
  RAN-->MEC[🖥️ Multi-Access Edge\nComputing MEC]
</div>

<h3>② Network Slicing Flow</h3>
<div class="mermaid">
flowchart LR
  REQ[Service Request]-->NSSF[Network Slice\nSelection Function\nNSSF]
  NSSF-->|eMBB|S1[eMBB Slice\n≥20 Gbps]
  NSSF-->|URLLC|S2[URLLC Slice\n≤1ms]
  NSSF-->|mMTC|S3[mMTC Slice\n1M dev/km²]
  S1-->UPF1[UPF 1]
  S2-->UPF2[UPF 2]
  S3-->UPF3[UPF 3]
  UPF1 & UPF2 & UPF3 --> DN[Data Network]
</div>

<h3>③ KNet Framework Decision Flowchart</h3>
<div class="mermaid">
flowchart TD
  START([🚀 New 5G Deployment])-->UC{Use Case?}
  UC-->|Video / AR / VR|eMBB[eMBB Path\nSub-6 or mmWave\nSA Mode]
  UC-->|Autonomous / Surgery|URLLC[URLLC Path\nFR1 + MEC\nSA + Network Slice]
  UC-->|IoT / Smart City|mMTC[mMTC Path\nSub-6 GHz\nNB-IoT / RedCap]
  eMBB-->BAND{Band?}
  BAND-->|Dense urban|mmWave[mmWave 24-100 GHz\nSmall cells]
  BAND-->|Suburban|Sub6[Sub-6 GHz\nMacro cells]
  URLLC-->MEC2[Deploy MEC\n≤10ms to UE]
  mMTC-->LPWA[Use Low Power\nNarrow Band IoT]
  mmWave & Sub6 & MEC2 & LPWA-->KPI[Measure KPIs\nThroughput / Latency\nQoS / Energy]
  KPI-->PASS{KPI ≥ ITU Target?}
  PASS-->|Yes|DEPLOY([✅ Deploy])
  PASS-->|No|OPT[Optimise:\nBeamforming\nSON\nSlice reconfig]
  OPT-->KPI
</div>

<h3>④ 5G Core NF Interactions (SBA)</h3>
<div class="mermaid">
graph LR
  AMF[AMF\nAccess & Mobility] <--> SMF[SMF\nSession Mgmt]
  SMF <--> UPF[UPF\nUser Plane]
  AMF <--> AUSF[AUSF\nAuthentication]
  AMF <--> UDM[UDM\nSubscriber Data]
  SMF <--> PCF[PCF\nPolicy Control]
  AMF <--> NRF[NRF\nService Registry]
  SMF <--> NRF
  PCF <--> NRF
</div>

<h3>⑤ QoS Scoring Formula Flow</h3>
<div class="mermaid">
flowchart LR
  A[Throughput\nMbps] -->|÷2000×40|W1[40%\nWeight]
  B[Latency\nms] -->|max 0 10-L ÷10×30|W2[30%\nWeight]
  C[Availability\n%] -->|A-99×20|W3[20%\nWeight]
  D[Packet Loss\n%] -->|1-PL×10|W4[10%\nWeight]
  W1 & W2 & W3 & W4 --> SUM[Σ QoS Score\n0–100]
</div>

<h3>⑥ Spectrum Strategy Map</h3>
<div class="mermaid">
quadrantChart
  title 5G Frequency Band Strategy
  x-axis Low Coverage --> High Coverage
  y-axis Low Capacity --> High Capacity
  quadrant-1 High Cap High Cov
  quadrant-2 High Cap Low Cov
  quadrant-3 Low Cap Low Cov
  quadrant-4 Low Cap High Cov
  mmWave 28GHz: [0.1, 0.9]
  FR2 39GHz: [0.15, 0.85]
  Sub6 3.5GHz: [0.65, 0.55]
  FR1 700MHz: [0.9, 0.2]
</div>

</body></html>
"""

with tabs[3]:
    st.markdown('<p class="section-hdr">5G Architecture & KNet Framework Diagrams</p>',
                unsafe_allow_html=True)
    st.info("ℹ️ Mermaid diagrams render via embedded HTML below. All 6 diagrams included.")
    st.components.v1.html(MERMAID_HTML, height=2800, scrolling=True)

# ─────────────────────────────────────────────────────────────────────────────
# TAB 5 — FIELD DICTIONARY
# ─────────────────────────────────────────────────────────────────────────────
with tabs[4]:
    st.markdown('<p class="section-hdr">Field Definitions & Formulas</p>',
                unsafe_allow_html=True)
    for fname, (desc, formula) in FIELD_DOCS.items():
        st.markdown(f"""
<div class="field-card">
  <span class="field-name">🔹 {fname}</span><br>
  {desc}<br>
  <span class="formula-tag">📐 {formula}</span>
</div>""", unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────────────────────
# TAB 6 — EXPORT
# ─────────────────────────────────────────────────────────────────────────────
with tabs[5]:
    st.markdown('<p class="section-hdr">Export Filtered Dataset</p>',
                unsafe_allow_html=True)
    st.info(f"Will export **{len(df_f)}** records in **{export_fmt}** format.")

    if st.button("⬇️ Generate & Download", use_container_width=True):
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")

        # ── JSON ──────────────────────────────────────────────────────────
        if export_fmt == "JSON":
            payload = {
                "metadata": {
                    "framework": "KNet 5G Intelligence Framework",
                    "author": "Randy Singh",
                    "organization": "Kalsnet (KNet) Consulting Group",
                    "generated": datetime.now().isoformat(),
                    "records": len(df_f),
                },
                "taxonomy": {
                    "eMBB": "Enhanced Mobile Broadband — peak ≥20 Gbps",
                    "URLLC": "Ultra-Reliable Low-Latency — ≤1ms latency",
                    "mMTC": "Massive Machine-Type — 1M devices/km²",
                },
                "data": df_f.to_dict(orient="records"),
            }
            buf = io.BytesIO(json.dumps(payload, indent=2).encode())
            st.download_button("⬇️ Download JSON", buf, f"KNet5G_{ts}.json",
                               "application/json")

        # ── TEXT ──────────────────────────────────────────────────────────
        elif export_fmt == "Text (.txt)":
            lines = [
                "=" * 70,
                "  KNet 5G Intelligence Framework",
                "  Randy Singh | Kalsnet (KNet) Consulting Group",
                f"  Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
                f"  Records: {len(df_f)}",
                "=" * 70,
                "",
                "5G TAXONOMY",
                "-" * 40,
                "eMBB  — Enhanced Mobile Broadband (peak ≥ 20 Gbps)",
                "URLLC — Ultra-Reliable Low-Latency (≤ 1 ms)",
                "mMTC  — Massive Machine-Type Comms (1M dev/km²)",
                "",
                "DATA",
                "-" * 40,
            ]
            lines.append(df_f.to_string(index=False))
            buf = io.BytesIO("\n".join(lines).encode())
            st.download_button("⬇️ Download TXT", buf, f"KNet5G_{ts}.txt", "text/plain")

        # ── WORD ──────────────────────────────────────────────────────────
        elif export_fmt == "Word (.docx)":
            if not DOCX_OK:
                st.error("python-docx not installed.")
            else:
                doc = Document()
                # Title
                t = doc.add_heading("KNet 5G Intelligence Framework", 0)
                t.runs[0].font.color.rgb = RGBColor(0x1E, 0x90, 0xFF)
                doc.add_paragraph("Developed by Randy Singh | Kalsnet (KNet) Consulting Group")
                doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                doc.add_paragraph(f"Records: {len(df_f)}")
                doc.add_heading("5G Taxonomy", level=1)
                for item in [
                    "eMBB — Enhanced Mobile Broadband: Peak downlink ≥ 20 Gbps.",
                    "URLLC — Ultra-Reliable Low-Latency: End-to-end latency ≤ 1 ms.",
                    "mMTC — Massive Machine-Type: Supports ≥ 1 million devices/km².",
                ]:
                    doc.add_paragraph(item, style="List Bullet")
                doc.add_heading("Field Definitions", level=1)
                for fname, (desc, formula) in FIELD_DOCS.items():
                    p = doc.add_paragraph()
                    run = p.add_run(fname + ": ")
                    run.bold = True
                    run.font.color.rgb = RGBColor(0x1E, 0x90, 0xFF)
                    p.add_run(f"{desc}  [Formula: {formula}]")
                doc.add_heading("Filtered Data", level=1)
                cols = list(df_f.columns)
                table = doc.add_table(rows=1, cols=len(cols))
                table.style = "Light Shading Accent 1"
                for i, c in enumerate(cols):
                    table.rows[0].cells[i].text = c
                for _, row in df_f.head(200).iterrows():
                    cells = table.add_row().cells
                    for i, c in enumerate(cols):
                        cells[i].text = str(row[c])
                buf = io.BytesIO()
                doc.save(buf)
                buf.seek(0)
                st.download_button("⬇️ Download DOCX", buf, f"KNet5G_{ts}.docx",
                                   "application/vnd.openxmlformats-officedocument"
                                   ".wordprocessingml.document")

        # ── PDF ──────────────────────────────────────────────────────────
        elif export_fmt == "PDF":
            if not REPORTLAB_OK:
                st.error("reportlab not installed.")
            else:
                buf = io.BytesIO()
                doc = SimpleDocTemplate(buf, pagesize=landscape(letter),
                                        rightMargin=30, leftMargin=30,
                                        topMargin=30, bottomMargin=30)
                styles = getSampleStyleSheet()
                blue = colors.HexColor("#1E90FF")
                title_style = ParagraphStyle("KNetTitle", parent=styles["Title"],
                                             textColor=blue, fontSize=18, spaceAfter=6)
                h1 = ParagraphStyle("KNetH1", parent=styles["Heading1"],
                                    textColor=blue, fontSize=13, spaceBefore=12)
                body = styles["Normal"]
                story = [
                    Paragraph("KNet 5G Intelligence Framework", title_style),
                    Paragraph("Randy Singh | Kalsnet (KNet) Consulting Group", body),
                    Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
                              f" | Records: {len(df_f)}", body),
                    HRFlowable(width="100%", color=blue, spaceAfter=8),
                    Paragraph("5G Taxonomy", h1),
                    Paragraph("eMBB — Enhanced Mobile Broadband: Peak downlink ≥ 20 Gbps.", body),
                    Paragraph("URLLC — Ultra-Reliable Low-Latency: ≤ 1 ms end-to-end delay.", body),
                    Paragraph("mMTC — Massive Machine-Type: ≥ 1 million devices/km².", body),
                    Spacer(1, 10),
                    Paragraph("Dataset (first 150 rows)", h1),
                ]
                show_cols = ["Record_ID","Region","Use_Case","Throughput_Mbps",
                             "Latency_ms","QoS_Score","Network_Availability_Pct",
                             "Frequency_Band","Deployment_Mode","Vendor"]
                tdata = [show_cols] + [
                    [str(r[c])[:18] for c in show_cols]
                    for _, r in df_f.head(150).iterrows()
                ]
                col_widths = [80, 80, 75, 70, 60, 55, 90, 85, 80, 70]
                t = Table(tdata, colWidths=col_widths, repeatRows=1)
                t.setStyle(TableStyle([
                    ("BACKGROUND",  (0,0), (-1,0), blue),
                    ("TEXTCOLOR",   (0,0), (-1,0), colors.white),
                    ("FONTSIZE",    (0,0), (-1,0), 7),
                    ("FONTSIZE",    (0,1), (-1,-1), 6),
                    ("ROWBACKGROUNDS", (0,1), (-1,-1),
                     [colors.HexColor("#0D2148"), colors.HexColor("#0A1628")]),
                    ("TEXTCOLOR",   (0,1), (-1,-1), colors.HexColor("#CFE2FF")),
                    ("GRID",        (0,0), (-1,-1), 0.3, colors.HexColor("#1565C0")),
                    ("VALIGN",      (0,0), (-1,-1), "MIDDLE"),
                ]))
                story.append(t)
                doc.build(story)
                buf.seek(0)
                st.download_button("⬇️ Download PDF", buf, f"KNet5G_{ts}.pdf",
                                   "application/pdf")

    st.markdown("---")
    st.markdown("**© KNet 5G Intelligence Framework | Randy Singh | "
                "Kalsnet (KNet) Consulting Group**")

