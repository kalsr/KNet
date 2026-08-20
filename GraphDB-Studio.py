# GraphDB Insights Studio
# Developed by Randy Singh - Kalsnet (KNet) Consulting
# A Streamlit application demonstrating graph-database concepts using an
# in-memory graph engine (NetworkX). Supports synthetic data generation,
# real data upload, interactive visualization, graph queries, and export
# of results to PDF, Word, TXT and CSV.
import io
import random
from datetime import date, timedelta
import pandas as pd
import networkx as nx
import plotly.express as px
import plotly.graph_objects as go
import streamlit as st
# Optional export libraries
from docx import Document
from docx.shared import Pt, RGBColor
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
st.set_page_config(page_title="GraphDB Insights Studio", layout="wide")
# ----------------------------------------------------------------------------
# TITLE BAR
# ----------------------------------------------------------------------------
st.markdown(
    """
    <style>
    .title-bar {font-size:54px; font-weight:900; color:#0033CC; margin-bottom:0; text-align:center;}
    .subtitle-bar {font-size:24px; font-weight:900; color:#0033CC; margin-top:0; text-align:center;}
    .section-header {font-size:24px; font-weight:700; color:#0033CC; margin-top:25px;}
    </style>
    <p class="title-bar">GraphDB Insights Studio</p>
    <p class="subtitle-bar">Developed by Randy Singh from Kalsnet (KNet) Consulting</p>
    <hr style="border:2px solid #0033CC;">
    """,
    unsafe_allow_html=True,
)
REQUIRED_COLS = ["record_id", "source_type", "source_id", "source_name", "relationship",
                  "target_type", "target_id", "target_name", "weight", "record_date"]
# ----------------------------------------------------------------------------
# EXPLAINERS
# ----------------------------------------------------------------------------
with st.expander("What is a Graph Database? (click to expand)", expanded=True):
    st.markdown("""
A **graph database** stores data as a network of **nodes** (entities) and **edges**
(relationships), instead of rows and tables like a relational database.
| Concept | Relational DB | Graph DB |
|---|---|---|
| Core unit | Table / Row | Node / Edge |
| Relationships | Foreign keys + JOINs | First-class edges, traversed directly |
| Best for | Structured, tabular reporting | Highly connected data: social networks, fraud detection, recommendation engines, knowledge graphs |
| Query style | SQL (`SELECT ... JOIN`) | Graph traversal (e.g. Cypher: `MATCH (a)-[:KNOWS]->(b)`) |
**Why it matters:** relationships are stored *directly* on disk as pointers between
nodes, so multi-hop questions ("who does my colleague's colleague know?") run in
near-constant time per hop, instead of requiring expensive JOINs that slow down as
data grows. Popular graph databases include **Neo4j**, **Amazon Neptune**, **ArangoDB**,
and **TigerGraph**. This app simulates the same node/edge model in-memory using the
Python library **NetworkX**, so the concepts transfer directly to a production graph DB.
""")
with st.expander("Schema Used in This App (click to expand)"):
    st.markdown("""
This app models a small **enterprise relationship graph** with three node types and
four relationship (edge) types.
**Node types**
- **Person** — an individual employee/customer (`id`, `name`, `age`, `city`, `department`)
- **Company** — an organization (`id`, `name`, `industry`)
- **Product** — a good or service (`id`, `name`, `category`, `price`)
**Relationship types (edges)**
- `WORKS_AT` — Person → Company
- `KNOWS` — Person → Person (social connection)
- `PURCHASED` — Person → Product
- `SUPPLIES` — Company → Product
**CSV column meaning (each row = one relationship / edge):**
| Column | Meaning |
|---|---|
| `record_id` | Unique identifier for the relationship record |
| `source_type` | Node type of the source (`Person`, `Company`) |
| `source_id` | Unique ID of the source node |
| `source_name` | Display name of the source node |
| `relationship` | Edge type (`WORKS_AT`, `KNOWS`, `PURCHASED`, `SUPPLIES`) |
| `target_type` | Node type of the target (`Person`, `Company`, `Product`) |
| `target_id` | Unique ID of the target node |
| `target_name` | Display name of the target node |
| `weight` | Numeric strength of the relationship (years employed, purchase price, friendship strength, quantity supplied) |
| `record_date` | Date the relationship/event was recorded |
""")
# ----------------------------------------------------------------------------
# SYNTHETIC DATA GENERATOR
# ----------------------------------------------------------------------------
FIRST_NAMES = ["James","Mary","Robert","Patricia","John","Jennifer","Michael","Linda","David","Elizabeth",
"William","Barbara","Richard","Susan","Joseph","Jessica","Thomas","Sarah","Charles","Karen",
"Ravi","Priya","Wei","Fatima","Carlos","Sophia","Liam","Olivia","Noah","Ava"]
LAST_NAMES = ["Smith","Johnson","Williams","Brown","Jones","Garcia","Miller","Davis","Rodriguez","Martinez",
"Lee","Perez","Thompson","White","Harris","Clark","Ramirez","Lewis","Robinson","Taylor"]
CITIES = ["New York","Toronto","London","San Francisco","Chicago","Boston","Seattle","Austin","Denver","Vancouver"]
DEPARTMENTS = ["Engineering","Sales","Marketing","Finance","HR","Operations","Legal","Support","Product","Design"]
INDUSTRIES = ["Technology","Finance","Healthcare","Retail","Manufacturing","Education","Energy","Media"]
PRODUCT_CATS = ["Software","Hardware","Consulting Services","Cloud Storage","Analytics Suite","Security Tools","Training"]
COMPANY_NAMES = ["Acme Corp","NovaTech","BlueRiver Inc","Summit Systems","Pinnacle Group","Quantum Works",
"Bright Path","Cedar Analytics","Horizon Labs","Vertex Solutions"]
PRODUCT_NAMES = ["DataSync Pro","CloudVault","SecureShield","InsightBoard","AutoFlow","GraphLens",
"StreamPipe","MetricHub","VisionAI","TeamSpace","CodeForge","NetGuard","PulseCRM","LedgerX","SwiftDeploy"]
def generate_synthetic_data(n_persons=30, n_edges=150, seed=None):
    rng = random.Random(seed)
    persons = [{"id": f"P{i:03d}", "name": f"{rng.choice(FIRST_NAMES)} {rng.choice(LAST_NAMES)}",
                "age": rng.randint(22, 62), "city": rng.choice(CITIES),
                "department": rng.choice(DEPARTMENTS)} for i in range(1, n_persons + 1)]
    companies = [{"id": f"C{i:02d}", "name": name, "industry": rng.choice(INDUSTRIES)}
                 for i, name in enumerate(COMPANY_NAMES, start=1)]
    products = [{"id": f"PR{i:02d}", "name": name, "category": rng.choice(PRODUCT_CATS),
                 "price": round(rng.uniform(49, 4999), 2)} for i, name in enumerate(PRODUCT_NAMES, start=1)]
    def rand_date():
        start = date(2023, 1, 1)
        return (start + timedelta(days=rng.randint(0, 900))).isoformat()
    rows, rid = [], 1
    for p in persons:
        c = rng.choice(companies)
        rows.append([rid, "Person", p["id"], p["name"], "WORKS_AT", "Company", c["id"], c["name"],
                     rng.randint(1, 15), rand_date()]); rid += 1
    while len(rows) < n_edges:
        choice = rng.random()
        if choice < 0.45 and len(persons) >= 2:
            p1, p2 = rng.sample(persons, 2)
            rows.append([rid, "Person", p1["id"], p1["name"], "KNOWS", "Person", p2["id"], p2["name"],
                         rng.randint(1, 10), rand_date()])
        elif choice < 0.8:
            p, pr = rng.choice(persons), rng.choice(products)
            rows.append([rid, "Person", p["id"], p["name"], "PURCHASED", "Product", pr["id"], pr["name"],
                         pr["price"], rand_date()])
        else:
            c, pr = rng.choice(companies), rng.choice(products)
            rows.append([rid, "Company", c["id"], c["name"], "SUPPLIES", "Product", pr["id"], pr["name"],
                         rng.randint(1, 500), rand_date()])
        rid += 1
    return pd.DataFrame(rows, columns=REQUIRED_COLS)
def build_graph(df):
    G = nx.DiGraph()
    for _, r in df.iterrows():
        G.add_node(r["source_id"], label=r["source_name"], type=r["source_type"])
        G.add_node(r["target_id"], label=r["target_name"], type=r["target_type"])
        G.add_edge(r["source_id"], r["target_id"], relationship=r["relationship"], weight=r["weight"])
    return G
NODE_COLORS = {"Person": "#1f77b4", "Company": "#ff7f0e", "Product": "#2ca02c"}
def plotly_network(G, max_nodes=120):
    if G.number_of_nodes() > max_nodes:
        nodes_subset = list(G.nodes)[:max_nodes]
        G = G.subgraph(nodes_subset)
    pos = nx.spring_layout(G, seed=42, k=0.6)
    edge_x, edge_y = [], []
    for u, v in G.edges():
        x0, y0 = pos[u]; x1, y1 = pos[v]
        edge_x += [x0, x1, None]; edge_y += [y0, y1, None]
    edge_trace = go.Scatter(x=edge_x, y=edge_y, line=dict(width=0.6, color="#AAAAAA"),
                             hoverinfo="none", mode="lines")
    node_x, node_y, node_color, node_text = [], [], [], []
    for n in G.nodes():
        x, y = pos[n]
        node_x.append(x); node_y.append(y)
        ntype = G.nodes[n].get("type", "Person")
        node_color.append(NODE_COLORS.get(ntype, "#999999"))
        node_text.append(f"{G.nodes[n].get('label', n)} ({ntype})<br>Connections: {G.degree(n)}")
    node_trace = go.Scatter(x=node_x, y=node_y, mode="markers", hoverinfo="text", text=node_text,
                             marker=dict(size=12, color=node_color, line=dict(width=1, color="#333")))
    fig = go.Figure(data=[edge_trace, node_trace],
                     layout=go.Layout(showlegend=False, hovermode="closest",
                                       margin=dict(l=10, r=10, t=10, b=10),
                                       xaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
                                       yaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
                                       height=520, plot_bgcolor="white"))
    return fig
# ----------------------------------------------------------------------------
# EXPORT HELPERS
# ----------------------------------------------------------------------------
def to_csv_bytes(df):
    return df.to_csv(index=False).encode("utf-8")
def to_txt_bytes(df, title):
    buf = io.StringIO()
    buf.write(f"{title}\nGenerated by GraphDB Insights Studio - Randy Singh, Kalsnet (KNet) Consulting\n")
    buf.write("=" * 80 + "\n\n")
    buf.write(df.to_string(index=False))
    return buf.getvalue().encode("utf-8")
def to_docx_bytes(df, title):
    doc = Document()
    h = doc.add_heading(title, level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x00, 0x33, 0xCC)
    p = doc.add_paragraph("Developed by Randy Singh from Kalsnet (KNet) Consulting")
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.color.rgb = RGBColor(0x00, 0x33, 0xCC)
    doc.add_paragraph(" ")
    cols = list(df.columns)
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Light Grid Accent 1"
    for i, c in enumerate(cols):
        table.rows[0].cells[i].text = str(c)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, c in enumerate(cols):
            cells[i].text = str(row[c])
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()
def to_pdf_bytes(df, title):
    bio = io.BytesIO()
    doc = SimpleDocTemplate(bio, pagesize=letter)
    styles = getSampleStyleSheet()
    title_style = styles["Heading1"]
    title_style.textColor = colors.HexColor("#0033CC")
    elems = [Paragraph(title, title_style),
             Paragraph("Developed by Randy Singh from Kalsnet (KNet) Consulting", styles["Heading3"]),
             Spacer(1, 12)]
    max_cols = 6
    cols = list(df.columns)[:max_cols]
    data = [cols] + df[cols].astype(str).values.tolist()
    data = [row[:max_cols] for row in data][:60]  # cap rows for a clean PDF page count
    t = Table(data, repeatRows=1)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0033CC")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTSIZE", (0, 0), (-1, -1), 7),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.white]),
    ]))
    elems.append(t)
    doc.build(elems)
    return bio.getvalue()
def export_block(df, key_prefix, title):
    st.markdown("**Export these results:**")
    c1, c2, c3, c4 = st.columns(4)
    with c1:
        st.download_button("CSV", to_csv_bytes(df), f"{key_prefix}.csv", "text/csv", key=key_prefix + "csv")
    with c2:
        st.download_button("TXT", to_txt_bytes(df, title), f"{key_prefix}.txt", "text/plain", key=key_prefix + "txt")
    with c3:
        st.download_button("Word", to_docx_bytes(df, title), f"{key_prefix}.docx",
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=key_prefix + "docx")
    with c4:
        st.download_button("PDF", to_pdf_bytes(df, title), f"{key_prefix}.pdf", "application/pdf",
                            key=key_prefix + "pdf")
# ----------------------------------------------------------------------------
# SIDEBAR — DATA SOURCE CONTROLS
# ----------------------------------------------------------------------------
st.sidebar.markdown("## Data Source")
data_mode = st.sidebar.radio("Choose data source:", ["Synthetic Data", "Upload Real Data"])
if data_mode == "Synthetic Data":
    st.sidebar.markdown("### Synthetic Data Generator")
    n_persons = st.sidebar.slider("Number of Person nodes", 10, 80, 30)
    n_edges = st.sidebar.slider("Number of relationship records", 50, 300, 150)
    seed = st.sidebar.number_input("Random seed", value=42, step=1)
    if st.sidebar.button("Generate Synthetic Data", type="primary"):
        st.session_state["df"] = generate_synthetic_data(n_persons, n_edges, seed)
    if "df" not in st.session_state:
        st.session_state["df"] = generate_synthetic_data(n_persons, n_edges, seed)
else:
    st.sidebar.markdown("### Upload Your Own Data")
    st.sidebar.caption("CSV must contain columns: " + ", ".join(REQUIRED_COLS))
    uploaded = st.sidebar.file_uploader("Upload relationship CSV", type=["csv"])
    if uploaded is not None:
        try:
            udf = pd.read_csv(uploaded)
            missing = [c for c in REQUIRED_COLS if c not in udf.columns]
            if missing:
                st.sidebar.error(f"Missing required columns: {missing}")
            else:
                st.session_state["df"] = udf
                st.sidebar.success(f"Loaded {len(udf)} records.")
        except Exception as e:
            st.sidebar.error(f"Could not read file: {e}")
    if "df" not in st.session_state:
        st.session_state["df"] = generate_synthetic_data(30, 150, 42)
        st.sidebar.info("No file uploaded yet — showing default synthetic data.")
df = st.session_state["df"]
G = build_graph(df)
st.sidebar.markdown("---")
st.sidebar.metric("Total Relationship Records", len(df))
st.sidebar.metric("Total Nodes (entities)", G.number_of_nodes())
st.sidebar.metric("Total Edges (relationships)", G.number_of_edges())
# ----------------------------------------------------------------------------
# TABS — USE CASES
# ----------------------------------------------------------------------------
tab1, tab2, tab3, tab4 = st.tabs([
    "1. Network Explorer",
    "2. Shortest Path Finder",
    "3. Influence & Centrality",
    "4. Relationship Analytics",
])
# --- TAB 1: Network Explorer ---
with tab1:
    st.markdown('<p class="section-header">Use Case 1: Interactive Network Explorer</p>', unsafe_allow_html=True)
    st.write("""This view renders the raw graph — every **node** (Person - blue, Company - orange, Product - green) and
    every **edge** (relationship) connecting them. In a real graph database this is equivalent to
    running `MATCH (n)-[r]-(m) RETURN n, r, m` and visualizing the result.""")
    fig1 = plotly_network(G)
    st.plotly_chart(fig1, use_container_width=True)
    st.dataframe(df, use_container_width=True, height=250)
    export_block(df, "network_explorer", "GraphDB Insights Studio - Network Explorer Results")
# --- TAB 2: Shortest Path ---
with tab2:
    st.markdown('<p class="section-header">Use Case 2: Shortest Path Finder</p>', unsafe_allow_html=True)
    st.write("""Graph databases excel at answering "how are these two things connected?" This use case
    finds the **shortest path** between any two nodes — a classic graph traversal query
    (`MATCH p = shortestPath((a)-[*]-(b))` in Cypher).""")
    node_options = {f"{G.nodes[n].get('label', n)} ({G.nodes[n].get('type', '')})": n for n in G.nodes()}
    colA, colB = st.columns(2)
    with colA:
        src_label = st.selectbox("From node:", list(node_options.keys()), index=0)
    with colB:
        dst_label = st.selectbox("To node:", list(node_options.keys()),
                                  index=min(1, len(node_options) - 1))
    src, dst = node_options[src_label], node_options[dst_label]
    path_df = pd.DataFrame()
    if st.button("Find Shortest Path"):
        try:
            UG = G.to_undirected()
            path = nx.shortest_path(UG, src, dst)
            st.success(f"Path found with {len(path) - 1} hop(s): " +
                       " → ".join(G.nodes[n].get("label", n) for n in path))
            path_rows = []
            for i in range(len(path) - 1):
                a, b = path[i], path[i + 1]
                rel = G.get_edge_data(a, b) or G.get_edge_data(b, a) or {}
                path_rows.append({"step": i + 1, "from": G.nodes[a].get("label", a),
                                   "relationship": rel.get("relationship", "CONNECTED"),
                                   "to": G.nodes[b].get("label", b)})
            path_df = pd.DataFrame(path_rows)
            st.table(path_df)
        except nx.NetworkXNoPath:
            st.warning("No path exists between these two nodes in the current graph.")
        except Exception as e:
            st.error(f"Error: {e}")
    if not path_df.empty:
        export_block(path_df, "shortest_path", "GraphDB Insights Studio - Shortest Path Results")
    else:
        st.info("Run a search above to enable export of the path results.")
# --- TAB 3: Centrality ---
with tab3:
    st.markdown('<p class="section-header">Use Case 3: Influence & Centrality Analysis</p>', unsafe_allow_html=True)
    st.write("""**Centrality** measures identify the most "important" nodes in a network — used for fraud
    detection, key-influencer identification, and recommendation systems.
    - **Degree centrality**: how many direct connections a node has.
    - **Betweenness centrality**: how often a node sits on the shortest path between others (a "broker").""")
    UG = G.to_undirected()
    deg_cent = nx.degree_centrality(UG)
    bet_cent = nx.betweenness_centrality(UG)
    cent_df = pd.DataFrame([
        {"node_id": n, "name": G.nodes[n].get("label", n), "type": G.nodes[n].get("type", ""),
         "degree_centrality": round(deg_cent[n], 4), "betweenness_centrality": round(bet_cent[n], 4),
         "connections": UG.degree(n)}
        for n in G.nodes()
    ]).sort_values("degree_centrality", ascending=False).reset_index(drop=True)
    top_n = st.slider("Show top N most-connected nodes", 5, 30, 10)
    top_df = cent_df.head(top_n)
    fig2 = px.bar(top_df, x="name", y="connections", color="type",
                  color_discrete_map=NODE_COLORS, title="Top Nodes by Number of Connections")
    fig2.update_layout(xaxis_tickangle=-40)
    st.plotly_chart(fig2, use_container_width=True)
    st.dataframe(cent_df, use_container_width=True, height=250)
    export_block(cent_df, "centrality_analysis", "GraphDB Insights Studio - Centrality Analysis Results")
# --- TAB 4: Relationship Analytics ---
with tab4:
    st.markdown('<p class="section-header">Use Case 4: Relationship Type Analytics</p>', unsafe_allow_html=True)
    st.write("""Aggregate view of relationship types across the graph — useful for understanding overall
    graph composition, e.g. how many `PURCHASED` vs `KNOWS` vs `WORKS_AT` edges exist.""")
    rel_counts = df["relationship"].value_counts().reset_index()
    rel_counts.columns = ["relationship", "count"]
    c1, c2 = st.columns(2)
    with c1:
        fig3 = px.pie(rel_counts, names="relationship", values="count", title="Relationship Type Distribution",
                       color_discrete_sequence=px.colors.qualitative.Bold)
        st.plotly_chart(fig3, use_container_width=True)
    with c2:
        type_counts = pd.concat([df["source_type"], df["target_type"]]).value_counts().reset_index()
        type_counts.columns = ["node_type", "count"]
        fig4 = px.bar(type_counts, x="node_type", y="count", title="Node Type Frequency in Relationships",
                       color="node_type", color_discrete_map=NODE_COLORS)
        st.plotly_chart(fig4, use_container_width=True)
    if "weight" in df.columns:
        purchase_df = df[df["relationship"] == "PURCHASED"]
        if not purchase_df.empty:
            fig5 = px.histogram(purchase_df, x="weight", nbins=20, title="Distribution of Purchase Amounts ($)")
            st.plotly_chart(fig5, use_container_width=True)
    summary_df = rel_counts.merge(
        df.groupby("relationship")["weight"].agg(["mean", "min", "max"]).reset_index(),
        on="relationship"
    ).round(2)
    summary_df.columns = ["relationship", "count", "avg_weight", "min_weight", "max_weight"]
    st.dataframe(summary_df, use_container_width=True)
    export_block(summary_df, "relationship_analytics", "GraphDB Insights Studio - Relationship Analytics Summary")
st.markdown("---")
st.caption("GraphDB Insights Studio - In-memory graph engine powered by NetworkX - "
           "Developed by Randy Singh, Kalsnet (KNet) Consulting")
