

# =========================================================
# AI Governance & Public Policy Intelligence Platform
# Professional Streamlit GUI with Groq Integration
# Developed by Randy Singh from Kalsnet (KNet)
# =========================================================

import streamlit as st
import json
import pandas as pd
from io import BytesIO
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from groq import Groq

# =========================================================
# PAGE CONFIG
# =========================================================

st.set_page_config(
    page_title="AI Governance Platform",
    layout="wide",
    page_icon="🏛️"
)

# =========================================================
# STYLING (PROFESSIONAL BLUE HEADER)
# =========================================================

st.markdown("""
<style>

body {
    background-color: #0E1117;
}

/* TITLE BAR */
.main-title {
    font-size: 40px;
    font-weight: 900;
    color: #1E90FF;
    text-align: center;
    padding: 10px;
}

.sub-title {
    text-align: center;
    color: #AAB4C0;
    font-size: 16px;
    margin-bottom: 20px;
}

/* BUTTON STYLE */
.stButton>button {
    background-color: #1E90FF;
    color: white;
    border-radius: 10px;
    font-weight: bold;
    padding: 0.5rem 1rem;
}

/* SIDEBAR */
section[data-testid="stSidebar"] {
    background-color: #111827;
    color: white;
}

</style>
""", unsafe_allow_html=True)

# =========================================================
# HEADER
# =========================================================

st.markdown('<div class="main-title">🏛️ AI Governance & Policy Intelligence Platform</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-title">Developed by Randy Singh from Kalsnet (KNet) Consulting Group</div>', unsafe_allow_html=True)

# =========================================================
# SIDEBAR - GROQ CONFIG
# =========================================================

st.sidebar.header("⚙️ Groq Configuration")

api_key = st.sidebar.text_input("Enter Groq API Key", type="password")

st.sidebar.markdown("""
### 🔑 How to get Groq API Key
1. Go to https://console.groq.com
2. Create a free account
3. Navigate to API Keys
4. Generate a new key
5. Paste it above
""")

model = st.sidebar.selectbox(
    "Select Model",
    ["llama-3.3-70b-versatile", "mixtral-8x7b-32768", "gemma2-9b-it"]
)

# =========================================================
# INPUT SECTION
# =========================================================

col1, col2 = st.columns(2)

with col1:
    org = st.text_input("Organization Name")
    policy_type = st.selectbox("Policy Type", [
        "Cybersecurity Policy",
        "AI Governance Policy",
        "Data Privacy Policy",
        "Incident Response Policy",
        "HR Policy"
    ])

with col2:
    industry = st.text_input("Industry")
    risk = st.selectbox("Risk Level", ["Low", "Medium", "High", "Critical"])

objective = st.text_area("Policy Objective / Requirement")

# =========================================================
# GROQ PROMPT
# =========================================================

def build_prompt():
    return f"""
Create a professional enterprise-grade policy document.

Organization: {org}
Industry: {industry}
Policy Type: {policy_type}
Risk Level: {risk}

Objective:
{objective}

Include:
- Purpose
- Scope
- Responsibilities
- Controls
- Compliance mapping
- Risk considerations
"""

# =========================================================
# RESET BUTTON
# =========================================================

if st.button("🔄 Reset"):
    st.session_state.clear()
    st.rerun()

# =========================================================
# GENERATE POLICY
# =========================================================

if st.button("🚀 Generate Policy", use_container_width=True):

    if not api_key:
        st.error("Please enter Groq API Key")
    else:
        client = Groq(api_key=api_key)

        with st.spinner("Generating policy using Groq AI..."):
            response = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": "You are a senior policy and governance expert."},
                    {"role": "user", "content": build_prompt()}
                ],
                temperature=0.3,
                max_tokens=4000
            )

        st.session_state["policy"] = response.choices[0].message.content

# =========================================================
# OUTPUT DISPLAY
# =========================================================

if "policy" in st.session_state:

    st.subheader("📄 Generated Policy Output")

    edited_text = st.text_area(
        "Edit Policy",
        st.session_state["policy"],
        height=400
    )

    st.session_state["policy"] = edited_text

    # =====================================================
    # EXPORT FUNCTIONS
    # =====================================================

    col1, col2, col3, col4 = st.columns(4)

    # ---------------- JSON ----------------
    with col1:
        st.download_button(
            "⬇️ JSON",
            data=json.dumps({"policy": edited_text}, indent=4),
            file_name="policy.json",
            mime="application/json"
        )

    # ---------------- CSV ----------------
    with col2:
        df = pd.DataFrame({"Policy": [edited_text]})
        st.download_button(
            "⬇️ CSV",
            data=df.to_csv(index=False),
            file_name="policy.csv",
            mime="text/csv"
        )

    # ---------------- DOCX ----------------
    with col3:
        doc = Document()
        doc.add_heading("AI Policy Document", 0)
        doc.add_paragraph(edited_text)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.download_button(
            "⬇️ Word",
            data=buffer,
            file_name="policy.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    # ---------------- PDF ----------------
    with col4:
        buffer = BytesIO()
        pdf = SimpleDocTemplate(buffer)

        styles = getSampleStyleSheet()
        story = [Paragraph(edited_text, styles["BodyText"])]

        pdf.build(story)
        buffer.seek(0)

        st.download_button(
            "⬇️ PDF",
            data=buffer,
            file_name="policy.pdf",
            mime="application/pdf"
        )

# =========================================================
# FILE UPLOAD SECTION
# =========================================================

st.subheader("📤 Upload Existing Policy File")

uploaded_file = st.file_uploader(
    "Upload TXT / JSON / CSV (basic support)",
    type=["txt", "json", "csv"]
)

if uploaded_file:
    file_type = uploaded_file.name.split(".")[-1]

    if file_type == "txt":
        content = uploaded_file.read().decode("utf-8")
    elif file_type == "json":
        content = json.dumps(json.load(uploaded_file), indent=4)
    elif file_type == "csv":
        content = uploaded_file.read().decode("utf-8")
    else:
        content = "Unsupported format"

    st.text_area("Uploaded Content", content, height=250)