

# CyberSentinel Intelligence Platform
# AI-powered cyber threat detection, threat intelligence, and predictive risk scoring
# Developed by Randy Singh | Kalsnet (KNet) Consulting
# EXPORT CAPABILITIES: PDF, Word (.docx), CSV, JSON, TXT
import io
import json
import logging
import os
import random
import uuid
import datetime as dt
from typing import Dict, List
import pandas as pd
import numpy as np
import streamlit as st

try:
    from fpdf import FPDF
    HAS_PDF = True
except ImportError:
    HAS_PDF = False
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# =========================================================================
# LOGGING
# =========================================================================
LOG_DIR = os.environ.get("CYBERSENTINEL_LOG_DIR", "./logs")
os.makedirs(LOG_DIR, exist_ok=True)
logger = logging.getLogger("cybersentinel")
if not logger.handlers:
    logger.setLevel(logging.INFO)
    _fh = logging.FileHandler(os.path.join(LOG_DIR, "cybersentinel.log"))
    _fh.setFormatter(logging.Formatter("%(asctime)s | %(levelname)s | %(message)s"))
    logger.addHandler(_fh)

# =========================================================================
# PAGE CONFIG
# =========================================================================
st.set_page_config(
    page_title="CyberSentinel Intelligence Platform",
    page_icon="",
    layout="wide",
    initial_sidebar_state="expanded",
)

# =========================================================================
# PROFESSIONAL STYLING
# =========================================================================
st.markdown(
    """
    <style>
    .main-title {
        text-align: center;
        color: #0047AB;
        font-weight: 900;
        font-size: 2.8rem;
        margin-bottom: 0.3rem;
        letter-spacing: 0.8px;
        font-family: 'Arial Black', sans-serif;
    }
    .subtitle-line {
        text-align: center;
        color: #0047AB;
        font-weight: 800;
        font-size: 1.2rem;
        margin-top: 0;
        margin-bottom: 1.2rem;
        font-family: 'Arial', sans-serif;
    }
    .section-header {
        color: #0047AB;
        font-weight: 800;
        font-size: 1.8rem;
        margin-top: 1.5rem;
        margin-bottom: 0.8rem;
        border-bottom: 3px solid #0047AB;
        padding-bottom: 0.5rem;
    }
    .subsection-header {
        color: #003380;
        font-weight: 700;
        font-size: 1.3rem;
        margin-top: 1rem;
        margin-bottom: 0.6rem;
    }
    .formula-box {
        background: #fff9e6;
        border: 2px solid #ffd700;
        padding: 1rem;
        border-radius: 6px;
        font-family: 'Courier New', monospace;
        font-size: 0.9rem;
        margin: 0.8rem 0;
    }
    .divider {
        border-top: 2px solid #0047AB;
        margin: 1.5rem 0;
        opacity: 0.3;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# =========================================================================
# TITLE BAR
# =========================================================================
st.markdown('<div class="main-title"> CyberSentinel Intelligence Platform</div>', unsafe_allow_html=True)
st.markdown('<div class="subtitle-line">Developed by Randy Singh from Kalsnet (KNet) Consulting</div>', unsafe_allow_html=True)
st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
st.markdown(
    """
    **Unified Cyber + Physical/Operational Threat Intelligence Platform** — Real-time AI for
    threat detection, IOC confidence scoring, predictive risk forecasting, and vulnerability/asset
    risk analysis across both digital and physical security domains.

    **Focus:** Detect faster • Prioritize smarter • Predict escalation before it happens — across every domain
    """
)

# =========================================================================
# ENVIRONMENT CONFIG
# =========================================================================
DEPLOY_ENV = os.environ.get("CYBERSENTINEL_ENV", "production")
MAX_UPLOAD_MB = int(os.environ.get("CYBERSENTINEL_MAX_UPLOAD_MB", "25"))

# =========================================================================
# CONSTANTS
# =========================================================================
CYBER_THREAT_TYPES = {
    "Malware Infection": {"severity": 9, "category": "Endpoint", "domain": "Cyber", "description": "Malicious software detected on host, may spread laterally"},
    "Ransomware": {"severity": 10, "category": "Endpoint", "domain": "Cyber", "description": "File encryption / extortion attempt, business-critical"},
    "Phishing Attempt": {"severity": 6, "category": "Email", "domain": "Cyber", "description": "Credential harvesting or malicious attachment via email"},
    "Brute Force Attack": {"severity": 7, "category": "Network", "domain": "Cyber", "description": "Repeated authentication attempts against a service"},
    "DDoS Attack": {"severity": 8, "category": "Network", "domain": "Cyber", "description": "Volumetric or application-layer availability attack"},
    "SQL Injection": {"severity": 9, "category": "Application", "domain": "Cyber", "description": "Malicious query injection against a data-backed application"},
    "Unauthorized Access (Digital)": {"severity": 8, "category": "Identity", "domain": "Cyber", "description": "Access outside of normal permission or behavior baseline"},
    "Data Exfiltration": {"severity": 10, "category": "Data Loss", "domain": "Cyber", "description": "Abnormal outbound transfer of sensitive data"},
    "Insider Threat (Digital)": {"severity": 8, "category": "Identity", "domain": "Cyber", "description": "Anomalous activity by an authenticated internal user"},
    "Port Scanning": {"severity": 3, "category": "Network", "domain": "Cyber", "description": "Reconnaissance activity probing open services"},
    "Privilege Escalation": {"severity": 9, "category": "Identity", "domain": "Cyber", "description": "Attempt to gain elevated system or account rights"},
    "Command & Control (C2) Beacon": {"severity": 10, "category": "Network", "domain": "Cyber", "description": "Outbound traffic consistent with known C2 infrastructure"},
}
PHYSICAL_THREAT_TYPES = {
    "Unauthorized Facility Access": {"severity": 8, "category": "Perimeter", "domain": "Physical/Operational", "description": "Badge tailgating, forced entry, or unrecognized credential at a controlled point"},
    "Perimeter Breach": {"severity": 8, "category": "Perimeter", "domain": "Physical/Operational", "description": "Fence, gate, or barrier compromise detected by sensor or patrol"},
    "Suspicious Package / Object": {"severity": 9, "category": "Safety", "domain": "Physical/Operational", "description": "Unattended or unidentified item reported near a facility or asset"},
    "Insider Threat (Physical)": {"severity": 8, "category": "Personnel", "domain": "Physical/Operational", "description": "Employee or contractor behavior inconsistent with role or access needs"},
    "Equipment Tampering / Sabotage": {"severity": 9, "category": "Operational", "domain": "Physical/Operational", "description": "Evidence of deliberate interference with machinery, utilities, or controls"},
    "Fire / Hazmat Incident": {"severity": 10, "category": "Safety", "domain": "Physical/Operational", "description": "Fire, chemical spill, or hazardous material exposure event"},
    "Workplace Violence Indicator": {"severity": 9, "category": "Personnel", "domain": "Physical/Operational", "description": "Threatening behavior, altercation, or escalation reported on-site"},
    "Supply Chain / Vendor Disruption": {"severity": 6, "category": "Operational", "domain": "Physical/Operational", "description": "Compromise or failure at a third-party vendor affecting operations"},
    "Surveillance / Reconnaissance Activity": {"severity": 5, "category": "Perimeter", "domain": "Physical/Operational", "description": "Repeated observation of a facility consistent with pre-attack planning"},
    "Utility / Environmental Failure": {"severity": 7, "category": "Operational", "domain": "Physical/Operational", "description": "Power, HVAC, or environmental control failure affecting critical operations"},
    "Lost/Stolen Access Credential": {"severity": 7, "category": "Personnel", "domain": "Physical/Operational", "description": "Badge, key, or physical token reported lost or stolen"},
    "Vehicle / Perimeter Ramming Concern": {"severity": 9, "category": "Perimeter", "domain": "Physical/Operational", "description": "Vehicle behavior or bollard/barrier alert near a facility perimeter"},
}
THREAT_TYPES = {**CYBER_THREAT_TYPES, **PHYSICAL_THREAT_TYPES}
DOMAIN_OPTIONS = ["Mixed (Cyber + Physical/Operational)", "Cyber Only", "Physical/Operational Only"]
ASSET_TYPES_CYBER = [
    "Web Server", "Database Server", "Domain Controller", "Endpoint-Workstation",
    "Firewall/Edge Device", "Email Server", "Cloud Storage Bucket", "IoT Device", "VPN Gateway"
]
ASSET_TYPES_PHYSICAL = [
    "Facility - Main Entrance", "Perimeter Fence/Gate", "Data Center (Physical)", "Warehouse/Loading Dock",
    "Executive Office", "Production Floor", "Utility/Mechanical Room", "Parking Structure", "Server Closet"
]
ASSET_TYPES = ASSET_TYPES_CYBER + ASSET_TYPES_PHYSICAL
STATUS_OPTIONS = ["Detected", "Investigating", "Contained", "Resolved", "False Positive"]
IOC_TYPES_CYBER = ["IP Address", "Domain", "File Hash (SHA-256)", "URL", "Email Address"]
IOC_TYPES_PHYSICAL = ["Badge/Credential ID", "Vehicle License Plate", "CCTV Detection Tag", "Access Control Alert", "Tip Line Report"]
IOC_TYPES = IOC_TYPES_CYBER + IOC_TYPES_PHYSICAL
SOURCE_FEEDS_CYBER = ["Internal SIEM", "Commercial Feed A", "Commercial Feed B", "OSINT Feed", "ISAC Sharing", "Government Advisory"]
SOURCE_FEEDS_PHYSICAL = ["Security Guard Report", "CCTV Analytics", "Access Control System", "Local Law Enforcement Tip", "Employee Report"]
SOURCE_FEEDS = SOURCE_FEEDS_CYBER + SOURCE_FEEDS_PHYSICAL

# MITRE ATT&CK mapping (cyber threat types only - physical/operational threats
# fall outside the ATT&CK Enterprise matrix and are labeled N/A)
MITRE_ATTACK_MAP = {
    "Malware Infection": {"tactic": "Execution", "technique_id": "T1204", "technique_name": "User Execution"},
    "Ransomware": {"tactic": "Impact", "technique_id": "T1486", "technique_name": "Data Encrypted for Impact"},
    "Phishing Attempt": {"tactic": "Initial Access", "technique_id": "T1566", "technique_name": "Phishing"},
    "Brute Force Attack": {"tactic": "Credential Access", "technique_id": "T1110", "technique_name": "Brute Force"},
    "DDoS Attack": {"tactic": "Impact", "technique_id": "T1498", "technique_name": "Network Denial of Service"},
    "SQL Injection": {"tactic": "Initial Access", "technique_id": "T1190", "technique_name": "Exploit Public-Facing Application"},
    "Unauthorized Access (Digital)": {"tactic": "Initial Access", "technique_id": "T1078", "technique_name": "Valid Accounts"},
    "Data Exfiltration": {"tactic": "Exfiltration", "technique_id": "T1041", "technique_name": "Exfiltration Over C2 Channel"},
    "Insider Threat (Digital)": {"tactic": "Collection", "technique_id": "T1213", "technique_name": "Data from Information Repositories"},
    "Port Scanning": {"tactic": "Reconnaissance", "technique_id": "T1595", "technique_name": "Active Scanning"},
    "Privilege Escalation": {"tactic": "Privilege Escalation", "technique_id": "T1068", "technique_name": "Exploitation for Privilege Escalation"},
    "Command & Control (C2) Beacon": {"tactic": "Command and Control", "technique_id": "T1071", "technique_name": "Application Layer Protocol"},
}

# Incident case management
CASE_PRIORITIES = {"Low": 72, "Medium": 24, "High": 8, "Critical": 2}  # SLA hours by priority
CASE_STATUSES = ["Open", "In Progress", "Escalated", "Closed"]
ANALYSTS = ["A. Rao", "J. Fernandes", "M. Chen", "S. Patel", "T. Okafor", "R. Novak"]

# Alert rules / notification simulation
ALERT_CHANNELS = ["Email", "SMS", "Pager", "SOC Dashboard"]

# Watchlist
WATCHLIST_TYPES = ["IP Address", "Zone/Facility Area", "Asset Type", "User/Employee"]
WATCHLIST_PRIORITIES = ["Standard", "Elevated", "High Priority"]

# =========================================================================
# SESSION STATE
# =========================================================================
DEFAULTS = {
    "synthetic_df": None,
    "real_df": None,
    "ioc_df": None,
    "case_df": None,
    "alert_log_df": None,
    "watchlist_df": None,
}
for k, v in DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = v
if "session_id" not in st.session_state:
    st.session_state.session_id = str(uuid.uuid4())[:8]

# =========================================================================
# THREAT INTELLIGENCE FORMULAS
# =========================================================================
class ThreatFormulas:
    """Collection of threat intelligence / prediction formulas with explanations"""

    @staticmethod
    def threat_risk_score(threat_type: str, exploitability: float, asset_criticality: float) -> float:
        """
        THREAT RISK SCORE (TRS)

        TRS = (Severity/10 × 0.4 + Exploitability/10 × 0.3 + Asset_Criticality/10 × 0.3) × 100

        Where:
        - Severity: Inherent severity of the threat type (1-10)
        - Exploitability: How easily this threat can be exploited right now (1-10)
        - Asset_Criticality: Business importance of the targeted asset (1-10)

        Interpretation:
        - Score < 30: LOW - Log and monitor
        - Score 30-60: MEDIUM - Investigate within SLA
        - Score 60-85: HIGH - Escalate to analyst immediately
        - Score > 85: CRITICAL - Incident response activation
        """
        if threat_type not in THREAT_TYPES:
            return 0
        severity = THREAT_TYPES[threat_type]["severity"]
        score = ((severity / 10) * 0.4 + (exploitability / 10) * 0.3 + (asset_criticality / 10) * 0.3) * 100
        return round(min(100, max(0, score)), 2)

    @staticmethod
    def ioc_confidence_score(source_reliability: float, corroboration_count: int, recency_days: int) -> float:
        """
        IOC CONFIDENCE SCORE

        Confidence = (Source_Reliability × 0.40 + Corroboration_Factor × 0.35 + Recency_Factor × 0.25) × 100

        Where:
        - Source_Reliability: 0-1 rating of the intel feed's historical accuracy
        - Corroboration_Factor: min(1, corroborating_sources / 5)
        - Recency_Factor: decays as the indicator ages (fresher = more confident)

        Interpretation:
        - < 40%: LOW - Informational only
        - 40-70%: MEDIUM - Monitor, do not auto-block
        - > 70%: HIGH - Actionable, safe to auto-block
        """
        corroboration_factor = min(1.0, corroboration_count / 5)
        recency_factor = max(0.0, 1 - (recency_days / 90))
        confidence = (source_reliability * 0.40 + corroboration_factor * 0.35 + recency_factor * 0.25) * 100
        return round(min(100, max(0, confidence)), 2)

    @staticmethod
    def escalation_probability(event_frequency_trend: float, avg_confidence: float, avg_severity: float) -> float:
        """
        PREDICTIVE THREAT ESCALATION PROBABILITY

        Escalation% = (Frequency_Trend × 0.5 + Confidence_Norm × 0.3 + Severity_Norm × 0.2) × 100

        Where:
        - Frequency_Trend: normalized rate-of-change in event volume over the observed window (0-1)
        - Confidence_Norm: average IOC confidence, normalized (0-1)
        - Severity_Norm: average threat severity, normalized (0-1)

        This estimates the likelihood a cluster of related events escalates into
        a confirmed incident in the next observation window, based on momentum
        rather than a single point-in-time score.

        Interpretation:
        - < 25%: Unlikely to escalate
        - 25-50%: Possible - keep watching
        - 50-75%: Likely - proactively prepare response
        - > 75%: Highly likely - escalate now
        """
        freq_norm = min(1.0, max(0.0, event_frequency_trend))
        conf_norm = min(1.0, max(0.0, avg_confidence / 100))
        sev_norm = min(1.0, max(0.0, avg_severity / 10))
        escalation = (freq_norm * 0.5 + conf_norm * 0.3 + sev_norm * 0.2) * 100
        return round(min(100, max(0, escalation)), 2)

    @staticmethod
    def attack_surface_risk_index(open_ports: int, known_vulns: int, days_since_patch: int, internet_facing: bool) -> float:
        """
        ATTACK SURFACE RISK INDEX (ASRI)

        ASRI = (Port_Score × 0.25 + Vuln_Score × 0.40 + Patch_Score × 0.20 + Exposure_Score × 0.15)

        Where:
        - Port_Score: min(10, open_ports / 2)
        - Vuln_Score: min(10, known_vulns × 1.5)
        - Patch_Score: min(10, days_since_patch / 30)
        - Exposure_Score: 10 if internet-facing else 3

        Interpretation (0-10 scale):
        - 0-3: Low exposure
        - 3-6: Moderate - schedule hardening
        - 6-8: High - prioritize remediation
        - 8-10: Critical - immediate remediation required
        """
        port_score = min(10, open_ports / 2)
        vuln_score = min(10, known_vulns * 1.5)
        patch_score = min(10, days_since_patch / 30)
        exposure_score = 10 if internet_facing else 3
        asri = port_score * 0.25 + vuln_score * 0.40 + patch_score * 0.20 + exposure_score * 0.15
        return round(min(10, max(0, asri)), 2)

    @staticmethod
    def sla_hours_remaining(created: dt.datetime, sla_hours: float) -> float:
        """
        SLA HOURS REMAINING

        Remaining = SLA_Hours - (Now - Created_At in hours)

        A negative value means the case has breached its SLA window.
        """
        due = created + dt.timedelta(hours=sla_hours)
        remaining = (due - dt.datetime.now()).total_seconds() / 3600
        return round(remaining, 1)


# =========================================================================
# EXPORT FUNCTIONS
# =========================================================================
_PDF_CHAR_MAP = {
    "🟢": "[LOW] ", "🟡": "[MEDIUM] ", "🟠": "[HIGH] ", "🔴": "[CRITICAL] ",
    "✅": "[OK] ", "⚠️": "[WARNING] ", "⚠": "[WARNING] ", "❌": "[X] ",
    "🛡️": "", "🔄": "", "📊": "", "📋": "", "📄": "", "📝": "", "🔧": "", "🔮": "", "⚙️": "", "📈": "", "🌐": "",
    "₹": "Rs. ", "–": "-", "—": "-",
    "’": "'", "‘": "'", "“": '"', "”": '"', "…": "...",
}

def _pdf_safe(value) -> str:
    """Make text safe for FPDF's Latin-1-only core fonts."""
    text = str(value)
    for original, replacement in _PDF_CHAR_MAP.items():
        text = text.replace(original, replacement)
    return text.encode("latin-1", "ignore").decode("latin-1").strip()

def export_to_pdf(df: pd.DataFrame, title: str, formulas: Dict = None) -> bytes:
    if not HAS_PDF:
        return None
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", "B", 16)
        pdf.cell(0, 10, _pdf_safe(title), ln=True, align="C")
        pdf.set_font("Arial", "B", 10)
        pdf.cell(0, 5, f"Generated: {dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", ln=True)
        pdf.cell(0, 5, "CyberSentinel Intelligence Platform | Kalsnet (KNet) Consulting", ln=True)
        pdf.set_font("Arial", "", 8)
        pdf.ln(5)

        col_width = 190 / len(df.columns)
        row_height = 7
        page_bottom = pdf.h - pdf.b_margin

        def draw_header():
            pdf.set_font("Arial", "B", 8)
            for col in df.columns:
                pdf.cell(col_width, row_height, _pdf_safe(col)[:20], border=1, align="C")
            pdf.ln()
            pdf.set_font("Arial", "", 8)

        draw_header()
        for _, row in df.head(50).iterrows():
            if pdf.get_y() + row_height > page_bottom:
                pdf.add_page()
                draw_header()
            for val in row:
                pdf.cell(col_width, row_height, _pdf_safe(val)[:20], border=1, align="L")
            pdf.ln()

        if formulas:
            pdf.add_page()
            pdf.set_font("Arial", "B", 12)
            pdf.cell(0, 10, "Formulas & Explanations", ln=True)
            pdf.set_font("Arial", "", 9)
            for formula_name, formula_text in formulas.items():
                pdf.cell(0, 5, _pdf_safe(formula_name) + ":", ln=True)
                pdf.multi_cell(0, 4, _pdf_safe(formula_text)[:500])
                pdf.ln(2)

        raw_output = pdf.output(dest="S")
        if isinstance(raw_output, str):
            return raw_output.encode("latin-1")
        return bytes(raw_output)
    except Exception as e:
        st.error(f"PDF export error: {e}")
        logger.exception("PDF export failed")
        return None

def export_to_docx(df: pd.DataFrame, title: str, formulas: Dict = None) -> bytes:
    if not HAS_DOCX:
        return None
    try:
        doc = Document()
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 71, 171)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        meta_para = doc.add_paragraph()
        meta_run = meta_para.add_run(f"Generated: {dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        meta_run.font.size = Pt(10)
        meta_run = meta_para.add_run("CyberSentinel Intelligence Platform | Developed by Randy Singh | Kalsnet (KNet) Consulting")
        meta_run.font.size = Pt(10)
        meta_run.italic = True

        doc.add_paragraph()
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr_cells[i].text = str(col)
        for _, row in df.head(50).iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val)[:100]

        if formulas:
            doc.add_page_break()
            formulas_heading = doc.add_paragraph()
            formulas_run = formulas_heading.add_run("Formulas & Explanations")
            formulas_run.font.size = Pt(14)
            formulas_run.font.bold = True
            formulas_run.font.color.rgb = RGBColor(0, 71, 171)
            for formula_name, formula_text in formulas.items():
                doc.add_paragraph(formula_name, style='Heading 3')
                doc.add_paragraph(formula_text[:500])

        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        return docx_bytes.getvalue()
    except Exception as e:
        st.error(f"Word export error: {e}")
        return None

def export_to_text(df: pd.DataFrame, title: str, formulas: Dict = None) -> str:
    text = f"\n{'='*80}\n{title}\n{'='*80}\n\n"
    text += f"Generated: {dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
    text += "CyberSentinel Intelligence Platform | Developed by Randy Singh | Kalsnet (KNet) Consulting\n\n"
    text += df.to_string()
    if formulas:
        text += f"\n\n{'='*80}\nFormulas & Explanations\n{'='*80}\n\n"
        for formula_name, formula_text in formulas.items():
            text += f"\n{formula_name}:\n{formula_text}\n"
    return text

def export_report_to_pdf(narrative: str, title: str) -> bytes:
    """Export a plain-text narrative report (no dataframe) to PDF."""
    if not HAS_PDF:
        return None
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", "B", 16)
        pdf.cell(0, 10, _pdf_safe(title), ln=True, align="C")
        pdf.set_font("Arial", "", 9)
        pdf.cell(0, 5, f"Generated: {dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", ln=True)
        pdf.cell(0, 5, "CyberSentinel Intelligence Platform | Kalsnet (KNet) Consulting", ln=True)
        pdf.ln(4)
        pdf.set_font("Arial", "", 10)
        for line in narrative.split("\n"):
            pdf.multi_cell(0, 5, _pdf_safe(line))
        raw_output = pdf.output(dest="S")
        if isinstance(raw_output, str):
            return raw_output.encode("latin-1")
        return bytes(raw_output)
    except Exception as e:
        st.error(f"PDF export error: {e}")
        return None

def export_report_to_docx(narrative: str, title: str) -> bytes:
    """Export a plain-text narrative report (no dataframe) to Word."""
    if not HAS_DOCX:
        return None
    try:
        doc = Document()
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 71, 171)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        meta_para = doc.add_paragraph()
        meta_run = meta_para.add_run(f"Generated: {dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        meta_run.font.size = Pt(10)
        meta_run = meta_para.add_run("CyberSentinel Intelligence Platform | Developed by Randy Singh | Kalsnet (KNet) Consulting")
        meta_run.font.size = Pt(10)
        meta_run.italic = True
        doc.add_paragraph()

        for block in narrative.split("\n\n"):
            lines = block.split("\n")
            if lines and lines[0].endswith(":") and len(lines[0]) < 60:
                doc.add_paragraph(lines[0], style="Heading 3")
                for l in lines[1:]:
                    doc.add_paragraph(l)
            else:
                doc.add_paragraph(block)

        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        return docx_bytes.getvalue()
    except Exception as e:
        st.error(f"Word export error: {e}")
        return None

def build_executive_summary(df: pd.DataFrame, case_df: pd.DataFrame, alert_df: pd.DataFrame) -> str:
    """Compose a plain-language narrative summary from current session data."""
    lines = []
    lines.append("EXECUTIVE SUMMARY:")
    if df is None or len(df) == 0:
        lines.append("No threat detection data is currently loaded. Generate synthetic data or upload real events in the Threat Detection tab to populate this report.")
        return "\n\n".join(lines)

    total = len(df)
    avg_score = df["Risk_Score"].mean() if "Risk_Score" in df.columns else 0
    critical_ct = (df.get("Risk_Level", pd.Series(dtype=str)) == "🔴 CRITICAL").sum() if "Risk_Level" in df.columns else 0
    high_ct = (df.get("Risk_Level", pd.Series(dtype=str)) == "🟠 HIGH").sum() if "Risk_Level" in df.columns else 0
    cyber_ct = (df.get("Domain", pd.Series(dtype=str)) == "Cyber").sum() if "Domain" in df.columns else 0
    phys_ct = (df.get("Domain", pd.Series(dtype=str)) == "Physical/Operational").sum() if "Domain" in df.columns else 0

    lines.append(f"During the observed period, {total} threat events were recorded across the environment, "
                 f"with an average risk score of {avg_score:.1f}/100. Of these, {critical_ct} were rated CRITICAL "
                 f"and {high_ct} were rated HIGH, indicating a combined {critical_ct + high_ct} events requiring "
                 f"immediate or near-term analyst attention.")

    lines.append("DOMAIN BREAKDOWN:")
    lines.append(f"Cyber-domain events accounted for {cyber_ct} of {total} total events, while "
                 f"physical/operational events accounted for {phys_ct}. This split should inform whether "
                 f"security investment is currently balanced across both domains.")

    if "Threat_Type" in df.columns:
        top_threats = df["Threat_Type"].value_counts().head(3)
        lines.append("TOP THREAT TYPES:")
        for ttype, cnt in top_threats.items():
            lines.append(f"- {ttype}: {cnt} occurrences")

    if "Asset_Type" in df.columns:
        top_assets = df["Asset_Type"].value_counts().head(3)
        lines.append("MOST TARGETED ASSETS:")
        for atype, cnt in top_assets.items():
            lines.append(f"- {atype}: {cnt} occurrences")

    if case_df is not None and len(case_df) > 0:
        open_cases = (case_df["Status"].isin(["Open", "In Progress", "Escalated"])).sum()
        overdue = (case_df["Hours_Remaining"] < 0).sum() if "Hours_Remaining" in case_df.columns else 0
        lines.append("CASE MANAGEMENT STATUS:")
        lines.append(f"{open_cases} cases are currently open across the team, of which {overdue} have breached "
                     f"their SLA window and should be reprioritized immediately.")

    if alert_df is not None and len(alert_df) > 0:
        lines.append("ALERTING ACTIVITY:")
        lines.append(f"{len(alert_df)} alerts were triggered by configured rules during this period, "
                     f"distributed across {alert_df['Channel'].nunique() if 'Channel' in alert_df.columns else 'multiple'} notification channels.")

    lines.append("RECOMMENDED ACTIONS:")
    recs = []
    if critical_ct > 0:
        recs.append(f"Prioritize triage of the {critical_ct} CRITICAL-rated event(s) before end of shift.")
    if phys_ct > cyber_ct:
        recs.append("Physical/operational events currently outweigh cyber events - confirm physical security staffing matches this volume.")
    elif cyber_ct > phys_ct:
        recs.append("Cyber events currently outweigh physical events - confirm SOC staffing matches this volume.")
    if case_df is not None and len(case_df) > 0 and (case_df["Hours_Remaining"] < 0).sum() > 0:
        recs.append("Reassign or escalate all SLA-breached cases immediately.")
    if not recs:
        recs.append("No urgent action items identified from current data; continue routine monitoring.")
    for r in recs:
        lines.append(f"- {r}")

    return "\n\n".join(lines)

def render_export_row(df: pd.DataFrame, title: str, formulas_dict: Dict, key_prefix: str):
    col1, col2, col3, col4, col5 = st.columns(5)
    with col1:
        csv = df.to_csv(index=False)
        st.download_button("📊 CSV", csv.encode(), f"{key_prefix}.csv", "text/csv", key=f"{key_prefix}_csv")
    with col2:
        json_str = df.to_json(orient="records", indent=2)
        st.download_button("📋 JSON", json_str.encode(), f"{key_prefix}.json", "application/json", key=f"{key_prefix}_json")
    with col3:
        txt = export_to_text(df, title, formulas_dict)
        st.download_button("📄 TXT", txt.encode(), f"{key_prefix}.txt", "text/plain", key=f"{key_prefix}_txt")
    with col4:
        if HAS_DOCX:
            docx_bytes = export_to_docx(df, title, formulas_dict)
            if docx_bytes:
                st.download_button("📝 WORD", docx_bytes, f"{key_prefix}.docx",
                                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key=f"{key_prefix}_docx")
        else:
            st.write("Word export unavailable")
    with col5:
        if HAS_PDF:
            pdf_bytes = export_to_pdf(df, title, formulas_dict)
            if pdf_bytes:
                st.download_button("🔴 PDF", pdf_bytes, f"{key_prefix}.pdf", "application/pdf", key=f"{key_prefix}_pdf")
        else:
            st.write("PDF export unavailable")

# =========================================================================
# SIDEBAR
# =========================================================================
with st.sidebar:
    st.header("⚙️ Configuration")
    with st.expander("📊 Data Settings", expanded=True):
        st.markdown(f"**Max Upload**: {MAX_UPLOAD_MB} MB")
        st.markdown(f"**Environment**: {DEPLOY_ENV.upper()}")
        st.markdown(f"**Session**: {st.session_state.session_id}")
    with st.expander("ℹ️ About", expanded=False):
        st.markdown("""
        CyberSentinel combines rule-based threat scoring, IOC confidence
        rating, and trend-based escalation prediction into a single
        analyst workflow. All scoring logic is transparent and documented
        under each tab's formula box.
        """)
    st.markdown("---")
    st.caption("v1.0 | Kalsnet (KNet) Consulting")

# =========================================================================
# MAIN INTERFACE - TABS
# =========================================================================
tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8, tab9 = st.tabs([
    "🚨 Threat Detection",
    "🌐 Threat Intelligence (IOC)",
    "🔮 Predictive Forecasting",
    "🖥️ Asset & Vulnerability Risk",
    "📈 Dashboard & Analytics",
    "🗂️ Case Management",
    "🔔 Alert Rules",
    "⭐ Watchlist",
    "🧾 Executive Summary",
])

# =========================================================================
# TAB 1: THREAT DETECTION & SCORING
# =========================================================================
with tab1:
    st.markdown('<div class="section-header">🚨 Threat Detection & Risk Scoring</div>', unsafe_allow_html=True)
    st.markdown("""
    Score inbound security events by combining threat severity, exploitability, and the
    business criticality of the targeted asset into a single prioritized risk score.
    """)

    col1, col2, col3, col4 = st.columns([2, 2, 1, 1])
    with col1:
        n_records = st.slider("Synthetic records", 0, 500, 150, step=10)
    with col2:
        domain_choice = st.selectbox("Threat domain", DOMAIN_OPTIONS)
    with col3:
        st.write(""); st.write("")
        generate = st.button("🔄 Generate Demo", key="gen_demo")
    with col4:
        st.write(""); st.write("")
        reset_demo = st.button("🗑️ Reset", key="reset_demo")

    if reset_demo:
        st.session_state.synthetic_df = None
        st.success("✅ Synthetic threat data reset")

    if generate:
        if domain_choice == "Cyber Only":
            pool, asset_pool = CYBER_THREAT_TYPES, ASSET_TYPES_CYBER
        elif domain_choice == "Physical/Operational Only":
            pool, asset_pool = PHYSICAL_THREAT_TYPES, ASSET_TYPES_PHYSICAL
        else:
            pool, asset_pool = THREAT_TYPES, ASSET_TYPES

        rows = []
        for i in range(n_records):
            threat_type = random.choice(list(pool.keys()))
            domain = THREAT_TYPES[threat_type]["domain"]
            exploitability = round(random.uniform(1, 10), 1)
            asset_criticality = round(random.uniform(1, 10), 1)
            score = ThreatFormulas.threat_risk_score(threat_type, exploitability, asset_criticality)

            if score < 30:
                risk_level = "🟢 LOW"
            elif score < 60:
                risk_level = "🟡 MEDIUM"
            elif score < 85:
                risk_level = "🟠 HIGH"
            else:
                risk_level = "🔴 CRITICAL"

            # Physical events use a facility zone reference instead of an IP address
            if domain == "Cyber":
                origin = f"{random.randint(1,223)}.{random.randint(0,255)}.{random.randint(0,255)}.{random.randint(1,254)}"
                mitre = MITRE_ATTACK_MAP.get(threat_type, {})
                mitre_technique = f"{mitre.get('technique_id','N/A')} - {mitre.get('technique_name','N/A')}"
                mitre_tactic = mitre.get("tactic", "N/A")
            else:
                origin = f"ZONE-{random.choice(['A','B','C','D'])}{random.randint(1,12)}"
                mitre_technique = "N/A - Physical Domain"
                mitre_tactic = "N/A - Physical Domain"

            rows.append({
                "Event_ID": f"EVT-{random.randint(100000, 999999)}",
                "Timestamp": (dt.datetime.now() - dt.timedelta(hours=random.randint(0, 720))).strftime("%Y-%m-%d %H:%M"),
                "Domain": domain,
                "Source_IP_or_Zone": origin,
                "Threat_Type": threat_type,
                "Category": THREAT_TYPES[threat_type]["category"],
                "MITRE_Tactic": mitre_tactic,
                "MITRE_Technique": mitre_technique,
                "Asset_Type": random.choice(asset_pool if asset_pool else ASSET_TYPES),
                "Exploitability": exploitability,
                "Asset_Criticality": asset_criticality,
                "Risk_Score": score,
                "Risk_Level": risk_level,
                "Status": random.choice(STATUS_OPTIONS),
            })
        st.session_state.synthetic_df = pd.DataFrame(rows)
        st.success(f"✅ Generated {n_records} synthetic threat events ({domain_choice})")

    st.markdown("**OR upload real security event data:**")
    uploaded = st.file_uploader("Upload CSV/Excel file", type=["csv", "xlsx"], key="upload_events")
    if uploaded:
        try:
            if uploaded.name.endswith(".csv"):
                st.session_state.real_df = pd.read_csv(uploaded)
            else:
                st.session_state.real_df = pd.read_excel(uploaded)
            st.success(f"✅ Loaded {len(st.session_state.real_df)} records")
        except Exception as e:
            st.error(f"Error: {e}")

    df_show = st.session_state.real_df if st.session_state.real_df is not None else st.session_state.synthetic_df

    if df_show is not None:
        st.markdown("**Data Preview:**")
        st.dataframe(df_show.head(20), use_container_width=True, height=250)

        st.markdown('<div class="subsection-header">📈 Analysis Summary</div>', unsafe_allow_html=True)
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Total Events", len(df_show))
        with col2:
            critical = (df_show.get("Risk_Level", pd.Series(dtype=str)) == "🔴 CRITICAL").sum() if "Risk_Level" in df_show.columns else 0
            st.metric("Critical", critical)
        with col3:
            high = (df_show.get("Risk_Level", pd.Series(dtype=str)) == "🟠 HIGH").sum() if "Risk_Level" in df_show.columns else 0
            st.metric("High", high)
        with col4:
            if "Risk_Score" in df_show.columns:
                st.metric("Avg Risk Score", f"{df_show['Risk_Score'].mean():.1f}")

        if "Domain" in df_show.columns:
            st.markdown('<div class="subsection-header">Cyber vs Physical/Operational Split</div>', unsafe_allow_html=True)
            domain_summary = df_show.groupby("Domain").size().reset_index(name="Count")
            st.bar_chart(domain_summary.set_index("Domain"))

        if "Threat_Type" in df_show.columns:
            st.markdown('<div class="subsection-header">Threat Type Breakdown</div>', unsafe_allow_html=True)
            threat_summary = df_show.groupby("Threat_Type").size().reset_index(name="Count")
            st.bar_chart(threat_summary.set_index("Threat_Type"))

        if "Status" in df_show.columns:
            st.markdown('<div class="subsection-header">Status Breakdown</div>', unsafe_allow_html=True)
            status_summary = df_show.groupby("Status").size().reset_index(name="Count")
            st.bar_chart(status_summary.set_index("Status"))

        st.markdown('<div class="formula-box">', unsafe_allow_html=True)
        st.markdown("""
        **THREAT RISK SCORE (TRS) FORMULA**

        **TRS = (Severity/10 × 0.4 + Exploitability/10 × 0.3 + Asset_Criticality/10 × 0.3) × 100**

        **Severity Weights:**
        """)
        st.markdown("**🌐 Cyber Threats:**")
        for ttype, info in CYBER_THREAT_TYPES.items():
            st.write(f"- **{ttype}** ({info['severity']}/10, {info['category']}): {info['description']}")
        st.markdown("**🏭 Physical/Operational Threats:**")
        for ttype, info in PHYSICAL_THREAT_TYPES.items():
            st.write(f"- **{ttype}** ({info['severity']}/10, {info['category']}): {info['description']}")
        st.markdown("""
        **Interpretation:**
        - Score < 30: 🟢 LOW - Log and monitor
        - Score 30-60: 🟡 MEDIUM - Investigate within SLA
        - Score 60-85: 🟠 HIGH - Escalate to analyst immediately
        - Score > 85: 🔴 CRITICAL - Incident response activation
        """)
        st.markdown('</div>', unsafe_allow_html=True)

        with st.expander("🎯 MITRE ATT&CK Technique Mapping (Cyber Threats)"):
            mitre_rows = [
                {"Threat_Type": k, "Tactic": v["tactic"], "Technique_ID": v["technique_id"], "Technique_Name": v["technique_name"]}
                for k, v in MITRE_ATTACK_MAP.items()
            ]
            st.dataframe(pd.DataFrame(mitre_rows), use_container_width=True, hide_index=True)
            st.caption("Physical/operational threats fall outside the ATT&CK Enterprise matrix and are tagged N/A.")
            if "MITRE_Technique" in df_show.columns:
                st.markdown("**Technique Frequency in Current Data:**")
                tech_summary = df_show[df_show["MITRE_Technique"] != "N/A - Physical Domain"].groupby("MITRE_Technique").size().reset_index(name="Count")
                if len(tech_summary) > 0:
                    st.bar_chart(tech_summary.set_index("MITRE_Technique"))

        st.markdown('<div class="subsection-header">💾 Export Results</div>', unsafe_allow_html=True)
        formulas_dict = {
            "Threat Risk Score": "TRS = (Severity/10 x 0.4 + Exploitability/10 x 0.3 + Asset_Criticality/10 x 0.3) x 100",
            "Threat Severity Reference": "\n".join([f"{k}: {v['severity']}/10 - {v['description']}" for k, v in THREAT_TYPES.items()]),
        }
        render_export_row(df_show, "Threat Detection Analysis", formulas_dict, "threat_detect")

# =========================================================================
# TAB 2: THREAT INTELLIGENCE (IOC)
# =========================================================================
with tab2:
    st.markdown('<div class="section-header">🌐 Threat Intelligence & IOC Confidence</div>', unsafe_allow_html=True)
    st.markdown("""
    Rate incoming Indicators of Compromise (IOCs) from intelligence feeds by combining
    source reliability, cross-source corroboration, and recency into a single confidence score.
    """)

    col1, col2, col3, col4 = st.columns([2, 2, 1, 1])
    with col1:
        n_iocs = st.slider("Synthetic IOC records", 0, 500, 120, step=10)
    with col2:
        ioc_domain_choice = st.selectbox("Intel domain", DOMAIN_OPTIONS, key="ioc_domain")
    with col3:
        st.write(""); st.write("")
        gen_ioc = st.button("🔄 Generate", key="gen_ioc")
    with col4:
        st.write(""); st.write("")
        reset_ioc = st.button("🗑️ Reset", key="reset_ioc")

    if reset_ioc:
        st.session_state.ioc_df = None
        st.success("✅ IOC data reset")

    if gen_ioc:
        if ioc_domain_choice == "Cyber Only":
            type_pool, feed_pool, dom = IOC_TYPES_CYBER, SOURCE_FEEDS_CYBER, "Cyber"
        elif ioc_domain_choice == "Physical/Operational Only":
            type_pool, feed_pool, dom = IOC_TYPES_PHYSICAL, SOURCE_FEEDS_PHYSICAL, "Physical/Operational"
        else:
            type_pool, feed_pool, dom = IOC_TYPES, SOURCE_FEEDS, None

        rows = []
        for i in range(n_iocs):
            ioc_type = random.choice(type_pool)
            source_feed = random.choice(feed_pool)
            domain = dom if dom else ("Cyber" if ioc_type in IOC_TYPES_CYBER else "Physical/Operational")
            source_reliability = round(random.uniform(0.3, 1.0), 2)
            corroboration = random.randint(0, 8)
            recency = random.randint(0, 120)
            confidence = ThreatFormulas.ioc_confidence_score(source_reliability, corroboration, recency)

            if confidence < 40:
                action = "🟢 Informational"
            elif confidence < 70:
                action = "🟡 Monitor"
            else:
                action = "🔴 Actionable - Block"

            rows.append({
                "IOC_ID": f"IOC-{random.randint(10000, 99999)}",
                "Domain": domain,
                "IOC_Type": ioc_type,
                "Source_Feed": source_feed,
                "Source_Reliability": source_reliability,
                "Corroborating_Sources": corroboration,
                "Age_Days": recency,
                "Confidence_%": confidence,
                "Recommended_Action": action,
            })
        st.session_state.ioc_df = pd.DataFrame(rows)
        st.success(f"✅ Generated {n_iocs} synthetic IOC records ({ioc_domain_choice})")

    if st.session_state.ioc_df is not None:
        ioc_df = st.session_state.ioc_df
        st.dataframe(ioc_df, use_container_width=True, height=250)

        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Total IOCs", len(ioc_df))
        with col2:
            st.metric("Avg Confidence", f"{ioc_df['Confidence_%'].mean():.1f}%")
        with col3:
            actionable = (ioc_df["Recommended_Action"] == "🔴 Actionable - Block").sum()
            st.metric("Actionable", actionable)
        with col4:
            informational = (ioc_df["Recommended_Action"] == "🟢 Informational").sum()
            st.metric("Informational", informational)

        st.markdown('<div class="subsection-header">Source Feed Breakdown</div>', unsafe_allow_html=True)
        feed_summary = ioc_df.groupby("Source_Feed").size().reset_index(name="Count")
        st.bar_chart(feed_summary.set_index("Source_Feed"))

        st.markdown('<div class="formula-box">', unsafe_allow_html=True)
        st.markdown("""
        **IOC CONFIDENCE SCORE FORMULA**

        **Confidence = (Source_Reliability × 0.40 + Corroboration_Factor × 0.35 + Recency_Factor × 0.25) × 100**

        - Corroboration_Factor = min(1, corroborating_sources / 5)
        - Recency_Factor = max(0, 1 - (age_days / 90))

        **Interpretation:**
        - < 40%: 🟢 Informational only
        - 40-70%: 🟡 Monitor, do not auto-block
        - > 70%: 🔴 Actionable, safe to auto-block
        """)
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown('<div class="subsection-header">💾 Export Results</div>', unsafe_allow_html=True)
        formulas_dict = {
            "IOC Confidence Score": "Confidence = (Source_Reliability x 0.40 + Corroboration_Factor x 0.35 + Recency_Factor x 0.25) x 100",
        }
        render_export_row(ioc_df, "Threat Intelligence IOC Analysis", formulas_dict, "ioc_intel")

# =========================================================================
# TAB 3: PREDICTIVE FORECASTING
# =========================================================================
with tab3:
    st.markdown('<div class="section-header">🔮 Predictive Threat Escalation Forecasting</div>', unsafe_allow_html=True)
    st.markdown("""
    Estimate the probability that a currently-observed cluster of threat activity escalates
    into a confirmed incident, based on event momentum rather than a single point-in-time score.
    """)

    col1, col2, col3 = st.columns(3)
    with col1:
        freq_trend = st.slider("Event Frequency Trend (0=flat, 1=rapidly rising)", 0.0, 1.0, 0.4, step=0.05)
    with col2:
        avg_conf = st.slider("Average IOC Confidence (%)", 0, 100, 55, step=5)
    with col3:
        avg_sev = st.slider("Average Threat Severity (1-10)", 1.0, 10.0, 6.0, step=0.5)

    escalation = ThreatFormulas.escalation_probability(freq_trend, avg_conf, avg_sev)

    if escalation < 25:
        status, action, color = "🟢 UNLIKELY", "Continue routine monitoring", "#1a8a3e"
    elif escalation < 50:
        status, action, color = "🟡 POSSIBLE", "Keep watching, review in next shift", "#ff9800"
    elif escalation < 75:
        status, action, color = "🟠 LIKELY", "Proactively prepare incident response", "#ffa500"
    else:
        status, action, color = "🔴 HIGHLY LIKELY", "Escalate to incident response now", "#c0392b"

    st.markdown(f'<div style="background:{color}; color:white; padding:1.5rem; border-radius:8px; text-align:center;">'
                f'<h2>Escalation Probability: {escalation:.1f}%</h2>'
                f'<h3>{status}</h3><p>{action}</p></div>', unsafe_allow_html=True)

    st.markdown('<div class="formula-box">', unsafe_allow_html=True)
    st.markdown(f"""
    **PREDICTIVE THREAT ESCALATION FORMULA**

    **Escalation% = (Frequency_Trend × 0.5 + Confidence_Norm × 0.3 + Severity_Norm × 0.2) × 100**

    **Your Values:**
    - Frequency Trend: {freq_trend:.2f} → weighted contribution {freq_trend*0.5*100:.1f}
    - Confidence: {avg_conf}% → weighted contribution {(avg_conf/100)*0.3*100:.1f}
    - Severity: {avg_sev:.1f}/10 → weighted contribution {(avg_sev/10)*0.2*100:.1f}

    **Predicted Escalation Probability: {escalation:.1f}%**

    **Interpretation:**
    - < 25%: Unlikely to escalate
    - 25-50%: Possible - keep watching
    - 50-75%: Likely - proactively prepare response
    - > 75%: Highly likely - escalate now
    """)
    st.markdown('</div>', unsafe_allow_html=True)

    st.markdown('<div class="subsection-header">📈 30-Day Simulated Threat Volume Trend</div>', unsafe_allow_html=True)
    days = pd.date_range(end=dt.datetime.now(), periods=30)
    base = 20 + freq_trend * 40
    trend_vals = [max(0, base + i * freq_trend * 2 + random.uniform(-5, 5)) for i in range(30)]
    trend_df = pd.DataFrame({"Date": days, "Simulated_Event_Volume": trend_vals}).set_index("Date")
    st.line_chart(trend_df)

# =========================================================================
# TAB 4: ASSET & VULNERABILITY RISK
# =========================================================================
with tab4:
    st.markdown('<div class="section-header">🖥️ Attack Surface & Asset Vulnerability Risk</div>', unsafe_allow_html=True)
    st.markdown("""
    Estimate how exposed a given asset is based on open ports, known unpatched vulnerabilities,
    patch age, and internet exposure. Applies equally to digital assets (servers, endpoints)
    and physical assets (facilities, control systems) — for a physical asset, treat "open ports"
    as unmonitored entry points and "internet-facing" as public-facing/street-accessible.
    """)

    col1, col2, col3, col4 = st.columns(4)
    with col1:
        open_ports = st.slider("Open Ports", 0, 30, 6, step=1)
    with col2:
        known_vulns = st.slider("Known Unpatched Vulnerabilities", 0, 15, 3, step=1)
    with col3:
        patch_age = st.slider("Days Since Last Patch", 0, 365, 45, step=5)
    with col4:
        internet_facing = st.toggle("Internet-Facing Asset", value=True)

    asri = ThreatFormulas.attack_surface_risk_index(open_ports, known_vulns, patch_age, internet_facing)

    if asri < 3:
        status, color = "🟢 LOW EXPOSURE", "#1a8a3e"
    elif asri < 6:
        status, color = "🟡 MODERATE - SCHEDULE HARDENING", "#ff9800"
    elif asri < 8:
        status, color = "🟠 HIGH - PRIORITIZE REMEDIATION", "#ffa500"
    else:
        status, color = "🔴 CRITICAL - IMMEDIATE REMEDIATION", "#c0392b"

    st.markdown(f'<div style="background:{color}; color:white; padding:2rem; border-radius:8px; text-align:center;">'
                f'<h1>ASRI: {asri:.2f}/10</h1><h2>{status}</h2></div>', unsafe_allow_html=True)

    st.markdown('<div class="formula-box">', unsafe_allow_html=True)
    st.markdown(f"""
    **ATTACK SURFACE RISK INDEX (ASRI) FORMULA**

    **ASRI = Port_Score×0.25 + Vuln_Score×0.40 + Patch_Score×0.20 + Exposure_Score×0.15**

    - Port_Score = min(10, open_ports / 2) → {min(10, open_ports/2):.1f}
    - Vuln_Score = min(10, known_vulns × 1.5) → {min(10, known_vulns*1.5):.1f}
    - Patch_Score = min(10, days_since_patch / 30) → {min(10, patch_age/30):.1f}
    - Exposure_Score = 10 if internet-facing else 3 → {10 if internet_facing else 3}

    **Predicted ASRI: {asri:.2f}/10**

    **Reference:**
    - 0-3: Low exposure
    - 3-6: Moderate - schedule hardening
    - 6-8: High - prioritize remediation
    - 8-10: Critical - immediate remediation required
    """)
    st.markdown('</div>', unsafe_allow_html=True)

# =========================================================================
# TAB 5: DASHBOARD & ANALYTICS
# =========================================================================
with tab5:
    st.markdown('<div class="section-header">📈 Dashboard & Analytics Overview</div>', unsafe_allow_html=True)

    df_dash = st.session_state.real_df if st.session_state.real_df is not None else st.session_state.synthetic_df

    if df_dash is None:
        st.info("Generate or upload threat detection data in the 'Threat Detection' tab to populate this dashboard.")
    else:
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Total Events", len(df_dash))
        with col2:
            if "Risk_Score" in df_dash.columns:
                st.metric("Avg Risk Score", f"{df_dash['Risk_Score'].mean():.1f}")
        with col3:
            if "Status" in df_dash.columns:
                resolved = (df_dash["Status"] == "Resolved").sum()
                st.metric("Resolved", resolved)
        with col4:
            if "Status" in df_dash.columns:
                open_events = (df_dash["Status"].isin(["Detected", "Investigating"])).sum()
                st.metric("Open", open_events)

        if "Domain" in df_dash.columns:
            st.markdown('<div class="subsection-header">Cyber vs Physical/Operational Split</div>', unsafe_allow_html=True)
            domain_summary = df_dash.groupby("Domain").size().reset_index(name="Count")
            st.bar_chart(domain_summary.set_index("Domain"))

        if "Category" in df_dash.columns:
            st.markdown('<div class="subsection-header">Threats by Category</div>', unsafe_allow_html=True)
            cat_summary = df_dash.groupby("Category").size().reset_index(name="Count")
            st.bar_chart(cat_summary.set_index("Category"))

        if "Asset_Type" in df_dash.columns:
            st.markdown('<div class="subsection-header">Threats by Targeted Asset Type</div>', unsafe_allow_html=True)
            asset_summary = df_dash.groupby("Asset_Type").size().reset_index(name="Count")
            st.bar_chart(asset_summary.set_index("Asset_Type"))

        if "Risk_Score" in df_dash.columns and "Timestamp" in df_dash.columns:
            st.markdown('<div class="subsection-header">Risk Score Over Time</div>', unsafe_allow_html=True)
            try:
                temp = df_dash.copy()
                temp["Timestamp"] = pd.to_datetime(temp["Timestamp"], errors="coerce")
                temp = temp.dropna(subset=["Timestamp"]).sort_values("Timestamp")
                st.line_chart(temp.set_index("Timestamp")["Risk_Score"])
            except Exception:
                st.caption("Timestamp column could not be parsed for the time-series chart.")

        if "Asset_Type" in df_dash.columns and "Threat_Type" in df_dash.columns and "Risk_Score" in df_dash.columns:
            st.markdown('<div class="subsection-header">🔥 Risk Heatmap: Asset Type × Threat Type</div>', unsafe_allow_html=True)
            st.caption("Cell values show average Risk_Score for that asset/threat combination. Darker red = higher concentrated risk.")
            try:
                heatmap_data = df_dash.pivot_table(
                    index="Asset_Type", columns="Threat_Type", values="Risk_Score", aggfunc="mean"
                ).round(1)
                st.dataframe(
                    heatmap_data.style.background_gradient(cmap="Reds", axis=None).format(precision=1, na_rep=""),
                    use_container_width=True,
                )
            except Exception as e:
                st.caption(f"Heatmap could not be rendered: {e}")

# =========================================================================
# TAB 6: INCIDENT CASE MANAGEMENT
# =========================================================================
with tab6:
    st.markdown('<div class="section-header">🗂️ Incident Case Management</div>', unsafe_allow_html=True)
    st.markdown("""
    Track investigations as formal cases with an assigned analyst, priority-driven SLA,
    and status workflow. Priority determines the SLA window: Low=72h, Medium=24h, High=8h, Critical=2h.
    """)

    col1, col2, col3 = st.columns([2, 1, 1])
    with col1:
        n_cases = st.slider("Synthetic cases", 0, 200, 40, step=5)
    with col2:
        st.write(""); st.write("")
        gen_cases = st.button("🔄 Generate", key="gen_cases")
    with col3:
        st.write(""); st.write("")
        reset_cases = st.button("🗑️ Reset", key="reset_cases")

    if reset_cases:
        st.session_state.case_df = None
        st.success("✅ Case data reset")

    if gen_cases:
        source_df = st.session_state.real_df if st.session_state.real_df is not None else st.session_state.synthetic_df
        rows = []
        for i in range(n_cases):
            priority = random.choices(list(CASE_PRIORITIES.keys()), weights=[3, 4, 2, 1])[0]
            sla_hours = CASE_PRIORITIES[priority]
            created = dt.datetime.now() - dt.timedelta(hours=random.uniform(0, sla_hours * 2.5))
            status = random.choices(CASE_STATUSES, weights=[3, 3, 1, 3])[0]
            remaining = ThreatFormulas.sla_hours_remaining(created, sla_hours) if status != "Closed" else 0.0
            linked_event = ""
            if source_df is not None and "Event_ID" in source_df.columns and len(source_df) > 0:
                linked_event = source_df["Event_ID"].sample(1).values[0]

            rows.append({
                "Case_ID": f"CASE-{random.randint(10000, 99999)}",
                "Linked_Event_ID": linked_event,
                "Title": f"{random.choice(list(THREAT_TYPES.keys()))} investigation",
                "Priority": priority,
                "Assigned_Analyst": random.choice(ANALYSTS),
                "Status": status,
                "Created_At": created.strftime("%Y-%m-%d %H:%M"),
                "SLA_Hours": sla_hours,
                "Hours_Remaining": remaining,
                "SLA_State": "🔴 BREACHED" if remaining < 0 and status != "Closed" else ("🟡 AT RISK" if 0 <= remaining < sla_hours * 0.25 else "🟢 ON TRACK"),
            })
        st.session_state.case_df = pd.DataFrame(rows)
        st.success(f"✅ Generated {n_cases} synthetic cases")

    if st.session_state.case_df is not None:
        case_df = st.session_state.case_df
        st.dataframe(case_df, use_container_width=True, height=280)

        col1, col2, col3, col4 = st.columns(4)
        with col1:
            open_ct = case_df["Status"].isin(["Open", "In Progress", "Escalated"]).sum()
            st.metric("Open Cases", open_ct)
        with col2:
            breached = (case_df["SLA_State"] == "🔴 BREACHED").sum()
            st.metric("SLA Breached", breached)
        with col3:
            at_risk = (case_df["SLA_State"] == "🟡 AT RISK").sum()
            st.metric("SLA At Risk", at_risk)
        with col4:
            closed = (case_df["Status"] == "Closed").sum()
            st.metric("Closed", closed)

        st.markdown('<div class="subsection-header">Cases by Analyst</div>', unsafe_allow_html=True)
        analyst_summary = case_df.groupby("Assigned_Analyst").size().reset_index(name="Case_Count")
        st.bar_chart(analyst_summary.set_index("Assigned_Analyst"))

        st.markdown('<div class="formula-box">', unsafe_allow_html=True)
        st.markdown("""
        **SLA HOURS REMAINING FORMULA**

        **Remaining = SLA_Hours − (Now − Created_At in hours)**

        **SLA Windows by Priority:** Low = 72h | Medium = 24h | High = 8h | Critical = 2h

        A negative remaining value means the case has breached its SLA and should be reprioritized immediately.
        """)
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown('<div class="subsection-header">💾 Export Results</div>', unsafe_allow_html=True)
        formulas_dict = {"SLA Hours Remaining": "Remaining = SLA_Hours - (Now - Created_At in hours)"}
        render_export_row(case_df, "Incident Case Management", formulas_dict, "case_mgmt")

# =========================================================================
# TAB 7: ALERT RULES & NOTIFICATION SIMULATION
# =========================================================================
with tab7:
    st.markdown('<div class="section-header">🔔 Alert Rules & Notification Simulation</div>', unsafe_allow_html=True)
    st.markdown("""
    Define threshold-based alert rules and simulate the notification log that would be
    generated against the currently loaded threat detection data.
    """)

    col1, col2, col3 = st.columns(3)
    with col1:
        risk_threshold = st.slider("Alert if Risk_Score ≥", 0, 100, 75, step=5)
    with col2:
        channel_choice = st.multiselect("Notification channels", ALERT_CHANNELS, default=["Email", "SOC Dashboard"])
    with col3:
        st.write(""); st.write("")
        run_sim = st.button("🔔 Run Alert Simulation")

    df_for_alerts = st.session_state.real_df if st.session_state.real_df is not None else st.session_state.synthetic_df

    if run_sim:
        if df_for_alerts is None or "Risk_Score" not in df_for_alerts.columns:
            st.warning("No threat detection data loaded. Generate or upload data in the Threat Detection tab first.")
        elif not channel_choice:
            st.warning("Select at least one notification channel.")
        else:
            triggered = df_for_alerts[df_for_alerts["Risk_Score"] >= risk_threshold]
            rows = []
            for _, ev in triggered.iterrows():
                channel = random.choice(channel_choice)
                rows.append({
                    "Alert_ID": f"ALRT-{random.randint(100000, 999999)}",
                    "Event_ID": ev.get("Event_ID", ""),
                    "Rule_Triggered": f"Risk_Score >= {risk_threshold}",
                    "Risk_Score": ev.get("Risk_Score", ""),
                    "Domain": ev.get("Domain", ""),
                    "Channel": channel,
                    "Sent_At": dt.datetime.now().strftime("%Y-%m-%d %H:%M"),
                    "Message": f"[ALERT] {ev.get('Threat_Type','Unknown')} on {ev.get('Asset_Type','Unknown')} - Risk {ev.get('Risk_Score','')}",
                })
            st.session_state.alert_log_df = pd.DataFrame(rows)
            st.success(f"✅ Simulation complete: {len(rows)} alert(s) triggered")

    if st.session_state.alert_log_df is not None:
        alert_df = st.session_state.alert_log_df
        st.dataframe(alert_df, use_container_width=True, height=250)

        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Total Alerts", len(alert_df))
        with col2:
            if "Channel" in alert_df.columns and len(alert_df) > 0:
                top_channel = alert_df["Channel"].mode().values[0]
                st.metric("Top Channel", top_channel)
        with col3:
            if "Domain" in alert_df.columns:
                cyber_alerts = (alert_df["Domain"] == "Cyber").sum()
                st.metric("Cyber Alerts", cyber_alerts)

        if len(alert_df) > 0:
            st.markdown('<div class="subsection-header">Alerts by Channel</div>', unsafe_allow_html=True)
            channel_summary = alert_df.groupby("Channel").size().reset_index(name="Count")
            st.bar_chart(channel_summary.set_index("Channel"))

        st.markdown('<div class="subsection-header">💾 Export Results</div>', unsafe_allow_html=True)
        formulas_dict = {"Alert Rule": f"Trigger alert when Risk_Score >= configured threshold, dispatched to selected channels"}
        render_export_row(alert_df, "Alert Log", formulas_dict, "alert_log")

# =========================================================================
# TAB 8: WATCHLIST
# =========================================================================
with tab8:
    st.markdown('<div class="section-header">⭐ Watchlist</div>', unsafe_allow_html=True)
    st.markdown("""
    Flag specific IPs, zones, assets, or users for priority monitoring. Any current threat
    detection events matching a watchlist entry are counted as hits below.
    """)

    with st.form("add_watchlist_form"):
        st.markdown("**Add a Watchlist Entry**")
        c1, c2, c3, c4 = st.columns(4)
        with c1:
            wl_type = st.selectbox("Type", WATCHLIST_TYPES)
        with c2:
            wl_value = st.text_input("Value (e.g. IP, zone, asset type, user)")
        with c3:
            wl_priority = st.selectbox("Priority", WATCHLIST_PRIORITIES)
        with c4:
            wl_reason = st.text_input("Reason", value="Analyst flagged")
        submitted = st.form_submit_button("➕ Add to Watchlist")

    if submitted and wl_value.strip():
        new_entry = pd.DataFrame([{
            "Watchlist_ID": f"WL-{random.randint(10000, 99999)}",
            "Type": wl_type,
            "Value": wl_value.strip(),
            "Priority": wl_priority,
            "Reason": wl_reason,
            "Added_By": "Current Analyst",
            "Date_Added": dt.datetime.now().strftime("%Y-%m-%d %H:%M"),
        }])
        if st.session_state.watchlist_df is None:
            st.session_state.watchlist_df = new_entry
        else:
            st.session_state.watchlist_df = pd.concat([st.session_state.watchlist_df, new_entry], ignore_index=True)
        st.success(f"✅ Added '{wl_value.strip()}' to watchlist")

    col1, col2 = st.columns([3, 1])
    with col1:
        n_wl = st.slider("Or generate synthetic watchlist entries", 0, 50, 10, step=5)
    with col2:
        st.write(""); st.write("")
        gen_wl = st.button("🔄 Generate Sample Watchlist")

    if gen_wl:
        rows = []
        for i in range(n_wl):
            wtype = random.choice(WATCHLIST_TYPES)
            if wtype == "IP Address":
                value = f"{random.randint(1,223)}.{random.randint(0,255)}.{random.randint(0,255)}.{random.randint(1,254)}"
            elif wtype == "Zone/Facility Area":
                value = f"ZONE-{random.choice(['A','B','C','D'])}{random.randint(1,12)}"
            elif wtype == "Asset Type":
                value = random.choice(ASSET_TYPES)
            else:
                value = f"user{random.randint(100,999)}@company.com"
            rows.append({
                "Watchlist_ID": f"WL-{random.randint(10000, 99999)}",
                "Type": wtype,
                "Value": value,
                "Priority": random.choice(WATCHLIST_PRIORITIES),
                "Reason": random.choice(["Prior incident", "Threat intel match", "Anomalous behavior", "Analyst flagged", "Repeat offender"]),
                "Added_By": random.choice(ANALYSTS),
                "Date_Added": (dt.datetime.now() - dt.timedelta(days=random.randint(0, 60))).strftime("%Y-%m-%d %H:%M"),
            })
        st.session_state.watchlist_df = pd.DataFrame(rows)
        st.success(f"✅ Generated {n_wl} watchlist entries")

    if st.session_state.watchlist_df is not None:
        wl_df = st.session_state.watchlist_df
        st.dataframe(wl_df, use_container_width=True, height=250)

        df_for_hits = st.session_state.real_df if st.session_state.real_df is not None else st.session_state.synthetic_df
        hit_count = 0
        if df_for_hits is not None:
            watch_values = set(wl_df["Value"].astype(str))
            for col in ["Source_IP_or_Zone", "Asset_Type"]:
                if col in df_for_hits.columns:
                    hit_count += df_for_hits[col].astype(str).isin(watch_values).sum()

        col1, col2 = st.columns(2)
        with col1:
            st.metric("Watchlist Entries", len(wl_df))
        with col2:
            st.metric("Matching Events in Current Data", int(hit_count))

        st.markdown('<div class="subsection-header">💾 Export Results</div>', unsafe_allow_html=True)
        render_export_row(wl_df, "Watchlist", {}, "watchlist")

# =========================================================================
# TAB 9: EXECUTIVE SUMMARY REPORT
# =========================================================================
with tab9:
    st.markdown('<div class="section-header">🧾 Executive Summary Report</div>', unsafe_allow_html=True)
    st.markdown("""
    Auto-generate a written narrative summarizing current threat detection, case management,
    and alerting activity - suitable for leadership or shift-handoff reporting.
    """)

    if st.button("🧾 Generate Executive Summary"):
        df_for_summary = st.session_state.real_df if st.session_state.real_df is not None else st.session_state.synthetic_df
        narrative = build_executive_summary(df_for_summary, st.session_state.case_df, st.session_state.alert_log_df)
        st.session_state["exec_summary_text"] = narrative

    if "exec_summary_text" in st.session_state:
        st.markdown('<div class="formula-box">', unsafe_allow_html=True)
        st.markdown(st.session_state["exec_summary_text"].replace("\n", "  \n"))
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown('<div class="subsection-header">💾 Export Report</div>', unsafe_allow_html=True)
        col1, col2, col3 = st.columns(3)
        narrative_text = st.session_state["exec_summary_text"]
        with col1:
            st.download_button("📄 TXT", narrative_text.encode(), "executive_summary.txt", "text/plain", key="exec_txt")
        with col2:
            if HAS_DOCX:
                docx_bytes = export_report_to_docx(narrative_text, "Executive Summary Report")
                if docx_bytes:
                    st.download_button("📝 WORD", docx_bytes, "executive_summary.docx",
                                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                        key="exec_docx")
        with col3:
            if HAS_PDF:
                pdf_bytes = export_report_to_pdf(narrative_text, "Executive Summary Report")
                if pdf_bytes:
                    st.download_button("🔴 PDF", pdf_bytes, "executive_summary.pdf", "application/pdf", key="exec_pdf")
    else:
        st.info("Click 'Generate Executive Summary' to compose the report from current session data.")

st.markdown("---")
st.caption(f"🛡️ CyberSentinel Intelligence Platform v1.0 | Session {st.session_state.session_id} | {DEPLOY_ENV.upper()} | Developed by Randy Singh | Kalsnet (KNet) Consulting")
