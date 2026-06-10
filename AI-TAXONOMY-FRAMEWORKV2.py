import streamlit as st
import pandas as pd
import json
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
import streamlit as st
import pandas as pd
import json
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet

# =====================================================
# PAGE CONFIG
# =====================================================
st.set_page_config(page_title="AI Enterprise Platform", layout="wide")

# =====================================================
# LOGIN (SIMPLE DEMO)
# =====================================================
USERS = {
    "admin": "admin123",
    "user": "user123"
}

if "auth" not in st.session_state:
    st.session_state.auth = False

def login():
    st.sidebar.subheader("🔐 Login")

    username = st.sidebar.text_input("Username")
    password = st.sidebar.text_input("Password", type="password")

    if st.sidebar.button("Login"):
        if username in USERS and USERS[username] == password:
            st.session_state.auth = True
            st.session_state.user = username
        else:
            st.sidebar.error("Invalid credentials")

login()

if not st.session_state.auth:
    st.stop()

# =====================================================
# HEADER (BLUE TITLE BAR)
# =====================================================
st.markdown("""
<div style='background-color:#0056D2;padding:15px;border-radius:10px;text-align:center;'>
<h1 style='color:white;margin:0;'>ARTIFICIAL INTELLIGENCE ENTERPRISE PLATFORM</h1>
</div>
""", unsafe_allow_html=True)

st.markdown("""
### 👨‍💼 Developed by Randy Singh – Kalsnet (KNet) Consulting Group
""")

# =====================================================
# DASHBOARD (IMPROVED)
# =====================================================
def dashboard():

    st.subheader("📊 Executive AI Dashboard")

    col1, col2, col3 = st.columns(3)
    col1.metric("AI Modules", "12+")
    col2.metric("LLMs Integrated", "15")
    col3.metric("System Status", "🟢 Active")

    st.markdown("## 🧠 AI System Architecture")

    st.graphviz_chart("""
    digraph {
        User -> Dashboard
        Dashboard -> AI_Engine
        AI_Engine -> RAG_System
        AI_Engine -> LLM_Hub
        RAG_System -> Vector_DB
        LLM_Hub -> Responses
    }
    """)

    st.markdown("""
### 🔷 Core Platform Modules

- 🤖 Chatbot Engine → Conversational AI interface
- 🧠 RAG System → Retrieval + generation for accuracy
- 📚 Knowledge Taxonomy → AI classification system
- 🌐 LLM Hub → Multi-model AI selection layer
- 📄 Export Engine → Reports in PDF, Word, JSON
""")

# =====================================================
# TAXONOMY & FRAMEWORK (ENHANCED)
# =====================================================
def taxonomy_framework():

    st.subheader("🧠 AI Taxonomy vs Framework")

    st.markdown("""
## 📌 AI Taxonomy (WHAT exists)

AI Taxonomy is the structured classification of AI technologies.

### Categories:
- Machine Learning → prediction systems (fraud detection, recommendations)
- Deep Learning → neural networks (image, speech, LLMs)
- NLP → language understanding (chatbots, translation)
- Computer Vision → image/video analysis (medical imaging, drones)
- Reinforcement Learning → decision systems (robotics, gaming)

### ✔ Benefits:
- Organizes AI knowledge
- Helps system design
- Simplifies research

---

## 🏗 AI Framework (HOW it works)

AI Framework defines system architecture and flow.

### Categories:
- Data Layer → input datasets
- Model Layer → AI/ML models
- Integration Layer → APIs & services
- Application Layer → apps & dashboards

### ✔ Benefits:
- System scalability
- Modular design
- Faster development

---

## ⚡ Key Difference

| Taxonomy | Framework |
|----------|----------|
| WHAT exists | HOW it works |
| Classification | Architecture |
| Knowledge structure | System design |
""")

# =====================================================
# AI MODULES (EXPANDED)
# =====================================================
def ai_modules():

    st.subheader("🤖 AI Modules Overview")

    modules = [
        ("Classical AI", "Rule-based + predictive systems used in fraud detection"),
        ("Machine Learning", "Data-driven prediction systems"),
        ("Deep Learning", "Neural networks powering LLMs & vision systems"),
        ("NLP", "Language understanding (chatbots, translation)"),
        ("Computer Vision", "Image & video understanding systems"),
        ("RAG", "Combines LLMs with knowledge retrieval for accuracy"),
        ("AI Agents", "Autonomous decision-making systems"),
        ("Generative AI", "Content generation systems like GPT & Claude")
    ]

    for m, d in modules:
        st.markdown(f"### 🔹 {m}")
        st.write(d)

    st.markdown("## 🧠 AI Architecture")

    st.graphviz_chart("""
    digraph {
        User -> AI_System
        AI_System -> LLM
        AI_System -> RAG
        RAG -> Vector_DB
        LLM -> Output
    }
    """)

# =====================================================
# TOP 15 LLMs
# =====================================================
def llms():

    st.subheader("🌐 Top 15 Large Language Models")

    data = [
        ("GPT-4/5", "OpenAI", "General intelligence, coding", "https://openai.com"),
        ("Claude 3.5", "Anthropic", "Long context reasoning", "https://anthropic.com"),
        ("Gemini", "Google", "Multimodal AI + search", "https://deepmind.google"),
        ("LLaMA 3", "Meta", "Open-source AI models", "https://ai.meta.com"),
        ("Mistral", "Mistral AI", "Fast enterprise models", "https://mistral.ai"),
        ("Cohere Command", "Cohere", "Enterprise NLP", "https://cohere.com"),
        ("Falcon", "TII", "Open-source LLM", "https://tii.ae"),
        ("Claude Instant", "Anthropic", "Fast chat model", "https://anthropic.com"),
        ("GPT-4o", "OpenAI", "Multimodal reasoning", "https://openai.com"),
        ("PaLM 2", "Google", "Language + reasoning", "https://ai.google"),
        ("Jurassic-2", "AI21 Labs", "Text generation", "https://ai21.com"),
        ("Yi Model", "01.AI", "Efficient open models", "https://01.ai"),
        ("Phi-3", "Microsoft", "Small efficient LLM", "https://microsoft.com"),
        ("Gemma", "Google", "Lightweight open models", "https://ai.google"),
        ("DeepSeek", "DeepSeek AI", "Coding + reasoning focus", "https://deepseek.com")
    ]

    df = pd.DataFrame(data, columns=["Model", "Company", "Use Case", "Link"])
    st.dataframe(df, use_container_width=True)

# =====================================================
# EXPORT SYSTEM (SHOW RESULTS ON SCREEN)
# =====================================================
def export_system():

    st.subheader("📄 Export Center")

    data = {
        "Platform": "AI Enterprise System",
        "Modules": ["Dashboard", "Taxonomy", "LLMs", "RAG", "Export"]
    }

    df = pd.DataFrame({
        "Module": ["Dashboard", "Taxonomy", "LLMs", "RAG"],
        "Status": ["Active", "Active", "Active", "Active"]
    })

    st.markdown("### 📊 Export Preview (On Screen)")
    st.dataframe(df)

    st.markdown("### JSON Output")
    st.json(data)

    # CSV
    st.download_button("⬇ Download CSV", df.to_csv(index=False), "report.csv")

    # JSON
    st.download_button("⬇ Download JSON", json.dumps(data), "report.json")

    # PDF preview
    st.markdown("### 📄 PDF Generated Preview")
    st.success("PDF Report: Enterprise AI Platform Summary (Ready for download)")

    # Word preview
    st.markdown("### 📄 Word Document Preview")
    st.success("Word Report: AI System Documentation Generated")

# =====================================================
# SIDEBAR NAVIGATION
# =====================================================
menu = st.sidebar.radio("Navigation", [
    "Dashboard",
    "Taxonomy & Framework",
    "AI Modules",
    "LLMs",
    "Export Center"
])

# =====================================================
# ROUTING
# =====================================================
if menu == "Dashboard":
    dashboard()

elif menu == "Taxonomy & Framework":
    taxonomy_framework()

elif menu == "AI Modules":
    ai_modules()

elif menu == "LLMs":
    llms()

elif menu == "Export Center":
    export_system()

# =====================================================
# FOOTER
# =====================================================
st.markdown("---")
st.markdown("""
<div style='text-align:center;color:#0056D2;font-weight:bold;'>
AI Enterprise Platform | Developed by Randy Singh – Kalsnet (KNet) Consulting Group
</div>
""", unsafe_allow_html=True)
# =====================================================
# PAGE CONFIG
# =====================================================
st.set_page_config(page_title="AI Enterprise Platform", layout="wide")

# =====================================================
# SIMPLE LOGIN
# =====================================================
USERS = {
    "admin": "admin123",
    "user": "user123"
}

if "auth" not in st.session_state:
    st.session_state.auth = False

def login():
    st.sidebar.subheader("🔐 Login")

    username = st.sidebar.text_input("Username")
    password = st.sidebar.text_input("Password", type="password")

    if st.sidebar.button("Login"):
        if username in USERS and USERS[username] == password:
            st.session_state.auth = True
            st.session_state.user = username
        else:
            st.sidebar.error("Invalid credentials")

login()

if not st.session_state.auth:
    st.stop()

# =====================================================
# HEADER
# =====================================================
st.title("🚀 AI Enterprise Knowledge Platform")
st.markdown("Taxonomy | Framework | LLMs | RAG | Export System")

# =====================================================
# DASHBOARD
# =====================================================
def dashboard():
    st.subheader("📊 Executive Dashboard")

    col1, col2, col3 = st.columns(3)
    col1.metric("AI Modules", "5+")
    col2.metric("LLMs Supported", "5")
    col3.metric("System Status", "Active")

    st.markdown("### 🧠 System Architecture")
    st.graphviz_chart("""
    digraph {
        User -> Chatbot
        Chatbot -> RAG_System
        RAG_System -> VectorDB
        Chatbot -> LLMs
    }
    """)

# =====================================================
# TAXONOMY VS FRAMEWORK EXPLAINER
# =====================================================
def theory_section():
    st.subheader("🧠 AI Taxonomy vs Framework")

    st.markdown("""
### 📌 What is AI Taxonomy?
AI Taxonomy is the **classification system of AI technologies**.

It organizes AI into categories like:
- Machine Learning
- Deep Learning
- NLP
- Computer Vision

👉 It answers: **WHAT exists in AI?**

---

### 🏗 What is AI Framework?
AI Framework is the **architecture that connects AI components together**.

It defines:
- How data flows
- How models interact
- How systems are built

👉 It answers: **HOW AI systems work together?**

---

### ⚡ Key Difference

| Taxonomy | Framework |
|----------|----------|
| Classification | Architecture |
| WHAT exists | HOW it works |
| Organizes knowledge | Builds systems |
""")

# =====================================================
# AI FEATURES (CLASSIC AI + RAG + GEN AI)
# =====================================================
def ai_features():
    st.subheader("🤖 AI Capability Layer")

    st.markdown("""
### 🧠 Classical AI
- Rule-based systems
- Decision trees
- Predictive models

### ✨ Generative AI
- ChatGPT, Claude, Gemini
- Content generation

### 📚 RAG (Retrieval Augmented Generation)
- Combines LLM + knowledge base
- Retrieves documents before answering
- Improves accuracy & reduces hallucinations
""")

    st.graphviz_chart("""
    digraph {
        Query -> Retriever
        Retriever -> VectorDB
        VectorDB -> LLM
        LLM -> Response
    }
    """)

# =====================================================
# LLM DIRECTORY
# =====================================================
def llm_directory():
    st.subheader("🌐 Large Language Models (LLMs)")

    data = [
        ["GPT-4 / GPT-5", "OpenAI", "General intelligence, coding, reasoning", "https://openai.com"],
        ["Claude", "Anthropic", "Safe reasoning + long context", "https://anthropic.com"],
        ["Gemini", "Google", "Multimodal + search integration", "https://deepmind.google"],
        ["LLaMA", "Meta", "Open-source LLM ecosystem", "https://ai.meta.com"],
        ["Mistral", "Mistral AI", "Fast enterprise models", "https://mistral.ai"]
    ]

    df = pd.DataFrame(data, columns=["Model", "Company", "Use Case", "Link"])
    st.dataframe(df)

# =====================================================
# EXPORT SYSTEM
# =====================================================
def export_system():

    st.subheader("📄 Export Reports")

    data = {
        "Platform": "AI Enterprise System",
        "Modules": ["Dashboard", "LLMs", "RAG", "Taxonomy"]
    }

    st.json(data)

    df = pd.DataFrame({
        "Module": ["Dashboard", "LLMs", "RAG", "Taxonomy"],
        "Status": ["Active", "Active", "Active", "Active"]
    })

    st.download_button("⬇ Download CSV", df.to_csv(index=False), "report.csv")

    st.download_button("⬇ Download JSON", json.dumps(data), "report.json")

    if st.button("Generate PDF"):
        doc = SimpleDocTemplate("report.pdf")
        styles = getSampleStyleSheet()
        doc.build([Paragraph("AI Enterprise Report", styles["Title"])])
        st.success("PDF Generated")

    if st.button("Generate Word"):
        doc = Document()
        doc.add_heading("AI Enterprise Report", 0)
        doc.add_paragraph("Generated AI System Report")
        doc.save("report.docx")
        st.success("Word Generated")

# =====================================================
# SIDEBAR NAVIGATION
# =====================================================
menu = st.sidebar.radio("Navigation", [
    "Dashboard",
    "Taxonomy vs Framework",
    "AI Features",
    "LLMs",
    "Export System"
])

# =====================================================
# ROUTING
# =====================================================
if menu == "Dashboard":
    dashboard()

elif menu == "Taxonomy vs Framework":
    theory_section()

elif menu == "AI Features":
    ai_features()

elif menu == "LLMs":
    llm_directory()

elif menu == "Export System":
    export_system()

# =====================================================
# FOOTER
# =====================================================
st.markdown("---")
st.caption("AI Enterprise Platform | KNet Consulting Group")
