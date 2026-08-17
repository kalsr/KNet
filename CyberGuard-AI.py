


# CyberGuard AI
# Developed by Randy Singh from Kalsnet (KNet) Consulting group
# -----------------------------------------------------------
# A unified Streamlit application demonstrating 5 AI/ML-driven
# cybersecurity use cases:
# 1. Autonomous SOC / AI Security Analyst
# 2. AI Vulnerability & Exposure Management
# 3. Behavioral Threat & Anomaly Detection
# 4. AI Incident Response & Attack-Path Analysis
# 5. AI Identity & Agent Security

# Each module supports synthetic data generation (~200 records by
# default), real-data CSV upload, ML-driven analysis with documented
# scoring formulas, visualizations, and export of results to PDF,
# Word, CSV and Text.


import io
import random
import base64
from datetime import datetime, timedelta

import numpy as np
import pandas as pd
import streamlit as st
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import networkx as nx

from sklearn.ensemble import IsolationForest

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors as rl_colors
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image as RLImage
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

# =============================================================
# PAGE CONFIG
# =============================================================
st.set_page_config(
    page_title="CyberGuard AI",
    page_icon="🛡️",
    layout="wide",
    initial_sidebar_state="expanded",
)

APP_NAME = "CyberGuard AI"
DEVELOPER_LINE = "Developed by Randy Singh from Kalsnet (KNet) Consulting group"

PRIMARY = "#0B2545"
ACCENT = "#13C4A3"
ACCENT2 = "#FF6B6B"
LIGHT_BG = "#F4F7FB"
DEV_BLUE = "#4FA8FF"

# =============================================================
# GLOBAL STYLING + TITLE BAR
# =============================================================
def render_title_bar(subtitle: str):
    st.markdown(
        f"""
        <style>
            .block-container {{ padding-top: 1.2rem; }}
            #MainMenu {{visibility: hidden;}}
            footer {{visibility: hidden;}}

            .knet-header {{
                background: linear-gradient(90deg, {PRIMARY} 0%, #123B6D 60%, {ACCENT} 150%);
                padding: 18px 28px;
                border-radius: 10px;
                display: flex;
                justify-content: space-between;
                align-items: center;
                box-shadow: 0 4px 14px rgba(0,0,0,0.25);
                margin-bottom: 6px;
            }}
            .knet-title {{
                color: white;
                font-size: 28px;
                font-weight: 800;
                letter-spacing: 0.3px;
                margin: 0;
            }}
            .knet-developer {{
                color: {DEV_BLUE};
                font-size: 15px;
                font-weight: 700;
                margin: 4px 0 2px 0;
            }}
            .knet-subtitle {{
                color: #D8E6F5;
                font-size: 13.5px;
                margin-top: 2px;
            }}
            .knet-badge {{
                background: rgba(255,255,255,0.12);
                border: 1px solid rgba(255,255,255,0.35);
                color: white;
                padding: 10px 16px;
                border-radius: 8px;
                text-align: right;
                font-size: 12.5px;
                line-height: 1.4;
            }}
            .knet-badge b {{ color: {ACCENT}; font-size: 13.5px; }}
            .module-banner {{
                background: white;
                border-left: 6px solid {ACCENT};
                padding: 14px 20px;
                border-radius: 6px;
                margin: 14px 0 18px 0;
                box-shadow: 0 2px 8px rgba(0,0,0,0.06);
            }}
            .module-banner h3 {{ margin: 0 0 6px 0; color: {PRIMARY}; }}
            .module-banner p {{ margin: 0; color: #333; font-size: 14.5px; }}
            .metric-card {{
                background: white;
                border-radius: 10px;
                padding: 14px;
                box-shadow: 0 2px 6px rgba(0,0,0,0.08);
                text-align: center;
                border-top: 4px solid {ACCENT};
            }}
            .stTabs [data-baseweb="tab-list"] {{ gap: 6px; }}
            .stTabs [data-baseweb="tab"] {{
                background-color: white;
                border-radius: 8px 8px 0 0;
                padding: 8px 16px;
                font-weight: 600;
            }}
            .footer-bar {{
                margin-top: 40px;
                padding: 12px;
                text-align: center;
                color: #7c8ba1;
                font-size: 12px;
                border-top: 1px solid #e3e8ef;
            }}
        </style>

        <div class="knet-header">
            <div>
                <p class="knet-title">{APP_NAME}</p>
                <p class="knet-developer">{DEVELOPER_LINE}</p>
                <p class="knet-subtitle">{subtitle}</p>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_footer():
    st.markdown(
        f"""
        <div class="footer-bar">
            CyberGuard AI &nbsp;|&nbsp; © {datetime.now().year} Kalsnet (KNet) Consulting group
            &nbsp;|&nbsp; Developed by Randy Singh &nbsp;|&nbsp; For demonstration &amp; training purposes only
        </div>
        """,
        unsafe_allow_html=True,
    )


def module_banner(title, description):
    st.markdown(
        f"""
        <div class="module-banner">
            <h3>{title}</h3>
            <p>{description}</p>
        </div>
        """,
        unsafe_allow_html=True,
    )


def field_glossary(schema: dict):
    with st.expander("Schema Field Reference — What each column means", expanded=False):
        df = pd.DataFrame(
            [{"Field": k, "Description": v} for k, v in schema.items()]
        )
        st.table(df)


def formula_box(explanation: str, formulas: list):
    """Displays a brief explanation plus the exact scoring formulas used in the module."""
    with st.expander("AI Methodology & Formulas Used", expanded=False):
        st.markdown(explanation)
        for label, formula in formulas:
            st.markdown(f"**{label}**")
            st.code(formula, language="text")


def show_all_records(df: pd.DataFrame, key_prefix: str):
    """Confirms the full record count and gives access to every loaded/generated record."""
    st.success(f"{len(df)} records loaded — all {len(df)} records are included in the analysis below.")
    with st.expander(f"View All {len(df)} Loaded Records (raw data)", expanded=False):
        st.dataframe(df, use_container_width=True, height=min(600, 80 + 28 * min(len(df), 18)))


def metric_row(items):
    cols = st.columns(len(items))
    for c, (label, value) in zip(cols, items):
        c.markdown(
            f"""<div class="metric-card"><div style="font-size:22px;font-weight:800;color:{PRIMARY}">{value}</div>
            <div style="font-size:12.5px;color:#667;">{label}</div></div>""",
            unsafe_allow_html=True,
        )


# =============================================================
# GENERIC EXPORT HELPERS (CSV / TEXT / WORD / PDF)
# =============================================================
def export_csv_bytes(df: pd.DataFrame) -> bytes:
    return df.to_csv(index=False).encode("utf-8")


def export_text_bytes(title: str, narrative: str, df: pd.DataFrame) -> bytes:
    buf = io.StringIO()
    buf.write(f"{title}\n")
    buf.write("=" * len(title) + "\n")
    buf.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
    buf.write(f"{APP_NAME} | {DEVELOPER_LINE}\n\n")
    buf.write(narrative.strip() + "\n\n")
    buf.write("-- Result Data --\n")
    buf.write(df.to_string(index=False))
    return buf.getvalue().encode("utf-8")


def export_word_bytes(title: str, narrative: str, df: pd.DataFrame, chart_png: bytes = None) -> bytes:
    doc = Document()

    # Header / title block
    h = doc.add_heading(level=0)
    run = h.add_run(APP_NAME)
    run.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
    sub = doc.add_paragraph(DEVELOPER_LINE)
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub.runs[0].font.size = Pt(11)
    sub.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph(f"Report generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_heading(title, level=1)

    for para in narrative.strip().split("\n"):
        if para.strip():
            doc.add_paragraph(para.strip())

    if chart_png:
        doc.add_paragraph("")
        doc.add_picture(io.BytesIO(chart_png), width=Inches(6))

    doc.add_heading("Result Data", level=2)
    if not df.empty:
        max_rows = min(len(df), 200)
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr[i].text = str(col)
        for _, row in df.head(max_rows).iterrows():
            cells = table.add_row().cells
            for i, col in enumerate(df.columns):
                cells[i].text = str(row[col])
        if len(df) > max_rows:
            doc.add_paragraph(f"...showing first {max_rows} of {len(df)} rows.")

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def export_pdf_bytes(title: str, narrative: str, df: pd.DataFrame, chart_png: bytes = None) -> bytes:
    out = io.BytesIO()
    doc = SimpleDocTemplate(out, pagesize=letter, topMargin=0.6 * inch, bottomMargin=0.6 * inch)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("TitleStyle", parent=styles["Title"], textColor=rl_colors.HexColor("#0B2545"))
    sub_style = ParagraphStyle("SubStyle", parent=styles["Normal"], textColor=rl_colors.HexColor("#555555"), fontSize=10)
    body_style = styles["Normal"]

    elems = []
    elems.append(Paragraph(APP_NAME, title_style))
    elems.append(Paragraph(DEVELOPER_LINE, sub_style))
    elems.append(Paragraph(f"Report generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", sub_style))
    elems.append(Spacer(1, 14))
    elems.append(Paragraph(title, styles["Heading1"]))
    for para in narrative.strip().split("\n"):
        if para.strip():
            elems.append(Paragraph(para.strip(), body_style))
    elems.append(Spacer(1, 10))

    if chart_png:
        elems.append(RLImage(io.BytesIO(chart_png), width=6 * inch, height=3.2 * inch))
        elems.append(Spacer(1, 10))

    elems.append(Paragraph("Result Data", styles["Heading2"]))
    if not df.empty:
        max_rows = min(len(df), 60)
        max_cols = min(len(df.columns), 8)
        cols = list(df.columns[:max_cols])
        data = [cols] + df[cols].head(max_rows).astype(str).values.tolist()
        tbl = Table(data, repeatRows=1)
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), rl_colors.HexColor("#0B2545")),
            ("TEXTCOLOR", (0, 0), (-1, 0), rl_colors.white),
            ("FONTSIZE", (0, 0), (-1, -1), 7),
            ("GRID", (0, 0), (-1, -1), 0.4, rl_colors.grey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [rl_colors.white, rl_colors.HexColor("#F4F7FB")]),
        ]))
        elems.append(tbl)
        if len(df) > max_rows or len(df.columns) > max_cols:
            elems.append(Paragraph(
                f"Showing first {max_rows} rows / {max_cols} columns of {len(df)} rows / {len(df.columns)} columns.",
                sub_style))

    doc.build(elems)
    return out.getvalue()


def fig_to_png_bytes(fig) -> bytes:
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    buf.seek(0)
    return buf.getvalue()


def export_panel(key_prefix: str, title: str, narrative: str, df: pd.DataFrame, fig=None):
    st.markdown("#### Export Results")
    chart_png = fig_to_png_bytes(fig) if fig is not None else None
    c1, c2, c3, c4 = st.columns(4)
    c1.download_button("⬇️ CSV", export_csv_bytes(df), file_name=f"{key_prefix}.csv",
                        mime="text/csv", use_container_width=True, key=f"{key_prefix}_csv")
    c2.download_button("⬇️ Text", export_text_bytes(title, narrative, df), file_name=f"{key_prefix}.txt",
                        mime="text/plain", use_container_width=True, key=f"{key_prefix}_txt")
    c3.download_button("⬇️ Word (.docx)", export_word_bytes(title, narrative, df, chart_png),
                        file_name=f"{key_prefix}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True, key=f"{key_prefix}_docx")
    c4.download_button("⬇️ PDF", export_pdf_bytes(title, narrative, df, chart_png), file_name=f"{key_prefix}.pdf",
                        mime="application/pdf", use_container_width=True, key=f"{key_prefix}_pdf")


def data_source_controls(key_prefix, schema, gen_func, default_n=200, template_cols=None,
                          slider_label="Number of synthetic records", slider_min=50,
                          slider_max=2000, slider_step=50):
    """Common widget block: choose synthetic vs upload, returns a dataframe or None."""
    st.markdown("#### Data Source")
    mode = st.radio(
        "Choose data source",
        ["Generate Synthetic Data", "Upload Real Data (CSV)"],
        horizontal=True,
        key=f"{key_prefix}_mode",
    )
    df = None
    if mode.startswith("Generate"):
        n = st.slider(slider_label, slider_min, slider_max, default_n, step=slider_step, key=f"{key_prefix}_n")
        if st.button("Generate Synthetic Data", key=f"{key_prefix}_gen_btn", type="primary"):
            df = gen_func(n)
            st.session_state[f"{key_prefix}_df"] = df
        if f"{key_prefix}_df" in st.session_state:
            df = st.session_state[f"{key_prefix}_df"]
    else:
        if template_cols is not None:
            template_df = pd.DataFrame(columns=template_cols)
            st.download_button(
                "⬇️ Download CSV Template", export_csv_bytes(template_df),
                file_name=f"{key_prefix}_template.csv", mime="text/csv", key=f"{key_prefix}_tmpl"
            )
        up = st.file_uploader("Upload CSV file", type=["csv"], key=f"{key_prefix}_upload")
        if up is not None:
            df = pd.read_csv(up)
            st.session_state[f"{key_prefix}_df"] = df
        if f"{key_prefix}_df" in st.session_state:
            df = st.session_state[f"{key_prefix}_df"]

    field_glossary(schema)
    return df


# =============================================================
# USE CASE 1: AUTONOMOUS SOC / AI SECURITY ANALYST
# =============================================================
SOC_SCHEMA = {
    "timestamp": "Date/time the security event was observed.",
    "event_id": "Unique identifier for the raw event/alert.",
    "source_ip": "Originating IP address of the activity.",
    "dest_ip": "Destination IP address / targeted asset.",
    "user": "User or service account associated with the event.",
    "event_type": "Category of activity (login, file_access, network_conn, process_exec, dns_query).",
    "alert_source": "Telemetry/tool that generated the alert (EDR, SIEM, Firewall, IDS, Cloud).",
    "severity": "Analyst/vendor assigned severity (Low, Medium, High, Critical).",
    "mitre_tactic": "Mapped MITRE ATT&CK tactic (e.g., Initial Access, Persistence, Exfiltration).",
    "description": "Free-text description of what was observed.",
    "status": "Current triage status (New, Investigating, Escalated, Closed).",
}

def gen_soc_data(n):
    random.seed(); np.random.seed()
    users = [f"user{n_:03d}@corp.local" for n_ in range(1, 40)] + ["svc_backup", "svc_etl", "admin_jsmith"]
    event_types = ["login", "file_access", "network_conn", "process_exec", "dns_query"]
    alert_sources = ["EDR", "SIEM", "Firewall", "IDS", "CloudTrail", "IdentityGuard"]
    severities = np.random.choice(["Low", "Medium", "High", "Critical"], size=n, p=[0.40, 0.32, 0.20, 0.08])
    tactics = ["Initial Access", "Execution", "Persistence", "Privilege Escalation",
               "Defense Evasion", "Credential Access", "Lateral Movement", "Exfiltration", "Impact"]
    now = datetime.now()
    rows = []
    for i in range(n):
        ts = now - timedelta(minutes=random.randint(0, 60 * 24 * 7))
        rows.append({
            "timestamp": ts.strftime("%Y-%m-%d %H:%M:%S"),
            "event_id": f"EVT-{100000+i}",
            "source_ip": f"10.{random.randint(0,255)}.{random.randint(0,255)}.{random.randint(1,254)}",
            "dest_ip": f"172.16.{random.randint(0,255)}.{random.randint(1,254)}",
            "user": random.choice(users),
            "event_type": random.choice(event_types),
            "alert_source": random.choice(alert_sources),
            "severity": severities[i],
            "mitre_tactic": random.choice(tactics),
            "description": "Automated telemetry event captured by monitoring stack.",
            "status": random.choice(["New", "Investigating", "Escalated", "Closed"]),
        })
    return pd.DataFrame(rows)


def analyze_soc(df):
    sev_weight = {"Low": 10, "Medium": 35, "High": 65, "Critical": 90}
    df = df.copy()
    df["severity_score"] = df["severity"].map(sev_weight).fillna(20)

    # Correlate events into incidents by user + source_ip
    grp = df.groupby(["user", "source_ip"]).agg(
        event_count=("event_id", "count"),
        max_severity_score=("severity_score", "max"),
        avg_severity_score=("severity_score", "mean"),
        tactics_seen=("mitre_tactic", lambda x: ", ".join(sorted(set(x)))[:80]),
        alert_sources=("alert_source", lambda x: ", ".join(sorted(set(x)))),
        first_seen=("timestamp", "min"),
        last_seen=("timestamp", "max"),
    ).reset_index()

    grp["ai_risk_score"] = (
        grp["max_severity_score"] * 0.5
        + grp["avg_severity_score"] * 0.2
        + grp["event_count"].clip(upper=20) * 1.5
    ).round(1).clip(upper=100)

    def priority(score):
        if score >= 75: return "🔴 Critical — Auto-Escalate"
        if score >= 55: return "🟠 High — Analyst Review"
        if score >= 35: return "🟡 Medium — Monitor"
        return "🟢 Low — Informational"

    grp["ai_recommended_action"] = grp["ai_risk_score"].apply(priority)
    grp = grp.sort_values("ai_risk_score", ascending=False).reset_index(drop=True)
    return df, grp


def render_soc():
    module_banner(
        "Autonomous SOC / AI Security Analyst",
        "AI ingests raw alerts from EDR, SIEM, firewalls and cloud logs, then <b>detects</b> suspicious "
        "activity, <b>correlates</b> related events into incidents, <b>investigates</b> by scoring risk, "
        "and <b>recommends response</b> actions — cutting analyst triage time from hours to seconds. "
        "Business value: reduces alert fatigue, accelerates Mean-Time-to-Detect/Respond (MTTD/MTTR), "
        "and lets a small team cover enterprise-scale telemetry."
    )
    formula_box(
        "Raw events are grouped by **user + source IP** to approximate incident correlation. Each group "
        "is then scored for AI risk using a weighted blend of peak severity, average severity, and event "
        "frequency (capped at 100).",
        [
            ("Severity weight lookup", "Low=10, Medium=35, High=65, Critical=90"),
            ("AI Risk Score",
             "ai_risk_score = (0.5 x max_severity_score) + (0.2 x avg_severity_score)\n"
             "                + (1.5 x min(event_count, 20))\ncapped at 100"),
            ("Recommended action thresholds",
             ">=75: Critical - Auto-Escalate\n>=55: High - Analyst Review\n"
             ">=35: Medium - Monitor\n<35: Low - Informational"),
        ],
    )
    df = data_source_controls("soc", SOC_SCHEMA, gen_soc_data, default_n=200,
                               template_cols=list(SOC_SCHEMA.keys()))
    if df is None or df.empty:
        st.info("Generate synthetic data or upload a CSV to run the AI analysis.")
        return
    show_all_records(df, "soc")

    raw, incidents = analyze_soc(df)

    metric_row([
        ("Total Events", len(raw)),
        ("Correlated Incidents", len(incidents)),
        ("Critical Incidents", int((incidents["ai_risk_score"] >= 75).sum())),
        ("Avg Risk Score", round(incidents["ai_risk_score"].mean(), 1)),
    ])

    tab1, tab2, tab3 = st.tabs(["Analytics", "AI-Correlated Incidents", "Export"])

    with tab1:
        fig, axes = plt.subplots(1, 2, figsize=(11, 4))
        raw["severity"].value_counts().reindex(["Critical", "High", "Medium", "Low"]).plot(
            kind="bar", ax=axes[0], color=[ACCENT2, "#F4A259", "#F7D060", ACCENT])
        axes[0].set_title("Event Severity Distribution")
        axes[0].set_ylabel("Count")

        raw["event_type"].value_counts().plot(kind="barh", ax=axes[1], color=PRIMARY)
        axes[1].set_title("Events by Type")
        st.pyplot(fig)

        raw["ts_hour"] = pd.to_datetime(raw["timestamp"]).dt.floor("H")
        ts_counts = raw.groupby("ts_hour").size()
        fig2, ax2 = plt.subplots(figsize=(11, 3.5))
        ts_counts.plot(ax=ax2, color=ACCENT2)
        ax2.set_title("Event Volume Over Time")
        ax2.set_ylabel("Events")
        st.pyplot(fig2)

    with tab2:
        st.markdown("**AI-correlated incidents ranked by risk score** (events grouped by user + source IP, "
                     "scored using severity, frequency and blast radius).")
        st.dataframe(incidents, use_container_width=True, height=420)

    with tab3:
        narrative = (
            f"Autonomous SOC Analysis Summary\n"
            f"Total raw events analyzed: {len(raw)}\n"
            f"AI-correlated incidents: {len(incidents)}\n"
            f"Critical incidents requiring auto-escalation: {int((incidents['ai_risk_score']>=75).sum())}\n"
            f"Average AI risk score across incidents: {round(incidents['ai_risk_score'].mean(),1)}\n\n"
            "Methodology: Events are grouped by user + source IP to approximate incident correlation. "
            "Each group is scored using a weighted combination of maximum severity, average severity and "
            "event frequency, producing an AI risk score (0-100) and a recommended triage action."
        )
        export_panel("soc_incidents", "Autonomous SOC — Incident Report", narrative, incidents, fig)


# =============================================================
# USE CASE 2: AI VULNERABILITY & EXPOSURE MANAGEMENT
# =============================================================
VULN_SCHEMA = {
    "asset_id": "Unique identifier for the IT asset/host.",
    "hostname": "Human-readable name of the asset.",
    "ip_address": "Network address of the asset.",
    "os": "Operating system running on the asset.",
    "cve_id": "CVE identifier of the detected vulnerability.",
    "cvss_score": "Common Vulnerability Scoring System base score (0-10); higher = more severe.",
    "exploit_available": "Whether a public/working exploit exists (True/False).",
    "asset_criticality": "Business importance of the asset (Low, Medium, High, Critical).",
    "exposure": "Network exposure of the asset (Internet-Facing or Internal).",
    "patch_available": "Whether a vendor patch/fix currently exists (True/False).",
    "days_open": "Number of days the vulnerability has remained unremediated.",
}

def gen_vuln_data(n):
    random.seed(); np.random.seed()
    os_list = ["Windows Server 2019", "Windows 11", "Ubuntu 22.04", "RHEL 8", "macOS 14", "IOS-XE"]
    crit = np.random.choice(["Low", "Medium", "High", "Critical"], size=n, p=[0.25, 0.35, 0.28, 0.12])
    exposure = np.random.choice(["Internet-Facing", "Internal"], size=n, p=[0.25, 0.75])
    rows = []
    for i in range(n):
        rows.append({
            "asset_id": f"AST-{2000+i}",
            "hostname": f"host-{2000+i}.corp.local",
            "ip_address": f"192.168.{random.randint(0,255)}.{random.randint(1,254)}",
            "os": random.choice(os_list),
            "cve_id": f"CVE-2025-{random.randint(1000,9999)}",
            "cvss_score": round(random.uniform(2.0, 10.0), 1),
            "exploit_available": random.random() < 0.28,
            "asset_criticality": crit[i],
            "exposure": exposure[i],
            "patch_available": random.random() < 0.70,
            "days_open": random.randint(1, 240),
        })
    return pd.DataFrame(rows)


def analyze_vuln(df):
    df = df.copy()
    crit_w = {"Low": 5, "Medium": 10, "High": 18, "Critical": 25}
    df["asset_criticality"] = df["asset_criticality"].astype(str)
    df["exploit_available"] = df["exploit_available"].astype(str).str.lower().isin(["true", "1", "yes"])
    df["patch_available"] = df["patch_available"].astype(str).str.lower().isin(["true", "1", "yes"])

    df["ai_priority_score"] = (
        df["cvss_score"] * 6
        + df["exploit_available"].astype(int) * 18
        + df["asset_criticality"].map(crit_w).fillna(8)
        + (df["exposure"] == "Internet-Facing").astype(int) * 15
        + (~df["patch_available"]).astype(int) * 8
        + (df["days_open"] / 240 * 10)
    ).round(1).clip(upper=100)

    def tier(s):
        if s >= 75: return "🔴 Critical — Remediate Now"
        if s >= 55: return "🟠 High — This Sprint"
        if s >= 35: return "🟡 Medium — Scheduled"
        return "🟢 Low — Backlog"

    df["risk_tier"] = df["ai_priority_score"].apply(tier)

    def recommend(row):
        if row["exploit_available"] and row["exposure"] == "Internet-Facing":
            return "Auto-isolate asset + emergency patch / virtual patch via WAF-IPS"
        if not row["patch_available"]:
            return "Apply compensating control (segmentation) until patch is released"
        if row["ai_priority_score"] >= 55:
            return "Schedule patch deployment within 7 days"
        return "Include in standard monthly patch cycle"

    df["ai_recommended_remediation"] = df.apply(recommend, axis=1)
    df = df.sort_values("ai_priority_score", ascending=False).reset_index(drop=True)
    return df


def render_vuln():
    module_banner(
        "AI Vulnerability & Exposure Management",
        "AI continuously ingests vulnerability scan data, correlates it with exploit intelligence, "
        "asset criticality and internet exposure, then <b>prioritizes true business risk</b> — not just "
        "raw CVSS — and recommends or automates remediation. Business value: focuses limited patching "
        "resources on what attackers would actually exploit, shrinking the exploitable attack surface faster."
    )
    formula_box(
        "Each vulnerability's raw CVSS score is re-weighted using exploit availability, asset criticality, "
        "internet exposure, patch availability and age to produce a single business-risk priority score.",
        [
            ("Criticality weight lookup", "Low=5, Medium=10, High=18, Critical=25"),
            ("AI Priority Score",
             "ai_priority_score = (cvss_score x 6) + (exploit_available x 18)\n"
             "                   + criticality_weight + (internet_facing x 15)\n"
             "                   + (no_patch_available x 8) + (days_open / 240 x 10)\n"
             "capped at 100"),
            ("Risk tier thresholds",
             ">=75: Critical - Remediate Now\n>=55: High - This Sprint\n"
             ">=35: Medium - Scheduled\n<35: Low - Backlog"),
        ],
    )
    df = data_source_controls("vuln", VULN_SCHEMA, gen_vuln_data, default_n=200,
                               template_cols=list(VULN_SCHEMA.keys()))
    if df is None or df.empty:
        st.info("Generate synthetic data or upload a CSV to run the AI analysis.")
        return
    show_all_records(df, "vuln")

    result = analyze_vuln(df)

    metric_row([
        ("Total Vulnerabilities", len(result)),
        ("Critical Priority", int((result["ai_priority_score"] >= 75).sum())),
        ("Exploitable & Internet-Facing", int(((result["exploit_available"]) & (result["exposure"] == "Internet-Facing")).sum())),
        ("Avg Priority Score", round(result["ai_priority_score"].mean(), 1)),
    ])

    tab1, tab2, tab3 = st.tabs(["Analytics", "AI-Prioritized Vulnerabilities", "Export"])

    with tab1:
        fig, axes = plt.subplots(1, 2, figsize=(11, 4))
        colors_map = {"Internet-Facing": ACCENT2, "Internal": ACCENT}
        for exp, sub in result.groupby("exposure"):
            axes[0].scatter(sub["cvss_score"], sub["ai_priority_score"], label=exp,
                             alpha=0.6, color=colors_map.get(exp, PRIMARY))
        axes[0].set_xlabel("CVSS Score"); axes[0].set_ylabel("AI Priority Score")
        axes[0].set_title("CVSS vs AI-Weighted Priority"); axes[0].legend()

        result["risk_tier"].value_counts().plot(kind="bar", ax=axes[1], color=PRIMARY)
        axes[1].set_title("Vulnerabilities by Risk Tier")
        st.pyplot(fig)

    with tab2:
        st.markdown("**AI-prioritized remediation queue** — ranked by business risk, not raw CVSS alone.")
        st.dataframe(result, use_container_width=True, height=420)

    with tab3:
        narrative = (
            f"AI Vulnerability & Exposure Management Summary\n"
            f"Total vulnerabilities assessed: {len(result)}\n"
            f"Critical-priority items: {int((result['ai_priority_score']>=75).sum())}\n"
            f"Exploitable AND internet-facing: {int(((result['exploit_available'])&(result['exposure']=='Internet-Facing')).sum())}\n"
            f"Average AI priority score: {round(result['ai_priority_score'].mean(),1)}\n\n"
            "Methodology: AI priority score blends CVSS severity, public exploit availability, asset "
            "business criticality, internet exposure, patch availability and vulnerability age into a "
            "single 0-100 business-risk score, driving remediation recommendations."
        )
        export_panel("vuln_priorities", "AI Vulnerability Management — Priority Report", narrative, result, fig)


# =============================================================
# USE CASE 3: BEHAVIORAL THREAT & ANOMALY DETECTION
# =============================================================
BEHAV_SCHEMA = {
    "timestamp": "Date/time the metric was recorded.",
    "entity_id": "Identifier of the monitored entity (user, device, API client, AI agent).",
    "entity_type": "Type of entity (user, device, api, ai_agent).",
    "metric_name": "Behavioral metric being measured (login_count, data_transfer_mb, api_calls, geo_velocity_kmh).",
    "value": "Observed value of the metric for this entity/time window.",
    "baseline_mean": "Historical average value for this entity (its normal behavior).",
    "baseline_std": "Historical standard deviation for this entity (normal variability).",
}

def gen_behav_data(n):
    random.seed(); np.random.seed()
    entity_types = ["user", "device", "api", "ai_agent"]
    metrics_by_type = {
        "user": ["login_count", "data_transfer_mb", "failed_logins"],
        "device": ["network_conn_count", "data_transfer_mb", "cpu_util_pct"],
        "api": ["api_calls", "error_rate_pct", "data_transfer_mb"],
        "ai_agent": ["api_calls", "tool_invocations", "data_transfer_mb"],
    }
    now = datetime.now()
    rows = []
    for i in range(n):
        etype = random.choice(entity_types)
        metric = random.choice(metrics_by_type[etype])
        baseline_mean = round(random.uniform(5, 200), 1)
        baseline_std = round(baseline_mean * random.uniform(0.08, 0.25), 1)
        is_anomalous = random.random() < 0.08
        if is_anomalous:
            value = round(baseline_mean + baseline_std * random.uniform(4, 9) * random.choice([1, -1]), 1)
            value = max(value, 0)
        else:
            value = round(max(np.random.normal(baseline_mean, baseline_std), 0), 1)
        ts = now - timedelta(minutes=random.randint(0, 60 * 24 * 7))
        rows.append({
            "timestamp": ts.strftime("%Y-%m-%d %H:%M:%S"),
            "entity_id": f"{etype[:3].upper()}-{random.randint(1000,1200)}",
            "entity_type": etype,
            "metric_name": metric,
            "value": value,
            "baseline_mean": baseline_mean,
            "baseline_std": baseline_std,
        })
    return pd.DataFrame(rows)


def analyze_behav(df):
    df = df.copy()
    df["deviation"] = (df["value"] - df["baseline_mean"]) / df["baseline_std"].replace(0, 1)
    df["z_abs"] = df["deviation"].abs()

    features = df[["value", "baseline_mean", "baseline_std", "deviation"]].fillna(0)
    contamination = min(max(0.03, (df["z_abs"] > 3).mean() + 0.02), 0.25)
    model = IsolationForest(n_estimators=200, contamination=contamination, random_state=42)
    df["ml_anomaly_flag"] = model.fit_predict(features)
    df["ml_anomaly_flag"] = df["ml_anomaly_flag"].map({-1: "🚨 Anomaly", 1: "Normal"})
    df["anomaly_score"] = (-model.decision_function(features)).round(3)

    df = df.sort_values("anomaly_score", ascending=False).reset_index(drop=True)
    return df


def render_behav():
    module_banner(
        "Behavioral Threat &amp; Anomaly Detection",
        "AI builds a behavioral baseline for every user, device, API client and AI agent, then uses "
        "unsupervised machine learning (Isolation Forest) to flag statistically unusual activity — "
        "impossible travel, data-volume spikes, abnormal API/agent behavior — <b>before</b> a signature "
        "or rule exists for it. Business value: catches novel and insider threats that rule-based tools miss."
    )
    formula_box(
        "Each observation's deviation from its own historical baseline is combined with its raw feature "
        "values and fed into an unsupervised Isolation Forest model, which isolates statistically "
        "anomalous points without needing predefined rules or signatures.",
        [
            ("Deviation (z-score)", "deviation = (value - baseline_mean) / baseline_std"),
            ("Feature vector fed to model", "[value, baseline_mean, baseline_std, deviation]"),
            ("Model", "IsolationForest(n_estimators=200, contamination=auto-tuned 3%-25%, random_state=42)"),
            ("Anomaly score", "anomaly_score = -1 x model.decision_function(features)\n"
             "flag = 'Anomaly' if model predicts -1, else 'Normal'"),
        ],
    )
    df = data_source_controls("behav", BEHAV_SCHEMA, gen_behav_data, default_n=200,
                               template_cols=list(BEHAV_SCHEMA.keys()))
    if df is None or df.empty:
        st.info("Generate synthetic data or upload a CSV to run the AI analysis.")
        return
    show_all_records(df, "behav")

    result = analyze_behav(df)
    n_anom = int((result["ml_anomaly_flag"] == "🚨 Anomaly").sum())

    metric_row([
        ("Entities/Events Analyzed", len(result)),
        ("Anomalies Detected", n_anom),
        ("Anomaly Rate", f"{round(100*n_anom/len(result),1)}%"),
        ("Entity Types Monitored", result["entity_type"].nunique()),
    ])

    tab1, tab2, tab3 = st.tabs(["Analytics", "AI-Flagged Anomalies", "Export"])

    with tab1:
        fig, axes = plt.subplots(1, 2, figsize=(11, 4))
        normal = result[result["ml_anomaly_flag"] == "Normal"]
        anom = result[result["ml_anomaly_flag"] == "🚨 Anomaly"]
        axes[0].scatter(normal["baseline_mean"], normal["value"], alpha=0.4, color=ACCENT, label="Normal")
        axes[0].scatter(anom["baseline_mean"], anom["value"], alpha=0.9, color=ACCENT2, label="Anomaly")
        axes[0].set_xlabel("Baseline Mean"); axes[0].set_ylabel("Observed Value")
        axes[0].set_title("Observed vs Baseline (Anomalies Highlighted)"); axes[0].legend()

        result.groupby("entity_type")["ml_anomaly_flag"].apply(
            lambda x: (x == "🚨 Anomaly").sum()).plot(kind="bar", ax=axes[1], color=ACCENT2)
        axes[1].set_title("Anomalies by Entity Type")
        st.pyplot(fig)

    with tab2:
        st.markdown("**AI-flagged anomalies**, ranked by anomaly score (Isolation Forest, unsupervised).")
        st.dataframe(anom.drop(columns=["z_abs"]) if not anom.empty else result.head(0),
                     use_container_width=True, height=380)

    with tab3:
        narrative = (
            f"Behavioral Threat & Anomaly Detection Summary\n"
            f"Total entity/event observations analyzed: {len(result)}\n"
            f"Anomalies detected by Isolation Forest model: {n_anom} ({round(100*n_anom/len(result),1)}%)\n"
            f"Entity types monitored: {', '.join(result['entity_type'].unique())}\n\n"
            "Methodology: Each observation's deviation from its own historical baseline (z-score) is "
            "combined with raw feature values and fed into an unsupervised Isolation Forest model, which "
            "isolates statistically anomalous behavior without relying on predefined signatures."
        )
        export_panel("behavioral_anomalies", "Behavioral Threat & Anomaly Detection — Report",
                      narrative, anom if not anom.empty else result, fig)


# =============================================================
# USE CASE 4: AI INCIDENT RESPONSE & ATTACK-PATH ANALYSIS
# =============================================================
IR_SCHEMA = {
    "incident_id": "Identifier grouping related attack steps into one incident/kill-chain.",
    "step_order": "Sequence number of the step within the attack chain.",
    "source_entity": "Entity (host/user/service) the step originated from.",
    "target_entity": "Entity the step acted upon / moved to.",
    "technique": "MITRE ATT&CK technique ID associated with the step (e.g., T1078, T1021).",
    "action": "Description of the action performed at this step.",
    "timestamp": "Date/time the step occurred.",
    "contained": "Whether this step/entity has been automatically or manually contained (True/False).",
}

MITRE_TECHNIQUES = [
    ("T1566", "Phishing"), ("T1078", "Valid Accounts"), ("T1059", "Command Execution"),
    ("T1021", "Remote Services / Lateral Movement"), ("T1003", "Credential Dumping"),
    ("T1486", "Data Encrypted for Impact"), ("T1041", "Exfiltration Over C2 Channel"),
    ("T1547", "Boot/Logon Autostart Persistence"), ("T1068", "Privilege Escalation Exploit"),
]

def gen_ir_data(n_incidents):
    random.seed(); np.random.seed()
    rows = []
    now = datetime.now()
    for i in range(1, n_incidents + 1):
        incident_id = f"INC-{5000+i}"
        chain_len = random.randint(3, 7)
        entities = [f"host-{random.randint(100,999)}"] + \
                   [f"host-{random.randint(100,999)}" for _ in range(chain_len)]
        ts = now - timedelta(hours=random.randint(1, 240))
        for step in range(chain_len):
            tech_id, tech_name = random.choice(MITRE_TECHNIQUES)
            ts = ts + timedelta(minutes=random.randint(5, 90))
            rows.append({
                "incident_id": incident_id,
                "step_order": step + 1,
                "source_entity": entities[step],
                "target_entity": entities[step + 1],
                "technique": f"{tech_id} - {tech_name}",
                "action": f"{tech_name} observed from {entities[step]} to {entities[step+1]}",
                "timestamp": ts.strftime("%Y-%m-%d %H:%M:%S"),
                "contained": random.random() < 0.35,
            })
    return pd.DataFrame(rows)


def analyze_ir(df):
    df = df.copy()
    df["contained"] = df["contained"].astype(str).str.lower().isin(["true", "1", "yes"])
    summary = df.groupby("incident_id").agg(
        chain_length=("step_order", "max"),
        entities_involved=("target_entity", "nunique"),
        techniques_used=("technique", lambda x: x.nunique()),
        contained_steps=("contained", "sum"),
        first_seen=("timestamp", "min"),
        last_seen=("timestamp", "max"),
    ).reset_index()
    summary["containment_pct"] = (summary["contained_steps"] / summary["chain_length"] * 100).round(1)

    def sev(row):
        if row["chain_length"] >= 6 or row["entities_involved"] >= 5:
            return "🔴 Critical"
        if row["chain_length"] >= 4:
            return "🟠 High"
        return "🟡 Medium"
    summary["ai_severity"] = summary.apply(sev, axis=1)

    def rec(row):
        if row["containment_pct"] < 50 and "Critical" in row["ai_severity"]:
            return "Auto-isolate all involved hosts + force credential reset + block C2 egress"
        if row["containment_pct"] < 100:
            return "Complete containment on remaining uncontained hosts; validate with EDR"
        return "Fully contained — proceed to root-cause & lessons-learned review"
    summary["ai_response_recommendation"] = summary.apply(rec, axis=1)
    summary = summary.sort_values("chain_length", ascending=False).reset_index(drop=True)
    return df, summary


def draw_attack_graph(steps_df):
    G = nx.DiGraph()
    for _, r in steps_df.iterrows():
        G.add_edge(r["source_entity"], r["target_entity"], label=r["technique"].split(" - ")[0])
    fig, ax = plt.subplots(figsize=(8, 5))
    pos = nx.spring_layout(G, seed=42, k=1.2)
    nx.draw_networkx_nodes(G, pos, node_color=ACCENT, node_size=1200, ax=ax)
    nx.draw_networkx_labels(G, pos, font_size=8, ax=ax)
    nx.draw_networkx_edges(G, pos, edge_color=PRIMARY, arrows=True, arrowsize=18, ax=ax,
                            connectionstyle="arc3,rad=0.08")
    edge_labels = nx.get_edge_attributes(G, "label")
    nx.draw_networkx_edge_labels(G, pos, edge_labels=edge_labels, font_size=7, ax=ax)
    ax.set_title("AI-Reconstructed Attack Path")
    ax.axis("off")
    return fig


def render_ir():
    module_banner(
        "AI Incident Response &amp; Attack-Path Analysis",
        "AI stitches together individual security events into a full <b>attack chain / kill chain</b>, "
        "mapping each hop to MITRE ATT&amp;CK techniques, then automatically recommends or executes "
        "containment (isolate hosts, revoke credentials, block C2). Business value: compresses "
        "investigation time from days to minutes and limits blast radius through faster containment."
    )
    formula_box(
        "Individual telemetry steps are chained by incident ID into a directed attack graph (each "
        "incident has 3-7 steps, ~5 on average, so ~40 incidents yields roughly 200 attack-step records). "
        "Severity is derived from chain length and blast radius (unique entities touched).",
        [
            ("Containment coverage", "containment_pct = (contained_steps / chain_length) x 100"),
            ("AI severity rule",
             "Critical: chain_length >= 6 OR entities_involved >= 5\n"
             "High: chain_length >= 4\nMedium: otherwise"),
            ("Response recommendation rule",
             "If containment_pct < 50 AND severity = Critical:\n"
             "  Auto-isolate hosts + force credential reset + block C2 egress\n"
             "If containment_pct < 100: Complete containment on remaining hosts\n"
             "Else: Fully contained - proceed to root-cause review"),
        ],
    )
    df = data_source_controls("ir", IR_SCHEMA, gen_ir_data, default_n=40,
                               template_cols=list(IR_SCHEMA.keys()),
                               slider_label="Number of synthetic incidents (~5 attack-chain records each)",
                               slider_min=10, slider_max=200, slider_step=5)
    if df is None or df.empty:
        st.info("Generate synthetic data or upload a CSV to run the AI analysis.")
        return
    show_all_records(df, "ir")

    steps, summary = analyze_ir(df)

    metric_row([
        ("Incidents Analyzed", len(summary)),
        ("Critical Severity", int((summary["ai_severity"] == "🔴 Critical").sum())),
        ("Avg Chain Length", round(summary["chain_length"].mean(), 1)),
        ("Avg Containment %", f"{round(summary['containment_pct'].mean(),1)}%"),
    ])

    tab1, tab2, tab3 = st.tabs(["Attack Path Visualization", "Incident Summary", "Export"])

    with tab1:
        chosen = st.selectbox("Select an incident to visualize its attack path", summary["incident_id"].tolist())
        sub_steps = steps[steps["incident_id"] == chosen].sort_values("step_order")
        st.dataframe(sub_steps, use_container_width=True, height=200)
        fig = draw_attack_graph(sub_steps)
        st.pyplot(fig)

    with tab2:
        st.markdown("**AI incident severity & containment summary**, ranked by attack-chain length.")
        st.dataframe(summary, use_container_width=True, height=380)

    with tab3:
        narrative = (
            f"AI Incident Response & Attack-Path Analysis Summary\n"
            f"Incidents analyzed: {len(summary)}\n"
            f"Critical-severity incidents: {int((summary['ai_severity']=='🔴 Critical').sum())}\n"
            f"Average attack-chain length: {round(summary['chain_length'].mean(),1)} steps\n"
            f"Average containment coverage: {round(summary['containment_pct'].mean(),1)}%\n\n"
            "Methodology: Individual telemetry steps are chained by incident ID into a directed attack "
            "graph, mapped to MITRE ATT&CK techniques. Severity is derived from chain length and blast "
            "radius (unique entities touched); containment recommendations are generated per incident."
        )
        export_panel("incident_response", "AI Incident Response — Attack-Path Report", narrative, summary, fig)


# =============================================================
# USE CASE 5: AI IDENTITY & AGENT SECURITY
# =============================================================
ID_SCHEMA = {
    "identity_id": "Unique identifier for the identity (human user, service account, or AI agent).",
    "identity_type": "Type of identity: human, service_account, or ai_agent.",
    "privilege_count": "Number of roles/permissions/entitlements granted to the identity.",
    "mfa_enabled": "Whether multi-factor authentication is enforced (True/False).",
    "last_activity_days_ago": "Days since the identity was last active (higher = more dormant).",
    "anomalous_access_count": "Number of anomalous access events attributed to this identity in the period.",
    "is_privileged_admin": "Whether the identity holds admin/root/owner-level privileges (True/False).",
}

def gen_identity_data(n):
    random.seed(); np.random.seed()
    id_types = np.random.choice(["human", "service_account", "ai_agent"], size=n, p=[0.65, 0.20, 0.15])
    rows = []
    for i in range(n):
        itype = id_types[i]
        priv = random.randint(1, 6) if itype == "human" else random.randint(1, 15)
        rows.append({
            "identity_id": f"{itype[:3].upper()}-{4000+i}",
            "identity_type": itype,
            "privilege_count": priv,
            "mfa_enabled": random.random() < (0.85 if itype == "human" else 0.45),
            "last_activity_days_ago": random.choice([0,1,2,3,5,7,14,30,60,90,180]),
            "anomalous_access_count": np.random.choice([0,0,0,0,1,2,3,5], p=[0.45,0.15,0.1,0.1,0.08,0.05,0.04,0.03]),
            "is_privileged_admin": random.random() < 0.15,
        })
    return pd.DataFrame(rows)


def analyze_identity(df):
    df = df.copy()
    df["mfa_enabled"] = df["mfa_enabled"].astype(str).str.lower().isin(["true", "1", "yes"])
    df["is_privileged_admin"] = df["is_privileged_admin"].astype(str).str.lower().isin(["true", "1", "yes"])

    df["ai_identity_risk_score"] = (
        (~df["mfa_enabled"]).astype(int) * 22
        + df["is_privileged_admin"].astype(int) * 20
        + df["privilege_count"].clip(upper=15) * 2
        + df["anomalous_access_count"].clip(upper=10) * 6
        + (df["last_activity_days_ago"] >= 60).astype(int) * 15
    ).round(1).clip(upper=100)

    def flag(row):
        tags = []
        if not row["mfa_enabled"]:
            tags.append("No MFA")
        if row["is_privileged_admin"] and row["privilege_count"] > 8:
            tags.append("Over-Privileged Admin")
        if row["last_activity_days_ago"] >= 60:
            tags.append("Dormant")
        if row["anomalous_access_count"] >= 2:
            tags.append("Anomalous Access")
        return ", ".join(tags) if tags else "Nominal"

    df["ai_risk_flags"] = df.apply(flag, axis=1)

    def rec(score):
        if score >= 70: return "Suspend / require re-certification + enforce MFA immediately"
        if score >= 45: return "Reduce privileges to least-privilege + enforce MFA"
        return "Monitor — no immediate action required"
    df["ai_recommended_action"] = df["ai_identity_risk_score"].apply(rec)

    df = df.sort_values("ai_identity_risk_score", ascending=False).reset_index(drop=True)
    return df


def render_identity():
    module_banner(
        "AI Identity &amp; Agent Security",
        "AI continuously monitors human users, service accounts and non-human <b>AI agents</b>, "
        "evaluating privilege levels, MFA posture, dormancy and anomalous access to compute a real-time "
        "<b>identity risk score</b> — then recommends least-privilege remediation. Business value: shrinks "
        "the identity attack surface, a leading cause of breaches, including new risks from autonomous AI agents."
    )
    formula_box(
        "Every identity (human, service account, or AI agent) is scored using its MFA posture, "
        "privilege level, entitlement count, dormancy and anomalous access history.",
        [
            ("AI Identity Risk Score",
             "ai_identity_risk_score = (no_mfa x 22) + (is_privileged_admin x 20)\n"
             "  + (min(privilege_count, 15) x 2) + (min(anomalous_access_count, 10) x 6)\n"
             "  + (last_activity_days_ago >= 60 x 15)\ncapped at 100"),
            ("Recommended action thresholds",
             ">=70: Suspend / require re-certification + enforce MFA immediately\n"
             ">=45: Reduce privileges to least-privilege + enforce MFA\n"
             "<45: Monitor - no immediate action required"),
        ],
    )
    df = data_source_controls("identity", ID_SCHEMA, gen_identity_data, default_n=200,
                               template_cols=list(ID_SCHEMA.keys()))
    if df is None or df.empty:
        st.info("Generate synthetic data or upload a CSV to run the AI analysis.")
        return
    show_all_records(df, "identity")

    result = analyze_identity(df)
    n_high = int((result["ai_identity_risk_score"] >= 70).sum())

    metric_row([
        ("Identities Monitored", len(result)),
        ("High-Risk Identities", n_high),
        ("No-MFA Identities", int((~result["mfa_enabled"]).sum())),
        ("AI Agents Monitored", int((result["identity_type"] == "ai_agent").sum())),
    ])

    tab1, tab2, tab3 = st.tabs(["Analytics", "AI Identity Risk Register", "Export"])

    with tab1:
        fig, axes = plt.subplots(1, 2, figsize=(11, 4))
        result.groupby("identity_type")["ai_identity_risk_score"].mean().plot(
            kind="bar", ax=axes[0], color=[PRIMARY, ACCENT, ACCENT2])
        axes[0].set_title("Avg Risk Score by Identity Type")

        result["ai_risk_flags"].apply(lambda s: "Nominal" if s == "Nominal" else "Flagged").value_counts().plot(
            kind="pie", ax=axes[1], autopct="%1.0f%%", colors=[ACCENT, ACCENT2])
        axes[1].set_ylabel("")
        axes[1].set_title("Flagged vs Nominal Identities")
        st.pyplot(fig)

    with tab2:
        st.markdown("**AI identity risk register**, ranked by computed risk score.")
        st.dataframe(result, use_container_width=True, height=420)

    with tab3:
        narrative = (
            f"AI Identity & Agent Security Summary\n"
            f"Total identities monitored: {len(result)} "
            f"(human: {int((result['identity_type']=='human').sum())}, "
            f"service accounts: {int((result['identity_type']=='service_account').sum())}, "
            f"AI agents: {int((result['identity_type']=='ai_agent').sum())})\n"
            f"High-risk identities (score >= 70): {n_high}\n"
            f"Identities without MFA: {int((~result['mfa_enabled']).sum())}\n\n"
            "Methodology: AI identity risk score combines MFA posture, admin/privilege level, total "
            "entitlement count, dormancy (days since last activity) and anomalous access events into a "
            "single 0-100 score, with tailored least-privilege remediation recommendations."
        )
        export_panel("identity_risk", "AI Identity & Agent Security — Risk Register", narrative, result, fig)


# =============================================================
# HOME / OVERVIEW PAGE
# =============================================================
def render_home():
    module_banner(
        f"Welcome to {APP_NAME}",
        "A unified demonstration platform showcasing five production-representative AI/ML cybersecurity "
        "use cases. Select a module from the left sidebar to explore synthetic-data demos, upload your "
        "own data, and export professional reports."
    )

    overview = pd.DataFrame([
        ["1", "Autonomous SOC / AI Security Analyst", "Detect, investigate, correlate, respond", "Very High"],
        ["2", "AI Vulnerability & Exposure Management", "Find vulnerabilities, prioritize risk, recommend remediation", "Very High"],
        ["3", "Behavioral Threat & Anomaly Detection", "Detect unusual users, devices, APIs, network traffic and AI agents", "Very High"],
        ["4", "AI Incident Response & Attack-Path Analysis", "Determine attack chain and automatically contain threats", "Very High"],
        ["5", "AI Identity & Agent Security", "Monitor human/non-human identities, privileges and agent activity", "Very High"],
    ], columns=["#", "AI/ML Use Case", "What AI Does", "Business Value"])
    st.table(overview)

    st.markdown("### How to use this application")
    st.markdown(
        """
        1. Choose a use case from the **sidebar**.
        2. Either click **Generate Synthetic Data** to see the AI in action instantly (~200 records by
           default, and every generated/uploaded record is shown in the results), or **upload your own CSV**
           (a downloadable template is provided in each module showing the required schema).
        3. Open **AI Methodology & Formulas Used** in each module to see exactly how each score is calculated.
        4. Review the AI-generated **analytics, scoring, and recommendations**.
        5. Use the **Export** tab to download results as **PDF, Word, CSV, or Text** for reporting.
        """
    )
    st.markdown("### About")
    st.markdown(
        f"This application was designed and developed by **Randy Singh** from **Kalsnet (KNet) Consulting "
        "group** to demonstrate practical, explainable AI/ML applications across the modern Security "
        "Operations Center. All data shown by default is synthetically generated for demonstration purposes only."
    )


# =============================================================
# MAIN APP / NAVIGATION
# =============================================================
def main():
    with st.sidebar:
        st.markdown("## Navigation")
        page = st.radio(
            "Select a module",
            [
                "Home / Overview",
                "1. Autonomous SOC / AI Security Analyst",
                "2. AI Vulnerability & Exposure Management",
                "3. Behavioral Threat & Anomaly Detection",
                "4. AI Incident Response & Attack-Path Analysis",
                "5. AI Identity & Agent Security",
            ],
        )
        st.markdown("---")
        st.caption(APP_NAME)
        st.caption(DEVELOPER_LINE)

    subtitle_map = {
        "Home / Overview": "Unified AI/ML Cybersecurity Demonstration Platform",
        "1. Autonomous SOC / AI Security Analyst": "Module 1 of 5 - Detect, Investigate, Correlate, Respond",
        "2. AI Vulnerability & Exposure Management": "Module 2 of 5 - Find, Prioritize, Remediate",
        "3. Behavioral Threat & Anomaly Detection": "Module 3 of 5 - Users, Devices, APIs, AI Agents",
        "4. AI Incident Response & Attack-Path Analysis": "Module 4 of 5 - Attack Chain Reconstruction & Containment",
        "5. AI Identity & Agent Security": "Module 5 of 5 - Human & Non-Human Identity Risk",
    }
    render_title_bar(subtitle_map[page])

    if page.startswith("Home"):
        render_home()
    elif page.startswith("1."):
        render_soc()
    elif page.startswith("2."):
        render_vuln()
    elif page.startswith("3."):
        render_behav()
    elif page.startswith("4."):
        render_ir()
    elif page.startswith("5."):
        render_identity()

    render_footer()


if __name__ == "__main__":
    main()