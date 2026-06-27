# ChainIntel Pro
# AI-powered supply chain intelligence system for procurement, demand forecasting,
# inventory optimization, and vendor risk management.
# Developed by Randy Singh | Kalsnet (KNet) Consulting
# Focus: Reduce stockouts • Reduce dead stock • Faster procurement • Vendor risk control
# ROI: 8-14 months payback through inventory reduction + vendor optimization
import io
import json
import logging
import os
import random
import uuid
import datetime as dt
from typing import Dict, List, Tuple
import pandas as pd
import numpy as np
import streamlit as st
try:
    from fpdf import FPDF
    HAS_PDF = True
except ImportError:
    HAS_PDF = False
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
# =========================================================================
# LOGGING
# =========================================================================
LOG_DIR = os.environ.get("CHAINintel_LOG_DIR", "./logs")
os.makedirs(LOG_DIR, exist_ok=True)
logger = logging.getLogger("chaintel_pro")
if not logger.handlers:
    logger.setLevel(logging.INFO)
    _fh = logging.FileHandler(os.path.join(LOG_DIR, "chaintel_pro.log"))
    _fh.setFormatter(logging.Formatter("%(asctime)s | %(levelname)s | %(message)s"))
    logger.addHandler(_fh)
# =========================================================================
# PAGE CONFIG
# =========================================================================
st.set_page_config(
    page_title="ChainIntel Pro | Supply Chain AI",
    page_icon="⚙️",
    layout="wide",
    initial_sidebar_state="expanded",
)
# =========================================================================
# PROFESSIONAL STYLING (SAME AS SHOPFLOOR)
# =========================================================================
st.markdown(
    """
    <style>
    .main-title {
        text-align: center;
        color: #0047AB;
        font-weight: 900;
        font-size: 2.8rem;
        margin-bottom: 0.3rem;
        letter-spacing: 0.8px;
        font-family: 'Arial Black', sans-serif;
    }
    .subtitle-line {
        text-align: center;
        color: #0047AB;
        font-weight: 800;
        font-size: 1.2rem;
        margin-top: 0;
        margin-bottom: 1.2rem;
        font-family: 'Arial', sans-serif;
    }
    .section-header {
        color: #0047AB;
        font-weight: 800;
        font-size: 1.8rem;
        margin-top: 1.5rem;
        margin-bottom: 0.8rem;
        border-bottom: 3px solid #0047AB;
        padding-bottom: 0.5rem;
    }
    .subsection-header {
        color: #003380;
        font-weight: 700;
        font-size: 1.3rem;
        margin-top: 1rem;
        margin-bottom: 0.6rem;
    }
    .metric-card {
        background: #f0f7ff;
        border-left: 5px solid #0047AB;
        padding: 1rem;
        border-radius: 8px;
        margin: 0.8rem 0;
    }
    .formula-box {
        background: #fff9e6;
        border: 2px solid #ffd700;
        padding: 1rem;
        border-radius: 6px;
        font-family: 'Courier New', monospace;
        font-size: 0.9rem;
        margin: 0.8rem 0;
    }
    .info-box {
        background: #e8f4f8;
        border-left: 4px solid #0099cc;
        padding: 0.8rem;
        border-radius: 4px;
        font-size: 0.85rem;
        margin: 0.6rem 0;
    }
    .divider {
        border-top: 2px solid #0047AB;
        margin: 1.5rem 0;
        opacity: 0.3;
    }
    </style>
    """,
    unsafe_allow_html=True,
)
# =========================================================================
# TITLE BAR
# =========================================================================
st.markdown('<div class="main-title">⚙️ ChainIntel Pro</div>', unsafe_allow_html=True)
st.markdown('<div class="subtitle-line">Developed by Randy Singh from Kalsnet (KNet) Consulting</div>', unsafe_allow_html=True)
st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
st.markdown(
    """
    **Supply Chain Intelligence Platform** — AI-powered demand forecasting, inventory optimization,
    vendor risk scoring, and procurement automation. Reduce stockouts, reduce dead stock, faster quotes.

    **ROI Focus:** Better forecasts • Smarter inventory • Vendor control • Faster procurement
    """
)
# =========================================================================
# ENVIRONMENT CONFIG
# =========================================================================
DEPLOY_ENV = os.environ.get("CHAINTEL_ENV", "production")
MAX_UPLOAD_MB = int(os.environ.get("CHAINTEL_MAX_UPLOAD_MB", "25"))
# =========================================================================
# CONSTANTS
# =========================================================================
MATERIAL_TYPES = ["Raw Material", "Work-in-Progress", "Finished Goods", "Packaging", "Consumables"]
OEM_NAMES = ["Maruti Suzuki", "Tata Motors", "Hyundai", "Kia Motors", "MG Motor", "Skoda", "Volkswagen"]
MATERIAL_NAMES = {
    "Aluminum": {"unit": "kg", "density": 2.7, "base_price": 150},
    "Cast Iron": {"unit": "kg", "density": 7.8, "base_price": 80},
    "Steel": {"unit": "kg", "density": 7.85, "base_price": 100},
    "Brass": {"unit": "kg", "density": 8.5, "base_price": 400},
    "Plastic Resin": {"unit": "kg", "density": 1.2, "base_price": 200},
}
# =========================================================================
# SESSION STATE
# =========================================================================
DEFAULTS = {
    "forecast_df": None,
    "real_forecast_df": None,
    "inventory_df": None,
    "inventory_sku_df": None,
    "vendor_df": None,
    "real_vendor_df": None,
    "working_capital_result": None,
    "audit_log": [],
}
for k, v in DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = v
if "session_id" not in st.session_state:
    st.session_state.session_id = str(uuid.uuid4())[:8]
# =========================================================================
# SUPPLY CHAIN FORMULAS
# =========================================================================
class SupplyChainFormulas:
    """Supply chain AI formulas with business logic"""
    @staticmethod
    def forecast_demand(historical_avg: float, seasonality_factor: float = 1.0,
                       trend: float = 0.0, safety_stock_days: int = 7) -> float:
        """
        DEMAND FORECASTING FORMULA

        Forecast = Historical_Avg × Seasonality × (1 + Trend) + Safety_Stock

        Where:
        - Historical_Avg: Average daily/weekly demand (from last 12 months)
        - Seasonality_Factor: Month/quarter multiplier (1.0 = normal, 1.2 = 20% higher)
        - Trend: Percentage change (-0.1 to +0.3 typical)
        - Safety_Stock_Days: Buffer for lead time variance (typically 5-14 days)

        Example (Aluminum castings):
        - Avg 1000 units/week
        - Q1 seasonality: 0.85 (slower Jan-Mar)
        - Trend: +0.10 (10% YoY growth)
        - Safety stock: 7 days

        Forecast = 1000 × 0.85 × 1.10 + (1000/7 × 7) = 935 + 1000 = 1935 units for month

        ROI: Better forecast → 15-25% inventory reduction
        """
        forecast = historical_avg * seasonality_factor * (1 + trend)
        safety = (historical_avg / 7) * safety_stock_days
        return round(forecast + safety, 0)
    @staticmethod
    def economic_order_quantity(annual_demand: float, order_cost: float = 5000,
                               unit_cost: float = 100, holding_cost_percent: float = 0.20) -> Dict:
        """
        ECONOMIC ORDER QUANTITY (EOQ) FORMULA

        EOQ = √(2 × D × S / H)

        Where:
        - D = Annual demand (units)
        - S = Order cost per PO (typically ₹2000-10000)
        - H = Holding cost per unit per year (Unit_Cost × Holding_Cost_Percent)

        Then:
        - Reorder Point = (Lead_Time_Days × Daily_Demand) + Safety_Stock
        - Number of Orders/Year = Annual_Demand / EOQ

        Example (Steel bar stock):
        - Annual demand: 100,000 units
        - Order cost: ₹5000/PO
        - Unit cost: ₹100
        - Holding cost: 100 × 0.20 = ₹20/unit/year

        EOQ = √(2 × 100000 × 5000 / 20) = √50,000,000 = 7,071 units

        Reorder every ~70 days (100,000 / 7,071 ≈ 14 orders/year)

        ROI: Optimal order size reduces total inventory cost by 10-15%

        NOTE: holding cost is driven by the actual unit_cost argument passed in
        (previously this was hardcoded to assume ₹100/unit regardless of the
        unit cost entered in the UI, which silently made the EOQ ignore that input).
        """
        holding_cost_per_unit = unit_cost * holding_cost_percent
        eoq = (2 * annual_demand * order_cost / holding_cost_per_unit) ** 0.5 if holding_cost_per_unit > 0 else 0

        orders_per_year = annual_demand / eoq if eoq > 0 else 0
        days_between_orders = 365 / orders_per_year if orders_per_year > 0 else 0

        return {
            "EOQ_units": round(eoq, 0),
            "Orders_per_year": round(orders_per_year, 1),
            "Days_between_orders": round(days_between_orders, 1),
            "Holding_Cost_Per_Unit": round(holding_cost_per_unit, 2),
        }
    @staticmethod
    def vendor_risk_score(otif_percent: float, quality_defect_ppm: int,
                         payment_days_late: int, price_variance_percent: float) -> Dict:
        """
        VENDOR RISK SCORE FORMULA (0-100, higher = riskier)

        Risk = (OTIF_Risk × 0.35) + (Quality_Risk × 0.35) + (Payment_Risk × 0.20) + (Price_Risk × 0.10)

        Where:
        - OTIF Risk: On-Time-In-Full delivery performance
          100% OTIF → 0 risk points
          95% OTIF → 5 risk points
          <90% OTIF → 20+ risk points

        - Quality Risk: Defect rate in PPM
          <100 PPM → 0 risk (Excellent)
          100-500 PPM → 5-10 risk
          >1000 PPM → 30+ risk (Unacceptable)

        - Payment Risk: Average days late
          On-time → 0 risk
          1-7 days late → 2-5 risk
          >30 days late → 30+ risk

        - Price Risk: Variance from baseline
          ±2% → 0 risk (stable)
          ±5-10% → 5-10 risk
          >±15% → 30+ risk

        Example (Supplier A):
        - OTIF: 94% → Risk 6 × 0.35 = 2.1
        - PPM: 250 → Risk 8 × 0.35 = 2.8
        - Late: 5 days → Risk 4 × 0.20 = 0.8
        - Price var: +3% → Risk 5 × 0.10 = 0.5
        - Total Risk = 2.1 + 2.8 + 0.8 + 0.5 = 6.2/100 → GREEN (Reliable)

        Action:
        - 0-15: Green (reliable, preferred supplier)
        - 15-35: Yellow (monitor, contingency planning)
        - 35-60: Orange (high risk, find alternatives)
        - 60-100: Red (critical risk, immediate action)

        ROI: Reduces supply chain disruptions, prevents line stops
        """
        # OTIF risk (0-100 scale where 100 OTIF = 0 risk)
        otif_risk = max(0, (100 - otif_percent) / 100 * 20)

        # Quality risk (PPM to 0-30 scale)
        if quality_defect_ppm < 100:
            quality_risk = 0
        elif quality_defect_ppm < 500:
            quality_risk = (quality_defect_ppm - 100) / 400 * 10
        elif quality_defect_ppm < 1000:
            quality_risk = 10 + (quality_defect_ppm - 500) / 500 * 10
        else:
            quality_risk = 20 + min(10, quality_defect_ppm / 1000 * 10)

        # Payment risk (days late to 0-30 scale)
        if payment_days_late <= 0:
            payment_risk = 0
        elif payment_days_late <= 7:
            payment_risk = payment_days_late / 7 * 5
        elif payment_days_late <= 30:
            payment_risk = 5 + (payment_days_late - 7) / 23 * 10
        else:
            payment_risk = 15 + min(15, payment_days_late / 30 * 15)

        # Price variance risk (percent to 0-30 scale)
        abs_variance = abs(price_variance_percent)
        if abs_variance <= 2:
            price_risk = 0
        elif abs_variance <= 10:
            price_risk = (abs_variance - 2) / 8 * 10
        else:
            price_risk = 10 + min(20, (abs_variance - 10) / 10 * 20)

        total_risk = (otif_risk * 0.35 + quality_risk * 0.35 +
                     payment_risk * 0.20 + price_risk * 0.10)

        if total_risk <= 15:
            status = "🟢 GREEN - Reliable"
        elif total_risk <= 35:
            status = "🟡 YELLOW - Monitor"
        elif total_risk <= 60:
            status = "🟠 ORANGE - High Risk"
        else:
            status = "🔴 RED - Critical"

        return {
            "Total_Risk_Score": round(total_risk, 1),
            "Status": status,
            "OTIF_Risk": round(otif_risk, 1),
            "Quality_Risk": round(quality_risk, 1),
            "Payment_Risk": round(payment_risk, 1),
            "Price_Risk": round(price_risk, 1),
        }
    @staticmethod
    def inventory_turnover(cogs: float, avg_inventory_value: float) -> float:
        """
        INVENTORY TURNOVER RATIO FORMULA

        Turnover = COGS / Avg_Inventory_Value

        Where:
        - COGS: Cost of Goods Sold (annual)
        - Avg_Inventory: Average inventory value on hand

        Example:
        - Annual COGS: ₹100 crore
        - Avg inventory: ₹10 crore
        - Turnover = 100/10 = 10x/year

        Interpretation:
        - Turnover = 4-6x: Typical manufacturing (good balance)
        - Turnover = 8-10x: Efficient (low carrying cost)
        - Turnover = >12x: Lean (risky stockout potential)
        - Turnover = <3x: Heavy (excess capital tied up)

        Days of Inventory = 365 / Turnover

        ROI: Improve turnover 1x → Frees up ₹10 crore working capital
        """
        if avg_inventory_value <= 0:
            return 0
        return round(cogs / avg_inventory_value, 2)
    @staticmethod
    def working_capital_impact(inventory_value: float, payables_days: int,
                              receivables_days: int) -> Dict:
        """
        WORKING CAPITAL CASH CONVERSION CYCLE FORMULA

        Cash_Cycle_Days = Inventory_Days + Receivables_Days - Payables_Days

        Working_Capital_Need = (COGS / 365) × Cash_Cycle_Days

        Where:
        - Inventory_Days: Days of inventory on hand (365 / Turnover)
        - Receivables_Days: Average collection period
        - Payables_Days: Average payment terms from suppliers

        Example:
        - Inventory days: 60 (sell every 60 days)
        - Receivables: 45 (collect from OEMs in 45 days)
        - Payables: 30 (pay suppliers in 30 days)
        - Cycle = 60 + 45 - 30 = 75 days

        If COGS = ₹100 crore/year:
        - Daily cost = ₹27.4 lakh
        - Working capital need = ₹27.4 lakh × 75 = ₹20.5 crore

        Each 10-day cycle improvement = ₹2.74 crore freed up

        ROI: Negotiate longer payables + improve inventory = cash released for growth
        """
        daily_cogs = 365 / 365  # Normalized to 1 per day for example

        cash_cycle = inventory_value + receivables_days - payables_days

        return {
            "Cash_Cycle_Days": cash_cycle,
            "Status": "Good" if cash_cycle < 45 else ("Average" if cash_cycle < 75 else "Poor"),
            "Improvement_Potential": "Reduce inventory by 10 days" if inventory_value > 60 else "Extend payables by 10 days"
        }
# =========================================================================
# EXPORT FUNCTIONS
# =========================================================================
# FPDF's built-in core fonts only support Latin-1. Emoji status markers
# (🟢/🟡/🟠/🔴), the ₹ symbol, and smart quotes/dashes all crash PDF export
# with a Unicode error unless sanitized first — _pdf_safe() does that, and
# export_to_pdf() also normalizes whatever type pdf.output() returns (str,
# bytearray, or bytes, depending on the installed fpdf/fpdf2 version) so
# export never throws regardless of which package version is deployed.
_PDF_CHAR_MAP = {
    "🟢": "[GOOD] ", "🟡": "[MONITOR] ", "🟠": "[WARN] ", "🔴": "[CRITICAL] ",
    "✅": "[OK] ", "⚠️": "[WARNING] ", "⚠": "[WARNING] ", "❌": "[X] ",
    "⚙️": "", "📈": "", "📦": "", "🤝": "", "💰": "", "🎯": "",
    "🔄": "", "📊": "", "📋": "", "📄": "", "📝": "", "📁": "", "🗑️": "",
    "₹": "Rs. ", "–": "-", "—": "-",
    "’": "'", "‘": "'", "“": '"', "”": '"', "…": "...",
}
def _pdf_safe(value) -> str:
    """Make text safe for FPDF's Latin-1-only core fonts."""
    text = str(value)
    for original, replacement in _PDF_CHAR_MAP.items():
        text = text.replace(original, replacement)
    # Drop anything still outside Latin-1 instead of letting FPDF crash on it
    return text.encode("latin-1", "ignore").decode("latin-1").strip()
def export_to_pdf(df: pd.DataFrame, title: str, formulas: Dict = None) -> bytes:
    """Export dataframe and formulas to PDF"""
    if not HAS_PDF:
        return None

    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", "B", 16)
        pdf.cell(0, 10, _pdf_safe(title), ln=True, align="C")

        pdf.set_font("Arial", "B", 10)
        pdf.cell(0, 5, f"Generated: {dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", ln=True)
        pdf.cell(0, 5, "ChainIntel Pro | Kalsnet (KNet) Consulting", ln=True)

        # Table
        pdf.set_font("Arial", "", 8)
        pdf.ln(5)

        col_width = 190 / len(df.columns)
        row_height = 7
        page_bottom = pdf.h - pdf.b_margin

        def draw_header():
            pdf.set_font("Arial", "B", 8)
            for col in df.columns:
                pdf.cell(col_width, row_height, _pdf_safe(col)[:20], border=1, align="C")
            pdf.ln()
            pdf.set_font("Arial", "", 8)

        draw_header()

        for _, row in df.head(50).iterrows():
            # Start a new page (with repeated header) instead of letting
            # rows silently run off the bottom of the page
            if pdf.get_y() + row_height > page_bottom:
                pdf.add_page()
                draw_header()
            for val in row:
                pdf.cell(col_width, row_height, _pdf_safe(val)[:20], border=1, align="L")
            pdf.ln()

        if formulas:
            pdf.add_page()
            pdf.set_font("Arial", "B", 12)
            pdf.cell(0, 10, "Formulas & Explanations", ln=True)
            pdf.set_font("Arial", "", 9)
            for formula_name, formula_text in formulas.items():
                pdf.cell(0, 5, _pdf_safe(formula_name) + ":", ln=True)
                pdf.multi_cell(0, 4, _pdf_safe(formula_text)[:500])
                pdf.ln(2)

        # Different fpdf/fpdf2 versions return str, bytearray, or bytes from
        # output(dest='S'); normalize all three so this never throws.
        raw_output = pdf.output(dest="S")
        if isinstance(raw_output, str):
            return raw_output.encode("latin-1")
        return bytes(raw_output)
    except Exception as e:
        st.error(f"PDF export error: {e}")
        logger.exception("PDF export failed")
        return None
def export_to_docx(df: pd.DataFrame, title: str, formulas: Dict = None) -> bytes:
    """Export dataframe and formulas to Word document"""
    if not HAS_DOCX:
        return None

    try:
        doc = Document()

        # Title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 71, 171)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        meta_para = doc.add_paragraph()
        meta_run = meta_para.add_run(f"Generated: {dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        meta_run.font.size = Pt(10)
        meta_run = meta_para.add_run("ChainIntel Pro | Developed by Randy Singh | Kalsnet (KNet) Consulting")
        meta_run.font.size = Pt(10)
        meta_run.italic = True

        doc.add_paragraph()

        # Table
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr_cells[i].text = str(col)

        for _, row in df.head(50).iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val)[:100]

        # Formulas section
        if formulas:
            doc.add_page_break()
            formulas_heading = doc.add_paragraph()
            formulas_run = formulas_heading.add_run("Formulas & Explanations")
            formulas_run.font.size = Pt(14)
            formulas_run.font.bold = True
            formulas_run.font.color.rgb = RGBColor(0, 71, 171)

            for formula_name, formula_text in formulas.items():
                p = doc.add_paragraph(formula_name, style='Heading 3')
                doc.add_paragraph(formula_text[:500])

        # Save to bytes
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        return docx_bytes.getvalue()
    except Exception as e:
        st.error(f"Word export error: {e}")
        return None
def export_to_text(df: pd.DataFrame, title: str, formulas: Dict = None) -> str:
    """Export to plain text format"""
    text = f"\n{'='*80}\n"
    text += f"{title}\n"
    text += f"{'='*80}\n\n"
    text += f"Generated: {dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
    text += "ChainIntel Pro | Developed by Randy Singh | Kalsnet (KNet) Consulting\n\n"
    text += df.to_string()

    if formulas:
        text += f"\n\n{'='*80}\nFormulas & Explanations\n{'='*80}\n\n"
        for formula_name, formula_text in formulas.items():
            text += f"\n{formula_name}:\n{formula_text}\n"

    return text
def load_uploaded_table(uploaded_file) -> pd.DataFrame:
    """Read an uploaded CSV/Excel file into a DataFrame with a friendly error on failure."""
    try:
        if uploaded_file.name.lower().endswith(".csv"):
            return pd.read_csv(uploaded_file)
        return pd.read_excel(uploaded_file)
    except Exception as e:
        st.error(f"Could not read file: {e}")
        logger.exception("Failed to read uploaded file")
        return None
# =========================================================================
# SIDEBAR CONFIG
# =========================================================================
with st.sidebar:
    st.header("⚙️ Configuration")

    with st.expander("📊 Data Settings", expanded=True):
        st.markdown(f"**Max Upload**: {MAX_UPLOAD_MB} MB")
        st.markdown(f"**Environment**: {DEPLOY_ENV.upper()}")
        st.markdown(f"**Session**: {st.session_state.session_id}")

    st.markdown("---")
    st.caption("v1.0 | Kalsnet (KNet) Consulting")
# =========================================================================
# MAIN TABS
# =========================================================================
tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "📈 Demand Forecast",
    "📦 Inventory Optimization",
    "🤝 Vendor Risk",
    "💰 Working Capital",
    "🎯 Analytics"
])
# =========================================================================
# TAB 1: DEMAND FORECASTING
# =========================================================================
with tab1:
    st.markdown('<div class="section-header">📈 AI Demand Forecasting</div>', unsafe_allow_html=True)

    st.markdown("""
    Predict demand using historical trends, OEM schedules, and seasonality.
    Reduce stockouts and dead stock through better planning.
    """)

    col1, col2, col3, col4 = st.columns([2, 2, 1, 1])

    with col1:
        n_forecast = st.slider("Generate forecast records", 0, 500, 100, step=10)

    with col2:
        selected_oem = st.selectbox("Select OEM", OEM_NAMES)

    with col3:
        st.write("")
        st.write("")
        gen_forecast = st.button("🔄 Generate", key="gen_forecast")

    with col4:
        st.write("")
        st.write("")
        reset_forecast = st.button("🗑️ Reset", key="reset_forecast")

    if reset_forecast:
        st.session_state.forecast_df = None
        st.success("✅ Synthetic forecast data reset")

    if gen_forecast:
        try:
            rows = []
            base_demand = random.randint(500, 2000)

            for i in range(n_forecast):
                month = (i % 12) + 1

                # Seasonality (Q1 slower, Q3 busier)
                if month in [1, 2, 3]:
                    seasonality = 0.85
                elif month in [7, 8, 9]:
                    seasonality = 1.15
                else:
                    seasonality = 1.0

                trend = random.uniform(-0.05, 0.15)
                actual_demand = int(base_demand * seasonality * (1 + trend * (i / n_forecast)))

                forecast = SupplyChainFormulas.forecast_demand(base_demand, seasonality, trend)

                accuracy = max(70, 95 - abs(actual_demand - forecast) / forecast * 100)

                rows.append({
                    "Month": f"2024-{month:02d}",
                    "OEM": selected_oem,
                    "Historical_Avg": base_demand,
                    "Seasonality": round(seasonality, 2),
                    "Trend_%": round(trend * 100, 1),
                    "Forecast_Units": int(forecast),
                    "Actual_Units": actual_demand,
                    "Accuracy_%": round(accuracy, 1),
                })

            st.session_state.forecast_df = pd.DataFrame(rows)
            st.success(f"✅ Generated {n_forecast} demand forecast records")
        except Exception as e:
            st.error(f"Could not generate forecast data: {e}")
            logger.exception("Forecast generation failed")

    # Upload real data
    st.markdown("**OR upload real demand/actuals data:**")
    col_u1, col_u2 = st.columns([4, 1])
    with col_u1:
        uploaded_forecast = st.file_uploader("Upload CSV/Excel file", type=["csv", "xlsx"], key="upload_forecast")
    with col_u2:
        if st.session_state.real_forecast_df is not None:
            st.write("")
            if st.button("❌ Clear upload", key="clear_real_forecast"):
                st.session_state.real_forecast_df = None
                st.success("Cleared uploaded forecast data")

    if uploaded_forecast:
        real_df = load_uploaded_table(uploaded_forecast)
        if real_df is not None:
            st.session_state.real_forecast_df = real_df
            st.success(f"✅ Loaded {len(real_df)} real forecast records")

    # Display & analyze (real data overrides synthetic, same pattern as ShopFloor Sentinel)
    df = st.session_state.real_forecast_df if st.session_state.real_forecast_df is not None else st.session_state.forecast_df

    if df is not None:
        st.markdown("**Forecast Data:**")
        st.dataframe(df, use_container_width=True, height=250)

        col1, col2, col3, col4 = st.columns(4)
        with col1:
            if "Forecast_Units" in df.columns:
                st.metric("Avg Forecast", f"{df['Forecast_Units'].mean():.0f} units")
            else:
                st.metric("Avg Forecast", "N/A")
        with col2:
            if "Actual_Units" in df.columns:
                st.metric("Avg Actual", f"{df['Actual_Units'].mean():.0f} units")
            else:
                st.metric("Avg Actual", "N/A")
        with col3:
            if "Accuracy_%" in df.columns:
                st.metric("Accuracy", f"{df['Accuracy_%'].mean():.1f}%")
            else:
                st.metric("Accuracy", "N/A")
        with col4:
            if "Actual_Units" in df.columns and "Forecast_Units" in df.columns:
                stockouts = (df['Actual_Units'] > df['Forecast_Units']).sum()
                st.metric("Potential Stockouts", stockouts)
            else:
                st.metric("Potential Stockouts", "N/A")

        st.markdown('<div class="formula-box">', unsafe_allow_html=True)
        st.markdown("""
        **DEMAND FORECASTING FORMULA**

        **Forecast = Historical_Avg × Seasonality × (1 + Trend) + Safety_Stock**

        **Components:**
        - Historical_Avg: 12-month average demand
        - Seasonality: Month/quarter multiplier (0.85-1.15 typical)
          * Q1 (Jan-Mar): 0.85 (slower)
          * Q2 (Apr-Jun): 1.0 (average)
          * Q3 (Jul-Sep): 1.15 (busier)
          * Q4 (Oct-Dec): 1.0 (average)

        - Trend: YoY growth rate (-5% to +15% typical)
        - Safety_Stock: Buffer for lead time variance (7-14 days typical)

        **Example (Maruti supply):**
        - Avg demand: 1000 units/week
        - Jan forecast: 1000 × 0.85 × 1.05 = 893 units
        - Sep forecast: 1000 × 1.15 × 1.05 = 1206 units

        **ROI:** Better forecast → 15-25% inventory reduction
        """)
        st.markdown('</div>', unsafe_allow_html=True)

        # Export
        st.markdown('<div class="subsection-header">💾 Export Results</div>', unsafe_allow_html=True)

        formulas_dict_fc = {
            "Demand Forecast Formula": "Forecast = Historical_Avg x Seasonality x (1 + Trend) + Safety_Stock",
        }

        col1, col2, col3, col4, col5 = st.columns(5)
        with col1:
            st.download_button("📊 CSV", df.to_csv(index=False).encode(), "demand_forecast.csv", "text/csv", key="csv_fc")
        with col2:
            st.download_button("📋 JSON", df.to_json(orient="records", indent=2).encode(), "demand_forecast.json", "application/json", key="json_fc")
        with col3:
            txt = export_to_text(df, "Demand Forecast Analysis", formulas_dict_fc)
            st.download_button("📄 TXT", txt.encode(), "demand_forecast.txt", "text/plain", key="txt_fc")
        with col4:
            if HAS_DOCX:
                docx_bytes = export_to_docx(df, "Demand Forecast Analysis", formulas_dict_fc)
                if docx_bytes:
                    st.download_button("📝 WORD", docx_bytes, "demand_forecast.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="docx_fc")
            else:
                st.caption("Word export unavailable")
        with col5:
            if HAS_PDF:
                pdf_bytes = export_to_pdf(df, "Demand Forecast Analysis", formulas_dict_fc)
                if pdf_bytes:
                    st.download_button("🔴 PDF", pdf_bytes, "demand_forecast.pdf", "application/pdf", key="pdf_fc")
            else:
                st.caption("PDF export unavailable")
# =========================================================================
# TAB 2: INVENTORY OPTIMIZATION
# =========================================================================
with tab2:
    st.markdown('<div class="section-header">📦 Economic Order Quantity & Reorder Points</div>', unsafe_allow_html=True)

    st.markdown("""
    Calculate optimal order quantities and reorder points. Reduce holding costs while preventing stockouts.
    """)

    st.markdown('<div class="subsection-header">Single Scenario Calculator</div>', unsafe_allow_html=True)

    col1, col2, col3, col4 = st.columns(4)

    with col1:
        annual_demand = st.number_input("Annual Demand (units)", 1000, 1000000, 50000, step=1000)

    with col2:
        order_cost = st.number_input("Order Cost per PO (₹)", 1000, 50000, 5000, step=500)

    with col3:
        unit_cost = st.number_input("Unit Cost (₹)", 10, 10000, 100, step=10)

    with col4:
        lead_time_days = st.number_input("Lead Time (days)", 1, 60, 14, step=1)

    # Calculate EOQ (unit_cost is now actually used — previously hardcoded to ₹100 regardless of input)
    holding_cost_percent = 0.20  # Standard 20% of unit cost
    eoq_data = SupplyChainFormulas.economic_order_quantity(annual_demand, order_cost, unit_cost, holding_cost_percent)

    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("EOQ (units)", int(eoq_data["EOQ_units"]))
    with col2:
        st.metric("Orders/Year", eoq_data["Orders_per_year"])
    with col3:
        st.metric("Days Between Orders", int(eoq_data["Days_between_orders"]))
    with col4:
        daily_demand = annual_demand / 365
        reorder_point = int(daily_demand * lead_time_days)
        st.metric("Reorder Point", reorder_point)

    st.markdown('<div class="formula-box">', unsafe_allow_html=True)
    st.markdown(f"""
    **ECONOMIC ORDER QUANTITY (EOQ) FORMULA**

    **EOQ = √(2 × D × S / H)**

    Where:
    - D (Annual Demand) = {annual_demand:,} units
    - S (Order Cost) = ₹{order_cost:,}/PO
    - H (Holding Cost) = Unit_Cost × 20% = ₹{unit_cost} × 0.20 = ₹{unit_cost * 0.20:.2f}

    **Calculation:**
    EOQ = √(2 × {annual_demand:,} × {order_cost:,} / {unit_cost * 0.20:.2f})
    EOQ = √({2 * annual_demand * order_cost:,.0f} / {unit_cost * 0.20:.2f})
    EOQ = **{eoq_data["EOQ_units"]:.0f} units**

    **Reorder Logic:**
    - Reorder Point = Daily_Demand × Lead_Time_Days
    - Daily Demand = {annual_demand:,} / 365 = {daily_demand:.0f} units/day
    - Reorder Point = {daily_demand:.0f} × {lead_time_days} = **{int(daily_demand * lead_time_days)} units**

    **Action:**
    When inventory reaches {int(daily_demand * lead_time_days)} units → Place PO for {int(eoq_data["EOQ_units"])} units

    **ROI:** Optimal sizing reduces total inventory cost by 10-15%
    """)
    st.markdown('</div>', unsafe_allow_html=True)

    # Export this single scenario
    scenario_df = pd.DataFrame([{
        "Annual_Demand": annual_demand,
        "Order_Cost": order_cost,
        "Unit_Cost": unit_cost,
        "Lead_Time_Days": lead_time_days,
        "EOQ_Units": int(eoq_data["EOQ_units"]),
        "Orders_per_Year": eoq_data["Orders_per_year"],
        "Days_Between_Orders": eoq_data["Days_between_orders"],
        "Reorder_Point": reorder_point,
    }])

    formulas_dict_eoq = {
        "EOQ Formula": "EOQ = sqrt(2 x Annual_Demand x Order_Cost / Holding_Cost_Per_Unit), Holding_Cost_Per_Unit = Unit_Cost x 20%",
        "Reorder Point": "Reorder_Point = Daily_Demand x Lead_Time_Days",
    }

    st.markdown('<div class="subsection-header">💾 Export This Scenario</div>', unsafe_allow_html=True)
    col1, col2, col3, col4, col5 = st.columns(5)
    with col1:
        st.download_button("📊 CSV", scenario_df.to_csv(index=False).encode(), "eoq_scenario.csv", "text/csv", key="csv_eoq")
    with col2:
        st.download_button("📋 JSON", scenario_df.to_json(orient="records", indent=2).encode(), "eoq_scenario.json", "application/json", key="json_eoq")
    with col3:
        txt = export_to_text(scenario_df, "Inventory EOQ Scenario", formulas_dict_eoq)
        st.download_button("📄 TXT", txt.encode(), "eoq_scenario.txt", "text/plain", key="txt_eoq")
    with col4:
        if HAS_DOCX:
            docx_bytes = export_to_docx(scenario_df, "Inventory EOQ Scenario", formulas_dict_eoq)
            if docx_bytes:
                st.download_button("📝 WORD", docx_bytes, "eoq_scenario.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="docx_eoq")
    with col5:
        if HAS_PDF:
            pdf_bytes = export_to_pdf(scenario_df, "Inventory EOQ Scenario", formulas_dict_eoq)
            if pdf_bytes:
                st.download_button("🔴 PDF", pdf_bytes, "eoq_scenario.pdf", "application/pdf", key="pdf_eoq")

    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
    st.markdown('<div class="subsection-header">📁 Bulk SKU Upload (Multiple Items at Once)</div>', unsafe_allow_html=True)
    st.markdown(f"""
    Upload a file with one row per SKU to calculate EOQ and reorder point for your whole catalog in one go.
    Expected columns (case-insensitive): `SKU_ID`, `Annual_Demand`. Optional: `Order_Cost`, `Unit_Cost`, `Lead_Time_Days`
    — any missing optional column falls back to the values set above (₹{order_cost:,}, ₹{unit_cost:,}, {lead_time_days} days).
    """)

    col1, col2 = st.columns([3, 1])
    with col1:
        uploaded_sku = st.file_uploader("Upload SKU file (CSV/Excel)", type=["csv", "xlsx"], key="upload_sku")
    with col2:
        st.write("")
        st.write("")
        clear_sku = st.button("🗑️ Clear", key="clear_sku")

    if clear_sku:
        st.session_state.inventory_sku_df = None
        st.success("✅ Bulk SKU data cleared")

    if uploaded_sku:
        raw_sku_df = load_uploaded_table(uploaded_sku)
        if raw_sku_df is not None:
            try:
                cols_lower = {c.lower().strip(): c for c in raw_sku_df.columns}

                def get_col(*names):
                    for n in names:
                        if n in cols_lower:
                            return raw_sku_df[cols_lower[n]]
                    return None

                sku_ids = get_col("sku_id", "sku", "item", "material")
                demand_col = get_col("annual_demand", "demand")

                if demand_col is None:
                    st.error("Could not find an 'Annual_Demand' column in the uploaded file.")
                else:
                    order_cost_col = get_col("order_cost")
                    unit_cost_col = get_col("unit_cost")
                    lead_time_col = get_col("lead_time_days", "lead_time")

                    result_rows = []
                    for i in range(len(raw_sku_df)):
                        try:
                            row_demand = float(demand_col.iloc[i])
                            row_order_cost = float(order_cost_col.iloc[i]) if order_cost_col is not None else float(order_cost)
                            row_unit_cost = float(unit_cost_col.iloc[i]) if unit_cost_col is not None else float(unit_cost)
                            row_lead_time = float(lead_time_col.iloc[i]) if lead_time_col is not None else float(lead_time_days)

                            row_eoq = SupplyChainFormulas.economic_order_quantity(row_demand, row_order_cost, row_unit_cost, holding_cost_percent)
                            row_daily_demand = row_demand / 365
                            row_reorder_point = int(row_daily_demand * row_lead_time)

                            result_rows.append({
                                "SKU_ID": sku_ids.iloc[i] if sku_ids is not None else f"SKU-{i+1}",
                                "Annual_Demand": row_demand,
                                "Order_Cost": row_order_cost,
                                "Unit_Cost": row_unit_cost,
                                "Lead_Time_Days": row_lead_time,
                                "EOQ_Units": int(row_eoq["EOQ_units"]),
                                "Orders_per_Year": row_eoq["Orders_per_year"],
                                "Days_Between_Orders": row_eoq["Days_between_orders"],
                                "Reorder_Point": row_reorder_point,
                            })
                        except Exception as row_err:
                            logger.warning(f"Skipping SKU row {i}: {row_err}")
                            continue

                    if result_rows:
                        st.session_state.inventory_sku_df = pd.DataFrame(result_rows)
                        st.success(f"✅ Calculated EOQ for {len(result_rows)} SKUs")
                    else:
                        st.error("No valid rows could be processed from the uploaded file.")
            except Exception as e:
                st.error(f"Could not process SKU file: {e}")
                logger.exception("Bulk SKU EOQ failed")

    if st.session_state.inventory_sku_df is not None:
        sku_df = st.session_state.inventory_sku_df
        st.dataframe(sku_df, use_container_width=True, height=300)

        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Total SKUs", len(sku_df))
        with col2:
            st.metric("Avg EOQ", f"{sku_df['EOQ_Units'].mean():.0f} units")
        with col3:
            st.metric("Avg Reorder Point", f"{sku_df['Reorder_Point'].mean():.0f} units")

        st.markdown('<div class="subsection-header">💾 Export Bulk SKU Results</div>', unsafe_allow_html=True)
        col1, col2, col3, col4, col5 = st.columns(5)
        with col1:
            st.download_button("📊 CSV", sku_df.to_csv(index=False).encode(), "bulk_eoq.csv", "text/csv", key="csv_sku")
        with col2:
            st.download_button("📋 JSON", sku_df.to_json(orient="records", indent=2).encode(), "bulk_eoq.json", "application/json", key="json_sku")
        with col3:
            txt = export_to_text(sku_df, "Bulk SKU EOQ Analysis", formulas_dict_eoq)
            st.download_button("📄 TXT", txt.encode(), "bulk_eoq.txt", "text/plain", key="txt_sku")
        with col4:
            if HAS_DOCX:
                docx_bytes = export_to_docx(sku_df, "Bulk SKU EOQ Analysis", formulas_dict_eoq)
                if docx_bytes:
                    st.download_button("📝 WORD", docx_bytes, "bulk_eoq.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="docx_sku")
        with col5:
            if HAS_PDF:
                pdf_bytes = export_to_pdf(sku_df, "Bulk SKU EOQ Analysis", formulas_dict_eoq)
                if pdf_bytes:
                    st.download_button("🔴 PDF", pdf_bytes, "bulk_eoq.pdf", "application/pdf", key="pdf_sku")
# =========================================================================
# TAB 3: VENDOR RISK SCORING
# =========================================================================
with tab3:
    st.markdown('<div class="section-header">🤝 Vendor Risk Assessment & Scoring</div>', unsafe_allow_html=True)

    st.markdown("""
    Score and rank vendors based on OTIF, quality, payment reliability, and price stability.
    Identify supply chain risks early.
    """)

    col1, col2, col3, col4 = st.columns([2, 2, 1, 1])

    with col1:
        n_vendors = st.slider("Generate vendor records", 0, 500, 50, step=10)

    with col2:
        region = st.selectbox("Region", ["North", "South", "East", "West", "Multi-region"])

    with col3:
        st.write("")
        st.write("")
        gen_vendors = st.button("🔄 Generate", key="gen_vendors")

    with col4:
        st.write("")
        st.write("")
        reset_vendors = st.button("🗑️ Reset", key="reset_vendors")

    if reset_vendors:
        st.session_state.vendor_df = None
        st.success("✅ Synthetic vendor data reset")

    if gen_vendors:
        try:
            rows = []
            for i in range(n_vendors):
                otif = random.randint(85, 100)
                ppm = random.randint(50, 2000)
                days_late = random.randint(-5, 30)
                price_var = random.uniform(-8, 12)

                risk_data = SupplyChainFormulas.vendor_risk_score(otif, ppm, days_late, price_var)

                rows.append({
                    "Vendor_ID": f"VEN-{random.randint(1000, 9999)}",
                    "Vendor_Name": f"Supplier-{i+1}",
                    "OTIF_%": otif,
                    "Quality_PPM": ppm,
                    "Days_Late_Avg": days_late,
                    "Price_Variance_%": round(price_var, 1),
                    # FIX: vendor_risk_score() returns "Total_Risk_Score", not
                    # "Failure_Risk_Score" — the old key name here is what
                    # caused the KeyError that crashed the whole app.
                    "Risk_Score": risk_data["Total_Risk_Score"],
                    "Status": risk_data["Status"],
                })

            st.session_state.vendor_df = pd.DataFrame(rows)
            st.success(f"✅ Generated {n_vendors} vendor records")
        except Exception as e:
            st.error(f"Could not generate vendor data: {e}")
            logger.exception("Vendor generation failed")

    # Upload real data
    st.markdown("**OR upload real vendor performance data:**")
    st.caption("Expected columns (case-insensitive): Vendor_ID, Vendor_Name, OTIF_%, Quality_PPM, Days_Late_Avg, Price_Variance_%")
    col_u1, col_u2 = st.columns([4, 1])
    with col_u1:
        uploaded_vendor = st.file_uploader("Upload CSV/Excel file", type=["csv", "xlsx"], key="upload_vendor")
    with col_u2:
        if st.session_state.real_vendor_df is not None:
            st.write("")
            if st.button("❌ Clear upload", key="clear_real_vendor"):
                st.session_state.real_vendor_df = None
                st.success("Cleared uploaded vendor data")

    if uploaded_vendor:
        raw_vendor_df = load_uploaded_table(uploaded_vendor)
        if raw_vendor_df is not None:
            try:
                cols_lower = {c.lower().strip(): c for c in raw_vendor_df.columns}

                def get_val(row, *names, default=0):
                    for n in names:
                        if n in cols_lower:
                            return row[cols_lower[n]]
                    return default

                result_rows = []
                for i, row in raw_vendor_df.iterrows():
                    try:
                        otif = float(get_val(row, "otif_%", "otif", default=95))
                        ppm = float(get_val(row, "quality_ppm", "ppm", default=200))
                        days_late = float(get_val(row, "days_late_avg", "days_late", default=0))
                        price_var = float(get_val(row, "price_variance_%", "price_variance", default=0))

                        risk_data = SupplyChainFormulas.vendor_risk_score(otif, ppm, days_late, price_var)

                        vendor_id = get_val(row, "vendor_id", default=f"VEN-{1000+i}")
                        vendor_name = get_val(row, "vendor_name", "vendor", default=f"Supplier-{i+1}")

                        result_rows.append({
                            "Vendor_ID": vendor_id,
                            "Vendor_Name": vendor_name,
                            "OTIF_%": otif,
                            "Quality_PPM": ppm,
                            "Days_Late_Avg": days_late,
                            "Price_Variance_%": round(price_var, 1),
                            "Risk_Score": risk_data["Total_Risk_Score"],
                            "Status": risk_data["Status"],
                        })
                    except Exception as row_err:
                        logger.warning(f"Skipping vendor row {i}: {row_err}")
                        continue

                if result_rows:
                    st.session_state.real_vendor_df = pd.DataFrame(result_rows)
                    st.success(f"✅ Scored {len(result_rows)} real vendor records")
                else:
                    st.error("No valid rows could be scored from the uploaded file.")
            except Exception as e:
                st.error(f"Could not process vendor file: {e}")
                logger.exception("Real vendor scoring failed")

    # Display (real data overrides synthetic)
    df = st.session_state.real_vendor_df if st.session_state.real_vendor_df is not None else st.session_state.vendor_df

    if df is not None:
        st.markdown("**Vendor Risk Data:**")
        st.dataframe(df.sort_values("Risk_Score"), use_container_width=True, height=250)

        col1, col2, col3, col4 = st.columns(4)
        green = (df["Status"].str.contains("GREEN")).sum()
        yellow = (df["Status"].str.contains("YELLOW")).sum()
        orange = (df["Status"].str.contains("ORANGE")).sum()
        red = (df["Status"].str.contains("RED")).sum()

        with col1:
            st.metric("🟢 Green (Reliable)", green)
        with col2:
            st.metric("🟡 Yellow (Monitor)", yellow)
        with col3:
            st.metric("🟠 Orange (Risk)", orange)
        with col4:
            st.metric("🔴 Red (Critical)", red)

        st.markdown('<div class="formula-box">', unsafe_allow_html=True)
        st.markdown("""
        **VENDOR RISK SCORE FORMULA (0-100 scale)**

        **Risk = (OTIF_Risk × 0.35) + (Quality_Risk × 0.35) + (Payment_Risk × 0.20) + (Price_Risk × 0.10)**

        **OTIF Risk (35% weight):**
        - 100% on-time → 0 risk
        - 95% → 5 risk
        - <90% → 20+ risk

        **Quality Risk (35% weight):**
        - <100 PPM → 0 risk (Excellent)
        - 100-500 PPM → 5-10 risk
        - >1000 PPM → 30+ risk

        **Payment Risk (20% weight):**
        - On-time → 0 risk
        - 1-7 days late → 2-5 risk
        - >30 days late → 30+ risk

        **Price Risk (10% weight):**
        - ±2% → 0 risk (stable)
        - ±5-10% → 5-10 risk
        - >±15% → 30+ risk

        **Risk Status:**
        - 0-15: 🟢 GREEN - Reliable, preferred supplier
        - 15-35: 🟡 YELLOW - Monitor, contingency planning
        - 35-60: 🟠 ORANGE - High risk, find alternatives
        - 60-100: 🔴 RED - Critical risk, immediate action
        """)
        st.markdown('</div>', unsafe_allow_html=True)

        # Export
        st.markdown('<div class="subsection-header">💾 Export Results</div>', unsafe_allow_html=True)

        formulas_dict_vendor = {
            "Vendor Risk Formula": "Risk = (OTIF_Risk x 0.35) + (Quality_Risk x 0.35) + (Payment_Risk x 0.20) + (Price_Risk x 0.10)",
            "Risk Status": "0-15 GREEN | 15-35 YELLOW | 35-60 ORANGE | 60-100 RED",
        }

        col1, col2, col3, col4, col5 = st.columns(5)
        with col1:
            st.download_button("📊 CSV", df.to_csv(index=False).encode(), "vendor_risk.csv", "text/csv", key="csv_vendor")
        with col2:
            st.download_button("📋 JSON", df.to_json(orient="records", indent=2).encode(), "vendor_risk.json", "application/json", key="json_vendor")
        with col3:
            txt = export_to_text(df, "Vendor Risk Analysis", formulas_dict_vendor)
            st.download_button("📄 TXT", txt.encode(), "vendor_risk.txt", "text/plain", key="txt_vendor")
        with col4:
            if HAS_DOCX:
                docx_bytes = export_to_docx(df, "Vendor Risk Analysis", formulas_dict_vendor)
                if docx_bytes:
                    st.download_button("📝 WORD", docx_bytes, "vendor_risk.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="docx_vendor")
        with col5:
            if HAS_PDF:
                pdf_bytes = export_to_pdf(df, "Vendor Risk Analysis", formulas_dict_vendor)
                if pdf_bytes:
                    st.download_button("🔴 PDF", pdf_bytes, "vendor_risk.pdf", "application/pdf", key="pdf_vendor")
# =========================================================================
# TAB 4: WORKING CAPITAL
# =========================================================================
with tab4:
    st.markdown('<div class="section-header">💰 Working Capital & Cash Cycle Analysis</div>', unsafe_allow_html=True)

    col1, col2, col3 = st.columns(3)

    with col1:
        inventory_days = st.slider("Days of Inventory on Hand", 10, 180, 60, step=5)

    with col2:
        receivables_days = st.slider("Days to Collect from OEMs", 15, 90, 45, step=5)

    with col3:
        payables_days = st.slider("Days to Pay Suppliers", 15, 90, 30, step=5)

    # Calculate cash cycle
    cash_cycle = inventory_days + receivables_days - payables_days

    col1, col2, col3, col4 = st.columns(4)

    with col1:
        st.metric("Inventory Days", inventory_days)
    with col2:
        st.metric("Receivable Days", receivables_days)
    with col3:
        st.metric("Payable Days", payables_days)
    with col4:
        if cash_cycle < 45:
            status = "🟢 Good"
        elif cash_cycle < 75:
            status = "🟡 Average"
        else:
            status = "🔴 Poor"
        st.metric("Cash Cycle Days", f"{cash_cycle} {status}")

    st.markdown("---")

    # Annual COGS
    annual_cogs = st.slider("Annual COGS (₹ crore)", 10, 500, 100, step=10)

    daily_cost = (annual_cogs * 1_00_00_000) / 365  # Convert to actual rupees
    working_capital_need = (daily_cost * cash_cycle) / 1_00_00_000  # Back to crore

    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Annual COGS", f"₹{annual_cogs} cr")
    with col2:
        st.metric("Daily Cost", f"₹{daily_cost/10_00_000:.2f} cr")
    with col3:
        st.metric("Working Capital Need", f"₹{working_capital_need:.1f} cr")

    st.markdown('<div class="formula-box">', unsafe_allow_html=True)
    st.markdown(f"""
    **CASH CONVERSION CYCLE FORMULA**

    **Cash_Cycle_Days = Inventory_Days + Receivables_Days - Payables_Days**

    **Your Calculation:**
    - Inventory: {inventory_days} days
    - Receivables: {receivables_days} days
    - Payables: {payables_days} days
    - **Cash Cycle: {cash_cycle} days**

    **Working Capital Need = (COGS / 365) × Cash_Cycle_Days**

    = (₹{annual_cogs} cr / 365) × {cash_cycle} days
    = ₹{daily_cost/10_00_000:.2f} cr × {cash_cycle}
    = **₹{working_capital_need:.1f} crore**

    **Improvement Potential:**
    """)

    if inventory_days > 60:
        st.markdown(f"- Reduce inventory by 10 days → Frees ₹{(daily_cost * 10)/1_00_00_000:.1f} crore")
    if receivables_days > 45:
        st.markdown(f"- Collect 5 days faster → Frees ₹{(daily_cost * 5)/1_00_00_000:.1f} crore")
    if payables_days < 45:
        st.markdown(f"- Extend payables by 10 days → Frees ₹{(daily_cost * 10)/1_00_00_000:.1f} crore")

    st.markdown(f"""
    **ROI:** Improve cash cycle by 15 days = ₹{(daily_cost * 15)/1_00_00_000:.1f} crore freed for growth/debt repayment
    """)
    st.markdown('</div>', unsafe_allow_html=True)

    # Store this scenario so the Analytics tab can show it live, and export it
    st.session_state.working_capital_result = {
        "Inventory_Days": inventory_days,
        "Receivables_Days": receivables_days,
        "Payables_Days": payables_days,
        "Cash_Cycle_Days": cash_cycle,
        "Annual_COGS_cr": annual_cogs,
        "Working_Capital_Need_cr": round(working_capital_need, 2),
    }

    wc_df = pd.DataFrame([st.session_state.working_capital_result])

    formulas_dict_wc = {
        "Cash Conversion Cycle": "Cash_Cycle_Days = Inventory_Days + Receivables_Days - Payables_Days",
        "Working Capital Need": "Working_Capital_Need = (COGS / 365) x Cash_Cycle_Days",
    }

    st.markdown('<div class="subsection-header">💾 Export This Scenario</div>', unsafe_allow_html=True)
    col1, col2, col3, col4, col5 = st.columns(5)
    with col1:
        st.download_button("📊 CSV", wc_df.to_csv(index=False).encode(), "working_capital.csv", "text/csv", key="csv_wc")
    with col2:
        st.download_button("📋 JSON", wc_df.to_json(orient="records", indent=2).encode(), "working_capital.json", "application/json", key="json_wc")
    with col3:
        txt = export_to_text(wc_df, "Working Capital Analysis", formulas_dict_wc)
        st.download_button("📄 TXT", txt.encode(), "working_capital.txt", "text/plain", key="txt_wc")
    with col4:
        if HAS_DOCX:
            docx_bytes = export_to_docx(wc_df, "Working Capital Analysis", formulas_dict_wc)
            if docx_bytes:
                st.download_button("📝 WORD", docx_bytes, "working_capital.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="docx_wc")
    with col5:
        if HAS_PDF:
            pdf_bytes = export_to_pdf(wc_df, "Working Capital Analysis", formulas_dict_wc)
            if pdf_bytes:
                st.download_button("🔴 PDF", pdf_bytes, "working_capital.pdf", "application/pdf", key="pdf_wc")
# =========================================================================
# TAB 5: ANALYTICS
# =========================================================================
with tab5:
    st.markdown('<div class="section-header">🎯 Supply Chain Analytics Dashboard</div>', unsafe_allow_html=True)

    st.markdown("""
    Key metrics pulled live from the data you've generated or uploaded in the other tabs —
    no more placeholder numbers.
    """)

    forecast_source = st.session_state.real_forecast_df if st.session_state.real_forecast_df is not None else st.session_state.forecast_df
    vendor_source = st.session_state.real_vendor_df if st.session_state.real_vendor_df is not None else st.session_state.vendor_df
    wc_result = st.session_state.working_capital_result

    col1, col2, col3, col4 = st.columns(4)

    with col1:
        if forecast_source is not None and "Accuracy_%" in forecast_source.columns:
            st.metric("Avg Forecast Accuracy", f"{forecast_source['Accuracy_%'].mean():.1f}%")
        else:
            st.metric("Avg Forecast Accuracy", "No data")
            st.caption("Generate/upload data in Demand Forecast tab")

    with col2:
        st.metric("Inventory Turnover", "See calculator below")

    with col3:
        if vendor_source is not None and "Risk_Score" in vendor_source.columns:
            reliability = max(0, 100 - vendor_source["Risk_Score"].mean())
            st.metric("Vendor Reliability", f"{reliability:.1f}%")
        else:
            st.metric("Vendor Reliability", "No data")
            st.caption("Generate/upload data in Vendor Risk tab")

    with col4:
        if wc_result is not None:
            st.metric("Working Capital", f"₹{wc_result['Working_Capital_Need_cr']:.1f} cr")
        else:
            st.metric("Working Capital", "No data")
            st.caption("Set a scenario in Working Capital tab")

    st.markdown("---")
    st.markdown('<div class="subsection-header">📐 Inventory Turnover Calculator</div>', unsafe_allow_html=True)
    st.markdown("How efficiently inventory converts to sales — higher means less cash tied up in stock.")

    col1, col2 = st.columns(2)
    with col1:
        cogs_cr = st.number_input("Annual COGS (₹ crore)", 1.0, 1000.0, 100.0, step=5.0, key="analytics_cogs")
    with col2:
        avg_inv_cr = st.number_input("Avg Inventory Value (₹ crore)", 0.1, 500.0, 16.0, step=1.0, key="analytics_inv")

    turnover = SupplyChainFormulas.inventory_turnover(cogs_cr, avg_inv_cr)
    days_of_inventory = 365 / turnover if turnover > 0 else 0

    col1, col2 = st.columns(2)
    with col1:
        st.metric("Inventory Turnover", f"{turnover:.1f}x/year")
    with col2:
        st.metric("Days of Inventory", f"{days_of_inventory:.0f} days")

    if turnover >= 8:
        st.success("Efficient — low carrying cost.")
    elif turnover >= 4:
        st.info("Typical manufacturing range — good balance.")
    else:
        st.warning("Heavy — excess capital tied up in stock.")

    st.markdown("---")
    st.markdown("**Key Metrics Reference:**")

    st.info("""
    **Forecast Accuracy** (Target: >85%)
    - Measures prediction vs actual demand
    - Better accuracy → Less inventory waste

    **Inventory Turnover** (Target: 6-10x/year)
    - Higher = More efficient
    - Industry avg: 4-6x
    - Best-in-class: 8-12x

    **Vendor Reliability** (Target: >95%)
    - Combination of OTIF + Quality + Payment + Price risk (100 - avg Risk Score)
    - <90% = Supply risk

    **Working Capital** (Target: Minimize)
    - Cash tied up in operations
    - Every ₹1 cr freed = cash for growth
    """)
st.markdown("---")
st.caption(f"⚙️ ChainIntel Pro v1.0 | Session {st.session_state.session_id} | {DEPLOY_ENV.upper()} | Developed by Randy Singh | Kalsnet (KNet) Consulting")
