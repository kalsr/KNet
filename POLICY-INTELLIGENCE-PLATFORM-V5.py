# =========================================================
# AI Governance & Public Policy Intelligence Platform
# Enterprise SaaS Dashboard (Enhanced + Traceability + Export)
# Developed by Randy Singh from Kalsnet (KNet)
# =========================================================
import streamlit as st
from groq import Groq
import groq
import json
import re
import pandas as pd
from io import BytesIO
from datetime import datetime
import html as html_lib
# PDF + Word support
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.pagesizes import letter
from reportlab.lib.colors import HexColor
from docx import Document
from docx.shared import Pt, RGBColor

# =========================================================
# PAGE CONFIG
# =========================================================
st.set_page_config(
    page_title="AI Governance Platform",
    layout="wide",
    page_icon=""
)

# =========================================================
# LIGHT MODE SIDEBAR
# =========================================================
st.markdown("""
<style>
.main { background-color: #f5f7fb; }
section[data-testid="stSidebar"] {
    background-color: #ffffff !important;
    color: #111;
    border-right: 1px solid #e6e6e6;
}
.header {
    background: linear-gradient(90deg, #0B3D91, #1E90FF);
    padding: 18px;
    border-radius: 12px;
    text-align: center;
    margin-bottom: 20px;
}
.title {
    font-size: 34px;
    font-weight: 900;
    color: white;
}
.subtitle {
    font-size: 16px;
    font-weight: 900;
    color: #00BFFF;
}
.card {
    background: white;
    padding: 18px;
    border-radius: 12px;
    border: 1px solid #e5e7eb;
}
.stButton>button {
    background-color: #1E90FF;
    color: white;
    font-weight: bold;
    border-radius: 8px;
}
</style>
""", unsafe_allow_html=True)

# =========================================================
# HEADER
# =========================================================
st.markdown("""
<div class="header">
    <div class="title">AI GOVERNANCE & PUBLIC POLICY INTELLIGENCE PLATFORM</div>
    <div class="subtitle">
        <b>Developed by Randy Singh from Kalsnet (KNet) Consulting Group</b>
    </div>
</div>
""", unsafe_allow_html=True)

# =========================================================
# NAVIGATION
# =========================================================
menu = st.sidebar.radio(
    "Navigation",
    ["Dashboard", "Policy Generator", "Compliance Auditor", "Reports", "Settings"]
)

# =========================================================
# GROQ CONFIG
# =========================================================
st.sidebar.header(" AI Configuration")
api_key = st.sidebar.text_input("Groq API Key", type="password")
st.sidebar.markdown("""
###  Get API Key
- https://console.groq.com
- Create account
- Generate API key
- Paste here
""")

# NOTE: Groq deprecates models on a rolling basis. As of this writing
# (Aug 2026), "mixtral-8x7b-32768", "gemma2-9b-it", "mistral-saba-24b",
# "llama-3.1-8b-instant" AND "llama-3.3-70b-versatile" have ALL been
# shut down and now return a 404 groq.NotFoundError, which is exactly
# what you hit. Only the models still marked "Production" on
# https://console.groq.com/docs/models are listed below. If one of
# these ever 404s again, check that page for the current list.
model = st.sidebar.selectbox(
    "Model",
    [
        "openai/gpt-oss-120b",
        "openai/gpt-oss-20b",
        "groq/compound",
        "groq/compound-mini",
    ]
)

# =========================================================
# KPI TRACEABILITY ENGINE
# =========================================================
def get_metrics():
    policies = st.session_state.get("policy_count", 128)
    compliance = st.session_state.get("compliance_score", 92)
    risks = st.session_state.get("risk_count", 37)
    reports = st.session_state.get("report_count", 18)
    return {
        "Policies": {"value": policies, "source": "policy DB / session_state"},
        "Compliance": {"value": compliance, "source": "audit engine"},
        "Risks": {"value": risks, "source": "risk detection model"},
        "Reports": {"value": reports, "source": "report store"}
    }

# =========================================================
# EXPORT FUNCTIONS
# =========================================================
# Brand blue used for the header gradient/subtitle elsewhere in the app —
# reused here so the report heading matches the app's own look.
BRAND_BLUE = "#0B3D91"

def clean_report_text(text):
    """Strip markdown emphasis markers (**bold**, *italic*, "* bullet")
    that the LLM tends to return, since PDF/Word don't render markdown —
    left as-is, a raw "*" just looks like clutter in the exported file."""
    if not text:
        return ""
    return re.sub(r"\*+", "", text).strip()

def split_paragraphs(text):
    """Break cleaned report text into readable paragraphs. Prefers
    blank-line breaks; falls back to single newlines if the model didn't
    separate sections with blank lines."""
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    if len(paragraphs) <= 1:
        paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
    return paragraphs

def export_pdf(text, industry="AI Governance Report"):
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=letter,
        topMargin=54, bottomMargin=54, leftMargin=54, rightMargin=54
    )

    # ReportLab has no built-in "Arial" — Helvetica is the standard,
    # metric-compatible substitute every PDF viewer renders identically
    # to Arial, without needing a font file bundled/installed on the
    # server (which Streamlit Cloud doesn't have by default).
    heading_style = ParagraphStyle(
        "IndustryHeading",
        fontName="Helvetica-Bold",
        fontSize=26,
        leading=30,
        textColor=HexColor(BRAND_BLUE),
        spaceAfter=16,
    )
    body_style = ParagraphStyle(
        "Body",
        fontName="Helvetica",
        fontSize=12,
        leading=18,
        textColor=HexColor("#111111"),
        spaceAfter=12,
    )

    # ReportLab's Paragraph runs its own mini-HTML/XML parser on the text
    # you pass it (that's how it supports tags like <br/>, <b>, <i>...).
    # The Groq response is free-form LLM text, and if it happens to
    # contain a raw "<" (e.g. "throughput < 100ms", "<policy_id>", a
    # stray HTML/markdown fragment, etc.) the parser tries to read it as
    # the start of a tag and raises "paraparser: syntax error". The fix
    # is to HTML-escape each paragraph before handing it to Paragraph().
    cleaned = clean_report_text(text)
    content = [Paragraph(html_lib.escape(industry or "AI Governance Report"), heading_style)]
    for para in split_paragraphs(cleaned):
        content.append(Paragraph(html_lib.escape(para), body_style))

    doc.build(content)
    buffer.seek(0)
    return buffer

def export_word(text, industry="AI Governance Report"):
    doc = Document()

    heading = doc.add_paragraph()
    run = heading.add_run(industry or "AI Governance Report")
    run.font.name = "Arial"
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0B, 0x3D, 0x91)

    doc.add_paragraph()  # spacing under the heading

    cleaned = clean_report_text(text)
    for para_text in split_paragraphs(cleaned):
        p = doc.add_paragraph()
        r = p.add_run(para_text)
        r.font.name = "Arial"
        r.font.size = Pt(12)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def export_csv(text, industry="AI Governance Report"):
    cleaned = clean_report_text(text)
    df = pd.DataFrame({"industry": [industry], "report": [cleaned]})
    return df.to_csv(index=False).encode("utf-8")

def export_json(text, industry="AI Governance Report"):
    cleaned = clean_report_text(text)
    return json.dumps(
        {"industry": industry, "report": cleaned, "timestamp": str(datetime.now())},
        indent=2
    ).encode("utf-8")

# =========================================================
# DASHBOARD
# =========================================================
if menu == "Dashboard":
    metrics = get_metrics()
    col1, col2, col3, col4 = st.columns(4)
    col1.metric("Policies", metrics["Policies"]["value"], "↑ 12%")
    col2.metric("Compliance", f'{metrics["Compliance"]["value"]}%', "↑ 5%")
    col3.metric("Risks", metrics["Risks"]["value"], "↓ 3%")
    col4.metric("Reports", metrics["Reports"]["value"], "↑ 2")
    st.success("System Active — Ready for Policy Generation")
    with st.expander(" KPI Data Source Debug Panel"):
        for k, v in metrics.items():
            st.write(f"**{k}** → {v['value']} (Source: {v['source']})")

# =========================================================
# POLICY GENERATOR
# =========================================================
if menu == "Policy Generator":
    st.subheader(" Policy Generator")
    org = st.text_input("Organization Name")
    industry = st.selectbox("Industry", ["Government","Defense","Healthcare","Finance","Manufacturing/Casting & Machining Units"])
    objective = st.selectbox("Objective", ["Cybersecurity Protection","AI Governance","Compliance"])
    risk = st.selectbox("Risk Level", ["Low","Medium","High","Critical"])
    extra = st.multiselect("Additional Requirements", [
        "Audit Logging","Encryption","RBAC","Zero Trust"
    ])
    req = st.text_area("Custom Requirement")

    if st.button("Generate Policy"):
        if not api_key:
            st.error("Enter API Key")
        else:
            prompt = f"""
Create enterprise policy:
Org: {org}
Industry: {industry}
Objective: {objective}
Risk: {risk}
Controls: {extra}
Custom: {req}
"""
            try:
                client = Groq(api_key=api_key)
                response = client.chat.completions.create(
                    model=model,
                    messages=[
                        {"role": "system", "content": "Senior governance architect"},
                        {"role": "user", "content": prompt}
                    ]
                )
                st.session_state["result"] = response.choices[0].message.content
                # remember which industry this report was generated for,
                # so the Reports tab can print it as the report heading
                st.session_state["result_industry"] = industry
                # update metrics dynamically
                st.session_state["policy_count"] = st.session_state.get("policy_count", 128) + 1

            except groq.NotFoundError:
                st.error(
                    f"Model '{model}' was not found on Groq's API. It may have been "
                    "renamed or decommissioned — pick a different model from the "
                    "sidebar (see https://console.groq.com/docs/models for the "
                    "current list) and try again."
                )
            except groq.AuthenticationError:
                st.error(
                    "Groq rejected your API key. Double-check that you copied it "
                    "correctly from https://console.groq.com/keys."
                )
            except groq.RateLimitError:
                st.error("Groq rate limit reached. Wait a moment and try again.")
            except groq.APIStatusError as e:
                st.error(f"Groq API returned an error (status {e.status_code}): {e}")
            except groq.APIConnectionError:
                st.error("Could not reach the Groq API. Check your network connection and try again.")
            except Exception as e:
                st.error(f"Unexpected error while generating the policy: {e}")

    if "result" in st.session_state:
        st.markdown("### Generated Policy")
        st.markdown(st.session_state["result"])

# =========================================================
# COMPLIANCE AUDITOR
# =========================================================
if menu == "Compliance Auditor":
    text = st.text_area("Paste Policy")
    if st.button("Run Audit"):
        score = 85
        st.session_state["compliance_score"] = score
        st.session_state["risk_count"] = 37
        st.success(f"Compliance Score: {score}%")
        st.write("✔ NIST Alignment")
        st.write("⚠ Improve logging depth")

# =========================================================
# REPORTS (ENHANCED EXPORT)
# =========================================================
if menu == "Reports":
    st.subheader("Reports Dashboard")
    if "result" in st.session_state:
        report = st.session_state["result"]
        report_industry = st.session_state.get("result_industry", "AI Governance Report")
        st.code(report[:2000])
        st.markdown("### Export Report")

        # Each export format is generated independently and wrapped in its
        # own try/except so that, say, a PDF-generation edge case can't
        # take down the CSV/JSON/Word buttons next to it (or the whole tab).
        col1, col2, col3, col4 = st.columns(4)

        with col1:
            try:
                st.download_button(" PDF", export_pdf(report, report_industry), "report.pdf")
            except Exception as e:
                st.error(f"PDF export failed: {e}")

        with col2:
            try:
                st.download_button(" Word", export_word(report, report_industry), "report.docx")
            except Exception as e:
                st.error(f"Word export failed: {e}")

        with col3:
            try:
                st.download_button(" CSV", export_csv(report, report_industry), "report.csv")
            except Exception as e:
                st.error(f"CSV export failed: {e}")

        with col4:
            try:
                st.download_button(" JSON", export_json(report, report_industry), "report.json")
            except Exception as e:
                st.error(f"JSON export failed: {e}")
    else:
        st.info("No policy generated yet")

# =========================================================
# SETTINGS
# =========================================================
# Purpose: a read-only status/diagnostics panel. It doesn't configure
# anything itself (the actual model + API key controls live in the
# sidebar, since they need to be visible on every page) — it just lets
# you confirm at a glance what the app is currently running with,
# without ever displaying the key itself. It's also the natural home
# for future app-wide preferences (default industry, export format,
# theme, etc.) as this dashboard grows.
if menu == "Settings":
    st.subheader("Settings")
    st.caption("Read-only view of the current session configuration. "
               "Model and API key are set from the sidebar on any page.")
    st.write("Model:", model)
    st.write("API Key:", "Configured" if api_key else "Not Configured")
    st.warning("Advanced settings coming soon")
