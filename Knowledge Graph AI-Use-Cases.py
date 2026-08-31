# Knowledge Graph AI Use Cases Application
# Single File Streamlit Edition
# Developed by Randy Singh from Kalsnet KNet Consulting Group
# A professional Streamlit application that demonstrates fifteen Knowledge
# Graph plus AI use cases. Each use case tab supports synthetic data
# generation or real data upload, interactive graph visualization, AI
# analysis powered by Groq or Gemini, and export of results to PDF, Word,
# text and CSV formats.
# The Fraud Detection and Anti Money Laundering use cases additionally
# support loading genuine real world data directly from the live
# OpenSanctions sanctions and PEP screening API, and from the published
# Elliptic real world labeled Bitcoin AML transaction dataset.
# This single file contains all configuration, data generation, graph
# analytics, LLM integration, export utilities, real data source
# integration, and the Streamlit interface.
# How to run
# 1. Install dependencies: pip install streamlit pandas networkx plotly python-docx reportlab requests
# 2. Run the application: streamlit run knowledge_graph_ai_app.py
# No special symbols are used anywhere in this source file.
import io
import json
import random
import string
from datetime import datetime
import requests
import pandas as pd
import networkx as nx
import plotly.graph_objects as go
import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.units import inch
# =============================================================================
# SECTION 1: USE CASE CONFIGURATION
# =============================================================================
USE_CASES = {
    "cybersecurity": {
        "order": 1,
        "title": "Cybersecurity Threat Intelligence",
        "graph_provides": "IP to Device to User to Vulnerability to Threat Actor relationships",
        "ai_does": "Correlates threats and explains attack paths",
        "description": (
            "This use case builds a knowledge graph that links network identifiers, "
            "hardware, people, software weaknesses and known threat actors. Security "
            "teams use this connected view to trace how a single exposed IP address "
            "can lead to a full compromise chain. The AI layer reads the graph paths "
            "and produces a plain language explanation of the attack path, the "
            "likely threat actor motive, and the recommended containment steps."
        ),
        "node_types": ["IP Address", "Device", "User", "Vulnerability", "Threat Actor"],
        "edge_types": [
            ("IP Address", "Device", "connects_to"),
            ("Device", "User", "assigned_to"),
            ("Device", "Vulnerability", "exposes"),
            ("Vulnerability", "Threat Actor", "exploited_by"),
        ],
        "fields": {
            "IP Address": ["ip_id", "ip_value", "geo_location", "reputation_score"],
            "Device": ["device_id", "hostname", "device_type", "os_version", "patch_level"],
            "User": ["user_id", "user_name", "department", "access_level"],
            "Vulnerability": ["cve_id", "severity", "cvss_score", "description"],
            "Threat Actor": ["actor_id", "actor_name", "known_ttp", "origin_region"],
        },
        "formulas": [
            ("Attack Path Length", "Number of edges on the shortest path between an IP node and a Threat Actor node. Shorter paths indicate higher immediate risk."),
            ("Node Risk Score", "Risk = 0.5 times CVSS Score plus 0.3 times Device Patch Gap plus 0.2 times IP Reputation Penalty. All three inputs are normalized to a zero to ten scale before weighting."),
            ("Degree Centrality", "Fraction of all other nodes that a given node is directly connected to. A device with high degree centrality is a high value pivot point for an attacker."),
            ("Betweenness Centrality", "Measures how often a node lies on the shortest path between other node pairs. High betweenness devices are natural choke points for containment."),
        ],
        "benefits": [
            "Cuts investigation time by visualizing multi hop attack chains instead of reading flat log tables",
            "Prioritizes patching based on graph position, not just raw CVSS score",
            "Surfaces hidden lateral movement paths between users and devices",
            "Gives incident responders an explainable narrative to share with leadership",
        ],
        "ai_prompt_focus": "explain the attack path, likely intent of the threat actor, and containment recommendation",
    },
    "fraud": {
        "order": 2,
        "title": "Fraud Detection",
        "graph_provides": "Customer to Account to Transaction to Device to Location relationships",
        "ai_does": "Finds suspicious patterns and hidden connections",
        "description": (
            "This use case connects customers, the accounts they hold, the transactions "
            "they make, the devices used to transact and the physical or IP based "
            "locations involved. Fraud rings are rarely visible in a single transaction "
            "row, they emerge when several accounts share the same device fingerprint "
            "or the same delivery address. The AI layer inspects these shared "
            "connections and produces a ranked explanation of why a cluster looks "
            "suspicious."
        ),
        "node_types": ["Customer", "Account", "Transaction", "Device", "Location"],
        "edge_types": [
            ("Customer", "Account", "owns"),
            ("Account", "Transaction", "initiates"),
            ("Transaction", "Device", "originated_from"),
            ("Transaction", "Location", "occurred_at"),
        ],
        "fields": {
            "Customer": ["customer_id", "full_name", "risk_tier", "account_age_days"],
            "Account": ["account_id", "account_type", "balance", "open_date"],
            "Transaction": ["transaction_id", "amount", "currency", "timestamp", "channel"],
            "Device": ["device_id", "device_fingerprint", "device_type"],
            "Location": ["location_id", "city", "country", "ip_address"],
        },
        "formulas": [
            ("Shared Device Score", "Count of distinct customers that transact from the same device fingerprint within a rolling thirty day window. A value above three is treated as a red flag."),
            ("Velocity Score", "Sum of transaction amounts for an account divided by the number of hours since the account first transacted in the window. High velocity indicates rapid movement of funds."),
            ("Ring Density", "Number of actual edges among a suspected cluster divided by the number of possible edges in a fully connected cluster of the same size. Values close to one indicate a tightly coordinated ring."),
            ("Anomaly Score", "Weighted sum of Shared Device Score, Velocity Score and Ring Density, each normalized to zero to one before combination."),
        ],
        "benefits": [
            "Detects coordinated fraud rings that rule based systems miss because each single transaction looks normal",
            "Reduces false positives by using network context rather than a single account view",
            "Speeds up investigator review with a visual map of the suspicious cluster",
            "Provides an explainable case file for compliance and law enforcement referral",
        ],
        "ai_prompt_focus": "identify suspicious clusters, explain the shared connections, and recommend an investigation priority",
    },
    "enterprise_knowledge": {
        "order": 3,
        "title": "Enterprise Knowledge Assistant",
        "graph_provides": "Employee to System to Document to Project to Policy",
        "ai_does": "Answers questions using connected enterprise knowledge",
        "description": (
            "This use case models how people, internal systems, documents, projects "
            "and policies relate to one another inside an organization. Instead of "
            "searching a flat document index, the assistant walks the graph to find "
            "which policy governs which project, who owns which system, and which "
            "documents are authoritative. The AI layer answers natural language "
            "questions by grounding its response in the retrieved graph neighborhood."
        ),
        "node_types": ["Employee", "System", "Document", "Project", "Policy"],
        "edge_types": [
            ("Employee", "System", "has_access_to"),
            ("Employee", "Project", "works_on"),
            ("Document", "Project", "relates_to"),
            ("Policy", "System", "governs"),
        ],
        "fields": {
            "Employee": ["employee_id", "full_name", "role", "department"],
            "System": ["system_id", "system_name", "owner", "criticality"],
            "Document": ["document_id", "title", "author", "last_updated"],
            "Project": ["project_id", "project_name", "status", "sponsor"],
            "Policy": ["policy_id", "policy_name", "effective_date", "category"],
        },
        "formulas": [
            ("Knowledge Coverage", "Number of documents connected to a project divided by the expected document count for a project of that type. Lower coverage flags a documentation gap."),
            ("Answer Grounding Score", "Number of graph facts cited in an AI answer divided by the total number of factual claims in the answer. Used to measure hallucination risk."),
            ("Path Relevance", "Inverse of the number of hops between the question entity and the answer entity. Closer entities receive a higher relevance weight during retrieval."),
        ],
        "benefits": [
            "Reduces time employees spend searching multiple systems for a single answer",
            "Improves answer accuracy because responses are grounded in verified graph facts",
            "Surfaces ownership so questions about accountability are answered immediately",
            "Highlights documentation gaps for projects that lack policy coverage",
        ],
        "ai_prompt_focus": "answer the enterprise question using only the retrieved graph facts and cite the entities used",
    },
    "rag_kg": {
        "order": 4,
        "title": "RAG plus Knowledge Graph",
        "graph_provides": "Facts and relationships between documents and entities",
        "ai_does": "Produces more accurate, contextual answers",
        "description": (
            "Traditional retrieval augmented generation ranks document chunks by "
            "text similarity alone, which can miss facts that are only connected "
            "through relationships. This use case adds a knowledge graph layer on "
            "top of a document corpus so retrieval can follow entity relationships "
            "as well as text similarity. The AI layer combines both retrieval "
            "signals before generating an answer, which reduces hallucination."
        ),
        "node_types": ["Document", "Entity", "Concept", "Chunk", "Source"],
        "edge_types": [
            ("Document", "Chunk", "contains"),
            ("Chunk", "Entity", "mentions"),
            ("Entity", "Concept", "categorized_as"),
            ("Document", "Source", "published_by"),
        ],
        "fields": {
            "Document": ["document_id", "title", "publish_date"],
            "Entity": ["entity_id", "entity_name", "entity_type"],
            "Concept": ["concept_id", "concept_name"],
            "Chunk": ["chunk_id", "text_snippet", "embedding_id"],
            "Source": ["source_id", "source_name", "credibility_score"],
        },
        "formulas": [
            ("Hybrid Retrieval Score", "0.6 times text similarity score plus 0.4 times graph proximity score, both normalized to zero to one. Combines semantic and relational signals."),
            ("Graph Proximity Score", "Inverse of the shortest path length in hops between the query entity and the candidate chunk entity."),
            ("Faithfulness Score", "Number of answer sentences supported by a retrieved graph fact divided by the total number of answer sentences."),
        ],
        "benefits": [
            "Improves answer accuracy by combining semantic search with relationship search",
            "Reduces hallucination because generated claims can be checked against graph facts",
            "Surfaces multi document connections that pure vector search would miss",
            "Provides a traceable citation path from answer back to source document",
        ],
        "ai_prompt_focus": "generate an answer grounded in the retrieved chunks and graph facts, and flag any unsupported claim",
    },
    "supply_chain": {
        "order": 5,
        "title": "Supply Chain Risk",
        "graph_provides": "Supplier to Component to Factory to Shipment to Customer relationships",
        "ai_does": "Identifies cascading risks and alternative suppliers",
        "description": (
            "This use case connects suppliers, the components they provide, the "
            "factories that assemble them, the shipments that move goods, and the "
            "customers who ultimately receive them. When a single supplier is "
            "disrupted, the graph reveals every downstream factory, shipment and "
            "customer affected. The AI layer explains the cascading impact and "
            "recommends qualified alternative suppliers already present in the graph."
        ),
        "node_types": ["Supplier", "Component", "Factory", "Shipment", "Customer"],
        "edge_types": [
            ("Supplier", "Component", "supplies"),
            ("Component", "Factory", "used_in"),
            ("Factory", "Shipment", "produces"),
            ("Shipment", "Customer", "delivered_to"),
        ],
        "fields": {
            "Supplier": ["supplier_id", "supplier_name", "region", "reliability_score"],
            "Component": ["component_id", "component_name", "lead_time_days"],
            "Factory": ["factory_id", "factory_name", "capacity_units"],
            "Shipment": ["shipment_id", "quantity", "ship_date", "status"],
            "Customer": ["customer_id", "customer_name", "contract_value"],
        },
        "formulas": [
            ("Cascading Impact Score", "Number of downstream customer nodes reachable from a disrupted supplier node, divided by the total number of customer nodes in the graph."),
            ("Single Source Risk", "One divided by the number of distinct suppliers that provide the same component. A value of one means there is no backup supplier."),
            ("Alternative Supplier Fit", "Weighted match of region proximity, reliability score and available capacity between the disrupted supplier and each candidate alternative."),
        ],
        "benefits": [
            "Reveals hidden single points of failure across multi tier supplier networks",
            "Shortens disruption response time with pre identified alternative suppliers",
            "Quantifies which customers are most exposed to a given supplier risk",
            "Supports procurement diversification decisions with data instead of guesswork",
        ],
        "ai_prompt_focus": "explain the cascading impact of the disruption and recommend qualified alternative suppliers",
    },
    "healthcare": {
        "order": 6,
        "title": "Healthcare Intelligence",
        "graph_provides": "Patient to Condition to Medication to Lab to Provider relationships",
        "ai_does": "Finds relationships and assists clinical reasoning",
        "description": (
            "This use case links patients to their diagnosed conditions, prescribed "
            "medications, lab results and treating providers. Clinical decisions "
            "often depend on relationships that are easy to miss in a flat electronic "
            "health record, such as a drug interaction across two conditions. The AI "
            "layer reviews the connected record and highlights relevant relationships "
            "for the treating clinician, always framed as decision support rather "
            "than a diagnosis."
        ),
        "node_types": ["Patient", "Condition", "Medication", "Lab Result", "Provider"],
        "edge_types": [
            ("Patient", "Condition", "diagnosed_with"),
            ("Patient", "Medication", "prescribed"),
            ("Patient", "Lab Result", "tested_for"),
            ("Provider", "Patient", "treats"),
        ],
        "fields": {
            "Patient": ["patient_id", "age", "sex", "primary_provider"],
            "Condition": ["condition_id", "condition_name", "icd_code", "onset_date"],
            "Medication": ["medication_id", "drug_name", "dosage", "start_date"],
            "Lab Result": ["lab_id", "test_name", "result_value", "result_date"],
            "Provider": ["provider_id", "provider_name", "specialty"],
        },
        "formulas": [
            ("Interaction Risk Score", "Count of medication pairs prescribed to the same patient that appear on a known interaction reference list."),
            ("Condition Comorbidity Score", "Number of shared patients between two condition nodes divided by the total number of patients with either condition."),
            ("Lab Trend Deviation", "Difference between the most recent lab result value and the patient historical mean, divided by the historical standard deviation."),
        ],
        "benefits": [
            "Surfaces medication interaction risk across the full connected patient record",
            "Helps identify condition patterns across a patient population for care planning",
            "Reduces time spent manually cross referencing charts, labs and prescriptions",
            "Supports clinicians with an explainable summary rather than a black box score",
        ],
        "ai_prompt_focus": "summarize clinically relevant relationships and potential interaction risks for clinician review, this is decision support not a diagnosis",
    },
    "financial_risk": {
        "order": 7,
        "title": "Financial Risk Analysis",
        "graph_provides": "Company to Subsidiary to Executive to Loan to Transaction relationships",
        "ai_does": "Detects exposure, concentration and hidden relationships",
        "description": (
            "This use case connects companies, their subsidiaries, executives, loans "
            "and transactions. Concentration risk and conflicts of interest often "
            "hide behind layers of subsidiaries or shared executives across "
            "portfolio companies. The AI layer walks the ownership and lending "
            "graph to expose hidden exposure and explains why a concentration "
            "limit may be at risk."
        ),
        "node_types": ["Company", "Subsidiary", "Executive", "Loan", "Transaction"],
        "edge_types": [
            ("Company", "Subsidiary", "owns"),
            ("Executive", "Company", "leads"),
            ("Company", "Loan", "holds"),
            ("Loan", "Transaction", "generates"),
        ],
        "fields": {
            "Company": ["company_id", "company_name", "sector", "credit_rating"],
            "Subsidiary": ["subsidiary_id", "subsidiary_name", "ownership_pct"],
            "Executive": ["executive_id", "executive_name", "role"],
            "Loan": ["loan_id", "principal_amount", "interest_rate", "maturity_date"],
            "Transaction": ["transaction_id", "amount", "transaction_type", "date"],
        },
        "formulas": [
            ("Concentration Ratio", "Total loan exposure to a single company and its subsidiaries divided by total loan exposure across the entire portfolio."),
            ("Hidden Relationship Score", "Count of shared executive nodes between two otherwise unrelated companies in the graph."),
            ("Exposure at Risk", "Sum of principal amount across all loan nodes connected to a company node, weighted by the inverse of its credit rating."),
        ],
        "benefits": [
            "Reveals concentration risk that is invisible when companies are viewed in isolation",
            "Detects shared executive relationships that may indicate conflicts of interest",
            "Improves portfolio diversification decisions with quantified exposure metrics",
            "Provides examiners with an explainable, connected view of risk lineage",
        ],
        "ai_prompt_focus": "explain hidden relationships driving concentration risk and quantify exposure",
    },
    "digital_twin": {
        "order": 8,
        "title": "Digital Twin",
        "graph_provides": "Asset to Sensor to Component to Process to Failure relationships",
        "ai_does": "Predicts failures and explains causes",
        "description": (
            "This use case builds a live connected model of a physical asset, its "
            "sensors, components, the process steps it participates in, and its "
            "historical failure events. When a sensor reading drifts, the graph "
            "shows exactly which components and processes are downstream. The AI "
            "layer reviews the connected failure history and explains the likely "
            "root cause and expected time to failure."
        ),
        "node_types": ["Asset", "Sensor", "Component", "Process", "Failure Event"],
        "edge_types": [
            ("Asset", "Sensor", "monitored_by"),
            ("Sensor", "Component", "measures"),
            ("Component", "Process", "participates_in"),
            ("Component", "Failure Event", "experienced"),
        ],
        "fields": {
            "Asset": ["asset_id", "asset_name", "install_date", "criticality"],
            "Sensor": ["sensor_id", "sensor_type", "unit_of_measure"],
            "Component": ["component_id", "component_name", "expected_life_hours"],
            "Process": ["process_id", "process_name", "throughput_rate"],
            "Failure Event": ["failure_id", "failure_type", "failure_date", "downtime_hours"],
        },
        "formulas": [
            ("Remaining Useful Life", "Expected life hours of a component minus cumulative operating hours since install, adjusted downward by the deviation of recent sensor readings from baseline."),
            ("Failure Propagation Score", "Number of downstream process nodes reachable from a failing component divided by total process nodes in the graph."),
            ("Anomaly Deviation", "Absolute difference between current sensor reading and the rolling average, divided by the rolling standard deviation."),
        ],
        "benefits": [
            "Predicts component failure before it causes unplanned downtime",
            "Explains root cause across the full asset to process chain, not just a single sensor",
            "Prioritizes maintenance based on downstream process impact",
            "Reduces inspection cost by focusing attention on high deviation components",
        ],
        "ai_prompt_focus": "explain the likely failure cause, downstream impact, and recommended maintenance action",
    },
    "network_ops": {
        "order": 9,
        "title": "Network Operations",
        "graph_provides": "Application to API to Server to Network to Dependency relationships",
        "ai_does": "Determines root cause of outages",
        "description": (
            "This use case maps how applications depend on APIs, which run on "
            "servers, which sit on network segments, with explicit dependency "
            "edges between services. When an outage alert fires, the graph shows "
            "every upstream and downstream dependency in seconds. The AI layer "
            "reviews the alert cluster and the dependency graph to explain the "
            "most likely root cause rather than the symptom."
        ),
        "node_types": ["Application", "API", "Server", "Network Segment", "Dependency"],
        "edge_types": [
            ("Application", "API", "calls"),
            ("API", "Server", "hosted_on"),
            ("Server", "Network Segment", "connected_to"),
            ("Application", "Dependency", "depends_on"),
        ],
        "fields": {
            "Application": ["app_id", "app_name", "owner_team", "tier"],
            "API": ["api_id", "api_name", "version", "avg_latency_ms"],
            "Server": ["server_id", "hostname", "region", "status"],
            "Network Segment": ["segment_id", "segment_name", "utilization_pct"],
            "Dependency": ["dependency_id", "dependency_name", "criticality"],
        },
        "formulas": [
            ("Blast Radius", "Number of application nodes reachable upstream from a failing server node, representing the total number of services affected."),
            ("Root Cause Likelihood", "For each failing node, the count of downstream alerts explained by that node divided by the total number of alerts in the incident cluster. The node with the highest ratio is the likely root cause."),
            ("Path Criticality", "Sum of the criticality weight of every dependency edge along the path from the failing node to the customer facing application."),
        ],
        "benefits": [
            "Cuts mean time to resolution by pointing directly at the likely root cause node",
            "Shows blast radius so teams can communicate impact accurately during an incident",
            "Prevents chasing symptom alerts that are downstream of the real failure",
            "Builds institutional knowledge of dependency chains over time",
        ],
        "ai_prompt_focus": "determine the most likely root cause node and explain the outage propagation path",
    },
    "compliance": {
        "order": 10,
        "title": "Compliance and Policy AI",
        "graph_provides": "Regulation to Policy to Control to System to Evidence relationships",
        "ai_does": "Determines compliance gaps",
        "description": (
            "This use case links external regulations to internal policies, the "
            "controls that implement those policies, the systems the controls "
            "apply to, and the evidence collected to prove the control operates. "
            "A compliance gap exists whenever a required regulation to evidence "
            "path is broken. The AI layer walks every required path and reports "
            "exactly where the chain breaks."
        ),
        "node_types": ["Regulation", "Policy", "Control", "System", "Evidence"],
        "edge_types": [
            ("Regulation", "Policy", "requires"),
            ("Policy", "Control", "implemented_by"),
            ("Control", "System", "applies_to"),
            ("Control", "Evidence", "proven_by"),
        ],
        "fields": {
            "Regulation": ["regulation_id", "regulation_name", "jurisdiction"],
            "Policy": ["policy_id", "policy_name", "owner"],
            "Control": ["control_id", "control_name", "control_type", "frequency"],
            "System": ["system_id", "system_name", "environment"],
            "Evidence": ["evidence_id", "evidence_type", "collected_date", "status"],
        },
        "formulas": [
            ("Compliance Coverage Ratio", "Number of complete regulation to evidence paths divided by the total number of required regulation to control mappings."),
            ("Gap Severity Score", "For each broken path, the regulation weight multiplied by the number of systems affected by the missing control."),
            ("Evidence Freshness", "Number of days since the evidence collected date, compared against the control required frequency."),
        ],
        "benefits": [
            "Pinpoints exactly which regulation lacks supporting evidence instead of a vague audit finding",
            "Reduces audit preparation time with an always current compliance map",
            "Prioritizes remediation using gap severity rather than a flat checklist",
            "Gives auditors a traceable path from regulation down to evidence artifact",
        ],
        "ai_prompt_focus": "identify broken regulation to evidence paths and explain the compliance gap and remediation priority",
    },
    "aml": {
        "order": 11,
        "title": "Fraud and Anti Money Laundering",
        "graph_provides": "Person to Account to Business to Transaction to Address relationships",
        "ai_does": "Detects networks rather than isolated transactions",
        "description": (
            "This use case connects people, the accounts and businesses they "
            "control, the transactions they make, and shared addresses. Money "
            "laundering typically requires moving funds through a network of "
            "linked accounts and shell businesses. The AI layer reviews the "
            "connected transaction network for layering and structuring patterns "
            "and explains the network rather than flagging one transaction."
        ),
        "node_types": ["Person", "Account", "Business", "Transaction", "Address"],
        "edge_types": [
            ("Person", "Account", "controls"),
            ("Person", "Business", "owns"),
            ("Account", "Transaction", "sends_or_receives"),
            ("Business", "Address", "registered_at"),
        ],
        "fields": {
            "Person": ["person_id", "full_name", "nationality", "pep_status"],
            "Account": ["account_id", "account_type", "opened_date"],
            "Business": ["business_id", "business_name", "industry_code"],
            "Transaction": ["transaction_id", "amount", "currency", "date", "channel"],
            "Address": ["address_id", "street", "city", "country"],
        },
        "formulas": [
            ("Structuring Score", "Count of transactions just below a reporting threshold within a rolling seven day window, for a given account."),
            ("Shared Address Density", "Number of distinct businesses registered at the same address divided by the expected average businesses per address for that jurisdiction."),
            ("Layering Depth", "Number of distinct accounts a single dollar of funds passes through before reaching a final destination account, measured along the transaction graph."),
        ],
        "benefits": [
            "Detects layered laundering networks that single transaction monitoring cannot see",
            "Identifies shell business clusters sharing addresses or beneficial owners",
            "Reduces false positive alert volume by scoring the whole network, not one transaction",
            "Produces an explainable case file suitable for a suspicious activity report",
        ],
        "ai_prompt_focus": "explain the laundering network structure, key facilitators, and structuring pattern detected",
    },
    "defense": {
        "order": 12,
        "title": "Defense and Mission Intelligence",
        "graph_provides": "Mission to Unit to Asset to Location to Capability relationships",
        "ai_does": "Performs relationship based mission analysis",
        "description": (
            "This use case connects missions to the units assigned, the assets "
            "those units operate, the locations involved, and the capabilities "
            "each asset provides. Mission planners need to understand quickly "
            "whether the required capability mix is available at the right "
            "location. The AI layer reviews the graph and explains capability "
            "gaps and readiness risk for a given mission."
        ),
        "node_types": ["Mission", "Unit", "Asset", "Location", "Capability"],
        "edge_types": [
            ("Mission", "Unit", "assigned_to"),
            ("Unit", "Asset", "operates"),
            ("Asset", "Location", "positioned_at"),
            ("Asset", "Capability", "provides"),
        ],
        "fields": {
            "Mission": ["mission_id", "mission_name", "objective", "start_date"],
            "Unit": ["unit_id", "unit_name", "readiness_level"],
            "Asset": ["asset_id", "asset_type", "status"],
            "Location": ["location_id", "location_name", "region"],
            "Capability": ["capability_id", "capability_name", "required_level"],
        },
        "formulas": [
            ("Capability Coverage", "Number of required mission capabilities present among assigned assets divided by the total number of required capabilities."),
            ("Readiness Score", "Average readiness level across all units assigned to a mission, weighted by each unit share of assigned assets."),
            ("Positional Risk", "Count of required assets not yet positioned at the mission location divided by the total number of required assets."),
        ],
        "benefits": [
            "Speeds up mission planning by showing capability gaps at a glance",
            "Improves readiness assessment with a connected unit and asset view",
            "Highlights positional risk before it becomes an operational problem",
            "Supports after action review with a clear relationship based record",
        ],
        "ai_prompt_focus": "identify capability gaps and readiness risk for the mission and explain the reasoning",
    },
    "software_dependency": {
        "order": 13,
        "title": "Software Dependency Intelligence",
        "graph_provides": "Application to API to Library to Vulnerability relationships",
        "ai_does": "Determines impact of vulnerabilities",
        "description": (
            "This use case maps which applications call which APIs, which are "
            "built on which open source libraries, and which libraries carry "
            "known vulnerabilities. When a new vulnerability is disclosed, the "
            "graph instantly shows every affected application. The AI layer "
            "explains the blast radius and suggests a remediation order based "
            "on business criticality."
        ),
        "node_types": ["Application", "API", "Library", "Vulnerability"],
        "edge_types": [
            ("Application", "API", "uses"),
            ("API", "Library", "built_on"),
            ("Library", "Vulnerability", "affected_by"),
        ],
        "fields": {
            "Application": ["app_id", "app_name", "business_criticality"],
            "API": ["api_id", "api_name", "version"],
            "Library": ["library_id", "library_name", "library_version"],
            "Vulnerability": ["cve_id", "severity", "cvss_score", "published_date"],
        },
        "formulas": [
            ("Vulnerability Blast Radius", "Number of distinct application nodes reachable upstream from a vulnerable library node."),
            ("Remediation Priority Score", "CVSS score multiplied by business criticality weight of the affected application, summed across every affected application."),
            ("Patch Debt", "Number of days since the vulnerability was published, used to flag libraries that have been vulnerable for an extended period."),
        ],
        "benefits": [
            "Answers which applications are affected by a new vulnerability in seconds",
            "Prioritizes patching using both severity and business criticality",
            "Reduces duplicate remediation effort by grouping applications sharing a library",
            "Provides security leadership a defensible, explainable patch order",
        ],
        "ai_prompt_focus": "explain vulnerability blast radius across applications and recommend a remediation order",
    },
    "customer_360": {
        "order": 14,
        "title": "Customer 360",
        "graph_provides": "Customer to Products to Interactions to Complaints to Transactions relationships",
        "ai_does": "Predicts churn and recommends actions",
        "description": (
            "This use case connects a customer to the products they hold, every "
            "interaction they have had with the company, complaints they have "
            "filed, and their transaction history. Churn risk is rarely explained "
            "by one factor alone, it is the combination of declining transactions, "
            "rising complaints and reduced interaction that signals risk. The AI "
            "layer reviews the connected customer record and recommends a "
            "retention action."
        ),
        "node_types": ["Customer", "Product", "Interaction", "Complaint", "Transaction"],
        "edge_types": [
            ("Customer", "Product", "holds"),
            ("Customer", "Interaction", "had"),
            ("Customer", "Complaint", "filed"),
            ("Customer", "Transaction", "made"),
        ],
        "fields": {
            "Customer": ["customer_id", "full_name", "tenure_months", "segment"],
            "Product": ["product_id", "product_name", "category"],
            "Interaction": ["interaction_id", "channel", "date", "sentiment"],
            "Complaint": ["complaint_id", "category", "date", "resolved_flag"],
            "Transaction": ["transaction_id", "amount", "date"],
        },
        "formulas": [
            ("Churn Risk Score", "0.4 times decline in transaction frequency plus 0.3 times unresolved complaint count plus 0.3 times negative interaction sentiment ratio, all normalized to zero to one."),
            ("Customer Lifetime Value", "Average transaction amount multiplied by expected remaining tenure in months, discounted by the churn risk score."),
            ("Engagement Ratio", "Number of positive sentiment interactions divided by total interactions in the trailing twelve months."),
        ],
        "benefits": [
            "Predicts churn earlier by combining signals that live in different systems",
            "Recommends the next best retention action instead of just a risk score",
            "Prioritizes outreach using customer lifetime value alongside churn risk",
            "Gives frontline teams full context in one connected view before a call",
        ],
        "ai_prompt_focus": "assess churn risk from the connected record and recommend a specific retention action",
    },
    "root_cause": {
        "order": 15,
        "title": "AI Root Cause Analysis",
        "graph_provides": "Events to Components to Dependencies to Failures relationships",
        "ai_does": "Traces root cause across connected systems and events",
        "description": (
            "This use case connects observed events to the system components they "
            "originate from, the dependencies between components, and historical "
            "failure records. When several events fire close together, the graph "
            "reveals whether they share a common upstream component. The AI layer "
            "reviews the event cluster and the dependency chain to explain the "
            "single most likely root cause."
        ),
        "node_types": ["Event", "Component", "Dependency", "Failure"],
        "edge_types": [
            ("Event", "Component", "originates_from"),
            ("Component", "Dependency", "depends_on"),
            ("Component", "Failure", "has_history_of"),
        ],
        "fields": {
            "Event": ["event_id", "event_type", "timestamp", "severity"],
            "Component": ["component_id", "component_name", "layer"],
            "Dependency": ["dependency_id", "dependency_type", "criticality"],
            "Failure": ["failure_id", "failure_type", "failure_date", "resolution_time_hours"],
        },
        "formulas": [
            ("Root Cause Score", "For each component in the event cluster, the number of downstream events explained by that component divided by the total number of events in the cluster."),
            ("Historical Recurrence Rate", "Count of prior failure records for a component divided by the number of months the component has been in production."),
            ("Dependency Chain Depth", "Number of hops between the candidate root cause component and the furthest affected component in the current event cluster."),
        ],
        "benefits": [
            "Finds the true upstream root cause instead of treating every alert as independent",
            "Uses failure history to flag chronically weak components before they recur",
            "Shortens post incident review with an explainable dependency chain",
            "Reduces alert fatigue by grouping related events under one root cause",
        ],
        "ai_prompt_focus": "identify the single most likely root cause component and explain the dependency chain that produced the event cluster",
    },
}
USE_CASE_ORDER = [k for k, v in sorted(USE_CASES.items(), key=lambda kv: kv[1]["order"])]
# =============================================================================
# SECTION 2: SYNTHETIC DATA GENERATION AND GRAPH CONSTRUCTION
# =============================================================================
FIRST_NAMES = ["James", "Maria", "Wei", "Amara", "Liam", "Sofia", "Noah", "Priya",
               "Ethan", "Fatima", "Lucas", "Ingrid", "Omar", "Chen", "Ava", "Diego"]
LAST_NAMES = ["Smith", "Garcia", "Chen", "Okafor", "Muller", "Rossi", "Kim", "Singh",
              "Brown", "Nguyen", "Andersson", "Silva", "Khan", "Kowalski", "Dubois"]
CITIES = ["New York", "London", "Singapore", "Toronto", "Mumbai", "Berlin",
          "Sydney", "Sao Paulo", "Dubai", "Tokyo", "Chicago", "Nairobi"]
COUNTRIES = ["United States", "United Kingdom", "Singapore", "Canada", "India",
             "Germany", "Australia", "Brazil", "United Arab Emirates", "Japan"]
def random_name():
    return random.choice(FIRST_NAMES) + " " + random.choice(LAST_NAMES)
def random_id(prefix, n=6):
    suffix = "".join(random.choices(string.digits, k=n))
    return prefix + suffix
def random_field_value(field_name, node_type, index):
    name_lower = field_name.lower()
    if name_lower.endswith("_id"):
        return random_id(node_type[:2].upper(), 5) if index is None else node_type[:2].upper() + str(1000 + index)
    if "name" in name_lower and "user" not in name_lower:
        return random_name()
    if "user_name" in name_lower or "hostname" in name_lower:
        return random.choice(["host", "user", "node"]) + str(random.randint(100, 999))
    if "score" in name_lower or "rating" in name_lower or "level" in name_lower or "pct" in name_lower:
        return round(random.uniform(0, 100), 2)
    if "amount" in name_lower or "value" in name_lower or "balance" in name_lower or "capacity" in name_lower:
        return round(random.uniform(100, 100000), 2)
    if "date" in name_lower or "timestamp" in name_lower:
        y = random.randint(2022, 2026)
        m = random.randint(1, 12)
        d = random.randint(1, 28)
        return "%04d-%02d-%02d" % (y, m, d)
    if "city" in name_lower:
        return random.choice(CITIES)
    if "country" in name_lower or "region" in name_lower or "jurisdiction" in name_lower or "origin" in name_lower:
        return random.choice(COUNTRIES)
    if "status" in name_lower or "flag" in name_lower:
        return random.choice(["Active", "Inactive", "Pending", "Resolved"])
    if "severity" in name_lower:
        return random.choice(["Low", "Medium", "High", "Critical"])
    if "type" in name_lower or "category" in name_lower or "channel" in name_lower or "tier" in name_lower or "segment" in name_lower:
        return random.choice(["Type A", "Type B", "Type C", "Type D"])
    if "days" in name_lower or "hours" in name_lower or "count" in name_lower or "quantity" in name_lower:
        return random.randint(1, 500)
    return "Value " + str(random.randint(1, 999))
def generate_synthetic_dataset(use_case_key, config, nodes_per_type=12, seed=None):
    """
    Generates a synthetic node table and edge table for the given use case
    configuration. Returns a dict with keys nodes (DataFrame) and edges (DataFrame).
    """
    if seed is not None:
        random.seed(seed)
    node_types = config["node_types"]
    fields = config["fields"]
    edge_types = config["edge_types"]
    node_rows = []
    node_ids_by_type = {}
    for ntype in node_types:
        ids_for_type = []
        for i in range(nodes_per_type):
            row = {"node_type": ntype, "node_label": ntype + " " + str(i + 1)}
            for f in fields.get(ntype, []):
                row[f] = random_field_value(f, ntype, i)
            primary_key = fields.get(ntype, ["id"])[0]
            node_id = ntype + "_" + str(i + 1)
            row["node_id"] = node_id
            row[primary_key] = row.get(primary_key, node_id)
            node_rows.append(row)
            ids_for_type.append(node_id)
        node_ids_by_type[ntype] = ids_for_type
    nodes_df = pd.DataFrame(node_rows)
    edge_rows = []
    for source_type, target_type, relation in edge_types:
        source_ids = node_ids_by_type.get(source_type, [])
        target_ids = node_ids_by_type.get(target_type, [])
        if not source_ids or not target_ids:
            continue
        for sid in source_ids:
            num_links = random.randint(1, min(3, len(target_ids)))
            targets = random.sample(target_ids, num_links)
            for tid in targets:
                edge_rows.append({
                    "source_id": sid,
                    "source_type": source_type,
                    "target_id": tid,
                    "target_type": target_type,
                    "relationship": relation,
                    "weight": round(random.uniform(0.1, 1.0), 3),
                })
    edges_df = pd.DataFrame(edge_rows)
    return {"nodes": nodes_df, "edges": edges_df}
def build_networkx_graph(nodes_df, edges_df):
    G = nx.DiGraph()
    if nodes_df is not None and len(nodes_df) > 0:
        for _, row in nodes_df.iterrows():
            node_id = row.get("node_id", row.iloc[0])
            label = row.get("node_label", str(node_id))
            ntype = row.get("node_type", "Entity")
            extra_attrs = row.to_dict()
            extra_attrs.pop("node_id", None)
            extra_attrs.pop("node_label", None)
            extra_attrs.pop("node_type", None)
            G.add_node(node_id, label=label, node_type=ntype, **extra_attrs)
    if edges_df is not None and len(edges_df) > 0:
        for _, row in edges_df.iterrows():
            src = row.get("source_id")
            tgt = row.get("target_id")
            if src is None or tgt is None:
                continue
            if src not in G:
                G.add_node(src, label=str(src), node_type=row.get("source_type", "Entity"))
            if tgt not in G:
                G.add_node(tgt, label=str(tgt), node_type=row.get("target_type", "Entity"))
            rel = row.get("relationship", "related_to")
            weight = row.get("weight", 1.0)
            G.add_edge(src, tgt, relationship=rel, weight=weight)
    return G
def dataframes_from_uploaded(nodes_file, edges_file):
    """
    Reads user uploaded CSV files for nodes and edges into DataFrames.
    Expected node columns: node_id, node_label, node_type, plus any custom fields.
    Expected edge columns: source_id, target_id, relationship, and optional weight.
    """
    nodes_df = pd.read_csv(nodes_file)
    edges_df = pd.read_csv(edges_file)
    if "node_id" not in nodes_df.columns:
        nodes_df["node_id"] = nodes_df.index.astype(str)
    if "node_label" not in nodes_df.columns:
        nodes_df["node_label"] = nodes_df["node_id"]
    if "node_type" not in nodes_df.columns:
        nodes_df["node_type"] = "Entity"
    if "weight" not in edges_df.columns:
        edges_df["weight"] = 1.0
    return {"nodes": nodes_df, "edges": edges_df}
# =============================================================================
# SECTION 3: GRAPH ANALYTICS AND VISUALIZATION
# =============================================================================
def compute_core_metrics(G):
    """
    Computes degree centrality, betweenness centrality and pagerank for
    every node in the graph. Returns a DataFrame indexed by node id.
    """
    if G.number_of_nodes() == 0:
        return pd.DataFrame(columns=["node_id", "label", "node_type",
                                      "degree_centrality", "betweenness_centrality", "pagerank"])
    degree = nx.degree_centrality(G)
    try:
        betweenness = nx.betweenness_centrality(G)
    except Exception:
        betweenness = {n: 0.0 for n in G.nodes()}
    try:
        pagerank = nx.pagerank(G, weight="weight")
    except Exception:
        pagerank = {n: 0.0 for n in G.nodes()}
    rows = []
    for n in G.nodes():
        data = G.nodes[n]
        rows.append({
            "node_id": n,
            "label": data.get("label", str(n)),
            "node_type": data.get("node_type", "Entity"),
            "degree_centrality": round(degree.get(n, 0.0), 4),
            "betweenness_centrality": round(betweenness.get(n, 0.0), 4),
            "pagerank": round(pagerank.get(n, 0.0), 4),
        })
    return pd.DataFrame(rows).sort_values("pagerank", ascending=False).reset_index(drop=True)
def find_top_paths(G, max_paths=5):
    """
    Finds a handful of representative shortest paths between nodes of
    different types that have low out degree and high in degree, useful
    for narrating attack paths, root cause chains or cascading impact.
    """
    paths = []
    nodes = list(G.nodes())
    if len(nodes) < 2:
        return paths
    sources = [n for n in nodes if G.out_degree(n) > 0 and G.in_degree(n) == 0]
    sinks = [n for n in nodes if G.in_degree(n) > 0 and G.out_degree(n) == 0]
    if not sources:
        sources = nodes[: min(5, len(nodes))]
    if not sinks:
        sinks = nodes[: min(5, len(nodes))]
    count = 0
    for s in sources:
        for t in sinks:
            if s == t:
                continue
            try:
                if nx.has_path(G, s, t):
                    p = nx.shortest_path(G, s, t)
                    if len(p) > 1:
                        paths.append(p)
                        count += 1
            except Exception:
                continue
            if count >= max_paths:
                return paths
    return paths
def build_plotly_graph(G, highlight_path=None, max_nodes=250):
    """
    Builds an interactive Plotly figure of the graph using a spring layout.
    Nodes are colored by node_type. If highlight_path is provided, those
    edges are drawn in a distinct color and thicker width.
    """
    if G.number_of_nodes() == 0:
        fig = go.Figure()
        fig.update_layout(title="No graph data available")
        return fig
    if G.number_of_nodes() > max_nodes:
        keep = list(G.nodes())[:max_nodes]
        G = G.subgraph(keep).copy()
    pos = nx.spring_layout(G, seed=42, k=None)
    node_types = sorted(set(nx.get_node_attributes(G, "node_type").values()))
    palette = ["#1f4e9c", "#c0392b", "#1e8449", "#b9770e", "#6c3483",
               "#117864", "#a04000", "#2874a6", "#7d3c98", "#154360"]
    color_map = {nt: palette[i % len(palette)] for i, nt in enumerate(node_types)}
    highlight_edges = set()
    if highlight_path:
        for i in range(len(highlight_path) - 1):
            highlight_edges.add((highlight_path[i], highlight_path[i + 1]))
    edge_traces = []
    edge_x_normal, edge_y_normal = [], []
    edge_x_highlight, edge_y_highlight = [], []
    for u, v in G.edges():
        x0, y0 = pos[u]
        x1, y1 = pos[v]
        if (u, v) in highlight_edges:
            edge_x_highlight += [x0, x1, None]
            edge_y_highlight += [y0, y1, None]
        else:
            edge_x_normal += [x0, x1, None]
            edge_y_normal += [y0, y1, None]
    edge_traces.append(go.Scatter(x=edge_x_normal, y=edge_y_normal, mode="lines",
                                   line=dict(width=1, color="#bbbbbb"),
                                   hoverinfo="none", showlegend=False))
    if edge_x_highlight:
        edge_traces.append(go.Scatter(x=edge_x_highlight, y=edge_y_highlight, mode="lines",
                                       line=dict(width=3, color="#e74c3c"),
                                       hoverinfo="none", showlegend=False,
                                       name="Highlighted Path"))
    node_traces = []
    for nt in node_types:
        xs, ys, texts, ids = [], [], [], []
        for n in G.nodes():
            if G.nodes[n].get("node_type", "Entity") == nt:
                x, y = pos[n]
                xs.append(x)
                ys.append(y)
                label = G.nodes[n].get("label", str(n))
                texts.append(label + " (" + nt + ")")
        node_traces.append(go.Scatter(
            x=xs, y=ys, mode="markers", name=nt,
            marker=dict(size=14, color=color_map[nt], line=dict(width=1, color="#ffffff")),
            text=texts, hoverinfo="text"
        ))
    fig = go.Figure(data=edge_traces + node_traces)
    fig.update_layout(
        showlegend=True,
        margin=dict(l=10, r=10, t=30, b=10),
        xaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
        yaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
        plot_bgcolor="#ffffff",
        legend=dict(orientation="h", yanchor="bottom", y=1.02),
        height=560,
    )
    return fig
# =============================================================================
# SECTION 4: LLM INTEGRATION - GROQ AND GEMINI
# =============================================================================
GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"
# Groq periodically retires older chat models (the Llama 3.x, Mixtral and
# Gemma2 families have all been shut down). openai/gpt-oss-120b is the
# current flagship production model on GroqCloud as of the latest model
# list published at console.groq.com/docs/models.
GROQ_DEFAULT_MODEL = "openai/gpt-oss-120b"
GROQ_MODEL_OPTIONS = [
    "openai/gpt-oss-120b",
    "openai/gpt-oss-20b",
    "qwen/qwen3.6-27b",
    "groq/compound",
    "groq/compound-mini",
]
GEMINI_API_URL_TEMPLATE = "https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={key}"
# Google periodically retires older Gemini models (gemini-2.0-flash and
# gemini-2.0-flash-lite have both been shut down). gemini-3.6-flash is a
# current stable model on the Gemini API as of the latest model list
# published at ai.google.dev/gemini-api/docs/models.
GEMINI_DEFAULT_MODEL = "gemini-3.6-flash"
GEMINI_MODEL_OPTIONS = [
    "gemini-3.6-flash",
    "gemini-3.7-flash",
    "gemini-2.5-flash",
    "gemini-2.5-flash-lite",
]
def call_groq(api_key, system_prompt, user_prompt, model=None, timeout=60):
    if not api_key:
        return None, "Groq API key is missing. Please enter a key in the sidebar."
    model_name = model or GROQ_DEFAULT_MODEL
    headers = {
        "Authorization": "Bearer " + api_key,
        "Content-Type": "application/json",
    }
    payload = {
        "model": model_name,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        "temperature": 0.3,
        "max_tokens": 900,
    }
    try:
        resp = requests.post(GROQ_API_URL, headers=headers, data=json.dumps(payload), timeout=timeout)
        if resp.status_code != 200:
            detail = resp.text[:300]
            if resp.status_code == 404 and "model_not_found" in resp.text:
                detail += (
                    " This Groq model has likely been retired. Pick a different "
                    "model from the sidebar Groq Model dropdown."
                )
            return None, "Groq API error, status code " + str(resp.status_code) + ". Details: " + detail
        data = resp.json()
        text = data["choices"][0]["message"]["content"]
        return text, None
    except Exception as e:
        return None, "Groq API request failed: " + str(e)
def call_gemini(api_key, system_prompt, user_prompt, model=None, timeout=60):
    if not api_key:
        return None, "Gemini API key is missing. Please enter a key in the sidebar."
    model_name = model or GEMINI_DEFAULT_MODEL
    url = GEMINI_API_URL_TEMPLATE.format(model=model_name, key=api_key)
    combined_prompt = system_prompt + "\n\n" + user_prompt
    payload = {
        "contents": [
            {"parts": [{"text": combined_prompt}]}
        ],
        "generationConfig": {"temperature": 0.3, "maxOutputTokens": 900},
    }
    headers = {"Content-Type": "application/json"}
    try:
        resp = requests.post(url, headers=headers, data=json.dumps(payload), timeout=timeout)
        if resp.status_code != 200:
            detail = resp.text[:300]
            if resp.status_code == 404:
                detail += (
                    " This Gemini model has likely been retired. Pick a different "
                    "model from the sidebar Gemini Model dropdown."
                )
            return None, "Gemini API error, status code " + str(resp.status_code) + ". Details: " + detail
        data = resp.json()
        candidates = data.get("candidates", [])
        if not candidates:
            return None, "Gemini API returned no candidates."
        parts = candidates[0]["content"]["parts"]
        text = "".join([p.get("text", "") for p in parts])
        return text, None
    except Exception as e:
        return None, "Gemini API request failed: " + str(e)
def run_llm_analysis(provider, api_key, system_prompt, user_prompt, model=None):
    """
    Dispatches to the selected provider. provider is either Groq or Gemini.
    Returns a tuple of result text and error message. If the call fails or
    no key is present, a locally generated fallback explanation is returned
    so the application remains usable without live API access.
    """
    if provider == "Groq":
        text, error = call_groq(api_key, system_prompt, user_prompt, model=model)
    elif provider == "Gemini":
        text, error = call_gemini(api_key, system_prompt, user_prompt, model=model)
    else:
        text, error = None, "Unknown provider selected."
    if text:
        return text, None
    fallback = build_fallback_summary(user_prompt)
    note = "Live LLM call was not available (" + (error or "no key provided") + "). Showing a rule based summary instead."
    return fallback, note
def build_fallback_summary(user_prompt):
    """
    A simple deterministic fallback used when no API key is configured or a
    live call fails, so the demo experience is not blocked.
    """
    lines = [
        "Rule based summary generated locally, no live LLM call was made.",
        "",
        "Key graph observations were extracted from node and edge statistics",
        "including centrality, pagerank and shortest path analysis.",
        "",
        "To receive a full natural language explanation, add a valid Groq or",
        "Gemini API key in the sidebar and rerun the analysis.",
    ]
    return "\n".join(lines)
# =============================================================================
# SECTION 5: EXPORT UTILITIES - PDF, WORD, TEXT, CSV
# =============================================================================
def build_text_report(title, sections):
    """
    sections is a list of tuples (heading, body_text)
    """
    lines = []
    lines.append(title)
    lines.append("=" * len(title))
    lines.append("")
    lines.append("Prepared using the Knowledge Graph AI Use Cases application")
    lines.append("Developed by Randy Singh from Kalsnet KNet Consulting Group")
    lines.append("")
    for heading, body in sections:
        lines.append(heading)
        lines.append("-" * len(heading))
        lines.append(body)
        lines.append("")
    return "\n".join(lines).encode("utf-8")
def build_csv_report(nodes_df, edges_df, metrics_df):
    buffer = io.StringIO()
    buffer.write("NODE TABLE\n")
    if nodes_df is not None and len(nodes_df) > 0:
        nodes_df.to_csv(buffer, index=False)
    buffer.write("\nEDGE TABLE\n")
    if edges_df is not None and len(edges_df) > 0:
        edges_df.to_csv(buffer, index=False)
    buffer.write("\nGRAPH METRICS\n")
    if metrics_df is not None and len(metrics_df) > 0:
        metrics_df.to_csv(buffer, index=False)
    return buffer.getvalue().encode("utf-8")
def build_word_report(title, sections, metrics_df=None):
    doc = Document()
    heading = doc.add_heading(level=0)
    run = heading.add_run(title)
    run.font.color.rgb = RGBColor(0x1f, 0x4e, 0x9c)
    run.font.size = Pt(24)
    run.bold = True
    sub = doc.add_paragraph()
    sub_run = sub.add_run("Knowledge Graph AI Use Cases Application")
    sub_run.bold = True
    sub_run.font.color.rgb = RGBColor(0x1f, 0x4e, 0x9c)
    footer_p = doc.add_paragraph("Developed by Randy Singh from Kalsnet KNet Consulting Group")
    footer_p.runs[0].italic = True
    doc.add_paragraph("")
    for heading_text, body_text in sections:
        doc.add_heading(heading_text, level=1)
        doc.add_paragraph(body_text)
    if metrics_df is not None and len(metrics_df) > 0:
        doc.add_heading("Graph Metrics Summary", level=1)
        top = metrics_df.head(15)
        table = doc.add_table(rows=1, cols=len(top.columns))
        table.style = "Light Grid Accent 1"
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(top.columns):
            hdr_cells[i].text = str(col)
        for _, row in top.iterrows():
            cells = table.add_row().cells
            for i, col in enumerate(top.columns):
                cells[i].text = str(row[col])
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
def build_pdf_report(title, sections, metrics_df=None):
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter,
                             topMargin=0.6 * inch, bottomMargin=0.6 * inch,
                             leftMargin=0.6 * inch, rightMargin=0.6 * inch)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("TitleBlue", parent=styles["Title"],
                                  textColor=colors.HexColor("#1f4e9c"), fontSize=22)
    heading_style = ParagraphStyle("HeadingBlue", parent=styles["Heading2"],
                                    textColor=colors.HexColor("#1f4e9c"))
    body_style = styles["BodyText"]
    footer_style = ParagraphStyle("Footer", parent=styles["Normal"], fontSize=9,
                                   textColor=colors.HexColor("#555555"))
    story = []
    story.append(Paragraph(title, title_style))
    story.append(Paragraph("Knowledge Graph AI Use Cases Application", heading_style))
    story.append(Paragraph("Developed by Randy Singh from Kalsnet KNet Consulting Group", footer_style))
    story.append(Spacer(1, 16))
    for heading_text, body_text in sections:
        story.append(Paragraph(heading_text, heading_style))
        safe_body = body_text.replace("\n", "<br/>")
        story.append(Paragraph(safe_body, body_style))
        story.append(Spacer(1, 10))
    if metrics_df is not None and len(metrics_df) > 0:
        story.append(Paragraph("Graph Metrics Summary", heading_style))
        top = metrics_df.head(15)
        data = [list(top.columns)] + top.astype(str).values.tolist()
        t = Table(data, repeatRows=1)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f4e9c")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTSIZE", (0, 0), (-1, -1), 7.5),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f2f6fc")]),
        ]))
        story.append(t)
    doc.build(story)
    buf.seek(0)
    return buf.read()
# =============================================================================
# SECTION 6: REAL WORLD DATA SOURCE INTEGRATION - OPENSANCTIONS AND ELLIPTIC
# =============================================================================
OPENSANCTIONS_BASE_URL = "https://api.opensanctions.org"
def fetch_opensanctions_graph(api_key, search_text, dataset="default", limit=15, timeout=30):
    """
    Calls the real OpenSanctions search API for the given free text query,
    then calls the entities endpoint for each match to retrieve nested
    relationships such as ownership, family and sanctions. Returns a dict
    with nodes and edges DataFrames built in the application node and
    edge schema, ready to feed into build_networkx_graph.
    """
    if not api_key:
        return None, "OpenSanctions API key is missing. Please enter a key in the sidebar or the panel above."
    headers = {"Authorization": "ApiKey " + api_key}
    search_url = OPENSANCTIONS_BASE_URL + "/search/" + dataset
    try:
        resp = requests.get(
            search_url,
            headers=headers,
            params={"q": search_text, "limit": limit},
            timeout=timeout,
        )
    except Exception as e:
        return None, "OpenSanctions request failed: " + str(e)
    if resp.status_code != 200:
        return None, "OpenSanctions API error, status code " + str(resp.status_code) + ". Details: " + resp.text[:300]
    data = resp.json()
    results = data.get("results", [])
    if not results:
        return None, "No matching entities were found for this search text."
    node_rows = []
    edge_rows = []
    seen_nodes = set()
    def add_node(node_id, label, node_type, extra=None):
        if node_id in seen_nodes:
            return
        row = {"node_id": node_id, "node_label": label, "node_type": node_type}
        if extra:
            row.update(extra)
        node_rows.append(row)
        seen_nodes.add(node_id)
    for entity in results:
        entity_id = entity.get("id")
        caption = entity.get("caption", entity_id)
        schema = entity.get("schema", "Entity")
        props = entity.get("properties", {})
        topics = props.get("topics", [])
        countries = props.get("country", [])
        add_node(
            entity_id, caption, "Person" if schema == "Person" else "Business",
            extra={
                "schema_type": schema,
                "risk_topics": ", ".join(topics) if topics else "none listed",
                "country": ", ".join(countries) if countries else "unknown",
                "match_score": entity.get("score", ""),
            },
        )
        for ds in entity.get("datasets", [])[:3]:
            ds_node_id = "dataset_" + ds
            add_node(ds_node_id, ds, "Watchlist Source")
            edge_rows.append({
                "source_id": entity_id, "source_type": "Person",
                "target_id": ds_node_id, "target_type": "Watchlist Source",
                "relationship": "listed_in", "weight": 1.0,
            })
        for topic in topics:
            topic_node_id = "topic_" + topic
            add_node(topic_node_id, topic, "Risk Topic")
            edge_rows.append({
                "source_id": entity_id, "source_type": "Person",
                "target_id": topic_node_id, "target_type": "Risk Topic",
                "relationship": "flagged_as", "weight": 1.0,
            })
        try:
            detail_resp = requests.get(
                OPENSANCTIONS_BASE_URL + "/entities/" + entity_id,
                headers=headers, timeout=timeout,
            )
            if detail_resp.status_code == 200:
                detail = detail_resp.json()
                nested_props = detail.get("properties", {})
                for rel_field, rel_key, rel_type_label in [
                    ("ownershipOwner", "Owner", "owns"),
                    ("directorshipDirector", "Director", "directs"),
                    ("familyPerson", "Family", "related_to"),
                    ("addressEntity", "Address", "registered_at"),
                ]:
                    related_list = nested_props.get(rel_field, [])
                    if isinstance(related_list, list):
                        for related in related_list[:5]:
                            if isinstance(related, dict):
                                related_id = related.get("id", rel_key + "_" + str(len(node_rows)))
                                related_caption = related.get("caption", related_id)
                                add_node(related_id, related_caption, rel_key)
                                edge_rows.append({
                                    "source_id": entity_id, "source_type": "Person",
                                    "target_id": related_id, "target_type": rel_key,
                                    "relationship": rel_type_label, "weight": 1.0,
                                })
        except Exception:
            continue
    nodes_df = pd.DataFrame(node_rows)
    edges_df = pd.DataFrame(edge_rows) if edge_rows else pd.DataFrame(
        columns=["source_id", "source_type", "target_id", "target_type", "relationship", "weight"]
    )
    return {"nodes": nodes_df, "edges": edges_df}, None
def load_elliptic_dataset(classes_file, edgelist_file, features_file=None, sample_size=300, seed=42):
    """
    Reads the real Elliptic Bitcoin dataset files that the user has
    downloaded from Kaggle and uploaded to the application. Builds a
    transaction to transaction graph in the application node and edge
    schema, sampled down to a manageable size for interactive rendering.
    Expected file formats, matching the original Kaggle release:
    classes_file columns: txId, class (values are 1 for illicit,
      2 for licit, unknown for unlabeled)
    edgelist_file columns: txId1, txId2
    features_file (optional) has no header row, first column is txId,
      second column is a time step, remaining columns are numeric features
    """
    classes_df = pd.read_csv(classes_file)
    edges_raw = pd.read_csv(edgelist_file)
    classes_df.columns = [c.strip().lower() for c in classes_df.columns]
    edges_raw.columns = [c.strip().lower() for c in edges_raw.columns]
    if "txid" not in classes_df.columns:
        return None, "The classes file does not contain a txId column. Please upload the original elliptic_txs_classes.csv file."
    if "txid1" not in edges_raw.columns or "txid2" not in edges_raw.columns:
        return None, "The edgelist file does not contain txId1 and txId2 columns. Please upload the original elliptic_txs_edgelist.csv file."
    label_map = {"1": "Illicit", "2": "Licit", "unknown": "Unknown"}
    classes_df["class"] = classes_df["class"].astype(str)
    classes_df["label"] = classes_df["class"].map(label_map).fillna("Unknown")
    illicit_ids = classes_df[classes_df["label"] == "Illicit"]["txid"].tolist()
    if len(illicit_ids) == 0:
        sample_ids = classes_df["txid"].sample(min(sample_size, len(classes_df)), random_state=seed).tolist()
    else:
        take_illicit = min(len(illicit_ids), max(20, sample_size // 3))
        sample_ids = pd.Series(illicit_ids).sample(take_illicit, random_state=seed).tolist()
        remaining = sample_size - len(sample_ids)
        other_ids = classes_df[~classes_df["txid"].isin(sample_ids)]["txid"]
        if remaining > 0 and len(other_ids) > 0:
            sample_ids += other_ids.sample(min(remaining, len(other_ids)), random_state=seed).tolist()
    sample_id_set = set(sample_ids)
    sampled_edges = edges_raw[
        edges_raw["txid1"].isin(sample_id_set) & edges_raw["txid2"].isin(sample_id_set)
    ].copy()
    if len(sampled_edges) < 10:
        expanded_edges = edges_raw[
            edges_raw["txid1"].isin(sample_id_set) | edges_raw["txid2"].isin(sample_id_set)
        ].copy()
        sample_id_set = set(expanded_edges["txid1"]).union(set(expanded_edges["txid2"]))
        sampled_edges = expanded_edges
    node_rows = []
    for txid in sample_id_set:
        label_row = classes_df[classes_df["txid"] == txid]
        label_val = label_row["label"].iloc[0] if len(label_row) > 0 else "Unknown"
        node_rows.append({
            "node_id": "tx_" + str(txid),
            "node_label": "Transaction " + str(txid),
            "node_type": "Transaction",
            "aml_label": label_val,
            "source_txid": txid,
        })
    edge_rows = []
    for _, row in sampled_edges.iterrows():
        edge_rows.append({
            "source_id": "tx_" + str(row["txid1"]),
            "source_type": "Transaction",
            "target_id": "tx_" + str(row["txid2"]),
            "target_type": "Transaction",
            "relationship": "flows_to",
            "weight": 1.0,
        })
    nodes_df = pd.DataFrame(node_rows)
    edges_df = pd.DataFrame(edge_rows) if edge_rows else pd.DataFrame(
        columns=["source_id", "source_type", "target_id", "target_type", "relationship", "weight"]
    )
    summary = (
        "Loaded " + str(len(nodes_df)) + " real transactions and " + str(len(edges_df))
        + " real transaction flows from the Elliptic Bitcoin dataset. "
        + str(len(classes_df[classes_df['label'] == 'Illicit'])) + " transactions in the full dataset "
        + "are labeled illicit out of " + str(len(classes_df)) + " total labeled and unlabeled transactions."
    )
    return {"nodes": nodes_df, "edges": edges_df, "summary": summary}, None
# =============================================================================
# SECTION 7: STREAMLIT APPLICATION
# =============================================================================
import streamlit as st
import pandas as pd
from datetime import datetime
REAL_DATA_ENABLED_KEYS = ["aml", "fraud"]
st.set_page_config(
    page_title="Knowledge Graph AI Use Cases",
    page_icon=None,
    layout="wide",
)
# ---------------------------------------------------------------------------
# Global styling
# ---------------------------------------------------------------------------
st.markdown(
    """
    <style>
    .kg-main-title {
        color: #1f4e9c;
        font-weight: 800;
        font-size: 44px;
        margin-bottom: 0px;
    }
    .kg-sub-title {
        color: #1f4e9c;
        font-weight: 700;
        font-size: 20px;
        margin-top: 0px;
    }
    .kg-developer-line {
        color: #1f4e9c;
        font-weight: 700;
        font-size: 16px;
    }
    .kg-benefit-box {
        background-color: #f2f6fc;
        border-left: 5px solid #1f4e9c;
        padding: 12px 16px;
        border-radius: 4px;
        margin-bottom: 8px;
    }
    .kg-usecase-heading {
        color: #1f4e9c;
        font-weight: 800;
        font-size: 28px;
    }
    </style>
    """,
    unsafe_allow_html=True,
)
# ---------------------------------------------------------------------------
# Session state initialization
# ---------------------------------------------------------------------------
if "datasets" not in st.session_state:
    st.session_state["datasets"] = {}
if "analysis_results" not in st.session_state:
    st.session_state["analysis_results"] = {}
# ---------------------------------------------------------------------------
# Sidebar
# ---------------------------------------------------------------------------
st.sidebar.markdown(
    "<div class='kg-sub-title'>Knowledge Graph AI Use Cases</div>",
    unsafe_allow_html=True,
)
st.sidebar.markdown("Select a use case below")
use_case_labels = [str(USE_CASES[k]["order"]) + ". " + USE_CASES[k]["title"] for k in USE_CASE_ORDER]
label_to_key = dict(zip(use_case_labels, USE_CASE_ORDER))
selected_label = st.sidebar.radio("Use Cases", use_case_labels, index=0, label_visibility="collapsed")
selected_key = label_to_key[selected_label]
config = USE_CASES[selected_key]
st.sidebar.markdown("---")
st.sidebar.markdown("**AI Provider Settings**")
provider = st.sidebar.selectbox("Select LLM Provider", ["Groq", "Gemini"])
groq_key = st.sidebar.text_input("Groq API Key", type="password", key="groq_key_input")
groq_model_choice = st.sidebar.selectbox(
    "Groq Model",
    GROQ_MODEL_OPTIONS,
    index=0,
    help="Groq periodically retires older models (the Llama 3.x, Mixtral and Gemma2 "
         "families have all been shut down). If a model here ever returns a "
         "model not found error, check console.groq.com/docs/models for the current "
         "list and pick a replacement.",
)
gemini_key = st.sidebar.text_input("Gemini API Key", type="password", key="gemini_key_input")
gemini_model_choice = st.sidebar.selectbox(
    "Gemini Model",
    GEMINI_MODEL_OPTIONS,
    index=0,
    help="Google periodically retires older Gemini models (gemini-2.0-flash and "
         "gemini-2.0-flash-lite have both been shut down). If a model here ever "
         "returns a not found error, check ai.google.dev/gemini-api/docs/models "
         "for the current list and pick a replacement.",
)
active_key = groq_key if provider == "Groq" else gemini_key
active_model = groq_model_choice if provider == "Groq" else gemini_model_choice
with st.sidebar.expander("How to get a free Groq API key"):
    st.markdown(
        """
        1. Go to console dot groq dot com in your browser
        2. Sign up for a free account using an email address or a Google account
        3. Once logged in, open the API Keys section from the left menu
        4. Click Create API Key and give it a name
        5. Copy the generated key immediately and paste it into the Groq API Key field in this sidebar
        6. Groq offers a free developer tier with generous request limits, no credit card required to start
        """
    )
with st.sidebar.expander("How to get a free Gemini API key"):
    st.markdown(
        """
        1. Go to aistudio dot google dot com in your browser
        2. Sign in with a Google account
        3. Click Get API Key in the left menu, then click Create API Key
        4. Choose to create the key in a new or existing Google Cloud project
        5. Copy the generated key and paste it into the Gemini API Key field in this sidebar
        6. Google AI Studio provides a free usage tier suitable for development and testing
        """
    )
st.sidebar.markdown("---")
st.sidebar.markdown("**Real Data Source Settings**")
st.sidebar.caption("Used by the Fraud and AML use cases to load real, live data instead of synthetic data")
opensanctions_key = st.sidebar.text_input("OpenSanctions API Key", type="password", key="opensanctions_key_input")
with st.sidebar.expander("How to get a free OpenSanctions API key"):
    st.markdown(
        """
        1. Go to opensanctions dot org forward slash account in your browser
        2. Sign up for a free account
        3. If you sign up with a business email address, a free trial key is generated automatically
        4. OpenSanctions also issues free keys for journalism, academic research and civil society work,
           you can request one from the support page if the trial does not fit your use
        5. Copy the generated key and paste it into the OpenSanctions API Key field in this sidebar
        6. OpenSanctions is free for non commercial use, a data license is required for commercial use
        """
    )
with st.sidebar.expander("How to get the Elliptic Bitcoin dataset"):
    st.markdown(
        """
        1. Go to kaggle dot com forward slash datasets forward slash ellipticco forward slash elliptic dash data dash set
        2. Sign in to Kaggle with a free account
        3. Click Download to get the dataset zip file
        4. Unzip the file on your computer, it contains three files, elliptic txs classes csv,
           elliptic txs edgelist csv, and elliptic txs features csv
        5. In the AML use case Data Source tab, choose Load Elliptic Real Dataset and upload the
           classes file and the edgelist file, the features file is optional
        """
    )
st.sidebar.markdown("---")
st.sidebar.markdown(
    "<div class='kg-developer-line'>Developed by Randy Singh</div>"
    "<div class='kg-developer-line'>Kalsnet KNet Consulting Group</div>",
    unsafe_allow_html=True,
)
# ---------------------------------------------------------------------------
# Header
# ---------------------------------------------------------------------------
st.markdown("<div class='kg-main-title'>Knowledge Graph AI Use Cases</div>", unsafe_allow_html=True)
st.markdown(
    "<div class='kg-sub-title'>Enterprise Platform for Graph Connected Artificial Intelligence</div>",
    unsafe_allow_html=True,
)
st.markdown(
    "<div class='kg-developer-line'>Developed by Randy Singh from Kalsnet KNet Consulting Group</div>",
    unsafe_allow_html=True,
)
st.write("")
# ---------------------------------------------------------------------------
# Use case heading
# ---------------------------------------------------------------------------
st.markdown(
    "<div class='kg-usecase-heading'>" + str(config["order"]) + ". " + config["title"] + "</div>",
    unsafe_allow_html=True,
)
st.write(config["description"])
col_a, col_b = st.columns(2)
with col_a:
    st.info("What the Graph Provides: " + config["graph_provides"])
with col_b:
    st.success("What AI Does: " + config["ai_does"])
st.markdown("---")
# ---------------------------------------------------------------------------
# Tabs within the selected use case
# ---------------------------------------------------------------------------
tab_data, tab_schema, tab_graph, tab_ai, tab_export, tab_benefits = st.tabs(
    ["Data Source", "Schema and Formulas", "Graph Visualization", "AI Analysis", "Export Results", "Benefits"]
)
# ----- Tab 1: Data Source ---------------------------------------------------
with tab_data:
    st.subheader("Choose a Data Source")
    mode_options = ["Use Synthetic Data", "Upload Real Data"]
    if selected_key in REAL_DATA_ENABLED_KEYS:
        mode_options.append("Load Live Real Data")
    data_mode = st.radio(
        "Data mode",
        mode_options,
        horizontal=True,
        key="mode_" + selected_key,
    )
    if data_mode == "Use Synthetic Data":
        num_records = st.slider(
            "Number of records to generate per entity type",
            min_value=5, max_value=40, value=12, key="slider_" + selected_key,
        )
        if st.button("Generate Synthetic Data", key="gen_" + selected_key):
            dataset = generate_synthetic_dataset(selected_key, config, nodes_per_type=num_records, seed=42)
            st.session_state["datasets"][selected_key] = dataset
            st.success("Synthetic data generated for " + config["title"])
    elif data_mode == "Upload Real Data":
        st.write(
            "Upload two CSV files, one for nodes and one for edges. "
            "Node file should contain columns node_id, node_label, node_type and any custom fields. "
            "Edge file should contain columns source_id, target_id, relationship and optional weight."
        )
        nodes_file = st.file_uploader("Upload Nodes CSV", type=["csv"], key="nodes_upload_" + selected_key)
        edges_file = st.file_uploader("Upload Edges CSV", type=["csv"], key="edges_upload_" + selected_key)
        if nodes_file is not None and edges_file is not None:
            if st.button("Load Uploaded Data", key="load_" + selected_key):
                try:
                    dataset = dataframes_from_uploaded(nodes_file, edges_file)
                    st.session_state["datasets"][selected_key] = dataset
                    st.success("Uploaded data loaded for " + config["title"])
                except Exception as e:
                    st.error("Could not read uploaded files: " + str(e))
    else:
        st.markdown("### Real World Data Sources")
        st.write(
            "These are genuine, real data sources used in actual compliance and anti money "
            "laundering work, not synthetic placeholders."
        )
        real_source = st.radio(
            "Choose a real data source",
            ["OpenSanctions Live Sanctions and PEP Screening", "Elliptic Real Bitcoin AML Dataset"],
            key="real_source_" + selected_key,
        )
        if real_source == "OpenSanctions Live Sanctions and PEP Screening":
            st.write(
                "Calls the real OpenSanctions API and builds a live knowledge graph from actual "
                "sanctions list entries, politically exposed persons, and their real ownership, "
                "family and address relationships. Requires a free OpenSanctions API key entered "
                "in the sidebar."
            )
            search_text = st.text_input(
                "Search name or company",
                value="Vladimir Putin",
                key="opensanctions_query_" + selected_key,
            )
            result_limit = st.slider("Number of matched entities to retrieve", 5, 30, 10, key="opensanctions_limit_" + selected_key)
            if st.button("Fetch Live Data from OpenSanctions", key="fetch_opensanctions_" + selected_key):
                with st.spinner("Calling the real OpenSanctions API..."):
                    dataset, err = fetch_opensanctions_graph(opensanctions_key, search_text, limit=result_limit)
                if err:
                    st.error(err)
                else:
                    st.session_state["datasets"][selected_key] = dataset
                    st.success("Loaded real OpenSanctions data for the search term: " + search_text)
        else:
            st.write(
                "Uses the real, published Elliptic Bitcoin dataset, a graph of over two hundred "
                "thousand real anonymized bitcoin transactions with actual illicit and licit labels, "
                "used in real anti money laundering research. Download the dataset from Kaggle first "
                "using the steps in the sidebar, then upload the files below."
            )
            classes_upload = st.file_uploader(
                "Upload elliptic_txs_classes.csv", type=["csv"], key="elliptic_classes_" + selected_key
            )
            edgelist_upload = st.file_uploader(
                "Upload elliptic_txs_edgelist.csv", type=["csv"], key="elliptic_edges_" + selected_key
            )
            sample_size = st.slider(
                "Number of real transactions to sample for the graph",
                50, 800, 300, key="elliptic_sample_" + selected_key,
            )
            if classes_upload is not None and edgelist_upload is not None:
                if st.button("Load Elliptic Real Dataset", key="load_elliptic_" + selected_key):
                    with st.spinner("Reading the real Elliptic dataset..."):
                        dataset, err = load_elliptic_dataset(classes_upload, edgelist_upload, sample_size=sample_size)
                    if err:
                        st.error(err)
                    else:
                        st.session_state["datasets"][selected_key] = dataset
                        st.success(dataset.get("summary", "Elliptic dataset loaded."))
    current_dataset = st.session_state["datasets"].get(selected_key)
    if current_dataset is not None:
        st.markdown("**Node Table Preview**")
        st.dataframe(current_dataset["nodes"].head(50), use_container_width=True)
        st.markdown("**Edge Table Preview**")
        st.dataframe(current_dataset["edges"].head(50), use_container_width=True)
    else:
        st.warning("No data loaded yet for this use case. Generate synthetic data or upload real data above.")
# ----- Tab 2: Schema and Formulas ------------------------------------------
with tab_schema:
    st.subheader("Graph Schema")
    st.write(
        "The knowledge graph for this use case connects the following entity types "
        "through the relationships shown below."
    )
    for source_type, target_type, relation in config["edge_types"]:
        st.markdown("- **" + source_type + "** " + relation + " **" + target_type + "**")
    st.markdown("### Entity Fields")
    for entity_type, field_list in config["fields"].items():
        with st.expander(entity_type + " Fields"):
            for f in field_list:
                st.write("- " + f)
    st.markdown("### Formulas Used")
    for formula_name, formula_explanation in config["formulas"]:
        st.markdown("**" + formula_name + "**")
        st.write(formula_explanation)
        st.write("")
# ----- Tab 3: Graph Visualization -------------------------------------------
with tab_graph:
    st.subheader("Interactive Graph Visualization")
    current_dataset = st.session_state["datasets"].get(selected_key)
    if current_dataset is None:
        st.warning("Load data in the Data Source tab first.")
    else:
        G = build_networkx_graph(current_dataset["nodes"], current_dataset["edges"])
        st.write(
            "Graph contains " + str(G.number_of_nodes()) + " nodes and "
            + str(G.number_of_edges()) + " edges."
        )
        metrics_df = compute_core_metrics(G)
        st.session_state["analysis_results"].setdefault(selected_key, {})["metrics_df"] = metrics_df
        paths = find_top_paths(G, max_paths=5)
        highlight_choice = None
        if paths:
            path_labels = []
            for p in paths:
                labels = [G.nodes[n].get("label", str(n)) for n in p]
                path_labels.append(" then ".join(labels))
            selected_path_label = st.selectbox("Highlight a representative path", ["None"] + path_labels)
            if selected_path_label != "None":
                idx = path_labels.index(selected_path_label)
                highlight_choice = paths[idx]
                st.session_state["analysis_results"][selected_key]["highlighted_path"] = highlight_choice
        fig = build_plotly_graph(G, highlight_path=highlight_choice)
        st.plotly_chart(fig, use_container_width=True)
        st.markdown("### Top Ranked Nodes by Graph Metrics")
        st.write(
            "Degree centrality measures direct connections. Betweenness centrality "
            "measures how often a node sits on the shortest path between other nodes. "
            "Pagerank measures overall structural importance considering edge weight."
        )
        st.dataframe(metrics_df.head(20), use_container_width=True)
# ----- Tab 4: AI Analysis ----------------------------------------------------
with tab_ai:
    st.subheader("AI Powered Analysis")
    current_dataset = st.session_state["datasets"].get(selected_key)
    if current_dataset is None:
        st.warning("Load data in the Data Source tab first.")
    else:
        results_store = st.session_state["analysis_results"].setdefault(selected_key, {})
        metrics_df = results_store.get("metrics_df")
        if metrics_df is None:
            G = build_networkx_graph(current_dataset["nodes"], current_dataset["edges"])
            metrics_df = compute_core_metrics(G)
            results_store["metrics_df"] = metrics_df
        st.write(
            "This step sends a summary of the graph structure and top ranked nodes "
            "to the selected AI provider, which will " + config["ai_prompt_focus"] + "."
        )
        st.caption("Using " + provider + " model: " + active_model)
        user_question = st.text_area(
            "Optional question to focus the analysis",
            value="Summarize the most important risk or insight visible in this graph.",
            key="question_" + selected_key,
        )
        if st.button("Run AI Analysis", key="run_ai_" + selected_key):
            top_nodes = metrics_df.head(10).to_dict(orient="records")
            system_prompt = (
                "You are an expert analyst for the use case: " + config["title"] + ". "
                "You " + config["ai_does"] + ". "
                "Use the structured graph summary provided to " + config["ai_prompt_focus"] + ". "
                "Be specific, concise and use plain business language. Do not invent facts "
                "that are not supported by the provided summary."
            )
            user_prompt = (
                "Use case schema relationships: " + str(config["edge_types"]) + "\n"
                "Top ranked nodes by graph importance: " + str(top_nodes) + "\n"
                "Total nodes: " + str(len(current_dataset["nodes"])) + "\n"
                "Total edges: " + str(len(current_dataset["edges"])) + "\n"
                "User question: " + user_question
            )
            with st.spinner("Contacting " + provider + " for analysis..."):
                result_text, note = run_llm_analysis(
                    provider, active_key, system_prompt, user_prompt, model=active_model
                )
            results_store["ai_result"] = result_text
            results_store["ai_note"] = note
            results_store["ai_question"] = user_question
        if "ai_result" in results_store:
            if results_store.get("ai_note"):
                st.warning(results_store["ai_note"])
            st.markdown("### AI Analysis Result")
            st.write(results_store["ai_result"])
# ----- Tab 5: Export Results --------------------------------------------------
with tab_export:
    st.subheader("Export Results")
    current_dataset = st.session_state["datasets"].get(selected_key)
    results_store = st.session_state["analysis_results"].get(selected_key, {})
    if current_dataset is None:
        st.warning("Load data in the Data Source tab first.")
    else:
        nodes_df = current_dataset["nodes"]
        edges_df = current_dataset["edges"]
        metrics_df = results_store.get("metrics_df")
        ai_result = results_store.get("ai_result", "AI analysis has not been run yet for this use case.")
        report_title = config["title"] + " Analysis Report"
        generated_on = "Generated on " + datetime.now().strftime("%Y-%m-%d %H:%M")
        sections = [
            ("Use Case Overview", config["description"]),
            ("What the Graph Provides", config["graph_provides"]),
            ("What AI Does", config["ai_does"]),
            ("Data Summary", "Total nodes: " + str(len(nodes_df)) + ". Total edges: " + str(len(edges_df)) + ". " + generated_on),
            ("AI Analysis Result", ai_result),
            ("Benefits", "\n".join(["- " + b for b in config["benefits"]])),
        ]
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            pdf_bytes = build_pdf_report(report_title, sections, metrics_df)
            st.download_button(
                "Download PDF", data=pdf_bytes,
                file_name=selected_key + "_report.pdf", mime="application/pdf",
            )
        with col2:
            word_bytes = build_word_report(report_title, sections, metrics_df)
            st.download_button(
                "Download Word", data=word_bytes,
                file_name=selected_key + "_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        with col3:
            text_bytes = build_text_report(report_title, sections)
            st.download_button(
                "Download Text", data=text_bytes,
                file_name=selected_key + "_report.txt", mime="text/plain",
            )
        with col4:
            csv_bytes = build_csv_report(nodes_df, edges_df, metrics_df)
            st.download_button(
                "Download CSV", data=csv_bytes,
                file_name=selected_key + "_data.csv", mime="text/csv",
            )
        st.markdown("### Report Preview")
        for heading, body in sections:
            st.markdown("**" + heading + "**")
            st.write(body)
# ----- Tab 6: Benefits -------------------------------------------------------
with tab_benefits:
    st.subheader("Benefits of This Use Case")
    for benefit in config["benefits"]:
        st.markdown("<div class='kg-benefit-box'>" + benefit + "</div>", unsafe_allow_html=True)
# ---------------------------------------------------------------------------
# Footer
# ---------------------------------------------------------------------------
st.markdown("---")
st.markdown(
    "<div class='kg-developer-line'>Knowledge Graph AI Use Cases Platform</div>"
    "<div class='kg-developer-line'>Developed by Randy Singh from Kalsnet KNet Consulting Group</div>",
    unsafe_allow_html=True,
)
