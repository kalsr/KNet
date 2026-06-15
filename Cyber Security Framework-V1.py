


#=============================================================================
# CYBERSECURITY FRAMEWORK APPLICATION FOR DoD CRITICAL INFRASTRUCTURE
# @Developed by Randy Singh | KalSnet (KNet) Consulting Group
# Original Java code: September 2017 — Python/Streamlit conversion: 2024
#=============================================================================
# Run with:
#   pip install streamlit plotly pandas python-docx
#   (optional) pip install fpdf2
#   streamlit run cybersecurity_framework.py
#=============================================================================

import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import random
import json
import io
from datetime import datetime, timedelta

# ─── Optional export dependencies ─────────────────────────────────────────────
FPDF_AVAILABLE = False
try:
    from fpdf import FPDF
    FPDF_AVAILABLE = True
except ImportError:
    FPDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    from docx.shared import RGBColor, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ══════════════════════════════════════════════════════════════════════════════
#  PAGE CONFIG
# ══════════════════════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="CyberSecurity Framework | KNet Consulting",
    page_icon="🛡️",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ══════════════════════════════════════════════════════════════════════════════
#  GLOBAL STYLES
# ══════════════════════════════════════════════════════════════════════════════
st.markdown(
    """
<style>
  body {
      background-color:#f4f7fc;
      font-family:"Segoe UI",system-ui,-apple-system,BlinkMacSystemFont,sans-serif;
  }
  .cs-title {
      text-align:center; font-size:2.6rem; font-weight:900;
      color:#0047AB; letter-spacing:1px; margin-bottom:0;
  }
  .cs-subtitle {
      text-align:center; font-size:1.1rem; font-weight:700;
      color:#0047AB; margin-top:2px; margin-bottom:2px;
  }
  .cs-tagline {
      text-align:center; color:#555; font-size:0.9rem; margin-bottom:16px;
  }
  .cs-hr { border:2px solid #0047AB; margin:8px 0 18px 0; }
  .section-header {
      background:linear-gradient(90deg,#0047AB 0%,#1a6fe8 100%);
      color:white; padding:10px 18px; border-radius:8px;
      font-weight:800; font-size:1.1rem; margin:12px 0 10px 0;
      box-shadow:0 2px 6px rgba(0,0,0,0.18);
  }
  .req-card {
      background:#f0f6ff; border-left:4px solid #0047AB;
      border-radius:6px; padding:9px 13px; margin:5px 0; font-size:0.92rem;
  }
  .rec-card {
      background:#fff8e1; border-left:4px solid #f0a500;
      border-radius:6px; padding:9px 13px; margin:5px 0; font-size:0.92rem;
  }
  .sol-card {
      background:#e8f5e9; border-left:4px solid #2e7d32;
      border-radius:6px; padding:9px 13px; margin:5px 0; font-size:0.92rem;
  }
  .ref-card {
      background:#fce4ec; border-left:4px solid #c62828;
      border-radius:6px; padding:9px 13px; margin:5px 0; font-size:0.85rem; font-family:monospace;
  }
  div[data-testid="metric-container"] {
      background:#f0f6ff; border:1px solid #c2d9ff; border-radius:8px; padding:10px;
  }
  section[data-testid="stSidebar"] {
      background:#001a6b;
      overflow-y:auto;
  }
  section[data-testid="stSidebar"] * { color:#e8f0ff !important; }
  .sidebar-box {
      background:rgba(255,255,255,0.06);
      padding:10px 12px;
      border-radius:8px;
      margin-bottom:12px;
      border:1px solid rgba(255,255,255,0.12);
  }
  .sidebar-footer {
      font-size:0.8rem;
      color:#cbd4ff;
      text-align:center;
      margin-top:10px;
  }
</style>
""",
    unsafe_allow_html=True,
)

# ══════════════════════════════════════════════════════════════════════════════
#  HEADER
# ══════════════════════════════════════════════════════════════════════════════
st.markdown(
    '<p class="cs-title">CYBERSECURITY FRAMEWORK APPLICATION</p>',
    unsafe_allow_html=True,
)
st.markdown(
    '<p class="cs-subtitle">FOR US CRITICAL INFRASTRUCTURE</p>',
    unsafe_allow_html=True,
)
st.markdown(
    '<p class="cs-subtitle">Developed by Randy Singh | KalSnet (KNet) Consulting Group</p>',
    unsafe_allow_html=True,
)
st.markdown(
    '<p class="cs-tagline">Helping Security Analysts, Planners & Administrators address IoT Security Vulnerabilities in DISA, DoD & USA Critical Infrastructure &nbsp;|&nbsp; Originally authored June 2026</p>',
    unsafe_allow_html=True,
)
st.markdown('<hr class="cs-hr">', unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
#  FRAMEWORK DATA  (first 12 checklist items shown; extend as needed)
# ══════════════════════════════════════════════════════════════════════════════
FRAMEWORK_ITEMS = {
    1: {
        "title": "Identify - Asset Management",
        "function": "IDENTIFY",
        "summary": "Data, personnel, devices, systems, and facilities that enable the organization to achieve business purposes are identified and managed consistent with their relative importance to business objectives and the organization's risk strategy.",
        "subcategories": [
            "Physical devices and systems within the organization are inventoried.",
            "Software platforms and applications within the organization are inventoried.",
            "Organizational communication and data flows are mapped.",
            "External information systems are catalogued.",
            "Resources are prioritized based on classification, criticality, and business value.",
            "Cybersecurity roles and responsibilities for the workforce and third parties are established.",
        ],
        "example": "Use a CMDB like ServiceNow to auto-discover and track all hardware/software assets. Tag assets by classification (CUI, FOUO, UNCLASS).",
        "solution": "Deploy an automated asset discovery tool integrated with SIEM. Establish a quarterly asset reconciliation process.",
        "recommendation": "Require all new devices to be registered in the CMDB before network access is granted. Review asset inventory monthly.",
        "references": [
            "CCS CSC 1, COBIT 5 BAI09.01/02, ISA 62443-2-1:2009 4.2.3.4, NIST SP 800-53 Rev.4 CM-8",
        ],
    },
    2: {
        "title": "Identify - Business Environment",
        "function": "IDENTIFY",
        "summary": "The organization's mission, objectives, stakeholders, and activities are understood and prioritized. The organization's role in the supply chain is identified and communicated.",
        "subcategories": [
            "The organization's role in the supply chain is identified and communicated.",
            "The organization's place in critical infrastructure and its industry sector is identified.",
            "Priorities for mission, objectives, and activities are established and communicated.",
            "Dependencies and critical functions for delivery of critical services are established.",
            "Resilience requirements to support delivery of critical services are established.",
        ],
        "example": "Map the agency's position in the DoD supply chain using a Business Impact Analysis. Identify mission-critical systems.",
        "solution": "Conduct an annual Business Environment Review using CISA's Critical Infrastructure Framework.",
        "recommendation": "Align mission priorities to CISA's 16 Critical Infrastructure Sectors. Document third-party dependencies.",
        "references": [
            "COBIT 5 APO08.04/05/APO10.03-05, ISO/IEC 27001:2013 A.15.1.3/A.15.2.1-2, NIST SP 800-53 CP-2/SA-12",
        ],
    },
    3: {
        "title": "Identify - Governance Policies & Procedures",
        "function": "IDENTIFY",
        "summary": "Policies, procedures, and processes to manage and monitor regulatory, legal, risk, environmental, and operational requirements are understood and inform cybersecurity risk management.",
        "subcategories": [
            "Information security policy is established.",
            "Roles and responsibilities are coordinated and aligned with internal and external partners.",
            "Legal and regulatory requirements regarding cybersecurity are understood and managed.",
            "Governance and risk management processes address cybersecurity risks.",
        ],
        "example": "Publish a DoD-aligned Information Security Policy covering AUP, data classification, incident reporting, and third-party access.",
        "solution": "Establish a Cybersecurity Governance Board and use GRC tools to track policy compliance.",
        "recommendation": "Align policies to FISMA, DFARS, and CMMC. Conduct policy gap assessments annually.",
        "references": [
            "ISA 62443-2-1:2009 4.3.2.6, ISO/IEC 27001:2013 A.5.1.1, NIST SP 800-53 (all families).",
        ],
    },
    4: {
        "title": "Identify - Risk Assessment",
        "function": "IDENTIFY",
        "summary": "The organization understands cybersecurity risk to operations, assets, and individuals.",
        "subcategories": [
            "Asset vulnerabilities are identified and documented.",
            "Threat and vulnerability information is received from information sharing forums.",
            "Threats are identified and documented.",
            "Potential business impacts and likelihoods are identified.",
            "Threats, vulnerabilities, likelihoods, and impacts are used to determine risk.",
            "Risk responses are identified and prioritized.",
        ],
        "example": "Run quarterly vulnerability scans and produce risk-ranked remediation plans.",
        "solution": "Implement a Risk Register using NIST SP 800-30 methodology.",
        "recommendation": "Perform annual Risk Assessments per OMB A-130. Prioritize Critical/High vulnerabilities.",
        "references": [
            "CCS CSC 4, COBIT 5 APO12.01-04, NIST SP 800-53 CA-2/CA-7/RA-3/RA-5.",
        ],
    },
    5: {
        "title": "Identify - Risk Management Strategy",
        "function": "IDENTIFY",
        "summary": "Priorities, constraints, risk tolerances, and assumptions are established and used to support operational risk decisions.",
        "subcategories": [
            "Risk management processes are established and agreed to by stakeholders.",
            "Organizational risk tolerance is determined and clearly expressed.",
            "Risk tolerance is informed by the organization's role in critical infrastructure.",
            "Resilience requirements to support delivery of critical services are established.",
        ],
        "example": "Define a Risk Tolerance Statement tied to ATO renewal criteria.",
        "solution": "Develop an Enterprise Risk Management framework aligned to OMB Circular A-11.",
        "recommendation": "Review risk tolerance annually with senior leadership.",
        "references": [
            "COBIT 5 APO12.04/05/APO13.02, NIST SP 800-53 PM-9.",
        ],
    },
    6: {
        "title": "Protect - Access Control",
        "function": "PROTECT",
        "summary": "Access to assets and facilities is limited to authorized users, processes, or devices.",
        "subcategories": [
            "Identities and credentials are managed for authorized devices and users.",
            "Physical access to assets is managed and protected.",
            "Remote access is managed.",
            "Access permissions are managed with least privilege and separation of duties.",
            "Network integrity is protected, incorporating network segregation.",
        ],
        "example": "Implement CAC/PIV card authentication and PAM for privileged accounts.",
        "solution": "Deploy Zero Trust Network Access with MFA for all remote access.",
        "recommendation": "Review access rights quarterly and log all privileged access to SIEM.",
        "references": [
            "CCS CSC 16, NIST SP 800-53 AC-2/IA family.",
        ],
    },
    7: {
        "title": "Protect - Awareness & Training",
        "function": "PROTECT",
        "summary": "Personnel and partners are provided cybersecurity awareness education and training.",
        "subcategories": [
            "All users are informed and trained.",
            "Privileged users understand roles and responsibilities.",
            "Third-party stakeholders understand roles and responsibilities.",
            "Senior executives understand roles and responsibilities.",
            "Physical and information security personnel understand roles and responsibilities.",
        ],
        "example": "Mandate annual Cyber Awareness training and phishing simulations.",
        "solution": "Implement a Security Awareness Training platform and track completion rates.",
        "recommendation": "Require 100% annual training completion before account renewal.",
        "references": [
            "CCS CSC 9, NIST SP 800-53 AT-3/PM-13.",
        ],
    },
    8: {
        "title": "Protect - Data Security",
        "function": "PROTECT",
        "summary": "Information and records are managed to protect confidentiality, integrity, and availability.",
        "subcategories": [
            "Data-at-rest is protected.",
            "Data-in-transit is protected.",
            "Assets are managed throughout removal, transfers, and disposition.",
            "Adequate capacity to ensure availability is maintained.",
            "Protections against data leaks are implemented.",
            "Integrity checking mechanisms are used.",
            "Development and testing environments are separate from production.",
        ],
        "example": "Encrypt all data-at-rest and use TLS 1.3 for data-in-transit.",
        "solution": "Deploy Data Classification and DLP solutions.",
        "recommendation": "Classify all data per DoD guidance and encrypt CUI at rest and in transit.",
        "references": [
            "CCS CSC 17, NIST SP 800-53 SC-28/SC-8.",
        ],
    },
    9: {
        "title": "Protect - Information Protection Processes & Procedures",
        "function": "PROTECT",
        "summary": "Security policies, processes, and procedures are maintained and used to manage protection of information systems and assets.",
        "subcategories": [
            "Baseline configuration of IT/ICS is created and maintained.",
            "SDLC to manage systems is implemented.",
            "Configuration change control processes are in place.",
            "Backups are conducted, maintained, and tested.",
            "Physical operating environment policies are met.",
            "Data is destroyed according to policy.",
            "Protection processes are continuously improved.",
            "Effectiveness of protection technologies is shared.",
            "Response and recovery plans are in place and managed.",
            "Response and recovery plans are tested.",
            "Cybersecurity is included in HR practices.",
            "A vulnerability management plan is implemented.",
        ],
        "example": "Maintain STIG-hardened baselines and test backup restoration regularly.",
        "solution": "Implement Configuration Management and backup tools with DR plans.",
        "recommendation": "Apply STIGs before deployment and enforce change management via ITSM.",
        "references": [
            "CCS CSC 3/10, NIST SP 800-53 CM-2/3/4/5/6/7/9.",
        ],
    },
    10: {
        "title": "Protect - Maintenance of ICS & Information Systems",
        "function": "PROTECT",
        "summary": "Maintenance and repairs of ICS and information system components are performed consistent with policies and procedures.",
        "subcategories": [
            "Maintenance and repair of assets is performed and logged with approved tools.",
            "Remote maintenance is approved, logged, and performed securely.",
        ],
        "example": "Use jump servers for remote maintenance and record all sessions.",
        "solution": "Require dual authorization for critical system maintenance.",
        "recommendation": "Review maintenance logs monthly for anomalies.",
        "references": [
            "NIST SP 800-53 MA-2/3/5.",
        ],
    },
    11: {
        "title": "Protect - Protective Technology",
        "function": "PROTECT",
        "summary": "Technical security solutions are managed to ensure security and resilience of systems and assets.",
        "subcategories": [
            "Audit/log records are determined, documented, implemented, and reviewed.",
            "Removable media is protected and restricted.",
            "Access to systems and assets is controlled with least functionality.",
            "Communications and control networks are protected.",
        ],
        "example": "Deploy SIEM, disable USB ports, and use application whitelisting.",
        "solution": "Implement NGFW with IPS and network segmentation.",
        "recommendation": "Audit logs daily and review firewall rules quarterly.",
        "references": [
            "NIST SP 800-53 AU family/AC-3/CM-7/SC-7.",
        ],
    },
    12: {
        "title": "Detect - Anomalies & Events",
        "function": "DETECT",
        "summary": "Anomalous activity is detected in a timely manner and the potential impact of events is understood.",
        "subcategories": [
            "Baseline of network operations and expected data flows is established.",
            "Detected events are analyzed to understand attack targets and methods.",
            "Event data are aggregated and correlated from multiple sources.",
            "Impact of events is determined.",
            "Incident alert thresholds are established.",
        ],
        "example": "Use network flow analysis tools to establish behavioral baselines.",
        "solution": "Deploy UEBA integrated with threat intel feeds.",
        "recommendation": "Map detection rules to MITRE ATT&CK and conduct threat hunting quarterly.",
        "references": [
            "NIST SP 800-53 SI-4/CA-3/AC-4.",
        ],
    },
    # You can continue adding items 13–21 here following the same structure.
}

# ══════════════════════════════════════════════════════════════════════════════
#  EXPORT HELPERS
# ══════════════════════════════════════════════════════════════════════════════
def build_text_report(item_id: int, item: dict) -> str:
    lines = [
        "=" * 80,
        "CYBERSECURITY FRAMEWORK APPLICATION",
        "Developed by Randy Singh | KalSnet (KNet) Consulting Group",
        f"Checklist Item: {item_id} - {item['title']}",
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "=" * 80,
        "",
        f"Function: {item['function']}",
        "",
        "Summary:",
        f"  {item['summary']}",
        "",
        "Subcategories:",
    ]
    for sc in item["subcategories"]:
        lines.append(f"  - {sc}")
    lines.append("")
    lines.append("Example:")
    lines.append(f"  {item['example']}")
    lines.append("")
    lines.append("Solution:")
    lines.append(f"  {item['solution']}")
    lines.append("")
    lines.append("Recommendation:")
    lines.append(f"  {item['recommendation']}")
    lines.append("")
    lines.append("References:")
    for ref in item["references"]:
        lines.append(f"  {ref}")
    lines.append("")
    return "\n".join(lines)


def build_json_report(item_id: int, item: dict) -> str:
    payload = {
        "framework": "CyberSecurity Framework Application",
        "author": "Randy Singh | KalSnet (KNet) Consulting Group",
        "checklist_item_id": item_id,
        "title": item["title"],
        "function": item["function"],
        "generated": datetime.now().isoformat(),
        "data": item,
    }
    return json.dumps(payload, indent=2, default=str)


def _pdf_safe(text: str) -> str:
    replacements = [
        ("\u2013", "-"),
        ("\u2014", "--"),
        ("\u2018", "'"),
        ("\u2019", "'"),
        ("\u201c", '"'),
        ("\u201d", '"'),
        ("\u2026", "..."),
        ("\u2022", "-"),
        ("\u00a0", " "),
    ]
    for char, replacement in replacements:
        text = text.replace(char, replacement)
    return text.encode("latin-1", errors="ignore").decode("latin-1")


def build_pdf_bytes(item_id: int, item: dict) -> bytes:
    # If fpdf2 is available, use it; otherwise, build a minimal PDF fallback.
    text_report = build_text_report(item_id, item)

    if FPDF_AVAILABLE:
        try:
            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()
            pdf.set_font("Helvetica", "B", 14)
            pdf.set_text_color(0, 71, 171)
            pdf.multi_cell(0, 8, _pdf_safe("CYBERSECURITY FRAMEWORK APPLICATION"), align="C")

            pdf.set_font("Helvetica", "", 9)
            pdf.set_text_color(30, 30, 30)
            pdf.ln(4)
            pdf.multi_cell(0, 5, _pdf_safe(text_report))

            raw = pdf.output()
            if isinstance(raw, (bytes, bytearray)):
                return bytes(raw)
            return raw.encode("latin-1")
        except Exception:
            pass

    # Minimal PDF fallback (simple text-only PDF)
    content = _pdf_safe(text_report)
    pdf_bytes = (
        b"%PDF-1.4\n"
        b"1 0 obj<<>>endobj\n"
        b"2 0 obj<< /Length 3 0 R >>\n"
        b"stream\n"
        + content.encode("latin-1", errors="ignore")
        + b"\nendstream\nendobj\n"
        b"3 0 obj " + str(len(content)).encode("ascii") + b"\n"
        b"4 0 obj<< /Type /Page /Parent 5 0 R /Contents 2 0 R >>endobj\n"
        b"5 0 obj<< /Type /Pages /Kids [4 0 R] /Count 1 >>endobj\n"
        b"6 0 obj<< /Type /Catalog /Pages 5 0 R >>endobj\n"
        b"xref\n0 7\n0000000000 65535 f \n"
        b"trailer<< /Size 7 /Root 6 0 R >>\nstartxref\n0\n%%EOF"
    )
    return pdf_bytes


def build_docx_bytes(item_id: int, item: dict) -> bytes:
    if not DOCX_AVAILABLE:
        return b""
    try:
        doc = DocxDocument()

        t = doc.add_heading("CYBERSECURITY FRAMEWORK APPLICATION", 0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in t.runs:
            run.font.color.rgb = RGBColor(0, 71, 171)

        sub = doc.add_paragraph(
            "Developed by Randy Singh | KalSnet (KNet) Consulting Group"
        )
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if sub.runs:
            sub.runs[0].bold = True
            sub.runs[0].font.color.rgb = RGBColor(0, 71, 171)

        ts_para = doc.add_paragraph(
            f"Checklist Item: {item_id} - {item['title']}   |   {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        )
        ts_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        h = doc.add_heading("Summary", level=2)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0, 71, 171)
        doc.add_paragraph(item["summary"])

        h = doc.add_heading("Subcategories", level=2)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0, 71, 171)
        for sc in item["subcategories"]:
            doc.add_paragraph(sc, style="List Bullet")

        h = doc.add_heading("Example", level=2)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0, 71, 171)
        doc.add_paragraph(item["example"])

        h = doc.add_heading("Solution", level=2)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0, 71, 171)
        doc.add_paragraph(item["solution"])

        h = doc.add_heading("Recommendation", level=2)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0, 71, 171)
        doc.add_paragraph(item["recommendation"])

        h = doc.add_heading("References", level=2)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0, 71, 171)
        for ref in item["references"]:
            doc.add_paragraph(ref, style="List Bullet")

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except Exception:
        return b""

# ══════════════════════════════════════════════════════════════════════════════
#  SIDEBAR — NAVIGATION & EXPORT (RADIO BUTTONS)
# ══════════════════════════════════════════════════════════════════════════════
with st.sidebar:
    st.markdown("## Browse Checklist Items")
    st.markdown("---")

    item_labels = [
        f"{i}. {FRAMEWORK_ITEMS[i]['title']}"
        for i in sorted(FRAMEWORK_ITEMS.keys())
    ]
    selected_label = st.radio(
        "Select checklist item",
        item_labels,
        index=0,
        key="item_radio",
    )

    selected_id = int(selected_label.split(".")[0])
    selected_item = FRAMEWORK_ITEMS[selected_id]

    st.markdown('<div class="sidebar-box">', unsafe_allow_html=True)
    st.markdown("**About**")
    st.markdown(
        "Randy Singh  \nComputer Scientist  \nKalSnet (KNet) Consulting Group"
    )
    st.markdown("`(301) 225-9535`")
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown('<div class="sidebar-box">', unsafe_allow_html=True)
    st.markdown("**Export This Checklist Item**")

    export_choice = st.radio(
        "Choose export format",
        ["Text (.txt)", "JSON (.json)", "PDF (.pdf)", "Word (.docx)"],
        index=0,
        key="export_radio",
    )

    text_report = build_text_report(selected_id, selected_item)
    json_report = build_json_report(selected_id, selected_item)
    pdf_bytes = build_pdf_bytes(selected_id, selected_item)
    docx_bytes = build_docx_bytes(selected_id, selected_item)

    if export_choice == "Text (.txt)":
        st.download_button(
            "Download Text",
            data=text_report,
            file_name=f"cyber_framework_item_{selected_id}.txt",
            mime="text/plain",
            key="dl_txt",
        )
    elif export_choice == "JSON (.json)":
        st.download_button(
            "Download JSON",
            data=json_report,
            file_name=f"cyber_framework_item_{selected_id}.json",
            mime="application/json",
            key="dl_json",
        )
    elif export_choice == "PDF (.pdf)":
        st.download_button(
            "Download PDF",
            data=pdf_bytes,
            file_name=f"cyber_framework_item_{selected_id}.pdf",
            mime="application/pdf",
            key="dl_pdf",
        )
    elif export_choice == "Word (.docx)":
        if docx_bytes:
            st.download_button(
                "Download Word",
                data=docx_bytes,
                file_name=f"cyber_framework_item_{selected_id}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl_docx",
            )
        else:
            st.caption("Word export not available — install python-docx in the environment.")

    st.markdown(
        '<div class="sidebar-footer">CyberSecurity Framework | KalSnet (KNet) Consulting</div>',
        unsafe_allow_html=True,
    )

# ══════════════════════════════════════════════════════════════════════════════
#  MAIN CONTENT — SELECTED CHECKLIST ITEM
# ══════════════════════════════════════════════════════════════════════════════
st.markdown(
    '<div class="section-header">Browse CyberSecurity Framework Checklist Items</div>',
    unsafe_allow_html=True,
)

st.markdown(f"### {selected_item['title']}")
st.markdown(f"**Function:** {selected_item['function']}")

st.markdown(
    f'<div class="req-card"><b>Summary:</b> {selected_item["summary"]}</div>',
    unsafe_allow_html=True,
)

st.markdown("#### Subcategories")
for sc in selected_item["subcategories"]:
    st.markdown(f"- {sc}")

st.markdown(
    f'<div class="sol-card"><b>Solution Pattern:</b> {selected_item["solution"]}</div>',
    unsafe_allow_html=True,
)
st.markdown(
    f'<div class="rec-card"><b>Recommendation:</b> {selected_item["recommendation"]}</div>',
    unsafe_allow_html=True,
)
st.markdown(
    '<div class="ref-card"><b>References:</b><br>'
    + "<br>".join(selected_item["references"])
    + "</div>",
    unsafe_allow_html=True,
)

# ══════════════════════════════════════════════════════════════════════════════
#  FOOTER
# ══════════════════════════════════════════════════════════════════════════════
st.markdown(
    """
<div style="text-align:center;color:#888;font-size:0.8rem;margin-top:20px;">
CyberSecurity Framework Application | Developed by Randy Singh | KalSnet (KNet) Consulting Group |
{year}
</div>
""".format(year=datetime.now().year),
    unsafe_allow_html=True,
)
