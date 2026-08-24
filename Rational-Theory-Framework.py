

# Rational Choice & Game Theory Framework
# Developed by Randy Singh - Kalsnet (KNet) Consulting


import io
import itertools
from datetime import datetime

import numpy as np
import pandas as pd
import streamlit as st
import matplotlib.pyplot as plt
import seaborn as sns

# ------------------------------------------------------------------
# Optional export engines (Word / PDF). App still works if missing.
# ------------------------------------------------------------------
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

try:
    from fpdf import FPDF
    PDF_OK = True
except ImportError:
    PDF_OK = False


# ====================================================================
# PAGE CONFIG + GLOBAL STYLE
# ====================================================================
st.set_page_config(
    page_title="Rational Choice & Game Theory Framework | KNet Consulting",
    page_icon="🧭",
    layout="wide",
    initial_sidebar_state="expanded",
)

PRIMARY_BLUE = "#0B3D91"
ACCENT_BLUE = "#1668E3"
LIGHT_BLUE = "#EAF2FF"

st.markdown(
    f"""
    <style>
    .knet-title {{
        font-size: 44px;
        font-weight: 900;
        color: {PRIMARY_BLUE};
        line-height: 1.15;
        margin-bottom: 0px;
        letter-spacing: 0.5px;
    }}
    .knet-subtitle {{
        font-size: 24px;
        font-weight: 800;
        color: {ACCENT_BLUE};
        margin-top: 2px;
        margin-bottom: 18px;
    }}
    .knet-banner {{
        background: linear-gradient(90deg, {LIGHT_BLUE} 0%, #FFFFFF 100%);
        padding: 18px 24px;
        border-left: 8px solid {PRIMARY_BLUE};
        border-radius: 6px;
        margin-bottom: 20px;
    }}
    .metric-card {{
        background-color: {LIGHT_BLUE};
        border-radius: 10px;
        padding: 14px 18px;
        border: 1px solid #cfe0fb;
    }}
    section[data-testid="stSidebar"] {{
        background-color: #0B1F3A;
    }}
    section[data-testid="stSidebar"] * {{
        color: #F0F4FF !important;
    }}
    </style>
    """,
    unsafe_allow_html=True,
)

st.markdown(
    f"""
    <div class="knet-banner">
        <div class="knet-title">Rational Choice &amp; Game Theory Framework</div>
        <div class="knet-subtitle">Developed by Randy Singh &mdash; Kalsnet (KNet) Consulting</div>
    </div>
    """,
    unsafe_allow_html=True,
)

# ====================================================================
# HOW-IT-WORKS / SCHEMA EXPLANATION PANEL
# ====================================================================
with st.expander("ℹ️  How this framework works & what each field means", expanded=True):
    st.markdown(
        """
**How it works, in brief:** This framework supports two classic decision-science
approaches. The **Rational Choice** module scores a set of *alternatives* against a
set of *criteria*, applies a *weight* to each criterion, and combines them into a
single **Expected Utility** score per alternative — the option that maximizes this
score is the "rational" choice. The **Game Theory** module models a strategic
interaction between two players, each with their own *payoff matrix*, and solves
for **Nash equilibria** (strategy pairs where neither player can do better by
switching unilaterally) and any **strictly dominant strategies**. You can try both
with the built-in synthetic data generator, or upload your own test data.
        """
    )
    st.markdown("**Rational Choice (MCDA) schema — field by field:**")
    st.markdown(
        """
- **Alternative / Option** *(row label)* — one of the choices being evaluated (e.g. Option A, Option B).
- **Criterion** *(column, e.g. "Crit 1")* — a dimension used to judge the alternatives (e.g. cost, quality, risk).
- **Score** *(cell value, 1–10 in the demo)* — how well an alternative performs on that criterion.
- **Weight** *(0–5 slider)* — how important that criterion is relative to the others; weights are auto-normalized before use.
- **Expected Utility** *(output)* — the weighted sum of an alternative's scores across all criteria: `Σ (score × normalized weight)`.
- **Rank** *(output)* — the alternative's position after sorting by Expected Utility, highest first.
        """
    )
    st.markdown("**Game Theory schema — field by field:**")
    st.markdown(
        """
- **Player A / Player B** — the two decision-makers in the game.
- **Strategy** *(e.g. "A-strat 1", "B-strat 2")* — one of the discrete choices available to a player.
- **Payoff matrix** — a grid where each cell holds the payoff a player receives for a given combination of both players' strategies.
- **Nash Equilibrium** *(output)* — a strategy pair where Player A's choice is a best response to Player B's, and vice versa; neither player benefits from switching alone.
- **Dominant Strategy** *(output)* — a strategy that yields a strictly better payoff for a player than any of their other strategies, no matter what the opponent does.
        """
    )

# ====================================================================
# SESSION STATE
# ====================================================================
if "mcda_df" not in st.session_state:
    st.session_state.mcda_df = None
if "mcda_results" not in st.session_state:
    st.session_state.mcda_results = None
if "game_results" not in st.session_state:
    st.session_state.game_results = None
if "payoff_a" not in st.session_state:
    st.session_state.payoff_a = None
if "payoff_b" not in st.session_state:
    st.session_state.payoff_b = None

# ====================================================================
# HELPER FUNCTIONS
# ====================================================================

def make_synthetic_mcda(n_alt=5, n_crit=4, seed=None):
    """Synthetic Rational Choice dataset: alternatives x criteria scores + weights."""
    rng = np.random.default_rng(seed)
    alt_names = [f"Option {chr(65+i)}" for i in range(n_alt)]
    crit_names = [f"Criterion {i+1}" for i in range(n_crit)]
    scores = rng.integers(1, 11, size=(n_alt, n_crit))
    df = pd.DataFrame(scores, columns=crit_names, index=alt_names)
    weights = rng.dirichlet(np.ones(n_crit)).round(3)
    weights = weights / weights.sum()
    return df, dict(zip(crit_names, weights))


def compute_weighted_utility(df, weights):
    w = np.array([weights[c] for c in df.columns])
    w = w / w.sum()
    utility = (df.values * w).sum(axis=1)
    out = pd.DataFrame(
        {"Expected Utility": utility},
        index=df.index,
    ).sort_values("Expected Utility", ascending=False)
    out["Rank"] = range(1, len(out) + 1)
    return out


def make_synthetic_game(size=3, seed=None):
    rng = np.random.default_rng(seed)
    payoff_a = rng.integers(-5, 10, size=(size, size))
    payoff_b = rng.integers(-5, 10, size=(size, size))
    return payoff_a, payoff_b


def find_pure_nash(payoff_a, payoff_b):
    """Brute-force pure-strategy Nash equilibria for a 2-player normal-form game."""
    rows, cols = payoff_a.shape
    equilibria = []
    for i in range(rows):
        for j in range(cols):
            row_best = payoff_a[:, j].max() == payoff_a[i, j]
            col_best = payoff_b[i, :].max() == payoff_b[i, j]
            if row_best and col_best:
                equilibria.append((i, j))
    return equilibria


def find_dominant_strategies(payoff_a, payoff_b):
    """Very simple strict-dominance check (row player dominance only, illustrative)."""
    dominant = []
    rows = payoff_a.shape[0]
    for i in range(rows):
        others = [k for k in range(rows) if k != i]
        if all((payoff_a[i, :] > payoff_a[k, :]).all() for k in others):
            dominant.append(i)
    return dominant


def df_to_csv_bytes(df):
    return df.to_csv().encode("utf-8")


def df_to_txt_bytes(df, title=""):
    buf = io.StringIO()
    buf.write(f"{title}\n")
    buf.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
    buf.write("Rational Choice & Game Theory Framework - Randy Singh, Kalsnet (KNet) Consulting\n\n")
    buf.write(df.to_string())
    return buf.getvalue().encode("utf-8")


def df_to_word_bytes(df, title=""):
    if not DOCX_OK:
        return None
    doc = Document()
    heading = doc.add_heading(title, level=1)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x0B, 0x3D, 0x91)
    sub = doc.add_paragraph("Developed by Randy Singh — Kalsnet (KNet) Consulting")
    sub.runs[0].font.bold = True
    sub.runs[0].font.color.rgb = RGBColor(0x16, 0x68, 0xE3)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=len(df.columns) + 1)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Item"
    for idx, col in enumerate(df.columns):
        hdr[idx + 1].text = str(col)
    for row_label, row in df.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row_label)
        for idx, val in enumerate(row):
            cells[idx + 1].text = str(val)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def df_to_pdf_bytes(df, title=""):
    if not PDF_OK:
        return None
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 18)
    pdf.set_text_color(11, 61, 145)
    pdf.cell(0, 12, title, ln=True)
    pdf.set_font("Helvetica", "B", 12)
    pdf.set_text_color(22, 104, 227)
    pdf.cell(0, 8, "Developed by Randy Singh - Kalsnet (KNet) Consulting", ln=True)
    pdf.set_text_color(0, 0, 0)
    pdf.set_font("Helvetica", "", 9)
    pdf.cell(0, 8, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", ln=True)
    pdf.ln(4)

    col_width = 190 / (len(df.columns) + 1)
    pdf.set_font("Helvetica", "B", 9)
    pdf.set_fill_color(234, 242, 255)
    pdf.cell(col_width, 8, "Item", border=1, fill=True)
    for col in df.columns:
        pdf.cell(col_width, 8, str(col)[:18], border=1, fill=True)
    pdf.ln()
    pdf.set_font("Helvetica", "", 9)
    for row_label, row in df.iterrows():
        pdf.cell(col_width, 8, str(row_label)[:18], border=1)
        for val in row:
            pdf.cell(col_width, 8, str(round(val, 3)) if isinstance(val, float) else str(val), border=1)
        pdf.ln()

    return bytes(pdf.output(dest="S"))


# ====================================================================
# SIDEBAR NAVIGATION
# ====================================================================
st.sidebar.markdown("## 🧭 Navigation")
page = st.sidebar.radio(
    "Go to",
    [
        "Overview",
        "1. Synthetic Data Demo",
        "2. Upload Your Own Data",
        "3. Rational Choice Model (MCDA)",
        "4. Game Theory Model (Nash)",
        "5. Export Results",
    ],
)
st.sidebar.markdown("---")
st.sidebar.markdown("**KNet Consulting**  \nRandy Singh — Framework Author")
st.sidebar.markdown("v1.0 · Rational Choice & Game Theory")

# ====================================================================
# PAGE: OVERVIEW
# ====================================================================
if page == "Overview":
    st.subheader("About this Framework")
    st.write(
        "This framework operationalizes two classical decision-science tools "
        "in one interactive workspace:"
    )
    c1, c2 = st.columns(2)
    with c1:
        st.markdown("#### 🧮 Rational Choice (MCDA)")
        st.write(
            "A weighted multi-criteria decision model. Alternatives are scored "
            "against criteria, weights are applied, and an **Expected Utility** "
            "ranking is produced — the classic rational-actor decision rule: "
            "choose the option that maximizes expected utility."
        )
    with c2:
        st.markdown("#### ♟️ Game Theory (Strategic Interaction)")
        st.write(
            "A two-player normal-form game module. Enter or generate payoff "
            "matrices for Player A and Player B, and the framework solves for "
            "**pure-strategy Nash equilibria** and strictly dominant strategies."
        )
    st.info(
        "Use the sidebar to explore a synthetic demo, upload your own dataset, "
        "run either model, and export a report in PDF, Word, Text, or CSV."
    )
    st.markdown("##### Workflow")
    st.markdown(
        "1. **Synthetic Data Demo** — generate a random dataset to see the framework in action.\n"
        "2. **Upload Your Own Data** — bring real test data (CSV/TXT).\n"
        "3. **Rational Choice Model** — compute weighted utility rankings.\n"
        "4. **Game Theory Model** — compute Nash equilibria on a payoff matrix.\n"
        "5. **Export Results** — download a report as PDF, Word, Text, or CSV."
    )

# ====================================================================
# PAGE: SYNTHETIC DATA DEMO
# ====================================================================
elif page == "1. Synthetic Data Demo":
    st.subheader("Synthetic Data Demo")
    st.write("Generate a random demo dataset to preview how the framework behaves, before using real data.")

    colA, colB, colC = st.columns(3)
    with colA:
        n_alt = st.slider("Number of alternatives / options", 2, 8, 5)
    with colB:
        n_crit = st.slider("Number of criteria", 2, 8, 4)
    with colC:
        seed = st.number_input("Random seed", value=42, step=1)

    if st.button("🎲 Generate Synthetic Rational Choice Data", type="primary"):
        df, weights = make_synthetic_mcda(n_alt, n_crit, seed)
        st.session_state.mcda_df = df
        st.session_state.mcda_weights = weights
        st.session_state.mcda_results = compute_weighted_utility(df, weights)

    if st.session_state.mcda_df is not None:
        st.markdown("**Synthetic scores (1–10 scale) by alternative and criterion:**")
        st.dataframe(st.session_state.mcda_df, use_container_width=True)

        st.markdown("**Randomly generated criteria weights:**")
        w_df = pd.DataFrame.from_dict(st.session_state.mcda_weights, orient="index", columns=["Weight"])
        st.bar_chart(w_df)

        st.markdown("**Resulting Expected Utility ranking:**")
        st.dataframe(st.session_state.mcda_results, use_container_width=True)

        fig, ax = plt.subplots(figsize=(6, 3.2))
        colors = sns.color_palette("Blues_r", len(st.session_state.mcda_results))
        ax.bar(
            st.session_state.mcda_results.index,
            st.session_state.mcda_results["Expected Utility"],
            color=colors,
        )
        ax.set_ylabel("Expected Utility")
        ax.set_title("Synthetic Demo — Expected Utility by Option", color=PRIMARY_BLUE, fontweight="bold")
        plt.xticks(rotation=20)
        st.pyplot(fig)

    st.markdown("---")
    st.write("Prefer strategic interaction instead of ranking? Generate a synthetic game:")
    size = st.slider("Number of strategies per player", 2, 5, 3)
    if st.button("🎲 Generate Synthetic Game Payoffs"):
        pa, pb = make_synthetic_game(size, seed)
        st.session_state.payoff_a = pa
        st.session_state.payoff_b = pb
        st.success("Synthetic game generated — open '4. Game Theory Model (Nash)' to analyze it.")

# ====================================================================
# PAGE: UPLOAD YOUR OWN DATA
# ====================================================================
elif page == "2. Upload Your Own Data":
    st.subheader("Upload Your Own Test Data")
    st.write(
        "Upload any CSV or text file with alternatives as rows and criteria as columns "
        "(numeric scores). This becomes the input to the Rational Choice model."
    )

    uploaded = st.file_uploader("Upload CSV or TXT file", type=["csv", "txt"])
    if uploaded is not None:
        try:
            if uploaded.name.endswith(".csv"):
                df = pd.read_csv(uploaded, index_col=0)
            else:
                df = pd.read_csv(uploaded, sep=None, engine="python", index_col=0)
            st.session_state.mcda_df = df
            st.success(f"Loaded '{uploaded.name}' — {df.shape[0]} rows × {df.shape[1]} columns.")
            st.dataframe(df, use_container_width=True)

            st.markdown("**Assign weights to each criterion (must sum to any positive total — auto-normalized):**")
            weights = {}
            cols = st.columns(min(4, len(df.columns)))
            for i, c in enumerate(df.columns):
                with cols[i % len(cols)]:
                    weights[c] = st.number_input(f"Weight: {c}", min_value=0.0, value=1.0, step=0.1, key=f"w_{c}")
            st.session_state.mcda_weights = weights

            if st.button("✅ Compute Expected Utility from Uploaded Data", type="primary"):
                st.session_state.mcda_results = compute_weighted_utility(df, weights)
                st.dataframe(st.session_state.mcda_results, use_container_width=True)
        except Exception as e:
            st.error(f"Could not parse file: {e}")
    else:
        st.info("No file uploaded yet. You can also use the Synthetic Data Demo tab to try the framework first.")

# ====================================================================
# PAGE: RATIONAL CHOICE MODEL
# ====================================================================
elif page == "3. Rational Choice Model (MCDA)":
    st.subheader("Rational Choice Model — Weighted Expected Utility")
    if st.session_state.mcda_df is None:
        st.warning("No dataset loaded yet. Go to 'Synthetic Data Demo' or 'Upload Your Own Data' first.")
    else:
        df = st.session_state.mcda_df
        weights = st.session_state.get("mcda_weights", {c: 1.0 for c in df.columns})

        st.markdown("**Current data:**")
        st.dataframe(df, use_container_width=True)

        st.markdown("**Adjust weights:**")
        cols = st.columns(min(4, len(df.columns)))
        new_weights = {}
        for i, c in enumerate(df.columns):
            with cols[i % len(cols)]:
                new_weights[c] = st.slider(c, 0.0, 5.0, float(weights.get(c, 1.0)), 0.1, key=f"rc_{c}")

        results = compute_weighted_utility(df, new_weights)
        st.session_state.mcda_results = results
        st.session_state.mcda_weights = new_weights

        st.markdown("**Expected Utility Ranking (Rational Choice output):**")
        st.dataframe(results, use_container_width=True)

        best = results.index[0]
        st.success(f"🏆 The rational choice, given current weights, is **{best}** "
                   f"(Expected Utility = {results.iloc[0]['Expected Utility']:.2f}).")

        fig, ax = plt.subplots(figsize=(6, 3.2))
        colors = sns.color_palette("Blues_r", len(results))
        ax.barh(results.index[::-1], results["Expected Utility"][::-1], color=colors[::-1])
        ax.set_xlabel("Expected Utility")
        ax.set_title("Expected Utility Ranking", color=PRIMARY_BLUE, fontweight="bold")
        st.pyplot(fig)

# ====================================================================
# PAGE: GAME THEORY MODEL
# ====================================================================
elif page == "4. Game Theory Model (Nash)":
    st.subheader("Game Theory Model — Two-Player Normal-Form Game")
    st.write("Provide payoff matrices for Player A (row) and Player B (column), or use synthetic data.")

    size = st.slider("Number of strategies per player", 2, 5,
                      st.session_state.payoff_a.shape[0] if st.session_state.payoff_a is not None else 3)

    if st.session_state.payoff_a is None or st.session_state.payoff_a.shape[0] != size:
        pa, pb = make_synthetic_game(size, seed=1)
        st.session_state.payoff_a, st.session_state.payoff_b = pa, pb

    st.markdown("**Player A payoffs (row player):**")
    df_a = pd.DataFrame(
        st.session_state.payoff_a,
        columns=[f"B-strat {j+1}" for j in range(size)],
        index=[f"A-strat {i+1}" for i in range(size)],
    )
    edited_a = st.data_editor(df_a, use_container_width=True, key="edit_a")

    st.markdown("**Player B payoffs (column player):**")
    df_b = pd.DataFrame(
        st.session_state.payoff_b,
        columns=[f"B-strat {j+1}" for j in range(size)],
        index=[f"A-strat {i+1}" for i in range(size)],
    )
    edited_b = st.data_editor(df_b, use_container_width=True, key="edit_b")

    if st.button("🧮 Solve Game (Nash Equilibria + Dominance)", type="primary"):
        pa = edited_a.values
        pb = edited_b.values
        equilibria = find_pure_nash(pa, pb)
        dominant = find_dominant_strategies(pa, pb)
        st.session_state.game_results = {
            "payoff_a": pa, "payoff_b": pb,
            "equilibria": equilibria, "dominant": dominant,
        }

    if st.session_state.game_results:
        res = st.session_state.game_results
        pa, pb = res["payoff_a"], res["payoff_b"]

        if res["equilibria"]:
            eq_text = ", ".join([f"(A-strat {i+1}, B-strat {j+1})" for i, j in res["equilibria"]])
            st.success(f"✅ Pure-strategy Nash equilibria found: {eq_text}")
        else:
            st.warning("No pure-strategy Nash equilibrium found (a mixed-strategy equilibrium likely exists).")

        if res["dominant"]:
            st.info("Strictly dominant strategy for Player A: "
                     + ", ".join([f"A-strat {i+1}" for i in res["dominant"]]))
        else:
            st.write("No strictly dominant strategy identified for Player A.")

        st.markdown("**Combined payoff matrix (A, B):**")
        combo = pd.DataFrame(
            [[f"({pa[i,j]}, {pb[i,j]})" for j in range(pa.shape[1])] for i in range(pa.shape[0])],
            columns=[f"B-strat {j+1}" for j in range(pa.shape[1])],
            index=[f"A-strat {i+1}" for i in range(pa.shape[0])],
        )
        st.dataframe(combo, use_container_width=True)

        fig, axes = plt.subplots(1, 2, figsize=(9, 3.5))
        sns.heatmap(pa, annot=True, fmt="d", cmap="Blues", ax=axes[0], cbar=False)
        axes[0].set_title("Player A payoffs", color=PRIMARY_BLUE, fontweight="bold")
        sns.heatmap(pb, annot=True, fmt="d", cmap="Blues", ax=axes[1], cbar=False)
        axes[1].set_title("Player B payoffs", color=PRIMARY_BLUE, fontweight="bold")
        st.pyplot(fig)

        st.session_state.game_summary_df = pd.DataFrame(
            {
                "Nash Equilibria (A,B strategy index)": [str(res["equilibria"])],
                "Dominant A Strategy": [str(res["dominant"])],
            }
        )

# ====================================================================
# PAGE: EXPORT RESULTS
# ====================================================================
elif page == "5. Export Results":
    st.subheader("Export Results")
    st.write("Download your current results in the format you need. Reports are generated on demand.")

    export_choice = st.selectbox(
        "Which result set do you want to export?",
        ["Rational Choice — Expected Utility Ranking", "Game Theory — Summary"],
    )

    if export_choice.startswith("Rational") and st.session_state.mcda_results is not None:
        df = st.session_state.mcda_results
        title = "Rational Choice Framework Report"
    elif export_choice.startswith("Game") and st.session_state.get("game_summary_df") is not None:
        df = st.session_state.game_summary_df
        title = "Game Theory Framework Report"
    else:
        df = None

    if df is None:
        st.warning("No results available yet for this option. Run the corresponding model first.")
    else:
        st.dataframe(df, use_container_width=True)
        c1, c2, c3, c4 = st.columns(4)

        with c1:
            st.download_button(
                "⬇️ CSV", data=df_to_csv_bytes(df),
                file_name="knet_results.csv", mime="text/csv",
            )
        with c2:
            st.download_button(
                "⬇️ Text", data=df_to_txt_bytes(df, title),
                file_name="knet_results.txt", mime="text/plain",
            )
        with c3:
            word_bytes = df_to_word_bytes(df, title)
            if word_bytes:
                st.download_button(
                    "⬇️ Word (.docx)", data=word_bytes,
                    file_name="knet_results.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            else:
                st.caption("python-docx not installed — Word export unavailable.")
        with c4:
            pdf_bytes = df_to_pdf_bytes(df, title)
            if pdf_bytes:
                st.download_button(
                    "⬇️ PDF", data=pdf_bytes,
                    file_name="knet_results.pdf", mime="application/pdf",
                )
            else:
                st.caption("fpdf2 not installed — PDF export unavailable.")

st.markdown("---")
st.caption("Rational Choice & Game Theory Framework · © Randy Singh, Kalsnet (KNet) Consulting")