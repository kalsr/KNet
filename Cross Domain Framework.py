# CROSS-DOMAIN-SOLUTION-FRAMEWORK
import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
import json, io, random
from datetime import datetime
try:
    from reportlab.lib.pagesizes import letter, landscape
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    Table, TableStyle, HRFlowable)
    from reportlab.lib.units import inch
    REPORTLAB_OK = True
except ImportError:
    REPORTLAB_OK = False
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    DOCX_OK = True
except ImportError:
    DOCX_OK = False
# ── Page Config ─────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="KNet CDS Intelligence Framework",
    page_icon="🔐",
    layout="wide",
    initial_sidebar_state="expanded",
)
# ── CSS ─────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
  @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700;900&display=swap');
  html,body,[class*="css"]{font-family:'Inter',sans-serif;}
  .knet-hero{background:linear-gradient(135deg,#0A1628 0%,#0D2B6B 55%,#1565C0 100%);
    border-radius:16px;padding:38px 44px 30px;margin-bottom:28px;
    box-shadow:0 8px 40px rgba(21,101,192,.45);}
  .knet-title{font-size:2.5rem;font-weight:900;color:#1E90FF;letter-spacing:-.5px;line-height:1.15;}
  .knet-sub{font-size:1rem;color:#90CAF9;margin-top:6px;}
  .knet-badge{display:inline-block;background:rgba(30,144,255,.18);
    border:1px solid #1E90FF;color:#64B5F6;padding:5px 16px;border-radius:20px;
    font-size:.82rem;font-weight:700;margin-top:14px;letter-spacing:.5px;}
  .metric-card{background:linear-gradient(135deg,#0D2B6B,#1565C0);border-radius:12px;
    padding:18px 16px;text-align:center;box-shadow:0 4px 16px rgba(21,101,192,.3);
    border:1px solid rgba(30,144,255,.3);}
  .metric-val{font-size:1.75rem;font-weight:900;color:#1E90FF;}
  .metric-lbl{font-size:.75rem;color:#90CAF9;margin-top:4px;}
  .sh{font-size:1.2rem;font-weight:700;color:#1E90FF;
    border-left:4px solid #1E90FF;padding-left:12px;margin:28px 0 14px;}
  .tax{background:#0A1628;border:1px solid #1565C0;border-radius:10px;
    padding:20px 24px;line-height:1.75;color:#CFE2FF;font-size:.9rem;margin-bottom:18px;}
  .tax b{color:#64B5F6;}
  .tax ul,ol{margin-left:20px;margin-top:6px;}
  .risk-high{color:#FF4444;font-weight:700;}
  .risk-med {color:#FFA726;font-weight:700;}
  .risk-low {color:#66BB6A;font-weight:700;}
  .field{background:#0D2148;border-left:3px solid #1E90FF;border-radius:6px;
    padding:13px 16px;margin-bottom:10px;font-size:.875rem;color:#B0C4DE;}
  .field-name{color:#1E90FF;font-weight:700;font-size:.95rem;}
  .formula{background:rgba(30,144,255,.12);color:#64B5F6;padding:2px 8px;
    border-radius:4px;font-family:monospace;font-size:.8rem;}
  .guard-card{background:#0D2148;border:1px solid #1565C030;border-radius:8px;
    padding:14px 16px;margin-bottom:10px;}
  .guard-name{color:#1E90FF;font-weight:700;}
  .guard-tag{display:inline-block;background:rgba(30,144,255,.15);
    color:#90CAF9;padding:2px 8px;border-radius:12px;font-size:.75rem;margin:2px;}
</style>
""", unsafe_allow_html=True)
# ════════════════════════════════════════════════════════════════════════════
# CONSTANTS
# ════════════════════════════════════════════════════════════════════════════
DOMAINS     = ["Unclassified (U)","Confidential (C)","Secret (S)","Top Secret (TS)","TS/SCI"]
DOMAIN_LVL  = {"Unclassified (U)":1,"Confidential (C)":2,"Secret (S)":3,
                "Top Secret (TS)":4,"TS/SCI":5}
DIRECTIONS  = ["Low-to-High","High-to-Low","Bidirectional"]
DATA_TYPES  = ["File Transfer","Video Stream","Email/Messaging","Database Query",
               "Voice/VoIP","Sensor/Telemetry","API/Web Service","Bulk Data"]
GUARDS      = ["Forcepoint CDS","Owl Cyber Defense","Everfox High Side",
               "Raytheon SurePath","BAE Trusted Guard","Genuine Solutions GEMINI",
               "L3Harris CDG","Radiant Mercury","Telos XACTA CDS","Verizon Federal CDS"]
PROTOCOLS   = ["HTTPS/TLS 1.3","SMTP Sanitised","FTP Sanitised","IPSEC","STANAG 5066",
               "NIAPM","XTS-AES-256","REST/JSON Filtered"]
FILTER_TYPES= ["Content Inspection","Anti-Virus","Data Tagging","DLP","Whitelist",
               "Format Check","Keyword Filter","AI/ML Anomaly","Regex Pattern","Steganography Check"]
CERTS       = ["NIAP CC EAL4+","NIAP CC EAL6+","DISA APL","NSA Suite B","FIPS 140-3",
               "DoD IL5","DoD IL6","NATO SECRET","UK OFFICIAL-SENSITIVE"]
STATUS_OPTS = ["Approved","Pending","Denied","Under Review"]
ORGS        = ["DoD","IC Community","NATO","Five Eyes","Federal Civilian",
               "State Dept","DHS","NSA","GCHQ","CSE Canada"]
# Guard detail catalogue
GUARD_DETAIL = {
    "Forcepoint CDS": {
        "vendor":"Forcepoint","cert":"NIAP CC EAL4+, DISA APL, FIPS 140-3",
        "direction":"Bidirectional","max_domains":"6",
        "notes":"Market-leading CDS with deep content inspection, XML/JSON filtering, AV, and DLP. Used widely across DoD and IC."
    },
    "Owl Cyber Defense": {
        "vendor":"Owl Cyber Defense","cert":"NIAP CC EAL4+, FIPS 140-3",
        "direction":"Unidirectional / Bidirectional","max_domains":"4",
        "notes":"Hardware-enforced data diode and software-based CDS. Specialises in OT/ICS environments."
    },
    "Everfox High Side": {
        "vendor":"Everfox (fmr. Forcepoint Gov)","cert":"NIAP CC EAL4+, DoD IL6",
        "direction":"Bidirectional","max_domains":"8",
        "notes":"Purpose-built for Intelligence Community high-side networks. Supports TS/SCI transfer with AI/ML inspection."
    },
    "Raytheon SurePath": {
        "vendor":"Raytheon Technologies","cert":"NIAP CC EAL5+, NSA Suite B",
        "direction":"High-to-Low, Low-to-High","max_domains":"5",
        "notes":"Certified for NSA/CSS environments. Strong cryptographic separation and hardware root of trust."
    },
    "BAE Trusted Guard": {
        "vendor":"BAE Systems","cert":"NIAP CC EAL6+, UK OFFICIAL-SENSITIVE",
        "direction":"Bidirectional","max_domains":"6",
        "notes":"UK/Five Eyes certified. Used by MoD and allied defence for cross-domain file transfer and email sanitisation."
    },
    "Genuine Solutions GEMINI": {
        "vendor":"Genuine Solutions","cert":"NIAP CC EAL4+, FIPS 140-3",
        "direction":"Bidirectional","max_domains":"4",
        "notes":"Compact appliance-based CDS for tactical edge deployments and forward-deployed units."
    },
    "L3Harris CDG": {
        "vendor":"L3Harris Technologies","cert":"NIAP CC EAL4+, DoD IL5",
        "direction":"Low-to-High, High-to-Low","max_domains":"6",
        "notes":"Integrated Cross Domain Gateway for multi-coalition environments. Supports STANAG and NATO data standards."
    },
    "Radiant Mercury": {
        "vendor":"Viasat (Radiant Logic)","cert":"NIAP CC EAL4+, DISA APL",
        "direction":"Bidirectional","max_domains":"8",
        "notes":"Widely deployed IC CDS. Supports structured and unstructured data, XML, video, and voice across classification levels."
    },
    "Telos XACTA CDS": {
        "vendor":"Telos Corporation","cert":"NIAP CC EAL4+, FedRAMP",
        "direction":"Low-to-High","max_domains":"3",
        "notes":"Compliance-integrated CDS with continuous monitoring, risk scoring, and automated accreditation workflow."
    },
    "Verizon Federal CDS": {
        "vendor":"Verizon Federal","cert":"DoD IL5, FIPS 140-3",
        "direction":"Bidirectional","max_domains":"4",
        "notes":"Cloud-native CDS service for federal agencies. Integrates with Verizon Managed Security for continuous threat monitoring."
    },
}
FIELD_DOCS = {
    "Record_ID":               ("Unique CDS session/transfer record identifier.",
                                "CDS-{NNNN} sequential"),
    "Source_Domain":           ("Security classification level of the originating domain.",
                                "Categorical — 5 classification levels"),
    "Destination_Domain":      ("Security classification level of the receiving domain.",
                                "Categorical — 5 levels, always ≠ Source"),
    "Direction":               ("Data flow direction relative to classification hierarchy.",
                                "Low-to-High if Dst_Lvl > Src_Lvl; else High-to-Low"),
    "Domain_Level_Src":        ("Numeric security level of source (1=U to 5=TS/SCI).",
                                "Integer 1–5 mapped from Source_Domain"),
    "Domain_Level_Dst":        ("Numeric security level of destination.",
                                "Integer 1–5 mapped from Destination_Domain"),
    "Guard_Product":           ("Commercial CDS guard appliance/software used.",
                                "Categorical — 10 NIAP/DISA certified vendors"),
    "Protocol":                ("Network protocol used for the transfer.",
                                "Categorical — 8 approved protocols"),
    "Data_Type":               ("Type of data traversing the cross-domain boundary.",
                                "Categorical — 8 data categories"),
    "Num_Filters":             ("Number of active filter engines applied to the transfer.",
                                "Uniform Integer(2, 8)"),
    "Filter_Types":            ("Semicolon-separated list of filter techniques applied.",
                                "Random sample of size=Num_Filters from 10 filter types"),
    "Risk_Score":              ("Composite risk of the transfer session (0=safe, 100=critical).",
                                "Normal(μ=70 if H→L else 30, σ=12) clipped [0,100]"),
    "Throughput_Mbps":         ("Effective data throughput after filtering overhead.",
                                "Normal(μ=500, σ=200) clipped ≥0.1 Mbps"),
    "Latency_ms":              ("End-to-end transfer latency including guard processing.",
                                "| Normal(Num_Filters×8, 15) | ms"),
    "False_Positive_Rate_Pct": ("% of benign transfers incorrectly blocked by filters.",
                                "| Normal(0.5, 0.3) | %"),
    "Inspection_Depth":        ("Guard inspection depth score (1–10 scale).",
                                "Uniform(3.0, 10.0)"),
    "Compliance_Score":        ("Composite policy compliance index (0–100).",
                                "(InspDepth/10×40) + ((100−Risk)/100×35) + (1−min(1,FPR))×25"),
    "Certification":           ("Primary accreditation/certification held by the guard.",
                                "Categorical — 9 US/NATO/UK certification standards"),
    "Organization":            ("Owning or operating organisation.",
                                "Categorical — 10 government/alliance organisations"),
    "Transfer_Status":         ("Outcome of the transfer request.",
                                "Categorical — Approved / Pending / Denied / Under Review"),
    "Year":                    ("Year of transfer session record.",
                                "Choice from [2022, 2023, 2024, 2025, 2026]"),
}
# ════════════════════════════════════════════════════════════════════════════
# DATA GENERATION
# ════════════════════════════════════════════════════════════════════════════
def generate_cds_data(n: int = 300, seed: int = 42) -> pd.DataFrame:
    np.random.seed(seed); random.seed(seed)
    rows = []
    for i in range(1, n + 1):
        src = random.choice(DOMAINS)
        dst_pool = [d for d in DOMAINS if d != src]
        dst = random.choice(dst_pool)
        src_lvl, dst_lvl = DOMAIN_LVL[src], DOMAIN_LVL[dst]
        direction = "Low-to-High" if dst_lvl > src_lvl else "High-to-Low"
        risk_mu = 68 if direction == "High-to-Low" else 28
        risk = round(float(np.clip(np.random.normal(risk_mu, 12), 0, 100)), 1)
        n_filt = random.randint(2, 8)
        latency = round(abs(float(np.random.normal(n_filt * 8, 15))), 1)
        tput = round(max(0.1, float(np.random.normal(500, 200))), 2)
        fpr = round(abs(float(np.random.normal(0.5, 0.3))), 3)
        insp = round(float(np.random.uniform(3, 10)), 1)
        compliance = round(
            (insp / 10 * 40) + ((100 - risk) / 100 * 35) + (1 - min(1.0, fpr)) * 25, 1
        )
        rows.append({
            "Record_ID":               f"CDS-{i:04d}",
            "Source_Domain":           src,
            "Destination_Domain":      dst,
            "Direction":               direction,
            "Domain_Level_Src":        src_lvl,
            "Domain_Level_Dst":        dst_lvl,
            "Guard_Product":           random.choice(GUARDS),
            "Protocol":                random.choice(PROTOCOLS),
            "Data_Type":               random.choice(DATA_TYPES),
            "Num_Filters":             n_filt,
            "Filter_Types":            "; ".join(random.sample(FILTER_TYPES,
                                           min(n_filt, len(FILTER_TYPES)))),
            "Risk_Score":              risk,
            "Throughput_Mbps":         tput,
            "Latency_ms":              latency,
            "False_Positive_Rate_Pct": fpr,
            "Inspection_Depth":        insp,
            "Compliance_Score":        compliance,
            "Certification":           random.choice(CERTS),
            "Organization":            random.choice(ORGS),
            "Transfer_Status":         random.choice(STATUS_OPTS),
            "Year":                    random.choice([2022, 2023, 2024, 2025, 2026]),
        })
    return pd.DataFrame(rows)
# ════════════════════════════════════════════════════════════════════════════
# SESSION STATE
# ════════════════════════════════════════════════════════════════════════════
if "cds_df" not in st.session_state:
    st.session_state.cds_df = generate_cds_data(300)
if "cds_n" not in st.session_state:
    st.session_state.cds_n = 300
# ════════════════════════════════════════════════════════════════════════════
# SIDEBAR
# ════════════════════════════════════════════════════════════════════════════
with st.sidebar:
    st.markdown("## ⚙️ Controls")
    n_rec  = st.slider("Records (0–300)", 0, 300, st.session_state.cds_n, step=10)
    seed_v = st.number_input("Random Seed", 0, 9999, 42)
    c1, c2 = st.columns(2)
    with c1:
        if st.button("🔄 Generate", use_container_width=True):
            st.session_state.cds_df = generate_cds_data(n_rec, seed=seed_v)
            st.session_state.cds_n = n_rec
            st.rerun()
    with c2:
        if st.button("🗑️ Reset", use_container_width=True):
            st.session_state.cds_df = generate_cds_data(300)
            st.session_state.cds_n = 300
            st.rerun()
    st.markdown("---")
    st.markdown("### 🔍 Filters")
    df_all = st.session_state.cds_df
    sel_dir    = st.multiselect("Direction",   DIRECTIONS[:2], default=DIRECTIONS[:2])
    sel_guard  = st.multiselect("Guard",       GUARDS, default=GUARDS)
    sel_src    = st.multiselect("Source Domain",  DOMAINS, default=DOMAINS)
    sel_dst    = st.multiselect("Dest Domain",    DOMAINS, default=DOMAINS)
    sel_status = st.multiselect("Status",         STATUS_OPTS, default=STATUS_OPTS)
    st.markdown("---")
    st.markdown("### 📤 Export Format")
    export_fmt = st.selectbox("Format", ["PDF","Word (.docx)","Text (.txt)","JSON","CSV"])
df = st.session_state.cds_df
if len(df):
    mask = (df["Direction"].isin(sel_dir) &
            df["Guard_Product"].isin(sel_guard) &
            df["Source_Domain"].isin(sel_src) &
            df["Destination_Domain"].isin(sel_dst) &
            df["Transfer_Status"].isin(sel_status))
    df_f = df[mask].copy()
else:
    df_f = df.copy()
# ════════════════════════════════════════════════════════════════════════════
# HERO
# ════════════════════════════════════════════════════════════════════════════
st.markdown("""
<div class="knet-hero">
  <p class="knet-title">🔐 KNet Cross Domain Solution (CDS) Intelligence Framework</p>
  <p class="knet-sub">Cross-Domain Security Analytics · Guard Catalogue · Risk Scoring · Policy Compliance</p>
  <span class="knet-badge">🏢 Developed by Randy Singh &nbsp;|&nbsp; Kalsnet (KNet) Consulting Group</span>
</div>
""", unsafe_allow_html=True)
# ════════════════════════════════════════════════════════════════════════════
# TABS
# ════════════════════════════════════════════════════════════════════════════
tabs = st.tabs([
    "🏠 Framework Overview",
    "📊 Data Explorer",
    "📈 Analytics",
    "🗺️ Diagrams & Flowcharts",
    "🛡️ Guard Catalogue",
    "📖 Field Dictionary",
    "📤 Export",
])
# ─────────────────────────────────────────────────────────────────────────────
# TAB 0 — OVERVIEW
# ─────────────────────────────────────────────────────────────────────────────
with tabs[0]:
    st.markdown('<p class="sh">📖 Cross Domain Solution (CDS) Taxonomy</p>',
                unsafe_allow_html=True)
    st.markdown("""
<div class="tax">
A <b>Cross Domain Solution (CDS)</b> is a form of controlled interface that provides the
ability to manually or automatically access and/or transfer information between two or more
differing security domains. CDSs enforce a security policy that prevents information from
flowing between domains unless explicitly authorised — protecting classified and sensitive
government and defence networks.<br><br>
<b>Security Domain Hierarchy (U.S. Classification):</b>
<ul>
  <li><b>Unclassified (U)</b> — Level 1: Public or non-sensitive information</li>
  <li><b>Confidential (C)</b> — Level 2: Disclosure could damage national security</li>
  <li><b>Secret (S)</b> — Level 3: Serious damage to national security if disclosed</li>
  <li><b>Top Secret (TS)</b> — Level 4: Grave damage; most sensitive operational data</li>
  <li><b>TS/SCI</b> — Level 5: Sensitive Compartmented Information; Intelligence sources &amp; methods</li>
</ul>
<br>
<b>Transfer Directions:</b>
<ul>
  <li><b>Low-to-High (L→H)</b>: Data moves from a lower classification to a higher one.
    Lower inherent risk — classified systems accept unclassified data with strict content inspection.</li>
  <li><b>High-to-Low (H→L)</b>: Data moves from a higher classification to a lower one.
    Higher risk — requires deep content filtering, redaction, DLP, and AI/ML anomaly detection
    to prevent data spillage.</li>
  <li><b>Bidirectional</b>: Simultaneous or alternating transfer in both directions via
    separate logical channels with independent filter stacks.</li>
</ul>
<br>
<b>Governing Standards &amp; Authorities:</b>
<ul>
  <li><b>NCDSMO</b> — National Cross Domain Strategy &amp; Management Office (oversees U.S. CDS policy)</li>
  <li><b>NIAP CC EAL4–EAL6+</b> — National Information Assurance Partnership Common Criteria evaluation</li>
  <li><b>DISA APL</b> — Defense Information Systems Agency Approved Products List</li>
  <li><b>NSA CSfC / Suite B</b> — Commercial Solutions for Classified program</li>
  <li><b>FIPS 140-3</b> — Federal cryptographic module validation standard</li>
  <li><b>DoD IL5 / IL6</b> — Department of Defense Impact Level 5 (Secret) and 6 (TS) cloud standards</li>
</ul>
</div>
""", unsafe_allow_html=True)
    st.markdown('<p class="sh">🏢 KNet CDS Framework Pillars</p>', unsafe_allow_html=True)
    st.markdown("""
<div class="tax">
The <b>KNet CDS Intelligence Framework</b> by <b>Randy Singh</b> at
<b>Kalsnet (KNet) Consulting Group</b> structures CDS selection, risk assessment, and
compliance monitoring across five pillars:
<ol>
  <li><b>Domain Risk Profiling</b> — Quantified risk scoring for every L→H and H→L transfer</li>
  <li><b>Guard Selection Intelligence</b> — Vendor-agnostic scoring across 10 certified guard products</li>
  <li><b>Filter Depth Analysis</b> — Measuring inspection coverage: AV, DLP, AI/ML, steganography</li>
  <li><b>Compliance Benchmarking</b> — Mapping transfers against NIAP, DISA APL, FIPS, and DoD IL standards</li>
  <li><b>Transfer Lifecycle Monitoring</b> — Throughput, latency, false-positive tracking per session</li>
</ol>
</div>
""", unsafe_allow_html=True)
    # KPI Cards
    st.markdown('<p class="sh">📊 Live KPI Summary</p>', unsafe_allow_html=True)
    if len(df_f):
        approved  = len(df_f[df_f["Transfer_Status"]=="Approved"])
        denied    = len(df_f[df_f["Transfer_Status"]=="Denied"])
        h2l_ct    = len(df_f[df_f["Direction"]=="High-to-Low"])
        l2h_ct    = len(df_f[df_f["Direction"]=="Low-to-High"])
        cols = st.columns(7)
        kpis = [
            (df_f["Record_ID"].count(), "Total Records"),
            (f'{df_f["Risk_Score"].mean():.1f}',        "Avg Risk Score"),
            (f'{df_f["Compliance_Score"].mean():.1f}',  "Avg Compliance"),
            (f'{df_f["Throughput_Mbps"].mean():.0f}',   "Avg Throughput Mbps"),
            (f'{df_f["Latency_ms"].mean():.1f} ms',     "Avg Latency"),
            (h2l_ct, "High→Low Transfers"),
            (l2h_ct, "Low→High Transfers"),
        ]
        for col, (val, lbl) in zip(cols, kpis):
            with col:
                st.markdown(f"""<div class="metric-card">
                  <div class="metric-val">{val}</div>
                  <div class="metric-lbl">{lbl}</div></div>""",
                            unsafe_allow_html=True)
# ─────────────────────────────────────────────────────────────────────────────
# TAB 1 — DATA EXPLORER
# ─────────────────────────────────────────────────────────────────────────────
with tabs[1]:
    st.markdown('<p class="sh">📋 CDS Transfer Dataset</p>', unsafe_allow_html=True)
    st.info(f"Showing **{len(df_f)}** of **{len(df)}** records after filters.")
    # Colour risk score column
    def colour_risk(val):
        if val >= 70: return "color:#FF4444;font-weight:700"
        elif val >= 40: return "color:#FFA726;font-weight:700"
        return "color:#66BB6A;font-weight:700"
    st.dataframe(df_f, use_container_width=True, height=500,
                 column_config={
                     "Risk_Score": st.column_config.ProgressColumn(
                         "Risk Score", min_value=0, max_value=100, format="%.1f"),
                     "Compliance_Score": st.column_config.ProgressColumn(
                         "Compliance Score", min_value=0, max_value=100, format="%.1f"),
                 })
    st.markdown('<p class="sh">📈 Descriptive Statistics</p>', unsafe_allow_html=True)
    num_c = df_f.select_dtypes(include=np.number).columns.tolist()
    st.dataframe(df_f[num_c].describe().T.style.format("{:.3f}"),
                 use_container_width=True)
# ─────────────────────────────────────────────────────────────────────────────
# TAB 2 — ANALYTICS
# ─────────────────────────────────────────────────────────────────────────────
with tabs[2]:
    if len(df_f) == 0:
        st.warning("No data — adjust filters.")
    else:
        DARK = dict(paper_bgcolor="#0A1628", plot_bgcolor="#0D2148")
        # Risk Distribution L→H vs H→L
        st.markdown('<p class="sh">Risk Score Distribution: Low→High vs High→Low</p>',
                    unsafe_allow_html=True)
        fig1 = px.histogram(df_f, x="Risk_Score", color="Direction",
                            barmode="overlay", nbins=40, opacity=.75,
                            color_discrete_map={"Low-to-High":"#1E90FF",
                                                "High-to-Low":"#FF4444"},
                            template="plotly_dark")
        fig1.update_layout(**DARK)
        st.plotly_chart(fig1, use_container_width=True)
        col_a, col_b = st.columns(2)
        with col_a:
            st.markdown('<p class="sh">Risk by Guard Product</p>', unsafe_allow_html=True)
            fig2 = px.box(df_f, x="Guard_Product", y="Risk_Score", color="Direction",
                          color_discrete_map={"Low-to-High":"#1E90FF","High-to-Low":"#FF4444"},
                          template="plotly_dark")
            fig2.update_layout(**DARK, xaxis_tickangle=-35)
            st.plotly_chart(fig2, use_container_width=True)
        with col_b:
            st.markdown('<p class="sh">Transfer Status Breakdown</p>',
                        unsafe_allow_html=True)
            sc = df_f["Transfer_Status"].value_counts().reset_index()
            sc.columns = ["Status","Count"]
            fig3 = px.pie(sc, names="Status", values="Count",
                          color_discrete_sequence=["#1E90FF","#FF4444","#FFA726","#64B5F6"],
                          template="plotly_dark")
            fig3.update_layout(**DARK)
            st.plotly_chart(fig3, use_container_width=True)
        st.markdown('<p class="sh">Compliance Score vs Risk Score (by Direction)</p>',
                    unsafe_allow_html=True)
        fig4 = px.scatter(df_f, x="Risk_Score", y="Compliance_Score",
                          color="Direction", size="Num_Filters",
                          hover_data=["Record_ID","Guard_Product","Source_Domain","Destination_Domain"],
                          color_discrete_map={"Low-to-High":"#1E90FF","High-to-Low":"#FF4444"},
                          template="plotly_dark")
        fig4.update_layout(**DARK)
        st.plotly_chart(fig4, use_container_width=True)
        col_c, col_d = st.columns(2)
        with col_c:
            st.markdown('<p class="sh">Avg Throughput by Data Type</p>',
                        unsafe_allow_html=True)
            fig5 = px.bar(df_f.groupby("Data_Type")["Throughput_Mbps"].mean()
                          .reset_index().sort_values("Throughput_Mbps", ascending=False),
                          x="Data_Type", y="Throughput_Mbps",
                          color="Throughput_Mbps", color_continuous_scale="Blues",
                          template="plotly_dark")
            fig5.update_layout(**DARK, xaxis_tickangle=-35, showlegend=False)
            st.plotly_chart(fig5, use_container_width=True)
        with col_d:
            st.markdown('<p class="sh">Latency vs Num Filters</p>', unsafe_allow_html=True)
            fig6 = px.scatter(df_f, x="Num_Filters", y="Latency_ms",
                              color="Direction",
                              color_discrete_map={"Low-to-High":"#1E90FF","High-to-Low":"#FF4444"},
                              template="plotly_dark")
            # Manual OLS trendline (avoids the "statsmodels" dependency that
            # px.scatter(..., trendline="ols") requires and that caused the
            # ModuleNotFoundError on Streamlit Cloud).
            for _dir, _color in [("Low-to-High", "#1E90FF"), ("High-to-Low", "#FF4444")]:
                _sub = df_f[df_f["Direction"] == _dir]
                if len(_sub) >= 2:
                    _x = _sub["Num_Filters"].to_numpy(dtype=float)
                    _y = _sub["Latency_ms"].to_numpy(dtype=float)
                    _slope, _intercept = np.polyfit(_x, _y, 1)
                    _x_line = np.linspace(_x.min(), _x.max(), 50)
                    _y_line = _slope * _x_line + _intercept
                    fig6.add_trace(go.Scatter(
                        x=_x_line, y=_y_line, mode="lines",
                        name=f"{_dir} OLS trend",
                        line=dict(color=_color, dash="dash", width=2),
                    ))
            fig6.update_layout(**DARK)
            st.plotly_chart(fig6, use_container_width=True)
        st.markdown('<p class="sh">Heatmap: Avg Risk Score (Source → Destination Domain)</p>',
                    unsafe_allow_html=True)
        pivot = df_f.pivot_table(values="Risk_Score",
                                  index="Source_Domain",
                                  columns="Destination_Domain",
                                  aggfunc="mean").fillna(0)
        fig7 = px.imshow(pivot, color_continuous_scale="RdBu_r",
                         text_auto=".1f", template="plotly_dark")
        fig7.update_layout(**DARK)
        st.plotly_chart(fig7, use_container_width=True)
        st.markdown('<p class="sh">Guard Usage Count</p>', unsafe_allow_html=True)
        gc = df_f["Guard_Product"].value_counts().reset_index()
        gc.columns = ["Guard","Count"]
        fig8 = px.bar(gc, x="Guard", y="Count", color="Count",
                      color_continuous_scale="Blues", template="plotly_dark")
        fig8.update_layout(**DARK, xaxis_tickangle=-30, showlegend=False)
        st.plotly_chart(fig8, use_container_width=True)
# ─────────────────────────────────────────────────────────────────────────────
# TAB 3 — DIAGRAMS (Mermaid via HTML component)
# ─────────────────────────────────────────────────────────────────────────────
MERMAID_HTML = """<!DOCTYPE html><html><head>
<script src="https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.min.js"></script>
<style>
  body{background:#0A1628;color:#90CAF9;font-family:Inter,sans-serif;padding:18px}
  h3{color:#1E90FF;border-left:4px solid #1E90FF;padding-left:10px;margin:24px 0 10px}
  .m{background:#0D2148;border-radius:10px;padding:20px;margin-bottom:20px;overflow-x:auto}
</style></head><body>
<script>mermaid.initialize({theme:'dark',themeVariables:{primaryColor:'#1E90FF',
  primaryTextColor:'#E3F2FD',primaryBorderColor:'#1565C0',lineColor:'#64B5F6',
  background:'#0D2148',nodeBorder:'#1E90FF'}});</script>
<h3>① CDS Reference Architecture</h3>
<div class="m"><div class="mermaid">
graph LR
  LO[🟢 Low Domain\nUnclassified/Confidential]-->GW[🔐 Cross Domain Guard\nFilter Stack]
  GW-->HI[🔴 High Domain\nSecret / TS / TS-SCI]
  GW-->AV[Anti-Virus\nEngine]
  GW-->DLP[Data Loss\nPrevention DLP]
  GW-->CT[Content\nInspection]
  GW-->ML[AI/ML Anomaly\nDetection]
  GW-->LOG[Audit Log\nSIEM]
</div></div>
<h3>② Low-to-High Data Flow</h3>
<div class="m"><div class="mermaid">
flowchart LR
  SRC[📄 Source\nLow Domain\nUnclassified]-->REQ{Transfer\nRequest}
  REQ-->P1[Format\nValidation]
  P1-->P2[Whitelist\nCheck]
  P2-->P3[Anti-Virus\nScan]
  P3-->P4[Content\nInspection]
  P4-->PASS{Pass?}
  PASS-->|Yes|DELIV[✅ Deliver to\nHigh Domain]
  PASS-->|No|BLOCK[🚫 Block &\nAlert]
  BLOCK-->LOG2[Audit Trail]
  DELIV-->LOG2
</div></div>
<h3>③ High-to-Low Data Flow (Critical Path)</h3>
<div class="m"><div class="mermaid">
flowchart LR
  SRC2[🔴 Source\nHigh Domain\nTS or TS-SCI]-->REQ2{Transfer\nRequest}
  REQ2-->P10[Classification\nLabel Check]
  P10-->P11[Redaction\nEngine]
  P11-->P12[Steganography\nDetection]
  P12-->P13[DLP Deep\nInspection]
  P13-->P14[AI/ML\nAnomaly Score]
  P14-->P15[Human\nReview Queue]
  P15-->PASS2{Approved?}
  PASS2-->|Yes|DELIV2[✅ Deliver to\nLow Domain]
  PASS2-->|No|BLOCK2[🚫 Deny &\nQuarantine]
  DELIV2 & BLOCK2 --> AUDIT[📋 SIEM Audit\nFIPS-compliant]
</div></div>
<h3>④ KNet CDS Guard Selection Flowchart</h3>
<div class="m"><div class="mermaid">
flowchart TD
  START([🚀 CDS Requirement])-->DIR{Direction?}
  DIR-->|Low-to-High|L2H[eMBB/File Transfer\nLower Risk Profile]
  DIR-->|High-to-Low|H2L[High Risk\nRequires deep filters]
  L2H-->CERT1{Certification\nNeeded?}
  H2L-->CERT2{Certification\nNeeded?}
  CERT1-->|NIAP EAL4|FP[Forcepoint CDS\nor Radiant Mercury]
  CERT1-->|DoD IL5|L3[L3Harris CDG]
  CERT2-->|NIAP EAL6|BAE[BAE Trusted Guard\nor Raytheon SurePath]
  CERT2-->|IC / TS-SCI|EF[Everfox High Side]
  FP & L3 & BAE & EF --> KPI2[Measure:\nRisk Score\nCompliance\nLatency\nFPR]
  KPI2-->PASS3{Pass\nThreshold?}
  PASS3-->|Yes|ACCRED[Submit for\nNCDSMO Accreditation]
  PASS3-->|No|TUNE[Tune Filter Stack:\nIncrease depth\nAdjust whitelist\nEnable AI-ML]
  TUNE-->KPI2
  ACCRED-->DEPLOY([✅ Operational Deployment])
</div></div>
<h3>⑤ Compliance Score Formula</h3>
<div class="m"><div class="mermaid">
flowchart LR
  A[Inspection\nDepth /10]-->|×40|W1[40% Weight]
  B[Risk Score\n0-100]-->|100-Risk ÷100 ×35|W2[35% Weight]
  C[False Positive\nRate %]-->|1-min1-FPR ×25|W3[25% Weight]
  W1 & W2 & W3-->SUM[Σ Compliance\nScore 0–100]
</div></div>
<h3>⑥ Domain Classification Hierarchy</h3>
<div class="m"><div class="mermaid">
graph BT
  U[🟢 Unclassified\nLevel 1]-->C[🔵 Confidential\nLevel 2]
  C-->S[🟡 Secret\nLevel 3]
  S-->TS[🟠 Top Secret\nLevel 4]
  TS-->SCI[🔴 TS/SCI\nLevel 5\nHighest]
  style U fill:#1B5E20,color:#fff
  style C fill:#0D47A1,color:#fff
  style S fill:#F9A825,color:#000
  style TS fill:#E65100,color:#fff
  style SCI fill:#B71C1C,color:#fff
</div></div>
<h3>⑦ CDS Deployment Lifecycle</h3>
<div class="m"><div class="mermaid">
timeline
  title KNet CDS Deployment Lifecycle
  Phase 1 : Requirements Definition
           : Domain Boundary Mapping
           : Data Flow Inventory
  Phase 2 : Guard Product Selection
           : NIAP/DISA Evaluation
           : Filter Stack Design
  Phase 3 : Lab Testing
           : Red Team Assessment
           : NCDSMO Submission
  Phase 4 : Accreditation
           : Operational Deployment
           : SIEM Integration
  Phase 5 : Continuous Monitoring
           : Annual Re-accreditation
           : AI/ML Tuning
</div></div>
</body></html>"""
with tabs[3]:
    st.markdown('<p class="sh">CDS Architecture Diagrams & Flowcharts</p>',
                unsafe_allow_html=True)
    st.info("ℹ️ 7 Mermaid diagrams — architecture, L→H & H→L flows, guard selection, "
            "compliance formula, classification hierarchy, and deployment lifecycle.")
    st.components.v1.html(MERMAID_HTML, height=3200, scrolling=True)
# ─────────────────────────────────────────────────────────────────────────────
# TAB 4 — GUARD CATALOGUE
# ─────────────────────────────────────────────────────────────────────────────
with tabs[4]:
    st.markdown('<p class="sh">🛡️ Commercial CDS Guard Catalogue</p>',
                unsafe_allow_html=True)
    search = st.text_input("🔍 Search guard name or vendor…", "")
    for gname, info in GUARD_DETAIL.items():
        if search.lower() in gname.lower() or search.lower() in info["vendor"].lower():
            with st.expander(f"🛡️ {gname}  —  {info['vendor']}"):
                c1, c2 = st.columns([2, 1])
                with c1:
                    st.markdown(f"**Description:** {info['notes']}")
                    st.markdown(f"**Certifications:** `{info['cert']}`")
                with c2:
                    st.markdown(f"**Direction:** `{info['direction']}`")
                    st.markdown(f"**Max Domains:** `{info['max_domains']}`")
                # Mini chart from filtered data
                gdf = df_f[df_f["Guard_Product"] == gname]
                if len(gdf) > 0:
                    m1, m2, m3, m4 = st.columns(4)
                    m1.metric("Avg Risk",    f"{gdf['Risk_Score'].mean():.1f}")
                    m2.metric("Avg Comply",  f"{gdf['Compliance_Score'].mean():.1f}")
                    m3.metric("Avg Tput",    f"{gdf['Throughput_Mbps'].mean():.0f} Mbps")
                    m4.metric("Records",     len(gdf))
# ─────────────────────────────────────────────────────────────────────────────
# TAB 5 — FIELD DICTIONARY
# ─────────────────────────────────────────────────────────────────────────────
with tabs[5]:
    st.markdown('<p class="sh">📖 Field Definitions & Formulas</p>',
                unsafe_allow_html=True)
    for fname, (desc, formula) in FIELD_DOCS.items():
        st.markdown(f"""
<div class="field">
  <span class="field-name">🔹 {fname}</span><br>
  {desc}<br>
  <span class="formula">📐 {formula}</span>
</div>""", unsafe_allow_html=True)
# ─────────────────────────────────────────────────────────────────────────────
# TAB 6 — EXPORT
# ─────────────────────────────────────────────────────────────────────────────
with tabs[6]:
    st.markdown('<p class="sh">📤 Export Filtered Dataset</p>', unsafe_allow_html=True)
    st.info(f"Exporting **{len(df_f)}** records in **{export_fmt}** format.")
    if st.button("⬇️ Generate & Download", use_container_width=True):
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        # ── CSV ──────────────────────────────────────────────────────────────
        if export_fmt == "CSV":
            buf = io.StringIO()
            df_f.to_csv(buf, index=False)
            st.download_button("⬇️ Download CSV", buf.getvalue().encode(),
                               f"KNetCDS_{ts}.csv", "text/csv")
        # ── JSON ─────────────────────────────────────────────────────────────
        elif export_fmt == "JSON":
            payload = {
                "metadata": {
                    "framework": "KNet CDS Intelligence Framework",
                    "author": "Randy Singh",
                    "organization": "Kalsnet (KNet) Consulting Group",
                    "generated": datetime.now().isoformat(),
                    "records": len(df_f),
                },
                "taxonomy": {
                    "Low-to-High": "Lower classification → Higher classification (lower risk)",
                    "High-to-Low": "Higher classification → Lower classification (higher risk, deep filter required)",
                    "Domains": {d: l for d, l in DOMAIN_LVL.items()},
                },
                "data": df_f.to_dict(orient="records"),
            }
            buf = io.BytesIO(json.dumps(payload, indent=2).encode())
            st.download_button("⬇️ Download JSON", buf, f"KNetCDS_{ts}.json",
                               "application/json")
        # ── TEXT ─────────────────────────────────────────────────────────────
        elif export_fmt == "Text (.txt)":
            lines = [
                "=" * 72,
                "  KNet Cross Domain Solution (CDS) Intelligence Framework",
                "  Randy Singh | Kalsnet (KNet) Consulting Group",
                f"  Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
                f"  Records: {len(df_f)}",
                "=" * 72,
                "",
                "CDS TAXONOMY",
                "-" * 40,
                "Low-to-High  : Data flows from lower to higher classification (Lower risk)",
                "High-to-Low  : Data flows from higher to lower classification (Higher risk)",
                "",
                "DOMAIN LEVELS",
                "-" * 40,
            ]
            for d, l in DOMAIN_LVL.items():
                lines.append(f"  Level {l}: {d}")
            lines += ["", "DATA", "-" * 40, df_f.to_string(index=False)]
            buf = io.BytesIO("\n".join(lines).encode())
            st.download_button("⬇️ Download TXT", buf, f"KNetCDS_{ts}.txt",
                               "text/plain")
        # ── WORD ─────────────────────────────────────────────────────────────
        elif export_fmt == "Word (.docx)":
            if not DOCX_OK:
                st.error("python-docx not installed.")
            else:
                doc = Document()
                t = doc.add_heading("KNet CDS Intelligence Framework", 0)
                t.runs[0].font.color.rgb = RGBColor(0x1E, 0x90, 0xFF)
                doc.add_paragraph("Developed by Randy Singh | Kalsnet (KNet) Consulting Group")
                doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
                                  f" | Records: {len(df_f)}")
                doc.add_heading("CDS Taxonomy", level=1)
                for line in [
                    "Low-to-High: Data from lower classification to higher — lower inherent risk.",
                    "High-to-Low: Data from higher classification to lower — requires deep content filtering, DLP, AI/ML anomaly detection.",
                    "Bidirectional: Both directions on independent filter stacks.",
                ]:
                    doc.add_paragraph(line, style="List Bullet")
                doc.add_heading("Domain Levels", level=1)
                for d, l in DOMAIN_LVL.items():
                    doc.add_paragraph(f"Level {l}: {d}", style="List Bullet")
                doc.add_heading("Field Definitions", level=1)
                for fname, (desc, formula) in FIELD_DOCS.items():
                    p = doc.add_paragraph()
                    r = p.add_run(fname + ": ")
                    r.bold = True
                    r.font.color.rgb = RGBColor(0x1E, 0x90, 0xFF)
                    p.add_run(f"{desc}  [Formula: {formula}]")
                doc.add_heading("Guard Catalogue Summary", level=1)
                for gname, info in GUARD_DETAIL.items():
                    p = doc.add_paragraph()
                    p.add_run(gname + ": ").bold = True
                    p.add_run(f"{info['notes']} Certs: {info['cert']}. Direction: {info['direction']}.")
                doc.add_heading("Filtered Transfer Data", level=1)
                show_c = ["Record_ID","Source_Domain","Destination_Domain","Direction",
                          "Guard_Product","Risk_Score","Compliance_Score",
                          "Transfer_Status","Throughput_Mbps","Latency_ms"]
                tbl = doc.add_table(rows=1, cols=len(show_c))
                tbl.style = "Light Shading Accent 1"
                for i, c in enumerate(show_c):
                    tbl.rows[0].cells[i].text = c
                for _, row in df_f.head(200).iterrows():
                    cells = tbl.add_row().cells
                    for i, c in enumerate(show_c):
                        cells[i].text = str(row[c])
                buf = io.BytesIO()
                doc.save(buf); buf.seek(0)
                st.download_button("⬇️ Download DOCX", buf, f"KNetCDS_{ts}.docx",
                                   "application/vnd.openxmlformats-officedocument"
                                   ".wordprocessingml.document")
        # ── PDF ──────────────────────────────────────────────────────────────
        elif export_fmt == "PDF":
            if not REPORTLAB_OK:
                st.error("reportlab not installed.")
            else:
                buf = io.BytesIO()
                pdoc = SimpleDocTemplate(buf, pagesize=landscape(letter),
                                         rightMargin=28, leftMargin=28,
                                         topMargin=28, bottomMargin=28)
                styles = getSampleStyleSheet()
                blue = colors.HexColor("#1E90FF")
                red  = colors.HexColor("#FF4444")
                ts_style = ParagraphStyle("T", parent=styles["Title"],
                                          textColor=blue, fontSize=17, spaceAfter=4)
                h1 = ParagraphStyle("H1", parent=styles["Heading1"],
                                    textColor=blue, fontSize=12, spaceBefore=10)
                body = styles["Normal"]
                story = [
                    Paragraph("KNet Cross Domain Solution (CDS) Intelligence Framework", ts_style),
                    Paragraph("Randy Singh | Kalsnet (KNet) Consulting Group", body),
                    Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
                              f"  |  Records: {len(df_f)}", body),
                    HRFlowable(width="100%", color=blue, spaceAfter=8),
                    Paragraph("CDS Transfer Directions", h1),
                    Paragraph("Low-to-High: Lower → Higher classification. Lower inherent risk.", body),
                    Paragraph("High-to-Low: Higher → Lower classification. High risk — requires deep content filtering, DLP, AI/ML.", body),
                    Spacer(1, 8),
                    Paragraph("Transfer Data (first 150 rows)", h1),
                ]
                show_c = ["Record_ID","Source_Domain","Destination_Domain","Direction",
                          "Guard_Product","Risk_Score","Compliance_Score",
                          "Transfer_Status","Throughput_Mbps","Latency_ms","Year"]
                wids = [65,80,80,70,85,48,58,65,60,55,38]
                tdata = [show_c] + [
                    [str(r[c])[:20] for c in show_c]
                    for _, r in df_f.head(150).iterrows()
                ]
                ptbl = Table(tdata, colWidths=wids, repeatRows=1)
                ptbl.setStyle(TableStyle([
                    ("BACKGROUND",  (0,0),(-1,0), blue),
                    ("TEXTCOLOR",   (0,0),(-1,0), colors.white),
                    ("FONTSIZE",    (0,0),(-1,0), 7),
                    ("FONTSIZE",    (0,1),(-1,-1), 6),
                    ("ROWBACKGROUNDS",(0,1),(-1,-1),
                     [colors.HexColor("#0D2148"),colors.HexColor("#0A1628")]),
                    ("TEXTCOLOR",   (0,1),(-1,-1), colors.HexColor("#CFE2FF")),
                    ("GRID",        (0,0),(-1,-1), 0.3, colors.HexColor("#1565C0")),
                    ("VALIGN",      (0,0),(-1,-1),"MIDDLE"),
                ]))
                story.append(ptbl)
                pdoc.build(story)
                buf.seek(0)
                st.download_button("⬇️ Download PDF", buf, f"KNetCDS_{ts}.pdf",
                                   "application/pdf")
    st.markdown("---")
    st.markdown("**© KNet CDS Intelligence Framework | Randy Singh | "
                "Kalsnet (KNet) Consulting Group**")
