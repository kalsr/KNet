

# =========================================================
# KNet PolicyForge AI
# Developed by Randy Singh from Kalsnet (KNet) Consulting Group
# =========================================================
#
# REQUIREMENTS:
# pip install streamlit groq pandas reportlab python-docx
#
# RUN:
# streamlit run knet_policyforge_ai_enhanced.py
# =========================================================

import streamlit as st
from groq import Groq
import json
import pandas as pd
from io import BytesIO
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from datetime import datetime

# =========================================================
# PAGE CONFIG
# =========================================================

st.set_page_config(
    page_title="KNet PolicyForge AI",
    layout="wide",
    page_icon="🛡️"
)

# =========================================================
# CUSTOM CSS
# =========================================================

st.markdown("""
<style>

.main {
    background-color: #0E1117;
}

h1, h2, h3, h4 {
    color: #4FC3F7;
}

.stButton>button {
    background-color: #1F6FEB;
    color: white;
    border-radius: 10px;
    padding: 0.5rem 1rem;
    border: none;
    font-weight: 600;
}

.stTextInput>div>div>input,
.stTextArea textarea,
[data-baseweb="select"] {
    border-radius: 8px;
}

.policy-header {
    text-align: center;
    padding: 25px;
    background: linear-gradient(135deg, #111827, #1F2937);
    border-radius: 18px;
    margin-bottom: 20px;
    border: 1px solid #374151;
}

.policy-title {
    color: #00BFFF;
    font-size: 3.2rem;
    font-weight: 900;
    margin: 0;
    text-shadow: 0 0 8px rgba(0,191,255,0.7);
}

.policy-subtitle {
    color: #D1D5DB;
    font-size: 1.1rem;
    font-weight: 600;
    margin-top: 10px;
}

.footer-box {
    background-color: #111827;
    padding: 20px;
    border-radius: 10px;
    margin-top: 20px;
    border: 1px solid #374151;
}

</style>
""", unsafe_allow_html=True)

# =========================================================
# HEADER
# =========================================================

st.markdown("""
<div class="policy-header">
    <div class="policy-title">🛡️ KNet PolicyForge AI</div>
    <div class="policy-subtitle">
        Developed by Randy Singh from Kalsnet (KNet) Consulting Group
    </div>
</div>
""", unsafe_allow_html=True)

# =========================================================
# SIDEBAR
# =========================================================

st.sidebar.title("⚙️ Configuration")

groq_api_key = st.sidebar.text_input("Groq API Key", type="password")

st.sidebar.markdown("""
### How to Get Groq API Key
1. Visit https://console.groq.com
2. Create a free account
3. Navigate to API Keys
4. Generate a new API key
5. Paste the key here
""")

model_name = st.sidebar.selectbox(
    "Select Groq Model",
    [
        "llama-3.3-70b-versatile",
        "mixtral-8x7b-32768",
        "gemma2-9b-it"
    ]
)

policy_area = st.sidebar.selectbox(
    "Policy Area",
    [
        "Cybersecurity Policy",
        "Acceptable Use Policy",
        "Remote Work Policy",
        "AI Governance Policy",
        "Incident Response Policy",
        "Disaster Recovery Policy",
        "Business Continuity Policy",
        "Data Retention Policy",
        "Privacy Policy",
        "HR Policy",
        "Zero Trust Policy",
        "Vendor Management Policy"
    ]
)

organization_type = st.sidebar.selectbox(
    "Organization Type",
    [
        "Government",
        "Enterprise",
        "Healthcare",
        "Financial",
        "Educational",
        "Small Business",
        "Technology Company"
    ]
)

detail_level = st.sidebar.selectbox(
    "Policy Detail Level",
    ["Basic", "Standard", "Advanced", "Enterprise Grade"]
)

# =========================================================
# DROPDOWN OPTIONS
# =========================================================

COMPLIANCE_OPTIONS = [
    "NIST CSF",
    "NIST SP 800-53",
    "NIST 800-171",
    "ISO 27001",
    "ISO 27701",
    "HIPAA",
    "GDPR",
    "PCI DSS",
    "SOC 2",
    "CMMC 2.0",
    "FedRAMP",
    "FISMA",
    "CJIS",
    "SOX",
    "GLBA"
]

POLICY_OBJECTIVE_OPTIONS = [
    "Protect Confidential Information",
    "Ensure Regulatory Compliance",
    "Reduce Cybersecurity Risk",
    "Define Roles and Responsibilities",
    "Establish Governance and Oversight",
    "Support Business Continuity",
    "Standardize Security Controls",
    "Strengthen Incident Response",
    "Improve Data Privacy",
    "Enable Zero Trust Architecture"
]

RISK_LEVEL_OPTIONS = [
    "Very Low",
    "Low",
    "Moderate",
    "High",
    "Critical",
    "Severe"
]

ADDITIONAL_REQUIREMENT_OPTIONS = [
    "Executive Approval Required",
    "Annual Policy Review",
    "Employee Training",
    "Audit Logging",
    "Encryption at Rest",
    "Encryption in Transit",
    "Multi-Factor Authentication",
    "Role-Based Access Control",
    "Third-Party Risk Assessment",
    "Continuous Monitoring",
    "Incident Reporting",
    "Records Retention",
    "Legal Review",
    "Board Oversight"
]

# =========================================================
# MAIN INPUTS
# =========================================================

col1, col2 = st.columns(2)

with col1:
    organization_name = st.text_input(
        "Organization Name",
        placeholder="Enter organization name"
    )

    industry = st.text_input(
        "Industry",
        placeholder="Healthcare, Finance, Government, etc."
    )

    compliance_selection = st.multiselect(
        "Compliance Requirements",
        COMPLIANCE_OPTIONS,
        default=["NIST SP 800-53"]
    )

    compliance_requirements = ", ".join(compliance_selection)

with col2:
    policy_objective = st.selectbox(
        "Policy Objective",
        POLICY_OBJECTIVE_OPTIONS
    )

    risk_level = st.selectbox(
        "Risk Level",
        RISK_LEVEL_OPTIONS,
        index=3
    )

    additional_selection = st.multiselect(
        "Additional Requirements",
        ADDITIONAL_REQUIREMENT_OPTIONS,
        default=["Annual Policy Review", "Employee Training"]
    )

    additional_requirements = ", ".join(additional_selection)

# =========================================================
# PROMPT TEMPLATE
# =========================================================

def build_prompt():
    return f"""
Generate a highly professional enterprise-grade policy document.

Policy Area:
{policy_area}

Organization Name:
{organization_name}

Organization Type:
{organization_type}

Industry:
{industry}

Compliance Requirements:
{compliance_requirements}

Policy Objective:
{policy_objective}

Risk Level:
{risk_level}

Detail Level:
{detail_level}

Additional Requirements:
{additional_requirements}

The policy must include:

1. Policy Title
2. Purpose
3. Scope
4. Definitions
5. Policy Statements
6. Roles and Responsibilities
7. Standards and Controls
8. Procedures
9. Exceptions
10. Compliance Requirements
11. Enforcement
12. Review Schedule
13. Approval Section
14. Version History

Write in formal professional enterprise policy format.
"""

# =========================================================
# GROQ GENERATION
# =========================================================

if st.button("🚀 Generate Policy", use_container_width=True):
    if not groq_api_key:
        st.error("Please enter Groq API Key.")
    else:
        try:
            client = Groq(api_key=groq_api_key)

            with st.spinner("Generating professional policy..."):
                response = client.chat.completions.create(
                    model=model_name,
                    messages=[
                        {
                            "role": "system",
                            "content": "You are an enterprise policy development expert."
                        },
                        {
                            "role": "user",
                            "content": build_prompt()
                        }
                    ],
                    temperature=0.3,
                    max_tokens=6000
                )

                st.session_state["generated_policy"] = (
                    response.choices[0].message.content
                )

        except Exception as e:
            st.error(f"Error: {e}")

# =========================================================
# DISPLAY GENERATED POLICY
# =========================================================

if "generated_policy" in st.session_state:
    st.subheader("📄 Generated Policy")

    edited_policy = st.text_area(
        "Policy Editor",
        value=st.session_state["generated_policy"],
        height=700
    )

    # =====================================================
    # EXPORT FUNCTIONS
    # =====================================================

    def generate_pdf(text):
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer)
        styles = getSampleStyleSheet()
        story = []

        for line in text.split("\n"):
            if line.strip():
                story.append(Paragraph(line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'), styles['BodyText']))
                story.append(Spacer(1, 6))

        doc.build(story)
        buffer.seek(0)
        return buffer

    def generate_docx(text):
        doc = Document()
        doc.add_heading("Generated Policy", level=1)
        doc.add_paragraph(text)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def generate_json(text):
        policy_json = {
            "generated_date": str(datetime.now()),
            "policy_area": policy_area,
            "organization_name": organization_name,
            "content": text
        }
        return json.dumps(policy_json, indent=4)

    def generate_csv(text):
        rows = [{"Policy_Content": line} for line in text.split("\n")]
        df = pd.DataFrame(rows)
        return df.to_csv(index=False).encode("utf-8")

    # =====================================================
    # EXPORT BUTTONS
    # =====================================================

    st.subheader("⬇️ Export Policy")

    c1, c2, c3, c4 = st.columns(4)

    with c1:
        st.download_button(
            label="📄 Download PDF",
            data=generate_pdf(edited_policy),
            file_name="policy.pdf",
            mime="application/pdf",
            use_container_width=True
        )

    with c2:
        st.download_button(
            label="📝 Download DOCX",
            data=generate_docx(edited_policy),
            file_name="policy.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

    with c3:
        st.download_button(
            label="📊 Download CSV",
            data=generate_csv(edited_policy),
            file_name="policy.csv",
            mime="text/csv",
            use_container_width=True
        )

    with c4:
        st.download_button(
            label="🧾 Download JSON",
            data=generate_json(edited_policy),
            file_name="policy.json",
            mime="application/json",
            use_container_width=True
        )

# =========================================================
# FILE UPLOAD SECTION
# =========================================================

st.subheader("📤 Upload Existing Policy")

uploaded_file = st.file_uploader(
    "Upload TXT Policy File",
    type=["txt"]
)

if uploaded_file is not None:
    content = uploaded_file.read().decode("utf-8")

    st.text_area(
        "Uploaded Policy Content",
        value=content,
        height=300
    )

# =========================================================
# FOOTER
# =========================================================

st.markdown("""
<div class="footer-box">

<h4>About KNet PolicyForge AI</h4>

<p>
This enterprise-grade AI policy generation platform helps organizations
create professional governance, cybersecurity, compliance,
risk management, HR, and operational policies using advanced AI models.
</p>

<b>Features:</b>
<ul>
<li>AI-powered policy development</li>
<li>Enterprise-grade formatting</li>
<li>Multi-format exports</li>
<li>Professional GUI</li>
<li>Editable policy workspace</li>
<li>Upload & review capability</li>
</ul>

<p><b>Developed by Randy Singh from Kalsnet (KNet) Consulting Group</b></p>

</div>
""", unsafe_allow_html=True)
