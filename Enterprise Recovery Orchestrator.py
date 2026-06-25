



#================================================================================
 # KNet Enterprise Recovery Orchestrator
# --------------------------------------------------------------------------
 #Developed by: Randy Singh
 #Organization: Kalsnet (KNet) Consulting Group

# DESCRIPTION
 --------------------------------------------------------------------------
 #A single-pane-of-glass Streamlit demonstration application that simulates
 #an Enterprise Recovery Orchestrator across five recovery domains:

    # 1. Disk Recovery
    # 2. Database Recovery
    # 3. Cloud Recovery
    # 4. Email Recovery
    # 5. Cyber Artifact Recovery

 # It generates synthetic recovery telemetry, persists run history to a local
 # SQLite database, renders interactive charts, simulates an "Agentic AI"
 # recovery advisor (rule-based decision engine standing in for an LLM), and
 # exports results to CSV, JSON, TXT, PDF, and DOCX.


import io
import json
import random
import sqlite3
from datetime import datetime, timedelta

import numpy as np
import pandas as pd
import streamlit as st

# --------------------------------------------------------------------------
# Optional dependency detection (graceful degradation)
# --------------------------------------------------------------------------
try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (
        SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
    )
    from reportlab.lib.units import inch
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import docx
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import plotly.express as px
    PLOTLY_AVAILABLE = True
except ImportError:
    PLOTLY_AVAILABLE = False


# ==========================================================================
# CONSTANTS / DOMAIN KNOWLEDGE BASE
# ==========================================================================

APP_TITLE = "Enterprise Recovery Orchestrator"
APP_AUTHOR = "Randy Singh"
APP_ORG = "Kalsnet (KNet) Consulting Group"
DB_PATH = "knet_recovery.db"

STATUS_OPTIONS = ["Recovered", "Partial", "Failed"]

# Sub-categories ("techniques") generated per domain
DOMAIN_SUBTYPES = {
    "Disk Recovery": [
        "Deleted File Recovery",
        "Partition Recovery",
        "Signature Carving",
        "Hash Verification",
    ],
    "Database Recovery": [
        "Transaction Log Replay",
        "Record Reconstruction",
        "Backup Comparison",
        "SQL Recovery",
    ],
    "Cloud Recovery": [
        "Snapshot Restoration",
        "Object Versioning",
        "Bucket Recovery",
        "IAM Validation",
    ],
    "Email Recovery": [
        "PST Recovery",
        "OST Recovery",
        "Exchange Export Recovery",
    ],
    "Cyber Artifact Recovery": [
        "Zeek Log Analysis",
        "PCAP Reconstruction",
        "Firewall Log Recovery",
        "API Log Recovery",
        "SIEM Event Recovery",
    ],
}

# Approximate success-rate weighting per domain: [Recovered, Partial, Failed]
DOMAIN_STATUS_WEIGHTS = {
    "Disk Recovery": [0.60, 0.25, 0.15],
    "Database Recovery": [0.55, 0.30, 0.15],
    "Cloud Recovery": [0.65, 0.20, 0.15],
    "Email Recovery": [0.50, 0.30, 0.20],
    "Cyber Artifact Recovery": [0.45, 0.30, 0.25],
}

# Tools referenced per domain (used by export + AI advisor)
DOMAIN_TOOLS = {
    "Disk Recovery": ["FTK Imager", "Autopsy", "X-Ways Forensics", "PhotoRec"],
    "Database Recovery": ["SQL Log Analyzer", "RecoverPoint", "pg_waldump", "ApexSQL Log"],
    "Cloud Recovery": ["AWS Backup", "Azure Site Recovery", "GCP Snapshot Manager", "Velero"],
    "Email Recovery": ["Kernel PST Repair", "Stellar OST to PST", "eDiscovery Export Tool"],
    "Cyber Artifact Recovery": ["Zeek", "Wireshark", "Splunk", "Elastic SIEM", "Suricata"],
}

# Suggested remediation / next-step actions per domain, keyed loosely to status mix
DOMAIN_ACTIONS = {
    "Disk Recovery": [
        "Create a forensic image before further analysis",
        "Verify recovered file hashes against known-good baselines",
        "Run signature carving on unallocated space",
        "Validate partition table integrity",
    ],
    "Database Recovery": [
        "Replay transaction logs from last known checkpoint",
        "Compare reconstructed records against latest backup",
        "Validate referential integrity post-recovery",
        "Run consistency checks (DBCC / pg_amcheck)",
    ],
    "Cloud Recovery": [
        "Restore from most recent validated snapshot",
        "Enable object versioning to prevent future loss",
        "Audit IAM policies tied to affected buckets",
        "Cross-region replicate critical recovered objects",
    ],
    "Email Recovery": [
        "Repair PST/OST containers prior to import",
        "Cross-validate recovered mail against Exchange export",
        "Re-index recovered mailbox for eDiscovery",
        "Notify compliance/legal of recovered custodial data",
    ],
    "Cyber Artifact Recovery": [
        "Correlate Zeek/PCAP artifacts with SIEM timeline",
        "Re-ingest recovered logs into SIEM for retroactive alerting",
        "Validate firewall log continuity (gap analysis)",
        "Preserve chain-of-custody for recovered artifacts",
    ],
}


# ==========================================================================
# DATABASE LAYER (SQLite persistence)
# ==========================================================================

def init_db():
    """Create the SQLite schema if it does not already exist."""
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("""
        CREATE TABLE IF NOT EXISTS runs (
            run_id INTEGER PRIMARY KEY AUTOINCREMENT,
            timestamp TEXT,
            domains TEXT,
            num_records INTEGER,
            recovered INTEGER,
            partial INTEGER,
            failed INTEGER,
            probability REAL,
            confidence TEXT
        )
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS records (
            record_id INTEGER PRIMARY KEY AUTOINCREMENT,
            run_id INTEGER,
            rec_index INTEGER,
            domain TEXT,
            subtype TEXT,
            status TEXT,
            size_mb REAL,
            confidence_score REAL,
            tool_used TEXT,
            timestamp TEXT,
            FOREIGN KEY (run_id) REFERENCES runs (run_id)
        )
    """)
    conn.commit()
    conn.close()


def save_run_to_db(df: pd.DataFrame, domains, ai_summary: dict) -> int:
    """Persist a generated run + its records to SQLite. Returns run_id."""
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    counts = df["Status"].value_counts().to_dict()
    cur.execute(
        """INSERT INTO runs
           (timestamp, domains, num_records, recovered, partial, failed, probability, confidence)
           VALUES (?, ?, ?, ?, ?, ?, ?, ?)""",
        (
            datetime.now().isoformat(timespec="seconds"),
            ", ".join(domains),
            len(df),
            counts.get("Recovered", 0),
            counts.get("Partial", 0),
            counts.get("Failed", 0),
            ai_summary.get("probability", 0.0),
            ai_summary.get("confidence", "Unknown"),
        ),
    )
    run_id = cur.lastrowid

    rows = [
        (
            run_id,
            int(r["ID"]),
            r["Domain"],
            r["SubType"],
            r["Status"],
            float(r["Size_MB"]),
            float(r["ConfidenceScore"]),
            r["ToolUsed"],
            r["Timestamp"],
        )
        for _, r in df.iterrows()
    ]
    cur.executemany(
        """INSERT INTO records
           (run_id, rec_index, domain, subtype, status, size_mb, confidence_score, tool_used, timestamp)
           VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)""",
        rows,
    )
    conn.commit()
    conn.close()
    return run_id


def load_run_history() -> pd.DataFrame:
    conn = sqlite3.connect(DB_PATH)
    df = pd.read_sql_query("SELECT * FROM runs ORDER BY run_id DESC", conn)
    conn.close()
    return df


def clear_history():
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("DELETE FROM records")
    cur.execute("DELETE FROM runs")
    conn.commit()
    conn.close()


# ==========================================================================
# SYNTHETIC DATA GENERATION
# ==========================================================================

def generate_synthetic_data(domains: list, num_records: int) -> pd.DataFrame:
    """Generate synthetic recovery telemetry across the selected domains."""
    if not domains or num_records <= 0:
        return pd.DataFrame(
            columns=["ID", "Domain", "SubType", "Status", "Size_MB",
                     "ConfidenceScore", "ToolUsed", "Timestamp"]
        )

    records = []
    now = datetime.now()
    for i in range(1, num_records + 1):
        domain = random.choice(domains)
        subtype = random.choice(DOMAIN_SUBTYPES[domain])
        weights = DOMAIN_STATUS_WEIGHTS[domain]
        status = random.choices(STATUS_OPTIONS, weights=weights, k=1)[0]

        # Confidence score correlates loosely with status
        if status == "Recovered":
            confidence_score = round(random.uniform(0.80, 0.99), 2)
        elif status == "Partial":
            confidence_score = round(random.uniform(0.45, 0.79), 2)
        else:
            confidence_score = round(random.uniform(0.05, 0.44), 2)

        size_mb = round(random.uniform(0.5, 4096.0), 2)
        tool_used = random.choice(DOMAIN_TOOLS[domain])
        ts = (now - timedelta(minutes=random.randint(0, 10_000))).isoformat(timespec="seconds")

        records.append({
            "ID": i,
            "Domain": domain,
            "SubType": subtype,
            "Status": status,
            "Size_MB": size_mb,
            "ConfidenceScore": confidence_score,
            "ToolUsed": tool_used,
            "Timestamp": ts,
        })

    return pd.DataFrame(records)


# ==========================================================================
# AGENTIC AI RECOVERY ADVISOR (rule-based decision engine)
# ==========================================================================
# This mimics the "Disk Analyzer / DB Analyzer / Cloud Analyzer / Log Analyzer
# -> LLM Recovery Advisor -> Confidence Score / Recovery Plan / Estimated Time
# / Suggested Tools" pipeline described in the architecture, using
# deterministic rules in place of a live LLM call. Swap `compute_ai_recommendation`
# for an actual LLM/agent call (e.g. the Anthropic API) in production.

def compute_ai_recommendation(df: pd.DataFrame, domains: list) -> dict:
    if df.empty:
        return {
            "probability": 0.0,
            "confidence": "N/A",
            "estimated_time": "N/A",
            "actions": [],
            "tools": [],
        }

    counts = df["Status"].value_counts()
    total = len(df)
    recovered = counts.get("Recovered", 0)
    partial = counts.get("Partial", 0)
    failed = counts.get("Failed", 0)

    probability = round(((recovered * 1.0) + (partial * 0.5)) / total * 100, 1)

    if probability >= 85:
        confidence = "High"
    elif probability >= 60:
        confidence = "Medium"
    else:
        confidence = "Low"

    # Estimated time scales with record volume and number of failed/partial cases
    base_minutes = total * 0.8 + failed * 3 + partial * 1.5
    hours = int(base_minutes // 60)
    minutes = int(base_minutes % 60)
    estimated_time = f"{hours}h {minutes}m" if hours else f"{minutes}m"

    # Aggregate suggested actions & tools across the active domains,
    # prioritized by domains with the most non-Recovered records.
    domain_failure_counts = (
        df[df["Status"] != "Recovered"]["Domain"].value_counts().to_dict()
    )
    priority_domains = sorted(
        domains, key=lambda d: domain_failure_counts.get(d, 0), reverse=True
    )

    actions = []
    tools = set()
    for d in priority_domains:
        for a in DOMAIN_ACTIONS.get(d, []):
            if a not in actions:
                actions.append(a)
        tools.update(DOMAIN_TOOLS.get(d, []))
        if len(actions) >= 6:
            break

    return {
        "probability": probability,
        "confidence": confidence,
        "estimated_time": estimated_time,
        "actions": actions[:6],
        "tools": sorted(tools)[:6],
        "recovered": int(recovered),
        "partial": int(partial),
        "failed": int(failed),
        "total": int(total),
    }


# ==========================================================================
# EXPORT HELPERS
# ==========================================================================

def export_csv(df: pd.DataFrame) -> bytes:
    return df.to_csv(index=False).encode("utf-8")


def export_json(df: pd.DataFrame, ai_summary: dict) -> bytes:
    payload = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "developer": APP_AUTHOR,
        "organization": APP_ORG,
        "record_count": len(df),
        "ai_recommendation": ai_summary,
        "records": df.to_dict(orient="records"),
    }
    return json.dumps(payload, indent=2).encode("utf-8")


def export_txt(df: pd.DataFrame, ai_summary: dict, domains: list) -> bytes:
    lines = []
    lines.append("=" * 70)
    lines.append(f"{APP_TITLE}")
    lines.append(f"Developed by {APP_AUTHOR} - {APP_ORG}")
    lines.append("=" * 70)
    lines.append(f"Generated: {datetime.now().isoformat(timespec='seconds')}")
    lines.append(f"Domains: {', '.join(domains)}")
    lines.append(f"Total Records: {len(df)}")
    lines.append("-" * 70)
    lines.append("RESULTS")
    lines.append("-" * 70)
    lines.append(df.to_string(index=False))
    lines.append("-" * 70)
    lines.append("AI RECOVERY ADVISOR")
    lines.append("-" * 70)
    lines.append(f"Recovery Probability : {ai_summary.get('probability', 0)}%")
    lines.append(f"Confidence            : {ai_summary.get('confidence', 'N/A')}")
    lines.append(f"Estimated Time        : {ai_summary.get('estimated_time', 'N/A')}")
    lines.append("Suggested Actions:")
    for a in ai_summary.get("actions", []):
        lines.append(f"  - {a}")
    lines.append("Suggested Tools:")
    for t in ai_summary.get("tools", []):
        lines.append(f"  - {t}")
    lines.append("=" * 70)
    return "\n".join(lines).encode("utf-8")


def export_pdf(df: pd.DataFrame, ai_summary: dict, domains: list) -> bytes:
    if not PDF_AVAILABLE:
        raise RuntimeError("reportlab is not installed. Run: pip install reportlab")

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter, topMargin=0.6 * inch, bottomMargin=0.6 * inch)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("TitleKNet", parent=styles["Title"], textColor=colors.HexColor("#0B3D91"))
    sub_style = ParagraphStyle("SubKNet", parent=styles["Normal"], textColor=colors.HexColor("#444444"))

    elements = []
    elements.append(Paragraph(APP_TITLE, title_style))
    elements.append(Paragraph(f"Developed by {APP_AUTHOR} &mdash; {APP_ORG}", sub_style))
    elements.append(Spacer(1, 12))
    elements.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles["Normal"]))
    elements.append(Paragraph(f"Domains: {', '.join(domains)}", styles["Normal"]))
    elements.append(Paragraph(f"Total Records: {len(df)}", styles["Normal"]))
    elements.append(Spacer(1, 16))

    elements.append(Paragraph("Recovery Results", styles["Heading2"]))
    table_data = [list(df.columns)] + df.astype(str).values.tolist()
    # Cap rows shown in PDF body for readability; full data still in CSV/JSON
    max_rows = 60
    truncated = len(table_data) - 1 > max_rows
    table_data = [table_data[0]] + table_data[1:max_rows + 1]
    t = Table(table_data, repeatRows=1)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0B3D91")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTSIZE", (0, 0), (-1, -1), 7),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F2F4F8")]),
    ]))
    elements.append(t)
    if truncated:
        elements.append(Spacer(1, 6))
        elements.append(Paragraph(f"(Showing first {max_rows} of {len(df)} records — see CSV/JSON for full data)", styles["Italic"]))

    elements.append(PageBreak())
    elements.append(Paragraph("AI Recovery Advisor", styles["Heading2"]))
    elements.append(Paragraph(f"Recovery Probability: {ai_summary.get('probability', 0)}%", styles["Normal"]))
    elements.append(Paragraph(f"Confidence: {ai_summary.get('confidence', 'N/A')}", styles["Normal"]))
    elements.append(Paragraph(f"Estimated Time to Complete: {ai_summary.get('estimated_time', 'N/A')}", styles["Normal"]))
    elements.append(Spacer(1, 10))
    elements.append(Paragraph("Suggested Actions:", styles["Heading3"]))
    for a in ai_summary.get("actions", []):
        elements.append(Paragraph(f"&bull; {a}", styles["Normal"]))
    elements.append(Spacer(1, 10))
    elements.append(Paragraph("Suggested Tools:", styles["Heading3"]))
    for tl in ai_summary.get("tools", []):
        elements.append(Paragraph(f"&bull; {tl}", styles["Normal"]))

    doc.build(elements)
    buf.seek(0)
    return buf.read()


def export_docx(df: pd.DataFrame, ai_summary: dict, domains: list) -> bytes:
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx")

    document = docx.Document()

    title = document.add_heading(APP_TITLE, level=0)
    sub = document.add_paragraph(f"Developed by {APP_AUTHOR} — {APP_ORG}")
    sub.runs[0].font.size = Pt(11)
    sub.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    document.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    document.add_paragraph(f"Domains: {', '.join(domains)}")
    document.add_paragraph(f"Total Records: {len(df)}")

    document.add_heading("Recovery Results", level=1)
    max_rows = 100
    display_df = df.head(max_rows)
    table = document.add_table(rows=1, cols=len(display_df.columns))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(display_df.columns):
        hdr_cells[i].text = str(col)
    for _, row in display_df.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    if len(df) > max_rows:
        document.add_paragraph(f"(Showing first {max_rows} of {len(df)} records — see CSV/JSON for full data)")

    document.add_page_break()
    document.add_heading("AI Recovery Advisor", level=1)
    document.add_paragraph(f"Recovery Probability: {ai_summary.get('probability', 0)}%")
    document.add_paragraph(f"Confidence: {ai_summary.get('confidence', 'N/A')}")
    document.add_paragraph(f"Estimated Time to Complete: {ai_summary.get('estimated_time', 'N/A')}")

    document.add_heading("Suggested Actions", level=2)
    for a in ai_summary.get("actions", []):
        document.add_paragraph(a, style="List Bullet")

    document.add_heading("Suggested Tools", level=2)
    for tl in ai_summary.get("tools", []):
        document.add_paragraph(tl, style="List Bullet")

    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)
    return buf.read()


# ==========================================================================
# UI HELPERS
# ==========================================================================

def render_header():
    st.markdown(
        """
        <div style="background-color:#0B3D91;padding:18px 24px;border-radius:8px;margin-bottom:18px;">
            <h1 style="color:white;margin:0;font-size:30px;">🛡️ Enterprise Recovery Orchestrator</h1>
            <p style="color:#D7E3FF;margin:4px 0 0 0;font-size:15px;">
                Developed by <b>Randy Singh</b> &nbsp;|&nbsp; Kalsnet (KNet) Consulting Group
            </p>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_status_badge(confidence: str):
    color_map = {"High": "#1A7F37", "Medium": "#B58105", "Low": "#B42318", "N/A": "#666666"}
    color = color_map.get(confidence, "#666666")
    st.markdown(
        f"""<span style="background-color:{color};color:white;padding:4px 12px;
        border-radius:14px;font-weight:600;font-size:13px;">{confidence}</span>""",
        unsafe_allow_html=True,
    )


# ==========================================================================
# MAIN APP
# ==========================================================================

def main():
    st.set_page_config(
        page_title="KNet Enterprise Recovery Orchestrator",
        page_icon="🛡️",
        layout="wide",
    )

    init_db()
    render_header()

    # ---------------- Session state defaults ----------------
    if "data" not in st.session_state:
        st.session_state.data = pd.DataFrame()
    if "generated" not in st.session_state:
        st.session_state.generated = False
    if "ai_summary" not in st.session_state:
        st.session_state.ai_summary = {}
    if "active_domains" not in st.session_state:
        st.session_state.active_domains = []

    # ---------------- Sidebar: configuration ----------------
    st.sidebar.header("⚙️ Recovery Configuration")

    domain_list = list(DOMAIN_SUBTYPES.keys())
    domains = st.sidebar.multiselect(
        "Recovery Domain(s)",
        options=domain_list,
        default=["Disk Recovery"],
        help="Select one or more recovery domains to simulate.",
    )

    num_records = st.sidebar.slider("Synthetic Records", min_value=0, max_value=300, value=100, step=1)

    st.sidebar.markdown("---")
    col_gen, col_reset = st.sidebar.columns(2)
    generate_clicked = col_gen.button("🚀 Generate Data", use_container_width=True)
    reset_clicked = col_reset.button("♻️ Reset", use_container_width=True)

    st.sidebar.markdown("---")
    st.sidebar.caption("Optional export engines:")
    st.sidebar.write(f"PDF export: {'✅ available' if PDF_AVAILABLE else '⚠️ install reportlab'}")
    st.sidebar.write(f"DOCX export: {'✅ available' if DOCX_AVAILABLE else '⚠️ install python-docx'}")
    st.sidebar.write(f"Rich charts: {'✅ plotly' if PLOTLY_AVAILABLE else 'ℹ️ using native charts'}")

    # ---------------- Button logic ----------------
    if reset_clicked:
        st.session_state.data = pd.DataFrame()
        st.session_state.generated = False
        st.session_state.ai_summary = {}
        st.session_state.active_domains = []
        st.rerun()

    if generate_clicked:
        if not domains:
            st.sidebar.error("Select at least one recovery domain.")
        elif num_records == 0:
            st.sidebar.warning("Set Synthetic Records above 0 to generate data.")
        else:
            df = generate_synthetic_data(domains, num_records)
            ai_summary = compute_ai_recommendation(df, domains)
            st.session_state.data = df
            st.session_state.generated = True
            st.session_state.ai_summary = ai_summary
            st.session_state.active_domains = domains
            save_run_to_db(df, domains, ai_summary)

    df = st.session_state.data
    ai_summary = st.session_state.ai_summary
    active_domains = st.session_state.active_domains or domains

    # ---------------- Tabs ----------------
    tab_results, tab_charts, tab_ai, tab_export, tab_history = st.tabs(
        ["📊 Results", "📈 Charts", "🤖 AI Recovery Advisor", "💾 Export", "🕘 History"]
    )

    # ===== Results tab =====
    with tab_results:
        if not st.session_state.generated or df.empty:
            st.info("Configure a recovery domain and click **Generate Data** in the sidebar to begin.")
        else:
            counts = df["Status"].value_counts()
            m1, m2, m3, m4 = st.columns(4)
            m1.metric("Total Records", len(df))
            m2.metric("Recovered", int(counts.get("Recovered", 0)))
            m3.metric("Partial", int(counts.get("Partial", 0)))
            m4.metric("Failed", int(counts.get("Failed", 0)))

            st.markdown("#### Recovery Results")
            st.dataframe(
                df[["ID", "Domain", "SubType", "Status", "Size_MB", "ConfidenceScore", "ToolUsed", "Timestamp"]],
                use_container_width=True,
                height=420,
            )

    # ===== Charts tab =====
    with tab_charts:
        if not st.session_state.generated or df.empty:
            st.info("Generate data first to view charts.")
        else:
            c1, c2 = st.columns(2)

            with c1:
                st.markdown("##### Status Distribution")
                status_counts = df["Status"].value_counts().reset_index()
                status_counts.columns = ["Status", "Count"]
                if PLOTLY_AVAILABLE:
                    fig = px.pie(status_counts, names="Status", values="Count", hole=0.4,
                                 color="Status",
                                 color_discrete_map={"Recovered": "#1A7F37", "Partial": "#B58105", "Failed": "#B42318"})
                    st.plotly_chart(fig, use_container_width=True)
                else:
                    st.bar_chart(status_counts.set_index("Status"))

            with c2:
                st.markdown("##### Status by Domain")
                pivot = df.pivot_table(index="Domain", columns="Status", aggfunc="size", fill_value=0)
                if PLOTLY_AVAILABLE:
                    pivot_reset = pivot.reset_index().melt(id_vars="Domain", var_name="Status", value_name="Count")
                    fig2 = px.bar(pivot_reset, x="Domain", y="Count", color="Status", barmode="stack",
                                  color_discrete_map={"Recovered": "#1A7F37", "Partial": "#B58105", "Failed": "#B42318"})
                    st.plotly_chart(fig2, use_container_width=True)
                else:
                    st.bar_chart(pivot)

            st.markdown("##### Confidence Score Distribution")
            if PLOTLY_AVAILABLE:
                fig3 = px.histogram(df, x="ConfidenceScore", color="Status", nbins=20,
                                     color_discrete_map={"Recovered": "#1A7F37", "Partial": "#B58105", "Failed": "#B42318"})
                st.plotly_chart(fig3, use_container_width=True)
            else:
                st.bar_chart(df["ConfidenceScore"])

    # ===== AI Advisor tab =====
    with tab_ai:
        if not st.session_state.generated or df.empty:
            st.info("Generate data first to receive an AI recovery recommendation.")
        else:
            st.markdown("#### 🤖 AI Recovery Advisor")
            st.caption(
                "Agentic decision-engine output (rule-based simulation standing in for an "
                "LLM-driven advisor in production — see `compute_ai_recommendation()`)."
            )

            a1, a2, a3 = st.columns(3)
            a1.metric("Recovery Probability", f"{ai_summary.get('probability', 0)}%")
            with a2:
                st.write("**Confidence**")
                render_status_badge(ai_summary.get("confidence", "N/A"))
            a3.metric("Estimated Time", ai_summary.get("estimated_time", "N/A"))

            st.markdown("##### Suggested Recovery Plan")
            for action in ai_summary.get("actions", []):
                st.markdown(f"- {action}")

            st.markdown("##### Suggested Tools")
            st.write(", ".join(ai_summary.get("tools", [])) or "—")

            st.markdown("##### Analyzer Pipeline (simulated)")
            st.code(
                "Input Sources\n"
                "   -> Disk / DB / Cloud / Email / Log Analyzers\n"
                "   -> LLM Recovery Advisor\n"
                "   -> Confidence Score + Recovery Plan + Estimated Time + Suggested Tools",
                language="text",
            )

    # ===== Export tab =====
    with tab_export:
        if not st.session_state.generated or df.empty:
            st.info("Generate data first to enable exports.")
        else:
            st.markdown("#### 💾 Export Results")
            e1, e2, e3, e4, e5 = st.columns(5)

            with e1:
                st.download_button(
                    "⬇️ CSV", data=export_csv(df), file_name="recovery_results.csv",
                    mime="text/csv", use_container_width=True,
                )
            with e2:
                st.download_button(
                    "⬇️ JSON", data=export_json(df, ai_summary), file_name="recovery_results.json",
                    mime="application/json", use_container_width=True,
                )
            with e3:
                st.download_button(
                    "⬇️ TXT", data=export_txt(df, ai_summary, active_domains), file_name="recovery_results.txt",
                    mime="text/plain", use_container_width=True,
                )
            with e4:
                if PDF_AVAILABLE:
                    try:
                        pdf_bytes = export_pdf(df, ai_summary, active_domains)
                        st.download_button(
                            "⬇️ PDF", data=pdf_bytes, file_name="recovery_report.pdf",
                            mime="application/pdf", use_container_width=True,
                        )
                    except Exception as exc:
                        st.error(f"PDF export failed: {exc}")
                else:
                    st.button("⬇️ PDF", disabled=True, use_container_width=True,
                               help="Install reportlab to enable: pip install reportlab")
            with e5:
                if DOCX_AVAILABLE:
                    try:
                        docx_bytes = export_docx(df, ai_summary, active_domains)
                        st.download_button(
                            "⬇️ DOCX", data=docx_bytes, file_name="recovery_report.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True,
                        )
                    except Exception as exc:
                        st.error(f"DOCX export failed: {exc}")
                else:
                    st.button("⬇️ DOCX", disabled=True, use_container_width=True,
                               help="Install python-docx to enable: pip install python-docx")

    # ===== History tab =====
    with tab_history:
        st.markdown("#### 🕘 Run History (SQLite)")
        history_df = load_run_history()
        if history_df.empty:
            st.info("No runs persisted yet. Generate data to create a history entry.")
        else:
            st.dataframe(history_df, use_container_width=True, height=350)
            if st.button("🗑️ Clear History"):
                clear_history()
                st.rerun()

    st.markdown("---")
    st.caption(
        f"{APP_TITLE} • {APP_ORG} • Synthetic demonstration data — no real recovery operations are performed."
    )


if __name__ == "__main__":
    main()
