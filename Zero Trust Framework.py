
# =============================================================================
 # ZERO-TRUST FRAMEWORK APPLICATION
 # Developed by Randy Singh | KalSnet (KNet) Consulting Group
# Original Java code: April 2019 — Python/Streamlit conversion: 2024
# =============================================================================
# Run with:import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import random
import json
import io
from datetime import datetime, timedelta

# ─── Optional export dependencies ─────────────────────────────────────────────
try:
    from fpdf import FPDF
    FPDF_AVAILABLE = True
except ImportError:
    FPDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    from docx.shared import RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ══════════════════════════════════════════════════════════════════════════════
#  PAGE CONFIG
# ══════════════════════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="Zero-Trust Framework | KNet Consulting",
    page_icon="🔒",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ══════════════════════════════════════════════════════════════════════════════
#  GLOBAL STYLES
# ══════════════════════════════════════════════════════════════════════════════
st.markdown("""
<style>
  .zt-title {
      text-align: center; font-size: 2.4rem; font-weight: 900;
      color: #0047AB; letter-spacing: 1px; margin-bottom: 0;
  }
  .zt-subtitle {
      text-align: center; font-size: 1.05rem; font-weight: 700;
      color: #0047AB; margin-top: 2px; margin-bottom: 4px;
  }
  .zt-tagline {
      text-align: center; color: #555; font-size: 0.88rem; margin-bottom: 18px;
  }
  .zt-hr { border: 2px solid #0047AB; margin: 8px 0 18px 0; }
  .section-header {
      background: linear-gradient(90deg, #0047AB 0%, #1a6fe8 100%);
      color: white; padding: 8px 16px; border-radius: 6px;
      font-weight: 700; font-size: 1.05rem; margin: 10px 0 6px 0;
  }
  .req-card {
      background: #f0f6ff; border-left: 4px solid #0047AB;
      border-radius: 4px; padding: 10px 14px; margin: 6px 0; font-size: 0.92rem;
  }
  .rec-card {
      background: #fff8e1; border-left: 4px solid #f0a500;
      border-radius: 4px; padding: 10px 14px; margin: 6px 0; font-size: 0.92rem;
  }
  .sol-card {
      background: #e8f5e9; border-left: 4px solid #2e7d32;
      border-radius: 4px; padding: 10px 14px; margin: 6px 0; font-size: 0.92rem;
  }
  div[data-testid="metric-container"] {
      background: #f0f6ff; border: 1px solid #c2d9ff; border-radius: 8px; padding: 10px;
  }
  section[data-testid="stSidebar"] { background: #001a6b; }
  section[data-testid="stSidebar"] * { color: #e8f0ff !important; }
</style>
""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
#  HEADER
# ══════════════════════════════════════════════════════════════════════════════
st.markdown('<p class="zt-title">ZERO-TRUST FRAMEWORK APPLICATION</p>', unsafe_allow_html=True)
st.markdown('<p class="zt-subtitle">Developed by Randy Singh | KalSnet (KNet) Consulting Group</p>', unsafe_allow_html=True)
st.markdown('<p class="zt-tagline">Based on DISA/BDE5 ZTA Working Group Requirements &amp; Use Cases &nbsp;|&nbsp; Originally authored April 2019</p>', unsafe_allow_html=True)
st.markdown('<hr class="zt-hr">', unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
#  DATA DEFINITIONS
# ══════════════════════════════════════════════════════════════════════════════
FUNCTIONAL_REQUIREMENTS = {
    "Identity & Authentication": {
        "requirement": "Need to associate identity to user and device.",
        "example": "Active Directory + certificate-based device identity via PKI/CAC cards.",
        "solution": "Deploy a centralised Identity Provider (IdP) such as Microsoft Entra ID or Okta integrated with LDAP/SAML.",
        "recommendation": "Enforce MFA for every user login. Bind device certificates to hardware TPM chips to prevent credential theft.",
    },
    "Health & Compliance": {
        "requirement": "Need to determine patch level and security posture of device prior to network authorization.",
        "example": "NAC solution checks OS STIG compliance, AV status, and missing agents before granting network access.",
        "solution": "Integrate a Network Access Control (NAC) platform such as Cisco ISE or Forescout with your IdP.",
        "recommendation": "Define a minimum compliance baseline (patch age <= 30 days, STIG score >= 80%). Non-compliant devices land in an auto-remediation VLAN.",
    },
    "Authorization": {
        "requirement": "Need to authorize user/device, based on identity and hygiene, at micro-perimeter or Just-In-Time (JIT) for administrators.",
        "example": "PAM tool grants time-limited SSH/RDP session to admin only after ITSM ticket is validated.",
        "solution": "Deploy a Privileged Access Management (PAM) solution (CyberArk, BeyondTrust) with JIT provisioning tied to ServiceNow tickets.",
        "recommendation": "Apply least-privilege RBAC. Rotate JIT credentials every session. Log every privileged action to SIEM.",
    },
    "Accounting": {
        "requirement": "Need to identify and document resources consumed by users and administrators for compliance and incident response.",
        "example": "Splunk SIEM ingests auth logs, network flow data, and privileged session recordings.",
        "solution": "Centralise logging via a SIEM (Splunk, Microsoft Sentinel). Feed NetFlow data and endpoint telemetry.",
        "recommendation": "Retain logs for at least 12 months online, 36 months archived. Automate anomaly alerting with UEBA rules.",
    },
    "Segmentation": {
        "requirement": "Need to limit communication (directional port/protocol) between all network flows including application tiers.",
        "example": "Micro-segmentation policy allows only TCP 443 from Web tier to App tier; all other flows are denied.",
        "solution": "Implement software-defined micro-segmentation (VMware NSX, Illumio, or Guardicore). Define allow-lists per workload.",
        "recommendation": "Start with ring-fence of crown-jewel systems. Expand micro-segmentation iteratively. Validate with automated penetration tests.",
    },
    "Orchestration": {
        "requirement": "Need capability to dynamically provision identity, hygiene, authorization, and segmentation policies.",
        "example": "Ansible playbook triggered by new VM spin-up auto-pushes firewall rules and registers device in IdP.",
        "solution": "Use a Security Orchestration platform (SOAR) or Infrastructure-as-Code (Terraform + Ansible) with REST API integration to all ZTA pillars.",
        "recommendation": "Treat policies as code. Store in Git. Enforce CI/CD pipeline validation before deployment. Run automated compliance checks post-deploy.",
    },
}

USE_CASES = {
    "Use Case 1 - End-User Authentication & Authorization": [
        "User leverages Multi-Factor Authentication (MFA) to establish identity via a centralised Identity Provider.",
        "Device establishes identity via certificate-based authentication (PKI/CAC).",
        "Health & Compliance tool confirms security posture, patch status, and required tool presence.",
        "Real-time contextual access control is determined by validated identity + device health score.",
        "Network authorisation granted based on contextual access control policy.",
        "All authentication & authorisation activities are logged to SIEM (Splunk/Sentinel).",
    ],
    "Use Case 2 - Application & Data Segmentation": [
        "Software-Defined Networking (SDN) enables MicroCore & Perimeter (MCAP) for switching/routing + firewall.",
        "Each application tier (Web, App, DB) enforces granular port/protocol firewall policy.",
        "Application processes are learned and whitelisted for legitimate port/protocol usage.",
        "Data is tagged and specifically associated with owning systems/applications.",
        "All authorisation activities are logged to SIEM.",
        "All authentication & authorisation occur against the centralised Identity Provider.",
    ],
    "Use Case 3 - Admin Authentication & Authorization": [
        "Administrator authenticates via MFA against a centralised Identity Provider.",
        "Just-In-Time (JIT) gateway grants time-limited access to specific systems upon valid ITSM ticket.",
        "Systems managed via JIT Gateway or API-driven orchestration tool over isolated management segment.",
        "All authorisation activities logged to SIEM.",
        "Centralised Identity Provider leveraged for all identities and authorisations.",
    ],
    "Use Case 4 - Security Policy Orchestration": [
        "SDE Global Orchestrator enables service delivery and cyber operations across the agency.",
        "Service-level orchestrators deploy specific systems and associated granular security policies via API.",
        "Policy Enforcement Engines accept policy via REST API and enforce on infrastructure.",
        "All orchestration activities logged to SIEM.",
        "Orchestrated policies include deployment of identity and roles on centralised Identity Provider.",
    ],
}

ARCHITECTURAL_REQUIREMENTS = {
    "Identity Requirements": [
        "1. The system shall NOT assume trust of the device or user.",
        "1.1 The system shall have the capability to identify the user.",
        "1.2 The system shall have the capability to identify the device.",
        "1.2.1 The system shall be able to identify client and server endpoints.",
        "1.2.2 The system shall be able to identify network devices.",
        "2. The system shall integrate with an enterprise identity provider.",
        "2.1 The system shall integrate with industry-standard providers: AD, RADIUS, LDAP, OAuth, OpenID Connect, SAML.",
        "3. The system shall support enterprise authentication methods.",
        "3.1 The system shall support certificate-based authentication.",
        "3.2 The system shall support Public Key Infrastructure (PKI).",
        "3.3 The system shall support or integrate with Multi-Factor Authentication (MFA).",
        "4. The system shall have the ability to continuously challenge identity.",
        "4.1 The system shall support multiple attributes (e.g., geolocation) to challenge identity.",
        "5. The system shall connect to a centralised identity source.",
        "6. The system shall redirect or drop unidentified users or machines.",
    ],
    "Health & Compliance Requirements": [
        "1. The system shall determine endpoint health & compliance.",
        "1.1 The system shall determine OS STIG compliance posture.",
        "1.2 The system shall determine patch level of the device.",
        "1.3 The system shall determine if required tools/agents are missing.",
        "1.4 The system shall detect unauthorised software on the device.",
        "1.5 The system shall determine FQDN, IP, MAC address, and OS of the device.",
        "2. The system shall place non-compliant machines in remediation.",
        "3. The system shall support health & compliance checks across multiple device types.",
    ],
    "Authorization Requirements": [
        "1. The system shall control authorisation of users and machines.",
        "1.1 The system shall assign role-based access controls (RBAC) to groups and users.",
        "2. The system shall support time-limited authorisation for administrators.",
        "2.1 The system shall provide a portal to request privileged access (JIT).",
        "3. The system shall associate identity to authorisation policy via ZTA tool integration.",
        "3.1 The system shall align network security policy to authorisation of users and machines.",
        "4. The system shall support dynamic/real-time access decisions.",
        "5. The system shall support querying to determine authorisation level of user or device.",
    ],
    "Accounting & Auditing Requirements": [
        "1. The system shall provide accounting of all elements within the ZTA.",
        "1.1 The system shall account for user and device identity events.",
        "1.2 The system shall account for each authorisation grant or change.",
        "1.3 The system shall account for each segmentation policy change.",
        "2. The system shall account for all network flows within the ZTA.",
        "2.1 The system shall provide packet capture capability.",
        "2.2 The system shall log and account for all production data flow transactions.",
        "2.3 The system shall support offloading to enterprise SIEM tools.",
        "3. The system shall provide log data via REST API to a log aggregator.",
    ],
    "Segmentation Requirements": [
        "1. The system shall provide network micro-segmentation of application systems.",
        "1.1 The system shall apply role-based rule sets to groups and individual systems.",
        "1.2 The system shall segment north-south network flows.",
        "1.3 The system shall segment east-west network flows.",
        "1.4 The system shall segment end-to-end / transport network flows.",
        "2. The system shall segment both stateless and stateful application flows.",
        "3. The system shall segment host-level processes for security policies.",
    ],
    "Orchestration Requirements": [
        "1. The system shall support automation through REST APIs.",
        "1.1 The system shall accept REST API calls to configure identity, security, and accounting policies.",
        "2. The system shall support dynamic policy enforcement and changes.",
    ],
    "Cloud, Data Tagging & Discovery Requirements": [
        "1. The system shall provide the same protections across cloud-hosted environments.",
        "1.1 The system shall support CSP IaaS offerings.",
        "1.2 The system shall support CSP PaaS offerings.",
        "1.3 The system shall support CSP SaaS offerings.",
        "1.4 The system shall protect/proxy CSP administrative access.",
        "2. The system shall support tagging of data.",
        "3. The system shall support asset discovery.",
        "3.1 The system shall discover assets on the network.",
        "3.2 The system shall discover flows between assets on the network.",
    ],
    "Optional Requirements": [
        "1. The system shall replace existing security architecture constructs.",
        "1.1 The system shall replace existing perimeter security capabilities.",
        "2. The system shall integrate with existing security architecture constructs.",
        "2.1 The system shall integrate with Web Application Firewalls (WAF).",
        "2.2 The system shall integrate with application-aware firewalls (NGFW).",
        "2.3 The system shall integrate with database firewalls.",
        "2.4 The system shall integrate with Intrusion Prevention Systems (IPS).",
    ],
}

EVALUATION_CRITERIA = {
    "Effectiveness": {
        "definition": "Ability to accomplish mission objectives and achievement of desired results -- aligned to ZTA requirements.",
        "metrics": ["% of ZTA pillars fully implemented", "Mean Time to Detect (MTTD)", "Mean Time to Respond (MTTR)", "% reduction in lateral movement incidents"],
        "target": ">= 90% pillar coverage; MTTD < 60 min; MTTR < 4 hrs",
        "recommendation": "Conduct quarterly ZTA maturity assessments against CISA ZT Maturity Model. Map every control to NIST SP 800-207.",
    },
    "Suitability": {
        "definition": "Ability of the solution to be supported in the intended operational environment, aligned to common operational requirements (automation, architecture, security).",
        "metrics": ["Vendor SLA uptime (%)", "Integration API coverage (%)", "STIG compliance score (%)", "Number of supported identity protocols"],
        "target": "99.9% uptime; 100% REST API coverage; STIG score >= 85%",
        "recommendation": "Evaluate solutions against DoD APL. Require vendor FIPS 140-2/3 certification. Include operational sustainment costs in TCO model.",
    },
    "Performance": {
        "definition": "Measure of system performance expressed in quantifiable form -- focused on deployment footprint and scalability.",
        "metrics": ["Authentication latency (ms)", "Policy enforcement throughput (Gbps)", "Simultaneous sessions supported", "Auto-scale response time (sec)"],
        "target": "Auth latency < 200ms; throughput >= 10 Gbps; >= 100K concurrent sessions",
        "recommendation": "Load-test at 150% of projected peak. Define auto-scaling triggers. Monitor with real-time dashboards (Grafana/Prometheus).",
    },
}

# ══════════════════════════════════════════════════════════════════════════════
#  SYNTHETIC DATA
# ══════════════════════════════════════════════════════════════════════════════
@st.cache_data
def generate_synthetic_data(n: int) -> pd.DataFrame:
    random.seed(42)
    pillars = ["Identity & Auth", "Health & Compliance", "Authorization", "Accounting", "Segmentation", "Orchestration"]
    device_types = ["Workstation", "Server", "Mobile", "IoT Sensor", "Network Device", "Cloud VM"]
    departments = ["CYBERCOM", "DISA/BDE5", "TRANSCOM", "DIA", "NSA", "SOCOM", "CENTCOM", "EUCOM"]
    statuses = ["Compliant", "Non-Compliant", "Remediation", "Pending Review"]
    status_weights = [0.60, 0.20, 0.12, 0.08]
    os_list = ["Windows 11 STIG", "RHEL 9 STIG", "macOS Ventura", "Ubuntu 22.04 LTS", "Windows Server 2022"]
    base_date = datetime(2024, 1, 1)
    records = []
    for i in range(1, n + 1):
        status = random.choices(statuses, weights=status_weights)[0]
        score = random.randint(60, 100) if status == "Compliant" else random.randint(20, 59)
        records.append({
            "Record_ID": f"ZT-{i:04d}",
            "Timestamp": base_date + timedelta(days=random.randint(0, 364), hours=random.randint(0, 23)),
            "Pillar": random.choice(pillars),
            "Device_Type": random.choice(device_types),
            "Department": random.choice(departments),
            "OS": random.choice(os_list),
            "Compliance_Status": status,
            "STIG_Score": score,
            "Auth_Latency_ms": random.randint(45, 380),
            "Patch_Age_Days": random.randint(0, 90),
            "MFA_Enabled": random.choice([True, True, True, False]),
            "Segmentation_Policy_Applied": random.choice([True, True, False]),
            "Incident_Flagged": status == "Non-Compliant" and random.random() < 0.4,
        })
    return pd.DataFrame(records)

# ══════════════════════════════════════════════════════════════════════════════
#  EXPORT HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def build_text_report(selected_category: str, content: dict) -> str:
    lines = [
        "=" * 80,
        "  ZERO-TRUST FRAMEWORK APPLICATION",
        "  Developed by Randy Singh | KalSnet (KNet) Consulting Group",
        f"  Category: {selected_category}",
        f"  Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "=" * 80, "",
    ]
    for key, val in content.items():
        lines.append(f"[{key}]")
        if isinstance(val, list):
            for item in val:
                lines.append(f"  * {item}")
        elif isinstance(val, dict):
            for k2, v2 in val.items():
                if isinstance(v2, list):
                    lines.append(f"  {k2}:")
                    for item in v2:
                        lines.append(f"    - {item}")
                else:
                    lines.append(f"  {k2}: {v2}")
        else:
            lines.append(f"  {val}")
        lines.append("")
    return "\n".join(lines)


def build_json_report(selected_category: str, content: dict) -> str:
    payload = {
        "framework": "Zero-Trust Framework Application",
        "author": "Randy Singh | KalSnet (KNet) Consulting Group",
        "category": selected_category,
        "generated": datetime.now().isoformat(),
        "data": content,
    }
    return json.dumps(payload, indent=2, default=str)


def _pdf_safe(text: str) -> str:
    """
    Sanitise a string so it contains only Latin-1-encodable characters.
    Replaces common Unicode symbols with ASCII equivalents, then drops
    anything still outside Latin-1 (including emoji, box-drawing chars, etc.)
    """
    replacements = [
        ("\u2013", "-"), ("\u2014", "--"),
        ("\u2018", "'"), ("\u2019", "'"),
        ("\u201c", '"'), ("\u201d", '"'),
        ("\u2026", "..."), ("\u2022", "*"),
        ("\u00a0", " "), ("\u2265", ">="),
        ("\u2264", "<="), ("\u00ae", "(R)"),
        ("\u2122", "(TM)"), ("\u00e9", "e"),
        ("\u00e8", "e"), ("\u00e0", "a"),
        ("\u00fc", "u"), ("\u00f6", "o"),
        ("\u00e4", "a"), ("\u00b0", " deg"),
        ("\u00b7", "-"), ("\u2212", "-"),
    ]
    for char, replacement in replacements:
        text = text.replace(char, replacement)
    # Strip anything still outside Latin-1
    return text.encode("latin-1", errors="ignore").decode("latin-1")


def build_pdf_bytes(selected_category: str, content: dict) -> bytes:
    """
    Build PDF and return bytes.
    Handles both old fpdf2 (output()->str) and new fpdf2 (output()->bytearray).
    Uses new_x/new_y API instead of deprecated ln= parameter.
    """
    if not FPDF_AVAILABLE:
        return b""
    try:
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        # Header
        pdf.set_font("Helvetica", "B", 16)
        pdf.set_text_color(0, 71, 171)
        pdf.cell(0, 10, _pdf_safe("ZERO-TRUST FRAMEWORK APPLICATION"),
                 new_x="LMARGIN", new_y="NEXT", align="C")

        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 7, _pdf_safe("Developed by Randy Singh | KalSnet (KNet) Consulting Group"),
                 new_x="LMARGIN", new_y="NEXT", align="C")

        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(80, 80, 80)
        ts = _pdf_safe(f"Category: {selected_category}   |   Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        pdf.cell(0, 6, ts, new_x="LMARGIN", new_y="NEXT", align="C")

        pdf.set_draw_color(0, 71, 171)
        pdf.set_line_width(0.8)
        y = pdf.get_y() + 2
        pdf.line(10, y, 200, y)
        pdf.ln(6)

        # Content
        for key, val in content.items():
            pdf.set_font("Helvetica", "B", 11)
            pdf.set_text_color(0, 71, 171)
            pdf.cell(0, 8, _pdf_safe(str(key)), new_x="LMARGIN", new_y="NEXT")

            pdf.set_font("Helvetica", "", 9)
            pdf.set_text_color(30, 30, 30)

            if isinstance(val, list):
                for item in val:
                    pdf.multi_cell(0, 6, _pdf_safe(f"  * {item}"))
            elif isinstance(val, dict):
                for k2, v2 in val.items():
                    pdf.set_font("Helvetica", "B", 9)
                    pdf.cell(0, 6, _pdf_safe(f"  {k2}:"), new_x="LMARGIN", new_y="NEXT")
                    pdf.set_font("Helvetica", "", 9)
                    if isinstance(v2, list):
                        for item in v2:
                            pdf.multi_cell(0, 5, _pdf_safe(f"    - {item}"))
                    else:
                        pdf.multi_cell(0, 6, _pdf_safe(f"    {v2}"))
            else:
                pdf.multi_cell(0, 6, _pdf_safe(f"  {val}"))

            pdf.ln(2)

        raw = pdf.output()
        # fpdf2 >= 2.x returns bytearray; older versions returned a latin-1 str
        if isinstance(raw, (bytes, bytearray)):
            return bytes(raw)
        return raw.encode("latin-1")

    except Exception as exc:
        # Return a minimal valid PDF with the error message instead of crashing
        try:
            fallback = FPDF()
            fallback.add_page()
            fallback.set_font("Helvetica", "", 11)
            fallback.multi_cell(0, 8, _pdf_safe(f"PDF error: {exc}"))
            raw = fallback.output()
            if isinstance(raw, (bytes, bytearray)):
                return bytes(raw)
            return raw.encode("latin-1")
        except Exception:
            return b""


def build_docx_bytes(selected_category: str, content: dict) -> bytes:
    if not DOCX_AVAILABLE:
        return b""
    try:
        doc = DocxDocument()

        # Title
        t = doc.add_heading("ZERO-TRUST FRAMEWORK APPLICATION", 0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in t.runs:
            run.font.color.rgb = RGBColor(0, 71, 171)

        sub = doc.add_paragraph("Developed by Randy Singh | KalSnet (KNet) Consulting Group")
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if sub.runs:
            sub.runs[0].bold = True
            sub.runs[0].font.color.rgb = RGBColor(0, 71, 171)

        ts_para = doc.add_paragraph(
            f"Category: {selected_category}   |   {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        )
        ts_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        for key, val in content.items():
            h = doc.add_heading(str(key), level=2)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0, 71, 171)

            if isinstance(val, list):
                for item in val:
                    doc.add_paragraph(str(item), style="List Bullet")
            elif isinstance(val, dict):
                for k2, v2 in val.items():
                    p = doc.add_paragraph()
                    run_label = p.add_run(f"{k2}: ")
                    run_label.bold = True
                    if isinstance(v2, list):
                        p.add_run(", ".join(str(x) for x in v2))
                    else:
                        p.add_run(str(v2))
            else:
                doc.add_paragraph(str(val))
            doc.add_paragraph()

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    except Exception as exc:
        # Return an error document instead of crashing
        try:
            doc = DocxDocument()
            doc.add_paragraph(f"Word export error: {exc}")
            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue()
        except Exception:
            return b""

# ══════════════════════════════════════════════════════════════════════════════
#  SIDEBAR — NAVIGATION
# ══════════════════════════════════════════════════════════════════════════════
with st.sidebar:
    st.markdown("## ZTA Navigator")
    st.markdown("---")
    category = st.selectbox(
        "Select Framework Category",
        [
            "Overview & Architecture Diagram",
            "1. Functional Requirements",
            "2. Proposed Use Cases",
            "3. Architectural Requirements",
            "4. Evaluation Criteria",
            "Synthetic Data & Analytics",
        ],
    )
    st.markdown("---")
    st.markdown("**About**")
    st.markdown("Randy Singh  \nComputer Scientist  \nDISA / BDE5  \nKalSnet (KNet) Consulting")
    st.markdown("`(301) 225-9535`")

# ══════════════════════════════════════════════════════════════════════════════
#  SIDEBAR — EXPORT
# ══════════════════════════════════════════════════════════════════════════════
def get_export_content(cat: str) -> dict:
    if "Functional" in cat:
        return {
            k: {
                "Requirement": v["requirement"],
                "Example": v["example"],
                "Solution": v["solution"],
                "Recommendation": v["recommendation"],
            }
            for k, v in FUNCTIONAL_REQUIREMENTS.items()
        }
    elif "Use Case" in cat:
        return USE_CASES
    elif "Architectural" in cat:
        return ARCHITECTURAL_REQUIREMENTS
    elif "Evaluation" in cat:
        return {
            k: {
                "Definition": v["definition"],
                "Metrics": v["metrics"],
                "Target": v["target"],
                "Recommendation": v["recommendation"],
            }
            for k, v in EVALUATION_CRITERIA.items()
        }
    else:
        return {"Framework": "Zero-Trust Architecture", "Version": "2.0", "Author": "Randy Singh / KNet"}


with st.sidebar:
    st.markdown("---")
    st.markdown("### Export Results")

    exp_content = get_export_content(category)
    clean_cat = (category
                 .replace("1. ", "").replace("2. ", "").replace("3. ", "").replace("4. ", ""))

    # Text — always available
    txt_data = build_text_report(clean_cat, exp_content)
    st.download_button(
        "Export as Text (.txt)",
        data=txt_data,
        file_name="zt_framework.txt",
        mime="text/plain",
        use_container_width=True,
    )

    # JSON — always available
    json_data = build_json_report(clean_cat, exp_content)
    st.download_button(
        "Export as JSON (.json)",
        data=json_data,
        file_name="zt_framework.json",
        mime="application/json",
        use_container_width=True,
    )

    # PDF
    if FPDF_AVAILABLE:
        pdf_bytes = build_pdf_bytes(clean_cat, exp_content)
        st.download_button(
            "Export as PDF (.pdf)",
            data=pdf_bytes,
            file_name="zt_framework.pdf",
            mime="application/pdf",
            use_container_width=True,
        )
    else:
        st.warning("PDF unavailable. Run: pip install fpdf2")

    # Word
    if DOCX_AVAILABLE:
        docx_bytes = build_docx_bytes(clean_cat, exp_content)
        st.download_button(
            "Export as Word (.docx)",
            data=docx_bytes,
            file_name="zt_framework.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
    else:
        st.warning("Word unavailable. Run: pip install python-docx")

# ══════════════════════════════════════════════════════════════════════════════
#  MAIN CONTENT
# ══════════════════════════════════════════════════════════════════════════════

# ── OVERVIEW ──────────────────────────────────────────────────────────────────
if "Overview" in category:
    st.markdown('<div class="section-header">Zero-Trust Architecture - Conceptual Overview</div>', unsafe_allow_html=True)

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("ZTA Pillars", "6", "Core Controls")
    c2.metric("Use Cases", "4", "Defined")
    c3.metric("Arch. Req. Areas", "7", "Categories")
    c4.metric("Eval. Criteria", "3", "Dimensions")

    st.markdown("---")
    st.markdown("#### Zero-Trust Architecture Flow Diagram")

    try:
        import graphviz
        dot = graphviz.Digraph(graph_attr={"rankdir": "LR", "bgcolor": "#f8f9ff", "fontname": "Helvetica"})
        dot.attr("node", shape="box", style="filled", fontname="Helvetica", fontsize="11")
        dot.node("User",  "User / Device",           fillcolor="#dce8ff", color="#0047AB")
        dot.node("MFA",   "MFA",                     fillcolor="#cce5ff", color="#0047AB")
        dot.node("IdP",   "Identity Provider (IdP)", fillcolor="#0047AB", fontcolor="white", color="#0047AB")
        dot.node("NAC",   "Health & Compliance",     fillcolor="#fff0cc", color="#f0a500")
        dot.node("PAM",   "JIT / PAM Authorization", fillcolor="#e8f5e9", color="#2e7d32")
        dot.node("PEP",   "Policy Enforcement (PEP)",fillcolor="#0047AB", fontcolor="white", color="#0047AB")
        dot.node("App",   "Application Tier",        fillcolor="#dce8ff", color="#0047AB")
        dot.node("SIEM",  "SIEM / Accounting",       fillcolor="#fce4ec", color="#c62828")
        dot.node("Orch",  "Orchestration (REST API)",fillcolor="#f3e5f5", color="#6a1b9a")
        dot.edge("User", "MFA",  label="Auth Request")
        dot.edge("MFA",  "IdP",  label="Credentials")
        dot.edge("User", "NAC",  label="Device Posture")
        dot.edge("NAC",  "IdP",  label="Health Signal")
        dot.edge("IdP",  "PAM",  label="Identity Assert")
        dot.edge("PAM",  "PEP",  label="Auth Decision")
        dot.edge("PEP",  "App",  label="Access Granted")
        dot.edge("PEP",  "SIEM", label="Log Events")
        dot.edge("Orch", "PEP",  label="Policy Push")
        dot.edge("Orch", "IdP",  label="Role Provision")
        dot.edge("App",  "SIEM", label="App Logs")
        st.graphviz_chart(dot, use_container_width=True)
    except Exception:
        st.code("""
USER/DEVICE ---> MFA ----------------> IdP ---------> PAM/JIT
     |                                  |                  |
     +---> NAC (Health Check) ----------+                  |
                                                           v
ORCHESTRATION (REST API) ----------------------> POLICY ENFORCEMENT POINT
                                                           |
                                          +----------------+----------------+
                                          v                                 v
                                    APPLICATION TIERS                 SIEM (ACCOUNTING)
        """)

    st.markdown("---")
    st.markdown("#### ZTA Pillar Maturity Radar")
    pillars_r = ["Identity & Auth", "Health & Compliance", "Authorization", "Accounting", "Segmentation", "Orchestration"]
    current   = [3.5, 2.8, 3.0, 3.8, 2.5, 2.2]
    target    = [5, 5, 5, 5, 5, 5]
    fig_radar = go.Figure()
    fig_radar.add_trace(go.Scatterpolar(
        r=current + [current[0]], theta=pillars_r + [pillars_r[0]],
        fill="toself", name="Current Maturity",
        line_color="#0047AB", fillcolor="rgba(0,71,171,0.2)"))
    fig_radar.add_trace(go.Scatterpolar(
        r=target + [target[0]], theta=pillars_r + [pillars_r[0]],
        fill="toself", name="Target Maturity",
        line_color="#f0a500", fillcolor="rgba(240,165,0,0.1)"))
    fig_radar.update_layout(
        polar=dict(radialaxis=dict(visible=True, range=[0, 5])),
        title="ZTA Pillar Maturity (1=Initial  5=Optimised)",
        height=420)
    st.plotly_chart(fig_radar, use_container_width=True)

    st.markdown("#### ZTA Core Principles")
    principles = [
        ("Never Trust, Always Verify", "Treat every request as untrusted regardless of network location."),
        ("Least-Privilege Access", "Grant minimum permissions necessary and expire them promptly."),
        ("Assume Breach", "Design controls assuming attackers may already be inside the perimeter."),
        ("Continuous Verification", "Re-evaluate trust on every transaction, not just at login."),
        ("Micro-Segmentation", "Divide the network into small zones to contain lateral movement."),
        ("Policy-as-Code", "Automate and version-control all security policies via REST APIs."),
    ]
    for title_p, desc in principles:
        st.markdown(f'<div class="req-card"><b>{title_p}</b><br>{desc}</div>', unsafe_allow_html=True)


# ── FUNCTIONAL REQUIREMENTS ───────────────────────────────────────────────────
elif "Functional" in category:
    st.markdown('<div class="section-header">Category 1 - Functional Requirements</div>', unsafe_allow_html=True)
    show_all = st.checkbox("Expand all pillars", value=True)
    for pillar, data in FUNCTIONAL_REQUIREMENTS.items():
        with st.expander(f"{pillar}", expanded=show_all):
            st.markdown(f'<div class="req-card"><b>Requirement:</b><br>{data["requirement"]}</div>', unsafe_allow_html=True)
            st.markdown(f'<div class="sol-card"><b>Example:</b><br>{data["example"]}</div>', unsafe_allow_html=True)
            st.markdown(f'<div class="sol-card"><b>Recommended Solution:</b><br>{data["solution"]}</div>', unsafe_allow_html=True)
            st.markdown(f'<div class="rec-card"><b>Recommendation:</b><br>{data["recommendation"]}</div>', unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("#### Pillar Coverage (Illustrative)")
    coverage = {"Identity & Authentication": 85, "Health & Compliance": 70, "Authorization": 78,
                "Accounting": 90, "Segmentation": 62, "Orchestration": 55}
    fig_bar = px.bar(x=list(coverage.keys()), y=list(coverage.values()),
                     labels={"x": "Pillar", "y": "Coverage (%)"},
                     color=list(coverage.values()),
                     color_continuous_scale=["#ff4444", "#f0a500", "#0047AB"],
                     title="Estimated Pillar Implementation Coverage (%)")
    fig_bar.update_layout(coloraxis_showscale=False, height=360)
    st.plotly_chart(fig_bar, use_container_width=True)


# ── USE CASES ─────────────────────────────────────────────────────────────────
elif "Use Case" in category:
    st.markdown('<div class="section-header">Category 2 - Proposed Use Cases</div>', unsafe_allow_html=True)
    uc_choice = st.radio("Select Use Case", list(USE_CASES.keys()))
    st.markdown(f"### {uc_choice}")
    for i, step in enumerate(USE_CASES[uc_choice], 1):
        st.markdown(f'<div class="req-card"><b>Step {i}:</b> {step}</div>', unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("#### Use Case Flow Diagram")
    try:
        import graphviz
        g = graphviz.Digraph(graph_attr={"rankdir": "TD"})
        g.attr("node", shape="box", style="filled", fontname="Helvetica", fontsize="10")
        steps_uc = USE_CASES[uc_choice]
        for idx, step_text in enumerate(steps_uc):
            label = step_text[:55] + "..." if len(step_text) > 55 else step_text
            g.node(str(idx), label, fillcolor="#dce8ff", color="#0047AB")
        for idx in range(len(steps_uc) - 1):
            g.edge(str(idx), str(idx + 1))
        st.graphviz_chart(g, use_container_width=True)
    except Exception:
        st.info("Install the graphviz system package for flow diagrams.")


# ── ARCHITECTURAL REQUIREMENTS ────────────────────────────────────────────────
elif "Architectural" in category:
    st.markdown('<div class="section-header">Category 3 - Architectural Requirements</div>', unsafe_allow_html=True)
    area_choice = st.selectbox("Select Requirement Area", list(ARCHITECTURAL_REQUIREMENTS.keys()))
    st.markdown(f"### {area_choice}")
    for req in ARCHITECTURAL_REQUIREMENTS[area_choice]:
        st.markdown(f'<div class="req-card">{req}</div>', unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("#### Requirement Count per Area")
    counts = {k: len(v) for k, v in ARCHITECTURAL_REQUIREMENTS.items()}
    fig_h = px.bar(x=list(counts.values()), y=list(counts.keys()), orientation="h",
                   labels={"x": "# Requirements", "y": "Area"},
                   color=list(counts.values()), color_continuous_scale=["#90caf9", "#0047AB"],
                   title="Architectural Requirements Count by Area")
    fig_h.update_layout(coloraxis_showscale=False, height=380, yaxis={"categoryorder": "total ascending"})
    st.plotly_chart(fig_h, use_container_width=True)

    st.markdown("#### Reference Standards")
    refs = [
        ("NIST SP 800-207", "Zero Trust Architecture - foundational ZTA standard."),
        ("CISA ZT Maturity Model", "5-pillar model for assessing ZTA implementation progress."),
        ("DoD ZT Strategy", "DoD mandate for ZT implementation across all components by FY27."),
        ("DISA STIG Library", "Security Technical Implementation Guides for OS/app hardening."),
        ("NIST SP 800-53 Rev 5", "Security and Privacy Controls - maps to ZTA control requirements."),
    ]
    for ref, desc in refs:
        st.markdown(f'<div class="rec-card"><b>{ref}</b> - {desc}</div>', unsafe_allow_html=True)


# ── EVALUATION CRITERIA ───────────────────────────────────────────────────────
elif "Evaluation" in category:
    st.markdown('<div class="section-header">Category 4 - Evaluation Criteria</div>', unsafe_allow_html=True)
    for criterion, data in EVALUATION_CRITERIA.items():
        with st.expander(f"{criterion}", expanded=True):
            st.markdown(f'<div class="req-card"><b>Definition:</b> {data["definition"]}</div>', unsafe_allow_html=True)
            st.markdown(f'<div class="sol-card"><b>Target:</b> {data["target"]}</div>', unsafe_allow_html=True)
            st.markdown("<b>Key Metrics:</b>", unsafe_allow_html=True)
            for m in data["metrics"]:
                st.markdown(f'<div class="req-card">{m}</div>', unsafe_allow_html=True)
            st.markdown(f'<div class="rec-card"><b>Recommendation:</b> {data["recommendation"]}</div>', unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("#### Sample Evaluation Scores (Illustrative)")
    gc1, gc2, gc3 = st.columns(3)
    for col, (crit, score, color) in zip(
        [gc1, gc2, gc3],
        [("Effectiveness", 78, "#0047AB"), ("Suitability", 85, "#2e7d32"), ("Performance", 71, "#f0a500")]
    ):
        fig_g = go.Figure(go.Indicator(
            mode="gauge+number", value=score, title={"text": crit},
            gauge={"axis": {"range": [0, 100]}, "bar": {"color": color},
                   "steps": [{"range": [0, 50], "color": "#ffcdd2"},
                              {"range": [50, 75], "color": "#fff9c4"},
                              {"range": [75, 100], "color": "#c8e6c9"}],
                   "threshold": {"line": {"color": "black", "width": 3}, "thickness": 0.8, "value": 80}}))
        fig_g.update_layout(height=260, margin=dict(t=40, b=10))
        col.plotly_chart(fig_g, use_container_width=True)


# ── SYNTHETIC DATA ────────────────────────────────────────────────────────────
elif "Synthetic" in category or "Analytics" in category or "Data" in category:
    st.markdown('<div class="section-header">Synthetic ZTA Compliance Data & Analytics</div>', unsafe_allow_html=True)
    num_records = st.slider("Number of Synthetic Records", min_value=0, max_value=300, value=150, step=10)

    if num_records == 0:
        st.info("Move the slider above 0 to generate synthetic data.")
    else:
        df = generate_synthetic_data(num_records)
        k1, k2, k3, k4, k5 = st.columns(5)
        k1.metric("Total Records", num_records)
        k2.metric("Compliant", int((df["Compliance_Status"] == "Compliant").sum()),
                  f'{(df["Compliance_Status"]=="Compliant").mean()*100:.0f}%')
        k3.metric("Non-Compliant", int((df["Compliance_Status"] == "Non-Compliant").sum()))
        k4.metric("Avg STIG Score", f"{df['STIG_Score'].mean():.1f}")
        k5.metric("Incidents Flagged", int(df["Incident_Flagged"].sum()))

        st.markdown("---")
        tab_a, tab_b, tab_c, tab_d = st.tabs(["Raw Data", "By Pillar", "By Department", "Trends"])

        STATUS_COLORS = {"Compliant": "#2e7d32", "Non-Compliant": "#c62828",
                         "Remediation": "#f0a500", "Pending Review": "#757575"}

        with tab_a:
            st.dataframe(df, use_container_width=True, height=380)
            st.download_button("Download CSV", df.to_csv(index=False).encode(),
                               "zt_synthetic_data.csv", "text/csv")

        with tab_b:
            c1, c2 = st.columns(2)
            pillar_status = df.groupby(["Pillar", "Compliance_Status"]).size().reset_index(name="Count")
            fig_ps = px.bar(pillar_status, x="Pillar", y="Count", color="Compliance_Status",
                            barmode="group", color_discrete_map=STATUS_COLORS,
                            title="Compliance Status by ZTA Pillar")
            c1.plotly_chart(fig_ps, use_container_width=True)
            stig_pillar = df.groupby("Pillar")["STIG_Score"].mean().reset_index()
            fig_sp = px.bar(stig_pillar, x="Pillar", y="STIG_Score",
                            color="STIG_Score", color_continuous_scale=["#ff4444", "#f0a500", "#2e7d32"],
                            title="Average STIG Score by Pillar", range_y=[0, 100])
            fig_sp.update_layout(coloraxis_showscale=False)
            c2.plotly_chart(fig_sp, use_container_width=True)

        with tab_c:
            dept_status = df.groupby(["Department", "Compliance_Status"]).size().reset_index(name="Count")
            fig_ds = px.bar(dept_status, x="Department", y="Count", color="Compliance_Status",
                            barmode="stack", color_discrete_map=STATUS_COLORS,
                            title="Compliance Status by Department")
            st.plotly_chart(fig_ds, use_container_width=True)
            dev_counts = df["Device_Type"].value_counts().reset_index()
            dev_counts.columns = ["Device_Type", "Count"]
            fig_pie = px.pie(dev_counts, values="Count", names="Device_Type",
                             title="Device Type Distribution",
                             color_discrete_sequence=px.colors.qualitative.Bold)
            st.plotly_chart(fig_pie, use_container_width=True)

        with tab_d:
            df["Month"] = pd.to_datetime(df["Timestamp"]).dt.to_period("M").astype(str)
            trend = df.groupby(["Month", "Compliance_Status"]).size().reset_index(name="Count")
            fig_trend = px.line(trend, x="Month", y="Count", color="Compliance_Status",
                                markers=True, color_discrete_map=STATUS_COLORS,
                                title="Compliance Trend Over Time")
            st.plotly_chart(fig_trend, use_container_width=True)
            fig_sc = px.scatter(df, x="STIG_Score", y="Auth_Latency_ms",
                                color="Compliance_Status", opacity=0.7,
                                color_discrete_map=STATUS_COLORS,
                                title="STIG Score vs Authentication Latency",
                                labels={"STIG_Score": "STIG Compliance Score",
                                        "Auth_Latency_ms": "Auth Latency (ms)"})
            st.plotly_chart(fig_sc, use_container_width=True)

# ══════════════════════════════════════════════════════════════════════════════
#  FOOTER
# ══════════════════════════════════════════════════════════════════════════════
st.markdown("---")
st.markdown(
    '<div style="text-align:center;color:#888;font-size:0.8rem;">'
    'Zero-Trust Framework Application &nbsp;|&nbsp; '
    'Developed by Randy Singh &nbsp;|&nbsp; '
    'KalSnet (KNet) Consulting Group &nbsp;|&nbsp; '
    f'DISA / BDE5 Technology Innovation Team &nbsp;|&nbsp; {datetime.now().year}'
    '</div>',
    unsafe_allow_html=True,
)
PYEOF
echo "Done — $(wc -l < /mnt/user-data/outputs/zero_trust_framework.py) lines written"
 #   pip install streamlit plotly pandas fpdf2 python-docx
 #   streamlit run zero_trust_framework.py
# =============================================================================

import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import random
import json
import io
from datetime import datetime, timedelta

# ─── Optional export dependencies ─────────────────────────────────────────────
try:
    from fpdf import FPDF
    FPDF_AVAILABLE = True
except ImportError:
    FPDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    from docx.shared import RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ══════════════════════════════════════════════════════════════════════════════
#  PAGE CONFIG
# ══════════════════════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="Zero-Trust Framework | KNet Consulting",
    page_icon="🔒",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ══════════════════════════════════════════════════════════════════════════════
#  GLOBAL STYLES
# ══════════════════════════════════════════════════════════════════════════════
st.markdown("""
<style>
  .zt-title {
      text-align: center; font-size: 2.4rem; font-weight: 900;
      color: #0047AB; letter-spacing: 1px; margin-bottom: 0;
  }
  .zt-subtitle {
      text-align: center; font-size: 1.05rem; font-weight: 700;
      color: #0047AB; margin-top: 2px; margin-bottom: 4px;
  }
  .zt-tagline {
      text-align: center; color: #555; font-size: 0.88rem; margin-bottom: 18px;
  }
  .zt-hr { border: 2px solid #0047AB; margin: 8px 0 18px 0; }
  .section-header {
      background: linear-gradient(90deg, #0047AB 0%, #1a6fe8 100%);
      color: white; padding: 8px 16px; border-radius: 6px;
      font-weight: 700; font-size: 1.05rem; margin: 10px 0 6px 0;
  }
  .req-card {
      background: #f0f6ff; border-left: 4px solid #0047AB;
      border-radius: 4px; padding: 10px 14px; margin: 6px 0; font-size: 0.92rem;
  }
  .rec-card {
      background: #fff8e1; border-left: 4px solid #f0a500;
      border-radius: 4px; padding: 10px 14px; margin: 6px 0; font-size: 0.92rem;
  }
  .sol-card {
      background: #e8f5e9; border-left: 4px solid #2e7d32;
      border-radius: 4px; padding: 10px 14px; margin: 6px 0; font-size: 0.92rem;
  }
  div[data-testid="metric-container"] {
      background: #f0f6ff; border: 1px solid #c2d9ff; border-radius: 8px; padding: 10px;
  }
  section[data-testid="stSidebar"] { background: #001a6b; }
  section[data-testid="stSidebar"] * { color: #e8f0ff !important; }
</style>
""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
#  HEADER
# ══════════════════════════════════════════════════════════════════════════════
st.markdown('<p class="zt-title">ZERO-TRUST FRAMEWORK APPLICATION</p>', unsafe_allow_html=True)
st.markdown('<p class="zt-subtitle">Developed by Randy Singh | KalSnet (KNet) Consulting Group</p>', unsafe_allow_html=True)
st.markdown('<p class="zt-tagline">Based on DISA/BDE5 ZTA Working Group Requirements &amp; Use Cases &nbsp;|&nbsp; Originally authored April 2019</p>', unsafe_allow_html=True)
st.markdown('<hr class="zt-hr">', unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
#  DATA DEFINITIONS
# ══════════════════════════════════════════════════════════════════════════════
FUNCTIONAL_REQUIREMENTS = {
    "Identity & Authentication": {
        "requirement": "Need to associate identity to user and device.",
        "example": "Active Directory + certificate-based device identity via PKI/CAC cards.",
        "solution": "Deploy a centralised Identity Provider (IdP) such as Microsoft Entra ID or Okta integrated with LDAP/SAML.",
        "recommendation": "Enforce MFA for every user login. Bind device certificates to hardware TPM chips to prevent credential theft.",
    },
    "Health & Compliance": {
        "requirement": "Need to determine patch level and security posture of device prior to network authorization.",
        "example": "NAC solution checks OS STIG compliance, AV status, and missing agents before granting network access.",
        "solution": "Integrate a Network Access Control (NAC) platform such as Cisco ISE or Forescout with your IdP.",
        "recommendation": "Define a minimum compliance baseline (patch age <= 30 days, STIG score >= 80%). Non-compliant devices land in an auto-remediation VLAN.",
    },
    "Authorization": {
        "requirement": "Need to authorize user/device, based on identity and hygiene, at micro-perimeter or Just-In-Time (JIT) for administrators.",
        "example": "PAM tool grants time-limited SSH/RDP session to admin only after ITSM ticket is validated.",
        "solution": "Deploy a Privileged Access Management (PAM) solution (CyberArk, BeyondTrust) with JIT provisioning tied to ServiceNow tickets.",
        "recommendation": "Apply least-privilege RBAC. Rotate JIT credentials every session. Log every privileged action to SIEM.",
    },
    "Accounting": {
        "requirement": "Need to identify and document resources consumed by users and administrators for compliance and incident response.",
        "example": "Splunk SIEM ingests auth logs, network flow data, and privileged session recordings.",
        "solution": "Centralise logging via a SIEM (Splunk, Microsoft Sentinel). Feed NetFlow data and endpoint telemetry.",
        "recommendation": "Retain logs for at least 12 months online, 36 months archived. Automate anomaly alerting with UEBA rules.",
    },
    "Segmentation": {
        "requirement": "Need to limit communication (directional port/protocol) between all network flows including application tiers.",
        "example": "Micro-segmentation policy allows only TCP 443 from Web tier to App tier; all other flows are denied.",
        "solution": "Implement software-defined micro-segmentation (VMware NSX, Illumio, or Guardicore). Define allow-lists per workload.",
        "recommendation": "Start with ring-fence of crown-jewel systems. Expand micro-segmentation iteratively. Validate with automated penetration tests.",
    },
    "Orchestration": {
        "requirement": "Need capability to dynamically provision identity, hygiene, authorization, and segmentation policies.",
        "example": "Ansible playbook triggered by new VM spin-up auto-pushes firewall rules and registers device in IdP.",
        "solution": "Use a Security Orchestration platform (SOAR) or Infrastructure-as-Code (Terraform + Ansible) with REST API integration to all ZTA pillars.",
        "recommendation": "Treat policies as code. Store in Git. Enforce CI/CD pipeline validation before deployment. Run automated compliance checks post-deploy.",
    },
}

USE_CASES = {
    "Use Case 1 - End-User Authentication & Authorization": [
        "User leverages Multi-Factor Authentication (MFA) to establish identity via a centralised Identity Provider.",
        "Device establishes identity via certificate-based authentication (PKI/CAC).",
        "Health & Compliance tool confirms security posture, patch status, and required tool presence.",
        "Real-time contextual access control is determined by validated identity + device health score.",
        "Network authorisation granted based on contextual access control policy.",
        "All authentication & authorisation activities are logged to SIEM (Splunk/Sentinel).",
    ],
    "Use Case 2 - Application & Data Segmentation": [
        "Software-Defined Networking (SDN) enables MicroCore & Perimeter (MCAP) for switching/routing + firewall.",
        "Each application tier (Web, App, DB) enforces granular port/protocol firewall policy.",
        "Application processes are learned and whitelisted for legitimate port/protocol usage.",
        "Data is tagged and specifically associated with owning systems/applications.",
        "All authorisation activities are logged to SIEM.",
        "All authentication & authorisation occur against the centralised Identity Provider.",
    ],
    "Use Case 3 - Admin Authentication & Authorization": [
        "Administrator authenticates via MFA against a centralised Identity Provider.",
        "Just-In-Time (JIT) gateway grants time-limited access to specific systems upon valid ITSM ticket.",
        "Systems managed via JIT Gateway or API-driven orchestration tool over isolated management segment.",
        "All authorisation activities logged to SIEM.",
        "Centralised Identity Provider leveraged for all identities and authorisations.",
    ],
    "Use Case 4 - Security Policy Orchestration": [
        "SDE Global Orchestrator enables service delivery and cyber operations across the agency.",
        "Service-level orchestrators deploy specific systems and associated granular security policies via API.",
        "Policy Enforcement Engines accept policy via REST API and enforce on infrastructure.",
        "All orchestration activities logged to SIEM.",
        "Orchestrated policies include deployment of identity and roles on centralised Identity Provider.",
    ],
}

ARCHITECTURAL_REQUIREMENTS = {
    "Identity Requirements": [
        "1. The system shall NOT assume trust of the device or user.",
        "1.1 The system shall have the capability to identify the user.",
        "1.2 The system shall have the capability to identify the device.",
        "1.2.1 The system shall be able to identify client and server endpoints.",
        "1.2.2 The system shall be able to identify network devices.",
        "2. The system shall integrate with an enterprise identity provider.",
        "2.1 The system shall integrate with industry-standard providers: AD, RADIUS, LDAP, OAuth, OpenID Connect, SAML.",
        "3. The system shall support enterprise authentication methods.",
        "3.1 The system shall support certificate-based authentication.",
        "3.2 The system shall support Public Key Infrastructure (PKI).",
        "3.3 The system shall support or integrate with Multi-Factor Authentication (MFA).",
        "4. The system shall have the ability to continuously challenge identity.",
        "4.1 The system shall support multiple attributes (e.g., geolocation) to challenge identity.",
        "5. The system shall connect to a centralised identity source.",
        "6. The system shall redirect or drop unidentified users or machines.",
    ],
    "Health & Compliance Requirements": [
        "1. The system shall determine endpoint health & compliance.",
        "1.1 The system shall determine OS STIG compliance posture.",
        "1.2 The system shall determine patch level of the device.",
        "1.3 The system shall determine if required tools/agents are missing.",
        "1.4 The system shall detect unauthorised software on the device.",
        "1.5 The system shall determine FQDN, IP, MAC address, and OS of the device.",
        "2. The system shall place non-compliant machines in remediation.",
        "3. The system shall support health & compliance checks across multiple device types.",
    ],
    "Authorization Requirements": [
        "1. The system shall control authorisation of users and machines.",
        "1.1 The system shall assign role-based access controls (RBAC) to groups and users.",
        "2. The system shall support time-limited authorisation for administrators.",
        "2.1 The system shall provide a portal to request privileged access (JIT).",
        "3. The system shall associate identity to authorisation policy via ZTA tool integration.",
        "3.1 The system shall align network security policy to authorisation of users and machines.",
        "4. The system shall support dynamic/real-time access decisions.",
        "5. The system shall support querying to determine authorisation level of user or device.",
    ],
    "Accounting & Auditing Requirements": [
        "1. The system shall provide accounting of all elements within the ZTA.",
        "1.1 The system shall account for user and device identity events.",
        "1.2 The system shall account for each authorisation grant or change.",
        "1.3 The system shall account for each segmentation policy change.",
        "2. The system shall account for all network flows within the ZTA.",
        "2.1 The system shall provide packet capture capability.",
        "2.2 The system shall log and account for all production data flow transactions.",
        "2.3 The system shall support offloading to enterprise SIEM tools.",
        "3. The system shall provide log data via REST API to a log aggregator.",
    ],
    "Segmentation Requirements": [
        "1. The system shall provide network micro-segmentation of application systems.",
        "1.1 The system shall apply role-based rule sets to groups and individual systems.",
        "1.2 The system shall segment north-south network flows.",
        "1.3 The system shall segment east-west network flows.",
        "1.4 The system shall segment end-to-end / transport network flows.",
        "2. The system shall segment both stateless and stateful application flows.",
        "3. The system shall segment host-level processes for security policies.",
    ],
    "Orchestration Requirements": [
        "1. The system shall support automation through REST APIs.",
        "1.1 The system shall accept REST API calls to configure identity, security, and accounting policies.",
        "2. The system shall support dynamic policy enforcement and changes.",
    ],
    "Cloud, Data Tagging & Discovery Requirements": [
        "1. The system shall provide the same protections across cloud-hosted environments.",
        "1.1 The system shall support CSP IaaS offerings.",
        "1.2 The system shall support CSP PaaS offerings.",
        "1.3 The system shall support CSP SaaS offerings.",
        "1.4 The system shall protect/proxy CSP administrative access.",
        "2. The system shall support tagging of data.",
        "3. The system shall support asset discovery.",
        "3.1 The system shall discover assets on the network.",
        "3.2 The system shall discover flows between assets on the network.",
    ],
    "Optional Requirements": [
        "1. The system shall replace existing security architecture constructs.",
        "1.1 The system shall replace existing perimeter security capabilities.",
        "2. The system shall integrate with existing security architecture constructs.",
        "2.1 The system shall integrate with Web Application Firewalls (WAF).",
        "2.2 The system shall integrate with application-aware firewalls (NGFW).",
        "2.3 The system shall integrate with database firewalls.",
        "2.4 The system shall integrate with Intrusion Prevention Systems (IPS).",
    ],
}

EVALUATION_CRITERIA = {
    "Effectiveness": {
        "definition": "Ability to accomplish mission objectives and achievement of desired results -- aligned to ZTA requirements.",
        "metrics": ["% of ZTA pillars fully implemented", "Mean Time to Detect (MTTD)", "Mean Time to Respond (MTTR)", "% reduction in lateral movement incidents"],
        "target": ">= 90% pillar coverage; MTTD < 60 min; MTTR < 4 hrs",
        "recommendation": "Conduct quarterly ZTA maturity assessments against CISA ZT Maturity Model. Map every control to NIST SP 800-207.",
    },
    "Suitability": {
        "definition": "Ability of the solution to be supported in the intended operational environment, aligned to common operational requirements (automation, architecture, security).",
        "metrics": ["Vendor SLA uptime (%)", "Integration API coverage (%)", "STIG compliance score (%)", "Number of supported identity protocols"],
        "target": "99.9% uptime; 100% REST API coverage; STIG score >= 85%",
        "recommendation": "Evaluate solutions against DoD APL. Require vendor FIPS 140-2/3 certification. Include operational sustainment costs in TCO model.",
    },
    "Performance": {
        "definition": "Measure of system performance expressed in quantifiable form -- focused on deployment footprint and scalability.",
        "metrics": ["Authentication latency (ms)", "Policy enforcement throughput (Gbps)", "Simultaneous sessions supported", "Auto-scale response time (sec)"],
        "target": "Auth latency < 200ms; throughput >= 10 Gbps; >= 100K concurrent sessions",
        "recommendation": "Load-test at 150% of projected peak. Define auto-scaling triggers. Monitor with real-time dashboards (Grafana/Prometheus).",
    },
}

# ══════════════════════════════════════════════════════════════════════════════
#  SYNTHETIC DATA
# ══════════════════════════════════════════════════════════════════════════════
@st.cache_data
def generate_synthetic_data(n: int) -> pd.DataFrame:
    random.seed(42)
    pillars = ["Identity & Auth", "Health & Compliance", "Authorization", "Accounting", "Segmentation", "Orchestration"]
    device_types = ["Workstation", "Server", "Mobile", "IoT Sensor", "Network Device", "Cloud VM"]
    departments = ["CYBERCOM", "DISA/BDE5", "TRANSCOM", "DIA", "NSA", "SOCOM", "CENTCOM", "EUCOM"]
    statuses = ["Compliant", "Non-Compliant", "Remediation", "Pending Review"]
    status_weights = [0.60, 0.20, 0.12, 0.08]
    os_list = ["Windows 11 STIG", "RHEL 9 STIG", "macOS Ventura", "Ubuntu 22.04 LTS", "Windows Server 2022"]
    base_date = datetime(2024, 1, 1)
    records = []
    for i in range(1, n + 1):
        status = random.choices(statuses, weights=status_weights)[0]
        score = random.randint(60, 100) if status == "Compliant" else random.randint(20, 59)
        records.append({
            "Record_ID": f"ZT-{i:04d}",
            "Timestamp": base_date + timedelta(days=random.randint(0, 364), hours=random.randint(0, 23)),
            "Pillar": random.choice(pillars),
            "Device_Type": random.choice(device_types),
            "Department": random.choice(departments),
            "OS": random.choice(os_list),
            "Compliance_Status": status,
            "STIG_Score": score,
            "Auth_Latency_ms": random.randint(45, 380),
            "Patch_Age_Days": random.randint(0, 90),
            "MFA_Enabled": random.choice([True, True, True, False]),
            "Segmentation_Policy_Applied": random.choice([True, True, False]),
            "Incident_Flagged": status == "Non-Compliant" and random.random() < 0.4,
        })
    return pd.DataFrame(records)

# ══════════════════════════════════════════════════════════════════════════════
#  EXPORT HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def build_text_report(selected_category: str, content: dict) -> str:
    lines = [
        "=" * 80,
        "  ZERO-TRUST FRAMEWORK APPLICATION",
        "  Developed by Randy Singh | KalSnet (KNet) Consulting Group",
        f"  Category: {selected_category}",
        f"  Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "=" * 80, "",
    ]
    for key, val in content.items():
        lines.append(f"[{key}]")
        if isinstance(val, list):
            for item in val:
                lines.append(f"  * {item}")
        elif isinstance(val, dict):
            for k2, v2 in val.items():
                if isinstance(v2, list):
                    lines.append(f"  {k2}:")
                    for item in v2:
                        lines.append(f"    - {item}")
                else:
                    lines.append(f"  {k2}: {v2}")
        else:
            lines.append(f"  {val}")
        lines.append("")
    return "\n".join(lines)


def build_json_report(selected_category: str, content: dict) -> str:
    payload = {
        "framework": "Zero-Trust Framework Application",
        "author": "Randy Singh | KalSnet (KNet) Consulting Group",
        "category": selected_category,
        "generated": datetime.now().isoformat(),
        "data": content,
    }
    return json.dumps(payload, indent=2, default=str)


def _pdf_safe(text: str) -> str:
    """
    Sanitise a string so it contains only Latin-1-encodable characters.
    Replaces common Unicode symbols with ASCII equivalents, then drops
    anything still outside Latin-1 (including emoji, box-drawing chars, etc.)
    """
    replacements = [
        ("\u2013", "-"), ("\u2014", "--"),
        ("\u2018", "'"), ("\u2019", "'"),
        ("\u201c", '"'), ("\u201d", '"'),
        ("\u2026", "..."), ("\u2022", "*"),
        ("\u00a0", " "), ("\u2265", ">="),
        ("\u2264", "<="), ("\u00ae", "(R)"),
        ("\u2122", "(TM)"), ("\u00e9", "e"),
        ("\u00e8", "e"), ("\u00e0", "a"),
        ("\u00fc", "u"), ("\u00f6", "o"),
        ("\u00e4", "a"), ("\u00b0", " deg"),
        ("\u00b7", "-"), ("\u2212", "-"),
    ]
    for char, replacement in replacements:
        text = text.replace(char, replacement)
    # Strip anything still outside Latin-1
    return text.encode("latin-1", errors="ignore").decode("latin-1")


def build_pdf_bytes(selected_category: str, content: dict) -> bytes:
    """
    Build PDF and return bytes.
    Handles both old fpdf2 (output()->str) and new fpdf2 (output()->bytearray).
    Uses new_x/new_y API instead of deprecated ln= parameter.
    """
    if not FPDF_AVAILABLE:
        return b""
    try:
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        # Header
        pdf.set_font("Helvetica", "B", 16)
        pdf.set_text_color(0, 71, 171)
        pdf.cell(0, 10, _pdf_safe("ZERO-TRUST FRAMEWORK APPLICATION"),
                 new_x="LMARGIN", new_y="NEXT", align="C")

        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 7, _pdf_safe("Developed by Randy Singh | KalSnet (KNet) Consulting Group"),
                 new_x="LMARGIN", new_y="NEXT", align="C")

        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(80, 80, 80)
        ts = _pdf_safe(f"Category: {selected_category}   |   Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        pdf.cell(0, 6, ts, new_x="LMARGIN", new_y="NEXT", align="C")

        pdf.set_draw_color(0, 71, 171)
        pdf.set_line_width(0.8)
        y = pdf.get_y() + 2
        pdf.line(10, y, 200, y)
        pdf.ln(6)

        # Content
        for key, val in content.items():
            pdf.set_font("Helvetica", "B", 11)
            pdf.set_text_color(0, 71, 171)
            pdf.cell(0, 8, _pdf_safe(str(key)), new_x="LMARGIN", new_y="NEXT")

            pdf.set_font("Helvetica", "", 9)
            pdf.set_text_color(30, 30, 30)

            if isinstance(val, list):
                for item in val:
                    pdf.multi_cell(0, 6, _pdf_safe(f"  * {item}"))
            elif isinstance(val, dict):
                for k2, v2 in val.items():
                    pdf.set_font("Helvetica", "B", 9)
                    pdf.cell(0, 6, _pdf_safe(f"  {k2}:"), new_x="LMARGIN", new_y="NEXT")
                    pdf.set_font("Helvetica", "", 9)
                    if isinstance(v2, list):
                        for item in v2:
                            pdf.multi_cell(0, 5, _pdf_safe(f"    - {item}"))
                    else:
                        pdf.multi_cell(0, 6, _pdf_safe(f"    {v2}"))
            else:
                pdf.multi_cell(0, 6, _pdf_safe(f"  {val}"))

            pdf.ln(2)

        raw = pdf.output()
        # fpdf2 >= 2.x returns bytearray; older versions returned a latin-1 str
        if isinstance(raw, (bytes, bytearray)):
            return bytes(raw)
        return raw.encode("latin-1")

    except Exception as exc:
        # Return a minimal valid PDF with the error message instead of crashing
        try:
            fallback = FPDF()
            fallback.add_page()
            fallback.set_font("Helvetica", "", 11)
            fallback.multi_cell(0, 8, _pdf_safe(f"PDF error: {exc}"))
            raw = fallback.output()
            if isinstance(raw, (bytes, bytearray)):
                return bytes(raw)
            return raw.encode("latin-1")
        except Exception:
            return b""


def build_docx_bytes(selected_category: str, content: dict) -> bytes:
    if not DOCX_AVAILABLE:
        return b""
    try:
        doc = DocxDocument()

        # Title
        t = doc.add_heading("ZERO-TRUST FRAMEWORK APPLICATION", 0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in t.runs:
            run.font.color.rgb = RGBColor(0, 71, 171)

        sub = doc.add_paragraph("Developed by Randy Singh | KalSnet (KNet) Consulting Group")
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if sub.runs:
            sub.runs[0].bold = True
            sub.runs[0].font.color.rgb = RGBColor(0, 71, 171)

        ts_para = doc.add_paragraph(
            f"Category: {selected_category}   |   {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        )
        ts_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        for key, val in content.items():
            h = doc.add_heading(str(key), level=2)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0, 71, 171)

            if isinstance(val, list):
                for item in val:
                    doc.add_paragraph(str(item), style="List Bullet")
            elif isinstance(val, dict):
                for k2, v2 in val.items():
                    p = doc.add_paragraph()
                    run_label = p.add_run(f"{k2}: ")
                    run_label.bold = True
                    if isinstance(v2, list):
                        p.add_run(", ".join(str(x) for x in v2))
                    else:
                        p.add_run(str(v2))
            else:
                doc.add_paragraph(str(val))
            doc.add_paragraph()

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    except Exception as exc:
        # Return an error document instead of crashing
        try:
            doc = DocxDocument()
            doc.add_paragraph(f"Word export error: {exc}")
            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue()
        except Exception:
            return b""

# ══════════════════════════════════════════════════════════════════════════════
#  SIDEBAR — NAVIGATION
# ══════════════════════════════════════════════════════════════════════════════
with st.sidebar:
    st.markdown("## ZTA Navigator")
    st.markdown("---")
    category = st.selectbox(
        "Select Framework Category",
        [
            "Overview & Architecture Diagram",
            "1. Functional Requirements",
            "2. Proposed Use Cases",
            "3. Architectural Requirements",
            "4. Evaluation Criteria",
            "Synthetic Data & Analytics",
        ],
    )
    st.markdown("---")
    st.markdown("**About**")
    st.markdown("Randy Singh  \nComputer Scientist  \nDISA / BDE5  \nKalSnet (KNet) Consulting")
    st.markdown("`(301) 225-9535`")

# ══════════════════════════════════════════════════════════════════════════════
#  SIDEBAR — EXPORT
# ══════════════════════════════════════════════════════════════════════════════
def get_export_content(cat: str) -> dict:
    if "Functional" in cat:
        return {
            k: {
                "Requirement": v["requirement"],
                "Example": v["example"],
                "Solution": v["solution"],
                "Recommendation": v["recommendation"],
            }
            for k, v in FUNCTIONAL_REQUIREMENTS.items()
        }
    elif "Use Case" in cat:
        return USE_CASES
    elif "Architectural" in cat:
        return ARCHITECTURAL_REQUIREMENTS
    elif "Evaluation" in cat:
        return {
            k: {
                "Definition": v["definition"],
                "Metrics": v["metrics"],
                "Target": v["target"],
                "Recommendation": v["recommendation"],
            }
            for k, v in EVALUATION_CRITERIA.items()
        }
    else:
        return {"Framework": "Zero-Trust Architecture", "Version": "2.0", "Author": "Randy Singh / KNet"}


with st.sidebar:
    st.markdown("---")
    st.markdown("### Export Results")

    exp_content = get_export_content(category)
    clean_cat = (category
                 .replace("1. ", "").replace("2. ", "").replace("3. ", "").replace("4. ", ""))

    # Text — always available
    txt_data = build_text_report(clean_cat, exp_content)
    st.download_button(
        "Export as Text (.txt)",
        data=txt_data,
        file_name="zt_framework.txt",
        mime="text/plain",
        use_container_width=True,
    )

    # JSON — always available
    json_data = build_json_report(clean_cat, exp_content)
    st.download_button(
        "Export as JSON (.json)",
        data=json_data,
        file_name="zt_framework.json",
        mime="application/json",
        use_container_width=True,
    )

    # PDF
    if FPDF_AVAILABLE:
        pdf_bytes = build_pdf_bytes(clean_cat, exp_content)
        st.download_button(
            "Export as PDF (.pdf)",
            data=pdf_bytes,
            file_name="zt_framework.pdf",
            mime="application/pdf",
            use_container_width=True,
        )
    else:
        st.warning("PDF unavailable. Run: pip install fpdf2")

    # Word
    if DOCX_AVAILABLE:
        docx_bytes = build_docx_bytes(clean_cat, exp_content)
        st.download_button(
            "Export as Word (.docx)",
            data=docx_bytes,
            file_name="zt_framework.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
    else:
        st.warning("Word unavailable. Run: pip install python-docx")

# ══════════════════════════════════════════════════════════════════════════════
#  MAIN CONTENT
# ══════════════════════════════════════════════════════════════════════════════

# ── OVERVIEW ──────────────────────────────────────────────────────────────────
if "Overview" in category:
    st.markdown('<div class="section-header">Zero-Trust Architecture - Conceptual Overview</div>', unsafe_allow_html=True)

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("ZTA Pillars", "6", "Core Controls")
    c2.metric("Use Cases", "4", "Defined")
    c3.metric("Arch. Req. Areas", "7", "Categories")
    c4.metric("Eval. Criteria", "3", "Dimensions")

    st.markdown("---")
    st.markdown("#### Zero-Trust Architecture Flow Diagram")

    try:
        import graphviz
        dot = graphviz.Digraph(graph_attr={"rankdir": "LR", "bgcolor": "#f8f9ff", "fontname": "Helvetica"})
        dot.attr("node", shape="box", style="filled", fontname="Helvetica", fontsize="11")
        dot.node("User",  "User / Device",           fillcolor="#dce8ff", color="#0047AB")
        dot.node("MFA",   "MFA",                     fillcolor="#cce5ff", color="#0047AB")
        dot.node("IdP",   "Identity Provider (IdP)", fillcolor="#0047AB", fontcolor="white", color="#0047AB")
        dot.node("NAC",   "Health & Compliance",     fillcolor="#fff0cc", color="#f0a500")
        dot.node("PAM",   "JIT / PAM Authorization", fillcolor="#e8f5e9", color="#2e7d32")
        dot.node("PEP",   "Policy Enforcement (PEP)",fillcolor="#0047AB", fontcolor="white", color="#0047AB")
        dot.node("App",   "Application Tier",        fillcolor="#dce8ff", color="#0047AB")
        dot.node("SIEM",  "SIEM / Accounting",       fillcolor="#fce4ec", color="#c62828")
        dot.node("Orch",  "Orchestration (REST API)",fillcolor="#f3e5f5", color="#6a1b9a")
        dot.edge("User", "MFA",  label="Auth Request")
        dot.edge("MFA",  "IdP",  label="Credentials")
        dot.edge("User", "NAC",  label="Device Posture")
        dot.edge("NAC",  "IdP",  label="Health Signal")
        dot.edge("IdP",  "PAM",  label="Identity Assert")
        dot.edge("PAM",  "PEP",  label="Auth Decision")
        dot.edge("PEP",  "App",  label="Access Granted")
        dot.edge("PEP",  "SIEM", label="Log Events")
        dot.edge("Orch", "PEP",  label="Policy Push")
        dot.edge("Orch", "IdP",  label="Role Provision")
        dot.edge("App",  "SIEM", label="App Logs")
        st.graphviz_chart(dot, use_container_width=True)
    except Exception:
        st.code("""
USER/DEVICE ---> MFA ----------------> IdP ---------> PAM/JIT
     |                                  |                  |
     +---> NAC (Health Check) ----------+                  |
                                                           v
ORCHESTRATION (REST API) ----------------------> POLICY ENFORCEMENT POINT
                                                           |
                                          +----------------+----------------+
                                          v                                 v
                                    APPLICATION TIERS                 SIEM (ACCOUNTING)
        """)

    st.markdown("---")
    st.markdown("#### ZTA Pillar Maturity Radar")
    pillars_r = ["Identity & Auth", "Health & Compliance", "Authorization", "Accounting", "Segmentation", "Orchestration"]
    current   = [3.5, 2.8, 3.0, 3.8, 2.5, 2.2]
    target    = [5, 5, 5, 5, 5, 5]
    fig_radar = go.Figure()
    fig_radar.add_trace(go.Scatterpolar(
        r=current + [current[0]], theta=pillars_r + [pillars_r[0]],
        fill="toself", name="Current Maturity",
        line_color="#0047AB", fillcolor="rgba(0,71,171,0.2)"))
    fig_radar.add_trace(go.Scatterpolar(
        r=target + [target[0]], theta=pillars_r + [pillars_r[0]],
        fill="toself", name="Target Maturity",
        line_color="#f0a500", fillcolor="rgba(240,165,0,0.1)"))
    fig_radar.update_layout(
        polar=dict(radialaxis=dict(visible=True, range=[0, 5])),
        title="ZTA Pillar Maturity (1=Initial  5=Optimised)",
        height=420)
    st.plotly_chart(fig_radar, use_container_width=True)

    st.markdown("#### ZTA Core Principles")
    principles = [
        ("Never Trust, Always Verify", "Treat every request as untrusted regardless of network location."),
        ("Least-Privilege Access", "Grant minimum permissions necessary and expire them promptly."),
        ("Assume Breach", "Design controls assuming attackers may already be inside the perimeter."),
        ("Continuous Verification", "Re-evaluate trust on every transaction, not just at login."),
        ("Micro-Segmentation", "Divide the network into small zones to contain lateral movement."),
        ("Policy-as-Code", "Automate and version-control all security policies via REST APIs."),
    ]
    for title_p, desc in principles:
        st.markdown(f'<div class="req-card"><b>{title_p}</b><br>{desc}</div>', unsafe_allow_html=True)


# ── FUNCTIONAL REQUIREMENTS ───────────────────────────────────────────────────
elif "Functional" in category:
    st.markdown('<div class="section-header">Category 1 - Functional Requirements</div>', unsafe_allow_html=True)
    show_all = st.checkbox("Expand all pillars", value=True)
    for pillar, data in FUNCTIONAL_REQUIREMENTS.items():
        with st.expander(f"{pillar}", expanded=show_all):
            st.markdown(f'<div class="req-card"><b>Requirement:</b><br>{data["requirement"]}</div>', unsafe_allow_html=True)
            st.markdown(f'<div class="sol-card"><b>Example:</b><br>{data["example"]}</div>', unsafe_allow_html=True)
            st.markdown(f'<div class="sol-card"><b>Recommended Solution:</b><br>{data["solution"]}</div>', unsafe_allow_html=True)
            st.markdown(f'<div class="rec-card"><b>Recommendation:</b><br>{data["recommendation"]}</div>', unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("#### Pillar Coverage (Illustrative)")
    coverage = {"Identity & Authentication": 85, "Health & Compliance": 70, "Authorization": 78,
                "Accounting": 90, "Segmentation": 62, "Orchestration": 55}
    fig_bar = px.bar(x=list(coverage.keys()), y=list(coverage.values()),
                     labels={"x": "Pillar", "y": "Coverage (%)"},
                     color=list(coverage.values()),
                     color_continuous_scale=["#ff4444", "#f0a500", "#0047AB"],
                     title="Estimated Pillar Implementation Coverage (%)")
    fig_bar.update_layout(coloraxis_showscale=False, height=360)
    st.plotly_chart(fig_bar, use_container_width=True)


# ── USE CASES ─────────────────────────────────────────────────────────────────
elif "Use Case" in category:
    st.markdown('<div class="section-header">Category 2 - Proposed Use Cases</div>', unsafe_allow_html=True)
    uc_choice = st.radio("Select Use Case", list(USE_CASES.keys()))
    st.markdown(f"### {uc_choice}")
    for i, step in enumerate(USE_CASES[uc_choice], 1):
        st.markdown(f'<div class="req-card"><b>Step {i}:</b> {step}</div>', unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("#### Use Case Flow Diagram")
    try:
        import graphviz
        g = graphviz.Digraph(graph_attr={"rankdir": "TD"})
        g.attr("node", shape="box", style="filled", fontname="Helvetica", fontsize="10")
        steps_uc = USE_CASES[uc_choice]
        for idx, step_text in enumerate(steps_uc):
            label = step_text[:55] + "..." if len(step_text) > 55 else step_text
            g.node(str(idx), label, fillcolor="#dce8ff", color="#0047AB")
        for idx in range(len(steps_uc) - 1):
            g.edge(str(idx), str(idx + 1))
        st.graphviz_chart(g, use_container_width=True)
    except Exception:
        st.info("Install the graphviz system package for flow diagrams.")


# ── ARCHITECTURAL REQUIREMENTS ────────────────────────────────────────────────
elif "Architectural" in category:
    st.markdown('<div class="section-header">Category 3 - Architectural Requirements</div>', unsafe_allow_html=True)
    area_choice = st.selectbox("Select Requirement Area", list(ARCHITECTURAL_REQUIREMENTS.keys()))
    st.markdown(f"### {area_choice}")
    for req in ARCHITECTURAL_REQUIREMENTS[area_choice]:
        st.markdown(f'<div class="req-card">{req}</div>', unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("#### Requirement Count per Area")
    counts = {k: len(v) for k, v in ARCHITECTURAL_REQUIREMENTS.items()}
    fig_h = px.bar(x=list(counts.values()), y=list(counts.keys()), orientation="h",
                   labels={"x": "# Requirements", "y": "Area"},
                   color=list(counts.values()), color_continuous_scale=["#90caf9", "#0047AB"],
                   title="Architectural Requirements Count by Area")
    fig_h.update_layout(coloraxis_showscale=False, height=380, yaxis={"categoryorder": "total ascending"})
    st.plotly_chart(fig_h, use_container_width=True)

    st.markdown("#### Reference Standards")
    refs = [
        ("NIST SP 800-207", "Zero Trust Architecture - foundational ZTA standard."),
        ("CISA ZT Maturity Model", "5-pillar model for assessing ZTA implementation progress."),
        ("DoD ZT Strategy", "DoD mandate for ZT implementation across all components by FY27."),
        ("DISA STIG Library", "Security Technical Implementation Guides for OS/app hardening."),
        ("NIST SP 800-53 Rev 5", "Security and Privacy Controls - maps to ZTA control requirements."),
    ]
    for ref, desc in refs:
        st.markdown(f'<div class="rec-card"><b>{ref}</b> - {desc}</div>', unsafe_allow_html=True)


# ── EVALUATION CRITERIA ───────────────────────────────────────────────────────
elif "Evaluation" in category:
    st.markdown('<div class="section-header">Category 4 - Evaluation Criteria</div>', unsafe_allow_html=True)
    for criterion, data in EVALUATION_CRITERIA.items():
        with st.expander(f"{criterion}", expanded=True):
            st.markdown(f'<div class="req-card"><b>Definition:</b> {data["definition"]}</div>', unsafe_allow_html=True)
            st.markdown(f'<div class="sol-card"><b>Target:</b> {data["target"]}</div>', unsafe_allow_html=True)
            st.markdown("<b>Key Metrics:</b>", unsafe_allow_html=True)
            for m in data["metrics"]:
                st.markdown(f'<div class="req-card">{m}</div>', unsafe_allow_html=True)
            st.markdown(f'<div class="rec-card"><b>Recommendation:</b> {data["recommendation"]}</div>', unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("#### Sample Evaluation Scores (Illustrative)")
    gc1, gc2, gc3 = st.columns(3)
    for col, (crit, score, color) in zip(
        [gc1, gc2, gc3],
        [("Effectiveness", 78, "#0047AB"), ("Suitability", 85, "#2e7d32"), ("Performance", 71, "#f0a500")]
    ):
        fig_g = go.Figure(go.Indicator(
            mode="gauge+number", value=score, title={"text": crit},
            gauge={"axis": {"range": [0, 100]}, "bar": {"color": color},
                   "steps": [{"range": [0, 50], "color": "#ffcdd2"},
                              {"range": [50, 75], "color": "#fff9c4"},
                              {"range": [75, 100], "color": "#c8e6c9"}],
                   "threshold": {"line": {"color": "black", "width": 3}, "thickness": 0.8, "value": 80}}))
        fig_g.update_layout(height=260, margin=dict(t=40, b=10))
        col.plotly_chart(fig_g, use_container_width=True)


# ── SYNTHETIC DATA ────────────────────────────────────────────────────────────
elif "Synthetic" in category or "Analytics" in category or "Data" in category:
    st.markdown('<div class="section-header">Synthetic ZTA Compliance Data & Analytics</div>', unsafe_allow_html=True)
    num_records = st.slider("Number of Synthetic Records", min_value=0, max_value=300, value=150, step=10)

    if num_records == 0:
        st.info("Move the slider above 0 to generate synthetic data.")
    else:
        df = generate_synthetic_data(num_records)
        k1, k2, k3, k4, k5 = st.columns(5)
        k1.metric("Total Records", num_records)
        k2.metric("Compliant", int((df["Compliance_Status"] == "Compliant").sum()),
                  f'{(df["Compliance_Status"]=="Compliant").mean()*100:.0f}%')
        k3.metric("Non-Compliant", int((df["Compliance_Status"] == "Non-Compliant").sum()))
        k4.metric("Avg STIG Score", f"{df['STIG_Score'].mean():.1f}")
        k5.metric("Incidents Flagged", int(df["Incident_Flagged"].sum()))

        st.markdown("---")
        tab_a, tab_b, tab_c, tab_d = st.tabs(["Raw Data", "By Pillar", "By Department", "Trends"])

        STATUS_COLORS = {"Compliant": "#2e7d32", "Non-Compliant": "#c62828",
                         "Remediation": "#f0a500", "Pending Review": "#757575"}

        with tab_a:
            st.dataframe(df, use_container_width=True, height=380)
            st.download_button("Download CSV", df.to_csv(index=False).encode(),
                               "zt_synthetic_data.csv", "text/csv")

        with tab_b:
            c1, c2 = st.columns(2)
            pillar_status = df.groupby(["Pillar", "Compliance_Status"]).size().reset_index(name="Count")
            fig_ps = px.bar(pillar_status, x="Pillar", y="Count", color="Compliance_Status",
                            barmode="group", color_discrete_map=STATUS_COLORS,
                            title="Compliance Status by ZTA Pillar")
            c1.plotly_chart(fig_ps, use_container_width=True)
            stig_pillar = df.groupby("Pillar")["STIG_Score"].mean().reset_index()
            fig_sp = px.bar(stig_pillar, x="Pillar", y="STIG_Score",
                            color="STIG_Score", color_continuous_scale=["#ff4444", "#f0a500", "#2e7d32"],
                            title="Average STIG Score by Pillar", range_y=[0, 100])
            fig_sp.update_layout(coloraxis_showscale=False)
            c2.plotly_chart(fig_sp, use_container_width=True)

        with tab_c:
            dept_status = df.groupby(["Department", "Compliance_Status"]).size().reset_index(name="Count")
            fig_ds = px.bar(dept_status, x="Department", y="Count", color="Compliance_Status",
                            barmode="stack", color_discrete_map=STATUS_COLORS,
                            title="Compliance Status by Department")
            st.plotly_chart(fig_ds, use_container_width=True)
            dev_counts = df["Device_Type"].value_counts().reset_index()
            dev_counts.columns = ["Device_Type", "Count"]
            fig_pie = px.pie(dev_counts, values="Count", names="Device_Type",
                             title="Device Type Distribution",
                             color_discrete_sequence=px.colors.qualitative.Bold)
            st.plotly_chart(fig_pie, use_container_width=True)

        with tab_d:
            df["Month"] = pd.to_datetime(df["Timestamp"]).dt.to_period("M").astype(str)
            trend = df.groupby(["Month", "Compliance_Status"]).size().reset_index(name="Count")
            fig_trend = px.line(trend, x="Month", y="Count", color="Compliance_Status",
                                markers=True, color_discrete_map=STATUS_COLORS,
                                title="Compliance Trend Over Time")
            st.plotly_chart(fig_trend, use_container_width=True)
            fig_sc = px.scatter(df, x="STIG_Score", y="Auth_Latency_ms",
                                color="Compliance_Status", opacity=0.7,
                                color_discrete_map=STATUS_COLORS,
                                title="STIG Score vs Authentication Latency",
                                labels={"STIG_Score": "STIG Compliance Score",
                                        "Auth_Latency_ms": "Auth Latency (ms)"})
            st.plotly_chart(fig_sc, use_container_width=True)

# ══════════════════════════════════════════════════════════════════════════════
#  FOOTER
# ══════════════════════════════════════════════════════════════════════════════
st.markdown("---")
st.markdown(
    '<div style="text-align:center;color:#888;font-size:0.8rem;">'
    'Zero-Trust Framework Application &nbsp;|&nbsp; '
    'Developed by Randy Singh &nbsp;|&nbsp; '
    'KalSnet (KNet) Consulting Group &nbsp;|&nbsp; '
    f'DISA / BDE5 Technology Innovation Team &nbsp;|&nbsp; {datetime.now().year}'
    '</div>',
    unsafe_allow_html=True,
)
PYEOF
echo "Done — $(wc -l < /mnt/user-data/outputs/zero_trust_framework.py) lines written"
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ══════════════════════════════════════════════════════════════════════════════
#  PAGE CONFIG
# ══════════════════════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="Zero-Trust Framework | KNet Consulting",
    page_icon="🔒",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ══════════════════════════════════════════════════════════════════════════════
#  GLOBAL STYLES
# ══════════════════════════════════════════════════════════════════════════════
st.markdown("""
<style>
  /* ---- Title block ---- */
  .zt-title {
      text-align: center;
      font-size: 2.4rem;
      font-weight: 900;
      color: #0047AB;
      letter-spacing: 1px;
      margin-bottom: 0;
  }
  .zt-subtitle {
      text-align: center;
      font-size: 1.05rem;
      font-weight: 700;
      color: #0047AB;
      margin-top: 2px;
      margin-bottom: 4px;
  }
  .zt-tagline {
      text-align: center;
      color: #555;
      font-size: 0.88rem;
      margin-bottom: 18px;
  }
  /* ---- Divider ---- */
  .zt-hr { border: 2px solid #0047AB; margin: 8px 0 18px 0; }
  /* ---- Section headers ---- */
  .section-header {
      background: linear-gradient(90deg, #0047AB 0%, #1a6fe8 100%);
      color: white;
      padding: 8px 16px;
      border-radius: 6px;
      font-weight: 700;
      font-size: 1.05rem;
      margin: 10px 0 6px 0;
  }
  /* ---- Requirement cards ---- */
  .req-card {
      background: #f0f6ff;
      border-left: 4px solid #0047AB;
      border-radius: 4px;
      padding: 10px 14px;
      margin: 6px 0;
      font-size: 0.92rem;
  }
  /* ---- Recommendation card ---- */
  .rec-card {
      background: #fff8e1;
      border-left: 4px solid #f0a500;
      border-radius: 4px;
      padding: 10px 14px;
      margin: 6px 0;
      font-size: 0.92rem;
  }
  /* ---- Solution card ---- */
  .sol-card {
      background: #e8f5e9;
      border-left: 4px solid #2e7d32;
      border-radius: 4px;
      padding: 10px 14px;
      margin: 6px 0;
      font-size: 0.92rem;
  }
  /* ---- Metric pill ---- */
  div[data-testid="metric-container"] {
      background: #f0f6ff;
      border: 1px solid #c2d9ff;
      border-radius: 8px;
      padding: 10px;
  }
  /* ---- Sidebar ---- */
  section[data-testid="stSidebar"] { background: #001a6b; }
  section[data-testid="stSidebar"] * { color: #e8f0ff !important; }
  section[data-testid="stSidebar"] .stSelectbox label { color: #cde !important; }
</style>
""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
#  HEADER
# ══════════════════════════════════════════════════════════════════════════════
st.markdown('<p class="zt-title">🔒 ZERO-TRUST FRAMEWORK APPLICATION</p>', unsafe_allow_html=True)
st.markdown('<p class="zt-subtitle">Developed by Randy Singh | KalSnet (KNet) Consulting Group</p>', unsafe_allow_html=True)
st.markdown('<p class="zt-tagline">Based on DISA/BDE5 ZTA Working Group Requirements &amp; Use Cases &nbsp;|&nbsp; Originally authored April 2019</p>', unsafe_allow_html=True)
st.markdown('<hr class="zt-hr">', unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
#  DATA DEFINITIONS
# ══════════════════════════════════════════════════════════════════════════════

FUNCTIONAL_REQUIREMENTS = {
    "Identity & Authentication": {
        "requirement": "Need to associate identity to user and device.",
        "example": "Active Directory + certificate-based device identity via PKI/CAC cards.",
        "solution": "Deploy a centralized Identity Provider (IdP) such as Microsoft Entra ID or Okta integrated with LDAP/SAML.",
        "recommendation": "Enforce MFA for every user login. Bind device certificates to hardware TPM chips to prevent credential theft.",
    },
    "Health & Compliance": {
        "requirement": "Need to determine patch level and security posture of device prior to network authorization.",
        "example": "NAC solution checks OS STIG compliance, AV status, and missing agents before granting network access.",
        "solution": "Integrate a Network Access Control (NAC) platform such as Cisco ISE or Forescout with your IdP.",
        "recommendation": "Define a minimum compliance baseline (patch age ≤30 days, STIG score ≥80%). Non-compliant devices land in an auto-remediation VLAN.",
    },
    "Authorization": {
        "requirement": "Need to authorize user/device, based on identity and hygiene, at micro-perimeter or Just-In-Time (JIT) for administrators.",
        "example": "PAM tool grants time-limited SSH/RDP session to admin only after ITSM ticket is validated.",
        "solution": "Deploy a Privileged Access Management (PAM) solution (CyberArk, BeyondTrust) with JIT provisioning tied to ServiceNow tickets.",
        "recommendation": "Apply least-privilege RBAC. Rotate JIT credentials every session. Log every privileged action to SIEM.",
    },
    "Accounting": {
        "requirement": "Need to identify and document resources consumed by users and administrators for compliance and incident response.",
        "example": "Splunk SIEM ingests auth logs, network flow data, and privileged session recordings.",
        "solution": "Centralise logging via a SIEM (Splunk, Microsoft Sentinel). Feed NetFlow data and endpoint telemetry.",
        "recommendation": "Retain logs for at least 12 months online, 36 months archived. Automate anomaly alerting with UEBA rules.",
    },
    "Segmentation": {
        "requirement": "Need to limit communication (directional port/protocol) between all network flows including application tiers.",
        "example": "Micro-segmentation policy allows only TCP 443 from Web tier to App tier; all other flows are denied.",
        "solution": "Implement software-defined micro-segmentation (VMware NSX, Illumio, or Guardicore). Define allow-lists per workload.",
        "recommendation": "Start with 'ring-fence' of crown-jewel systems. Expand micro-segmentation iteratively. Validate with automated penetration tests.",
    },
    "Orchestration": {
        "requirement": "Need capability to dynamically provision identity, hygiene, authorization, and segmentation policies.",
        "example": "Ansible playbook triggered by new VM spin-up auto-pushes firewall rules and registers device in IdP.",
        "solution": "Use a Security Orchestration platform (SOAR) or Infrastructure-as-Code (Terraform + Ansible) with REST API integration to all ZTA pillars.",
        "recommendation": "Treat policies as code. Store in Git. Enforce CI/CD pipeline validation before deployment. Run automated compliance checks post-deploy.",
    },
}

USE_CASES = {
    "Use Case 1 – End-User Authentication & Authorization": [
        "User leverages Multi-Factor Authentication (MFA) to establish identity via a centralised Identity Provider.",
        "Device establishes identity via certificate-based authentication (PKI/CAC).",
        "Health & Compliance tool confirms security posture, patch status, and required tool presence.",
        "Real-time contextual access control is determined by validated identity + device health score.",
        "Network authorisation granted based on contextual access control policy.",
        "All authentication & authorisation activities are logged to SIEM (Splunk/Sentinel).",
    ],
    "Use Case 2 – Application & Data Segmentation": [
        "Software-Defined Networking (SDN) enables MicroCore & Perimeter (MCAP) for switching/routing + firewall.",
        "Each application tier (Web → App → DB) enforces granular port/protocol firewall policy.",
        "Application processes are learned and whitelisted for legitimate port/protocol usage.",
        "Data is tagged and specifically associated with owning systems/applications.",
        "All authorisation activities are logged to SIEM.",
        "All authentication & authorisation occur against the centralised Identity Provider.",
    ],
    "Use Case 3 – Admin Authentication & Authorization": [
        "Administrator authenticates via MFA against a centralised Identity Provider.",
        "Just-In-Time (JIT) gateway grants time-limited access to specific systems upon valid ITSM ticket.",
        "Systems managed via JIT Gateway or API-driven orchestration tool over isolated management segment.",
        "All authorisation activities logged to SIEM.",
        "Centralised Identity Provider leveraged for all identities and authorisations.",
    ],
    "Use Case 4 – Security Policy Orchestration": [
        "SDE Global Orchestrator enables service delivery and cyber operations across the agency.",
        "Service-level orchestrators deploy specific systems and associated granular security policies via API.",
        "Policy Enforcement Engines accept policy via REST API and enforce on infrastructure.",
        "All orchestration activities logged to SIEM.",
        "Orchestrated policies include deployment of identity and roles on centralised Identity Provider.",
    ],
}

ARCHITECTURAL_REQUIREMENTS = {
    "Identity Requirements": [
        "1. The system shall NOT assume trust of the device or user.",
        "1.1 The system shall have the capability to identify the user.",
        "1.2 The system shall have the capability to identify the device.",
        "1.2.1 The system shall be able to identify client and server endpoints.",
        "1.2.2 The system shall be able to identify network devices.",
        "2. The system shall integrate with an enterprise identity provider.",
        "2.1 The system shall integrate with industry-standard providers: AD, RADIUS, LDAP, OAuth, OpenID Connect, SAML.",
        "3. The system shall support enterprise authentication methods.",
        "3.1 The system shall support certificate-based authentication.",
        "3.2 The system shall support Public Key Infrastructure (PKI).",
        "3.3 The system shall support or integrate with Multi-Factor Authentication (MFA).",
        "4. The system shall have the ability to continuously challenge identity.",
        "4.1 The system shall support multiple attributes (e.g., geolocation) to challenge identity.",
        "5. The system shall connect to a centralised identity source.",
        "6. The system shall redirect or drop unidentified users or machines.",
    ],
    "Health & Compliance Requirements": [
        "1. The system shall determine endpoint health & compliance.",
        "1.1 The system shall determine OS STIG compliance posture.",
        "1.2 The system shall determine patch level of the device.",
        "1.3 The system shall determine if required tools/agents are missing.",
        "1.4 The system shall detect unauthorised software on the device.",
        "1.5 The system shall determine FQDN, IP, MAC address, and OS of the device.",
        "2. The system shall place non-compliant machines in remediation.",
        "3. The system shall support health & compliance checks across multiple device types.",
    ],
    "Authorization Requirements": [
        "1. The system shall control authorisation of users and machines.",
        "1.1 The system shall assign role-based access controls (RBAC) to groups and users.",
        "2. The system shall support time-limited authorisation for administrators.",
        "2.1 The system shall provide a portal to request privileged access (JIT).",
        "3. The system shall associate identity to authorisation policy via ZTA tool integration.",
        "3.1 The system shall align network security policy to authorisation of users and machines.",
        "4. The system shall support dynamic/real-time access decisions.",
        "5. The system shall support querying to determine authorisation level of user or device.",
    ],
    "Accounting & Auditing Requirements": [
        "1. The system shall provide accounting of all elements within the ZTA.",
        "1.1 The system shall account for user and device identity events.",
        "1.2 The system shall account for each authorisation grant or change.",
        "1.3 The system shall account for each segmentation policy change.",
        "2. The system shall account for all network flows within the ZTA.",
        "2.1 The system shall provide packet capture capability.",
        "2.2 The system shall log and account for all production data flow transactions.",
        "2.3 The system shall support offloading to enterprise SIEM tools.",
        "3. The system shall provide log data via REST API to a log aggregator.",
    ],
    "Segmentation Requirements": [
        "1. The system shall provide network micro-segmentation of application systems.",
        "1.1 The system shall apply role-based rule sets to groups and individual systems.",
        "1.2 The system shall segment north–south network flows.",
        "1.3 The system shall segment east–west network flows.",
        "1.4 The system shall segment end-to-end / transport network flows.",
        "2. The system shall segment both stateless and stateful application flows.",
        "3. The system shall segment host-level processes for security policies.",
    ],
    "Orchestration Requirements": [
        "1. The system shall support automation through REST APIs.",
        "1.1 The system shall accept REST API calls to configure identity, security, and accounting policies.",
        "2. The system shall support dynamic policy enforcement and changes.",
    ],
    "Cloud, Data Tagging & Discovery Requirements": [
        "1. The system shall provide the same protections across cloud-hosted environments.",
        "1.1 The system shall support CSP IaaS offerings.",
        "1.2 The system shall support CSP PaaS offerings.",
        "1.3 The system shall support CSP SaaS offerings.",
        "1.4 The system shall protect/proxy CSP administrative access.",
        "2. The system shall support tagging of data.",
        "3. The system shall support asset discovery.",
        "3.1 The system shall discover assets on the network.",
        "3.2 The system shall discover flows between assets on the network.",
    ],
    "Optional Requirements": [
        "1. The system shall replace existing security architecture constructs.",
        "1.1 The system shall replace existing perimeter security capabilities.",
        "2. The system shall integrate with existing security architecture constructs.",
        "2.1 The system shall integrate with Web Application Firewalls (WAF).",
        "2.2 The system shall integrate with application-aware firewalls (NGFW).",
        "2.3 The system shall integrate with database firewalls.",
        "2.4 The system shall integrate with Intrusion Prevention Systems (IPS).",
    ],
}

EVALUATION_CRITERIA = {
    "Effectiveness": {
        "definition": "Ability to accomplish mission objectives and achievement of desired results — aligned to ZTA requirements.",
        "metrics": ["% of ZTA pillars fully implemented", "Mean Time to Detect (MTTD)", "Mean Time to Respond (MTTR)", "% reduction in lateral movement incidents"],
        "target": "≥ 90% pillar coverage; MTTD < 60 min; MTTR < 4 hrs",
        "recommendation": "Conduct quarterly ZTA maturity assessments against CISA ZT Maturity Model. Map every control to NIST SP 800-207.",
    },
    "Suitability": {
        "definition": "Ability of the solution to be supported in the intended operational environment, aligned to common operational requirements (automation, architecture, security).",
        "metrics": ["Vendor SLA uptime (%)", "Integration API coverage (%)", "STIG compliance score (%)", "Number of supported identity protocols"],
        "target": "99.9% uptime; 100% REST API coverage; STIG score ≥ 85%",
        "recommendation": "Evaluate solutions against DoD APL. Require vendor FIPS 140-2/3 certification. Include operational sustainment costs in TCO model.",
    },
    "Performance": {
        "definition": "Measure of system performance expressed in quantifiable form — focused on deployment footprint and scalability.",
        "metrics": ["Authentication latency (ms)", "Policy enforcement throughput (Gbps)", "Simultaneous sessions supported", "Auto-scale response time (sec)"],
        "target": "Auth latency < 200ms; throughput ≥ 10 Gbps; ≥ 100K concurrent sessions",
        "recommendation": "Load-test at 150% of projected peak. Define auto-scaling triggers. Monitor with real-time dashboards (Grafana/Prometheus).",
    },
}

# ══════════════════════════════════════════════════════════════════════════════
#  SYNTHETIC DATA GENERATOR  (0–300 records)
# ══════════════════════════════════════════════════════════════════════════════
@st.cache_data
def generate_synthetic_data(n: int) -> pd.DataFrame:
    """Generate n synthetic ZTA compliance records."""
    random.seed(42)
    pillars = ["Identity & Auth", "Health & Compliance", "Authorization", "Accounting", "Segmentation", "Orchestration"]
    device_types = ["Workstation", "Server", "Mobile", "IoT Sensor", "Network Device", "Cloud VM"]
    departments = ["CYBERCOM", "DISA/BDE5", "TRANSCOM", "DIA", "NSA", "SOCOM", "CENTCOM", "EUCOM"]
    statuses = ["Compliant", "Non-Compliant", "Remediation", "Pending Review"]
    status_weights = [0.60, 0.20, 0.12, 0.08]
    os_list = ["Windows 11 STIG", "RHEL 9 STIG", "macOS Ventura", "Ubuntu 22.04 LTS", "Windows Server 2022"]

    base_date = datetime(2024, 1, 1)
    records = []
    for i in range(1, n + 1):
        pillar = random.choice(pillars)
        status = random.choices(statuses, weights=status_weights)[0]
        score = random.randint(60, 100) if status == "Compliant" else random.randint(20, 59)
        records.append({
            "Record_ID": f"ZT-{i:04d}",
            "Timestamp": base_date + timedelta(days=random.randint(0, 364), hours=random.randint(0, 23)),
            "Pillar": pillar,
            "Device_Type": random.choice(device_types),
            "Department": random.choice(departments),
            "OS": random.choice(os_list),
            "Compliance_Status": status,
            "STIG_Score": score,
            "Auth_Latency_ms": random.randint(45, 380),
            "Patch_Age_Days": random.randint(0, 90),
            "MFA_Enabled": random.choice([True, True, True, False]),
            "Segmentation_Policy_Applied": random.choice([True, True, False]),
            "Incident_Flagged": status == "Non-Compliant" and random.random() < 0.4,
        })
    return pd.DataFrame(records)

# ══════════════════════════════════════════════════════════════════════════════
#  EXPORT HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def build_text_report(selected_category: str, content: dict) -> str:
    lines = [
        "=" * 80,
        "  ZERO-TRUST FRAMEWORK APPLICATION",
        "  Developed by Randy Singh | KalSnet (KNet) Consulting Group",
        f"  Category: {selected_category}",
        f"  Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "=" * 80, "",
    ]
    for key, val in content.items():
        lines.append(f"[{key}]")
        if isinstance(val, list):
            for item in val:
                lines.append(f"  • {item}")
        elif isinstance(val, dict):
            for k2, v2 in val.items():
                if isinstance(v2, list):
                    lines.append(f"  {k2}:")
                    for item in v2:
                        lines.append(f"    - {item}")
                else:
                    lines.append(f"  {k2}: {v2}")
        else:
            lines.append(f"  {val}")
        lines.append("")
    return "\n".join(lines)

def build_json_report(selected_category: str, content: dict) -> str:
    payload = {
        "framework": "Zero-Trust Framework Application",
        "author": "Randy Singh | KalSnet (KNet) Consulting Group",
        "category": selected_category,
        "generated": datetime.now().isoformat(),
        "data": content,
    }
    return json.dumps(payload, indent=2, default=str)

def _pdf_safe(text: str) -> str:
    """
    Convert a string so it contains only characters that FPDF's built-in
    Latin-1 fonts can encode.  Strategy:
      1. Replace common Unicode punctuation / symbols with ASCII equivalents.
      2. Drop any remaining non-latin-1 characters (including emoji).
    """
    replacements = {
        "\u2013": "-",   # en-dash
        "\u2014": "--",  # em-dash
        "\u2018": "'",   # left single quote
        "\u2019": "'",   # right single quote
        "\u201c": '"',   # left double quote
        "\u201d": '"',   # right double quote
        "\u2026": "...", # ellipsis
        "\u2022": "*",   # bullet
        "\u00a0": " ",   # non-breaking space
        "\u2265": ">=",  # >=
        "\u2264": "<=",  # <=
        "\u00ae": "(R)", # registered trademark
        "\u2122": "(TM)",# trademark
        "\u00e9": "e",   # é
        "\u00e8": "e",   # è
        "\u00e0": "a",   # à
        "\u00fc": "u",   # ü
        "\u00f6": "o",   # ö
        "\u00e4": "a",   # ä
        "&": "&",        # keep plain ampersand (latin-1 safe)
    }
    for char, replacement in replacements.items():
        text = text.replace(char, replacement)
    # Final pass: drop anything that still can't encode to latin-1
    return text.encode("latin-1", errors="ignore").decode("latin-1")


def build_pdf_bytes(selected_category: str, content: dict) -> bytes:
    if not FPDF_AVAILABLE:
        return b""

    pdf = FPDF()
    pdf.add_page()

    # ── Header ────────────────────────────────────────────────────────────────
    pdf.set_font("Helvetica", "B", 16)
    pdf.set_text_color(0, 71, 171)
    pdf.cell(0, 10, _pdf_safe("ZERO-TRUST FRAMEWORK APPLICATION"), ln=True, align="C")

    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(0, 7, _pdf_safe("Developed by Randy Singh | KalSnet (KNet) Consulting Group"), ln=True, align="C")

    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(80, 80, 80)
    pdf.cell(
        0, 6,
        _pdf_safe(f"Category: {selected_category}   |   Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"),
        ln=True, align="C",
    )

    pdf.set_draw_color(0, 71, 171)
    pdf.set_line_width(0.8)
    pdf.line(10, pdf.get_y() + 2, 200, pdf.get_y() + 2)
    pdf.ln(6)

    # ── Content ───────────────────────────────────────────────────────────────
    for key, val in content.items():
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(0, 71, 171)
        pdf.cell(0, 8, _pdf_safe(str(key)), ln=True)

        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(30, 30, 30)

        if isinstance(val, list):
            for item in val:
                pdf.multi_cell(0, 6, _pdf_safe(f"  * {item}"))

        elif isinstance(val, dict):
            for k2, v2 in val.items():
                pdf.set_font("Helvetica", "B", 9)
                pdf.cell(0, 6, _pdf_safe(f"  {k2}:"), ln=True)
                pdf.set_font("Helvetica", "", 9)
                if isinstance(v2, list):
                    for item in v2:
                        pdf.multi_cell(0, 5, _pdf_safe(f"    - {item}"))
                else:
                    pdf.multi_cell(0, 6, _pdf_safe(f"    {v2}"))
        else:
            pdf.multi_cell(0, 6, _pdf_safe(f"  {val}"))

        pdf.ln(2)

    return bytes(pdf.output())

def build_docx_bytes(selected_category: str, content: dict) -> bytes:
    if not DOCX_AVAILABLE:
        return b""
    doc = Document()
    # Title
    title = doc.add_heading("ZERO-TRUST FRAMEWORK APPLICATION", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 71, 171)
    sub = doc.add_paragraph("Developed by Randy Singh | KalSnet (KNet) Consulting Group")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True
    sub.runs[0].font.color.rgb = RGBColor(0, 71, 171)
    doc.add_paragraph(f"Category: {selected_category}   |   {datetime.now().strftime('%Y-%m-%d %H:%M')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    for key, val in content.items():
        h = doc.add_heading(key, level=2)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0, 71, 171)
        if isinstance(val, list):
            for item in val:
                doc.add_paragraph(item, style="List Bullet")
        elif isinstance(val, dict):
            for k2, v2 in val.items():
                p = doc.add_paragraph()
                p.add_run(f"{k2}: ").bold = True
                if isinstance(v2, list):
                    p.add_run(", ".join(str(x) for x in v2))
                else:
                    p.add_run(str(v2))
        else:
            doc.add_paragraph(str(val))
        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# ══════════════════════════════════════════════════════════════════════════════
#  SIDEBAR NAVIGATION
# ══════════════════════════════════════════════════════════════════════════════
with st.sidebar:
    st.markdown("## 🔒 ZTA Navigator")
    st.markdown("---")
    category = st.selectbox(
        "Select Framework Category",
        [
            "🏠 Overview & Architecture Diagram",
            "1️⃣  Functional Requirements",
            "2️⃣  Proposed Use Cases",
            "3️⃣  Architectural Requirements",
            "4️⃣  Evaluation Criteria",
            "📊 Synthetic Data & Analytics",
        ],
    )
    st.markdown("---")
    st.markdown("**About**")
    st.markdown("Randy Singh  \nComputer Scientist  \nDISA / BDE5  \nKalSnet (KNet) Consulting")
    st.markdown("`(301) 225-9535`")

# ══════════════════════════════════════════════════════════════════════════════
#  EXPORT SIDEBAR (always visible)
# ══════════════════════════════════════════════════════════════════════════════
with st.sidebar:
    st.markdown("---")
    st.markdown("### 📥 Export Results")
    export_category = category.replace("🏠 ", "").replace("1️⃣  ", "").replace("2️⃣  ", "").replace("3️⃣  ", "").replace("4️⃣  ", "").replace("📊 ", "")

    # Determine content to export based on current category
    def get_export_content():
        if "Functional" in category:
            return {k: {"Requirement": v["requirement"], "Example": v["example"],
                        "Solution": v["solution"], "Recommendation": v["recommendation"]}
                    for k, v in FUNCTIONAL_REQUIREMENTS.items()}
        elif "Use Case" in category:
            return USE_CASES
        elif "Architectural" in category:
            return ARCHITECTURAL_REQUIREMENTS
        elif "Evaluation" in category:
            return {k: {"Definition": v["definition"], "Metrics": v["metrics"],
                        "Target": v["target"], "Recommendation": v["recommendation"]}
                    for k, v in EVALUATION_CRITERIA.items()}
        else:
            return {"Framework": "Zero-Trust Architecture", "Version": "2.0", "Author": "Randy Singh / KNet"}

    exp_content = get_export_content()

    # Text
    txt_data = build_text_report(export_category, exp_content)
    st.download_button("📄 Export as Text", data=txt_data, file_name="zt_framework.txt", mime="text/plain", use_container_width=True)

    # JSON
    json_data = build_json_report(export_category, exp_content)
    st.download_button("🗃️ Export as JSON", data=json_data, file_name="zt_framework.json", mime="application/json", use_container_width=True)

    # PDF
    if FPDF_AVAILABLE:
        pdf_data = build_pdf_bytes(export_category, exp_content)
        st.download_button("📕 Export as PDF", data=pdf_data, file_name="zt_framework.pdf", mime="application/pdf", use_container_width=True)
    else:
        st.info("Install `fpdf2` for PDF export")

    # Word
    if DOCX_AVAILABLE:
        docx_data = build_docx_bytes(export_category, exp_content)
        st.download_button("📘 Export as Word", data=docx_data, file_name="zt_framework.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                           use_container_width=True)
    else:
        st.info("Install `python-docx` for Word export")


# ══════════════════════════════════════════════════════════════════════════════
#  MAIN CONTENT
# ══════════════════════════════════════════════════════════════════════════════

# ─── TAB 0: OVERVIEW ─────────────────────────────────────────────────────────
if "Overview" in category:
    st.markdown('<div class="section-header">🏛️ Zero-Trust Architecture — Conceptual Overview</div>', unsafe_allow_html=True)

    col1, col2, col3, col4 = st.columns(4)
    col1.metric("ZTA Pillars", "6", "Core Controls")
    col2.metric("Use Cases", "4", "Defined")
    col3.metric("Arch. Req. Areas", "7", "Categories")
    col4.metric("Eval. Criteria", "3", "Dimensions")

    st.markdown("---")

    # ── Mermaid-style ZTA flow rendered via Graphviz/Plotly Sankey alternative
    # Using Streamlit's native mermaid via st.graphviz_chart (ASCII fallback)
    st.markdown("#### 🔄 Zero-Trust Architecture Flow Diagram")

    try:
        import graphviz
        dot = graphviz.Digraph(comment="ZTA Flow", graph_attr={"rankdir": "LR", "bgcolor": "#f8f9ff", "fontname": "Helvetica"})
        dot.attr("node", shape="roundrectangle", style="filled", fontname="Helvetica", fontsize="11")

        # Nodes
        dot.node("User", "👤 User / Device", fillcolor="#dce8ff", color="#0047AB")
        dot.node("IdP", "🪪 Identity\nProvider (IdP)", fillcolor="#0047AB", fontcolor="white", color="#0047AB")
        dot.node("MFA", "🔐 MFA", fillcolor="#cce5ff", color="#0047AB")
        dot.node("NAC", "🩺 Health &\nCompliance (NAC)", fillcolor="#fff0cc", color="#f0a500")
        dot.node("PAM", "🔑 JIT / PAM\nAuthorization", fillcolor="#e8f5e9", color="#2e7d32")
        dot.node("PEP", "🛡️ Policy\nEnforcement\nPoint (PEP)", fillcolor="#0047AB", fontcolor="white", color="#0047AB")
        dot.node("App", "🖥️ Application\nTier", fillcolor="#dce8ff", color="#0047AB")
        dot.node("SIEM", "📋 SIEM /\nAccounting", fillcolor="#fce4ec", color="#c62828")
        dot.node("Orch", "⚙️ Orchestration\n(REST API)", fillcolor="#f3e5f5", color="#6a1b9a")

        # Edges
        dot.edge("User", "MFA", label="  Auth Request")
        dot.edge("MFA", "IdP", label="  Credentials")
        dot.edge("User", "NAC", label="  Device Posture")
        dot.edge("NAC", "IdP", label="  Health Signal")
        dot.edge("IdP", "PAM", label="  Identity Assert")
        dot.edge("PAM", "PEP", label="  Auth Decision")
        dot.edge("PEP", "App", label="  Access Granted")
        dot.edge("PEP", "SIEM", label="  Log Events")
        dot.edge("Orch", "PEP", label="  Policy Push")
        dot.edge("Orch", "IdP", label="  Role Provision")
        dot.edge("App", "SIEM", label="  App Logs")

        st.graphviz_chart(dot, use_container_width=True)
    except Exception:
        st.code("""
USER/DEVICE ──► MFA ──────────────► IdP ──────────► PAM/JIT
     │                                │                  │
     └──► NAC (Health Check) ─────────┘                  │
                                                          ▼
ORCHESTRATION (REST API) ──────────────────► POLICY ENFORCEMENT POINT
                                                          │
                                              ┌───────────┴────────────┐
                                              ▼                        ▼
                                        APPLICATION                   SIEM
                                          TIERS                  (ACCOUNTING)
        """, language="text")

    st.markdown("---")

    # ── Architecture Layers diagram (Plotly)
    st.markdown("#### 🏗️ ZTA Pillar Maturity Radar")
    radar_data = {
        "Pillar": ["Identity & Auth", "Health & Compliance", "Authorization", "Accounting", "Segmentation", "Orchestration"],
        "Current Maturity": [3.5, 2.8, 3.0, 3.8, 2.5, 2.2],
        "Target Maturity": [5, 5, 5, 5, 5, 5],
    }
    fig_radar = go.Figure()
    fig_radar.add_trace(go.Scatterpolar(
        r=radar_data["Current Maturity"] + [radar_data["Current Maturity"][0]],
        theta=radar_data["Pillar"] + [radar_data["Pillar"][0]],
        fill="toself", name="Current Maturity", line_color="#0047AB", fillcolor="rgba(0,71,171,0.2)"
    ))
    fig_radar.add_trace(go.Scatterpolar(
        r=radar_data["Target Maturity"] + [radar_data["Target Maturity"][0]],
        theta=radar_data["Pillar"] + [radar_data["Pillar"][0]],
        fill="toself", name="Target Maturity", line_color="#f0a500", fillcolor="rgba(240,165,0,0.1)"
    ))
    fig_radar.update_layout(
        polar=dict(radialaxis=dict(visible=True, range=[0, 5])),
        title="ZTA Pillar Maturity (1=Initial, 5=Optimised)",
        legend=dict(x=0.8, y=1.2),
        height=420,
    )
    st.plotly_chart(fig_radar, use_container_width=True)

    st.markdown("#### 📖 ZTA Principles Summary")
    principles = [
        ("Never Trust, Always Verify", "Treat every request as untrusted regardless of network location."),
        ("Least-Privilege Access", "Grant minimum permissions necessary and expire them promptly."),
        ("Assume Breach", "Design controls assuming attackers may already be inside the perimeter."),
        ("Continuous Verification", "Re-evaluate trust on every transaction, not just at login."),
        ("Micro-Segmentation", "Divide the network into small zones to contain lateral movement."),
        ("Policy-as-Code", "Automate and version-control all security policies via REST APIs."),
    ]
    for title_p, desc in principles:
        st.markdown(f'<div class="req-card"><b>{title_p}</b><br>{desc}</div>', unsafe_allow_html=True)


# ─── TAB 1: FUNCTIONAL REQUIREMENTS ─────────────────────────────────────────
elif "Functional" in category:
    st.markdown('<div class="section-header">⚙️ Category 1 — Functional Requirements</div>', unsafe_allow_html=True)

    show_all = st.checkbox("Show all pillars expanded", value=True)

    for pillar, data in FUNCTIONAL_REQUIREMENTS.items():
        with st.expander(f"🔹 {pillar}", expanded=show_all):
            st.markdown(f'<div class="req-card"><b>📌 Requirement:</b><br>{data["requirement"]}</div>', unsafe_allow_html=True)
            st.markdown(f'<div class="sol-card"><b>💡 Example:</b><br>{data["example"]}</div>', unsafe_allow_html=True)
            st.markdown(f'<div class="sol-card"><b>🛠️ Recommended Solution:</b><br>{data["solution"]}</div>', unsafe_allow_html=True)
            st.markdown(f'<div class="rec-card"><b>✅ Recommendation:</b><br>{data["recommendation"]}</div>', unsafe_allow_html=True)

    # Bar chart of pillar coverage
    st.markdown("---")
    st.markdown("#### 📊 Functional Requirement Pillar Coverage (Illustrative)")
    coverage = {"Identity & Authentication": 85, "Health & Compliance": 70, "Authorization": 78,
                "Accounting": 90, "Segmentation": 62, "Orchestration": 55}
    fig_bar = px.bar(
        x=list(coverage.keys()), y=list(coverage.values()),
        labels={"x": "Pillar", "y": "Coverage (%)"},
        color=list(coverage.values()), color_continuous_scale=["#ff4444", "#f0a500", "#0047AB"],
        title="Estimated Pillar Implementation Coverage (%)",
    )
    fig_bar.update_layout(coloraxis_showscale=False, height=360)
    st.plotly_chart(fig_bar, use_container_width=True)


# ─── TAB 2: USE CASES ────────────────────────────────────────────────────────
elif "Use Case" in category:
    st.markdown('<div class="section-header">📋 Category 2 — Proposed Use Cases</div>', unsafe_allow_html=True)

    uc_choice = st.radio("Select Use Case", list(USE_CASES.keys()), horizontal=False)

    st.markdown(f"### {uc_choice}")
    steps = USE_CASES[uc_choice]
    for i, step in enumerate(steps, 1):
        st.markdown(f'<div class="req-card"><b>Step {i}:</b> {step}</div>', unsafe_allow_html=True)

    # Flow diagram for the chosen use case
    st.markdown("---")
    st.markdown("#### 🔄 Use Case Flow")
    if "End-User" in uc_choice or "End User" in uc_choice:
        try:
            import graphviz
            g = graphviz.Digraph(graph_attr={"rankdir": "TD"})
            nodes = [("A", "User Login"), ("B", "MFA Challenge"), ("C", "IdP Identity Assert"),
                     ("D", "NAC Health Check"), ("E", "Contextual Policy Decision"), ("F", "Network Auth Granted"), ("G", "SIEM Logging")]
            for nid, label in nodes:
                g.node(nid, label, shape="box", style="filled", fillcolor="#dce8ff")
            for a, b in zip("ABCDEFG", "BCDEFG"):
                g.edge(a, b)
            st.graphviz_chart(g, use_container_width=True)
        except Exception:
            st.info("Install `graphviz` system package for flow diagrams.")

    elif "Segmentation" in uc_choice:
        try:
            import graphviz
            g = graphviz.Digraph(graph_attr={"rankdir": "LR"})
            tiers = [("Web", "#cce5ff"), ("App", "#c8e6c9"), ("DB", "#ffe0b2")]
            for name, color in tiers:
                g.node(name, name + " Tier", shape="cylinder", style="filled", fillcolor=color)
            g.node("SDN", "SDN / MCAP", shape="hexagon", style="filled", fillcolor="#0047AB", fontcolor="white")
            g.node("SIEM", "SIEM Logging", shape="note", style="filled", fillcolor="#fce4ec")
            for t, _ in tiers:
                g.edge("SDN", t)
            g.edge("SDN", "SIEM")
            st.graphviz_chart(g, use_container_width=True)
        except Exception:
            pass


# ─── TAB 3: ARCHITECTURAL REQUIREMENTS ──────────────────────────────────────
elif "Architectural" in category:
    st.markdown('<div class="section-header">🏗️ Category 3 — Architectural Requirements</div>', unsafe_allow_html=True)

    area_choice = st.selectbox("Select Requirement Area", list(ARCHITECTURAL_REQUIREMENTS.keys()))
    reqs = ARCHITECTURAL_REQUIREMENTS[area_choice]

    st.markdown(f"### {area_choice}")
    for req in reqs:
        indent = req.startswith(("1.", "2.", "3.")) and "." in req[2:4]
        prefix = "&nbsp;&nbsp;&nbsp;&nbsp;" if indent else ""
        st.markdown(f'<div class="req-card">{prefix}{req}</div>', unsafe_allow_html=True)

    # Count chart
    st.markdown("---")
    st.markdown("#### 📊 Requirement Count per Area")
    counts = {k: len(v) for k, v in ARCHITECTURAL_REQUIREMENTS.items()}
    fig_h = px.bar(
        x=list(counts.values()), y=list(counts.keys()),
        orientation="h", labels={"x": "# Requirements", "y": "Area"},
        color=list(counts.values()), color_continuous_scale=["#90caf9", "#0047AB"],
        title="Number of Architectural Requirements by Area",
    )
    fig_h.update_layout(coloraxis_showscale=False, height=380, yaxis={"categoryorder": "total ascending"})
    st.plotly_chart(fig_h, use_container_width=True)

    # Reference architecture note
    st.markdown("---")
    st.markdown("#### 🔗 Reference Standards")
    refs = [
        ("NIST SP 800-207", "Zero Trust Architecture — foundational ZTA standard."),
        ("CISA ZT Maturity Model", "5-pillar model for assessing ZTA implementation progress."),
        ("DoD ZT Strategy", "DoD mandate for ZT implementation across all components by FY27."),
        ("DISA STIG Library", "Security Technical Implementation Guides for OS/app hardening."),
        ("NIST SP 800-53 Rev 5", "Security and Privacy Controls — maps to ZTA control requirements."),
    ]
    for ref, desc in refs:
        st.markdown(f'<div class="rec-card"><b>{ref}</b> — {desc}</div>', unsafe_allow_html=True)


# ─── TAB 4: EVALUATION CRITERIA ──────────────────────────────────────────────
elif "Evaluation" in category:
    st.markdown('<div class="section-header">📐 Category 4 — Evaluation Criteria</div>', unsafe_allow_html=True)

    for criterion, data in EVALUATION_CRITERIA.items():
        with st.expander(f"📏 {criterion}", expanded=True):
            st.markdown(f'<div class="req-card"><b>Definition:</b> {data["definition"]}</div>', unsafe_allow_html=True)
            st.markdown(f'<div class="sol-card"><b>🎯 Target:</b> {data["target"]}</div>', unsafe_allow_html=True)
            st.markdown('<b>Key Metrics:</b>', unsafe_allow_html=True)
            for m in data["metrics"]:
                st.markdown(f'<div class="req-card">• {m}</div>', unsafe_allow_html=True)
            st.markdown(f'<div class="rec-card"><b>✅ Recommendation:</b> {data["recommendation"]}</div>', unsafe_allow_html=True)

    # Gauge charts
    st.markdown("---")
    st.markdown("#### 🎯 Sample Evaluation Scores (Illustrative)")
    gc1, gc2, gc3 = st.columns(3)
    for col, (crit, score, color) in zip(
        [gc1, gc2, gc3],
        [("Effectiveness", 78, "#0047AB"), ("Suitability", 85, "#2e7d32"), ("Performance", 71, "#f0a500")]
    ):
        fig_g = go.Figure(go.Indicator(
            mode="gauge+number",
            value=score,
            title={"text": crit},
            gauge={
                "axis": {"range": [0, 100]},
                "bar": {"color": color},
                "steps": [
                    {"range": [0, 50], "color": "#ffcdd2"},
                    {"range": [50, 75], "color": "#fff9c4"},
                    {"range": [75, 100], "color": "#c8e6c9"},
                ],
                "threshold": {"line": {"color": "black", "width": 3}, "thickness": 0.8, "value": 80},
            }
        ))
        fig_g.update_layout(height=260, margin=dict(t=40, b=10))
        col.plotly_chart(fig_g, use_container_width=True)


# ─── TAB 5: SYNTHETIC DATA & ANALYTICS ───────────────────────────────────────
elif "Synthetic" in category or "Analytics" in category or "Data" in category:
    st.markdown('<div class="section-header">📊 Synthetic ZTA Compliance Data & Analytics</div>', unsafe_allow_html=True)

    num_records = st.slider("Number of Synthetic Records", min_value=0, max_value=300, value=150, step=10)

    if num_records == 0:
        st.info("Move the slider above 0 to generate synthetic data.")
    else:
        df = generate_synthetic_data(num_records)

        # KPI row
        k1, k2, k3, k4, k5 = st.columns(5)
        k1.metric("Total Records", num_records)
        k2.metric("Compliant", int((df["Compliance_Status"] == "Compliant").sum()),
                  f"{(df['Compliance_Status']=='Compliant').mean()*100:.0f}%")
        k3.metric("Non-Compliant", int((df["Compliance_Status"] == "Non-Compliant").sum()))
        k4.metric("Avg STIG Score", f"{df['STIG_Score'].mean():.1f}")
        k5.metric("Incidents Flagged", int(df["Incident_Flagged"].sum()))

        st.markdown("---")

        tab_a, tab_b, tab_c, tab_d = st.tabs(["📋 Raw Data", "📊 By Pillar", "🗂️ By Department", "📈 Trends"])

        with tab_a:
            st.dataframe(df, use_container_width=True, height=380)
            csv = df.to_csv(index=False).encode()
            st.download_button("⬇️ Download CSV", csv, "zt_synthetic_data.csv", "text/csv")

        with tab_b:
            col_p1, col_p2 = st.columns(2)
            # Status by pillar
            pillar_status = df.groupby(["Pillar", "Compliance_Status"]).size().reset_index(name="Count")
            fig_ps = px.bar(pillar_status, x="Pillar", y="Count", color="Compliance_Status",
                            barmode="group",
                            color_discrete_map={"Compliant": "#2e7d32", "Non-Compliant": "#c62828",
                                                "Remediation": "#f0a500", "Pending Review": "#757575"},
                            title="Compliance Status by ZTA Pillar")
            col_p1.plotly_chart(fig_ps, use_container_width=True)
            # Avg STIG by pillar
            stig_pillar = df.groupby("Pillar")["STIG_Score"].mean().reset_index()
            fig_sp = px.bar(stig_pillar, x="Pillar", y="STIG_Score",
                            color="STIG_Score", color_continuous_scale=["#ff4444", "#f0a500", "#2e7d32"],
                            title="Average STIG Score by Pillar", range_y=[0, 100])
            fig_sp.update_layout(coloraxis_showscale=False)
            col_p2.plotly_chart(fig_sp, use_container_width=True)

        with tab_c:
            dept_status = df.groupby(["Department", "Compliance_Status"]).size().reset_index(name="Count")
            fig_ds = px.bar(dept_status, x="Department", y="Count", color="Compliance_Status",
                            barmode="stack",
                            color_discrete_map={"Compliant": "#2e7d32", "Non-Compliant": "#c62828",
                                                "Remediation": "#f0a500", "Pending Review": "#757575"},
                            title="Compliance Status by Department")
            st.plotly_chart(fig_ds, use_container_width=True)

            # Pie: device types
            dev_counts = df["Device_Type"].value_counts().reset_index()
            dev_counts.columns = ["Device_Type", "Count"]
            fig_pie = px.pie(dev_counts, values="Count", names="Device_Type",
                             title="Device Type Distribution", color_discrete_sequence=px.colors.qualitative.Bold)
            st.plotly_chart(fig_pie, use_container_width=True)

        with tab_d:
            df["Month"] = pd.to_datetime(df["Timestamp"]).dt.to_period("M").astype(str)
            trend = df.groupby(["Month", "Compliance_Status"]).size().reset_index(name="Count")
            fig_trend = px.line(trend, x="Month", y="Count", color="Compliance_Status",
                                markers=True,
                                color_discrete_map={"Compliant": "#2e7d32", "Non-Compliant": "#c62828",
                                                    "Remediation": "#f0a500", "Pending Review": "#757575"},
                                title="Compliance Trend Over Time")
            st.plotly_chart(fig_trend, use_container_width=True)

            # Scatter: auth latency vs stig score
            fig_sc = px.scatter(df, x="STIG_Score", y="Auth_Latency_ms", color="Compliance_Status",
                                size_max=8, opacity=0.7,
                                color_discrete_map={"Compliant": "#2e7d32", "Non-Compliant": "#c62828",
                                                    "Remediation": "#f0a500", "Pending Review": "#757575"},
                                title="STIG Score vs Authentication Latency",
                                labels={"STIG_Score": "STIG Compliance Score", "Auth_Latency_ms": "Auth Latency (ms)"})
            st.plotly_chart(fig_sc, use_container_width=True)


# ══════════════════════════════════════════════════════════════════════════════
#  FOOTER
# ══════════════════════════════════════════════════════════════════════════════
st.markdown("---")
st.markdown(
    '<div style="text-align:center;color:#888;font-size:0.8rem;">'
    '🔒 Zero-Trust Framework Application &nbsp;|&nbsp; '
    'Developed by <b>Randy Singh</b> &nbsp;|&nbsp; '
    'KalSnet (KNet) Consulting Group &nbsp;|&nbsp; '
    'DISA / BDE5 Technology Innovation Team &nbsp;|&nbsp; '
    f'{datetime.now().year}'
    '</div>',
    unsafe_allow_html=True
)
