

# DEFENSE-WAR-PLANNING-PLATFORM-ENHANCED.py

import streamlit as st
import pandas as pd
import numpy as np
import networkx as nx
import matplotlib.pyplot as plt
from reportlab.platypus import SimpleDocTemplate, Paragraph
from docx import Document
from datetime import datetime

# Optional YOLO
try:
    from ultralytics import YOLO
    YOLO_AVAILABLE = True
except:
    YOLO_AVAILABLE = False

st.set_page_config(layout="wide")

# ---------------- STYLE ----------------
st.markdown("""
<style>
.title-bar {
    color:#0033cc;
    font-size:42px;
    font-weight:bold;
    text-align:center;
}
.subtitle {
    color:#0033cc;
    font-size:20px;
    font-weight:bold;
    text-align:center;
}
.section-header {
    color:#0047AB;
    font-size:24px;
    font-weight:bold;
}
</style>
""", unsafe_allow_html=True)

# ---------------- TITLE ----------------
st.markdown("<div class='title-bar'>Kalsnet (KNet) AI WAR COMMAND PLATFORM</div>", unsafe_allow_html=True)
st.markdown("<div class='subtitle'>Developed by Randy Singh – KNet Consulting Group</div>", unsafe_allow_html=True)

# ---------------- ROLE ----------------
role = st.sidebar.selectbox("Select Role", ["Commander", "Analyst"])

st.markdown("### 🔹 Role Description")

if role == "Commander":
    st.info("""
**Commander Role:**  
Responsible for mission success, strategic decisions, and operational execution.  
Uses AI outputs (COA, Threat Scores, Logistics) to take action.

👉 Focus: **Decision Making + Execution**
""")
else:
    st.info("""
**Analyst Role:**  
Responsible for analyzing data, identifying patterns, and generating intelligence.  
Supports commanders with insights.

👉 Focus: **Data + Intelligence**
""")

# ---------------- MENU ----------------
menu = st.sidebar.radio("Select Module", [
    "Battlefield Map",
    "COA Decision Engine",
    "Threat Detection",
    "Logistics Optimization"
])

# ---------------- EXPORT FUNCTION ----------------
def export_data(df, name="output"):

    # CSV / JSON
    st.download_button("Download CSV", df.to_csv(), f"{name}.csv")
    st.download_button("Download JSON", df.to_json(), f"{name}.json")

    # PDF
    pdf_file = f"{name}.pdf"
    doc = SimpleDocTemplate(pdf_file)
    doc.build([Paragraph(df.head().to_string())])
    with open(pdf_file, "rb") as f:
        st.download_button("Download PDF", f, file_name=pdf_file)

    # WORD
    word_file = f"{name}.docx"
    document = Document()
    document.add_heading(name.upper(), 0)
    document.add_paragraph(df.to_string())
    document.save(word_file)

    with open(word_file, "rb") as f:
        st.download_button("Download WORD", f, file_name=word_file)

# ==========================================================
# ---------------- BATTLEFIELD MAP (FIXED) ----------------
# ==========================================================
if menu == "Battlefield Map":

    st.markdown("<div class='section-header'>Battlefield Map</div>", unsafe_allow_html=True)

    st.info("""
**WHAT THIS MODULE DOES:**

The Battlefield Map provides a **real-time geospatial visualization** of operational environments.

**KEY CAPABILITIES:**
- Displays command centers, threat zones, and operational regions
- Helps commanders understand battlefield positioning
- Allows analysts to monitor geographic threat distribution

**TOOLS USED:**
- Folium (Interactive Maps)
- Streamlit Integration

**REAL-WORLD USE:**
Used in military and intelligence systems for:
- Situational awareness
- Mission planning
- Threat tracking

👉 This acts as a **live operational dashboard map**
""")

    # SAFE MAP LOADING
    try:
        import folium
        from streamlit_folium import st_folium

        m = folium.Map(location=[20, 0], zoom_start=2)

        # Sample markers
        folium.Marker([28.6, 77.2], tooltip="Command Center").add_to(m)
        folium.Marker([34.5, 69.2], tooltip="Threat Zone").add_to(m)

        st_folium(m, width=1200, height=600)

    except:
        st.warning("Interactive map not available (folium not installed). Showing fallback view.")

        # Fallback visualization
        df_map = pd.DataFrame({
            "lat": [28.6, 34.5],
            "lon": [77.2, 69.2]
        })

        st.map(df_map)

# ==========================================================
# ---------------- COA ENGINE ----------------
# ==========================================================
if menu == "COA Decision Engine":

    st.markdown("<div class='section-header'>COA Decision Engine</div>", unsafe_allow_html=True)

    st.info("""
**DATA FIELDS EXPLAINED:**
- Terrain → Type of battlefield (affects movement & combat)
- Enemy → Enemy force size (threat level)
- Logistics → Available resources (fuel, weapons, support)
- Weather → Environmental condition impact

👉 These fields simulate **real battlefield variables**
""")

    n = st.slider("Synthetic Scenarios", 0, 200, 50)

    if st.button("Generate COA Data"):
        df = pd.DataFrame({
            "Terrain": np.random.choice(["Urban","Desert","Mountain"], n),
            "Enemy": np.random.randint(50,500,n),
            "Logistics": np.random.randint(50,200,n),
            "Weather": np.random.rand(n)
        })
        st.session_state.coa = df

    df = st.session_state.get("coa", pd.DataFrame())
    st.write(df)

    if not df.empty:

        df["Success"] = (df["Logistics"]/(df["Enemy"]+1))*df["Weather"]

        st.markdown("""
**FORMULA EXPLANATION:**
Success = (Logistics / Enemy) × Weather  

- More logistics → higher success  
- More enemy → lower success  
- Weather → multiplier impact  

👉 Used to simulate **mission success probability**
""")

        st.bar_chart(df["Success"])

        st.markdown("""
**GRAPH EXPLANATION:**
- Bar chart shows success probability per COA
- Each bar = one mission scenario
""")

        # Graph Theory
        G = nx.DiGraph()
        for i in range(len(df)):
            G.add_edge("Start", f"COA_{i}", weight=df["Success"][i])

        fig, ax = plt.subplots()
        nx.draw(G, with_labels=True)
        st.pyplot(fig)

        st.markdown("""
**GRAPH THEORY EXPLANATION:**
- Directed Graph = decision flow
- Each node = COA option
- Edges represent decision paths

👉 Used in military planning for **decision trees & probability flow**
""")

        export_data(df, "coa")

# ==========================================================
# ---------------- THREAT DETECTION ----------------
# ==========================================================
elif menu == "Threat Detection":

    st.markdown("<div class='section-header'>Threat Detection</div>", unsafe_allow_html=True)

    st.info("""
**DATA FIELDS EXPLAINED:**
- Speed → How fast threat is approaching
- Distance → How far threat is
- Signal → Detection strength / confidence

👉 Used to calculate threat urgency
""")

    n = st.slider("Synthetic Threat Data", 0, 200, 50)

    if st.button("Generate Threat Data"):
        df = pd.DataFrame({
            "Speed": np.random.randint(10,100,n),
            "Distance": np.random.randint(1,50,n),
            "Signal": np.random.rand(n)
        })
        st.session_state.threat = df

    df = st.session_state.get("threat", pd.DataFrame())
    st.write(df)

    if not df.empty:

        df["ThreatScore"] = df["Speed"]/(df["Distance"]+1)

        st.markdown("""
**FORMULA EXPLANATION:**
Threat Score = Speed / Distance  

- Faster + closer = higher threat  
👉 Used for **target prioritization**
""")

        st.bar_chart(df["ThreatScore"])

        # Graph
        G = nx.Graph()
        for i in range(len(df)):
            G.add_edge("Base", f"T{i}", weight=df["ThreatScore"][i])

        fig, ax = plt.subplots()
        nx.draw(G, with_labels=True)
        st.pyplot(fig)

        st.markdown("""
**GRAPH THEORY EXPLANATION:**
- Network graph shows relationships between base & threats
- Helps visualize multiple threats simultaneously

👉 Used in **radar tracking & defense systems**
""")

        export_data(df, "threat")

    # YOLO
    st.subheader("YOLO Object Detection (Optional)")

    st.info("""
**YOLO (You Only Look Once):**
- Real-time object detection AI model
- Processes entire image in one pass
- Detects objects like vehicles, weapons, drones

**HOW IT WORKS:**
1. Image split into grid
2. Each grid predicts objects
3. AI assigns confidence scores
4. Bounding boxes drawn

👉 Used in defense for:
- Drone detection
- Surveillance
- Target identification
""")

    file = st.file_uploader("Upload Image")

    if file and YOLO_AVAILABLE:
        model = YOLO("yolo_model.pt")
        result = model(file)
        st.image(result[0].plot())
    elif file:
        st.warning("YOLO model not available.")

# ==========================================================
# ---------------- LOGISTICS ----------------
# ==========================================================
elif menu == "Logistics Optimization":

    st.markdown("<div class='section-header'>Logistics Optimization</div>", unsafe_allow_html=True)

    st.info("""
**DATA FIELDS EXPLAINED:**
- From → Origin base
- To → Destination zone
- Cost → Transportation cost
- Time → Delivery time

👉 Used to optimize supply chain
""")

    n = st.slider("Synthetic Logistics Data", 0, 200, 50)

    if st.button("Generate Logistics Data"):
        df = pd.DataFrame({
            "From": np.random.choice(["BaseA","BaseB"], n),
            "To": np.random.choice(["Zone1","Zone2"], n),
            "Cost": np.random.randint(100,500,n),
            "Time": np.random.randint(1,20,n)
        })
        st.session_state.log = df

    df = st.session_state.get("log", pd.DataFrame())
    st.write(df)

    if not df.empty:

        df["Efficiency"] = df["Cost"]/df["Time"]

        st.markdown("""
**FORMULA EXPLANATION:**
Efficiency = Cost / Time  

- Lower cost + faster delivery = better efficiency  

👉 Used for **optimal military logistics planning**
""")

        st.bar_chart(df["Efficiency"])

        G = nx.Graph()
        for i in range(len(df)):
            G.add_edge(df["From"][i], df["To"][i], weight=df["Cost"][i])

        fig, ax = plt.subplots()
        nx.draw(G, with_labels=True)
        st.pyplot(fig)

        st.markdown("""
**GRAPH THEORY EXPLANATION:**
- Nodes = bases/zones
- Edges = routes
- Weight = cost

👉 Used for **shortest path optimization (like Dijkstra algorithm)**
""")

        export_data(df, "logistics")

# ---------------- FOOTER ----------------
st.markdown("---")
st.caption(f"System Time: {datetime.now()}")
