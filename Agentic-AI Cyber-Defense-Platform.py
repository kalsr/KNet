



# Agentic-AI-Cyber-Defense-Platform

import streamlit as st

import pandas as pd

import numpy as np

import json

import random

from datetime import datetime, timedelta



import plotly.express as px



# Export libs

from io import BytesIO

from docx import Document

from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

from reportlab.lib.styles import getSampleStyleSheet



# =========================

# PAGE CONFIG

# =========================

st.set_page_config(page_title="KNet Agentic AI Cyber Defense", layout="wide")



# =========================

# HEADER (BRANDING)

# =========================

st.markdown(

    """

    <h1 style='color:#0B5ED7; text-align:center;'>

    🔵 KNet Agentic-AI Cyber Defense Platform

    </h1>

    <h3 style='color:#0B5ED7; text-align:center;'>

    Developed by Randy Singh — Kalsnet (KNet) Consulting Group

    </h3>

    <hr>

    """,

    unsafe_allow_html=True

)



# =========================

# SIDEBAR - CONFIG

# =========================

st.sidebar.header("⚙️ Control Panel")



llm_choice = st.sidebar.selectbox("Select LLM Mode", ["Gemini", "Groq", "Compare Both"])



gemini_key = st.sidebar.text_input("Gemini API Key", type="password")

groq_key = st.sidebar.text_input("Groq API Key", type="password")



st.sidebar.markdown("### 🔑 Get Free API Keys")

st.sidebar.markdown("[Google AI Studio](https://aistudio.google.com)")

st.sidebar.markdown("[Groq Console](https://console.groq.com)")



use_case = st.sidebar.selectbox(

    "Select Cyber Defense Use Case",

    [

        "Threat Detection",

        "Incident Response",

        "Phishing Detection",

        "Malware Detection",

        "SOC Automation",

        "Vulnerability Management",

        "Zero Trust Identity",

        "Threat Intelligence",

        "Fraud Detection",

        "Red Team Simulation"

    ]

)



data_size = st.sidebar.slider("Synthetic Data Size", 0, 300, 100)



# =========================

# SESSION STATE

# =========================

if "data" not in st.session_state:

    st.session_state.data = pd.DataFrame()



# =========================

# SYNTHETIC DATA GENERATOR

# =========================

def generate_data(n):

    attack_types = [

        "phishing", "malware", "ransomware",

        "insider", "ddos", "credential_theft", "exfiltration"

    ]



    data = []

    base_time = datetime.now()



    for i in range(n):

        row = {

            "timestamp": base_time - timedelta(minutes=random.randint(0, 10000)),

            "user": f"user_{random.randint(1,50)}",

            "ip": f"192.168.1.{random.randint(1,255)}",

            "failed_logins": random.randint(0,10),

            "data_volume": random.randint(10,1000),

            "geo_risk": random.choice(["low","medium","high"]),

            "attack_type": random.choice(attack_types),

            "severity": random.randint(1,100)

        }

        data.append(row)



    return pd.DataFrame(data)



# =========================

# LOAD / UPLOAD DATA

# =========================

uploaded_file = st.sidebar.file_uploader("Upload Real Data (CSV/JSON/XLSX)")



def load_file(file):

    if file is None:

        return None



    if file.name.endswith(".csv"):

        return pd.read_csv(file)

    if file.name.endswith(".json"):

        return pd.read_json(file)

    if file.name.endswith(".xlsx"):

        return pd.read_excel(file)



    return None



# =========================

# RESET / GENERATE

# =========================

col1, col2 = st.sidebar.columns(2)



if col1.button("🔄 Generate"):

    st.session_state.data = generate_data(data_size)



if col2.button("🧹 Reset"):

    st.session_state.data = pd.DataFrame()



# Load uploaded file

if uploaded_file:

    st.session_state.data = load_file(uploaded_file)



df = st.session_state.data



# =========================

# LLM STUB (Gemini / Groq)

# =========================

def run_llm(prompt):

    if llm_choice == "Gemini":

        return f"[Gemini Analysis]\n{prompt}\n\n→ Simulated cybersecurity insights generated."

    elif llm_choice == "Groq":

        return f"[Groq Analysis]\n{prompt}\n\n→ Fast inference security insights generated."

    else:

        return f"[Comparison Mode]\nGemini + Groq combined analysis\n\n{prompt}"



# =========================

# ANALYSIS ENGINE

# =========================

def analyze_data(df, case):

    if df is None or df.empty:

        return "No data available for analysis."



    summary = {

        "total_events": len(df),

        "avg_severity": float(df["severity"].mean()),

        "high_risk_count": int((df["severity"] > 70).sum())

    }



    prompt = f"""

Cybersecurity Use Case: {case}



Dataset Summary:

{summary}



Provide:

- Threat insights

- Risk interpretation

- Recommended mitigation actions

"""



    return run_llm(prompt)



# =========================

# VISUALIZATIONS

# =========================

def show_charts(df):

    if df.empty:

        return



    col1, col2 = st.columns(2)



    with col1:

        fig = px.histogram(df, x="severity", title="Threat Severity Distribution")

        st.plotly_chart(fig, use_container_width=True)



    with col2:

        fig2 = px.pie(df, names="attack_type", title="Attack Type Breakdown")

        st.plotly_chart(fig2, use_container_width=True)



# =========================

# EXPORT FUNCTIONS

# =========================

def export_csv(df):

    return df.to_csv(index=False).encode("utf-8")



def export_json(df):

    return df.to_json(orient="records").encode("utf-8")



def export_docx(text):

    doc = Document()

    doc.add_heading("KNet Cyber Defense Report", 0)

    doc.add_paragraph(text)

    buffer = BytesIO()

    doc.save(buffer)

    buffer.seek(0)

    return buffer



def export_pdf(text):

    buffer = BytesIO()

    doc = SimpleDocTemplate(buffer)

    styles = getSampleStyleSheet()

    content = [Paragraph("KNet Cyber Defense Report", styles["Title"]),

               Spacer(1, 12),

               Paragraph(text, styles["BodyText"])]

    doc.build(content)

    buffer.seek(0)

    return buffer



# =========================

# MAIN DASHBOARD

# =========================

st.subheader("📊 Cyber Defense Dashboard")



if df.empty:

    st.warning("No data available. Generate or upload data.")

else:

    st.dataframe(df.head(20))

    show_charts(df)



# =========================

# USE CASE ENGINE

# =========================

st.subheader(f"🧠 AI Analysis: {use_case}")



analysis_result = analyze_data(df, use_case)

st.text_area("AI Output", analysis_result, height=250)



# =========================

# EXPORT PANEL

# =========================

st.subheader("📦 Export Results")



col1, col2, col3, col4 = st.columns(4)



with col1:

    st.download_button("CSV", export_csv(df), "data.csv")



with col2:

    st.download_button("JSON", export_json(df), "data.json")



with col3:

    st.download_button("DOCX", export_docx(analysis_result), "report.docx")



with col4:

    st.download_button("PDF", export_pdf(analysis_result), "report.pdf")



# =========================

# FOOTER

# =========================

st.markdown("---")

st.markdown(

    "<center><b>KNet Agentic-AI Cyber Defense Platform © 2026</b></center>",

    unsafe_allow_html=True

)
