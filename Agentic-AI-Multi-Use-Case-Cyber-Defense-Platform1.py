# python
# Agentic-AI-Cyber-Defense-Platform-V3
# REAL Gemini + Groq Integration

import streamlit as st
import pandas as pd
import numpy as np
import json
import random
from datetime import datetime, timedelta

import plotly.express as px

# REAL LLM IMPORTS
from groq import Groq
import google.generativeai as genai

# Export libs
from io import BytesIO
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

# =========================
# PAGE CONFIG
# =========================

st.set_page_config(
    page_title="KNet Agentic AI Cyber Defense",
    layout="wide"
)

# =========================
# HEADER (BRANDING)
# =========================

st.markdown(
    """
    <h1 style='color:#0B5ED7; text-align:center;'>
     KNet Agentic-AI Cyber Defense Platform
    </h1>

    <h3 style='color:#0B5ED7; text-align:center;'>
    Developed by Randy Singh — Kalsnet (KNet) Consulting Group
    </h3>

    <hr>
    """,
    unsafe_allow_html=True
)

# =========================
# SIDEBAR - CONFIG
# =========================

st.sidebar.header("⚙️ Control Panel")

llm_choice = st.sidebar.selectbox(
    "Select LLM Mode",
    ["Gemini", "Groq", "Compare Both"]
)

gemini_key = st.sidebar.text_input(
    "Gemini API Key",
    type="password"
)

groq_key = st.sidebar.text_input(
    "Groq API Key",
    type="password"
)

# ======================================
# API KEY INSTRUCTIONS
# ======================================

st.sidebar.markdown("## 🔑 Step-by-Step API Key Setup")

with st.sidebar.expander("How to Get FREE Gemini API Key"):

    st.markdown("""
    ### Gemini API Key Instructions

    1. Go to: https://aistudio.google.com
    2. Sign in using your Google account
    3. Click **Get API Key**
    4. Select **Create API Key**
    5. Choose existing Google project or create a new one
    6. Copy generated API key
    7. Paste the key into the Gemini API field above

    ✅ Gemini provides free-tier access for testing and development.
    """)

with st.sidebar.expander("How to Get FREE Groq API Key"):

    st.markdown("""
    ### Groq API Key Instructions

    1. Go to: https://console.groq.com
    2. Create free account or login
    3. Navigate to **API Keys**
    4. Click **Create API Key**
    5. Enter key name
    6. Copy generated API key
    7. Paste into Groq API field above

    ✅ Groq offers extremely fast LPU inference for AI workloads.
    """)

# =========================
# USE CASES
# =========================

use_case = st.sidebar.selectbox(
    "Select Cyber Defense Use Case",
    [
        "Threat Detection",
        "Incident Response",
        "Phishing Detection",
        "Malware Detection",
        "SOC Automation",
        "Vulnerability Management",
        "Zero Trust Identity",
        "Threat Intelligence",
        "Fraud Detection",
        "Red Team Simulation"
    ]
)

# ======================================
# FIXED DATA SIZE SLIDER
# ======================================

data_size = st.sidebar.slider(
    "Synthetic Data Size",
    min_value=0,
    max_value=300,
    value=100,
    step=10
)

# =========================
# SESSION STATE
# =========================

if "data" not in st.session_state:
    st.session_state.data = pd.DataFrame()

# =========================
# SYNTHETIC DATA GENERATOR
# =========================

def generate_data(n, selected_case):

    use_case_mapping = {
        "Threat Detection": ["ddos", "malware", "ransomware"],
        "Incident Response": ["insider", "credential_theft"],
        "Phishing Detection": ["phishing"],
        "Malware Detection": ["malware", "ransomware"],
        "SOC Automation": ["ddos", "credential_theft"],
        "Vulnerability Management": ["exfiltration", "insider"],
        "Zero Trust Identity": ["credential_theft"],
        "Threat Intelligence": ["phishing", "ddos", "malware"],
        "Fraud Detection": ["insider", "credential_theft"],
        "Red Team Simulation": ["ransomware", "ddos", "exfiltration"]
    }

    relevant_attacks = use_case_mapping.get(
        selected_case,
        ["malware"]
    )

    data = []

    base_time = datetime.now()

    for i in range(n):

        attack = random.choice(relevant_attacks)

        severity = random.randint(60, 100)

        if attack == "phishing":
            severity = random.randint(50, 85)

        if attack == "ransomware":
            severity = random.randint(80, 100)

        row = {
            "timestamp": base_time - timedelta(
                minutes=random.randint(0, 10000)
            ),
            "user": f"user_{random.randint(1,200)}",
            "ip": f"192.168.{random.randint(1,10)}.{random.randint(1,255)}",
            "failed_logins": random.randint(0,25),
            "data_volume": random.randint(50,5000),
            "geo_risk": random.choice(
                ["low","medium","high"]
            ),
            "attack_type": attack,
            "severity": severity,
            "endpoint": f"server-{random.randint(1,20)}",
            "department": random.choice([
                "Finance",
                "HR",
                "Operations",
                "SOC",
                "Engineering"
            ]),
            "status": random.choice([
                "Investigating",
                "Blocked",
                "Contained",
                "Escalated"
            ])
        }

        data.append(row)

    return pd.DataFrame(data)

# =========================
# LOAD / UPLOAD DATA
# =========================

uploaded_file = st.sidebar.file_uploader(
    "Upload Real Data (CSV/JSON/XLSX)"
)

def load_file(file):

    if file is None:
        return None

    if file.name.endswith(".csv"):
        return pd.read_csv(file)

    if file.name.endswith(".json"):
        return pd.read_json(file)

    if file.name.endswith(".xlsx"):
        return pd.read_excel(file)

    return None

# =========================
# RESET / GENERATE
# =========================

col1, col2 = st.sidebar.columns(2)

if col1.button("🔄 Generate"):

    st.session_state.data = generate_data(
        data_size,
        use_case
    )

if col2.button("🧹 Reset"):

    st.session_state.data = pd.DataFrame()

# Load uploaded file

if uploaded_file:

    st.session_state.data = load_file(uploaded_file)

df = st.session_state.data

# =========================
# REAL LLM ENGINE
# =========================

def run_groq(prompt):

    if not groq_key:
        return "ERROR: Please provide Groq API Key."

    try:

        client = Groq(api_key=groq_key)

        completion = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {
                    "role": "system",
                    "content": """
You are an elite cybersecurity analyst and SOC advisor.
Provide enterprise cybersecurity analysis,
threat intelligence,
risk interpretation,
and mitigation guidance.
"""
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            temperature=0.3,
            max_tokens=2000
        )

        return completion.choices[0].message.content

    except Exception as e:

        return f"Groq API Error: {str(e)}"

# =========================
# GEMINI ENGINE
# =========================

def run_gemini(prompt):

    if not gemini_key:
        return "ERROR: Please provide Gemini API Key."

    try:

        genai.configure(api_key=gemini_key)

        model = genai.GenerativeModel(
            "gemini-1.5-pro"
        )

        response = model.generate_content(prompt)

        return response.text

    except Exception as e:

        return f"Gemini API Error: {str(e)}"

# =========================
# MASTER LLM ROUTER
# =========================

def run_llm(prompt):

    if llm_choice == "Gemini":

        return run_gemini(prompt)

    elif llm_choice == "Groq":

        return run_groq(prompt)

    else:

        groq_result = run_groq(prompt)

        gemini_result = run_gemini(prompt)

        return f"""

========================
GROQ ANALYSIS
========================

{groq_result}

========================================================

========================
GEMINI ANALYSIS
========================

{gemini_result}

"""

# =========================
# ANALYSIS ENGINE
# =========================

def analyze_data(df, case):

    if df is None or df.empty:
        return "No data available for analysis."

    filtered_df = df.copy()

    top_attacks = filtered_df[
        "attack_type"
    ].value_counts().to_dict()

    avg_severity = round(
        filtered_df["severity"].mean(),
        2
    )

    high_risk = filtered_df[
        filtered_df["severity"] > 80
    ]

    top_geo = filtered_df[
        "geo_risk"
    ].value_counts().to_dict()

    failed_login_avg = round(
        filtered_df["failed_logins"].mean(),
        2
    )

    detailed_records = filtered_df.head(
        20
    ).to_string(index=False)

    prompt = f"""
Cybersecurity Use Case: {case}

==================================================
DATASET SUMMARY
==================================================

Total Security Events: {len(filtered_df)}

Average Severity Score: {avg_severity}

High Severity Incidents (>80): {len(high_risk)}

Top Attack Distribution:
{top_attacks}

Geo Risk Distribution:
{top_geo}

Average Failed Logins:
{failed_login_avg}

==================================================
ASSOCIATED SECURITY RECORDS
==================================================

{detailed_records}

==================================================
TASK
==================================================

Provide:

1. Detailed Threat Insights

2. Detailed Risk Interpretation

3. Detailed Recommended Mitigation Actions

4. Executive Summary

5. SOC Analyst Recommendations

6. Zero Trust Security Recommendations

7. Compliance Risks

8. Incident Response Priorities

Generate enterprise-grade cybersecurity analysis.
"""

    return run_llm(prompt)

# =========================
# VISUALIZATIONS
# =========================

def show_charts(df):

    if df.empty:
        return

    col1, col2 = st.columns(2)

    with col1:

        fig = px.histogram(
            df,
            x="severity",
            title="Threat Severity Distribution"
        )

        st.plotly_chart(
            fig,
            use_container_width=True
        )

    with col2:

        fig2 = px.pie(
            df,
            names="attack_type",
            title="Attack Type Breakdown"
        )

        st.plotly_chart(
            fig2,
            use_container_width=True
        )

# =========================
# EXPORT FUNCTIONS
# =========================

def export_csv(df):

    return df.to_csv(
        index=False
    ).encode("utf-8")

def export_json(df):

    return df.to_json(
        orient="records"
    ).encode("utf-8")

def export_docx(text):

    doc = Document()

    doc.add_heading(
        "KNet Cyber Defense Report",
        0
    )

    doc.add_paragraph(text)

    buffer = BytesIO()

    doc.save(buffer)

    buffer.seek(0)

    return buffer

def export_pdf(text):

    buffer = BytesIO()

    doc = SimpleDocTemplate(buffer)

    styles = getSampleStyleSheet()

    content = [
        Paragraph(
            "KNet Cyber Defense Report",
            styles["Title"]
        ),

        Spacer(1, 12),

        Paragraph(
            text.replace("\n", "<br/>"),
            styles["BodyText"]
        )
    ]

    doc.build(content)

    buffer.seek(0)

    return buffer

# =========================
# MAIN DASHBOARD
# =========================

st.subheader("📊 Cyber Defense Dashboard")

if df.empty:

    st.warning(
        "No data available. Generate or upload data."
    )

else:

    # DISPLAY RECORD COUNT

    st.success(
        f"Displaying {len(df)} Cybersecurity Records"
    )

    # FULL SCROLLABLE TABLE

    st.dataframe(
        df,
        use_container_width=True,
        height=500
    )

    # VISUALIZATIONS

    show_charts(df)

# =========================
# USE CASE ENGINE
# =========================

st.subheader(
    f"🧠 AI Analysis: {use_case}"
)

analysis_result = analyze_data(df, use_case)

st.text_area(
    "AI Output",
    analysis_result,
    height=700
)

# =========================
# EXPORT PANEL
# =========================

st.subheader("📦 Export Results")

col1, col2, col3, col4 = st.columns(4)

with col1:

    st.download_button(
        "CSV",
        export_csv(df),
        "data.csv"
    )

with col2:

    st.download_button(
        "JSON",
        export_json(df),
        "data.json"
    )

with col3:

    st.download_button(
        "DOCX",
        export_docx(analysis_result),
        "report.docx"
    )

with col4:

    st.download_button(
        "PDF",
        export_pdf(analysis_result),
        "report.pdf"
    )

# =========================
# FOOTER
# =========================

st.markdown("---")

st.markdown(
    """
<center>
<b>
KNet Agentic-AI Cyber Defense Platform © 2026
</b>
</center>
""",
    unsafe_allow_html=True
)

