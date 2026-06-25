

# ===============================================================================
# KNet Data Recovery — Taxonomy & Framework Explorer
#
# Developed by: Randy Singh
# Organization: Kalsnet (KNet) Consulting Group
#
# PURPOSE
# --------------------------------------------------------------------------
# • TAXONOMY        -> how recovery work is classified (domains, techniques,
#                      outcome states, confidence bands, tooling)
# • FRAMEWORK       -> the end-to-end process model (KDRF: KNet Data
#                      Recovery Framework) and the Agentic AI decision layer
# • DATA DICTIONARY -> every field used in the underlying data model, with
#                      type, meaning, and example values
# • FORMULAS        -> every calculation used (Recovery Probability,
#                      Confidence Level, Estimated Time, etc.), explained
#                      and made interactive so users can test their own numbers
# • GLOSSARY        -> plain-English definitions of every technical term used
# • EXPORT          -> the complete reference (Taxonomy + Framework + Data
#                      Dictionary + Formulas + Glossary) as Markdown, JSON,
#                      plain Text, Word (.docx), and PDF
# ===============================================================================
import html
import io
import json
from datetime import datetime
import pandas as pd
import streamlit as st
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable,
    )
    REPORTLAB_OK = True
except ImportError:
    REPORTLAB_OK = False
try:
    from docx import Document
    from docx.shared import RGBColor
    DOCX_OK = True
except ImportError:
    DOCX_OK = False
APP_TITLE = "Data Recovery — Taxonomy & Framework Explorer"
APP_AUTHOR = "Randy Singh"
APP_ORG = "Kalsnet (KNet) Consulting Group"
# ==========================================================================
# CORE REFERENCE DATA
# ==========================================================================
def taxonomy_data() -> dict:
    """
    The taxonomy is deliberately modeled as FOUR independent classification
    axes. A single recovery record is tagged with one value from each axis
    simultaneously — e.g. {Domain: Disk Recovery, Technique: Signature
    Carving, Outcome: Partial, Confidence Band: Medium}. Treating these as
    separate axes (rather than one flat list) is what makes the taxonomy
    composable and queryable.
    """
    return {
        "Axis 1 — Recovery Domain": {
            "description": (
                "WHAT kind of system or data the recovery effort targets. "
                "This is the top-level grouping used everywhere else in the app."
            ),
            "items": {
                "Disk Recovery": {
                    "summary": "Recovering data from physical or logical storage media (HDD, SSD, USB, removable media) after deletion, corruption, or formatting.",
                    "techniques": {
                        "Deleted File Recovery": "Restoring files removed from a filesystem but not yet overwritten on disk.",
                        "Partition Recovery": "Rebuilding or restoring a damaged or deleted partition table / volume structure.",
                        "Signature Carving": "Scanning unallocated disk space for known file-type byte signatures to reconstruct files with no filesystem metadata.",
                        "Hash Verification": "Comparing a recovered file's cryptographic hash (e.g. SHA-256) against a known-good baseline to confirm integrity.",
                    },
                },
                "Database Recovery": {
                    "summary": "Restoring structured data and transactional consistency in a DBMS after a crash, corruption, or accidental data loss.",
                    "techniques": {
                        "Transaction Log Replay": "Re-applying committed transactions from the write-ahead/transaction log to bring a database back to a consistent state.",
                        "Record Reconstruction": "Rebuilding individual rows/records from page-level remnants when the catalog or index is damaged.",
                        "Backup Comparison": "Diffing a damaged database against the most recent valid backup to identify the minimal recovery delta.",
                        "SQL Recovery": "Using vendor-specific repair utilities or manual SQL scripts to recover a corrupted database file.",
                    },
                },
                "Cloud Recovery": {
                    "summary": "Restoring data, objects, or infrastructure state in cloud-hosted environments (IaaS/PaaS/SaaS).",
                    "techniques": {
                        "Snapshot Restoration": "Restoring a volume, disk, or instance from a point-in-time snapshot.",
                        "Object Versioning": "Recovering a prior version of an object from a versioning-enabled storage bucket.",
                        "Bucket Recovery": "Restoring an entire deleted or corrupted object-storage bucket/container.",
                        "IAM Validation": "Confirming that identity and access permissions on recovered resources are correct and not over-permissive post-restore.",
                    },
                },
                "Email Recovery": {
                    "summary": "Recovering mailbox data and message stores from local or server-side mail systems.",
                    "techniques": {
                        "PST Recovery": "Repairing or extracting messages from a corrupted Outlook Personal Storage Table (.pst) file.",
                        "OST Recovery": "Recovering cached mailbox data from an Outlook Offline Storage Table (.ost) file, typically when no live server copy exists.",
                        "Exchange Export Recovery": "Reconstructing mailbox content from Exchange/Microsoft 365 export or eDiscovery archives.",
                    },
                },
                "Cyber Artifact Recovery": {
                    "summary": "Recovering forensic and security telemetry artifacts needed to investigate or reconstruct an incident timeline.",
                    "techniques": {
                        "Zeek Log Analysis": "Recovering and parsing Zeek (Bro) network-traffic logs for connection and protocol metadata.",
                        "PCAP Reconstruction": "Rebuilding readable packet captures from fragmented or partially corrupted .pcap files.",
                        "Firewall Log Recovery": "Restoring firewall/network-device logs to close gaps in an audit or incident timeline.",
                        "API Log Recovery": "Recovering application/API gateway access logs lost to rotation, corruption, or deletion.",
                        "SIEM Event Recovery": "Restoring security events from a SIEM's storage tier (hot/warm/cold) for retroactive correlation.",
                    },
                },
            },
        },
        "Axis 2 — Recovery Outcome (Status)": {
            "description": (
                "The RESULT classification applied to every individual recovery record once "
                "an attempt has been made. This is the primary input to every formula in this app."
            ),
            "items": {
                "Recovered": {"summary": "Data was fully restored and its integrity could be verified (e.g. hash match, consistency check passed).", "color": "#1A7F37"},
                "Partial": {"summary": "Data was restored incompletely, or restored but with reduced confidence (e.g. partial overwrite, unverifiable integrity).", "color": "#B58105"},
                "Failed": {"summary": "Data could not be recovered with the methods and tools available for the attempt.", "color": "#B42318"},
            },
        },
        "Axis 3 — Confidence Band": {
            "description": (
                "A CERTAINTY classification derived from a record's numeric ConfidenceScore "
                "(0.0–1.0). This turns a continuous score into a human-readable category."
            ),
            "items": {
                "High": {"summary": "ConfidenceScore ≥ 0.80 — strong evidence the outcome label is correct.", "color": "#1A7F37"},
                "Medium": {"summary": "0.45 ≤ ConfidenceScore < 0.80 — moderate evidence; spot-checking recommended.", "color": "#B58105"},
                "Low": {"summary": "ConfidenceScore < 0.45 — weak evidence; manual review required before relying on the outcome.", "color": "#B42318"},
            },
        },
        "Axis 4 — Tooling": {
            "description": (
                "The HOW — the specific instrument or utility used to attempt recovery. "
                "Tooling is always scoped to a Domain (Axis 1); see the Data Dictionary for the field that stores this."
            ),
            "items": {
                "Disk Recovery tools": {"summary": "FTK Imager, Autopsy, X-Ways Forensics, PhotoRec"},
                "Database Recovery tools": {"summary": "SQL Log Analyzer, RecoverPoint, pg_waldump, ApexSQL Log"},
                "Cloud Recovery tools": {"summary": "AWS Backup, Azure Site Recovery, GCP Snapshot Manager, Velero"},
                "Email Recovery tools": {"summary": "Kernel PST Repair, Stellar OST to PST, eDiscovery Export Tool"},
                "Cyber Artifact Recovery tools": {"summary": "Zeek, Wireshark, Splunk, Elastic SIEM, Suricata"},
            },
        },
    }
def framework_data() -> dict:
    """
    KDRF — the KNet Data Recovery Framework: a five-phase process model that
    any recovery engagement (manual or automated) moves through, plus a
    cross-cutting 'Decision Layer' that the Agentic AI advisor occupies.
    """
    phases = [
        {
            "phase": "1. Identify & Triage",
            "goal": "Detect that a data-loss event occurred and classify it against the taxonomy.",
            "activities": [
                "Confirm a loss/corruption event and its scope",
                "Assign the affected Recovery Domain(s) (Axis 1)",
                "Set initial priority based on business/data criticality",
            ],
            "output": "A scoped, classified recovery case ready for acquisition",
        },
        {
            "phase": "2. Acquire & Preserve",
            "goal": "Capture the affected data/media without altering the original evidence.",
            "activities": [
                "Create a forensic image, snapshot, or export of the source",
                "Establish chain-of-custody / acquisition timestamp",
                "Hash the acquired copy for later integrity verification",
            ],
            "output": "A preserved, verifiable working copy",
        },
        {
            "phase": "3. Analyze",
            "goal": "Run domain-specific analyzers against the acquired copy to characterize what is recoverable.",
            "activities": [
                "Run the Disk / DB / Cloud / Email / Log Analyzer relevant to the domain",
                "Identify candidate techniques (Axis 1 sub-techniques) to apply",
                "Produce a preliminary per-item ConfidenceScore",
            ],
            "output": "An analyzed inventory of recoverable items with preliminary confidence",
        },
        {
            "phase": "4. Recover & Reconstruct",
            "goal": "Apply the selected technique(s) and produce a definitive outcome per item.",
            "activities": [
                "Execute the technique (carving, log replay, snapshot restore, etc.)",
                "Re-verify integrity (hash / consistency check) on the result",
                "Assign the final Outcome status (Axis 2): Recovered / Partial / Failed",
            ],
            "output": "A completed set of recovery records, one per item",
        },
        {
            "phase": "5. Validate, Score & Report",
            "goal": "Aggregate individual outcomes into run-level metrics and a recommended next-step plan.",
            "activities": [
                "Compute Recovery Probability, Confidence Level, Estimated Time (see Formulas tab)",
                "Generate a Recovery Plan and Suggested Tools via the Decision Layer",
                "Export / report results to stakeholders",
            ],
            "output": "A scored run with an actionable recovery plan and exportable report",
        },
    ]
    decision_layer = {
        "name": "Decision Layer (Agentic AI Recovery Advisor)",
        "runs_during": "Phases 3–5",
        "pipeline": [
            "Input Sources (acquired data / logs / images)",
            "Domain Analyzers (Disk / DB / Cloud / Email / Log)",
            "Recovery Advisor (rule-based engine, or an LLM/agent in production)",
            "Outputs: Confidence Score · Recovery Plan · Estimated Time · Suggested Tools",
        ],
        "note": (
            "In the reference implementation this layer is a deterministic, rule-based "
            "stand-in (see the Formulas tab). In a production deployment it can be "
            "replaced with a live LLM or agent call without changing any other phase."
        ),
    }
    return {"phases": phases, "decision_layer": decision_layer}
def data_dictionary_df() -> pd.DataFrame:
    rows = [
        ("ID", "Integer", "Sequential identifier for a single synthetic recovery record within a run.", "1, 2, 3 …"),
        ("Domain", "Categorical (Axis 1)", "Which recovery domain the record belongs to.", "Disk Recovery"),
        ("SubType", "Categorical (Axis 1 technique)", "The specific technique applied within the domain.", "Signature Carving"),
        ("Status", "Categorical (Axis 2)", "The outcome of the recovery attempt.", "Partial"),
        ("Size_MB", "Float", "Simulated size of the recovered/attempted item, in megabytes.", "1284.7"),
        ("ConfidenceScore", "Float [0.0–1.0]", "Numeric certainty in the assigned Status; bucketed into a Confidence Band (Axis 3).", "0.61"),
        ("ToolUsed", "Categorical (Axis 4)", "The tool/instrument used for the attempt.", "PhotoRec"),
        ("Timestamp", "ISO-8601 datetime", "When the attempt was simulated/recorded.", "2026-06-24T09:12:05"),
        ("run_id", "Integer", "Identifier for a full generation run, grouping many records together.", "17"),
        ("domains", "String (comma-joined)", "The set of domains selected for a given run.", "Disk Recovery, Cloud Recovery"),
        ("num_records", "Integer", "Total records generated in a run (sum of recovered + partial + failed).", "100"),
        ("recovered / partial / failed", "Integer (×3 fields)", "Run-level counts of each Outcome status — the direct inputs to the formulas below.", "61 / 24 / 15"),
        ("probability", "Float (%)", "Run-level Recovery Probability — see Formulas tab.", "73.5"),
        ("confidence", "Categorical", "Run-level Confidence Level derived from probability — see Formulas tab.", "Medium"),
        ("estimated_time", "String (Hh Mm)", "Run-level Estimated Recovery Time — see Formulas tab.", "2h 14m"),
    ]
    return pd.DataFrame(rows, columns=["Field", "Type", "Description", "Example"])
def formulas_data() -> list:
    """Each formula: id, name, latex, plain-English explanation, variable legend."""
    return [
        {
            "id": "F1",
            "name": "Recovery Probability",
            "latex": r"P_r = \frac{(R \times 1.0) + (P \times 0.5) + (F \times 0.0)}{N} \times 100\%",
            "explanation": (
                "The headline metric for a run. Every Recovered record earns full credit, every "
                "Partial record earns half credit, and Failed records earn no credit. The result "
                "is the percentage of 'recovery value' achieved out of what was theoretically possible."
            ),
            "variables": {"R": "count of Recovered records", "P": "count of Partial records",
                           "F": "count of Failed records", "N": "total records = R + P + F"},
        },
        {
            "id": "F2",
            "name": "Confidence Level",
            "latex": r"\text{Confidence} = \begin{cases} \text{High} & P_r \ge 85 \\ \text{Medium} & 60 \le P_r < 85 \\ \text{Low} & P_r < 60 \end{cases}",
            "explanation": (
                "Converts the numeric Recovery Probability into a three-tier label so stakeholders "
                "can act on it at a glance, without needing to interpret a raw percentage."
            ),
            "variables": {"P_r": "Recovery Probability from F1, as a percentage"},
        },
        {
            "id": "F3",
            "name": "Estimated Recovery Time",
            "latex": r"T = (N \times 0.8) + (F \times 3) + (P \times 1.5)",
            "explanation": (
                "Estimates total analyst/engineer minutes for the run. Every record costs a baseline "
                "0.8 minutes to process; Failed records add 3 extra minutes (deeper troubleshooting) "
                "and Partial records add 1.5 extra minutes (verification/cleanup)."
            ),
            "variables": {"N": "total records", "F": "count of Failed records", "P": "count of Partial records"},
        },
        {
            "id": "F4",
            "name": "Record-level Confidence Score Band",
            "latex": r"cs \in \begin{cases} [0.80,\,0.99] & \text{Status} = \text{Recovered} \\ [0.45,\,0.79] & \text{Status} = \text{Partial} \\ [0.05,\,0.44] & \text{Status} = \text{Failed} \end{cases}",
            "explanation": (
                "Defines the numeric range a record's ConfidenceScore is drawn from, given its "
                "Status. This is what links Axis 2 (Outcome) to Axis 3 (Confidence Band) at the "
                "individual-record level."
            ),
            "variables": {"cs": "an individual record's ConfidenceScore"},
        },
        {
            "id": "F5",
            "name": "Domain Weighted Success Rate",
            "latex": r"W_d = w_R + 0.5 \times w_P \quad \text{where } w_R + w_P + w_F = 1",
            "explanation": (
                "An a-priori KPI for a domain, computed from its expected status-weight vector "
                "(used to bias synthetic generation). It answers 'before running anything, what "
                "success rate do we expect from this domain?' — distinct from the empirical, "
                "after-the-fact Recovery Probability (F1)."
            ),
            "variables": {"w_R": "domain's expected Recovered weight", "w_P": "domain's expected Partial weight",
                           "w_F": "domain's expected Failed weight"},
        },
        {
            "id": "F6",
            "name": "Aggregate Multi-Domain Probability",
            "latex": r"P_r^{\,total} = \frac{\sum_d \left(N_d \times P_{r,d}\right)}{\sum_d N_d}",
            "explanation": (
                "When a run spans multiple domains, this combines each domain's individual "
                "Recovery Probability into one overall figure, weighted by how many records "
                "each domain contributed — so a domain with more records has proportionally "
                "more influence on the total."
            ),
            "variables": {"d": "a domain in the run", "N_d": "record count for domain d",
                           "P_{r,d}": "Recovery Probability (F1) computed for domain d alone"},
        },
    ]
def glossary_data() -> dict:
    return {
        "Bucket": "A container for storing objects in cloud object storage (e.g. Amazon S3, Google Cloud Storage).",
        "Chain of Custody": "A documented, unbroken record of who handled evidence/data and when, used to preserve its integrity for legal or audit purposes.",
        "Confidence Band": "A categorical (High/Medium/Low) grouping of a numeric ConfidenceScore; see Axis 3 of the taxonomy.",
        "ConfidenceScore": "A 0.0–1.0 numeric estimate of how certain the system is that a record's assigned Status is correct.",
        "Forensic Image": "A bit-for-bit copy of a storage device or volume, created so analysis never touches the original media.",
        "Hash Verification": "Confirming data integrity by comparing a cryptographic hash (e.g. SHA-256) of recovered data against a known-good value.",
        "IAM (Identity and Access Management)": "The set of policies and permissions controlling who/what can access a cloud resource.",
        "Object Versioning": "A cloud storage feature that retains prior versions of an object so an earlier version can be restored.",
        "OST file": "Outlook Offline Storage Table — a local cache of a mailbox used for offline access.",
        "PCAP": "Packet capture — a file format storing captured network traffic for later analysis.",
        "PST file": "Outlook Personal Storage Table — a file storing a local copy of mailbox folders and messages.",
        "Recovery Probability": "The headline run-level metric estimating overall recovery success; see Formula F1.",
        "Signature Carving": "Recovering files from raw, unstructured storage by scanning for known file-type byte signatures, without relying on filesystem metadata.",
        "SIEM": "Security Information and Event Management — a platform that aggregates and correlates security event logs.",
        "Snapshot": "A point-in-time copy of a disk, volume, or database used as a restore point.",
        "Transaction Log": "A sequential record of all changes made to a database, used to replay or roll back transactions.",
        "Zeek": "An open-source network traffic analysis framework (formerly 'Bro') that produces structured logs of network activity.",
    }
# ==========================================================================
# FORMULA HELPER FUNCTIONS (used by the interactive calculator)
# ==========================================================================
def calc_recovery_probability(r: int, p: int, f: int):
    n = r + p + f
    if n == 0:
        return 0.0, 0
    return round(((r * 1.0) + (p * 0.5)) / n * 100, 2), n
def calc_confidence(prob: float) -> str:
    if prob >= 85:
        return "High"
    elif prob >= 60:
        return "Medium"
    return "Low"
def calc_estimated_time(n: int, p: int, f: int):
    minutes = (n * 0.8) + (f * 3) + (p * 1.5)
    hours = int(minutes // 60)
    mins = int(minutes % 60)
    return minutes, f"{hours}h {mins}m" if hours else f"{mins}m"
def calc_domain_weighted_success(w_r: float, w_p: float) -> float:
    return round(w_r + 0.5 * w_p, 3)
# ==========================================================================
# UI HELPERS
# ==========================================================================
def render_header():
    st.markdown(
        """
        <div style="background:linear-gradient(135deg,#0B3D91,#3D63B8);padding:18px 26px;
                    border-radius:10px;margin-bottom:18px;">
            <h1 style="color:white;margin:0;font-size:28px;">📚 Data Recovery — Taxonomy &amp; Framework Explorer</h1>
            <p style="color:#D7E3FF;margin:4px 0 0 0;font-size:14px;">
                Developed by <b>Randy Singh</b> &nbsp;|&nbsp; Kalsnet (KNet) Consulting Group
                &nbsp;•&nbsp; companion reference to the Enterprise Recovery Orchestrator
            </p>
        </div>
        """,
        unsafe_allow_html=True,
    )
def colored_chip(text: str, color: str):
    st.markdown(
        f"""<span style="background-color:{color}1A;color:{color};border:1px solid {color}55;
        padding:3px 11px;border-radius:12px;font-size:12.5px;font-weight:700;">{text}</span>""",
        unsafe_allow_html=True,
    )
def phase_box(title: str, idx: int, total: int):
    """Render one framework phase box with a connecting arrow (except the last)."""
    st.markdown(
        f"""
        <div style="background:#FFFFFF;border:1px solid #DDE3ED;border-left:5px solid #0B3D91;
                    border-radius:8px;padding:10px 14px;text-align:center;font-weight:700;
                    font-size:13px;color:#0B3D91;min-height:54px;display:flex;align-items:center;
                    justify-content:center;">{title}</div>
        """,
        unsafe_allow_html=True,
    )
# ==========================================================================
# EXPORT / REFERENCE DOCUMENT BUILDERS
# ==========================================================================
def build_markdown_reference() -> str:
    tax = taxonomy_data()
    fw = framework_data()
    formulas = formulas_data()
    glossary = glossary_data()
    md = [f"# {APP_TITLE}", f"_{APP_ORG} — generated {datetime.now().strftime('%Y-%m-%d %H:%M')}_", ""]
    md.append("## Taxonomy\n")
    for axis_name, axis in tax.items():
        md.append(f"### {axis_name}")
        md.append(axis["description"] + "\n")
        for item_name, item in axis["items"].items():
            summary = item.get("summary", "")
            md.append(f"- **{item_name}** — {summary}")
            if "techniques" in item:
                for t_name, t_desc in item["techniques"].items():
                    md.append(f"    - *{t_name}*: {t_desc}")
        md.append("")
    md.append("## Framework (KDRF)\n")
    for ph in fw["phases"]:
        md.append(f"### {ph['phase']}")
        md.append(f"**Goal:** {ph['goal']}")
        for a in ph["activities"]:
            md.append(f"- {a}")
        md.append(f"**Output:** {ph['output']}\n")
    dl = fw["decision_layer"]
    md.append(f"### {dl['name']} (runs during {dl['runs_during']})")
    for step in dl["pipeline"]:
        md.append(f"- {step}")
    md.append(dl["note"] + "\n")
    md.append("## Data Dictionary\n")
    dd = data_dictionary_df()
    try:
        md.append(dd.to_markdown(index=False))
    except ImportError:
        # Fallback if the optional 'tabulate' package isn't installed
        header = "| " + " | ".join(dd.columns) + " |"
        sep = "| " + " | ".join(["---"] * len(dd.columns)) + " |"
        body = "\n".join("| " + " | ".join(str(v) for v in row) + " |" for row in dd.values)
        md.append("\n".join([header, sep, body]))
    md.append("")
    md.append("## Formulas\n")
    for f in formulas:
        md.append(f"### {f['id']} — {f['name']}")
        md.append(f"`{f['latex']}`")
        md.append(f["explanation"])
        for v, d in f["variables"].items():
            md.append(f"- `{v}`: {d}")
        md.append("")
    md.append("## Glossary\n")
    for term, definition in sorted(glossary.items()):
        md.append(f"- **{term}**: {definition}")
    return "\n".join(md)
def build_json_reference() -> dict:
    return {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "title": APP_TITLE,
        "organization": APP_ORG,
        "taxonomy": taxonomy_data(),
        "framework": framework_data(),
        "data_dictionary": data_dictionary_df().to_dict(orient="records"),
        "formulas": formulas_data(),
        "glossary": glossary_data(),
    }
def build_text_reference() -> str:
    """Plain-text (.txt) rendering of the complete reference — taxonomy,
    framework, data dictionary, formulas, and glossary."""
    tax = taxonomy_data()
    fw = framework_data()
    dd = data_dictionary_df()
    formulas = formulas_data()
    glossary = glossary_data()
    bar = "=" * 78
    rule = "-" * 78
    lines = [
        bar,
        f"  {APP_TITLE}",
        f"  Developed by {APP_AUTHOR} | {APP_ORG}",
        f"  Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        bar,
        "",
        "TAXONOMY",
        rule,
    ]
    for axis_name, axis in tax.items():
        lines.append("")
        lines.append(axis_name)
        lines.append(axis["description"])
        for item_name, item in axis["items"].items():
            lines.append(f"  • {item_name}: {item.get('summary', '')}")
            if "techniques" in item:
                for t_name, t_desc in item["techniques"].items():
                    lines.append(f"      - {t_name}: {t_desc}")
    lines += ["", "FRAMEWORK (KDRF)", rule]
    for ph in fw["phases"]:
        lines.append("")
        lines.append(ph["phase"])
        lines.append(f"Goal: {ph['goal']}")
        for a in ph["activities"]:
            lines.append(f"  - {a}")
        lines.append(f"Output: {ph['output']}")
    dl = fw["decision_layer"]
    lines.append("")
    lines.append(f"{dl['name']} (runs during {dl['runs_during']})")
    for step in dl["pipeline"]:
        lines.append(f"  -> {step}")
    lines.append(dl["note"])
    lines += ["", "DATA DICTIONARY", rule, dd.to_string(index=False)]
    lines += ["", "FORMULAS", rule]
    for f in formulas:
        lines.append("")
        lines.append(f"{f['id']} — {f['name']}")
        lines.append(f["latex"])
        lines.append(f["explanation"])
        for v, d in f["variables"].items():
            lines.append(f"  {v}: {d}")
    lines += ["", "GLOSSARY", rule]
    for term, definition in sorted(glossary.items()):
        lines.append(f"{term}: {definition}")
    return "\n".join(lines)
def build_docx_reference():
    """Word (.docx) rendering of the complete reference. Returns a BytesIO
    buffer ready for st.download_button, or None if python-docx is missing."""
    if not DOCX_OK:
        return None
    doc = Document()
    title = doc.add_heading(APP_TITLE, 0)
    if title.runs:
        title.runs[0].font.color.rgb = RGBColor(0x0B, 0x3D, 0x91)
    doc.add_paragraph(f"Developed by {APP_AUTHOR} | {APP_ORG}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_heading("Taxonomy", level=1)
    for axis_name, axis in taxonomy_data().items():
        doc.add_heading(axis_name, level=2)
        doc.add_paragraph(axis["description"])
        for item_name, item in axis["items"].items():
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(item_name + ": ")
            r.bold = True
            p.add_run(item.get("summary", ""))
            if "techniques" in item:
                for t_name, t_desc in item["techniques"].items():
                    p2 = doc.add_paragraph(style="List Bullet 2")
                    r2 = p2.add_run(t_name + ": ")
                    r2.italic = True
                    p2.add_run(t_desc)
    doc.add_heading("Framework (KDRF)", level=1)
    fw = framework_data()
    for ph in fw["phases"]:
        doc.add_heading(ph["phase"], level=2)
        gp = doc.add_paragraph()
        gp.add_run("Goal: ").bold = True
        gp.add_run(ph["goal"])
        for a in ph["activities"]:
            doc.add_paragraph(a, style="List Bullet")
        op = doc.add_paragraph()
        op.add_run("Output: ").bold = True
        op.add_run(ph["output"])
    dl = fw["decision_layer"]
    doc.add_heading(dl["name"], level=2)
    doc.add_paragraph(f"Runs during: {dl['runs_during']}")
    for step in dl["pipeline"]:
        doc.add_paragraph(step, style="List Bullet")
    doc.add_paragraph(dl["note"])
    doc.add_heading("Data Dictionary", level=1)
    dd = data_dictionary_df()
    tbl = doc.add_table(rows=1, cols=len(dd.columns))
    tbl.style = "Light Shading Accent 1"
    for i, c in enumerate(dd.columns):
        tbl.rows[0].cells[i].text = c
    for _, row in dd.iterrows():
        cells = tbl.add_row().cells
        for i, c in enumerate(dd.columns):
            cells[i].text = str(row[c])
    doc.add_heading("Formulas", level=1)
    for f in formulas_data():
        doc.add_heading(f"{f['id']} — {f['name']}", level=2)
        mono_p = doc.add_paragraph()
        mono_run = mono_p.add_run(f["latex"])
        mono_run.font.name = "Courier New"
        doc.add_paragraph(f["explanation"])
        for v, d in f["variables"].items():
            vp = doc.add_paragraph(style="List Bullet")
            vr = vp.add_run(v + ": ")
            vr.bold = True
            vp.add_run(d)
    doc.add_heading("Glossary", level=1)
    for term, definition in sorted(glossary_data().items()):
        gp2 = doc.add_paragraph()
        gr2 = gp2.add_run(term + ": ")
        gr2.bold = True
        gp2.add_run(definition)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
def build_pdf_reference():
    """PDF rendering of the complete reference. Returns a BytesIO buffer
    ready for st.download_button, or None if reportlab is missing."""
    if not REPORTLAB_OK:
        return None
    def esc(s) -> str:
        return html.escape(str(s), quote=False)
    buf = io.BytesIO()
    pdoc = SimpleDocTemplate(
        buf, pagesize=letter,
        rightMargin=42, leftMargin=42, topMargin=42, bottomMargin=42,
    )
    styles = getSampleStyleSheet()
    navy = colors.HexColor("#0B3D91")
    title_style = ParagraphStyle("T", parent=styles["Title"], textColor=navy, fontSize=18, spaceAfter=4)
    h1 = ParagraphStyle("H1", parent=styles["Heading1"], textColor=navy, fontSize=14, spaceBefore=14, spaceAfter=6)
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], textColor=navy, fontSize=11.5, spaceBefore=10, spaceAfter=4)
    body = styles["Normal"]
    mono = ParagraphStyle("Mono", parent=styles["Normal"], fontName="Courier", fontSize=8.5,
                          textColor=colors.HexColor("#333333"), spaceAfter=4)
    story = [
        Paragraph(esc(APP_TITLE), title_style),
        Paragraph(esc(f"Developed by {APP_AUTHOR} | {APP_ORG}"), body),
        Paragraph(esc(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"), body),
        HRFlowable(width="100%", color=navy, spaceAfter=10),
    ]
    story.append(Paragraph("Taxonomy", h1))
    for axis_name, axis in taxonomy_data().items():
        story.append(Paragraph(esc(axis_name), h2))
        story.append(Paragraph(esc(axis["description"]), body))
        for item_name, item in axis["items"].items():
            story.append(Paragraph(f"<b>{esc(item_name)}</b> — {esc(item.get('summary', ''))}", body))
            if "techniques" in item:
                for t_name, t_desc in item["techniques"].items():
                    story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;&nbsp;<i>{esc(t_name)}</i>: {esc(t_desc)}", body))
        story.append(Spacer(1, 6))
    story.append(Paragraph("Framework (KDRF)", h1))
    fw = framework_data()
    for ph in fw["phases"]:
        story.append(Paragraph(esc(ph["phase"]), h2))
        story.append(Paragraph(f"<b>Goal:</b> {esc(ph['goal'])}", body))
        for a in ph["activities"]:
            story.append(Paragraph(f"• {esc(a)}", body))
        story.append(Paragraph(f"<b>Output:</b> {esc(ph['output'])}", body))
        story.append(Spacer(1, 4))
    dl = fw["decision_layer"]
    story.append(Paragraph(esc(dl["name"]), h2))
    story.append(Paragraph(esc(f"Runs during: {dl['runs_during']}"), body))
    for step in dl["pipeline"]:
        story.append(Paragraph(f"→ {esc(step)}", body))
    story.append(Paragraph(esc(dl["note"]), body))
    story.append(Spacer(1, 8))
    story.append(Paragraph("Data Dictionary", h1))
    dd = data_dictionary_df()
    header_row = [Paragraph(f"<b>{esc(c)}</b>", body) for c in dd.columns]
    data_rows = [[Paragraph(esc(cell), body) for cell in row] for row in dd.astype(str).values.tolist()]
    table_data = [header_row] + data_rows
    col_widths = [55, 75, 80, 195, 95]
    ptable = Table(table_data, colWidths=col_widths, repeatRows=1)
    ptable.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), navy),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTSIZE", (0, 0), (-1, -1), 7.5),
        ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#999999")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F2F5FA")]),
    ]))
    story.append(ptable)
    story.append(Spacer(1, 10))
    story.append(Paragraph("Formulas", h1))
    for f in formulas_data():
        story.append(Paragraph(esc(f"{f['id']} — {f['name']}"), h2))
        story.append(Paragraph(esc(f["latex"]), mono))
        story.append(Paragraph(esc(f["explanation"]), body))
        for v, d in f["variables"].items():
            story.append(Paragraph(f"<b>{esc(v)}</b>: {esc(d)}", body))
        story.append(Spacer(1, 4))
    story.append(Paragraph("Glossary", h1))
    for term, definition in sorted(glossary_data().items()):
        story.append(Paragraph(f"<b>{esc(term)}</b>: {esc(definition)}", body))
    pdoc.build(story)
    buf.seek(0)
    return buf
# ==========================================================================
# MAIN APP
# ==========================================================================
def main():
    st.set_page_config(page_title="KNet Taxonomy & Framework Explorer", page_icon="📚", layout="wide")
    render_header()
    tabs = st.tabs(
        ["🏠 Overview", "🧬 Taxonomy", "🧭 Framework", "📋 Data Dictionary",
         "🧮 Formulas (Interactive)", "📖 Glossary", "💾 Export Reference"]
    )
    # ---------------- Overview ----------------
    with tabs[0]:
        st.markdown("### What is this page?")
        st.write(
            "This is the reference companion to the **Enterprise Recovery Orchestrator** app. "
            "Where that app *runs* simulated recoveries, this app *explains the rules behind them* — "
            "how recovery work is classified (taxonomy), the process it follows end-to-end (framework), "
            "every data field involved (data dictionary), and every calculation used to score a run (formulas)."
        )
        st.markdown("### How the pieces fit together")
        c1, c2, c3, c4 = st.columns(4)
        with c1:
            st.markdown("**🧬 Taxonomy**")
            st.caption("Classifies *what* a recovery record is: domain, technique, outcome, confidence, tool.")
        with c2:
            st.markdown("**🧭 Framework**")
            st.caption("Describes *the process* a recovery case moves through, end to end.")
        with c3:
            st.markdown("**📋 Data Dictionary**")
            st.caption("Defines *every field* recorded about a recovery attempt.")
        with c4:
            st.markdown("**🧮 Formulas**")
            st.caption("Shows *the math* that turns raw records into run-level scores and plans.")
        st.info(
            "💡 Tip: start with **Taxonomy** to learn the vocabulary, then **Framework** to see the "
            "process those terms apply to, then try the **Formulas** tab's calculator with your own numbers.",
            icon="💡",
        )
    # ---------------- Taxonomy ----------------
    with tabs[1]:
        st.markdown("### Taxonomy — how a recovery record is classified")
        st.write(
            "A single recovery record carries a value from **each of four independent axes** "
            "simultaneously. Treating these as separate axes (rather than one flat list of categories) "
            "is what makes the taxonomy composable: any Domain can pair with any Outcome and any "
            "Confidence Band."
        )
        tax = taxonomy_data()
        for axis_name, axis in tax.items():
            with st.expander(f"**{axis_name}**", expanded=(axis_name.startswith("Axis 1"))):
                st.write(axis["description"])
                for item_name, item in axis["items"].items():
                    color = item.get("color")
                    if color:
                        colored_chip(item_name, color)
                        st.write(item.get("summary", ""))
                    else:
                        st.markdown(f"**{item_name}**")
                        st.write(item.get("summary", ""))
                    if "techniques" in item:
                        for t_name, t_desc in item["techniques"].items():
                            st.markdown(f"&nbsp;&nbsp;&nbsp;&nbsp;• **{t_name}** — {t_desc}", unsafe_allow_html=True)
                    st.markdown("")
    # ---------------- Framework ----------------
    with tabs[2]:
        st.markdown("### Framework — KDRF (KNet Data Recovery Framework)")
        st.write(
            "Every recovery engagement — automated or manual — moves through the same five phases. "
            "An **Agentic AI Decision Layer** runs across phases 3–5, turning analyzer output into a "
            "concrete recommendation."
        )
        fw = framework_data()
        cols = st.columns(len(fw["phases"]))
        for i, (col, ph) in enumerate(zip(cols, fw["phases"])):
            with col:
                phase_box(ph["phase"], i, len(fw["phases"]))
                if i < len(fw["phases"]) - 1:
                    st.markdown("<div style='text-align:center;font-size:20px;color:#3D63B8;'>→</div>", unsafe_allow_html=True)
        st.markdown("")
        for ph in fw["phases"]:
            with st.expander(f"{ph['phase']} — {ph['goal']}"):
                st.markdown("**Key activities:**")
                for a in ph["activities"]:
                    st.markdown(f"- {a}")
                st.markdown(f"**Output:** {ph['output']}")
        st.markdown("---")
        dl = fw["decision_layer"]
        st.markdown(f"#### 🤖 {dl['name']}")
        st.caption(f"Runs during: {dl['runs_during']}")
        pipeline_html = " &nbsp;→&nbsp; ".join(dl["pipeline"])
        st.markdown(
            f"""<div style="background:#0B3D91;color:#D7E3FF;border-radius:10px;padding:14px 18px;
            font-family:monospace;font-size:12.5px;line-height:1.8;">{pipeline_html}</div>""",
            unsafe_allow_html=True,
        )
        st.write(dl["note"])
    # ---------------- Data Dictionary ----------------
    with tabs[3]:
        st.markdown("### Data Dictionary")
        st.write("Every field used in the underlying data model, what it means, and an example value.")
        df = data_dictionary_df()
        st.dataframe(df, use_container_width=True, height=560)
        st.download_button(
            "⬇️ Download Data Dictionary (CSV)",
            data=df.to_csv(index=False).encode("utf-8"),
            file_name="data_dictionary.csv",
            mime="text/csv",
        )
    # ---------------- Formulas ----------------
    with tabs[4]:
        st.markdown("### Formulas — the math behind every score")
        st.write(
            "Each formula below shows the math, explains it in plain language, and defines every "
            "variable. Use the **calculator at the bottom** to try your own numbers."
        )
        for f in formulas_data():
            st.markdown(f"#### {f['id']} — {f['name']}")
            st.latex(f["latex"])
            st.write(f["explanation"])
            with st.expander("Variable legend"):
                for v, d in f["variables"].items():
                    st.markdown(f"- `{v}` — {d}")
            st.markdown("---")
        st.markdown("### 🧮 Try it yourself: live calculator")
        st.caption("Enter record counts to see Formulas F1–F3 computed live, with the rule that fired highlighted.")
        cc1, cc2, cc3 = st.columns(3)
        r = cc1.number_input("Recovered (R)", min_value=0, value=61, step=1)
        p = cc2.number_input("Partial (P)", min_value=0, value=24, step=1)
        fcount = cc3.number_input("Failed (F)", min_value=0, value=15, step=1)
        prob, n = calc_recovery_probability(int(r), int(p), int(fcount))
        confidence = calc_confidence(prob)
        minutes, time_str = calc_estimated_time(n, int(p), int(fcount))
        r1, r2, r3 = st.columns(3)
        r1.metric("Recovery Probability (F1)", f"{prob}%")
        r2.metric("Confidence Level (F2)", confidence)
        r3.metric("Estimated Time (F3)", time_str)
        st.caption(
            f"N = {n} total records · F1 = (({r}×1.0)+({p}×0.5)) / {n} × 100 = {prob}% · "
            f"F2 rule fired: "
            + ("P_r ≥ 85 → High" if prob >= 85 else "60 ≤ P_r < 85 → Medium" if prob >= 60 else "P_r < 60 → Low")
            + f" · F3 = ({n}×0.8) + ({fcount}×3) + ({p}×1.5) = {round(minutes,1)} minutes"
        )
    # ---------------- Glossary ----------------
    with tabs[5]:
        st.markdown("### Glossary")
        search = st.text_input("Filter terms", placeholder="e.g. snapshot, hash, IAM…")
        glossary = glossary_data()
        for term, definition in sorted(glossary.items()):
            if search and search.lower() not in term.lower() and search.lower() not in definition.lower():
                continue
            st.markdown(f"**{term}**")
            st.write(definition)
            st.markdown("")
    # ---------------- Export ----------------
    with tabs[6]:
        st.markdown("### Export the full reference")
        st.write(
            "Download the entire taxonomy, framework, data dictionary, formulas, and glossary for "
            "offline use — choose Markdown, JSON, plain Text, Word (.docx), or PDF."
        )
        md_doc = build_markdown_reference()
        json_doc = build_json_reference()
        txt_doc = build_text_reference()
        e1, e2, e3, e4, e5 = st.columns(5)
        with e1:
            st.download_button(
                "⬇️ Markdown",
                data=md_doc.encode("utf-8"),
                file_name="data_recovery_taxonomy_framework.md",
                mime="text/markdown",
                use_container_width=True,
            )
        with e2:
            st.download_button(
                "⬇️ JSON",
                data=json.dumps(json_doc, indent=2).encode("utf-8"),
                file_name="data_recovery_taxonomy_framework.json",
                mime="application/json",
                use_container_width=True,
            )
        with e3:
            st.download_button(
                "⬇️ Text (.txt)",
                data=txt_doc.encode("utf-8"),
                file_name="data_recovery_taxonomy_framework.txt",
                mime="text/plain",
                use_container_width=True,
            )
        with e4:
            if DOCX_OK:
                docx_buf = build_docx_reference()
                st.download_button(
                    "⬇️ Word (.docx)",
                    data=docx_buf,
                    file_name="data_recovery_taxonomy_framework.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            else:
                st.button(
                    "⬇️ Word (.docx)", disabled=True, use_container_width=True,
                    help="python-docx is not installed in this environment. Add 'python-docx' to requirements.txt.",
                )
        with e5:
            if REPORTLAB_OK:
                pdf_buf = build_pdf_reference()
                st.download_button(
                    "⬇️ PDF",
                    data=pdf_buf,
                    file_name="data_recovery_taxonomy_framework.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                )
            else:
                st.button(
                    "⬇️ PDF", disabled=True, use_container_width=True,
                    help="reportlab is not installed in this environment. Add 'reportlab' to requirements.txt.",
                )
        with st.expander("Preview Markdown reference"):
            st.markdown(md_doc)
    st.markdown("---")
    st.caption(f"{APP_TITLE} • {APP_ORG} • Reference/educational content — pairs with the Enterprise Recovery Orchestrator app.")
if __name__ == "__main__":
    main()