


# CogniReason AI Suite
# Developed by Randy Singh - Kalsnet (KNet) Consulting Group

# A professional multi-tab Streamlit application demonstrating 20 real-world
# Reasoning AI use cases, powered by live LLM reasoning (your choice of the
# Groq API or the Google Gemini API). Each tab includes:
# - A plain-English description of the use case
# - Formula / field / relevance documentation for the underlying schema
# - A synthetic data generator (with adjustable sample size)
# - A real-data CSV upload option that maps to the same schema
# - A graphic (Plotly) visualization of the data
# - A "Run AI Reasoning" button that sends a structured prompt + data
    # summary to the selected LLM and displays step-by-step reasoning
    # Export of the results to PDF, Word (.docx), Text (.txt) and CSV


import io
import os
import json
from datetime import datetime

import numpy as np
import pandas as pd
import requests
import streamlit as st
import plotly.express as px
import plotly.graph_objects as go

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import LETTER
from reportlab.lib.units import inch
from reportlab.lib import colors as rl_colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
)

# --------------------------------------------------------------------------
# PAGE CONFIG
# --------------------------------------------------------------------------
st.set_page_config(
    page_title="CogniReason AI Suite",
    layout="wide",
    initial_sidebar_state="expanded",
)

APP_NAME = "CogniReason AI Suite"
DEVELOPER_LINE = "Developed by Randy Singh &nbsp;|&nbsp; Kalsnet (KNet) Consulting Group"

# --------------------------------------------------------------------------
# GLOBAL STYLE
# --------------------------------------------------------------------------
st.markdown(
    """
    <style>
    .main-title {
        font-size: 54px;
        font-weight: 900;
        color: #0B3D91;
        text-align: center;
        letter-spacing: 1px;
        margin-bottom: 0px;
        text-shadow: 1px 1px 0px rgba(11,61,145,0.15);
    }
    .sub-title {
        text-align: center;
        font-size: 18px;
        color: #444444;
        font-weight: 500;
        margin-top: 4px;
        margin-bottom: 18px;
    }
    .kn-badge {
        display: inline-block;
        background: linear-gradient(90deg,#0B3D91,#1E6FEB);
        color: white;
        padding: 3px 12px;
        border-radius: 14px;
        font-size: 13px;
        font-weight: 700;
        letter-spacing: 0.5px;
    }
    .section-card {
        background-color: #F5F8FF;
        border: 1px solid #D6E0F5;
        border-radius: 10px;
        padding: 16px 20px;
        margin-bottom: 14px;
    }
    .field-row {
        border-bottom: 1px solid #E6E6E6;
        padding: 6px 0px;
    }
    .stTabs [data-baseweb="tab-list"] {
        gap: 4px;
    }
    .stTabs [data-baseweb="tab"] {
        background-color: #EEF3FC;
        border-radius: 8px 8px 0 0;
        padding: 10px 14px;
        font-weight: 600;
        color: #0B3D91;
    }
    .stTabs [aria-selected="true"] {
        background-color: #0B3D91 !important;
        color: white !important;
    }
    hr {margin-top: 6px; margin-bottom: 6px;}
    </style>
    """,
    unsafe_allow_html=True,
)

st.markdown(f"<div class='main-title'>{APP_NAME}</div>", unsafe_allow_html=True)
st.markdown(
    f"<div class='sub-title'>{DEVELOPER_LINE} &nbsp; "
    f"<span class='kn-badge'>Reasoning AI - Powered by Groq LLM</span></div>",
    unsafe_allow_html=True,
)
st.markdown("<hr>", unsafe_allow_html=True)

# --------------------------------------------------------------------------
# SIDEBAR - GLOBAL CONFIG
# --------------------------------------------------------------------------
with st.sidebar:
    st.markdown("## AI Engine Configuration")

    llm_provider = st.radio(
        "AI Provider", ["Groq", "Google Gemini"], horizontal=True,
        help="Pick whichever free API key you have (or want to create).",
    )

    if llm_provider == "Groq":
        groq_api_key = st.text_input(
            "Groq API Key",
            type="password",
            value=os.environ.get("GROQ_API_KEY", ""),
            help="Get a free key at console.groq.com. Never stored - used only for this session.",
        )
        gemini_api_key = ""
        groq_model = st.selectbox(
            "Reasoning Model",
            [
                "llama-3.3-70b-versatile",
                "llama-3.1-8b-instant",
                "mixtral-8x7b-32768",
                "gemma2-9b-it",
            ],
            index=0,
        )
        gemini_model = None
    else:
        gemini_api_key = st.text_input(
            "Gemini API Key",
            type="password",
            value=os.environ.get("GEMINI_API_KEY", ""),
            help="Get a free key at aistudio.google.com. Never stored - used only for this session.",
        )
        groq_api_key = ""
        gemini_model = st.selectbox(
            "Reasoning Model",
            [
                "gemini-2.5-flash",
                "gemini-2.5-flash-lite",
                "gemini-3-flash",
            ],
            index=0,
        )
        groq_model = None

    reasoning_depth = st.select_slider(
        "Reasoning Depth",
        options=["Concise", "Standard", "Deep Chain-of-Thought"],
        value="Standard",
    )

    with st.expander("How to get a FREE Groq or Gemini API key"):
        st.markdown(
            """
**Option A - Groq (fast, free, Llama/Mixtral/Gemma models)**
1. Go to **console.groq.com** and sign up (Google/GitHub login works).
2. Open **API Keys** in the left menu -> **Create API Key**.
3. Copy the key (starts with `gsk_...`) and paste it above.
4. Groq's free tier gives generous requests/minute - plenty for this app.

**Option B - Google Gemini (free, Google account)**
1. Go to **aistudio.google.com** and sign in with a Google account.
2. Click **Get API key** -> **Create API key** (choose a Google Cloud
   project, or let Studio create one for you - no credit card required
   for the free tier).
3. Copy the key (starts with `AIza...`) and paste it above.
4. The free tier covers the **Flash** / **Flash-Lite** models selected
   in the dropdown; quotas reset daily.

You only need **one** of the two keys - pick the provider above that
matches the key you have. Keys are used only in-memory for this
session and are never written to disk.
            """
        )

    st.markdown("---")
    st.caption(
        "Each tab below can run on **synthetic data** (generated instantly) "
        "or on **your uploaded CSV** matching the documented schema."
    )
    st.markdown("---")
    st.markdown("### Export")
    st.caption("Every tab has its own PDF / Word / Text / CSV export buttons "
                "once you've generated data and run AI reasoning.")
    st.markdown("---")
    st.caption("(c) " + str(datetime.now().year) + " Kalsnet (KNet) Consulting Group")

# --------------------------------------------------------------------------
# GROQ LLM CALL
# --------------------------------------------------------------------------
def call_groq(system_prompt: str, user_prompt: str, api_key: str, model: str,
               max_tokens: int = 1400) -> str:
    if not api_key:
        return (
            "**No Groq API key provided.** Enter your key in the left sidebar "
            "to generate live AI reasoning output. Below is where the model's "
            "step-by-step reasoning, key drivers, and recommendation will appear."
        )
    url = "https://api.groq.com/openai/v1/chat/completions"
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        "temperature": 0.3,
        "max_tokens": max_tokens,
    }
    try:
        resp = requests.post(url, headers=headers, json=payload, timeout=90)
        resp.raise_for_status()
        return resp.json()["choices"][0]["message"]["content"]
    except requests.exceptions.HTTPError as e:
        return f"Groq API error: {e} - {resp.text[:300]}"
    except Exception as e:
        return f"Error calling Groq API: {e}"


def call_gemini(system_prompt: str, user_prompt: str, api_key: str, model: str,
                 max_tokens: int = 1400) -> str:
    if not api_key:
        return (
            "**No Gemini API key provided.** Enter your key in the left sidebar "
            "to generate live AI reasoning output. Below is where the model's "
            "step-by-step reasoning, key drivers, and recommendation will appear."
        )
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent"
    headers = {"Content-Type": "application/json", "x-goog-api-key": api_key}
    payload = {
        "system_instruction": {"parts": [{"text": system_prompt}]},
        "contents": [{"role": "user", "parts": [{"text": user_prompt}]}],
        "generationConfig": {"temperature": 0.3, "maxOutputTokens": max_tokens},
    }
    try:
        resp = requests.post(url, headers=headers, json=payload, timeout=90)
        resp.raise_for_status()
        data = resp.json()
        parts = data["candidates"][0]["content"]["parts"]
        return "".join(p.get("text", "") for p in parts)
    except requests.exceptions.HTTPError as e:
        return f"Gemini API error: {e} - {resp.text[:300]}"
    except Exception as e:
        return f"Error calling Gemini API: {e}"


def call_llm(system_prompt: str, user_prompt: str) -> str:
    """Routes to whichever provider is selected in the sidebar."""
    if llm_provider == "Groq":
        return call_groq(system_prompt, user_prompt, groq_api_key, groq_model)
    else:
        return call_gemini(system_prompt, user_prompt, gemini_api_key, gemini_model)


def depth_instruction(depth: str) -> str:
    return {
        "Concise": "Give a brief, bullet-point reasoning summary (5-8 bullets) and a final recommendation.",
        "Standard": "Reason step-by-step through the key drivers found in the data, cite the specific "
                    "fields/values that matter most, then give a clear final recommendation with confidence level.",
        "Deep Chain-of-Thought": "Provide a detailed chain-of-thought: (1) restate the objective, "
                                  "(2) walk through each relevant field and what it implies, "
                                  "(3) weigh conflicting signals, (4) state assumptions, "
                                  "(5) give a final recommendation with a confidence level and caveats.",
    }[depth]

# --------------------------------------------------------------------------
# EXPORT HELPERS
# --------------------------------------------------------------------------
def df_to_csv_bytes(df: pd.DataFrame) -> bytes:
    return df.to_csv(index=False).encode("utf-8")


def text_bytes(title: str, description: str, reasoning: str) -> bytes:
    content = (
        f"{APP_NAME}\n{DEVELOPER_LINE.replace('&nbsp;', ' ')}\n"
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n"
        f"USE CASE: {title}\n{'-'*60}\n{description}\n\n"
        f"AI REASONING OUTPUT\n{'-'*60}\n{reasoning}\n"
    )
    return content.encode("utf-8")


def build_docx(title: str, description: str, df: pd.DataFrame, reasoning: str) -> bytes:
    doc = Document()
    h = doc.add_heading(APP_NAME, level=0)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0B, 0x3D, 0x91)
    p = doc.add_paragraph(DEVELOPER_LINE.replace("&nbsp;", " "))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_heading(title, level=1)
    doc.add_paragraph(description)

    if df is not None and len(df):
        doc.add_heading("Data Sample (first 15 rows)", level=2)
        sample = df.head(15)
        table = doc.add_table(rows=1, cols=len(sample.columns))
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, col in enumerate(sample.columns):
            hdr[i].text = str(col)
        for _, row in sample.iterrows():
            cells = table.add_row().cells
            for i, col in enumerate(sample.columns):
                cells[i].text = str(row[col])

    doc.add_heading("AI Reasoning Output", level=2)
    for line in reasoning.split("\n"):
        doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_pdf(title: str, description: str, df: pd.DataFrame, reasoning: str) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=LETTER,
                             topMargin=0.6 * inch, bottomMargin=0.6 * inch)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleBlue", parent=styles["Title"], textColor=rl_colors.HexColor("#0B3D91"),
        fontSize=22,
    )
    story = [
        Paragraph(APP_NAME, title_style),
        Paragraph(DEVELOPER_LINE.replace("&nbsp;", " "), styles["Normal"]),
        Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", styles["Normal"]),
        Spacer(1, 14),
        Paragraph(title, styles["Heading1"]),
        Paragraph(description, styles["Normal"]),
        Spacer(1, 10),
    ]

    if df is not None and len(df):
        story.append(Paragraph("Data Sample (first 12 rows)", styles["Heading2"]))
        sample = df.head(12)
        data = [list(sample.columns)] + sample.astype(str).values.tolist()
        tbl = Table(data, repeatRows=1)
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), rl_colors.HexColor("#0B3D91")),
            ("TEXTCOLOR", (0, 0), (-1, 0), rl_colors.white),
            ("FONTSIZE", (0, 0), (-1, -1), 7),
            ("GRID", (0, 0), (-1, -1), 0.5, rl_colors.grey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [rl_colors.white, rl_colors.HexColor("#F5F8FF")]),
        ]))
        story.append(tbl)
        story.append(Spacer(1, 14))

    story.append(Paragraph("AI Reasoning Output", styles["Heading2"]))
    for line in reasoning.split("\n"):
        if line.strip():
            story.append(Paragraph(line.replace("&", "&amp;").replace("<", "&lt;"), styles["Normal"]))
        else:
            story.append(Spacer(1, 6))

    doc.build(story)
    return buf.getvalue()

# --------------------------------------------------------------------------
# SYNTHETIC DATA GENERATORS  (one per use case)
# --------------------------------------------------------------------------
rng = np.random.default_rng()

def gen_financial_risk(n):
    df = pd.DataFrame({
        "account_id": [f"ACCT-{1000+i}" for i in range(n)],
        "account_age_years": rng.integers(1, 25, n),
        "credit_limit": rng.integers(1000, 30000, n),
        "current_balance": rng.integers(0, 25000, n),
        "monthly_income": rng.integers(2000, 15000, n),
        "monthly_debt_payments": rng.integers(100, 6000, n),
        "payment_history_score": rng.integers(300, 850, n),
    })
    df["credit_utilization_pct"] = (df["current_balance"] / df["credit_limit"] * 100).round(1)
    df["debt_to_income_pct"] = (df["monthly_debt_payments"] / df["monthly_income"] * 100).round(1)
    df["risk_score"] = (
        0.4 * (df["credit_utilization_pct"]) +
        0.4 * (df["debt_to_income_pct"]) -
        0.02 * (df["payment_history_score"] - 300)
    ).round(1)
    df["risk_category"] = pd.cut(df["risk_score"], bins=[-999, 20, 45, 999],
                                  labels=["Low", "Medium", "High"])
    return df

def gen_rca_manufacturing(n):
    planned = rng.integers(400, 480, n)
    downtime = rng.integers(5, 120, n)
    uptime = planned - downtime
    theoretical_max = rng.integers(800, 1200, n)
    actual_output = (theoretical_max * rng.uniform(0.6, 0.98, n)).astype(int)
    total_units = actual_output
    good_units = (total_units * rng.uniform(0.85, 0.999, n)).astype(int)
    df = pd.DataFrame({
        "machine_id": [f"MCH-{100+i}" for i in range(n)],
        "planned_minutes": planned,
        "downtime_minutes": downtime,
        "theoretical_max_units": theoretical_max,
        "actual_output_units": actual_output,
        "good_units": good_units,
        "avg_temperature_c": rng.uniform(55, 95, n).round(1),
        "avg_vibration_mm_s": rng.uniform(0.5, 6.0, n).round(2),
    })
    df["availability"] = (uptime / planned).round(3)
    df["performance"] = (df["actual_output_units"] / df["theoretical_max_units"]).round(3)
    df["quality"] = (df["good_units"] / df["actual_output_units"]).round(3)
    df["oee_pct"] = (df["availability"] * df["performance"] * df["quality"] * 100).round(1)
    return df

def gen_medical_diagnosis(n):
    height = rng.uniform(1.5, 1.95, n)
    weight = rng.uniform(50, 130, n)
    df = pd.DataFrame({
        "patient_id": [f"PT-{5000+i}" for i in range(n)],
        "age": rng.integers(18, 85, n),
        "height_m": height.round(2),
        "weight_kg": weight.round(1),
        "systolic_bp": rng.integers(95, 190, n),
        "glucose_mg_dl": rng.integers(70, 240, n),
        "symptom_severity_score": rng.integers(0, 10, n),
    })
    df["bmi"] = (df["weight_kg"] / (df["height_m"] ** 2)).round(1)
    df["risk_flag"] = np.where(
        (df["bmi"] > 30) | (df["systolic_bp"] > 140) | (df["glucose_mg_dl"] > 180),
        "Elevated", "Normal"
    )
    return df

def gen_legal_contract(n):
    clause_types = ["Indemnification", "Termination", "Limitation of Liability",
                     "Confidentiality", "Force Majeure", "Payment Terms"]
    df = pd.DataFrame({
        "contract_id": [f"CTR-{200+i}" for i in range(n)],
        "clause_type": rng.choice(clause_types, n),
        "ambiguity_score": rng.integers(1, 10, n),
        "liability_exposure_usd": rng.integers(5000, 2_000_000, n),
        "termination_notice_days": rng.choice([0, 15, 30, 60, 90], n),
        "counterparty_credit_rating": rng.choice(["AAA", "AA", "A", "BBB", "BB", "B"], n),
    })
    rating_weight = df["counterparty_credit_rating"].map(
        {"AAA": 1, "AA": 2, "A": 3, "BBB": 4, "BB": 5, "B": 6}
    )
    df["risk_score"] = (
        0.5 * df["ambiguity_score"] +
        0.0002 * df["liability_exposure_usd"] / 100 +
        rating_weight * 1.2
    ).round(1)
    return df

def gen_fraud_detection(n):
    df = pd.DataFrame({
        "txn_id": [f"TXN-{90000+i}" for i in range(n)],
        "amount_usd": rng.uniform(5, 9000, n).round(2),
        "txn_velocity_1hr": rng.integers(1, 20, n),
        "geo_mismatch_flag": rng.integers(0, 2, n),
        "device_change_flag": rng.integers(0, 2, n),
        "new_merchant_flag": rng.integers(0, 2, n),
    })
    df["fraud_score"] = (
        0.15 * df["txn_velocity_1hr"] +
        25 * df["geo_mismatch_flag"] +
        20 * df["device_change_flag"] +
        10 * df["new_merchant_flag"] +
        df["amount_usd"] / 200
    ).round(1)
    df["fraud_flag"] = np.where(df["fraud_score"] > 40, "Suspicious", "Normal")
    return df

def gen_churn(n):
    tenure = rng.integers(1, 72, n)
    charges = rng.uniform(20, 150, n)
    tickets = rng.integers(0, 8, n)
    nps = rng.integers(-100, 100, n)
    z = -0.04 * tenure + 0.01 * charges + 0.35 * tickets - 0.01 * nps - 1.2
    churn_prob = 1 / (1 + np.exp(-z))
    df = pd.DataFrame({
        "customer_id": [f"CUST-{4000+i}" for i in range(n)],
        "tenure_months": tenure,
        "monthly_charges_usd": charges.round(2),
        "support_tickets_90d": tickets,
        "nps_score": nps,
    })
    df["churn_probability_pct"] = (churn_prob * 100).round(1)
    return df

def gen_supply_chain(n):
    df = pd.DataFrame({
        "supplier_id": [f"SUP-{300+i}" for i in range(n)],
        "avg_lead_time_days": rng.integers(3, 90, n),
        "on_time_delivery_rate_pct": rng.uniform(60, 100, n).round(1),
        "geopolitical_risk_index": rng.uniform(0, 10, n).round(1),
        "inventory_days_cover": rng.integers(2, 60, n),
        "single_source_flag": rng.integers(0, 2, n),
    })
    df["disruption_risk_score"] = (
        (100 - df["on_time_delivery_rate_pct"]) * 0.5 +
        df["geopolitical_risk_index"] * 4 +
        df["single_source_flag"] * 15 -
        df["inventory_days_cover"] * 0.3
    ).round(1)
    return df

def gen_credit_underwriting(n):
    income = rng.integers(25000, 220000, n)
    loan = rng.integers(5000, 500000, n)
    asset_value = loan * rng.uniform(1.05, 1.6, n)
    debt = rng.integers(200, 4000, n)
    df = pd.DataFrame({
        "applicant_id": [f"APP-{7000+i}" for i in range(n)],
        "annual_income_usd": income,
        "loan_amount_usd": loan,
        "asset_value_usd": asset_value.astype(int),
        "monthly_debt_usd": debt,
        "credit_score": rng.integers(300, 850, n),
    })
    df["ltv_ratio_pct"] = (df["loan_amount_usd"] / df["asset_value_usd"] * 100).round(1)
    df["dti_ratio_pct"] = (df["monthly_debt_usd"] * 12 / df["annual_income_usd"] * 100).round(1)
    df["approval_recommendation"] = np.where(
        (df["ltv_ratio_pct"] < 90) & (df["dti_ratio_pct"] < 43) & (df["credit_score"] > 620),
        "Approve", "Refer for Manual Review"
    )
    return df

def gen_hr_fit(n):
    exp = rng.integers(0, 20, n)
    skill = rng.uniform(30, 100, n)
    culture = rng.uniform(30, 100, n)
    interview = rng.uniform(30, 100, n)
    df = pd.DataFrame({
        "candidate_id": [f"CAND-{600+i}" for i in range(n)],
        "years_experience": exp,
        "skill_match_pct": skill.round(1),
        "culture_fit_score": culture.round(1),
        "interview_score": interview.round(1),
    })
    df["overall_fit_score"] = (
        0.4 * df["skill_match_pct"] + 0.3 * df["culture_fit_score"] +
        0.3 * df["interview_score"]
    ).round(1)
    return df

def gen_strategy_swot(n):
    units = [f"BU-{chr(65+i)}" for i in range(n)]
    growth = rng.uniform(-5, 25, n)
    share = rng.uniform(1, 40, n)
    intensity = rng.uniform(1, 10, n)
    strengths = rng.integers(1, 10, n)
    weaknesses = rng.integers(1, 10, n)
    opportunities = rng.integers(1, 10, n)
    threats = rng.integers(1, 10, n)
    df = pd.DataFrame({
        "business_unit": units,
        "market_growth_rate_pct": growth.round(1),
        "market_share_pct": share.round(1),
        "competitive_intensity": intensity.round(1),
        "strengths_score": strengths,
        "weaknesses_score": weaknesses,
        "opportunities_score": opportunities,
        "threats_score": threats,
    })
    df["swot_net_score"] = (
        (df["strengths_score"] + df["opportunities_score"]) -
        (df["weaknesses_score"] + df["threats_score"])
    )
    return df

def gen_cyber_incident(n):
    sources = ["Firewall", "EDR", "SIEM", "IDS", "Proxy", "Email Gateway"]
    tactics = ["Recon", "Initial Access", "Lateral Movement", "Exfiltration", "C2", "Privilege Escalation"]
    df = pd.DataFrame({
        "alert_id": [f"ALRT-{8000+i}" for i in range(n)],
        "source_system": rng.choice(sources, n),
        "mitre_tactic": rng.choice(tactics, n),
        "severity_1to5": rng.integers(1, 6, n),
        "threat_ip_reputation_score": rng.integers(0, 100, n),
        "affected_hosts_count": rng.integers(1, 40, n),
        "correlated_alert_count": rng.integers(0, 25, n),
    })
    df["coordination_score"] = (
        df["severity_1to5"] * 6 +
        df["threat_ip_reputation_score"] * 0.3 +
        df["correlated_alert_count"] * 2.2 +
        df["affected_hosts_count"] * 1.1
    ).round(1).clip(upper=100)
    df["assessment"] = np.where(df["coordination_score"] > 60, "Likely Coordinated Attack", "Isolated Alert")
    return df

def gen_api_security(n):
    endpoints = ["/api/login", "/api/users", "/api/payments", "/api/orders",
                 "/api/search", "/api/admin", "/api/tokens"]
    df = pd.DataFrame({
        "api_endpoint": rng.choice(endpoints, n),
        "requests_per_min": rng.integers(5, 3000, n),
        "auth_failure_rate_pct": rng.uniform(0, 80, n).round(1),
        "unique_source_ips": rng.integers(1, 500, n),
        "payload_anomaly_score": rng.uniform(0, 10, n).round(1),
        "geo_diversity_index": rng.uniform(0, 10, n).round(1),
    })
    df["attack_likelihood_pct"] = (
        df["auth_failure_rate_pct"] * 0.5 +
        df["payload_anomaly_score"] * 5 +
        df["geo_diversity_index"] * 3 +
        df["requests_per_min"] / 50
    ).round(1).clip(upper=100)
    return df

def gen_software_troubleshoot(n):
    modules = ["Auth Service", "Payment Gateway", "Order API", "Search Index",
               "Notification Worker", "Database Layer", "Frontend Bundle"]
    df = pd.DataFrame({
        "ticket_id": [f"BUG-{1200+i}" for i in range(n)],
        "module_name": rng.choice(modules, n),
        "exceptions_last_24h": rng.integers(0, 200, n),
        "stack_trace_depth": rng.integers(2, 40, n),
        "code_churn_last_7d_commits": rng.integers(0, 60, n),
        "test_coverage_pct": rng.uniform(20, 98, n).round(1),
    })
    df["root_cause_likelihood_score"] = (
        df["exceptions_last_24h"] * 0.3 +
        df["code_churn_last_7d_commits"] * 0.6 +
        df["stack_trace_depth"] * 0.4 -
        df["test_coverage_pct"] * 0.3
    ).round(1)
    return df

def gen_medical_decision_support(n):
    symptoms = ["Fever + Cough", "Abdominal Pain", "Chest Pain", "Headache + Stiff Neck",
                "Joint Pain + Rash", "Shortness of Breath"]
    df = pd.DataFrame({
        "patient_id": [f"PT-{9000+i}" for i in range(n)],
        "primary_symptom_cluster": rng.choice(symptoms, n),
        "symptom_duration_days": rng.integers(1, 21, n),
        "fever_temp_c": rng.uniform(36.5, 40.5, n).round(1),
        "wbc_count_k_per_uL": rng.uniform(3.5, 22.0, n).round(1),
        "crp_level_mg_L": rng.uniform(1, 220, n).round(1),
        "comorbidity_count": rng.integers(0, 4, n),
    })
    df["differential_urgency_score"] = (
        (df["fever_temp_c"] - 36.5) * 8 +
        (df["wbc_count_k_per_uL"] - 3.5) * 2 +
        df["crp_level_mg_L"] * 0.25 +
        df["comorbidity_count"] * 5
    ).round(1)
    return df

def gen_financial_fraud_investigation(n):
    df = pd.DataFrame({
        "customer_id": [f"CIF-{55000+i}" for i in range(n)],
        "linked_accounts_count": rng.integers(1, 8, n),
        "txn_deviation_from_avg_pct": rng.uniform(-20, 400, n).round(1),
        "distance_from_home_km": rng.uniform(0, 9000, n).round(1),
        "account_age_days": rng.integers(10, 4000, n),
        "cross_border_txn_flag": rng.integers(0, 2, n),
    })
    df["network_suspicion_score"] = (
        df["txn_deviation_from_avg_pct"] * 0.2 +
        df["distance_from_home_km"] / 100 +
        df["cross_border_txn_flag"] * 15 +
        df["linked_accounts_count"] * 3 -
        df["account_age_days"] / 200
    ).round(1)
    return df

def gen_contract_policy_analysis(n):
    doc_types = ["Master Service Agreement", "Vendor Policy", "NDA", "SLA Addendum", "Compliance Policy"]
    df = pd.DataFrame({
        "document_id": [f"DOC-{400+i}" for i in range(n)],
        "document_type": rng.choice(doc_types, n),
        "obligation_count": rng.integers(2, 40, n),
        "exception_count": rng.integers(0, 15, n),
        "cross_reference_conflicts": rng.integers(0, 10, n),
        "penalty_severity_usd": rng.integers(0, 500000, n),
        "compliance_deadline_days": rng.integers(1, 180, n),
    })
    df["conflict_risk_score"] = (
        df["cross_reference_conflicts"] * 6 +
        df["exception_count"] * 2 +
        df["penalty_severity_usd"] / 20000 +
        (180 - df["compliance_deadline_days"]) * 0.1
    ).round(1)
    return df

def gen_predictive_maintenance(n):
    df = pd.DataFrame({
        "equipment_id": [f"EQ-{700+i}" for i in range(n)],
        "vibration_trend_slope_mm_s_per_wk": rng.uniform(-0.1, 1.5, n).round(3),
        "temperature_delta_c": rng.uniform(-5, 35, n).round(1),
        "run_hours_since_maintenance": rng.integers(10, 8000, n),
        "historical_failures_last_12mo": rng.integers(0, 6, n),
        "oil_particle_count_ppm": rng.integers(5, 400, n),
    })
    df["predicted_failure_probability_pct"] = (
        df["vibration_trend_slope_mm_s_per_wk"] * 40 +
        df["temperature_delta_c"] * 1.2 +
        df["run_hours_since_maintenance"] / 150 +
        df["historical_failures_last_12mo"] * 8 +
        df["oil_particle_count_ppm"] * 0.1
    ).round(1).clip(0, 100)
    return df

def gen_emergency_response(n):
    incident_types = ["Fire", "Medical", "HazMat", "Structural Collapse", "Flood"]
    df = pd.DataFrame({
        "incident_id": [f"INC-{50+i}" for i in range(n)],
        "incident_type": rng.choice(incident_types, n),
        "severity_level_1to5": rng.integers(1, 6, n),
        "population_at_risk": rng.integers(1, 5000, n),
        "distance_to_nearest_unit_km": rng.uniform(0.5, 40, n).round(1),
        "resource_availability_pct": rng.uniform(20, 100, n).round(1),
    })
    df["response_priority_score"] = (
        df["severity_level_1to5"] * 15 +
        np.log1p(df["population_at_risk"]) * 5 -
        df["distance_to_nearest_unit_km"] * 0.8 -
        df["resource_availability_pct"] * 0.2
    ).round(1)
    return df

def gen_aiops(n):
    services = ["checkout-svc", "auth-svc", "inventory-svc", "recommendation-svc",
                "payment-svc", "notification-svc"]
    df = pd.DataFrame({
        "service_name": rng.choice(services, n),
        "error_rate_pct": rng.uniform(0, 25, n).round(2),
        "latency_p95_ms": rng.integers(50, 4000, n),
        "cpu_utilization_pct": rng.uniform(10, 99, n).round(1),
        "recent_deployment_flag": rng.integers(0, 2, n),
        "dependency_failure_count": rng.integers(0, 10, n),
    })
    df["outage_root_cause_score"] = (
        df["error_rate_pct"] * 3 +
        df["latency_p95_ms"] / 100 +
        df["cpu_utilization_pct"] * 0.4 +
        df["recent_deployment_flag"] * 20 +
        df["dependency_failure_count"] * 5
    ).round(1)
    return df

def gen_autonomous_agents(n):
    df = pd.DataFrame({
        "task_id": [f"TASK-{900+i}" for i in range(n)],
        "task_complexity_score_1to10": rng.integers(1, 11, n),
        "available_tools_count": rng.integers(1, 8, n),
        "planned_tool_calls": rng.integers(1, 15, n),
        "context_window_used_pct": rng.uniform(5, 95, n).round(1),
        "fallback_triggered_flag": rng.integers(0, 2, n),
    })
    df["estimated_success_probability_pct"] = (
        100 - df["task_complexity_score_1to10"] * 6 +
        df["available_tools_count"] * 2 -
        df["context_window_used_pct"] * 0.15 -
        df["fallback_triggered_flag"] * 10
    ).round(1).clip(0, 100)
    return df

# --------------------------------------------------------------------------
# USE CASE CONFIGURATION
# --------------------------------------------------------------------------
USE_CASES = [
    {
        "id": "financial_risk",
        "title": "Financial Risk Assessment",
        "description": (
            "Reasons over an account's credit behavior to classify counterparty risk, "
            "the way a credit-risk analyst would triage a portfolio before deeper underwriting."
        ),
        "fields": [
            ("credit_utilization_pct", "current_balance / credit_limit x 100",
             "High utilization signals over-reliance on revolving credit and rising default risk."),
            ("debt_to_income_pct", "monthly_debt_payments / monthly_income x 100",
             "Core affordability metric; lenders typically flag DTI above ~40-43%."),
            ("payment_history_score", "Bureau-style score, 300-850",
             "Longer track record of on-time payments lowers probability of default."),
            ("risk_score", "0.4xUtilization + 0.4xDTI - 0.02x(Score-300)",
             "Composite score blending affordability and credit behavior into one risk index."),
        ],
        "generate": gen_financial_risk,
        "chart": lambda df: px.scatter(
            df, x="debt_to_income_pct", y="credit_utilization_pct",
            color="risk_category", size="risk_score",
            title="Debt-to-Income vs Credit Utilization (bubble = risk score)"
        ),
        "system_prompt": (
            "You are a senior credit-risk officer performing reasoning AI analysis on consumer "
            "account data. Explain WHY accounts land in each risk tier using the specific fields "
            "provided (utilization, DTI, payment history), then give a portfolio-level recommendation."
        ),
    },
    {
        "id": "rca",
        "title": "Root Cause Analysis (Manufacturing)",
        "description": (
            "Applies reasoning to machine telemetry to explain why Overall Equipment "
            "Effectiveness (OEE) is underperforming and what root cause is most likely."
        ),
        "fields": [
            ("availability", "(planned_minutes - downtime_minutes) / planned_minutes",
             "Captures how much of scheduled time the machine was actually running."),
            ("performance", "actual_output_units / theoretical_max_units",
             "Captures speed losses versus the machine's rated throughput."),
            ("quality", "good_units / actual_output_units",
             "Captures yield losses - defects and rework."),
            ("oee_pct", "Availability x Performance x Quality x 100",
             "Single industry-standard KPI; world-class OEE is typically 85%+."),
        ],
        "generate": gen_rca_manufacturing,
        "chart": lambda df: px.scatter(
            df, x="avg_temperature_c", y="oee_pct", color="avg_vibration_mm_s",
            size="downtime_minutes", title="OEE vs Temperature (color = vibration, size = downtime)"
        ),
        "system_prompt": (
            "You are a manufacturing reliability engineer performing root-cause reasoning on OEE "
            "telemetry. Identify whether availability, performance, or quality losses dominate, "
            "correlate with temperature/vibration anomalies, and recommend a corrective action plan."
        ),
    },
    {
        "id": "medical",
        "title": "Medical Differential Diagnosis Support",
        "description": (
            "Reasoning support (NOT a diagnosis) that highlights which vitals/labs are driving an "
            "elevated-risk flag, to help a clinician prioritize chart review."
        ),
        "fields": [
            ("bmi", "weight_kg / (height_m)^2",
             "Body Mass Index; >30 is classified as obese and raises cardiometabolic risk."),
            ("systolic_bp", "mmHg, upper blood-pressure reading",
             ">140 mmHg indicates Stage 2 hypertension per standard clinical thresholds."),
            ("glucose_mg_dl", "Fasting/random glucose reading",
             ">180 mg/dL suggests poor glycemic control."),
            ("risk_flag", "Elevated if BMI>30 OR SBP>140 OR Glucose>180",
             "Simple rule-based triage flag that the LLM is asked to reason about qualitatively."),
        ],
        "generate": gen_medical_diagnosis,
        "chart": lambda df: px.scatter(
            df, x="age", y="bmi", color="risk_flag", size="systolic_bp",
            title="Age vs BMI (color = risk flag, size = systolic BP)"
        ),
        "system_prompt": (
            "You are a clinical decision-support reasoning assistant (for licensed clinician review "
            "only, not a diagnosis). Explain which vitals are driving the risk flag for the given "
            "population, note important caveats, and suggest what a clinician should review next."
        ),
    },
    {
        "id": "legal",
        "title": "Legal Contract Clause Risk Reasoning",
        "description": (
            "Reasons about contract clauses to flag ambiguity and liability exposure the way a "
            "contracts-review associate would triage a redline batch."
        ),
        "fields": [
            ("ambiguity_score", "1-10 analyst-rated linguistic ambiguity",
             "Higher ambiguity increases litigation and misinterpretation risk."),
            ("liability_exposure_usd", "Estimated dollar exposure of the clause",
             "Direct financial materiality of the clause if triggered."),
            ("termination_notice_days", "Contractual notice period",
             "Shorter notice periods reduce a party's ability to react to termination."),
            ("risk_score", "0.5xAmbiguity + 0.002xExposure + 1.2xRatingWeight",
             "Composite clause-risk index blending language risk, dollar exposure, and counterparty credit."),
        ],
        "generate": gen_legal_contract,
        "chart": lambda df: px.box(
            df, x="clause_type", y="risk_score", color="clause_type",
            title="Risk Score Distribution by Clause Type"
        ),
        "system_prompt": (
            "You are a contracts-review reasoning assistant (informational, not legal advice). "
            "Explain which clause types carry the highest risk and why, referencing ambiguity, "
            "liability exposure, and counterparty credit rating, then suggest redline priorities."
        ),
    },
    {
        "id": "fraud",
        "title": "Fraud Detection Reasoning",
        "description": (
            "Reasons about transaction-level signals to explain why a transaction looks suspicious, "
            "beyond a single fraud-score number."
        ),
        "fields": [
            ("txn_velocity_1hr", "Count of transactions by the account in the trailing hour",
             "Unusually high velocity is a classic card-testing / account-takeover signal."),
            ("geo_mismatch_flag", "1 if transaction geo != home geo",
             "Geographic mismatch is a strong indicator of card-not-present fraud."),
            ("device_change_flag", "1 if a new/unrecognized device was used",
             "New-device use often precedes account-takeover fraud."),
            ("fraud_score", "0.15xVelocity + 25xGeo + 20xDevice + 10xNewMerchant + Amount/200",
             "Weighted composite score combining behavioral and contextual fraud signals."),
        ],
        "generate": gen_fraud_detection,
        "chart": lambda df: px.scatter(
            df, x="amount_usd", y="fraud_score", color="fraud_flag",
            title="Transaction Amount vs Fraud Score"
        ),
        "system_prompt": (
            "You are a fraud-operations reasoning analyst. Explain which behavioral signals "
            "(velocity, geo mismatch, device change) are driving flagged transactions, distinguish "
            "likely false positives from true fraud patterns, and recommend an action for the queue."
        ),
    },
    {
        "id": "churn",
        "title": "Customer Churn Reasoning & Retention Strategy",
        "description": (
            "Reasons about usage and satisfaction signals to explain WHY churn probability is high "
            "for a segment, and proposes targeted retention actions."
        ),
        "fields": [
            ("tenure_months", "Months since account start",
             "Newer customers historically churn at higher rates ('early-life churn')."),
            ("monthly_charges_usd", "Current monthly bill",
             "Higher bills relative to perceived value increase churn propensity."),
            ("support_tickets_90d", "Support tickets opened in the last 90 days",
             "Repeated service issues are one of the strongest churn predictors."),
            ("churn_probability_pct", "Logistic function of tenure, charges, tickets, NPS",
             "Probability estimate (0-100%) used to prioritize retention outreach."),
        ],
        "generate": gen_churn,
        "chart": lambda df: px.scatter(
            df, x="tenure_months", y="churn_probability_pct", color="support_tickets_90d",
            title="Tenure vs Churn Probability (color = support tickets)"
        ),
        "system_prompt": (
            "You are a customer-retention reasoning strategist. Explain which drivers (tenure, "
            "charges, support tickets, NPS) are pushing churn probability up for this segment, "
            "and propose 3-5 concrete, prioritized retention actions."
        ),
    },
    {
        "id": "supply_chain",
        "title": "Supply Chain Disruption Reasoning",
        "description": (
            "Reasons about supplier-level signals to explain disruption risk and recommend "
            "mitigation (dual-sourcing, buffer stock, expediting)."
        ),
        "fields": [
            ("on_time_delivery_rate_pct", "% of orders delivered on/before promise date",
             "Direct measure of supplier reliability."),
            ("geopolitical_risk_index", "0-10 composite geopolitical risk rating",
             "Captures exposure to trade, tariff, and political instability risk."),
            ("inventory_days_cover", "Days of inventory on hand at current usage rate",
             "Buffer available to absorb a disruption before stockout."),
            ("disruption_risk_score", "0.5x(100-OTD) + 4xGeoRisk + 15xSingleSource - 0.3xDaysCover",
             "Composite index used to prioritize supplier risk-mitigation efforts."),
        ],
        "generate": gen_supply_chain,
        "chart": lambda df: px.scatter(
            df, x="geopolitical_risk_index", y="disruption_risk_score",
            color="single_source_flag", size="inventory_days_cover",
            title="Geopolitical Risk vs Disruption Score (size = inventory cover)"
        ),
        "system_prompt": (
            "You are a supply-chain risk reasoning analyst. Explain which suppliers carry the "
            "highest disruption risk and why (delivery reliability, geopolitical exposure, "
            "single-sourcing, buffer inventory), and recommend mitigation actions."
        ),
    },
    {
        "id": "credit_underwriting",
        "title": "Credit Underwriting Reasoning",
        "description": (
            "Reasons through loan-application fields to explain an approve / refer decision the "
            "way an underwriter would justify a file note."
        ),
        "fields": [
            ("ltv_ratio_pct", "loan_amount_usd / asset_value_usd x 100",
             "Loan-to-Value; higher LTV means less collateral cushion for the lender."),
            ("dti_ratio_pct", "monthly_debt_usd x 12 / annual_income_usd x 100",
             "Debt-to-Income; measures the borrower's capacity to service new debt."),
            ("credit_score", "300-850 bureau-style score",
             "Reflects historical repayment behavior and default likelihood."),
            ("approval_recommendation", "Approve if LTV<90% AND DTI<43% AND Score>620",
             "Rule-of-thumb underwriting gate the LLM is asked to reason around, including edge cases."),
        ],
        "generate": gen_credit_underwriting,
        "chart": lambda df: px.scatter(
            df, x="dti_ratio_pct", y="ltv_ratio_pct", color="approval_recommendation",
            size="credit_score", title="DTI vs LTV (size = credit score)"
        ),
        "system_prompt": (
            "You are a mortgage/loan underwriting reasoning assistant. Explain why applicants were "
            "approved or referred using LTV, DTI, and credit score, discuss any borderline cases, "
            "and give a portfolio-level recommendation."
        ),
    },
    {
        "id": "hr_fit",
        "title": "HR Candidate Fit Reasoning",
        "description": (
            "Reasons about resume/interview signals to explain a candidate's overall fit score and "
            "surface strengths/gaps for the hiring panel."
        ),
        "fields": [
            ("skill_match_pct", "% overlap between candidate skills and job requirements",
             "Primary technical-fit signal extracted from resume/JD matching."),
            ("culture_fit_score", "0-100 structured-interview culture assessment",
             "Predicts long-term retention and team dynamics fit."),
            ("interview_score", "0-100 panel interview performance score",
             "Captures communication, problem-solving, and role-specific competency."),
            ("overall_fit_score", "0.4xSkill + 0.3xCulture + 0.3xInterview",
             "Weighted composite used to rank candidates for the next round."),
        ],
        "generate": gen_hr_fit,
        "chart": lambda df: px.scatter(
            df, x="years_experience", y="overall_fit_score", color="skill_match_pct",
            title="Experience vs Overall Fit (color = skill match %)"
        ),
        "system_prompt": (
            "You are a talent-acquisition reasoning assistant. Explain what is driving top and "
            "bottom candidates' fit scores across skill match, culture fit, and interview "
            "performance, and recommend which candidates to advance and why."
        ),
    },
    {
        "id": "strategy",
        "title": "Strategic Business Decision Reasoning (SWOT)",
        "description": (
            "Reasons across market and internal-capability signals to prioritize investment across "
            "business units, grounded in a quantified SWOT framework."
        ),
        "fields": [
            ("market_growth_rate_pct", "YoY growth rate of the addressable market",
             "Fast-growing markets reward share-gaining investment."),
            ("market_share_pct", "Business unit's current share of its market",
             "Combined with growth, indicates BCG-matrix style positioning."),
            ("competitive_intensity", "0-10 rating of rivalry intensity",
             "Higher intensity compresses margins and raises execution risk."),
            ("swot_net_score", "(Strengths+Opportunities) - (Weaknesses+Threats)",
             "Simple net positioning score to rank strategic priority across units."),
        ],
        "generate": gen_strategy_swot,
        "chart": lambda df: px.bar(
            df, x="business_unit", y="swot_net_score", color="market_growth_rate_pct",
            title="SWOT Net Score by Business Unit (color = market growth %)"
        ),
        "system_prompt": (
            "You are a corporate-strategy reasoning advisor. Use the SWOT-derived metrics and "
            "market growth/share data to reason about where to invest, where to defend, and where "
            "to divest, referencing the BCG growth-share framework where relevant."
        ),
    },
    {
        "id": "cyber_incident",
        "title": "Cybersecurity Incident Investigation",
        "description": (
            "Connects dozens of alerts/logs across firewalls, EDR, SIEM and IDS to reason about "
            "whether they represent one coordinated attack chain or isolated, unrelated noise."
        ),
        "fields": [
            ("severity_1to5", "Analyst/tool-assigned severity, 1 (info) to 5 (critical)",
             "Baseline weight for how dangerous a single alert is in isolation."),
            ("threat_ip_reputation_score", "0-100 external threat-intel reputation score",
             "Higher scores mean the source IP/domain is known-bad in threat feeds."),
            ("correlated_alert_count", "Count of other alerts sharing the same IOC/time window",
             "The strongest signal of coordination - many alerts sharing indicators suggest a campaign."),
            ("coordination_score", "6xSeverity + 0.3xIPRep + 2.2xCorrelated + 1.1xAffectedHosts",
             "Composite index used to separate a coordinated multi-stage attack from background noise."),
        ],
        "generate": gen_cyber_incident,
        "chart": lambda df: px.scatter(
            df, x="correlated_alert_count", y="coordination_score", color="mitre_tactic",
            size="affected_hosts_count", title="Correlated Alerts vs Coordination Score (by MITRE tactic)"
        ),
        "system_prompt": (
            "You are a SOC (Security Operations Center) incident-response reasoning analyst. "
            "Connect the alerts using severity, threat-intel reputation, correlated-alert counts, "
            "and MITRE ATT&CK tactics to determine whether they form one coordinated attack chain, "
            "identify the likely kill-chain stage, and recommend immediate containment actions."
        ),
    },
    {
        "id": "api_security",
        "title": "API Security Analysis",
        "description": (
            "Examines API call volume, authentication failures and traffic patterns to reason about "
            "the likelihood that an endpoint is under active attack (credential stuffing, scraping, abuse)."
        ),
        "fields": [
            ("auth_failure_rate_pct", "% of requests to the endpoint that fail authentication",
             "Spiking auth failures is the classic signature of credential stuffing / brute force."),
            ("payload_anomaly_score", "0-10 score from schema/size/content deviation from baseline",
             "Malformed or oversized payloads often indicate fuzzing or injection attempts."),
            ("geo_diversity_index", "0-10 score of how many distinct geographies are calling the endpoint",
             "Legitimate traffic is usually geographically clustered; high diversity suggests botnets."),
            ("attack_likelihood_pct", "0.5xAuthFail + 5xPayloadAnomaly + 3xGeoDiversity + RPM/50",
             "Composite likelihood score used to prioritize which endpoints need a closer look."),
        ],
        "generate": gen_api_security,
        "chart": lambda df: px.scatter(
            df, x="auth_failure_rate_pct", y="attack_likelihood_pct", color="api_endpoint",
            size="unique_source_ips", title="Auth Failure Rate vs Attack Likelihood by Endpoint"
        ),
        "system_prompt": (
            "You are an API security reasoning analyst. Explain which endpoints show the strongest "
            "signals of active attack (auth failures, payload anomalies, geographic dispersion, "
            "request volume), distinguish likely bot/credential-stuffing traffic from legitimate "
            "spikes, and recommend rate-limiting or WAF rule changes."
        ),
    },
    {
        "id": "software_troubleshoot",
        "title": "Software Troubleshooting",
        "description": (
            "Analyzes error logs, stack traces and recent code changes to reason about the most "
            "likely root cause of a defect, the way a senior engineer would triage a bug ticket."
        ),
        "fields": [
            ("exceptions_last_24h", "Count of exceptions/errors logged for the module in 24h",
             "Sudden spikes point directly at the module most likely responsible for an incident."),
            ("code_churn_last_7d_commits", "Number of commits touching the module in the last week",
             "High churn immediately before an incident is one of the best root-cause predictors."),
            ("stack_trace_depth", "Depth of the error's stack trace",
             "Very deep or very shallow traces hint at different classes of bugs (integration vs logic)."),
            ("root_cause_likelihood_score", "0.3xExceptions + 0.6xChurn + 0.4xDepth - 0.3xCoverage",
             "Composite score ranking which modules are the most probable root cause."),
        ],
        "generate": gen_software_troubleshoot,
        "chart": lambda df: px.bar(
            df.groupby("module_name", as_index=False)["root_cause_likelihood_score"].mean(),
            x="module_name", y="root_cause_likelihood_score",
            title="Average Root-Cause Likelihood by Module"
        ),
        "system_prompt": (
            "You are a senior software engineer performing root-cause reasoning on bug telemetry. "
            "Correlate exception spikes, recent code churn, stack-trace depth, and test coverage to "
            "identify the most likely root cause module, propose a diagnostic next step, and a fix priority."
        ),
    },
    {
        "id": "medical_decision_support",
        "title": "Medical Decision Support (Differential Generation)",
        "description": (
            "Combines presenting symptoms, vital signs and lab results with medical knowledge to "
            "reason about a ranked set of differential diagnosis possibilities (informational support "
            "only - distinct from the vitals-based chronic-risk triage in the Medical tab above)."
        ),
        "fields": [
            ("primary_symptom_cluster", "Chief complaint grouping reported at intake",
             "Anchors which differential category (infectious, cardiac, neuro, etc.) is most relevant."),
            ("crp_level_mg_L", "C-Reactive Protein, an acute inflammation marker",
             "Elevated CRP supports infectious/inflammatory causes over non-inflammatory ones."),
            ("wbc_count_k_per_uL", "White blood cell count (thousand/uL)",
             "Elevated WBC supports infection; very low WBC can suggest immunosuppression."),
            ("differential_urgency_score", "8x(Fever-36.5) + 2x(WBC-3.5) + 0.25xCRP + 5xComorbidities",
             "Composite urgency index used to help sequence which differential to rule out first."),
        ],
        "generate": gen_medical_decision_support,
        "chart": lambda df: px.scatter(
            df, x="crp_level_mg_L", y="differential_urgency_score", color="primary_symptom_cluster",
            size="wbc_count_k_per_uL", title="CRP vs Differential Urgency (color = symptom cluster)"
        ),
        "system_prompt": (
            "You are a clinical decision-support reasoning assistant (for licensed clinician review "
            "only - not a diagnosis). Combine the symptom cluster, vitals, and labs to reason about "
            "a ranked list of plausible differential diagnoses, note red-flag combinations, and "
            "suggest the next test a clinician should order to narrow the differential."
        ),
    },
    {
        "id": "financial_fraud_investigation",
        "title": "Financial Fraud Pattern Investigation",
        "description": (
            "Examines account networks, geolocation and historical spending baselines to reason about "
            "suspicious patterns across a customer relationship - distinct from the single-transaction "
            "fraud scoring in the Fraud Detection tab above."
        ),
        "fields": [
            ("txn_deviation_from_avg_pct", "% deviation of a transaction from the customer's historical average",
             "Large positive deviations from a stable baseline are a strong anomaly signal."),
            ("distance_from_home_km", "Distance between transaction location and home address",
             "Large distances combined with short time windows can indicate impossible-travel fraud."),
            ("linked_accounts_count", "Number of accounts sharing device/contact identifiers",
             "Fraud rings often operate through clusters of linked accounts."),
            ("network_suspicion_score", "0.2xDeviation + Distance/100 + 15xCrossBorder + 3xLinkedAccts - Age/200",
             "Composite score reasoning across behavior, geography and account-network context."),
        ],
        "generate": gen_financial_fraud_investigation,
        "chart": lambda df: px.scatter(
            df, x="distance_from_home_km", y="network_suspicion_score", color="cross_border_txn_flag",
            size="linked_accounts_count", title="Distance from Home vs Network Suspicion Score"
        ),
        "system_prompt": (
            "You are a financial-crimes investigation reasoning analyst. Reason across account-network, "
            "geolocation, and historical-behavior signals (not just a single transaction) to explain which "
            "customer relationships show organized fraud patterns versus isolated anomalies, and recommend "
            "investigative next steps."
        ),
    },
    {
        "id": "contract_policy_analysis",
        "title": "Contract / Policy Analysis (Portfolio-Level)",
        "description": (
            "Reasons across large sets of contracts/policies to surface conflicts, exceptions, "
            "obligations and consequences - a document-portfolio view, distinct from the single-clause "
            "risk scoring in the Legal Contract tab above."
        ),
        "fields": [
            ("obligation_count", "Number of distinct binding obligations extracted from the document",
             "More obligations mean more surface area for a breach or missed deadline."),
            ("cross_reference_conflicts", "Count of clauses that contradict another clause or linked document",
             "Conflicting obligations are a leading cause of disputes and audit findings."),
            ("penalty_severity_usd", "Estimated dollar penalty/consequence if an obligation is breached",
             "Direct financial materiality of non-compliance."),
            ("conflict_risk_score", "6xConflicts + 2xExceptions + Penalty/20000 + 0.1x(180-DeadlineDays)",
             "Composite score used to prioritize which documents need legal/compliance review first."),
        ],
        "generate": gen_contract_policy_analysis,
        "chart": lambda df: px.box(
            df, x="document_type", y="conflict_risk_score", color="document_type",
            title="Conflict Risk Score Distribution by Document Type"
        ),
        "system_prompt": (
            "You are a compliance/contracts-portfolio reasoning assistant (informational, not legal "
            "advice). Reason across the documents to explain which ones carry the highest risk from "
            "conflicting clauses, exceptions, obligations, and consequences, and recommend a review priority order."
        ),
    },
    {
        "id": "predictive_maintenance",
        "title": "Predictive Maintenance",
        "description": (
            "Reasons from sensor trends, equipment history and failure patterns to forecast the "
            "probability of a FUTURE failure - distinct from the after-the-fact OEE root-cause tab above."
        ),
        "fields": [
            ("vibration_trend_slope_mm_s_per_wk", "Weekly rate of change in vibration amplitude",
             "A rising slope is one of the earliest predictors of bearing/shaft failure."),
            ("temperature_delta_c", "Current operating temperature minus healthy baseline",
             "Sustained positive deltas indicate friction, lubrication, or cooling problems building up."),
            ("historical_failures_last_12mo", "Count of failures on this asset in the trailing 12 months",
             "Assets with a recent failure history are statistically more likely to fail again."),
            ("predicted_failure_probability_pct", "40xVibSlope + 1.2xTempDelta + Hours/150 + 8xPastFailures + 0.1xParticles",
             "Composite forecast used to schedule proactive maintenance before an unplanned failure."),
        ],
        "generate": gen_predictive_maintenance,
        "chart": lambda df: px.scatter(
            df, x="run_hours_since_maintenance", y="predicted_failure_probability_pct",
            color="historical_failures_last_12mo", size="oil_particle_count_ppm",
            title="Run Hours vs Predicted Failure Probability"
        ),
        "system_prompt": (
            "You are a reliability engineer performing predictive-maintenance reasoning. Use vibration "
            "trend, temperature delta, run hours, failure history and oil particle counts to reason about "
            "which assets are most likely to fail soon, estimate a time horizon, and recommend a maintenance schedule."
        ),
    },
    {
        "id": "emergency_response",
        "title": "Emergency Response Sequencing",
        "description": (
            "Evaluates multiple concurrent incidents and resource constraints to reason about the best "
            "sequence of dispatch and response actions."
        ),
        "fields": [
            ("severity_level_1to5", "Incident commander/triage severity rating",
             "Primary driver of response priority - life-safety incidents are weighted highest."),
            ("population_at_risk", "Estimated number of people affected or endangered",
             "Larger at-risk populations raise the urgency of a rapid, well-sequenced response."),
            ("distance_to_nearest_unit_km", "Distance from the nearest available responding unit",
             "Longer distances increase response time and may require re-routing other units."),
            ("response_priority_score", "15xSeverity + 5xln(1+Population) - 0.8xDistance - 0.2xResourceAvail",
             "Composite score used to sequence which incident gets resources dispatched first."),
        ],
        "generate": gen_emergency_response,
        "chart": lambda df: px.scatter(
            df, x="distance_to_nearest_unit_km", y="response_priority_score", color="incident_type",
            size="population_at_risk", title="Distance vs Response Priority by Incident Type"
        ),
        "system_prompt": (
            "You are an emergency-management reasoning assistant coordinating a multi-incident response. "
            "Evaluate severity, population at risk, unit distance and resource availability across all "
            "incidents to reason about the optimal dispatch sequence, and flag any incidents needing mutual-aid support."
        ),
    },
    {
        "id": "aiops",
        "title": "IT Operations / AIOps Root-Cause Correlation",
        "description": (
            "Correlates infrastructure events (errors, latency, deployments, dependency failures) across "
            "services to reason about the root cause of an outage - distinct from the security-focused "
            "cybersecurity-incident tab above."
        ),
        "fields": [
            ("error_rate_pct", "% of requests to the service returning an error",
             "The most direct signal that a service is degraded."),
            ("latency_p95_ms", "95th-percentile response latency in milliseconds",
             "Rising tail latency often precedes full outages as queues back up."),
            ("recent_deployment_flag", "1 if a deployment occurred in the correlation window",
             "A very common outage root cause - deploys are the top suspect until ruled out."),
            ("outage_root_cause_score", "3xErrorRate + Latency/100 + 0.4xCPU + 20xRecentDeploy + 5xDepFailures",
             "Composite score used to rank which service is the likely root cause versus a downstream symptom."),
        ],
        "generate": gen_aiops,
        "chart": lambda df: px.scatter(
            df, x="latency_p95_ms", y="outage_root_cause_score", color="service_name",
            size="dependency_failure_count", title="Latency vs Outage Root-Cause Score by Service"
        ),
        "system_prompt": (
            "You are an AIOps/SRE reasoning assistant correlating infrastructure telemetry across "
            "services. Reason about which service is the true root cause of the outage versus which "
            "services are merely downstream symptoms, weigh the recent-deployment signal heavily, and "
            "recommend a rollback/mitigation action."
        ),
    },
    {
        "id": "autonomous_agents",
        "title": "Autonomous Agents - Task Planning & Tool Selection",
        "description": (
            "Reasons about how an autonomous agent should plan a task, select which tools to call, and "
            "estimate its own probability of successful task completion - an agentic-planning use case "
            "rather than a data-classification one."
        ),
        "fields": [
            ("task_complexity_score_1to10", "Estimated complexity/ambiguity of the task",
             "Higher complexity tasks require more planning steps and carry more failure risk."),
            ("available_tools_count", "Number of tools/functions the agent can call for this task",
             "More applicable tools generally raise the odds of successfully completing the task."),
            ("context_window_used_pct", "% of the agent's context window already consumed",
             "High context usage raises the risk of losing track of earlier plan steps or instructions."),
            ("estimated_success_probability_pct", "100 - 6xComplexity + 2xTools - 0.15xContextUsed - 10xFallback",
             "Self-estimated confidence the agent uses to decide whether to proceed, ask for help, or replan."),
        ],
        "generate": gen_autonomous_agents,
        "chart": lambda df: px.scatter(
            df, x="task_complexity_score_1to10", y="estimated_success_probability_pct",
            color="fallback_triggered_flag", size="available_tools_count",
            title="Task Complexity vs Estimated Success Probability"
        ),
        "system_prompt": (
            "You are an autonomous-agent planning reasoning assistant. For each task, reason step-by-step "
            "about how you would decompose it, which tools you would select and in what order, where a "
            "fallback/human-in-the-loop step is warranted, and give a final confidence estimate for task success."
        ),
    },
]

# --------------------------------------------------------------------------
# GENERIC TAB RENDERER
# --------------------------------------------------------------------------
def render_use_case(cfg: dict):
    st.subheader(cfg['title'])
    st.write(cfg["description"])

    with st.expander("Formulas, Fields & Business Relevance", expanded=False):
        for field, formula, relevance in cfg["fields"]:
            st.markdown(
                f"<div class='field-row'><b>{field}</b><br>"
                f"<code>{formula}</code><br>"
                f"<span style='color:#555;'>{relevance}</span></div>",
                unsafe_allow_html=True,
            )

    st.markdown("#### Step 1: Data Source")
    src = st.radio(
        "Choose data source", ["Generate Synthetic Data", "Upload Real Data (CSV)"],
        key=f"src_{cfg['id']}", horizontal=True,
    )

    df = None
    if src == "Generate Synthetic Data":
        n = st.slider("Number of synthetic records", 20, 500, 100, key=f"n_{cfg['id']}")
        if st.button("Generate Synthetic Data", key=f"gen_{cfg['id']}"):
            st.session_state[f"df_{cfg['id']}"] = cfg["generate"](n)
        if f"df_{cfg['id']}" in st.session_state:
            df = st.session_state[f"df_{cfg['id']}"]
    else:
        expected_cols = ", ".join(f[0] for f in cfg["fields"])
        st.caption(f"Expected schema includes fields such as: *{expected_cols}* (extra columns are fine).")
        uploaded = st.file_uploader("Upload CSV", type=["csv"], key=f"up_{cfg['id']}")
        if uploaded is not None:
            df = pd.read_csv(uploaded)
            st.session_state[f"df_{cfg['id']}"] = df
        elif f"df_{cfg['id']}" in st.session_state:
            df = st.session_state[f"df_{cfg['id']}"]

    if df is None or len(df) == 0:
        st.info("Generate synthetic data or upload a CSV to continue.")
        return

    st.markdown("#### Step 2: Data Preview")
    st.dataframe(df.head(10), use_container_width=True)

    st.markdown("#### Step 3: Visualization")
    try:
        fig = cfg["chart"](df)
        st.plotly_chart(fig, use_container_width=True, key=f"chart_{cfg['id']}")
    except Exception as e:
        st.warning(f"Chart could not be rendered for this dataset (column mismatch?): {e}")

    st.markdown("#### Step 4: AI Reasoning")
    if st.button("Run AI Reasoning", key=f"run_{cfg['id']}"):
        summary = df.describe(include="all").to_string()
        user_prompt = (
            f"Use case: {cfg['title']}\n\n"
            f"Dataset statistical summary:\n{summary}\n\n"
            f"Sample rows (JSON):\n{df.head(8).to_json(orient='records')}\n\n"
            f"{depth_instruction(reasoning_depth)}"
        )
        with st.spinner(f"Reasoning over the data with {llm_provider}..."):
            reasoning = call_llm(cfg["system_prompt"], user_prompt)
        st.session_state[f"reasoning_{cfg['id']}"] = reasoning

    reasoning_text = st.session_state.get(f"reasoning_{cfg['id']}", "")
    if reasoning_text:
        st.markdown("##### Reasoning Output")
        st.markdown(
            f"<div class='section-card'>{reasoning_text}</div>",
            unsafe_allow_html=True,
        )

    st.markdown("#### Step 5: Export Results")
    c1, c2, c3, c4 = st.columns(4)
    reasoning_for_export = reasoning_text or "(No AI reasoning generated yet.)"

    with c1:
        st.download_button(
            "CSV", df_to_csv_bytes(df), file_name=f"{cfg['id']}_data.csv",
            mime="text/csv", key=f"csv_{cfg['id']}", use_container_width=True,
        )
    with c2:
        st.download_button(
            "Text", text_bytes(cfg["title"], cfg["description"], reasoning_for_export),
            file_name=f"{cfg['id']}_report.txt", mime="text/plain",
            key=f"txt_{cfg['id']}", use_container_width=True,
        )
    with c3:
        st.download_button(
            "Word", build_docx(cfg["title"], cfg["description"], df, reasoning_for_export),
            file_name=f"{cfg['id']}_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=f"docx_{cfg['id']}", use_container_width=True,
        )
    with c4:
        st.download_button(
            "PDF", build_pdf(cfg["title"], cfg["description"], df, reasoning_for_export),
            file_name=f"{cfg['id']}_report.pdf", mime="application/pdf",
            key=f"pdf_{cfg['id']}", use_container_width=True,
        )

# --------------------------------------------------------------------------
# MAIN TABS
# --------------------------------------------------------------------------
tab_labels = [c['title'] for c in USE_CASES]
tabs = st.tabs(tab_labels)
for tab, cfg in zip(tabs, USE_CASES):
    with tab:
        render_use_case(cfg)

st.markdown("<hr>", unsafe_allow_html=True)
st.caption(
    "CogniReason AI Suite - 20 embedded Reasoning AI use cases. "
    "Reasoning is generated live via your choice of the Groq or Google Gemini LLM API. "
    "Synthetic data is generated locally; uploaded data never leaves this session. "
    "Developed by Randy Singh, Kalsnet (KNet) Consulting Group.")
