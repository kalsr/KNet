


#CrossGuard AI
#A Streamlit demonstration application for agentic, AI-assisted cross-domain
#data filtering (Low -> High and High -> Low transfer simulation), with a
#two-stage filter engine (deterministic rule-based scan + LLM "agent" review
#pass using Claude or Groq), synthetic data generation for demos, and
#multi-format export (PDF, Word, TXT, JSON).

#Developed by Randy Singh | Kalsnet (KNet) Consulting

#IMPORTANT: This is an educational / business-document redaction demonstration
#tool. It is NOT an accredited Cross-Domain Solution (CDS) and must not be used
#to move real classified national-security information between security
#domains. Real CDS deployments require hardware-enforced guards and formal
#accreditation (e.g. via the appropriate national CDS authority).


import io
import json
import random
import re
import datetime as dt

import pandas as pd
import requests
import streamlit as st

# =========================================================================
# PAGE CONFIG
# =========================================================================
st.set_page_config(
    page_title="CrossDomainGuard AI | Cross-Domain Data Filter",
    page_icon="",
    layout="wide",
    initial_sidebar_state="expanded",
)

# =========================================================================
# GLOBAL STYLE
# =========================================================================
st.markdown(
    """
    <style>
    .block-container {padding-top: 1.5rem;}
    .title-line-1 {
        text-align: center;
        color: #0047AB;
        font-weight: 800;
        font-size: 2.2rem;
        margin-bottom: 0;
        letter-spacing: 0.5px;
    }
    .title-line-2 {
        text-align: center;
        color: #0047AB;
        font-weight: 700;
        font-size: 1.05rem;
        margin-top: 0;
    }
    .subtle-divider {
        border-top: 2px solid #0047AB;
        opacity: 0.25;
        margin: 0.6rem 0 1.2rem 0;
    }
    .status-pill {
        display:inline-block; padding:3px 10px; border-radius:12px;
        font-size:0.78rem; font-weight:700; color:white;
    }
    .pill-clean {background:#1a8a3e;}
    .pill-flag {background:#c0392b;}
    .pill-info {background:#0047AB;}
    div[data-testid="stMetricValue"] {color:#0047AB;}
    </style>
    """,
    unsafe_allow_html=True,
)

# =========================================================================
# TITLE BAR
# =========================================================================
st.markdown(
    """
    <div class="title-line-1"> CrossGuard AI — Cross-Domain Data Filter</div>
    <div class="title-line-2">Developed by Randy Singh from Kalsnet (KNet) Consulting</div>
    <div class="subtle-divider"></div>
    """,
    unsafe_allow_html=True,
)
st.caption(
    "Demonstration / educational tool for AI-assisted sensitive-data filtering. "
    "Not an accredited Cross-Domain Solution — do not use with real classified data."
)

# =========================================================================
# CONSTANTS
# =========================================================================
CLASSIFICATION_LEVELS = ["UNCLASSIFIED", "CONFIDENTIAL", "SECRET", "TOP SECRET"]
CLASSIFICATION_MARKERS = ["TOP SECRET", "SECRET", "CONFIDENTIAL", "RESTRICTED", "NOFORN", "FOUO", "SCI"]

FIRST_NAMES = ["James", "Maria", "Wei", "Fatima", "Liam", "Sara", "Carlos", "Aisha",
               "Noah", "Elena", "Raj", "Olivia", "Hassan", "Mei", "Lucas", "Amara"]
LAST_NAMES = ["Smith", "Johnson", "Chen", "Khan", "Garcia", "Singh", "Müller", "Rossi",
              "Tanaka", "Brown", "Patel", "Kim", "Nguyen", "Ibrahim", "Lopez", "Walker"]
CITIES = ["Ottawa", "Toronto", "Calgary", "Vancouver", "Halifax", "Winnipeg", "Quebec City", "Edmonton"]
DEPARTMENTS = ["Defence Logistics", "Cyber Operations", "Finance & Audit", "Diplomatic Affairs",
               "Border Security", "Health Intelligence", "Procurement", "Critical Infrastructure"]

NOTE_TEMPLATES = [
    "Subject {name} (SSN {ssn}) was reviewed under file ref {ref}. Contact: {email}, {phone}.",
    "Routine status update for {name}. Address on file: {address}, {city}. Marked {marker}.",
    "Procurement note: vendor contacted via {email}. Internal classification: {marker}.",
    "Personnel record {name} updated. Phone {phone}. No additional remarks.",
    "Incident summary regarding {name} ({ssn}). Distribution limited - {marker}.",
    "General correspondence with {email} regarding logistics in {city}. Unremarkable.",
    "Audit trail entry for case {ref}. No personal data referenced. Status: cleared.",
    "Financial reconciliation note. Card on file ending pattern: {cc}. Marked {marker}.",
]

PII_PATTERNS = {
    "SSN": r"\b\d{3}-\d{2}-\d{4}\b",
    "EMAIL": r"\b[\w.\-]+@[\w.\-]+\.\w+\b",
    "PHONE": r"\b\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b",
    "CREDIT_CARD": r"\b(?:\d[ -]?){13,16}\b",
}

CLAUDE_MODELS = ["claude-sonnet-4-6", "claude-opus-4-7", "claude-haiku-4-5-20251001"]
GROQ_MODELS = ["llama-3.3-70b-versatile", "llama-3.1-8b-instant", "mixtral-8x7b-32768",
               "gemma2-9b-it", "custom (type below)"]

# =========================================================================
# SESSION STATE DEFAULTS
# =========================================================================
DEFAULTS = {
    "synthetic_df": None,
    "synthetic_filtered_df": None,
    "synthetic_findings": [],
    "real_raw": None,
    "real_kind": None,        # "table" or "text"
    "real_filtered": None,
    "real_findings": [],
    "audit_log": [],
}
for k, v in DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = v


def log_audit(action, direction, records, findings_count, engine):
    st.session_state.audit_log.append(
        {
            "timestamp": dt.datetime.now().isoformat(timespec="seconds"),
            "action": action,
            "direction": direction,
            "records_processed": records,
            "findings_redacted": findings_count,
            "engine": engine,
        }
    )


# =========================================================================
# SIDEBAR — LLM CONFIGURATION
# =========================================================================
with st.sidebar:
    st.header(" LLM Configuration")

    provider = st.radio("LLM Provider", ["Claude (Anthropic)", "Groq"], index=0)

    if provider == "Claude (Anthropic)":
        model = st.selectbox("Claude model", CLAUDE_MODELS, index=0)
    else:
        model = st.selectbox("Groq model", GROQ_MODELS, index=0)
        if model == "custom (type below)":
            model = st.text_input("Custom Groq model name", value="llama-3.3-70b-versatile")

    api_key = st.text_input(
        f"{provider.split(' ')[0]} API Key",
        type="password",
        placeholder="Paste your API key here",
        help="Your key is kept only in this browser session's memory — it is never written to disk by this app.",
    )

    use_agent_review = st.checkbox(
        "Enable AI agent review pass (recommended)",
        value=True,
        help="After the deterministic rule-based filter runs, send a sample of records to the selected "
             "LLM for a contextual second-opinion pass that can catch leaks regex misses (e.g. names "
             "embedded in free text).",
    )

    with st.expander(" How do I get an API key?"):
        st.markdown(
            """
**Claude (Anthropic) API key**
1. Go to **console.anthropic.com** and sign up or log in.
2. Open **Settings → API Keys** in the left navigation.
3. Click **Create Key**, name it, and copy the value immediately
   (it is only shown once).
4. Add billing/credits under **Settings → Billing** if prompted —
   new accounts usually get a small free credit grant.
5. Paste the key into the field above.

---

**Groq API key**
1. Go to **console.groq.com** and sign up or log in (Groq currently
   offers a generous free tier for API access).
2. Open **API Keys** in the left navigation.
3. Click **Create API Key**, name it, and copy the value shown.
4. Paste the key into the field above.

---
⚠️ Treat both keys like passwords. Do not commit them to source control,
and rotate/revoke them from the respective console if you suspect exposure.
            """
        )

    st.markdown("---")
    st.caption("CrossGuard AI v1.0 — demo build")

PROVIDER_KEY = "Claude" if provider.startswith("Claude") else "Groq"


# =========================================================================
# LLM CALL HELPERS
# =========================================================================
def call_claude(key, model_name, system, user_msg):
    headers = {
        "x-api-key": key,
        "anthropic-version": "2023-06-01",
        "content-type": "application/json",
    }
    body = {
        "model": model_name,
        "max_tokens": 1024,
        "system": system,
        "messages": [{"role": "user", "content": user_msg}],
    }
    r = requests.post("https://api.anthropic.com/v1/messages", headers=headers, json=body, timeout=60)
    r.raise_for_status()
    data = r.json()
    return "".join(b.get("text", "") for b in data.get("content", []) if b.get("type") == "text")


def call_groq(key, model_name, system, user_msg):
    headers = {"Authorization": f"Bearer {key}", "Content-Type": "application/json"}
    body = {
        "model": model_name,
        "messages": [{"role": "system", "content": system}, {"role": "user", "content": user_msg}],
        "max_tokens": 1024,
        "temperature": 0,
    }
    r = requests.post("https://api.groq.com/openai/v1/chat/completions", headers=headers, json=body, timeout=60)
    r.raise_for_status()
    data = r.json()
    return data["choices"][0]["message"]["content"]


def call_llm(key, provider_key, model_name, system, user_msg):
    if provider_key == "Claude":
        return call_claude(key, model_name, system, user_msg)
    return call_groq(key, model_name, system, user_msg)


AGENT_SYSTEM_PROMPT = (
    "You are CrossGuard, a cross-domain security review agent. You are given a piece of text that has "
    "already passed through a deterministic redaction filter. Your job is to find ANYTHING that filter "
    "missed: personal names, addresses, government/military classification markers, identifiers, account "
    "numbers, or anything that looks like personally identifiable or classified information. "
    "Respond with STRICT JSON ONLY, no markdown fences, no commentary, in this exact shape: "
    '{"clean": true|false, "additional_findings": ["short description", ...], '
    '"final_redacted_text": "the text with any further necessary redactions applied as [REDACTED-...] tags"}'
)


def agent_review(key, provider_key, model_name, original_text, rule_redacted_text):
    user_msg = (
        f"ORIGINAL TEXT:\n{original_text}\n\n"
        f"RULE-FILTERED TEXT:\n{rule_redacted_text}\n\n"
        "Review the rule-filtered text for anything that should still be redacted. "
        "Return the JSON object described in your instructions."
    )
    raw = call_llm(key, provider_key, model_name, AGENT_SYSTEM_PROMPT, user_msg)
    cleaned = raw.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(json)?", "", cleaned).strip()
        cleaned = re.sub(r"```$", "", cleaned).strip()
    try:
        data = json.loads(cleaned)
        data.setdefault("clean", True)
        data.setdefault("additional_findings", [])
        data.setdefault("final_redacted_text", rule_redacted_text)
        return data
    except Exception:
        return {
            "clean": True,
            "additional_findings": [],
            "final_redacted_text": rule_redacted_text,
            "raw_response": raw,
        }


# =========================================================================
# RULE-BASED FILTER ENGINE
# =========================================================================
def rule_based_redact(text):
    if text is None:
        return text, []
    text = str(text)
    findings = []
    redacted = text
    for label, pattern in PII_PATTERNS.items():
        matches = re.findall(pattern, redacted)
        if matches:
            findings.append({"type": label, "count": len(matches)})
            redacted = re.sub(pattern, f"[REDACTED-{label}]", redacted)
    for marker in CLASSIFICATION_MARKERS:
        count = len(re.findall(re.escape(marker), redacted, flags=re.IGNORECASE))
        if count:
            findings.append({"type": f"CLASSIFICATION:{marker}", "count": count})
            redacted = re.sub(re.escape(marker), "[CLASSIFICATION-REDACTED]", redacted, flags=re.IGNORECASE)
    return redacted, findings


def redact_dataframe(df, columns_to_scan=None):
    """Run the rule-based filter over every cell of the given (or all string) columns."""
    df = df.copy()
    all_findings = []
    cols = columns_to_scan or [c for c in df.columns if df[c].dtype == object]
    for col in cols:
        new_vals = []
        for val in df[col]:
            red, finds = rule_based_redact(val)
            new_vals.append(red)
            all_findings.extend(finds)
        df[col] = new_vals
    return df, all_findings


def run_agent_pass_on_dataframe(df, columns_to_scan, key, provider_key, model_name, sample_size):
    """Run the LLM agent review on a sample of rows for the given text columns."""
    df = df.copy()
    extra_findings = []
    rows_to_check = df.index[:sample_size]
    for col in columns_to_scan:
        if col not in df.columns:
            continue
        for idx in rows_to_check:
            original = str(df.at[idx, col])
            rule_red, _ = rule_based_redact(original)
            try:
                result = agent_review(key, provider_key, model_name, original, rule_red)
                df.at[idx, col] = result.get("final_redacted_text", rule_red)
                if result.get("additional_findings"):
                    extra_findings.append(
                        {"row": idx, "column": col, "findings": result["additional_findings"]}
                    )
            except Exception as e:
                st.warning(f"Agent review skipped for row {idx}, column '{col}': {e}")
    return df, extra_findings


# =========================================================================
# SYNTHETIC DATA GENERATOR
# =========================================================================
def generate_synthetic_dataset(n, direction):
    rows = []
    for i in range(n):
        name = f"{random.choice(FIRST_NAMES)} {random.choice(LAST_NAMES)}"
        ssn = f"{random.randint(100,999)}-{random.randint(10,99)}-{random.randint(1000,9999)}"
        email = f"{name.split()[0].lower()}.{name.split()[1].lower()}@example.com"
        phone = f"({random.randint(200,999)}) {random.randint(200,999)}-{random.randint(1000,9999)}"
        address = f"{random.randint(10,9999)} {random.choice(['Main St','Elm Ave','King St','Oak Blvd'])}"
        city = random.choice(CITIES)
        cc = " ".join(str(random.randint(1000, 9999)) for _ in range(4))
        marker = random.choice(CLASSIFICATION_MARKERS)
        ref = f"REF-{random.randint(10000,99999)}"
        note = random.choice(NOTE_TEMPLATES).format(
            name=name, ssn=ssn, email=email, phone=phone, address=address,
            city=city, cc=cc, marker=marker, ref=ref,
        )
        rows.append(
            {
                "record_id": i + 1,
                "full_name": name,
                "ssn": ssn,
                "email": email,
                "phone": phone,
                "address": f"{address}, {city}",
                "department": random.choice(DEPARTMENTS),
                "classification": random.choice(CLASSIFICATION_LEVELS),
                "notes": note,
                "transfer_direction": direction,
            }
        )
    return pd.DataFrame(rows)


# =========================================================================
# EXPORT HELPERS
# =========================================================================
def _safe_latin1(text):
    return str(text).encode("latin-1", "replace").decode("latin-1")


def df_to_json_bytes(df):
    return df.to_json(orient="records", indent=2).encode("utf-8")


def df_to_text_bytes(df, title):
    buf = io.StringIO()
    buf.write(f"{title}\n")
    buf.write("=" * len(title) + "\n\n")
    buf.write(df.to_string(index=False))
    return buf.getvalue().encode("utf-8")


def df_to_docx_bytes(df, title):
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Generated: {dt.datetime.now().isoformat(timespec='seconds')}")
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            cells[i].text = str(row[col])
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue()


def df_to_pdf_bytes(df, title):
    from fpdf import FPDF

    pdf = FPDF(orientation="L", unit="mm", format="A4")
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 10, _safe_latin1(title), ln=True)
    pdf.set_font("Helvetica", "", 7)

    n_cols = max(len(df.columns), 1)
    col_width = 270 / n_cols

    pdf.set_font("Helvetica", "B", 7)
    for col in df.columns:
        pdf.cell(col_width, 7, _safe_latin1(str(col))[:22], border=1)
    pdf.ln()
    pdf.set_font("Helvetica", "", 7)
    for _, row in df.iterrows():
        for col in df.columns:
            pdf.cell(col_width, 7, _safe_latin1(str(row[col]))[:28], border=1)
        pdf.ln()
    out = pdf.output(dest="S")
    if isinstance(out, str):
        out = out.encode("latin-1", "replace")
    return bytes(out)


def text_to_docx_bytes(text, title):
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Generated: {dt.datetime.now().isoformat(timespec='seconds')}")
    for line in text.splitlines():
        doc.add_paragraph(line)
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue()


def text_to_pdf_bytes(text, title):
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 10, _safe_latin1(title), ln=True)
    pdf.set_font("Helvetica", "", 10)
    for line in text.splitlines() or [""]:
        pdf.multi_cell(0, 6, _safe_latin1(line))
    out = pdf.output(dest="S")
    if isinstance(out, str):
        out = out.encode("latin-1", "replace")
    return bytes(out)


def export_block(data, base_filename, title, key_prefix):
    """Render the 4 download buttons (PDF / DOCX / TXT / JSON) for a DataFrame or raw text."""
    is_df = isinstance(data, pd.DataFrame)
    c1, c2, c3, c4 = st.columns(4)

    with c1:
        try:
            payload = df_to_pdf_bytes(data, title) if is_df else text_to_pdf_bytes(data, title)
            st.download_button("⬇️ PDF", payload, file_name=f"{base_filename}.pdf",
                                mime="application/pdf", key=f"{key_prefix}_pdf")
        except Exception as e:
            st.error(f"PDF export error: {e}")

    with c2:
        try:
            payload = df_to_docx_bytes(data, title) if is_df else text_to_docx_bytes(data, title)
            st.download_button("⬇️ Word", payload, file_name=f"{base_filename}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"{key_prefix}_docx")
        except Exception as e:
            st.error(f"Word export error: {e}")

    with c3:
        try:
            payload = df_to_text_bytes(data, title) if is_df else data.encode("utf-8")
            st.download_button("⬇️ Text", payload, file_name=f"{base_filename}.txt",
                                mime="text/plain", key=f"{key_prefix}_txt")
        except Exception as e:
            st.error(f"Text export error: {e}")

    with c4:
        try:
            payload = df_to_json_bytes(data) if is_df else json.dumps({"text": data}, indent=2).encode("utf-8")
            st.download_button("⬇️ JSON", payload, file_name=f"{base_filename}.json",
                                mime="application/json", key=f"{key_prefix}_json")
        except Exception as e:
            st.error(f"JSON export error: {e}")


# =========================================================================
# MAIN TABS
# =========================================================================
tab_synth, tab_real, tab_audit = st.tabs(
    [" Synthetic Data Demo", " Real Data Upload", " Audit Log"]
)

# -------------------------------------------------------------------------
# TAB 1: SYNTHETIC DATA DEMO
# -------------------------------------------------------------------------
with tab_synth:
    st.subheader("Synthetic Data Generator & Filter Demonstration")
    st.write(
        "Generate fake-but-realistic records (with embedded PII and classification markers) "
        "to demonstrate the filter without touching any real data."
    )

    col_a, col_b, col_c = st.columns([3, 2, 1])
    with col_a:
        n_records = st.slider("Number of synthetic records", min_value=0, max_value=300, value=50, step=1)
    with col_b:
        synth_direction = st.selectbox(
            "Transfer direction", ["Low → High (Upload)", "High → Low (Downgrade/Export)"], key="synth_dir"
        )
    with col_c:
        st.write("")
        st.write("")
        gen_clicked = st.button(" Generate", use_container_width=True)

    reset_col, _ = st.columns([1, 5])
    with reset_col:
        if st.button(" Reset Data", use_container_width=True):
            st.session_state.synthetic_df = None
            st.session_state.synthetic_filtered_df = None
            st.session_state.synthetic_findings = []
            st.rerun()

    if gen_clicked:
        if n_records == 0:
            st.session_state.synthetic_df = pd.DataFrame()
            st.session_state.synthetic_filtered_df = None
            st.info("0 records selected — nothing to generate.")
        else:
            st.session_state.synthetic_df = generate_synthetic_dataset(n_records, synth_direction)
            st.session_state.synthetic_filtered_df = None
            st.success(f"Generated {n_records} synthetic records.")

    if st.session_state.synthetic_df is not None and not st.session_state.synthetic_df.empty:
        st.markdown("**Raw synthetic data (unfiltered):**")
        st.dataframe(st.session_state.synthetic_df, use_container_width=True, height=260)

        run_filter = st.button(" Run Cross-Domain Filter", type="primary")
        if run_filter:
            df = st.session_state.synthetic_df
            with st.spinner("Running deterministic rule-based filter..."):
                filtered_df, findings = redact_dataframe(df, columns_to_scan=["notes", "address", "email", "ssn", "phone"])

            engine_used = "Rule-based only"
            extra_findings = []
            if use_agent_review:
                if not api_key:
                    st.warning("No API key provided — skipping AI agent review pass (rule-based filter still applied).")
                else:
                    sample_n = min(20, len(filtered_df))
                    with st.spinner(f"Running AI agent review pass on {sample_n} sample record(s) via {provider}..."):
                        try:
                            filtered_df, extra_findings = run_agent_pass_on_dataframe(
                                filtered_df, ["notes"], api_key, PROVIDER_KEY, model, sample_n
                            )
                            engine_used = f"Rule-based + {provider} agent review"
                        except Exception as e:
                            st.error(f"Agent review pass failed: {e}")

            st.session_state.synthetic_filtered_df = filtered_df
            st.session_state.synthetic_findings = findings + extra_findings
            log_audit("Synthetic data filter", synth_direction, len(df), len(findings) + len(extra_findings), engine_used)
            st.success("Filter complete.")

    if st.session_state.synthetic_filtered_df is not None:
        st.markdown("**Filtered output:**")
        n_flags = len(st.session_state.synthetic_findings)
        pill = "pill-flag" if n_flags else "pill-clean"
        label = f"{n_flags} item(s) redacted" if n_flags else "No sensitive data detected"
        st.markdown(f'<span class="status-pill {pill}">{label}</span>', unsafe_allow_html=True)
        st.dataframe(st.session_state.synthetic_filtered_df, use_container_width=True, height=260)

        if st.session_state.synthetic_findings:
            with st.expander(" Findings detail"):
                st.json(st.session_state.synthetic_findings)

        st.markdown("**Download filtered data:**")
        export_block(
            st.session_state.synthetic_filtered_df,
            "crossguard_synthetic_filtered",
            "CrossGuard AI — Synthetic Filtered Data",
            "synth",
        )

# -------------------------------------------------------------------------
# TAB 2: REAL DATA UPLOAD
# -------------------------------------------------------------------------
with tab_real:
    st.subheader("Real Data Upload & Filter")
    st.warning(
        "⚠️ This demo tool is not an accredited Cross-Domain Solution. Do not upload real classified, "
        "regulated, or production personal data here. Use synthetic/test data only when evaluating this app."
    )

    real_direction = st.selectbox(
        "Transfer direction", ["Low → High (Upload)", "High → Low (Downgrade/Export)"], key="real_dir"
    )

    uploaded = st.file_uploader(
        "Upload a file to filter", type=["csv", "json", "txt", "xlsx"], key="real_uploader"
    )

    reset_col2, _ = st.columns([1, 5])
    with reset_col2:
        if st.button(" Reset Real Data", use_container_width=True):
            st.session_state.real_raw = None
            st.session_state.real_kind = None
            st.session_state.real_filtered = None
            st.session_state.real_findings = []
            st.rerun()

    if uploaded is not None:
        try:
            if uploaded.name.endswith(".csv"):
                st.session_state.real_raw = pd.read_csv(uploaded)
                st.session_state.real_kind = "table"
            elif uploaded.name.endswith(".xlsx"):
                st.session_state.real_raw = pd.read_excel(uploaded)
                st.session_state.real_kind = "table"
            elif uploaded.name.endswith(".json"):
                content = json.load(uploaded)
                st.session_state.real_raw = pd.json_normalize(content) if isinstance(content, list) else pd.json_normalize([content])
                st.session_state.real_kind = "table"
            else:  # .txt
                st.session_state.real_raw = uploaded.read().decode("utf-8", errors="replace")
                st.session_state.real_kind = "text"
            st.session_state.real_filtered = None
            st.success(f"Loaded '{uploaded.name}'.")
        except Exception as e:
            st.error(f"Could not read file: {e}")

    if st.session_state.real_raw is not None:
        st.markdown("**Raw uploaded data (unfiltered):**")
        if st.session_state.real_kind == "table":
            st.dataframe(st.session_state.real_raw, use_container_width=True, height=260)
        else:
            st.text_area("Raw text", st.session_state.real_raw, height=200, disabled=True)

        if st.button(" Run Cross-Domain Filter", type="primary", key="real_filter_btn"):
            if st.session_state.real_kind == "table":
                df = st.session_state.real_raw
                with st.spinner("Running deterministic rule-based filter..."):
                    filtered, findings = redact_dataframe(df)

                engine_used = "Rule-based only"
                extra_findings = []
                if use_agent_review:
                    if not api_key:
                        st.warning("No API key provided — skipping AI agent review pass.")
                    else:
                        text_cols = [c for c in filtered.columns if filtered[c].dtype == object]
                        sample_n = min(20, len(filtered))
                        with st.spinner(f"Running AI agent review pass via {provider}..."):
                            try:
                                filtered, extra_findings = run_agent_pass_on_dataframe(
                                    filtered, text_cols, api_key, PROVIDER_KEY, model, sample_n
                                )
                                engine_used = f"Rule-based + {provider} agent review"
                            except Exception as e:
                                st.error(f"Agent review pass failed: {e}")

                st.session_state.real_filtered = filtered
                st.session_state.real_findings = findings + extra_findings
                log_audit("Real data filter (table)", real_direction, len(df), len(findings) + len(extra_findings), engine_used)

            else:
                text = st.session_state.real_raw
                with st.spinner("Running deterministic rule-based filter..."):
                    rule_red, findings = rule_based_redact(text)

                final_text = rule_red
                extra_findings = []
                engine_used = "Rule-based only"
                if use_agent_review:
                    if not api_key:
                        st.warning("No API key provided — skipping AI agent review pass.")
                    else:
                        with st.spinner(f"Running AI agent review pass via {provider}..."):
                            try:
                                result = agent_review(api_key, PROVIDER_KEY, model, text, rule_red)
                                final_text = result.get("final_redacted_text", rule_red)
                                extra_findings = result.get("additional_findings", [])
                                engine_used = f"Rule-based + {provider} agent review"
                            except Exception as e:
                                st.error(f"Agent review pass failed: {e}")

                st.session_state.real_filtered = final_text
                st.session_state.real_findings = findings + [{"agent_findings": f} for f in extra_findings]
                log_audit("Real data filter (text)", real_direction, 1, len(st.session_state.real_findings), engine_used)

            st.success("Filter complete.")

    if st.session_state.real_filtered is not None:
        st.markdown("**Filtered output:**")
        n_flags = len(st.session_state.real_findings)
        pill = "pill-flag" if n_flags else "pill-clean"
        label = f"{n_flags} item(s) redacted" if n_flags else "No sensitive data detected"
        st.markdown(f'<span class="status-pill {pill}">{label}</span>', unsafe_allow_html=True)

        if isinstance(st.session_state.real_filtered, pd.DataFrame):
            st.dataframe(st.session_state.real_filtered, use_container_width=True, height=260)
        else:
            st.text_area("Filtered text", st.session_state.real_filtered, height=200, disabled=True)

        if st.session_state.real_findings:
            with st.expander(" Findings detail"):
                st.json(st.session_state.real_findings)

        st.markdown("**Download filtered data:**")
        export_block(
            st.session_state.real_filtered,
            "crossguard_real_filtered",
            "CrossGuard AI — Real Data Filtered Output",
            "real",
        )

# -------------------------------------------------------------------------
# TAB 3: AUDIT LOG
# -------------------------------------------------------------------------
with tab_audit:
    st.subheader("Filter Run Audit Log")
    if st.session_state.audit_log:
        audit_df = pd.DataFrame(st.session_state.audit_log)
        st.dataframe(audit_df, use_container_width=True, height=300)
        st.download_button(
            "⬇️ Download Audit Log (CSV)",
            audit_df.to_csv(index=False).encode("utf-8"),
            file_name="crossguard_audit_log.csv",
            mime="text/csv",
        )
        if st.button(" Clear Audit Log"):
            st.session_state.audit_log = []
            st.rerun()
    else:
        st.info("No filter runs logged yet. Run a filter in the Synthetic or Real Data tabs to populate this log.")

st.markdown("---")
st.caption("CrossGuard AI — demonstration build · Randy Singh · Kalsnet (KNet) Consulting")
