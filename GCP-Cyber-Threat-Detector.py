


# GCP Serverless Cyber Threat Detector
# Developed by Randy Singh | Kalsnet (KNet) Consulting

# A Streamlit demo application illustrating four serverless, GCP-style
# cyber-threat-detection patterns. Each tab is self-contained: it explains
# the use case and data schema, lets you generate synthetic data (or upload
# your own CSV), runs a detection function, visualizes results, and lets you
# export the findings to PDF, Word, TXT, or CSV.


import io
from datetime import datetime, timedelta

import numpy as np
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
)

try:
    from sklearn.ensemble import IsolationForest
    SKLEARN_OK = True
except Exception:
    SKLEARN_OK = False

# --------------------------------------------------------------------------
# PAGE CONFIG + TITLE BAR
# --------------------------------------------------------------------------
st.set_page_config(
    page_title="GCP Serverless Cyber Threat Detector",
    page_icon="🛡️",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.markdown(
    """
    <style>
    .title-block {
        padding: 18px 24px 14px 24px;
        border-radius: 10px;
        background: linear-gradient(90deg, #eaf1ff 0%, #f5f9ff 100%);
        border: 1px solid #c9dbff;
        margin-bottom: 18px;
    }
    .title-main {
        color: #0B3D91;
        font-weight: 800;
        font-size: 40px;
        letter-spacing: 0.3px;
        margin: 0;
        line-height: 1.15;
    }
    .title-sub {
        color: #0B3D91;
        font-weight: 800;
        font-size: 22px;
        margin-top: 6px;
        line-height: 1.2;
    }
    .schema-box {
        background-color: #f7f9fc;
        border-left: 4px solid #0B3D91;
        padding: 10px 14px;
        border-radius: 4px;
        margin-bottom: 10px;
    }
    .flag-high { background-color: #ffd6d6 !important; }
    div.stButton > button, div.stDownloadButton > button {
        border-radius: 6px;
        font-weight: 600;
    }
    </style>
    <div class="title-block">
        <p class="title-main">🛡️ GCP Serverless Cyber Threat Detector</p>
        <p class="title-sub">Developed by Randy Singh — Kalsnet (KNet) Consulting</p>
    </div>
    """,
    unsafe_allow_html=True,
)

st.caption(
    "Educational demo of serverless threat-detection patterns on Google Cloud "
    "(Pub/Sub → Cloud Functions/Cloud Run → BigQuery/Vertex AI). "
    "Use the synthetic-data buttons to generate demo traffic, or upload your own CSV in each tab."
)

# --------------------------------------------------------------------------
# EXPORT HELPERS (PDF / DOCX / TXT / CSV)
# --------------------------------------------------------------------------

def df_to_csv_bytes(df: pd.DataFrame) -> bytes:
    return df.to_csv(index=False).encode("utf-8")


def df_to_txt_bytes(df: pd.DataFrame, title: str, summary: str = "") -> bytes:
    buf = io.StringIO()
    buf.write(f"{title}\n")
    buf.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
    buf.write("Developed by Randy Singh | Kalsnet (KNet) Consulting\n")
    buf.write("=" * 70 + "\n\n")
    if summary:
        buf.write(summary + "\n\n")
    buf.write(df.to_string(index=False))
    return buf.getvalue().encode("utf-8")


def df_to_pdf_bytes(df: pd.DataFrame, title: str, summary: str = "") -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=landscape(letter),
        topMargin=0.5 * inch, bottomMargin=0.5 * inch,
        leftMargin=0.5 * inch, rightMargin=0.5 * inch,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleBlue", parent=styles["Title"], textColor=colors.HexColor("#0B3D91"),
        fontSize=18, spaceAfter=4,
    )
    sub_style = ParagraphStyle(
        "SubBlue", parent=styles["Normal"], textColor=colors.HexColor("#0B3D91"),
        fontSize=11, spaceAfter=10, fontName="Helvetica-Bold",
    )
    body_style = styles["Normal"]

    elements = [
        Paragraph(title, title_style),
        Paragraph("Developed by Randy Singh | Kalsnet (KNet) Consulting", sub_style),
        Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", body_style),
        Spacer(1, 10),
    ]
    if summary:
        elements.append(Paragraph(summary.replace("\n", "<br/>"), body_style))
        elements.append(Spacer(1, 12))

    # Limit columns/rows so the table fits reasonably on the page
    max_rows = 200
    show_df = df.head(max_rows)
    data = [list(show_df.columns)] + show_df.astype(str).values.tolist()
    table = Table(data, repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0B3D91")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 7),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f2f6ff")]),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    elements.append(table)
    if len(df) > max_rows:
        elements.append(Spacer(1, 8))
        elements.append(Paragraph(f"Showing first {max_rows} of {len(df)} rows.", body_style))
    doc.build(elements)
    return buf.getvalue()


def df_to_docx_bytes(df: pd.DataFrame, title: str, summary: str = "") -> bytes:
    doc = Document()

    h = doc.add_heading(title, level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0B, 0x3D, 0x91)

    sub = doc.add_paragraph()
    sub_run = sub.add_run("Developed by Randy Singh | Kalsnet (KNet) Consulting")
    sub_run.bold = True
    sub_run.font.color.rgb = RGBColor(0x0B, 0x3D, 0x91)
    sub_run.font.size = Pt(13)

    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    if summary:
        doc.add_paragraph(summary)

    max_rows = 300
    show_df = df.head(max_rows)
    table = doc.add_table(rows=1, cols=len(show_df.columns))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(show_df.columns):
        hdr_cells[i].text = str(col)
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True

    for _, row in show_df.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)

    if len(df) > max_rows:
        doc.add_paragraph(f"Showing first {max_rows} of {len(df)} rows.")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_bar(df: pd.DataFrame, key_prefix: str, title: str, summary: str = ""):
    """Render 4 download buttons (PDF, Word, TXT, CSV) for a results dataframe."""
    if df is None or df.empty:
        st.info("Run detection first to enable exports.")
        return
    c1, c2, c3, c4 = st.columns(4)
    with c1:
        st.download_button(
            "📄 Export PDF", data=df_to_pdf_bytes(df, title, summary),
            file_name=f"{key_prefix}_results.pdf", mime="application/pdf",
            use_container_width=True, key=f"{key_prefix}_pdf",
        )
    with c2:
        st.download_button(
            "📝 Export Word", data=df_to_docx_bytes(df, title, summary),
            file_name=f"{key_prefix}_results.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True, key=f"{key_prefix}_docx",
        )
    with c3:
        st.download_button(
            "🧾 Export TXT", data=df_to_txt_bytes(df, title, summary),
            file_name=f"{key_prefix}_results.txt", mime="text/plain",
            use_container_width=True, key=f"{key_prefix}_txt",
        )
    with c4:
        st.download_button(
            "📊 Export CSV", data=df_to_csv_bytes(df),
            file_name=f"{key_prefix}_results.csv", mime="text/csv",
            use_container_width=True, key=f"{key_prefix}_csv",
        )


def schema_box(schema: dict):
    st.markdown("<div class='schema-box'><b>Data Schema</b></div>", unsafe_allow_html=True)
    lines = "".join(f"- **`{k}`** — {v}\n" for k, v in schema.items())
    st.markdown(lines)


def data_source_controls(key_prefix, generator_fn, gen_label="Generate Synthetic Data", n_default=400):
    """Common widget block: synthetic-data button + slider + optional CSV upload."""
    c1, c2 = st.columns([2, 1])
    with c1:
        n = st.slider("Number of synthetic events", 50, 2000, n_default, 50, key=f"{key_prefix}_n")
        if st.button(f"🎲 {gen_label}", key=f"{key_prefix}_gen", use_container_width=True):
            st.session_state[f"{key_prefix}_data"] = generator_fn(n)
    with c2:
        up = st.file_uploader("Or upload your own CSV", type=["csv"], key=f"{key_prefix}_upload")
        if up is not None:
            try:
                st.session_state[f"{key_prefix}_data"] = pd.read_csv(up)
                st.success(f"Loaded {len(st.session_state[f'{key_prefix}_data'])} rows from upload.")
            except Exception as e:
                st.error(f"Could not read CSV: {e}")
    return st.session_state.get(f"{key_prefix}_data")


# --------------------------------------------------------------------------
# TAB 1 — BRUTE-FORCE LOGIN DETECTION
# --------------------------------------------------------------------------

def gen_login_data(n=400, seed=None):
    rng = np.random.default_rng(seed)
    users = [f"user{i}@company.com" for i in range(1, 26)]
    ips_normal = [f"10.0.{rng.integers(0,255)}.{rng.integers(0,255)}" for _ in range(30)]
    attacker_ips = ["203.0.113.7", "198.51.100.23", "185.220.101.5"]

    rows = []
    base_time = datetime.now() - timedelta(hours=6)

    # normal traffic
    for _ in range(int(n * 0.8)):
        t = base_time + timedelta(seconds=int(rng.integers(0, 6 * 3600)))
        rows.append({
            "timestamp": t,
            "user": rng.choice(users),
            "source_ip": rng.choice(ips_normal),
            "status": rng.choice(["SUCCESS", "SUCCESS", "SUCCESS", "FAILED"], p=[0.85, 0.05, 0.05, 0.05]),
            "country": rng.choice(["US", "CA", "US", "US"]),
        })

    # brute-force bursts from attacker IPs against a few target users
    targets = rng.choice(users, size=3, replace=False)
    for atk_ip in attacker_ips:
        target = rng.choice(targets)
        burst_start = base_time + timedelta(seconds=int(rng.integers(0, 6 * 3600)))
        for k in range(int(n * 0.2) // len(attacker_ips)):
            t = burst_start + timedelta(seconds=k * rng.integers(2, 8))
            rows.append({
                "timestamp": t,
                "user": target,
                "source_ip": atk_ip,
                "status": "FAILED" if k < 15 else rng.choice(["FAILED", "SUCCESS"], p=[0.9, 0.1]),
                "country": rng.choice(["RU", "CN", "NG", "BR"]),
            })

    df = pd.DataFrame(rows).sort_values("timestamp").reset_index(drop=True)
    return df


def detect_brute_force(df, fail_threshold=8, window_minutes=10):
    df = df.copy()
    df["timestamp"] = pd.to_datetime(df["timestamp"])
    fails = df[df["status"] == "FAILED"].sort_values("timestamp")

    flagged = []
    for (ip, user), grp in fails.groupby(["source_ip", "user"]):
        grp = grp.sort_values("timestamp")
        times = grp["timestamp"].tolist()
        i = 0
        for j in range(len(times)):
            while times[j] - times[i] > timedelta(minutes=window_minutes):
                i += 1
            count = j - i + 1
            if count >= fail_threshold:
                flagged.append({
                    "source_ip": ip, "user": user,
                    "failed_attempts": count,
                    "window_start": times[i], "window_end": times[j],
                    "risk": "HIGH" if count >= fail_threshold * 2 else "MEDIUM",
                })
                break
    return pd.DataFrame(flagged).drop_duplicates(subset=["source_ip", "user"])


# --------------------------------------------------------------------------
# TAB 2 — ANOMALOUS API USAGE DETECTION
# --------------------------------------------------------------------------

def gen_api_data(n=400, seed=None):
    rng = np.random.default_rng(seed)
    services = ["billing-api", "user-api", "orders-api", "inventory-api"]
    accounts = [f"svc-account-{i}@project.iam.gserviceaccount.com" for i in range(1, 11)]
    base_time = datetime.now() - timedelta(hours=24)

    rows = []
    # normal hourly baseline per account/service
    for acct in accounts:
        for svc in services:
            for hr in range(24):
                calls = max(0, int(rng.normal(20, 6)))
                rows.append({
                    "hour": base_time + timedelta(hours=hr),
                    "service_account": acct,
                    "api_service": svc,
                    "call_count": calls,
                    "error_rate": round(float(rng.uniform(0, 0.05)), 3),
                })

    # inject anomalies: sudden spikes / data-exfil-like patterns
    anomaly_acct = rng.choice(accounts)
    for hr in [6, 7, 14]:
        rows.append({
            "hour": base_time + timedelta(hours=hr),
            "service_account": anomaly_acct,
            "api_service": "user-api",
            "call_count": int(rng.integers(400, 900)),
            "error_rate": round(float(rng.uniform(0.2, 0.5)), 3),
        })

    df = pd.DataFrame(rows)
    return df


def detect_api_anomalies(df, z_thresh=3.0):
    df = df.copy()
    grp = df.groupby(["service_account", "api_service"])["call_count"]
    stats = grp.agg(["mean", "std"]).reset_index().rename(
        columns={"mean": "baseline_mean", "std": "baseline_std"})
    merged = df.merge(stats, on=["service_account", "api_service"], how="left")
    merged["baseline_std"] = merged["baseline_std"].replace(0, np.nan)
    merged["z_score"] = (merged["call_count"] - merged["baseline_mean"]) / merged["baseline_std"]
    merged["z_score"] = merged["z_score"].fillna(0)
    flagged = merged[merged["z_score"] >= z_thresh].copy()
    flagged["risk"] = np.where(flagged["z_score"] >= z_thresh * 1.7, "HIGH", "MEDIUM")
    return flagged.sort_values("z_score", ascending=False)


# --------------------------------------------------------------------------
# TAB 3 — SUSPICIOUS NETWORK / IP TRAFFIC DETECTION
# --------------------------------------------------------------------------

THREAT_INTEL_IPS = {"203.0.113.7", "198.51.100.23", "185.220.101.5", "45.83.64.1"}
HIGH_RISK_COUNTRIES = {"RU", "KP", "NG", "IR"}
SENSITIVE_PORTS = {22, 3389, 1433, 3306, 5432, 27017}


def gen_network_data(n=400, seed=None):
    rng = np.random.default_rng(seed)
    internal_ips = [f"10.0.{rng.integers(0,20)}.{rng.integers(0,255)}" for _ in range(15)]
    countries_normal = ["US", "US", "CA", "US", "DE"]
    ports_normal = [443, 443, 80, 8080, 443]

    rows = []
    base_time = datetime.now() - timedelta(hours=12)
    for _ in range(int(n * 0.85)):
        rows.append({
            "timestamp": base_time + timedelta(seconds=int(rng.integers(0, 12 * 3600))),
            "src_ip": f"192.168.{rng.integers(0,255)}.{rng.integers(0,255)}",
            "dst_ip": rng.choice(internal_ips),
            "dst_port": int(rng.choice(ports_normal)),
            "country": rng.choice(countries_normal),
            "bytes_transferred": int(rng.normal(50000, 15000)),
        })

    for _ in range(int(n * 0.15)):
        bad_ip = rng.choice(list(THREAT_INTEL_IPS))
        rows.append({
            "timestamp": base_time + timedelta(seconds=int(rng.integers(0, 12 * 3600))),
            "src_ip": bad_ip,
            "dst_ip": rng.choice(internal_ips),
            "dst_port": int(rng.choice(list(SENSITIVE_PORTS))),
            "country": rng.choice(list(HIGH_RISK_COUNTRIES)),
            "bytes_transferred": int(rng.integers(200000, 5_000_000)),
        })

    df = pd.DataFrame(rows).sort_values("timestamp").reset_index(drop=True)
    return df


def detect_network_threats(df):
    df = df.copy()
    df["reason"] = ""
    def reasons(row):
        r = []
        if row["src_ip"] in THREAT_INTEL_IPS:
            r.append("Known malicious IP (threat intel match)")
        if row.get("country") in HIGH_RISK_COUNTRIES:
            r.append("High-risk origin country")
        if row["dst_port"] in SENSITIVE_PORTS:
            r.append(f"Sensitive port {row['dst_port']}")
        if row["bytes_transferred"] > 200000:
            r.append("Abnormally large data transfer")
        return "; ".join(r)

    df["reason"] = df.apply(reasons, axis=1)
    flagged = df[df["reason"] != ""].copy()
    flagged["risk"] = np.where(
        flagged["reason"].str.contains("threat intel"), "HIGH", "MEDIUM"
    )
    return flagged.sort_values("bytes_transferred", ascending=False)


# --------------------------------------------------------------------------
# TAB 4 — ML BEHAVIORAL ANOMALY DETECTION (Isolation Forest)
# --------------------------------------------------------------------------

def gen_behavior_data(n=400, seed=None):
    rng = np.random.default_rng(seed)
    users = [f"user{i}" for i in range(1, 41)]
    rows = []
    for u in users:
        base_hour = rng.integers(7, 10)
        base_files = rng.normal(15, 4)
        base_locations = 1
        for _ in range(n // len(users)):
            rows.append({
                "user": u,
                "login_hour": int(np.clip(rng.normal(base_hour, 1.2), 0, 23)),
                "files_accessed": max(0, int(rng.normal(base_files, 4))),
                "data_downloaded_mb": max(0, float(rng.normal(50, 20))),
                "distinct_locations": int(np.clip(rng.normal(base_locations, 0.3), 1, 5)),
                "failed_logins": int(max(0, rng.normal(0.3, 0.5))),
            })

    # inject a handful of anomalous / insider-threat-like sessions
    anomalous_users = rng.choice(users, size=4, replace=False)
    for u in anomalous_users:
        rows.append({
            "user": u,
            "login_hour": int(rng.choice([2, 3, 4])),
            "files_accessed": int(rng.integers(200, 600)),
            "data_downloaded_mb": float(rng.integers(2000, 8000)),
            "distinct_locations": int(rng.integers(3, 6)),
            "failed_logins": int(rng.integers(4, 12)),
        })

    return pd.DataFrame(rows)


def detect_behavior_anomalies(df, contamination=0.05):
    features = ["login_hour", "files_accessed", "data_downloaded_mb",
                "distinct_locations", "failed_logins"]
    df = df.copy()
    for f in features:
        if f not in df.columns:
            st.error(f"Uploaded data is missing required column: {f}")
            return pd.DataFrame()

    X = df[features].fillna(0).values
    if SKLEARN_OK:
        model = IsolationForest(contamination=contamination, random_state=42)
        preds = model.fit_predict(X)
        scores = -model.score_samples(X)
        df["anomaly_score"] = scores
        df["is_anomaly"] = preds == -1
    else:
        # fallback: simple robust z-score ensemble if scikit-learn is unavailable
        z = (X - np.median(X, axis=0)) / (np.std(X, axis=0) + 1e-6)
        score = np.abs(z).mean(axis=1)
        df["anomaly_score"] = score
        thresh = np.quantile(score, 1 - contamination)
        df["is_anomaly"] = score >= thresh

    flagged = df[df["is_anomaly"]].sort_values("anomaly_score", ascending=False)
    flagged["risk"] = np.where(
        flagged["anomaly_score"] >= flagged["anomaly_score"].quantile(0.7), "HIGH", "MEDIUM"
    )
    return flagged


# --------------------------------------------------------------------------
# UI TABS
# --------------------------------------------------------------------------

tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "🔐 Brute-Force Login",
    "🔌 Anomalous API Usage",
    "🌐 Suspicious Network/IP",
    "🤖 ML Behavioral Anomaly",
    "☁️ Architecture",
])

# ---------------- TAB 1 ----------------
with tab1:
    st.subheader("Use Case 1: Brute-Force Login Detection")
    st.markdown(
        "Detects repeated failed authentication attempts from the same source IP "
        "against the same account within a short time window — the classic signature "
        "of credential-stuffing or brute-force attacks."
    )
    with st.expander("📘 GCP Serverless Pattern & Function Logic", expanded=False):
        st.markdown(
            "- **Ingestion:** Cloud Identity/Cloud Logging login events → **Pub/Sub** topic\n"
            "- **Processing:** A **Cloud Function** (or Cloud Run service) triggered per message "
            "groups recent events by `(source_ip, user)` and counts `FAILED` events inside a "
            "sliding time window (default 10 minutes)\n"
            "- **Detection function:** `detect_brute_force(df, fail_threshold, window_minutes)` — "
            "a sliding-window counter; if failed attempts ≥ threshold, the pair is flagged and "
            "risk is escalated to **HIGH** at 2× the threshold\n"
            "- **Response:** flagged IPs published to an alert topic → Cloud Function disables the "
            "account or updates a Cloud Armor block-list"
        )
    schema_box({
        "timestamp": "Event time of the login attempt (datetime)",
        "user": "Account/email that attempted to log in",
        "source_ip": "Origin IP address of the request",
        "status": "SUCCESS or FAILED",
        "country": "GeoIP-resolved origin country",
    })

    df1 = data_source_controls("bf", gen_login_data, "Generate Synthetic Login Events")
    c1, c2 = st.columns(2)
    with c1:
        fail_threshold = st.slider("Failed-attempt threshold", 3, 30, 8, key="bf_thresh")
    with c2:
        window_min = st.slider("Time window (minutes)", 1, 60, 10, key="bf_window")

    if df1 is not None:
        st.dataframe(df1.head(200), use_container_width=True, height=220)
        if st.button("🔎 Run Brute-Force Detection", key="bf_run", type="primary"):
            st.session_state["bf_results"] = detect_brute_force(df1, fail_threshold, window_min)

        results1 = st.session_state.get("bf_results")
        if results1 is not None and not results1.empty:
            st.success(f"Flagged {len(results1)} suspicious (IP, user) pair(s).")
            st.dataframe(results1, use_container_width=True)

            fig = px.bar(
                results1, x="source_ip", y="failed_attempts", color="risk",
                color_discrete_map={"HIGH": "#c0392b", "MEDIUM": "#e67e22"},
                title="Failed Login Attempts by Attacking IP",
            )
            st.plotly_chart(fig, use_container_width=True)

            timeline = df1.copy()
            timeline["timestamp"] = pd.to_datetime(timeline["timestamp"])
            timeline["minute"] = timeline["timestamp"].dt.floor("15min")
            ts = timeline[timeline["status"] == "FAILED"].groupby("minute").size().reset_index(name="failed_count")
            fig2 = px.line(ts, x="minute", y="failed_count", title="Failed Logins Over Time (15-min bins)")
            st.plotly_chart(fig2, use_container_width=True)

            export_bar(results1, "brute_force", "Brute-Force Login Detection Results",
                       f"Threshold: {fail_threshold} failed attempts within {window_min} minutes.")
        elif results1 is not None:
            st.info("No brute-force patterns detected with the current thresholds.")
    else:
        st.info("Generate synthetic data or upload a CSV to begin.")

# ---------------- TAB 2 ----------------
with tab2:
    st.subheader("Use Case 2: Anomalous API Usage Detection")
    st.markdown(
        "Baselines hourly call volume per service account/API and flags hours where "
        "usage deviates far beyond normal — indicative of a compromised service account, "
        "leaked API key, or automated data exfiltration."
    )
    with st.expander("📘 GCP Serverless Pattern & Function Logic", expanded=False):
        st.markdown(
            "- **Ingestion:** Cloud Audit Logs / API Gateway metrics streamed via **Pub/Sub**\n"
            "- **Processing:** **Cloud Run** job aggregates call counts per "
            "`(service_account, api_service, hour)` and computes a rolling baseline\n"
            "- **Detection function:** `detect_api_anomalies(df, z_thresh)` — a statistical "
            "**z-score** test: `z = (call_count − mean) / std`; hours with `z ≥ threshold` "
            "(default 3σ) are flagged, HIGH risk above ~5σ\n"
            "- **Response:** BigQuery sink for audit trail + alert to Security Command Center"
        )
    schema_box({
        "hour": "Hourly bucket timestamp",
        "service_account": "GCP service account making the calls",
        "api_service": "Target API/service name",
        "call_count": "Number of API calls in that hour",
        "error_rate": "Fraction of calls returning errors",
    })

    df2 = data_source_controls("api", gen_api_data, "Generate Synthetic API Logs")
    z_thresh = st.slider("Z-score threshold (standard deviations)", 1.5, 6.0, 3.0, 0.5, key="api_z")

    if df2 is not None:
        st.dataframe(df2.head(200), use_container_width=True, height=220)
        if st.button("🔎 Run Anomaly Detection", key="api_run", type="primary"):
            st.session_state["api_results"] = detect_api_anomalies(df2, z_thresh)

        results2 = st.session_state.get("api_results")
        if results2 is not None and not results2.empty:
            st.success(f"Flagged {len(results2)} anomalous usage hour(s).")
            st.dataframe(results2, use_container_width=True)

            fig = px.scatter(
                df2.merge(
                    results2[["service_account", "api_service", "hour"]].assign(flag="ANOMALY"),
                    on=["service_account", "api_service", "hour"], how="left"
                ).fillna({"flag": "normal"}),
                x="hour", y="call_count", color="flag",
                color_discrete_map={"ANOMALY": "#c0392b", "normal": "#3498db"},
                title="API Call Volume Over Time (flagged points in red)",
                hover_data=["service_account", "api_service"],
            )
            st.plotly_chart(fig, use_container_width=True)

            export_bar(results2, "api_usage", "Anomalous API Usage Detection Results",
                       f"Z-score threshold: {z_thresh}")
        elif results2 is not None:
            st.info("No anomalies detected at the current z-score threshold.")
    else:
        st.info("Generate synthetic data or upload a CSV to begin.")

# ---------------- TAB 3 ----------------
with tab3:
    st.subheader("Use Case 3: Suspicious Network / IP Traffic Detection")
    st.markdown(
        "Cross-references live traffic against threat-intelligence IP lists, "
        "high-risk countries, sensitive destination ports, and abnormal transfer sizes — "
        "a lightweight rule-based network detector."
    )
    with st.expander("📘 GCP Serverless Pattern & Function Logic", expanded=False):
        st.markdown(
            "- **Ingestion:** VPC Flow Logs → **Pub/Sub**\n"
            "- **Processing:** **Cloud Function** enriches each flow with GeoIP + a threat-intel "
            "IP set (could be refreshed from Cloud Storage or an external feed)\n"
            "- **Detection function:** `detect_network_threats(df)` — rule engine checking: "
            "(1) source IP in threat-intel list, (2) high-risk origin country, "
            "(3) connection to a sensitive port (SSH/RDP/DB), (4) unusually large transfer\n"
            "- **Response:** high-confidence hits auto-block via Cloud Armor / firewall rule update"
        )
    schema_box({
        "timestamp": "Time of the network flow",
        "src_ip": "Source IP address",
        "dst_ip": "Destination (internal) IP address",
        "dst_port": "Destination port",
        "country": "GeoIP country of the source IP",
        "bytes_transferred": "Bytes transferred in the flow",
    })

    df3 = data_source_controls("net", gen_network_data, "Generate Synthetic Network Flows")

    if df3 is not None:
        st.dataframe(df3.head(200), use_container_width=True, height=220)
        if st.button("🔎 Run Network Threat Detection", key="net_run", type="primary"):
            st.session_state["net_results"] = detect_network_threats(df3)

        results3 = st.session_state.get("net_results")
        if results3 is not None and not results3.empty:
            st.success(f"Flagged {len(results3)} suspicious flow(s).")
            st.dataframe(results3, use_container_width=True)

            fig = px.bar(
                results3.groupby("country").size().reset_index(name="flagged_flows"),
                x="country", y="flagged_flows", title="Flagged Flows by Origin Country",
                color="flagged_flows", color_continuous_scale="Reds",
            )
            st.plotly_chart(fig, use_container_width=True)

            fig2 = px.pie(results3, names="risk", title="Risk Level Distribution",
                          color="risk", color_discrete_map={"HIGH": "#c0392b", "MEDIUM": "#e67e22"})
            st.plotly_chart(fig2, use_container_width=True)

            export_bar(results3, "network_threats", "Suspicious Network/IP Traffic Detection Results")
        elif results3 is not None:
            st.info("No suspicious network flows detected.")
    else:
        st.info("Generate synthetic data or upload a CSV to begin.")

# ---------------- TAB 4 ----------------
with tab4:
    st.subheader("Use Case 4: ML Behavioral Anomaly Detection")
    st.markdown(
        "Uses an **Isolation Forest** model to learn 'normal' user behavior across several "
        "features simultaneously and flag sessions that don't fit the pattern — useful for "
        "insider-threat or account-takeover detection where no single rule would catch it."
    )
    with st.expander("📘 GCP Serverless Pattern & Function Logic", expanded=False):
        st.markdown(
            "- **Ingestion:** application/session logs → **Pub/Sub** → **BigQuery**\n"
            "- **Processing:** scheduled **Cloud Run job** (or Vertex AI Pipelines) retrains/"
            "scores an **Isolation Forest** model against recent sessions\n"
            "- **Detection function:** `detect_behavior_anomalies(df, contamination)` — Isolation "
            "Forest isolates points that are 'few and different'; the `contamination` parameter "
            "sets the expected anomaly fraction (default 5%)\n"
            "- **In production:** this would run on **Vertex AI** (managed training/serving) "
            "instead of in-process, with the model artifact stored in Cloud Storage\n"
            "- **Response:** top-risk sessions routed to a SOC dashboard for human review"
        )
    schema_box({
        "user": "User/session owner",
        "login_hour": "Hour of day (0-23) the session started",
        "files_accessed": "Number of files accessed in the session",
        "data_downloaded_mb": "Data volume downloaded (MB)",
        "distinct_locations": "Distinct geographic locations seen for this user recently",
        "failed_logins": "Failed login attempts preceding this session",
    })
    if not SKLEARN_OK:
        st.warning("scikit-learn not available in this environment — using a statistical fallback (robust z-score) instead of Isolation Forest.")

    df4 = data_source_controls("beh", gen_behavior_data, "Generate Synthetic Behavior Sessions")
    contamination = st.slider("Expected anomaly fraction (contamination)", 0.01, 0.20, 0.05, 0.01, key="beh_cont")

    if df4 is not None:
        st.dataframe(df4.head(200), use_container_width=True, height=220)
        if st.button("🔎 Run ML Anomaly Detection", key="beh_run", type="primary"):
            st.session_state["beh_results"] = detect_behavior_anomalies(df4, contamination)

        results4 = st.session_state.get("beh_results")
        if results4 is not None and not results4.empty:
            st.success(f"Flagged {len(results4)} anomalous session(s).")
            st.dataframe(results4, use_container_width=True)

            fig = px.scatter(
                df4.assign(flagged=df4.index.isin(results4.index)),
                x="files_accessed", y="data_downloaded_mb", color="flagged",
                color_discrete_map={True: "#c0392b", False: "#3498db"},
                title="Session Behavior: Files Accessed vs. Data Downloaded",
                hover_data=["user"],
            )
            st.plotly_chart(fig, use_container_width=True)

            fig2 = px.histogram(results4, x="anomaly_score", nbins=20, title="Anomaly Score Distribution (flagged)")
            st.plotly_chart(fig2, use_container_width=True)

            export_bar(results4, "ml_behavior", "ML Behavioral Anomaly Detection Results",
                       f"Model: {'IsolationForest' if SKLEARN_OK else 'Robust z-score fallback'}, contamination={contamination}")
        elif results4 is not None:
            st.info("No behavioral anomalies detected at the current sensitivity.")
    else:
        st.info("Generate synthetic data or upload a CSV to begin.")

# ---------------- TAB 5: ARCHITECTURE ----------------
with tab5:
    st.subheader("Reference Serverless Architecture on GCP")
    st.markdown(
        "All four use cases follow the same general serverless pipeline shape on Google Cloud:"
    )
    fig = go.Figure()
    stages = ["Log/Event\nSources", "Pub/Sub\n(ingest)", "Cloud Functions /\nCloud Run\n(detect)",
              "BigQuery\n(store & hunt)", "Alerting /\nAuto-Response"]
    x_positions = list(range(len(stages)))
    fig.add_trace(go.Scatter(
        x=x_positions, y=[0] * len(stages), mode="markers+text",
        marker=dict(size=60, color="#0B3D91"),
        text=stages, textposition="top center", textfont=dict(size=13, color="#0B3D91"),
    ))
    for i in range(len(stages) - 1):
        fig.add_annotation(
            x=x_positions[i + 1] - 0.15, y=0, ax=x_positions[i] + 0.15, ay=0,
            xref="x", yref="y", axref="x", ayref="y",
            showarrow=True, arrowhead=3, arrowsize=1.5, arrowwidth=2, arrowcolor="#0B3D91",
        )
    fig.update_layout(
        height=260, showlegend=False,
        xaxis=dict(visible=False, range=[-0.5, len(stages) - 0.5]),
        yaxis=dict(visible=False, range=[-1, 1]),
        margin=dict(l=10, r=10, t=10, b=10),
        plot_bgcolor="white",
    )
    st.plotly_chart(fig, use_container_width=True)

    st.markdown(
        """
        | Stage | GCP Service | Role in this demo |
        |---|---|---|
        | Ingest | **Pub/Sub** | Durable, at-least-once event bus for log/telemetry streams |
        | Detect (rule-based) | **Cloud Functions** | Brute-force + network rule engines (Tabs 1 & 3) |
        | Detect (stateful/stat) | **Cloud Run** | Windowed z-score aggregation for API usage (Tab 2) |
        | Detect (ML) | **Vertex AI** | Managed training/serving for the Isolation Forest model (Tab 4) |
        | Store & Hunt | **BigQuery** | SQL-based threat hunting over historical events |
        | Respond | **Cloud Monitoring + Functions** | Alerting, auto-remediation, Cloud Armor blocklisting |
        | Governance | **Security Command Center** | Centralized findings and posture management |

        This app simulates the *detection functions* locally in Python so you can explore the
        logic and outputs interactively; in production those functions would be deployed as
        the Cloud Functions/Cloud Run services shown above, triggered by real Pub/Sub messages.
        """
    )

st.markdown("---")
st.caption("© Kalsnet (KNet) Consulting — Demo application for educational purposes. Developed by Randy Singh.")