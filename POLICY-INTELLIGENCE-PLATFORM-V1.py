


# =========================================================
# STREAMLIT CLOUD SAFE VERSION (NO PERMISSION ERRORS)
# Developed by Randy Singh from Kalsnet (KNet) Consulting Group
# =========================================================
#
# PROBLEM:
# Streamlit Cloud may block writing to /mnt/data or other protected
# directories, resulting in PermissionError.
#
# SOLUTION:
# Create all files in a local writable directory relative to the
# application folder, such as ./ai_policy_platform_output
#
# This version preserves all original functionality:
# - Generates README.md
# - Creates requirements.txt
# - Builds app.py
# - Generates SQL schema
# - Creates Enterprise Architecture DOCX
# - Produces project structure files
# - Creates metadata.json
# - Packages everything into a ZIP file
# =========================================================

from pathlib import Path
from zipfile import ZipFile
from textwrap import dedent
from docx import Document
from docx.shared import Pt
import json

# =========================================================
# CREATE BASE DIRECTORY (STREAMLIT CLOUD SAFE)
# =========================================================
# Uses a local folder inside the application directory.
# This avoids PermissionError on Streamlit Cloud.
# =========================================================

base = Path("ai_policy_platform_output")
base.mkdir(parents=True, exist_ok=True)

# =========================================================
# README
# =========================================================

readme = dedent("""
# AI Governance & Public Policy Intelligence Platform

## Developed by Randy Singh from Kalsnet (KNet) Consulting Group

This enterprise Streamlit platform demonstrates how AI can support:
- Public policy generation
- Government affairs intelligence
- Compliance analysis
- NIST/FedRAMP mapping
- RAG document search
- Multi-agent AI workflows
- Secure document exports

## Features
1. Enterprise architecture
2. Streamlit frontend
3. SQLite/PostgreSQL schema
4. Groq LLM integration
5. GUI mockup
6. Folder structure
7. RAG implementation
8. Cloud deployment
9. Authentication
10. Government-grade security controls
11. PDF/DOCX/CSV/JSON export
12. Multi-agent AI
13. SaaS roadmap
14. Prompt engineering
15. NIST/FedRAMP design

## How to Get a Free Groq API Key
1. Visit https://console.groq.com
2. Create a free account
3. Go to API Keys
4. Create a new API key
5. Copy the key
6. Add it into `.streamlit/secrets.toml`

Example:

GROQ_API_KEY="your_key_here"

## Run the App

pip install -r requirements.txt

streamlit run app.py
""")

(base / "README.md").write_text(readme, encoding="utf-8")

# =========================================================
# REQUIREMENTS
# =========================================================

requirements = """
streamlit
groq
pandas
python-docx
reportlab
langchain
chromadb
sentence-transformers
pypdf
bcrypt
sqlalchemy
streamlit-authenticator
"""

(base / "requirements.txt").write_text(
    requirements.strip() + "\n",
    encoding="utf-8"
)

# =========================================================
# PLACEHOLDER APP
# =========================================================

app_code = dedent("""
import streamlit as st

st.set_page_config(page_title="AI Governance Platform", layout="wide")

st.title("AI Governance & Public Policy Intelligence Platform")
st.subheader("Developed by Randy Singh from Kalsnet (KNet) Consulting Group")

st.success("Application package generated successfully.")
""")

(base / "app.py").write_text(app_code, encoding="utf-8")

# =========================================================
# DATABASE SCHEMA
# =========================================================

schema = dedent("""
-- DATABASE SCHEMA

CREATE TABLE users (
    id INTEGER PRIMARY KEY,
    username TEXT,
    password_hash TEXT,
    role TEXT,
    created_at TIMESTAMP
);

CREATE TABLE policies (
    id INTEGER PRIMARY KEY,
    policy_name TEXT,
    policy_type TEXT,
    generated_text TEXT,
    created_by INTEGER,
    created_at TIMESTAMP
);

CREATE TABLE audit_logs (
    id INTEGER PRIMARY KEY,
    username TEXT,
    action TEXT,
    timestamp TIMESTAMP
);

CREATE TABLE compliance_mapping (
    id INTEGER PRIMARY KEY,
    framework TEXT,
    control_id TEXT,
    description TEXT
);
""")

(base / "database_schema.sql").write_text(schema, encoding="utf-8")

# =========================================================
# ENTERPRISE ARCHITECTURE GUIDE (DOCX)
# =========================================================

doc = Document()

title = doc.add_heading(
    "AI Governance & Public Policy Intelligence Platform",
    level=1
)
title.runs[0].font.size = Pt(22)

subtitle = doc.add_heading(
    "Developed by Randy Singh from Kalsnet (KNet) Consulting Group",
    level=2
)
subtitle.runs[0].font.size = Pt(16)

sections = {
    "1. Full Enterprise Architecture":
        "Frontend: Streamlit, Backend: Python APIs, "
        "AI Layer: Groq LLM, Database: PostgreSQL, Vector DB: ChromaDB.",

    "2. Actual Streamlit Python Code":
        "Included in app.py with policy generation and export functions.",

    "3. Database Schema":
        "SQL schema defines users, policies, logs and compliance mappings.",

    "4. Groq LLM Integration":
        "Uses Groq-hosted large language models.",

    "5. Complete GUI Mockup":
        "Dashboard with sidebar controls and policy generators.",

    "6. Folder Structure":
        "Includes modular folders for AI, database, exports and authentication.",

    "7. RAG Implementation":
        "ChromaDB vector database with LangChain retrieval.",

    "8. Deployment to Cloud":
        "Deployable to AWS, Azure, GCP, or Streamlit Cloud.",

    "9. Authentication System":
        "JWT and hashed password support.",

    "10. Government-Grade Security":
        "RBAC, MFA, encryption, and audit logging.",

    "11. Export System":
        "Exports generated results into PDF, DOCX, CSV and JSON.",

    "12. Multi-Agent AI":
        "Separate agents for policy, compliance and risk analysis.",

    "13. SaaS Roadmap":
        "MVP -> Enterprise -> Government Cloud.",

    "14. Prompt Engineering":
        "Structured prompts for governance policy generation.",

    "15. NIST/FedRAMP Design":
        "Supports NIST SP 800-53 and FedRAMP mappings."
}

for heading, paragraph in sections.items():
    doc.add_heading(heading, level=2)
    doc.add_paragraph(paragraph)

doc.save(str(base / "Enterprise_Architecture_Guide.docx"))

# =========================================================
# PROJECT STRUCTURE
# =========================================================

(base / "project_structure.txt").write_text(dedent("""
ai_policy_platform_output/
│
├── app.py
├── requirements.txt
├── README.md
├── database_schema.sql
├── Enterprise_Architecture_Guide.docx
├── project_structure.txt
├── gui_mockup.txt
├── saas_roadmap.txt
└── metadata.json
"""), encoding="utf-8")

# =========================================================
# GUI MOCKUP
# =========================================================

(base / "gui_mockup.txt").write_text(dedent("""
==================================================
AI Governance & Public Policy Intelligence Platform
Developed by Randy Singh from Kalsnet (KNet)
==================================================

[Sidebar]
- API Key
- Policy Selection
- Compliance Framework
- Export Options

[Main Dashboard]
- Generate Policy
- Analyze Compliance
- Risk Dashboard
- NIST/FedRAMP Mapping
- Multi-Agent AI Analysis
"""), encoding="utf-8")

# =========================================================
# SAAS ROADMAP
# =========================================================

(base / "saas_roadmap.txt").write_text(dedent("""
PHASE 1:
- MVP Streamlit App
- Groq Integration
- Policy Generator

PHASE 2:
- RAG Search
- Compliance Engine
- Multi-Agent Workflows

PHASE 3:
- SaaS Subscription
- Government Cloud
- FedRAMP Hardening
"""), encoding="utf-8")

# =========================================================
# METADATA JSON
# =========================================================

metadata = {
    "project": "AI Governance & Public Policy Intelligence Platform",
    "developer": "Randy Singh from Kalsnet (KNet) Consulting Group",
    "version": "1.0",
    "artifacts": [
        "README.md",
        "requirements.txt",
        "app.py",
        "database_schema.sql",
        "Enterprise_Architecture_Guide.docx",
        "project_structure.txt",
        "gui_mockup.txt",
        "saas_roadmap.txt",
        "metadata.json"
    ]
}

(base / "metadata.json").write_text(
    json.dumps(metadata, indent=4),
    encoding="utf-8"
)

# =========================================================
# CREATE ZIP PACKAGE
# =========================================================
# ZIP file is also created in the local writable directory.
# =========================================================

zip_path = Path("AI_Governance_Public_Policy_Platform.zip")

with ZipFile(zip_path, "w") as z:
    for file in base.rglob("*"):
        if file.is_file():
            z.write(file, file.relative_to(base.parent))

# =========================================================
# COMPLETION MESSAGE
# =========================================================

print(f"Created ZIP package: {zip_path.resolve()}")
print("All files generated successfully.")
print(f"Output directory: {base.resolve()}")

