# CyberGuard Pro — Enterprise Cybersecurity Domain Assessment Framework
# Refactored Streamlit App (NIST CSF 2.0 & CIS Controls v8)
# Developed by Randy Singh (Kalsnet / KNet Consulting)

import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
from datetime import datetime
import json
from io import BytesIO

# ── Optional export dependencies (graceful degradation) ────────────────────
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import (
        SimpleDocTemplate, Table, TableStyle,
        Paragraph, Spacer
    )
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib import colors as rl_colors
    REPORTLAB_OK = True
except ImportError:
    REPORTLAB_OK = False

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

# ════════════════════════════════════════════════════════════════════════════
# PAGE CONFIG
# ════════════════════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="CyberGuard Pro — Cybersecurity Assessment Framework",
    page_icon="🛡️",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ════════════════════════════════════════════════════════════════════════════
# GLOBAL CSS
# ════════════════════════════════════════════════════════════════════════════
st.markdown(
    """
<style>
.app-title{font-size:2.2rem;font-weight:900;color:#fff;text-align:center;letter-spacing:.5px;margin:0}
.app-subtitle{font-size:1rem;font-weight:700;color:#90CAF9;text-align:center;margin:3px 0 0}
.header-wrap{background:linear-gradient(100deg,#002B7F 0%,#1565C0 55%,#0288D1 100%);
  border-radius:14px;padding:24px 30px 20px;margin-bottom:18px;
  box-shadow:0 6px 24px rgba(0,43,127,.3)}
.kpi-card{border-radius:12px;padding:16px 18px;color:#fff;text-align:center;
  box-shadow:0 4px 16px rgba(0,0,0,.14)}
.kpi-val{font-size:1.9rem;font-weight:800;line-height:1.1}
.kpi-lbl{font-size:.8rem;opacity:.9;margin-top:2px}
.maturity-0{background:#C62828;color:#fff;padding:2px 10px;border-radius:12px;font-size:.75rem;font-weight:700}
.maturity-1{background:#E65100;color:#fff;padding:2px 10px;border-radius:12px;font-size:.75rem;font-weight:700}
.maturity-2{background:#F9A825;color:#fff;padding:2px 10px;border-radius:12px;font-size:.75rem;font-weight:700}
.maturity-3{background:#7CB342;color:#fff;padding:2px 10px;border-radius:12px;font-size:.75rem;font-weight:700}
.maturity-4{background:#2E7D32;color:#fff;padding:2px 10px;border-radius:12px;font-size:.75rem;font-weight:700}
.maturity-5{background:#1565C0;color:#fff;padding:2px 10px;border-radius:12px;font-size:.75rem;font-weight:700}
.stTabs [data-baseweb="tab"]{font-weight:600}
.domain-card{background:#fff;border:1px solid #D6E4F0;border-radius:10px;padding:12px 16px;margin-bottom:8px}
.gap-critical{border-left:4px solid #C62828}
.gap-high{border-left:4px solid #E65100}
.gap-medium{border-left:4px solid #F9A825}
.gap-low{border-left:4px solid #2E7D32}
.info-box{background:#EEF4FF;border:1px solid #BBD6F5;border-radius:10px;padding:14px 18px;margin-bottom:10px}
.formula-box{background:#FFF8E1;border:1px solid #F0D98C;border-radius:8px;padding:10px 14px;margin:6px 0;font-family:monospace;font-size:.85rem}
</style>
""",
    unsafe_allow_html=True,
)

# ════════════════════════════════════════════════════════════════════════════
# FRAMEWORK DATA — 108 CYBERSECURITY DOMAINS
# NIST CSF 2.0 Functions: GOVERN, IDENTIFY, PROTECT, DETECT, RESPOND, RECOVER
# ════════════════════════════════════════════════════════════════════════════
FRAMEWORK = {
    "GOVERN (GV)": {
        "color": "#4A148C",
        "description": "Establish and monitor cybersecurity risk management strategy, expectations, and policy.",
        "domains": [
            ("GV-01", "Organizational Cybersecurity Strategy", "GV.OC", "CIS 17",
             "Enterprise security strategy aligned with business objectives and risk appetite."),
            ("GV-02", "Cybersecurity Governance Structure", "GV.OC", "CIS 17",
             "Board oversight, CISO reporting lines, security steering committees."),
            ("GV-03", "Security Policy Management", "GV.PO", "CIS 14",
             "Development, approval, communication, and lifecycle of security policies."),
            ("GV-04", "Roles, Responsibilities & Authorities", "GV.RR", "CIS 17",
             "RACI for security functions, accountability frameworks, job descriptions."),
            ("GV-05", "Risk Management Strategy", "GV.RM", "CIS 17",
             "Enterprise risk framework, risk appetite statements, risk tolerance thresholds."),
            ("GV-06", "Risk Assessment Program", "GV.RM", "CIS 17",
             "Periodic risk assessments, threat modeling, risk registers, heat maps."),
            ("GV-07", "Cyber Risk Quantification", "GV.RM", "CIS 17",
             "FAIR analysis, monetary risk quantification, risk-based prioritization."),
            ("GV-08", "Third-Party Risk Management", "GV.SC", "CIS 15",
             "Vendor security assessments, contract security clauses, ongoing monitoring."),
            ("GV-09", "Supply Chain Security", "GV.SC", "CIS 15",
             "Software supply chain integrity, SBOM management, hardware provenance."),
            ("GV-10", "Compliance Management", "GV.OC", "CIS 17",
             "Regulatory mapping (GDPR, HIPAA, PCI-DSS, SOX), audit management."),
            ("GV-11", "Security Metrics & Reporting", "GV.OV", "CIS 17",
             "KPIs/KRIs, executive dashboards, board reporting, trend analysis."),
            ("GV-12", "Security Budget & Investment", "GV.RR", "CIS 17",
             "Security spend optimization, ROI analysis, resource allocation."),
            ("GV-13", "Cyber Insurance Management", "GV.RM", "CIS 17",
             "Coverage assessment, policy requirements, claims readiness."),
            ("GV-14", "Legal & Regulatory Liaison", "GV.OC", "CIS 17",
             "Breach notification laws, legal hold, regulator relationships."),
            ("GV-15", "Security Culture Program", "GV.RR", "CIS 14",
             "Culture measurement, behavioral change, security champions network."),
            ("GV-16", "Audit & Assurance Management", "GV.OV", "CIS 17",
             "Internal/external audit coordination, findings remediation tracking."),
        ],
    },
    "IDENTIFY (ID)": {
        "color": "#1565C0",
        "description": "Understand the organization's assets, risks, and vulnerabilities.",
        "domains": [
            ("ID-01", "Hardware Asset Inventory", "ID.AM", "CIS 1",
             "Discovery and inventory of all hardware: servers, endpoints, IoT, OT devices."),
            ("ID-02", "Software Asset Inventory", "ID.AM", "CIS 2",
             "Authorized software catalog, license management, shadow IT detection."),
            ("ID-03", "Data Asset Inventory & Classification", "ID.AM", "CIS 3",
             "Data discovery, classification schemes, data flow mapping, crown jewels."),
            ("ID-04", "Cloud Asset Management", "ID.AM", "CIS 1,2",
             "Multi-cloud inventory, CSPM, cloud resource tagging, orphaned resources."),
            ("ID-05", "Identity Inventory", "ID.AM", "CIS 5",
             "Human and non-human identity catalog, service accounts, API keys."),
            ("ID-06", "Network Architecture Documentation", "ID.AM", "CIS 12",
             "Network diagrams, data flows, trust boundaries, segmentation maps."),
            ("ID-07", "Vulnerability Management Program", "ID.RA", "CIS 7",
             "Scanning cadence, risk-based prioritization (CVSS/EPSS), remediation SLAs."),
            ("ID-08", "Penetration Testing Program", "ID.RA", "CIS 18",
             "Internal/external pentests, red team exercises, purple teaming."),
            ("ID-09", "Threat Intelligence Program", "ID.RA", "CIS 17",
             "TI feeds, IOC management, threat actor tracking, intelligence sharing (ISACs)."),
            ("ID-10", "Attack Surface Management", "ID.RA", "CIS 1,7",
             "External attack surface discovery, digital footprint, exposed services."),
            ("ID-11", "Business Impact Analysis", "ID.RA", "CIS 11",
             "Critical process identification, RTO/RPO definition, dependency mapping."),
            ("ID-12", "Threat Modeling", "ID.RA", "CIS 16",
             "STRIDE/PASTA analysis, design-phase security review, abuse cases."),
            ("ID-13", "Security Baseline Assessment", "ID.RA", "CIS 4",
             "Configuration benchmarks (CIS/STIG), baseline drift detection."),
            ("ID-14", "Insider Threat Program", "ID.RA", "CIS 14",
             "Insider risk indicators, UEBA, privileged user monitoring."),
            ("ID-15", "Crown Jewels Analysis", "ID.AM", "CIS 3",
             "Mission-critical asset identification, high-value target protection."),
            ("ID-16", "OT/ICS Asset Discovery", "ID.AM", "CIS 1",
             "Industrial control system inventory, SCADA visibility, passive discovery."),
        ],
    },
    "PROTECT (PR)": {
        "color": "#2E7D32",
        "description": "Implement safeguards to secure assets and limit the impact of cybersecurity events.",
        "domains": [
            ("PR-01", "Identity & Access Management (IAM)", "PR.AA", "CIS 5,6",
             "Identity lifecycle, provisioning/deprovisioning, access certification."),
            ("PR-02", "Multi-Factor Authentication", "PR.AA", "CIS 6",
             "MFA coverage for all users, phishing-resistant methods (FIDO2, passkeys)."),
            ("PR-03", "Privileged Access Management (PAM)", "PR.AA", "CIS 5",
             "Vault-based credential management, session recording, JIT elevation."),
            ("PR-04", "Single Sign-On & Federation", "PR.AA", "CIS 6",
             "SSO deployment, SAML/OIDC federation, identity provider hardening."),
            ("PR-05", "Zero Trust Architecture", "PR.AA", "CIS 6,12,13",
             "Never trust always verify, micro-perimeters, continuous verification."),
            ("PR-06", "Endpoint Protection (EPP/EDR)", "PR.PS", "CIS 10",
             "Anti-malware, EDR deployment, endpoint hardening, device control."),
            ("PR-07", "Mobile Device Management", "PR.PS", "CIS 4",
             "MDM/MAM enrollment, BYOD policy, mobile threat defense."),
            ("PR-08", "Secure Configuration Management", "PR.PS", "CIS 4",
             "Hardened images, CIS benchmarks, configuration drift monitoring."),
            ("PR-09", "Patch Management", "PR.PS", "CIS 7",
             "Patch deployment SLAs, emergency patching, virtual patching."),
            ("PR-10", "Application Whitelisting", "PR.PS", "CIS 2",
             "Allowlisting execution control, script blocking, LOLBin protections."),
            ("PR-11", "Network Segmentation", "PR.IR", "CIS 12",
             "VLAN/micro-segmentation, east-west traffic control, DMZ architecture."),
            ("PR-12", "Firewall & Perimeter Security", "PR.IR", "CIS 13",
             "NGFW rules management, egress filtering, rule review lifecycle."),
            ("PR-13", "Secure Remote Access (VPN/ZTNA)", "PR.AA", "CIS 12",
             "VPN hardening, ZTNA migration, split-tunnel policy, posture checks."),
            ("PR-14", "Wireless Network Security", "PR.IR", "CIS 12",
             "WPA3-Enterprise, rogue AP detection, guest network isolation."),
            ("PR-15", "Email Security", "PR.PS", "CIS 9",
             "SPF/DKIM/DMARC, anti-phishing gateway, attachment sandboxing."),
            ("PR-16", "Web Security Gateway", "PR.PS", "CIS 9",
             "URL filtering, SSL inspection, DNS security, browser isolation."),
            ("PR-17", "Data Loss Prevention (DLP)", "PR.DS", "CIS 3",
             "Endpoint/network/cloud DLP, content inspection, exfiltration prevention."),
            ("PR-18", "Encryption at Rest", "PR.DS", "CIS 3",
             "Full-disk encryption, database TDE, file-level encryption, key rotation."),
            ("PR-19", "Encryption in Transit", "PR.DS", "CIS 3",
             "TLS 1.2+ enforcement, certificate management, mTLS for services."),
            ("PR-20", "Key & Certificate Management", "PR.DS", "CIS 3",
             "HSM/KMS deployment, PKI lifecycle, certificate expiry monitoring."),
            ("PR-21", "Database Security", "PR.DS", "CIS 3",
             "DB activity monitoring, least-privilege DB accounts, query auditing."),
            ("PR-22", "Cloud Security Posture (CSPM)", "PR.PS", "CIS 4",
             "Misconfiguration detection, IaC scanning, compliance guardrails."),
            ("PR-23", "Cloud Workload Protection (CWPP)", "PR.PS", "CIS 10",
             "Container/VM runtime protection, serverless security, drift detection."),
            ("PR-24", "Container & Kubernetes Security", "PR.PS", "CIS 4,16",
             "Image scanning, admission control, pod security standards, RBAC."),
            ("PR-25", "Secure SDLC / DevSecOps", "PR.PS", "CIS 16",
             "Shift-left security, SAST/DAST/SCA pipeline gates, security requirements."),
            ("PR-26", "API Security", "PR.PS", "CIS 16",
             "API gateway controls, OAuth scopes, rate limiting, schema validation."),
            ("PR-27", "Secrets Management", "PR.DS", "CIS 16",
             "Vault adoption, secret rotation, hardcoded secret scanning."),
            ("PR-28", "Security Awareness Training", "PR.AT", "CIS 14",
             "Role-based training, phishing simulations, measured behavior change."),
            ("PR-29", "Developer Security Training", "PR.AT", "CIS 14,16",
             "Secure coding training, OWASP Top-10, security champions."),
            ("PR-30", "Physical Security", "PR.IR", "CIS 12",
             "Datacenter access controls, badge systems, CCTV, environmental controls."),
            ("PR-31", "Media Protection & Sanitization", "PR.DS", "CIS 3",
             "Secure disposal, degaussing, data destruction certificates."),
            ("PR-32", "Backup Security & Immutability", "PR.DS", "CIS 11",
             "Immutable backups, 3-2-1 rule, air-gapped copies, backup encryption."),
            ("PR-33", "Browser Security", "PR.PS", "CIS 9",
             "Browser hardening, extension control, enterprise browser management."),
            ("PR-34", "OT/ICS Security Controls", "PR.IR", "CIS 12",
             "Purdue model zoning, unidirectional gateways, OT protocol security."),
            ("PR-35", "AI/ML Security", "PR.PS", "CIS 16",
             "Model security, prompt injection defense, AI supply chain, LLM guardrails."),
        ],
    },
    "DETECT (DE)": {
        "color": "#E65100",
        "description": "Identify the occurrence of cybersecurity events through monitoring and detection.",
        "domains": [
            ("DE-01", "Security Operations Center (SOC)", "DE.CM", "CIS 8,13",
             "24×7 monitoring capability, tiered analyst model, SOC metrics."),
            ("DE-02", "SIEM Platform Management", "DE.CM", "CIS 8",
             "Log aggregation, correlation rules, use case development, tuning."),
            ("DE-03", "Log Management & Retention", "DE.CM", "CIS 8",
             "Centralized logging, retention policy, log integrity, time sync."),
            ("DE-04", "Endpoint Detection & Response", "DE.CM", "CIS 10,13",
             "EDR telemetry, behavioral detection, automated containment."),
            ("DE-05", "Network Detection & Response (NDR)", "DE.CM", "CIS 13",
             "Traffic analysis, lateral movement detection, encrypted traffic analytics."),
            ("DE-06", "Extended Detection & Response (XDR)", "DE.CM", "CIS 13",
             "Cross-layer correlation: endpoint + network + cloud + identity."),
            ("DE-07", "User & Entity Behavior Analytics", "DE.AE", "CIS 8",
             "Baseline behavior modeling, anomaly detection, insider threat detection."),
            ("DE-08", "Threat Hunting Program", "DE.AE", "CIS 13",
             "Hypothesis-driven hunting, MITRE ATT&CK coverage, hunt playbooks."),
            ("DE-09", "Intrusion Detection/Prevention", "DE.CM", "CIS 13",
             "IDS/IPS signatures, custom rules, inline blocking, alert tuning."),
            ("DE-10", "File Integrity Monitoring", "DE.CM", "CIS 3",
             "Critical file change detection, registry monitoring, baseline compare."),
            ("DE-11", "Cloud Security Monitoring", "DE.CM", "CIS 8,13",
             "CloudTrail/Activity Log analysis, cloud-native detections, CIEM alerts."),
            ("DE-12", "Container Runtime Detection", "DE.CM", "CIS 13",
             "Container behavioral monitoring, syscall analysis, escape detection."),
            ("DE-13", "Deception Technology", "DE.AE", "CIS 13",
             "Honeypots, honey tokens, decoy credentials, canary files."),
            ("DE-14", "Dark Web Monitoring", "DE.AE", "CIS 17",
             "Credential leak monitoring, brand abuse, data breach chatter."),
            ("DE-15", "Alert Triage & Case Management", "DE.AE", "CIS 8",
             "Alert workflow, false positive management, MTTD/MTTA tracking."),
            ("DE-16", "Detection Engineering", "DE.AE", "CIS 8",
             "Detection-as-code, Sigma rules, ATT&CK mapping, detection coverage."),
            ("DE-17", "Vulnerability Exploitation Detection", "DE.CM", "CIS 7,13",
             "Exploit attempt detection, virtual patching validation, IOA monitoring."),
            ("DE-18", "OT/ICS Anomaly Detection", "DE.CM", "CIS 13",
             "Industrial protocol monitoring, process anomaly detection, OT baselining."),
        ],
    },
    "RESPOND (RS)": {
        "color": "#C62828",
        "description": "Take action regarding detected cybersecurity incidents to contain their impact.",
        "domains": [
            ("RS-01", "Incident Response Plan & Playbooks", "RS.MA", "CIS 17",
             "IR plan, scenario playbooks (ransomware, BEC, data breach), runbooks."),
            ("RS-02", "Incident Response Team (CSIRT)", "RS.MA", "CIS 17",
             "Team structure, on-call rotation, retainer agreements, skills matrix."),
            ("RS-03", "Incident Classification & Escalation", "RS.MA", "CIS 17",
             "Severity matrix, escalation paths, declaration criteria, SLAs."),
            ("RS-04", "Digital Forensics Capability", "RS.AN", "CIS 17",
             "Evidence acquisition, chain of custody, memory/disk forensics, tooling."),
            ("RS-05", "Malware Analysis", "RS.AN", "CIS 17",
             "Sandbox analysis, reverse engineering, IOC extraction, YARA rules."),
            ("RS-06", "Containment & Eradication", "RS.MI", "CIS 17",
             "Isolation procedures, credential reset at scale, persistence removal."),
            ("RS-07", "SOAR & Response Automation", "RS.MI", "CIS 17",
             "Automated playbooks, orchestration workflows, enrichment automation."),
            ("RS-08", "Crisis Communication Plan", "RS.CO", "CIS 17",
             "Internal comms, media handling, customer notification templates."),
            ("RS-09", "Breach Notification Compliance", "RS.CO", "CIS 17",
             "72-hour GDPR readiness, state breach laws, regulator notifications."),
            ("RS-10", "Ransomware Response Readiness", "RS.MI", "CIS 11,17",
             "Ransomware playbook, negotiation policy, decryption capability, OFAC."),
            ("RS-11", "Tabletop Exercises & Simulations", "RS.MA", "CIS 17",
             "Executive tabletops, technical drills, full simulation exercises."),
            ("RS-12", "Post-Incident Review", "RS.AN", "CIS 17",
             "Lessons learned, root cause analysis, corrective action tracking."),
            ("RS-13", "Threat Actor Engagement Policy", "RS.CO", "CIS 17",
             "Law enforcement liaison, attribution policy, ransom payment stance."),
            ("RS-14", "Business Email Compromise Response", "RS.MI", "CIS 9,17",
             "BEC playbook, wire fraud recovery, mailbox forensics."),
        ],
    },
    "RECOVER (RC)": {
        "color": "#00695C",
        "description": "Restore capabilities and services impaired by cybersecurity incidents.",
        "domains": [
            ("RC-01", "Business Continuity Planning", "RC.RP", "CIS 11",
             "BCP documentation, alternate sites, workaround procedures."),
            ("RC-02", "Disaster Recovery Program", "RC.RP", "CIS 11",
             "DR plans, failover procedures, DR site management, RTO/RPO validation."),
            ("RC-03", "Backup & Restoration Testing", "RC.RP", "CIS 11",
             "Regular restore testing, recovery time measurement, backup validation."),
            ("RC-04", "Cyber Recovery Vault", "RC.RP", "CIS 11",
             "Isolated recovery environment, clean room, golden images."),
            ("RC-05", "Recovery Prioritization", "RC.RP", "CIS 11",
             "Tiered recovery sequencing, dependency-aware restoration order."),
            ("RC-06", "Post-Incident System Hardening", "RC.IM", "CIS 4",
             "Rebuild standards, compromise-informed hardening, validation scans."),
            ("RC-07", "Recovery Communication", "RC.CO", "CIS 17",
             "Stakeholder updates, service restoration announcements, trust rebuilding."),
            ("RC-08", "Reputation & Trust Recovery", "RC.CO", "CIS 17",
             "Customer confidence programs, transparency reports, PR strategy."),
            ("RC-09", "Recovery Metrics & Improvement", "RC.IM", "CIS 17",
             "MTTR tracking, recovery exercise results, plan update cadence."),
        ],
    },
}

MATURITY_LEVELS = {
    0: ("Not Performed", "#C62828", "No capability exists; control is absent."),
    1: ("Initial / Ad-hoc", "#E65100", "Informal, reactive, undocumented practices."),
    2: ("Developing", "#F9A825", "Documented but inconsistently applied."),
    3: ("Defined", "#7CB342", "Standardized, consistently implemented across the org."),
    4: ("Managed", "#2E7D32", "Measured, monitored with metrics, actively managed."),
    5: ("Optimized", "#1565C0", "Continuously improved, automated, industry-leading."),
}

FUNCTION_LIST = list(FRAMEWORK.keys())
TOTAL_DOMAINS = sum(len(f["domains"]) for f in FRAMEWORK.values())

FUNCTION_GUIDANCE = {
    "GOVERN (GV)": (
        "**What it covers:** GOVERN is the foundation function — it sets the cybersecurity risk management "
        "strategy, expectations, and policy that everything else operates under.\n\n"
        "**How to score it:** Look for a documented and board-approved security strategy, a named risk owner, "
        "current and enforced policies, a formal risk register, active third-party/supply-chain due diligence, "
        "and regular metrics reporting to leadership."
    ),
    "IDENTIFY (ID)": (
        "**What it covers:** IDENTIFY builds the organization's understanding of its assets, data, systems, risks, "
        "and vulnerabilities.\n\n"
        "**How to score it:** Look for automated inventories, data classification, vulnerability management, "
        "penetration testing, threat intelligence, and business impact analysis."
    ),
    "PROTECT (PR)": (
        "**What it covers:** PROTECT is the largest function — preventive safeguards: IAM, endpoint/network security, "
        "data protection, secure development, and training.\n\n"
        "**How to score it:** Look for enforced MFA, least-privilege, encryption, EDR, segmentation, secure SDLC, "
        "and measurable awareness training."
    ),
    "DETECT (DE)": (
        "**What it covers:** DETECT is monitoring, logging, and analysis across endpoint, network, cloud, and identity.\n\n"
        "**How to score it:** Look for 24×7 SOC/SIEM, centralized logs, EDR/NDR/XDR telemetry, tuned detections, "
        "and tracked MTTD."
    ),
    "RESPOND (RS)": (
        "**What it covers:** RESPOND is containment, communication, forensics, and execution of the IR plan.\n\n"
        "**How to score it:** Look for documented and tested IR plans, CSIRT, forensics, crisis comms, breach "
        "notification, and SOAR automation."
    ),
    "RECOVER (RC)": (
        "**What it covers:** RECOVER is restoring systems, data, and services and rebuilding trust.\n\n"
        "**How to score it:** Look for documented BCP/DR, tested backups, immutable copies, recovery sequencing, "
        "and post-incident hardening."
    ),
}

# ════════════════════════════════════════════════════════════════════════════
# SESSION STATE
# ════════════════════════════════════════════════════════════════════════════
def _init_state():
    if "scores" not in st.session_state:
        st.session_state.scores = {}
        for f in FRAMEWORK.values():
            for d in f["domains"]:
                st.session_state.scores[d[0]] = 0
    if "targets" not in st.session_state:
        st.session_state.targets = {}
        for f in FRAMEWORK.values():
            for d in f["domains"]:
                st.session_state.targets[d[0]] = 3
    if "org_name" not in st.session_state:
        st.session_state.org_name = "My Organization"
    if "assessor" not in st.session_state:
        st.session_state.assessor = ""
    if "notes" not in st.session_state:
        st.session_state.notes = {}
    if "timestamp" not in st.session_state:
        st.session_state.timestamp = datetime.utcnow().isoformat()


_init_state()

# ════════════════════════════════════════════════════════════════════════════
# HELPERS
# ════════════════════════════════════════════════════════════════════════════
def get_assessment_df() -> pd.DataFrame:
    rows = []
    for fname, f in FRAMEWORK.items():
        for d in f["domains"]:
            did, name, nist, cis, desc = d
            score = st.session_state.scores.get(did, 0)
            target = st.session_state.targets.get(did, 3)
            rows.append(
                {
                    "Domain ID": did,
                    "Domain": name,
                    "Function": fname,
                    "NIST CSF 2.0": nist,
                    "CIS Controls": cis,
                    "Current Maturity": score,
                    "Maturity Label": MATURITY_LEVELS[score][0],
                    "Target Maturity": target,
                    "Gap": max(0, target - score),
                    "Description": desc,
                    "Notes": st.session_state.notes.get(did, ""),
                }
            )
    return pd.DataFrame(rows)


def gap_priority(gap: int) -> str:
    if gap >= 3:
        return "Critical"
    if gap == 2:
        return "High"
    if gap == 1:
        return "Medium"
    return "None"


def function_scores(df: pd.DataFrame) -> pd.DataFrame:
    return (
        df.groupby("Function")
        .agg(
            Avg_Current=("Current Maturity", "mean"),
            Avg_Target=("Target Maturity", "mean"),
            Domains=("Domain ID", "count"),
            Total_Gap=("Gap", "sum"),
        )
        .reset_index()
    )


def function_abbrev(fname: str) -> str:
    # GOVERN (GV) → GV
    if "(" in fname and ")" in fname:
        return fname.split("(")[1].split(")")[0]
    return fname


# ════════════════════════════════════════════════════════════════════════════
# EXPORT FUNCTIONS
# ════════════════════════════════════════════════════════════════════════════
def export_json(df: pd.DataFrame):
    payload = {
        "organization": st.session_state.org_name,
        "assessor": st.session_state.assessor,
        "timestamp": st.session_state.timestamp,
        "data": df.to_dict(orient="records"),
    }
    data = json.dumps(payload, indent=2).encode("utf-8")
    return data, "CyberGuard-Assessment.json", "application/json"


def export_excel(df: pd.DataFrame):
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="xlsxwriter") as writer:
        df.to_excel(writer, sheet_name="Assessment", index=False)
    data = buffer.getvalue()
    buffer.close()
    return data, "CyberGuard-Assessment.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"


def export_pdf(df: pd.DataFrame):
    if not REPORTLAB_OK:
        return None, None, None

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    elements = []

    styles = getSampleStyleSheet()
    title_style = styles["Heading1"]
    normal_style = styles["BodyText"]

    elements.append(Paragraph("CyberGuard Pro – Gap Report", title_style))
    elements.append(Spacer(1, 12))
    elements.append(
        Paragraph(
            f"Organization: {st.session_state.org_name}<br/>Assessor: {st.session_state.assessor}",
            normal_style,
        )
    )
    elements.append(Spacer(1, 12))

    gdata = [["Domain ID", "Domain", "Function", "Current", "Target", "Gap"]]

    for _, r in df.iterrows():
        func_abbrev = function_abbrev(r["Function"])
        gdata.append(
            [
                r["Domain ID"],
                r["Domain"][:38],
                func_abbrev,
                str(r["Current Maturity"]),
                str(r["Target Maturity"]),
                str(r["Gap"]),
            ]
        )

    table = Table(gdata)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), rl_colors.lightgrey),
                ("GRID", (0, 0), (-1, -1), 0.5, rl_colors.black),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("ALIGN", (0, 0), (-1, -1), "LEFT"),
            ]
        )
    )

    elements.append(table)
    doc.build(elements)

    pdf_data = buffer.getvalue()
    buffer.close()

    return pdf_data, "CyberGuard-Gap-Report.pdf", "application/pdf"


def export_docx(df: pd.DataFrame):
    if not DOCX_OK:
        return None, None, None

    doc = Document()
    doc.add_heading("CyberGuard Pro – Cybersecurity Assessment", level=1)

    p = doc.add_paragraph()
    p.add_run(f"Organization: {st.session_state.org_name}\n").bold = True
    p.add_run(f"Assessor: {st.session_state.assessor}\n").bold = True
    p.add_run(f"Timestamp: {st.session_state.timestamp}\n")

    table = doc.add_table(rows=1, cols=6)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Domain ID"
    hdr_cells[1].text = "Domain"
    hdr_cells[2].text = "Function"
    hdr_cells[3].text = "Current"
    hdr_cells[4].text = "Target"
    hdr_cells[5].text = "Gap"

    for _, r in df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = r["Domain ID"]
        row_cells[1].text = r["Domain"]
        row_cells[2].text = function_abbrev(r["Function"])
        row_cells[3].text = str(r["Current Maturity"])
        row_cells[4].text = str(r["Target Maturity"])
        row_cells[5].text = str(r["Gap"])

    buffer = BytesIO()
    doc.save(buffer)
    data = buffer.getvalue()
    buffer.close()

    return data, "CyberGuard-Assessment.docx", (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


# ════════════════════════════════════════════════════════════════════════════
# UI COMPONENTS
# ════════════════════════════════════════════════════════════════════════════
def render_header():
    with st.container():
        st.markdown(
            """
<div class="header-wrap">
  <div class="app-title">CyberGuard Pro — Enterprise Cybersecurity Assessment</div>
  <div class="app-subtitle">NIST CSF 2.0 & CIS Controls v8 • 100+ Security Domains</div>
</div>
""",
            unsafe_allow_html=True,
        )


def render_sidebar(df: pd.DataFrame):
    with st.sidebar:
        st.markdown("### Assessment Metadata")
        st.session_state.org_name = st.text_input(
            "Organization Name", st.session_state.org_name
        )
        st.session_state.assessor = st.text_input(
            "Assessor Name", st.session_state.assessor
        )

        st.markdown("---")
        st.markdown("### Import / Export")

        uploaded = st.file_uploader(
            "Load previous assessment (JSON)", type=["json"], key="json_upload"
        )
        if uploaded is not None:
            try:
                payload = json.load(uploaded)
                for row in payload.get("data", []):
                    did = row.get("Domain ID")
                    if did in st.session_state.scores:
                        st.session_state.scores[did] = int(
                            row.get("Current Maturity", 0)
                        )
                        st.session_state.targets[did] = int(
                            row.get("Target Maturity", 3)
                        )
                        st.session_state.notes[did] = row.get("Notes", "")
                st.success("Assessment loaded from JSON.")
            except Exception as e:
                st.error(f"Failed to load JSON: {e}")

        st.markdown("#### Export current assessment")

        json_data, json_name, json_mime = export_json(df)
        st.download_button(
            "⬇️ Export JSON",
            data=json_data,
            file_name=json_name,
            mime=json_mime,
        )

        excel_data, excel_name, excel_mime = export_excel(df)
        st.download_button(
            "⬇️ Export Excel",
            data=excel_data,
            file_name=excel_name,
            mime=excel_mime,
        )

        pdf_data, pdf_name, pdf_mime = export_pdf(df)
        if pdf_data is not None:
            st.download_button(
                "⬇️ Export PDF Gap Report",
                data=pdf_data,
                file_name=pdf_name,
                mime=pdf_mime,
            )
        else:
            st.info("PDF export requires reportlab installed.")

        docx_data, docx_name, docx_mime = export_docx(df)
        if docx_data is not None:
            st.download_button(
                "⬇️ Export Word Report",
                data=docx_data,
                file_name=docx_name,
                mime=docx_mime,
            )
        else:
            st.info("Word export requires python-docx installed.")


def render_kpis(df: pd.DataFrame):
    avg_current = df["Current Maturity"].mean()
    avg_target = df["Target Maturity"].mean()
    total_gap = df["Gap"].sum()
    critical_count = (df["Gap"] >= 3).sum()

    c1, c2, c3, c4 = st.columns(4)

    with c1:
        st.markdown(
            f"""
<div class="kpi-card" style="background:#1565C0;">
  <div class="kpi-val">{avg_current:.2f}</div>
  <div class="kpi-lbl">Average Current Maturity</div>
</div>
""",
            unsafe_allow_html=True,
        )
    with c2:
        st.markdown(
            f"""
<div class="kpi-card" style="background:#2E7D32;">
  <div class="kpi-val">{avg_target:.2f}</div>
  <div class="kpi-lbl">Average Target Maturity</div>
</div>
""",
            unsafe_allow_html=True,
        )
    with c3:
        st.markdown(
            f"""
<div class="kpi-card" style="background:#E65100;">
  <div class="kpi-val">{int(total_gap)}</div>
  <div class="kpi-lbl">Total Maturity Gap (All Domains)</div>
</div>
""",
            unsafe_allow_html=True,
        )
    with c4:
        st.markdown(
            f"""
<div class="kpi-card" style="background:#C62828;">
  <div class="kpi-val">{int(critical_count)}</div>
  <div class="kpi-lbl">Critical Gap Domains (Gap ≥ 3)</div>
</div>
""",
            unsafe_allow_html=True,
        )


def render_charts(df: pd.DataFrame):
    fs = function_scores(df)

    c1, c2 = st.columns(2)

    with c1:
        st.markdown("#### Function-Level Maturity Radar")
        categories = [function_abbrev(f) for f in fs["Function"]]

        fig = go.Figure()
        fig.add_trace(
            go.Scatterpolar(
                r=fs["Avg_Current"],
                theta=categories,
                fill="toself",
                name="Current",
            )
        )
        fig.add_trace(
            go.Scatterpolar(
                r=fs["Avg_Target"],
                theta=categories,
                fill="toself",
                name="Target",
            )
        )
        fig.update_layout(
            polar=dict(radialaxis=dict(visible=True, range=[0, 5])),
            showlegend=True,
        )
        st.plotly_chart(fig, use_container_width=True)

    with c2:
        st.markdown("#### Total Gap by Function")
        bar = px.bar(
            fs,
            x=[function_abbrev(f) for f in fs["Function"]],
            y="Total_Gap",
            color="Total_Gap",
            color_continuous_scale="Reds",
            labels={"x": "Function", "Total_Gap": "Total Gap"},
        )
        st.plotly_chart(bar, use_container_width=True)


def render_domain_scoring(df: pd.DataFrame):
    st.markdown("### Domain-Level Scoring")

    fn_tabs = st.tabs([function_abbrev(f) for f in FUNCTION_LIST])

    for tab, fname in zip(fn_tabs, FUNCTION_LIST):
        with tab:
            st.markdown(f"#### {fname}")
            st.markdown(FUNCTION_GUIDANCE.get(fname, ""))

            fdata = FRAMEWORK[fname]
            for did, name, nist, cis, desc in fdata["domains"]:
                score = st.session_state.scores.get(did, 0)
                target = st.session_state.targets.get(did, 3)
                note = st.session_state.notes.get(did, "")

                gap = max(0, target - score)
                priority = gap_priority(gap)
                gap_class = (
                    "gap-critical"
                    if priority == "Critical"
                    else "gap-high"
                    if priority == "High"
                    else "gap-medium"
                    if priority == "Medium"
                    else "gap-low"
                )

                with st.expander(f"{did} — {name}", expanded=False):
                    st.markdown(
                        f"""
<div class="domain-card {gap_class}">
  <div><strong>NIST CSF 2.0:</strong> {nist} • <strong>CIS Controls:</strong> {cis}</div>
  <div style="margin-top:4px;">{desc}</div>
</div>
""",
                        unsafe_allow_html=True,
                    )

                    c1, c2 = st.columns(2)
                    with c1:
                        st.session_state.scores[did] = st.slider(
                            f"Current Maturity ({did})",
                            0,
                            5,
                            int(score),
                            key=f"score_{did}",
                        )
                    with c2:
                        st.session_state.targets[did] = st.slider(
                            f"Target Maturity ({did})",
                            0,
                            5,
                            int(target),
                            key=f"target_{did}",
                        )

                    st.session_state.notes[did] = st.text_area(
                        f"Notes ({did})",
                        value=note,
                        key=f"notes_{did}",
                    )

                    st.markdown(
                        f"**Gap:** {max(0, st.session_state.targets[did] - st.session_state.scores[did])} "
                        f"({gap_priority(max(0, st.session_state.targets[did] - st.session_state.scores[did]))})"
                    )


def render_gap_table(df: pd.DataFrame):
    st.markdown("### Gap Prioritization Table")

    df_gap = df.copy()
    df_gap["Priority"] = df_gap["Gap"].apply(gap_priority)
    df_gap = df_gap.sort_values(by=["Gap"], ascending=False)

    st.dataframe(
        df_gap[
            [
                "Domain ID",
                "Domain",
                "Function",
                "Current Maturity",
                "Target Maturity",
                "Gap",
                "Priority",
            ]
        ],
        use_container_width=True,
    )


# ════════════════════════════════════════════════════════════════════════════
# MAIN APP
# ════════════════════════════════════════════════════════════════════════════
def main():
    render_header()

    df = get_assessment_df()
    render_sidebar(df)

    render_kpis(df)
    render_charts(df)

    st.markdown("---")
    render_domain_scoring(df)

    st.markdown("---")
    render_gap_table(df)


if __name__ == "__main__":
    main()
