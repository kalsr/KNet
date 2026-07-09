# HealthID - Healthcare Identity Management System
# Developed by Randy Singh from Kalsnet (KNet) Consulting
import streamlit as st
import pandas as pd
import numpy as np
from faker import Faker
import random
from datetime import datetime, date, timedelta
import json
from io import BytesIO
import base64
import plotly.express as px
import plotly.graph_objects as go
# ── Optional heavy dependencies (graceful degradation) ─────────────────────────
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import (
        SimpleDocTemplate, Table, TableStyle,
        Paragraph, Spacer, HRFlowable,
    )
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors as rl_colors
    from reportlab.lib.units import inch
    REPORTLAB_OK = True
except ImportError:
    REPORTLAB_OK = False
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_OK = True
except ImportError:
    DOCX_OK = False
# ── Seed / helpers ─────────────────────────────────────────────────────────────
fake = Faker()
Faker.seed(42)
random.seed(42)
np.random.seed(42)
# ══════════════════════════════════════════════════════════════════════════════
# PAGE CONFIG
# ══════════════════════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="HealthID Pro — Healthcare Identity Management",
    page_icon="🏥",
    layout="wide",
    initial_sidebar_state="expanded",
)
# ══════════════════════════════════════════════════════════════════════════════
# GLOBAL CSS
# ══════════════════════════════════════════════════════════════════════════════
st.markdown("""
<style>
/* ── Title / header ─────────────────────────────── */
.app-title {
    font-size: 2.3rem;
    font-weight: 900;
    color: #ffffff;
    text-align: center;
    letter-spacing: .5px;
    margin: 0;
    text-shadow: 0 2px 6px rgba(0,0,0,.35);
}
.app-subtitle {
    font-size: 1.05rem;
    font-weight: 700;
    color: #FFF3C4;
    text-align: center;
    margin: 4px 0 0 0;
    letter-spacing: .3px;
}
.header-wrap {
    background: linear-gradient(100deg,#002B7F 0%,#1565C0 55%,#0288D1 100%);
    border-radius: 14px;
    padding: 22px 30px 18px;
    margin-bottom: 18px;
    box-shadow: 0 6px 24px rgba(0,43,127,.25);
    border: 1px solid rgba(255,255,255,.15);
}
/* ── Section / title bars (page headers) ────────── */
.section-header {
    background: linear-gradient(90deg,#002B7F 0%,#1565C0 60%,#0288D1 100%);
    color: #ffffff !important;
    padding: 14px 22px;
    border-radius: 10px;
    font-size: 1.5rem;
    font-weight: 800;
    letter-spacing: .3px;
    margin: 4px 0 18px 0;
    box-shadow: 0 4px 14px rgba(0,43,127,.28);
    border-left: 7px solid #FFD54F;
}
.subsection-header {
    background: #EEF4FF;
    color: #002B7F !important;
    padding: 8px 16px;
    border-radius: 8px;
    font-size: 1.05rem;
    font-weight: 800;
    margin: 10px 0 12px 0;
    border-left: 5px solid #1565C0;
}
/* ── Metric cards ───────────────────────────────── */
.kpi-card {
    border-radius: 12px;
    padding: 18px 20px;
    color: #fff;
    text-align: center;
    box-shadow: 0 4px 16px rgba(0,0,0,.14);
}
.kpi-val { font-size: 2rem; font-weight: 800; line-height: 1.1; }
.kpi-lbl { font-size: .82rem; opacity: .9; margin-top: 2px; }
/* ── Role badge ─────────────────────────────────── */
.badge {
    display: inline-block;
    padding: 3px 13px;
    border-radius: 20px;
    font-size: .8rem;
    font-weight: 700;
    color: #fff;
}
.badge-doctor      { background:#1B5E20; }
.badge-nurse       { background:#00838F; }
.badge-receptionist{ background:#E65100; }
.badge-admin       { background:#4A148C; }
/* ── Alerts / notices ───────────────────────────── */
.notice-hipaa {
    background:#FFF8E1;
    border-left:4px solid #F9A825;
    border-radius:6px;
    padding:10px 14px;
    font-size:.88rem;
    margin:8px 0;
}
.notice-restricted {
    background:#FFEBEE;
    border-left:4px solid #C62828;
    border-radius:6px;
    padding:10px 14px;
    font-size:.88rem;
    margin:8px 0;
}
/* ── Access-denied ──────────────────────────────── */
.deny-box {
    background:#FFEBEE;
    border:2px solid #EF5350;
    border-radius:10px;
    padding:20px;
    text-align:center;
    color:#B71C1C;
    font-weight:700;
    font-size:1.05rem;
}
/* ── Tab overrides ──────────────────────────────── */
.stTabs [data-baseweb="tab"] { font-weight:600; }
/* ── Login card ─────────────────────────────────── */
.login-card {
    background:#fff;
    border-radius:14px;
    box-shadow:0 4px 24px rgba(0,0,0,.1);
    padding:30px 32px;
}
</style>
""", unsafe_allow_html=True)
# ══════════════════════════════════════════════════════════════════════════════
# RBAC  ─  Default Users & Permissions (seed data; mutable copies live in
#          st.session_state so Admins can add/edit/remove staff at runtime)
# ══════════════════════════════════════════════════════════════════════════════
DEFAULT_USERS: dict = {
    "dr.smith":       {"password":"doctor123",  "role":"Doctor",       "name":"Dr. Emily Smith",  "dept":"Cardiology"},
    "dr.jones":       {"password":"doctor456",  "role":"Doctor",       "name":"Dr. Michael Jones","dept":"Neurology"},
    "nurse.patel":    {"password":"nurse123",   "role":"Nurse",        "name":"Nurse Anita Patel","dept":"General Medicine"},
    "nurse.kim":      {"password":"nurse456",   "role":"Nurse",        "name":"Nurse David Kim",  "dept":"Emergency"},
    "receptionist1":  {"password":"recept123",  "role":"Receptionist", "name":"Sarah Connor",     "dept":"Front Desk"},
    "receptionist2":  {"password":"recept456",  "role":"Receptionist", "name":"Tom Bradley",      "dept":"Front Desk"},
    "admin":          {"password":"admin123",   "role":"Admin",        "name":"Randy Singh",      "dept":"Administration"},
}
DEFAULT_PERMISSIONS: dict = {
    "Doctor": {
        "view_records":True, "view_confidential":True,
        "schedule":True,     "edit_notes":True,
        "export":True,       "view_audit":False,
    },
    "Nurse": {
        "view_records":True, "view_confidential":True,
        "schedule":True,     "edit_notes":False,
        "export":False,      "view_audit":False,
    },
    "Receptionist": {
        "view_records":True, "view_confidential":False,
        "schedule":True,     "edit_notes":False,
        "export":True,       "view_audit":False,
    },
    "Admin": {
        "view_records":True, "view_confidential":True,
        "schedule":True,     "edit_notes":True,
        "export":True,       "view_audit":True,
    },
}
ROLE_LIST = ["Doctor", "Nurse", "Receptionist", "Admin"]
# ══════════════════════════════════════════════════════════════════════════════
# REFERENCE DATA
# ══════════════════════════════════════════════════════════════════════════════
DIAGNOSES = [
    "Hypertension","Type 2 Diabetes","Asthma","Coronary Artery Disease",
    "Depression","Anxiety Disorder","COPD","Osteoarthritis",
    "Chronic Kidney Disease","Hypothyroidism","Migraine","GERD",
    "Atrial Fibrillation","Heart Failure","Pneumonia","UTI",
    "Rheumatoid Arthritis","Chronic Pain","Obesity","Hyperlipidemia",
]
MEDICATIONS = [
    "Lisinopril 10mg","Metformin 500mg","Atorvastatin 20mg","Amlodipine 5mg",
    "Omeprazole 20mg","Levothyroxine 50mcg","Albuterol Inhaler","Sertraline 50mg",
    "Aspirin 81mg","Warfarin 5mg","Metoprolol 25mg","Losartan 50mg",
    "Gabapentin 300mg","Pantoprazole 40mg","Furosemide 20mg",
]
DEFAULT_DOCTORS = ["Dr. Emily Smith","Dr. Michael Jones","Dr. Sarah Lee","Dr. James Wilson","Dr. Maria Garcia"]
DEPARTMENTS  = ["Cardiology","Neurology","Orthopedics","General Medicine","Pediatrics","Oncology","Psychiatry","Emergency"]
APPT_TYPES   = ["Routine Checkup","Follow-up","Consultation","Emergency","Lab Results","Pre-op","Physical Exam","Telehealth"]
STATUSES     = ["Scheduled","Completed","Cancelled","No-show","Rescheduled"]
BLOOD_TYPES  = ["A+","A-","B+","B-","AB+","AB-","O+","O-"]
INSURERS     = ["BlueCross BlueShield","Aetna","United Healthcare","Cigna","Humana","Medicare","Medicaid","Kaiser Permanente"]
PUBLIC_COLS = [
    "Patient ID","Name","Date of Birth","Age","Gender","Blood Type",
    "Phone","Email","Insurance","Primary Diagnosis","Attending Physician",
    "Department","Admission Date","Status","# Visits",
]
CONFIDENTIAL_COLS = [
    "Medical Notes","Medications","Allergies",
    "Mental Health","HIV Status","Substance Use","Secondary Diagnosis",
]
# ══════════════════════════════════════════════════════════════════════════════
# DATA GENERATORS
# ══════════════════════════════════════════════════════════════════════════════
def _med_note() -> str:
    templates = [
        f"Patient presents with {random.choice(['mild','moderate','severe'])} {random.choice(DIAGNOSES).lower()} symptoms. "
        f"{random.choice(['Prescribed','Continued','Adjusted'])} {random.choice(MEDICATIONS)}. "
        f"Follow-up in {random.randint(2,12)} weeks.",
        f"BP {random.randint(120,185)}/{random.randint(80,115)} mmHg. "
        f"Medication adjusted. Patient counselled on diet and lifestyle.",
        f"Lab review: {random.choice(['HbA1c','Cholesterol','Creatinine','TSH','CBC'])} "
        f"{random.choice(['within range','slightly elevated','improving'])}. "
        f"Continue current plan.",
        f"Patient reports {random.choice(['improvement','no change','worsening'])}. "
        f"Ordered {random.choice(['ECG','Chest X-ray','MRI','Ultrasound','CT Scan'])}. "
        f"Refer to {random.choice(DEPARTMENTS)} if no improvement.",
    ]
    return random.choice(templates)
def generate_patients(n: int, doctors: list = None) -> pd.DataFrame:
    fake_l = Faker(); Faker.seed(42); random.seed(42)
    doc_list = doctors if doctors else DEFAULT_DOCTORS
    rows = []
    for i in range(n):
        dob = fake_l.date_of_birth(minimum_age=5, maximum_age=90)
        age = (date.today() - dob).days // 365
        rows.append({
            "Patient ID":          f"PAT-{10000+i:05d}",
            "Name":                fake_l.name(),
            "Date of Birth":       str(dob),
            "Age":                 age,
            "Gender":              random.choice(["Male","Female","Non-binary"]),
            "Blood Type":          random.choice(BLOOD_TYPES),
            "Phone":               fake_l.phone_number(),
            "Email":               fake_l.email(),
            "Address":             fake_l.address().replace("\n", ", "),
            "Insurance":           random.choice(INSURERS),
            "Primary Diagnosis":   random.choice(DIAGNOSES),
            "Secondary Diagnosis": random.choice(DIAGNOSES + ["None"]),
            "Attending Physician": random.choice(doc_list),
            "Department":          random.choice(DEPARTMENTS),
            "Admission Date":      str(fake_l.date_between(start_date="-2y", end_date="today")),
            "Status":              random.choice(["Active","Discharged","Follow-up","Critical"]),
            "# Visits":            random.randint(1, 30),
            # ── confidential ──────────────────────────────────────────────
            "Medical Notes":  _med_note(),
            "Medications":    ", ".join(random.sample(MEDICATIONS, k=random.randint(1,4))),
            "Allergies":      random.choice(["None known","Penicillin","Sulfa drugs","Aspirin","Latex","Codeine","NSAIDs"]),
            "Mental Health":  random.choice(["No concerns","Anxiety","Depression","PTSD","Bipolar","N/A"]),
            "HIV Status":     random.choice(["Negative","Negative","Negative","Positive","Unknown"]),
            "Substance Use":  random.choice(["None","Tobacco","Alcohol","Cannabis","N/A"]),
        })
    return pd.DataFrame(rows)
def generate_appointments(n: int = 120, doctors: list = None) -> pd.DataFrame:
    fake_l = Faker(); Faker.seed(123); random.seed(123)
    doc_list = doctors if doctors else DEFAULT_DOCTORS
    rows = []
    for i in range(n):
        rows.append({
            "Appointment ID": f"APT-{20000+i:05d}",
            "Patient Name":   fake_l.name(),
            "Patient ID":     f"PAT-{random.randint(10000,10500):05d}",
            "Doctor":         random.choice(doc_list),
            "Department":     random.choice(DEPARTMENTS),
            "Date":           str(fake_l.date_between(start_date="-30d", end_date="+30d")),
            "Time":           f"{random.randint(8,17):02d}:{random.choice([0,15,30,45]):02d}",
            "Type":           random.choice(APPT_TYPES),
            "Status":         random.choice(STATUSES),
            "Duration (min)": random.choice([15,30,45,60]),
            "Room":           f"Room {random.randint(100,350)}",
            "Notes":          random.choice(["","Interpreter needed","Wheelchair room required","ER follow-up",""]),
        })
    return pd.DataFrame(rows)
# ══════════════════════════════════════════════════════════════════════════════
# EXPORT HELPERS
# ══════════════════════════════════════════════════════════════════════════════
def _export_json(df: pd.DataFrame, stem: str):
    raw = df.to_json(orient="records", indent=2, date_format="iso")
    return raw.encode(), f"{stem}.json", "application/json"
def _export_text(df: pd.DataFrame, stem: str, title: str, role: str):
    ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    sep = "=" * 80
    lines = [
        sep,
        f"  HealthID Pro  |  {title}",
        f"  Developed by Randy Singh | Kalsnet (KNet) Consulting",
        f"  Exported: {ts}  |  Role: {role}  |  Records: {len(df)}",
        sep, "",
    ]
    for idx, (_, row) in enumerate(df.iterrows(), 1):
        lines.append(f"── Record #{idx} ──")
        for col, val in row.items():
            lines.append(f"   {col}: {val}")
        lines.append("")
    return "\n".join(lines).encode(), f"{stem}.txt", "text/plain"
def _export_pdf(df: pd.DataFrame, stem: str, title: str, role: str):
    if not REPORTLAB_OK:
        return None, None, None
    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter,
                            leftMargin=.5*inch, rightMargin=.5*inch,
                            topMargin=.7*inch, bottomMargin=.7*inch)
    styles = getSampleStyleSheet()
    t_style  = ParagraphStyle("HTitle", parent=styles["Title"],
                               textColor=rl_colors.HexColor("#002B7F"),
                               fontSize=18, spaceAfter=4)
    st_style = ParagraphStyle("HSub", parent=styles["Normal"],
                               textColor=rl_colors.HexColor("#1565C0"),
                               fontSize=10, fontName="Helvetica-Bold", spaceAfter=8)
    story = [
        Paragraph("🏥  HealthID Pro — Healthcare Identity Management", t_style),
        Paragraph("Developed by Randy Singh | Kalsnet (KNet) Consulting", st_style),
        HRFlowable(width="100%", thickness=1, color=rl_colors.HexColor("#1565C0")),
        Spacer(1, 8),
        Paragraph(f"<b>{title}</b>", styles["Heading2"]),
        Paragraph(f"Role: <b>{role}</b>  |  Exported: {datetime.now():%Y-%m-%d %H:%M:%S}  |  Records: {len(df)}",
                  styles["Normal"]),
        Spacer(1, 12),
    ]
    cols = df.columns.tolist()[:10]
    col_w = (letter[0] - inch) / len(cols)
    tdata = [cols] + [[str(row[c])[:22] for c in cols] for _, row in df.head(50).iterrows()]
    tbl = Table(tdata, colWidths=[col_w]*len(cols), repeatRows=1)
    tbl.setStyle(TableStyle([
        ("BACKGROUND",  (0,0),(-1,0), rl_colors.HexColor("#002B7F")),
        ("TEXTCOLOR",   (0,0),(-1,0), rl_colors.white),
        ("FONTNAME",    (0,0),(-1,0), "Helvetica-Bold"),
        ("FONTSIZE",    (0,0),(-1,-1),7),
        ("ALIGN",       (0,0),(-1,-1),"LEFT"),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),[rl_colors.white, rl_colors.HexColor("#EEF4FF")]),
        ("GRID",        (0,0),(-1,-1),.3, rl_colors.grey),
        ("PADDING",     (0,0),(-1,-1),4),
    ]))
    story.append(tbl)
    doc.build(story)
    buf.seek(0)
    return buf.read(), f"{stem}.pdf", "application/pdf"
def _export_word(df: pd.DataFrame, stem: str, title: str, role: str):
    if not DOCX_OK:
        return None, None, None
    doc = Document()
    h = doc.add_heading("HealthID Pro — Healthcare Identity Management", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 43, 127)
    s = doc.add_paragraph("Developed by Randy Singh | Kalsnet (KNet) Consulting")
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in s.runs:
        run.font.bold = True; run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(21,101,192)
    doc.add_heading(title, 1)
    doc.add_paragraph(
        f"Role: {role}  |  Exported: {datetime.now():%Y-%m-%d %H:%M:%S}  |  Records: {len(df)}"
    )
    doc.add_paragraph()
    cols = df.columns.tolist()[:10]
    tbl = doc.add_table(rows=1, cols=len(cols)); tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, c in enumerate(cols):
        hdr[i].text = c
        run = hdr[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(0,43,127)
    for _, row in df.head(50).iterrows():
        cells = tbl.add_row().cells
        for i, c in enumerate(cols):
            cells[i].text = str(row[c])[:30]
    buf = BytesIO(); doc.save(buf); buf.seek(0)
    return buf.read(), f"{stem}.docx", \
           "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
# ══════════════════════════════════════════════════════════════════════════════
# SESSION STATE
# ══════════════════════════════════════════════════════════════════════════════
def _init_state():
    defaults = {
        "authenticated": False,
        "user": None,
        "patient_df": None,
        "appt_df": generate_appointments(),
        "n_synthetic": 50,
        "audit_log": [],
        "users_db": {k: dict(v) for k, v in DEFAULT_USERS.items()},
        "permissions_db": {k: dict(v) for k, v in DEFAULT_PERMISSIONS.items()},
        "doctors_list": list(DEFAULT_DOCTORS),
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v
_init_state()
def _log_audit(action: str):
    user = st.session_state.user
    if user:
        st.session_state.audit_log.append({
            "Timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "User": user["name"],
            "Role": user["role"],
            "Action": action,
            "IP": f"10.{random.randint(0,255)}.{random.randint(0,255)}.{random.randint(1,254)}",
            "Status": "Success",
        })
# ══════════════════════════════════════════════════════════════════════════════
# UI COMPONENTS
# ══════════════════════════════════════════════════════════════════════════════
def show_header():
    st.markdown("""
    <div class="header-wrap">
        <p class="app-title">🏥 HealthID Pro — Healthcare Identity Management System</p>
        <p class="app-subtitle">Developed By Randy Singh from Kalsnet (KNet) Consulting</p>
    </div>""", unsafe_allow_html=True)
def section_header(text: str):
    """Renders a prominent, high-visibility title bar for a page/section."""
    st.markdown(f'<div class="section-header">{text}</div>', unsafe_allow_html=True)
def subsection_header(text: str):
    """Renders a lighter-weight highlighted bar for sub-sections within a page."""
    st.markdown(f'<div class="subsection-header">{text}</div>', unsafe_allow_html=True)
def _kpi(col, label: str, val, color: str):
    col.markdown(
        f'<div class="kpi-card" style="background:linear-gradient(135deg,{color},{color}bb)">'
        f'<div class="kpi-lbl">{label}</div>'
        f'<div class="kpi-val">{val}</div></div>',
        unsafe_allow_html=True,
    )
# ══════════════════════════════════════════════════════════════════════════════
# LOGIN
# ══════════════════════════════════════════════════════════════════════════════
def page_login():
    show_header()
    _, col, _ = st.columns([1, 1.4, 1])
    with col:
        st.markdown('<div class="login-card">', unsafe_allow_html=True)
        st.markdown("### 🔐 Secure Staff Login")
        uname = st.text_input("Username", placeholder="e.g. dr.smith")
        pwd   = st.text_input("Password", type="password")
        c1, c2 = st.columns(2)
        with c1:
            if st.button("Login", use_container_width=True, type="primary"):
                users_db = st.session_state.users_db
                if uname in users_db and users_db[uname]["password"] == pwd:
                    st.session_state.authenticated = True
                    st.session_state.user = {**users_db[uname], "username": uname}
                    _log_audit("User logged in")
                    st.rerun()
                else:
                    st.error("❌ Invalid credentials.")
        with c2:
            if st.button("Show Demo Creds", use_container_width=True):
                st.info(
                    "**Doctor:** `dr.smith` / `doctor123`\n\n"
                    "**Nurse:** `nurse.patel` / `nurse123`\n\n"
                    "**Receptionist:** `receptionist1` / `recept123`\n\n"
                    "**Admin:** `admin` / `admin123`"
                )
        st.markdown("</div>", unsafe_allow_html=True)
        st.markdown("---")
        st.markdown("**Quick-reference credentials**")
        cred_rows = [
            {"Role": v["role"], "Username": u, "Password": v["password"]}
            for u, v in st.session_state.users_db.items()
        ]
        st.table(pd.DataFrame(cred_rows))
# ══════════════════════════════════════════════════════════════════════════════
# SIDEBAR
# ══════════════════════════════════════════════════════════════════════════════
def build_sidebar() -> str:
    role  = st.session_state.user["role"]
    perms = st.session_state.permissions_db[role]
    icons = {"Doctor":"🟢","Nurse":"🔵","Receptionist":"🟠","Admin":"🟣"}
    with st.sidebar:
        st.markdown("## 🏥 HealthID Pro")
        st.divider()
        user = st.session_state.user
        st.markdown(f"**{icons.get(role,'⚪')} {user['name']}**\n\n*{role} · {user['dept']}*")
        st.divider()
        pages = ["📊 Dashboard", "👤 Patient Records", "📅 Appointments",
                 "📈 Analytics", "⬆️ Data Management", "📤 Export"]
        if role == "Admin":
            pages.append("🧑‍⚕️ User Management")
        selected = st.radio("Navigation", pages, label_visibility="collapsed")
        st.divider()
        st.markdown("**🔑 Role Permissions**")
        perm_labels = [
            ("view_records",      "View Patient Records"),
            ("view_confidential", "View Medical Notes"),
            ("schedule",          "Schedule Appointments"),
            ("edit_notes",        "Edit Medical Notes"),
            ("export",            "Export Data"),
            ("view_audit",        "View Audit Logs"),
        ]
        for key, label in perm_labels:
            icon = "✅" if perms.get(key) else "🚫"
            st.markdown(f"{icon} {label}")
        st.divider()
        if st.button("🚪 Logout", use_container_width=True):
            for k in ["authenticated","user","patient_df"]:
                st.session_state[k] = False if k == "authenticated" else None
            st.rerun()
    return selected
# ══════════════════════════════════════════════════════════════════════════════
# DASHBOARD
# ══════════════════════════════════════════════════════════════════════════════
def page_dashboard():
    section_header("📊 Dashboard Overview")
    df   = st.session_state.patient_df
    appt = st.session_state.appt_df
    n_pat   = len(df)   if df   is not None else 0
    n_appt  = len(appt) if appt is not None else 0
    n_sched = len(appt[appt.Status=="Scheduled"])   if appt is not None else 0
    n_done  = len(appt[appt.Status=="Completed"])   if appt is not None else 0
    c1,c2,c3,c4 = st.columns(4)
    _kpi(c1, "👥 Total Patients",      n_pat,  "#002B7F")
    _kpi(c2, "📅 Total Appointments",  n_appt, "#00796B")
    _kpi(c3, "🗓 Scheduled",           n_sched,"#E65100")
    _kpi(c4, "✅ Completed",           n_done, "#4A148C")
    st.markdown("<br>", unsafe_allow_html=True)
    if df is not None and not df.empty:
        r1c1, r1c2 = st.columns(2)
        with r1c1:
            dc = df["Department"].value_counts()
            fig = px.pie(values=dc.values, names=dc.index,
                         title="🏥 Patients by Department",
                         color_discrete_sequence=px.colors.qualitative.Bold)
            fig.update_layout(height=320, margin=dict(t=40,b=10))
            st.plotly_chart(fig, use_container_width=True)
        with r1c2:
            fig = px.histogram(df, x="Age", nbins=20,
                               title="👥 Patient Age Distribution",
                               color_discrete_sequence=["#1565C0"])
            fig.update_layout(height=320, margin=dict(t=40,b=10),
                              xaxis_title="Age", yaxis_title="Count")
            st.plotly_chart(fig, use_container_width=True)
        r2c1, r2c2 = st.columns(2)
        with r2c1:
            dc2 = df["Primary Diagnosis"].value_counts().head(8)
            fig = px.bar(x=dc2.values, y=dc2.index, orientation="h",
                         title="🩺 Top 8 Diagnoses",
                         color=dc2.values,
                         color_continuous_scale="Blues")
            fig.update_layout(height=320, showlegend=False,
                              yaxis_title="", xaxis_title="Count")
            st.plotly_chart(fig, use_container_width=True)
        with r2c2:
            gc = df["Gender"].value_counts()
            fig = px.pie(values=gc.values, names=gc.index,
                         title="⚧ Gender Distribution",
                         color_discrete_map={
                             "Male":"#1565C0","Female":"#E91E63","Non-binary":"#9C27B0"
                         })
            fig.update_layout(height=320, margin=dict(t=40,b=10))
            st.plotly_chart(fig, use_container_width=True)
    if appt is not None and not appt.empty:
        st.markdown("**📅 Appointment Status Overview**")
        sc = appt["Status"].value_counts()
        cmap = {"Scheduled":"#1565C0","Completed":"#2E7D32",
                "Cancelled":"#C62828","No-show":"#F57F17","Rescheduled":"#7B1FA2"}
        fig = px.bar(x=sc.index, y=sc.values, color=sc.index,
                     color_discrete_map=cmap,
                     title="Appointments by Status")
        fig.update_layout(height=260, showlegend=False,
                          xaxis_title="Status", yaxis_title="Count")
        st.plotly_chart(fig, use_container_width=True)
    if df is None:
        st.info("💡 No patient data yet. Head to **⬆️ Data Management** to generate or upload records.")
# ══════════════════════════════════════════════════════════════════════════════
# PATIENT RECORDS
# ══════════════════════════════════════════════════════════════════════════════
def page_records():
    section_header("👤 Patient Records")
    df   = st.session_state.patient_df
    role = st.session_state.user["role"]
    can_see_conf = st.session_state.permissions_db[role]["view_confidential"]
    if df is None or df.empty:
        st.warning("⚠️ No patient data. Go to **⬆️ Data Management** to load records.")
        return
    _log_audit("Viewed patient records")
    # ── Search & filter row ────────────────────────────────────────────────
    fc1, fc2, fc3 = st.columns([2,1,1])
    with fc1:
        q = st.text_input("🔍 Search patients", placeholder="Name, ID, diagnosis …")
    with fc2:
        dept_f = st.selectbox("Department", ["All"] + DEPARTMENTS)
    with fc3:
        stat_f = st.selectbox("Status", ["All","Active","Discharged","Follow-up","Critical"])
    fdf = df.copy()
    if q:
        mask = fdf.apply(lambda r: r.astype(str).str.contains(q, case=False).any(), axis=1)
        fdf  = fdf[mask]
    if dept_f != "All": fdf = fdf[fdf["Department"] == dept_f]
    if stat_f != "All": fdf = fdf[fdf["Status"]     == stat_f]
    st.markdown(f"**{len(fdf):,} of {len(df):,} records**")
    if can_see_conf:
        # ── Doctor / Admin view ────────────────────────────────────────────
        st.markdown('<div class="notice-hipaa">⚠️ <b>HIPAA Notice</b>: '
                    'You are viewing confidential medical data. Handle per institutional policy.</div>',
                    unsafe_allow_html=True)
        t1, t2 = st.tabs(["📋 Patient Information", "🔒 Confidential Records"])
        with t1:
            st.dataframe(fdf[PUBLIC_COLS].reset_index(drop=True),
                         use_container_width=True, height=420)
        with t2:
            conf_cols_available = [c for c in CONFIDENTIAL_COLS if c in fdf.columns]
            st.dataframe(fdf[["Patient ID","Name"] + conf_cols_available].reset_index(drop=True),
                         use_container_width=True, height=420)
            if can_see_conf and st.session_state.permissions_db[role]["edit_notes"]:
                st.markdown("---")
                st.markdown("**✏️ Edit Medical Note**")
                pid = st.text_input("Patient ID to update notes", placeholder="PAT-10000")
                new_note = st.text_area("Updated medical note")
                if st.button("💾 Save Note"):
                    mask2 = st.session_state.patient_df["Patient ID"] == pid
                    if mask2.any():
                        st.session_state.patient_df.loc[mask2, "Medical Notes"] = new_note
                        _log_audit(f"Edited medical note for {pid}")
                        st.success(f"✅ Note updated for {pid}")
                    else:
                        st.error("Patient ID not found.")
    else:
        # ── Receptionist view ──────────────────────────────────────────────
        st.markdown('<div class="notice-restricted">🔒 <b>Access Restricted</b>: '
                    'Confidential medical information is hidden for your role (Receptionist). '
                    'This complies with HIPAA Privacy Rule §164.502.</div>',
                    unsafe_allow_html=True)
        pub_available = [c for c in PUBLIC_COLS if c in fdf.columns]
        st.dataframe(fdf[pub_available].reset_index(drop=True),
                     use_container_width=True, height=440)
        with st.expander("🔒 Hidden / Restricted Fields"):
            for c in CONFIDENTIAL_COLS:
                st.markdown(f"🚫 **{c}** — *Requires Doctor or Admin role*")
# ══════════════════════════════════════════════════════════════════════════════
# APPOINTMENTS
# ══════════════════════════════════════════════════════════════════════════════
def page_appointments():
    section_header("📅 Appointment Management")
    appt = st.session_state.appt_df
    role = st.session_state.user["role"]
    t1, t2 = st.tabs(["📋 View Appointments", "➕ Schedule New"])
    with t1:
        fc1, fc2, fc3 = st.columns(3)
        with fc1: date_f = st.date_input("Filter by date", value=None)
        with fc2: doc_f  = st.selectbox("Doctor", ["All"] + st.session_state.doctors_list)
        with fc3: stat_f = st.selectbox("Status", ["All"] + STATUSES)
        fdf = appt.copy()
        if date_f: fdf = fdf[fdf["Date"] == str(date_f)]
        if doc_f  != "All": fdf = fdf[fdf["Doctor"] == doc_f]
        if stat_f != "All": fdf = fdf[fdf["Status"] == stat_f]
        st.markdown(f"**{len(fdf):,} appointments**")
        STATUS_CSS = {
            "Scheduled":  "background-color:#E3F2FD;color:#1565C0",
            "Completed":  "background-color:#E8F5E9;color:#2E7D32",
            "Cancelled":  "background-color:#FFEBEE;color:#C62828",
            "No-show":    "background-color:#FFF3E0;color:#E65100",
            "Rescheduled":"background-color:#F3E5F5;color:#7B1FA2",
        }
        styled = fdf.reset_index(drop=True).style.map(
            lambda v: STATUS_CSS.get(v,""), subset=["Status"]
        )
        st.dataframe(styled, use_container_width=True, height=420)
    with t2:
        if not st.session_state.permissions_db[role]["schedule"]:
            st.markdown('<div class="deny-box">🚫 You do not have permission to schedule appointments.</div>',
                        unsafe_allow_html=True)
            return
        c1, c2 = st.columns(2)
        with c1:
            p_name = st.text_input("Patient Name *")
            p_id   = st.text_input("Patient ID", placeholder="PAT-XXXXX")
            doctor = st.selectbox("Doctor *", st.session_state.doctors_list)
            dept   = st.selectbox("Department", DEPARTMENTS)
        with c2:
            a_date = st.date_input("Appointment Date *", min_value=date.today())
            a_time = st.time_input("Time *", value=datetime.strptime("09:00","%H:%M").time())
            a_type = st.selectbox("Type", APPT_TYPES)
            dur    = st.selectbox("Duration (min)", [15,30,45,60])
        notes = st.text_area("Notes (optional)")
        if st.button("✅ Schedule Appointment", type="primary"):
            if p_name and doctor:
                new_row = {
                    "Appointment ID": f"APT-{random.randint(29000,39999):05d}",
                    "Patient Name": p_name, "Patient ID": p_id or "N/A",
                    "Doctor": doctor, "Department": dept,
                    "Date": str(a_date), "Time": a_time.strftime("%H:%M"),
                    "Type": a_type, "Status": "Scheduled",
                    "Duration (min)": dur, "Room": f"Room {random.randint(100,350)}",
                    "Notes": notes,
                }
                st.session_state.appt_df = pd.concat(
                    [st.session_state.appt_df, pd.DataFrame([new_row])], ignore_index=True
                )
                _log_audit(f"Scheduled appointment for {p_name}")
                st.success(f"✅ Appointment scheduled: **{p_name}** with **{doctor}** on **{a_date}** at **{a_time.strftime('%H:%M')}**")
            else:
                st.error("Please provide Patient Name and Doctor.")
# ══════════════════════════════════════════════════════════════════════════════
# ANALYTICS
# ══════════════════════════════════════════════════════════════════════════════
def page_analytics():
    section_header("📈 Analytics & Insights")
    df   = st.session_state.patient_df
    appt = st.session_state.appt_df
    role = st.session_state.user["role"]
    if df is None or df.empty:
        st.warning("No patient data. Please load data first.")
        return
    t1, t2, t3 = st.tabs(["👥 Patient Analytics", "📅 Appointment Analytics", "🔐 Audit Log"])
    with t1:
        c1, c2 = st.columns(2)
        with c1:
            ic = df["Insurance"].value_counts()
            fig = px.bar(x=ic.values, y=ic.index, orientation="h",
                         title="Insurance Provider Distribution",
                         color_discrete_sequence=["#002B7F"])
            fig.update_layout(height=320, yaxis_title="", xaxis_title="Patients")
            st.plotly_chart(fig, use_container_width=True)
        with c2:
            bc = df["Blood Type"].value_counts()
            fig = px.pie(values=bc.values, names=bc.index,
                         title="Blood Type Distribution",
                         color_discrete_sequence=px.colors.qualitative.Set2)
            fig.update_layout(height=320)
            st.plotly_chart(fig, use_container_width=True)
        c3, c4 = st.columns(2)
        with c3:
            fig = px.histogram(df, x="# Visits", nbins=15,
                               title="Visit Frequency Distribution",
                               color_discrete_sequence=["#00796B"])
            fig.update_layout(height=300)
            st.plotly_chart(fig, use_container_width=True)
        with c4:
            sc = df["Status"].value_counts()
            fig = px.pie(values=sc.values, names=sc.index,
                         title="Patient Status Breakdown",
                         color_discrete_sequence=["#1565C0","#2E7D32","#F57F17","#C62828"])
            fig.update_layout(height=300)
            st.plotly_chart(fig, use_container_width=True)
        if st.session_state.permissions_db[role]["view_confidential"]:
            st.markdown("---")
            st.markdown("**🔬 Clinical Insights** *(Medical Staff Only)*")
            c5, c6 = st.columns(2)
            with c5:
                fig = px.scatter(df, x="Age", y="# Visits", color="Department",
                                 title="Age vs Number of Visits",
                                 color_discrete_sequence=px.colors.qualitative.Bold)
                st.plotly_chart(fig, use_container_width=True)
            with c6:
                all_meds: list = []
                for m in df["Medications"].dropna():
                    all_meds.extend([x.strip() for x in m.split(",")])
                ms = pd.Series(all_meds).value_counts().head(10)
                fig = px.bar(x=ms.values, y=ms.index, orientation="h",
                             title="Top 10 Prescribed Medications",
                             color_discrete_sequence=["#7B1FA2"])
                fig.update_layout(height=320, yaxis_title="", xaxis_title="Count")
                st.plotly_chart(fig, use_container_width=True)
    with t2:
        c1, c2 = st.columns(2)
        with c1:
            da = appt["Department"].value_counts()
            fig = px.bar(x=da.index, y=da.values, title="Appointments by Department",
                         color_discrete_sequence=["#1565C0"])
            fig.update_layout(height=300, xaxis_tickangle=35)
            st.plotly_chart(fig, use_container_width=True)
        with c2:
            ta = appt["Type"].value_counts()
            fig = px.pie(values=ta.values, names=ta.index,
                         title="Appointment Types",
                         color_discrete_sequence=px.colors.qualitative.Pastel)
            fig.update_layout(height=300)
            st.plotly_chart(fig, use_container_width=True)
        dl = appt.groupby("Doctor")["Appointment ID"].count().sort_values(ascending=False)
        fig = px.bar(x=dl.values, y=dl.index, orientation="h",
                     title="Doctor Workload (Total Appointments)",
                     color=dl.values, color_continuous_scale="Blues")
        fig.update_layout(showlegend=False, height=280, yaxis_title="", xaxis_title="Appointments")
        st.plotly_chart(fig, use_container_width=True)
    with t3:
        if st.session_state.permissions_db[role]["view_audit"]:
            logs = st.session_state.audit_log
            if logs:
                log_df = pd.DataFrame(logs[::-1])
                st.dataframe(log_df, use_container_width=True, height=420)
            else:
                st.info("No audit events recorded yet.")
        else:
            st.markdown('<div class="deny-box">🚫 Audit logs are restricted to Administrators.</div>',
                        unsafe_allow_html=True)
# ══════════════════════════════════════════════════════════════════════════════
# DATA MANAGEMENT
# ══════════════════════════════════════════════════════════════════════════════
def page_data_mgmt():
    section_header("⬆️ Data Management")
    t1, t2 = st.tabs(["🔢 Synthetic Data Generator", "📁 Upload Real Data"])
    with t1:
        st.markdown("**Generate synthetic patient records for testing & demos.**")
        n = st.slider(
            "Number of Records to Generate",
            min_value=0, max_value=500,
            value=st.session_state.n_synthetic, step=10,
        )
        pct = n / 500
        # Visual progress bar
        st.markdown(f"""
        <div style="background:#E3F2FD;border-radius:8px;padding:3px;margin:6px 0 2px 0">
          <div style="background:linear-gradient(90deg,#002B7F,#1565C0,#0288D1);
                      border-radius:6px;height:22px;width:{max(pct*100,2):.1f}%;
                      display:flex;align-items:center;justify-content:center;
                      color:#fff;font-size:.8rem;font-weight:700;min-width:28px">
            {n}
          </div>
        </div>
        <p style="color:#555;font-size:.83rem;margin:0">{n} / 500 records selected</p>
        """, unsafe_allow_html=True)
        st.markdown("")
        if st.button("🔄 Generate Records", type="primary"):
            if n > 0:
                with st.spinner(f"Generating {n} synthetic patient records …"):
                    st.session_state.patient_df = generate_patients(n, doctors=st.session_state.doctors_list)
                    st.session_state.n_synthetic = n
                _log_audit(f"Generated {n} synthetic patient records")
                st.success(f"✅ {n:,} synthetic records generated!")
            else:
                st.warning("Select at least 1 record.")
        # ── Full dataset preview with scrollbar (shows ALL rows, not just a sample) ──
        if st.session_state.patient_df is not None and not st.session_state.patient_df.empty:
            st.markdown("---")
            subsection_header(
                f"📋 Full Dataset Preview — {len(st.session_state.patient_df):,} records "
                f"(scroll inside the table to view every row)"
            )
            st.dataframe(st.session_state.patient_df, use_container_width=True, height=520)
    with t2:
        st.markdown("**Upload a real patient CSV file (HIPAA-compliant environments only).**")
        st.info(
            "**Minimum required columns:** Patient ID, Name, Age, Gender, Primary Diagnosis, "
            "Attending Physician, Department, Status.\n\n"
            "Optional confidential columns: Medical Notes, Medications, Allergies, Mental Health, HIV Status, Substance Use."
        )
        subsection_header("📥 Sample Real-Data File")
        st.markdown(
            "Don't have a real dataset handy? Download a ready-made **400-record** sample "
            "file below (matches the exact upload schema), then re-upload it using the "
            "file picker underneath to try the full upload flow."
        )
        if "sample_400_df" not in st.session_state:
            st.session_state.sample_400_df = generate_patients(400, doctors=st.session_state.doctors_list)
        sample_csv = st.session_state.sample_400_df.to_csv(index=False).encode()
        st.download_button(
            "📥 Download Sample 400-Patient CSV", data=sample_csv,
            file_name="sample_400_patients.csv", mime="text/csv",
            use_container_width=True,
        )
        st.markdown("---")
        file = st.file_uploader("Choose CSV file", type=["csv"])
        if file:
            try:
                up_df = pd.read_csv(file)
                for c in CONFIDENTIAL_COLS:
                    if c not in up_df.columns:
                        up_df[c] = "N/A"
                st.success(f"✅ **{file.name}** — {len(up_df):,} rows, {len(up_df.columns)} columns")
                st.dataframe(up_df.head(), use_container_width=True)
                if st.button("✅ Load into System", type="primary"):
                    st.session_state.patient_df = up_df
                    _log_audit(f"Uploaded real data: {file.name} ({len(up_df)} records)")
                    st.success("Data loaded. Navigate to Patient Records to view.")
                    st.rerun()
            except Exception as exc:
                st.error(f"❌ Error reading file: {exc}")
# ══════════════════════════════════════════════════════════════════════════════
# EXPORT
# ══════════════════════════════════════════════════════════════════════════════
def page_export():
    section_header("📤 Export Data")
    df   = st.session_state.patient_df
    appt = st.session_state.appt_df
    role = st.session_state.user["role"]
    choice = st.radio("What to export?",
                      ["Patient Records","Appointments","Both (summary)"],
                      horizontal=True)
    # Apply role filtering
    if df is not None and not st.session_state.permissions_db[role]["view_confidential"]:
        pub_available = [c for c in PUBLIC_COLS if c in df.columns]
        export_pat = df[pub_available].copy()
        st.info("ℹ️ Confidential fields are excluded based on your role permissions.")
    else:
        export_pat = df.copy() if df is not None else pd.DataFrame()
    if choice == "Patient Records":
        export_df = export_pat; title = "Patient Records"; stem = "patient_records"
    elif choice == "Appointments":
        export_df = appt.copy(); title = "Appointment Records"; stem = "appointments"
    else:
        export_df = export_pat; title = "Healthcare Summary Report"; stem = "healthcare_summary"
    n_rec = len(export_df) if not export_df.empty else 0
    st.markdown(f"**{title}** — {n_rec:,} records ready for export")
    st.divider()
    c1, c2, c3, c4 = st.columns(4)
    with c1:
        st.markdown("#### 📄 PDF")
        if not REPORTLAB_OK:
            st.warning("Install: `pip install reportlab`")
        elif export_df.empty:
            st.warning("No data to export.")
        else:
            data, fname, mime = _export_pdf(export_df, stem, title, role)
            if data:
                st.download_button("⬇️ Download PDF", data=data,
                                   file_name=fname, mime=mime, use_container_width=True)
    with c2:
        st.markdown("#### 📝 Word")
        if not DOCX_OK:
            st.warning("Install: `pip install python-docx`")
        elif export_df.empty:
            st.warning("No data to export.")
        else:
            data, fname, mime = _export_word(export_df, stem, title, role)
            if data:
                st.download_button("⬇️ Download Word", data=data,
                                   file_name=fname, mime=mime, use_container_width=True)
    with c3:
        st.markdown("#### 📋 Text")
        if export_df.empty:
            st.warning("No data to export.")
        else:
            data, fname, mime = _export_text(export_df, stem, title, role)
            st.download_button("⬇️ Download Text", data=data,
                               file_name=fname, mime=mime, use_container_width=True)
    with c4:
        st.markdown("#### 🔧 JSON")
        if export_df.empty:
            st.warning("No data to export.")
        else:
            data, fname, mime = _export_json(export_df, stem)
            st.download_button("⬇️ Download JSON", data=data,
                               file_name=fname, mime=mime, use_container_width=True)
    if not export_df.empty:
        st.divider()
        st.markdown("**Preview (first 5 rows):**")
        st.dataframe(export_df.head(5), use_container_width=True)
# ══════════════════════════════════════════════════════════════════════════════
# USER & ROLE MANAGEMENT  (Admin only)
# ══════════════════════════════════════════════════════════════════════════════
def page_user_management():
    section_header("🧑‍⚕️ User & Role Management")
    role = st.session_state.user["role"]
    if role != "Admin":
        st.markdown('<div class="deny-box">🚫 User management is restricted to Administrators.</div>',
                    unsafe_allow_html=True)
        return
    users_db = st.session_state.users_db
    perms_db = st.session_state.permissions_db
    badge_cls = {"Doctor":"badge-doctor","Nurse":"badge-nurse",
                 "Receptionist":"badge-receptionist","Admin":"badge-admin"}
    subsection_header(f"👥 Current Staff Accounts — {len(users_db)} total")
    udf = pd.DataFrame([
        {"Username": u, "Full Name": v["name"], "Role": v["role"], "Department": v["dept"]}
        for u, v in users_db.items()
    ])
    st.dataframe(udf, use_container_width=True, height=min(360, 80 + 35*len(udf)))
    t1, t2, t3, t4 = st.tabs(["➕ Add Staff", "✏️ Edit Staff", "🗑️ Remove Staff", "🔑 Role Permissions"])
    # ── Add ──────────────────────────────────────────────────────────────
    with t1:
        st.markdown("Add a new **Doctor**, **Nurse**, **Receptionist**, or **Admin** account.")
        c1, c2 = st.columns(2)
        with c1:
            new_username = st.text_input("Username *", key="nu_user", placeholder="e.g. dr.patel")
            new_name     = st.text_input("Full Name *", key="nu_name", placeholder="e.g. Dr. Raj Patel")
            new_role     = st.selectbox("Role *", ROLE_LIST, key="nu_role")
        with c2:
            new_dept     = st.text_input("Department *", key="nu_dept", placeholder="e.g. Cardiology")
            new_password = st.text_input("Temporary Password *", key="nu_pwd", type="password")
        if st.button("✅ Create Account", type="primary", key="nu_create"):
            if not (new_username and new_name and new_dept and new_password):
                st.error("Please fill in all required fields.")
            elif new_username in users_db:
                st.error(f"Username '{new_username}' already exists.")
            else:
                users_db[new_username] = {
                    "password": new_password, "role": new_role,
                    "name": new_name, "dept": new_dept,
                }
                if new_role == "Doctor" and new_name not in st.session_state.doctors_list:
                    st.session_state.doctors_list.append(new_name)
                _log_audit(f"Created new {new_role} account: {new_username} ({new_name})")
                st.success(f"✅ {new_role} account created for **{new_name}** (`{new_username}`).")
                st.rerun()
    # ── Edit ─────────────────────────────────────────────────────────────
    with t2:
        if users_db:
            sel = st.selectbox("Select account to edit", list(users_db.keys()), key="eu_sel")
            cur = users_db[sel]
            c1, c2 = st.columns(2)
            with c1:
                e_name = st.text_input("Full Name", value=cur["name"], key="eu_name")
                e_role = st.selectbox("Role", ROLE_LIST, index=ROLE_LIST.index(cur["role"]), key="eu_role")
            with c2:
                e_dept = st.text_input("Department", value=cur["dept"], key="eu_dept")
                e_pwd  = st.text_input("Reset Password (leave blank to keep current)", key="eu_pwd", type="password")
            if st.button("💾 Save Changes", type="primary", key="eu_save"):
                old_name = cur["name"]
                old_role = cur["role"]
                users_db[sel]["name"] = e_name
                users_db[sel]["role"] = e_role
                users_db[sel]["dept"] = e_dept
                if e_pwd:
                    users_db[sel]["password"] = e_pwd
                doctors_list = st.session_state.doctors_list
                if old_role == "Doctor" and old_name in doctors_list and (e_role != "Doctor" or e_name != old_name):
                    doctors_list.remove(old_name)
                if e_role == "Doctor" and e_name not in doctors_list:
                    doctors_list.append(e_name)
                _log_audit(f"Updated account: {sel}")
                st.success(f"✅ Account '{sel}' updated.")
                st.rerun()
        else:
            st.info("No accounts to edit.")
    # ── Remove ───────────────────────────────────────────────────────────
    with t3:
        if users_db:
            sel_del = st.selectbox("Select account to remove", list(users_db.keys()), key="du_sel")
            st.warning(f"This will permanently remove account **{sel_del}**.")
            if st.button("🗑️ Remove Account", key="du_btn"):
                if sel_del == st.session_state.user["username"]:
                    st.error("You cannot remove the account you're currently logged in with.")
                elif users_db[sel_del]["role"] == "Admin" and \
                        sum(1 for v in users_db.values() if v["role"] == "Admin") <= 1:
                    st.error("Cannot remove the last remaining Admin account.")
                else:
                    removed = users_db.pop(sel_del)
                    if removed["name"] in st.session_state.doctors_list:
                        st.session_state.doctors_list.remove(removed["name"])
                    _log_audit(f"Removed account: {sel_del}")
                    st.success(f"✅ Account '{sel_del}' removed.")
                    st.rerun()
        else:
            st.info("No accounts to remove.")
    # ── Role Permissions ────────────────────────────────────────────────
    with t4:
        st.markdown("Control exactly what each role is allowed to do across the application.")
        perm_labels = [
            ("view_records",      "View Patient Records"),
            ("view_confidential", "View Medical / Confidential Notes"),
            ("schedule",          "Schedule Appointments"),
            ("edit_notes",        "Edit Medical Notes"),
            ("export",            "Export Data"),
            ("view_audit",        "View Audit Logs"),
        ]
        for r in ROLE_LIST:
            st.markdown(
                f'<span class="badge {badge_cls[r]}">{r}</span>',
                unsafe_allow_html=True,
            )
            cols = st.columns(len(perm_labels))
            for i, (key, label) in enumerate(perm_labels):
                with cols[i]:
                    val = st.checkbox(label, value=perms_db[r][key], key=f"perm_{r}_{key}")
                    perms_db[r][key] = val
            st.markdown("")
        if st.button("💾 Save Role Permissions", type="primary", key="save_perms"):
            _log_audit("Updated role permissions")
            st.success("✅ Role permissions updated.")
            st.rerun()
# ══════════════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════════════
def main():
    show_header()
    if not st.session_state.authenticated:
        page_login()
        return
    selected = build_sidebar()
    user = st.session_state.user
    role = user["role"]
    # Role badge top-right
    badge_cls = {"Doctor":"badge-doctor","Nurse":"badge-nurse",
                 "Receptionist":"badge-receptionist","Admin":"badge-admin"}
    st.markdown(
        f'<div style="display:flex;justify-content:flex-end;margin:-18px 0 8px">'
        f'<span class="badge {badge_cls.get(role,"badge-admin")}">👤 {user["name"]} | {role}</span></div>',
        unsafe_allow_html=True,
    )
    route = {
        "📊 Dashboard":         page_dashboard,
        "👤 Patient Records":   page_records,
        "📅 Appointments":      page_appointments,
        "📈 Analytics":         page_analytics,
        "⬆️ Data Management":   page_data_mgmt,
        "📤 Export":            page_export,
        "🧑‍⚕️ User Management": page_user_management,
    }
    fn = route.get(selected)
    if fn:
        fn()
if __name__ == "__main__":
    main()
