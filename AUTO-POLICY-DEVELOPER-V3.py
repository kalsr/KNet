# =========================================================
# KNet PolicyForge AI
# Developed by Randy Singh from Kalsnet (KNet) Consulting Group
# =========================================================
#
# REQUIREMENTS:
# pip install streamlit groq pandas reportlab python-docx
#
# RUN:
# streamlit run knet_policyforge_ai_enhanced.py
# =========================================================

import streamlit as st
from groq import Groq
import json
import pandas as pd
from io import BytesIO
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from datetime import datetime

# =========================================================
# PAGE CONFIG
# =========================================================

st.set_page_config(
    page_title="KNet PolicyForge AI",
    layout="wide",
    page_icon=" "
)

# =========================================================
# CUSTOM CSS (ONLY FOOTER + READABILITY FIXED)
# =========================================================

st.markdown("""
<style>

.main {
    background-color: #0E1117;
}

h1, h2, h3, h4 {
    color: #4FC3F7;
}

.stButton>button {
    background-color: #1F6FEB;
    color: white;
    border-radius: 10px;
    padding: 0.5rem 1rem;
    border: none;
    font-weight: 600;
}

.stTextInput>div>div>input,
.stTextArea textarea,
[data-baseweb="select"] {
    border-radius: 8px;
}

/* HEADER */
.policy-header {
    text-align: center;
    padding: 25px;
    background: linear-gradient(135deg, #111827, #1F2937);
    border-radius: 18px;
    margin-bottom: 20px;
    border: 1px solid #374151;
}

.policy-title {
    color: #00BFFF;
    font-size: 3.2rem;
    font-weight: 900;
    margin: 0;
    text-shadow: 0 0 8px rgba(0,191,255,0.7);
}

.policy-subtitle {
    color: #D1D5DB;
    font-size: 1.1rem;
    font-weight: 600;
    margin-top: 10px;
}

/* FOOTER (BRIGHT FIX) */
.footer-box {
    background-color: #ffffff !important;
    color: #111111 !important;
    padding: 22px;
    border-radius: 12px;
    margin-top: 20px;
    border: 1px solid #d0d0d0;
}

.footer-box h4,
.footer-box p,
.footer-box li,
.footer-box ul,
.footer-box b {
    color: #111111 !important;
}

</style>
""", unsafe_allow_html=True)

# =========================================================
# HEADER
# =========================================================

st.markdown("""
<div class="policy-header">
    <div class="policy-title"> KNet PolicyDeveloper AI</div>
    <div class="policy-subtitle">
        Developed by Randy Singh from Kalsnet (KNet) Consulting Group
    </div>
</div>
""", unsafe_allow_html=True)

# =========================================================
# SIDEBAR
# =========================================================

st.sidebar.title(" Configuration")

groq_api_key = st.sidebar.text_input("Groq API Key", type="password")

model_name = st.sidebar.selectbox(
    "Select Groq Model",
    ["llama-3.3-70b-versatile", "mixtral-8x7b-32768", "gemma2-9b-it"]
)

policy_area = st.sidebar.selectbox(
    "Policy Area",
    [
        "Cybersecurity Policy","Acceptable Use Policy","Remote Work Policy",
        "AI Governance Policy","Incident Response Policy","Disaster Recovery Policy",
        "Business Continuity Policy","Data Retention Policy","Privacy Policy",
        "HR Policy","Zero Trust Policy","Vendor Management Policy"
    ]
)

organization_type = st.sidebar.selectbox(
    "Organization Type",
    ["Government","Enterprise","Healthcare","Financial","Educational","Small Business","Technology Company"]
)

detail_level = st.sidebar.selectbox(
    "Policy Detail Level",
    ["Basic","Standard","Advanced","Enterprise Grade"]
)

# =========================================================
# DROPDOWNS
# =========================================================

COMPLIANCE_OPTIONS = ["NIST CSF","NIST SP 800-53","ISO 27001","HIPAA","GDPR","SOC 2"]

POLICY_OBJECTIVE_OPTIONS = [
    "Protect Confidential Information",
    "Ensure Regulatory Compliance",
    "Reduce Cybersecurity Risk"
]

RISK_LEVEL_OPTIONS = ["Very Low","Low","Moderate","High","Critical"]

ADDITIONAL_REQUIREMENT_OPTIONS = [
    "Audit Logging","Encryption at Rest","MFA","Role-Based Access Control"
]

# =========================================================
# MAIN INPUTS
# =========================================================

col1, col2 = st.columns(2)

with col1:
    organization_name = st.text_input("Organization Name")
    industry = st.text_input("Industry")

    compliance_selection = st.multiselect(
        "Compliance Requirements",
        COMPLIANCE_OPTIONS,
        default=["NIST SP 800-53"]
    )
    compliance_requirements = ", ".join(compliance_selection)

with col2:
    policy_objective = st.selectbox("Policy Objective", POLICY_OBJECTIVE_OPTIONS)
    risk_level = st.selectbox("Risk Level", RISK_LEVEL_OPTIONS, index=3)

    additional_selection = st.multiselect(
        "Additional Requirements",
        ADDITIONAL_REQUIREMENT_OPTIONS,
        default=["Audit Logging"]
    )
    additional_requirements = ", ".join(additional_selection)

# =========================================================
# PROMPT
# =========================================================

def build_prompt():
    return f"""
Generate enterprise-grade policy document.

Policy Area: {policy_area}
Organization: {organization_name}
Type: {organization_type}
Industry: {industry}
Compliance: {compliance_requirements}
Objective: {policy_objective}
Risk: {risk_level}
Detail Level: {detail_level}
Additional: {additional_requirements}
"""

# =========================================================
# GENERATION
# =========================================================

if st.button(" Generate Policy", use_container_width=True):

    if not groq_api_key:
        st.error("Please enter Groq API Key.")
    else:
        client = Groq(api_key=groq_api_key)

        with st.spinner("Generating..."):
            response = client.chat.completions.create(
                model=model_name,
                messages=[
                    {"role": "system", "content": "You are a policy expert."},
                    {"role": "user", "content": build_prompt()}
                ],
                temperature=0.3,
                max_tokens=6000
            )

        st.session_state["generated_policy"] = response.choices[0].message.content

# =========================================================
# OUTPUT
# =========================================================

if "generated_policy" in st.session_state:

    st.subheader(" Generated Policy")

    edited_policy = st.text_area(
        "Policy Editor",
        value=st.session_state["generated_policy"],
        height=700
    )

    # =====================================================
    #  EXPORT FUNCTIONS RESTORED
    # =====================================================

    def generate_pdf(text):
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer)
        styles = getSampleStyleSheet()
        story = []

        for line in text.split("\n"):
            if line.strip():
                story.append(Paragraph(line, styles["BodyText"]))
                story.append(Spacer(1, 6))

        doc.build(story)
        buffer.seek(0)
        return buffer

    def generate_docx(text):
        doc = Document()
        doc.add_heading("Generated Policy", 1)
        doc.add_paragraph(text)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def generate_csv(text):
        df = pd.DataFrame([{"Policy_Content": line} for line in text.split("\n")])
        return df.to_csv(index=False).encode("utf-8")

    def generate_json(text):
        return json.dumps({
            "generated_date": str(datetime.now()),
            "policy_area": policy_area,
            "organization": organization_name,
            "content": text
        }, indent=4)

    # =====================================================
    # DOWNLOAD BUTTONS
    # =====================================================

    st.subheader(" Export Policy")

    c1, c2, c3, c4 = st.columns(4)

    with c1:
        st.download_button(" PDF-Format", generate_pdf(edited_policy), "policy.pdf", "application/pdf")

    with c2:
        st.download_button("DOCX-Format", generate_docx(edited_policy),
                           "policy.docx",
                           "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    with c3:
        st.download_button("CSV-Format", generate_csv(edited_policy),
                           "policy.csv", "text/csv")

    with c4:
        st.download_button("JSON-Format", generate_json(edited_policy),
                           "policy.json", "application/json")

# =========================================================
# FILE UPLOAD
# =========================================================

st.subheader(" Upload Existing Policy")

uploaded_file = st.file_uploader("Upload TXT Policy File", type=["txt"])

if uploaded_file is not None:
    content = uploaded_file.read().decode("utf-8")
    st.text_area("Uploaded Policy Content", value=content, height=300)

# =========================================================
# FOOTER (BRIGHT FIX ALREADY APPLIED)
# =========================================================

st.markdown("""
<div class="footer-box">

<h4>About KNet PolicyForge AI</h4>

<p>
This enterprise-grade AI policy generation platform helps organizations
create governance, compliance, and risk policies.
</p>

<b>Developed by Randy Singh from Kalsnet (KNet) Consulting Group</b>

</div>
""", unsafe_allow_html=True)
