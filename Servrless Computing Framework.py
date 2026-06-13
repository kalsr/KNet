# serverless_streamlit_app.py – Enhanced Edition
# Developed by Randy Singh, Kalsnet (KNet) Consulting Group

import streamlit as st
import pandas as pd
import textwrap
from io import StringIO, BytesIO
import json

# Optional exports
try:
    from docx import Document
except ImportError:
    Document = None

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
except ImportError:
    canvas = None

# ---------------------------------------------------------
# Helper: synthetic data generators
# ---------------------------------------------------------
def generate_synthetic_faas_data(num_rows: int = 20):
    base = {
        "FunctionName": [
            "image_resize", "log_aggregator", "payment_webhook",
            "iot_sensor_ingest", "video_transcode_chunk",
            "auth_token_validator", "email_notification", "order_fulfillment",
        ],
        "TriggerType": [
            "HTTP API", "Log Event", "Webhook", "MQTT Message",
            "Object Storage Event", "JWT Auth", "Queue Message", "Order Event",
        ],
        "AvgExecTime_ms": [120, 35, 210, 80, 950, 45, 60, 320],
        "MonthlyInvocations": [50000, 120000, 15000, 250000, 8000, 500000, 75000, 30000],
        "CloudProvider": [
            "AWS Lambda", "Azure Functions", "GCP Cloud Functions", "AWS Lambda",
            "Cloudflare Workers", "AWS Lambda", "Azure Functions", "GCP Cloud Functions",
        ],
        "ColdStart_ms": [350, 420, 380, 300, 50, 280, 400, 360],
        "MemoryMB": [512, 256, 1024, 128, 128, 256, 256, 512],
    }
    rows = []
    for i in range(num_rows):
        idx = i % len(base["FunctionName"])
        rows.append({
            "FunctionName": base["FunctionName"][idx],
            "TriggerType": base["TriggerType"][idx],
            "AvgExecTime_ms": base["AvgExecTime_ms"][idx],
            "MonthlyInvocations": base["MonthlyInvocations"][idx],
            "CloudProvider": base["CloudProvider"][idx],
            "ColdStart_ms": base["ColdStart_ms"][idx],
            "MemoryMB": base["MemoryMB"][idx],
        })
    return pd.DataFrame(rows)


def generate_synthetic_serverless_eval_data():
    data = {
        "Platform": [
            "AWS Lambda", "Azure Functions", "GCP Cloud Functions",
            "OpenFaaS", "Knative", "Cloudflare Workers",
        ],
        "ColdStart_ms": [350, 420, 380, 600, 700, 5],
        "MaxConcurrent": [1000, 800, 900, 500, 400, 10000],
        "MaxTimeout_sec": [900, 600, 540, 300, 300, 30],
        "PricingModel": [
            "per-ms", "per-ms", "per-ms", "per-node", "per-node", "per-req",
        ],
        "FreeTier_M_reqs": [1, 1, 2, 0, 0, 10],
        "BestFor": [
            "Event-driven APIs", "Enterprise integrations", "Data pipelines",
            "On-prem FaaS", "Kubernetes-native apps", "Edge / ultra-low-latency",
        ],
        "VendorLockRisk": ["High", "High", "High", "Low", "Low", "Medium"],
    }
    return pd.DataFrame(data)


def generate_ai_pipeline_data():
    data = {
        "Stage": ["Ingest", "Preprocess", "Inference", "Post-process", "Store & Notify"],
        "Function": ["ingest_fn", "preprocess_fn", "model_api_call", "postprocess_fn", "store_notify_fn"],
        "AvgLatency_ms": [30, 80, 250, 40, 55],
        "Provider": ["AWS Lambda", "GCP Cloud Functions", "Hosted Model API", "AWS Lambda", "Azure Functions"],
        "ScalesTo": ["∞", "∞", "Rate-limited", "∞", "∞"],
    }
    return pd.DataFrame(data)


def generate_challenge_data():
    data = {
        "Challenge": [
            "Cold Starts", "Vendor Lock-in", "Observability", "State Management",
            "Execution Timeout", "Networking Complexity", "Testing Locally", "Cost at Scale",
        ],
        "Severity": ["High", "High", "Medium", "High", "Medium", "Medium", "Medium", "Medium"],
        "Mitigation": [
            "Provisioned concurrency / warm-up pings",
            "Abstraction layers (Serverless Framework, Pulumi)",
            "Distributed tracing (X-Ray, Jaeger, Honeycomb)",
            "External state stores (Redis, DynamoDB, Cosmos DB)",
            "Choreography patterns / async workflows",
            "Service mesh / VPC peering / private endpoints",
            "Local emulators (LocalStack, Azure Functions Core Tools)",
            "Hybrid architecture for high-volume steady workloads",
        ],
    }
    return pd.DataFrame(data)

# ---------------------------------------------------------
# Export helpers
# ---------------------------------------------------------
def export_df_as_docx(df: pd.DataFrame, title: str):
    if Document is None:
        return None
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph("Synthetic serverless dataset:")
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            row_cells[i].text = str(row[col])
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def export_df_as_pdf(df: pd.DataFrame, title: str):
    if canvas is None:
        return None
    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    width, height = letter
    y = height - 50

    def draw_line(text, y_pos):
        c.drawString(40, y_pos, text)
        return y_pos - 15

    y = draw_line(title, y)
    y = draw_line("", y)
    for col in df.columns:
        if y < 80:
            c.showPage()
            y = height - 50
        y = draw_line(f"{col}", y)
    y = draw_line("", y)
    for _, row in df.iterrows():
        line = ", ".join(f"{col}={row[col]}" for col in df.columns)
        if y < 80:
            c.showPage()
            y = height - 50
        y = draw_line(line[:120], y)
    c.showPage()
    c.save()
    buf.seek(0)
    return buf


def export_df_as_json(df: pd.DataFrame, title: str):
    payload = {
        "title": title,
        "rows": df.to_dict(orient="records"),
    }
    return json.dumps(payload, indent=2)

# ---------------------------------------------------------
# Layout helpers
# ---------------------------------------------------------
def section_header(title: str, subtitle: str = ""):
    st.markdown(f"<h2 style='color:#1a5fb4;'>{title}</h2>", unsafe_allow_html=True)
    if subtitle:
        st.caption(subtitle)


def show_mermaid(diagram_code: str, label: str = "Conceptual flow diagram"):
    st.markdown(f"**{label}**")
    st.markdown("```mermaid\n" + diagram_code + "\n```")


def show_ascii(diagram: str, label: str = "Structural diagram"):
    st.markdown(f"**{label}**")
    st.code(diagram, language="text")


def wrapped_text(text: str):
    return "\n".join(textwrap.wrap(text, width=110))


def rec_box(items: list):
    st.markdown("**★ Recommendations**")
    for item in items:
        st.markdown(f"- {item}")

# ---------------------------------------------------------
# Category 1 – FaaS
# ---------------------------------------------------------
def category_1_faas():
    section_header("1. Function-as-a-Service (FaaS)",
                   "Event-driven, stateless functions managed by the cloud provider.")

    st.write(
        "FaaS is the core building block of serverless computing. Developers write small, "
        "single-purpose functions; the platform handles provisioning, scaling, patching, and "
        "billing at millisecond granularity. No servers to manage — ever. In practice, FaaS "
        "becomes the glue for modern architectures: APIs, data pipelines, IoT ingestion, and "
        "workflow orchestration all converge on functions."
    )

    show_ascii("""
  FaaS CONCEPTUAL ARCHITECTURE

  ┌─────────────────────────────────────────────────────────────────┐
  │                   SERVERLESS PLATFORM                           │
  │                                                                 │
  │  Event Sources        Function Runtime       Outputs            │
  │  ─────────────        ────────────────       ───────            │
  │  HTTP / API  ──────►  [ Function A ]  ──────► Database         │
  │  Queue Msg   ──────►  [ Function B ]  ──────► Object Store     │
  │  Cron / Timer──────►  [ Function C ]  ──────► Event Bus        │
  │  Storage Event─────►  [ Function D ]  ──────► Notification     │
  │  IoT / MQTT  ──────►  [ Function E ]  ──────► Downstream Fn    │
  │                                                                 │
  │  Platform Services: Autoscale │ Billing │ Logs │ Tracing       │
  └─────────────────────────────────────────────────────────────────┘
""", "FaaS Platform Architecture")

    sub_choice = st.radio(
        "Select FaaS capability",
        [
            "Show full capability list",
            "Function execution capability",
            "Storage capability",
            "Container infrastructure capability",
            "Networking capability",
            "Fault tolerance capabilities",
            "Scaling of functions",
            "Data analytics at the edge",
            "Scientific computing",
            "Mobile computing",
        ],
    )

    # Synthetic data controls (global for FaaS)
    st.markdown("### Synthetic FaaS dataset controls")
    if "faas_rows" not in st.session_state:
        st.session_state.faas_rows = 20
    if "faas_df" not in st.session_state:
        st.session_state.faas_df = generate_synthetic_faas_data(st.session_state.faas_rows)

    col_slider, col_buttons = st.columns([2, 1])
    with col_slider:
        st.session_state.faas_rows = st.slider(
            "Number of synthetic FaaS records",
            min_value=0,
            max_value=300,
            value=st.session_state.faas_rows,
        )
    with col_buttons:
        if st.button("Generate synthetic FaaS data"):
            st.session_state.faas_df = generate_synthetic_faas_data(st.session_state.faas_rows)
        if st.button("Reset synthetic data"):
            st.session_state.faas_rows = 20
            st.session_state.faas_df = generate_synthetic_faas_data(st.session_state.faas_rows)

    st.dataframe(st.session_state.faas_df, use_container_width=True)

    # Export current synthetic dataset
    st.markdown("#### Export synthetic FaaS dataset")
    title = f"Synthetic FaaS dataset ({st.session_state.faas_rows} records)"
    csv = st.session_state.faas_df.to_csv(index=False).encode("utf-8")
    st.download_button("⬇ Download CSV", csv, "faas_synthetic_data.csv", "text/csv")

    json_buf = export_df_as_json(st.session_state.faas_df, title).encode("utf-8")
    st.download_button("⬇ Download JSON", json_buf, "faas_synthetic_data.json", "application/json")

    docx_buf = export_df_as_docx(st.session_state.faas_df, title)
    if docx_buf is not None:
        st.download_button(
            "⬇ Download Word (DOCX)",
            docx_buf,
            "faas_synthetic_data.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    else:
        st.info("Install `python-docx` to enable Word export.")

    pdf_buf = export_df_as_pdf(st.session_state.faas_df, title)
    if pdf_buf is not None:
        st.download_button(
            "⬇ Download PDF",
            pdf_buf,
            "faas_synthetic_data.pdf",
            "application/pdf",
        )
    else:
        st.info("Install `reportlab` to enable PDF export.")

    st.markdown("---")

    if sub_choice == "Show full capability list":
        st.markdown("### Full FaaS capability overview")
        st.write(
            "A serverless platform manages the full lifecycle of functions: provisioning, execution, "
            "autoscaling, and teardown. Developers focus exclusively on business logic. Typical real-world "
            "deployments include image processing APIs, webhook handlers, IoT data pipelines, real-time "
            "notification services, and scheduled data transformation jobs."
        )
        st.write(
            "In practice, an e-commerce platform uses FaaS to process payments, send confirmation emails, "
            "update inventory, generate invoices, and log analytics events — each as an independent function "
            "triggered by its own event, scaling independently. This separation of concerns reduces blast radius "
            "and makes it easier to reason about failures."
        )
        show_mermaid("""
flowchart LR
    UserRequest -->|HTTP/API| API_Gateway
    API_Gateway -->|Invoke| FaaS_Function
    FaaS_Function -->|Read/Write| DataStore[(Object Storage)]
    FaaS_Function --> EventBus((Event Bus))
    EventBus --> OtherFunctions
    OtherFunctions --> AnalyticsDB[(Analytics DB)]
    OtherFunctions --> NotifService((Push Notification))
""", "Full FaaS event flow")
        rec_box([
            "Decompose monolithic services into small, single-purpose functions.",
            "Use an API Gateway (AWS API GW, Azure APIM) in front of all HTTP-triggered functions.",
            "Tag functions by team and cost center for granular billing visibility.",
            "Set per-function memory and timeout limits to avoid runaway costs.",
        ])

    elif sub_choice == "Function execution capability":
        st.markdown("### Function execution capability")
        st.write(
            "Functions are short-lived, stateless compute units triggered by events: HTTP calls, queue messages, "
            "object storage changes, database streams, or cron schedules. The platform spins up micro-VMs or "
            "containers on demand, executes the function, and reclaims resources. Execution time is billed to the "
            "nearest millisecond — you pay only for active CPU time."
        )
        st.write(
            "Real-world execution patterns often include validation, enrichment, persistence, and event emission. "
            "A single function might validate a payment, enrich it with customer metadata, write to a ledger, and "
            "emit a downstream event for fulfillment."
        )
        st.write(
            "**Real-world example (Stripe):** Stripe processes billions of payment webhook calls per month using "
            "serverless functions. Each webhook validates the event signature, updates the order database, and "
            "publishes a fulfillment event — all in under 200 ms."
        )
        show_ascii("""
  FUNCTION EXECUTION LIFECYCLE

  Event ──► [Platform Scheduler]
                    │
            ┌───────▼────────┐
            │  Cold Start?   │
            └───┬────────┬───┘
               YES       NO
                │         │
          [Provision]  [Warm Pool]
                │         │
          [Init Runtime]  │
                └────┬────┘
                     ▼
              [Execute Function]
                     │
            ┌────────▼────────┐
            │  Write Outputs  │
            └────────┬────────┘
                     ▼
              [Release / Reclaim Resources]
""", "Function execution lifecycle")
        show_mermaid("""
flowchart LR
    PaymentGateway --> WebhookEndpoint
    WebhookEndpoint --> ValidateFn
    ValidateFn -->|Valid| FaaS_PaymentFn
    ValidateFn -->|Invalid| DLQ((Dead-Letter Queue))
    FaaS_PaymentFn --> OrdersDB[(Orders DB)]
    FaaS_PaymentFn --> FulfillmentQueue((Fulfillment Queue))
    FulfillmentQueue --> ShipFn
""", "Payment webhook function flow")
        rec_box([
            "Keep functions under 50 lines of logic; delegate heavy lifting to libraries or services.",
            "Use provisioned concurrency for latency-sensitive paths to eliminate cold starts.",
            "Validate and sanitize all inputs at the function boundary — never trust the event payload.",
            "Set a realistic timeout (e.g., 10 s for API, 5 min for batch) to cap runaway execution costs.",
        ])

    elif sub_choice == "Storage capability":
        st.markdown("### Storage capability")
        st.write(
            "Since functions are stateless, all durable state lives in external storage: object stores (S3, GCS, "
            "Azure Blob), relational/NoSQL databases, in-memory caches (Redis, Elasticache), or event logs "
            "(Kinesis, Kafka). Intermediate pipeline results flow through object storage, enabling loosely coupled, "
            "independently scalable processing stages."
        )
        st.write(
            "Designing storage tiers — hot, warm, and cold — lets you balance latency and cost. Hot paths use "
            "low-latency stores for user-facing operations; warm paths use object storage for batch processing; "
            "cold paths use warehouses for analytics."
        )
        st.write(
            "**Real-world example (Netflix):** Netflix uses AWS Lambda with S3 as an intermediary to transcode video "
            "files. Each Lambda reads from S3, processes a chunk, and writes output back — enabling massive parallelism "
            "without managing a transcoding fleet."
        )
        show_ascii("""
  SERVERLESS STORAGE TIERS

  ┌──────────────────────────────────────────────────────┐
  │  Hot Path (ms latency)                               │
  │  Redis / DynamoDB / Cosmos DB ◄── Function reads    │
  ├──────────────────────────────────────────────────────┤
  │  Warm Path (seconds)                                 │
  │  Object Storage (S3 / GCS / Blob) ◄── Intermediate  │
  ├──────────────────────────────────────────────────────┤
  │  Cold Path (minutes)                                 │
  │  Data Warehouse (Redshift / BigQuery / Synapse)      │
  │  ◄── Aggregated / analytical results                 │
  └──────────────────────────────────────────────────────┘
""", "Serverless storage tiers")
        show_mermaid("""
flowchart LR
    RawLogs[(Raw Logs in Object Storage)] --> TransformFn
    TransformFn --> AggregatedMetrics[(Aggregated Metrics DB)]
    AggregatedMetrics --> Dashboard
    TransformFn --> Cache[(Redis Hot Cache)]
    Cache --> RealTimeAPI
""", "Log processing pipeline with tiered storage")
        rec_box([
            "Never store state in function memory across invocations — use external stores.",
            "Use DynamoDB or Cosmos DB for low-latency state; S3 for bulk or intermediate data.",
            "Enable S3 Versioning and lifecycle policies to manage intermediate artifact costs.",
            "Cache frequently read configuration in Redis to avoid repeated database calls.",
        ])

    elif sub_choice == "Container infrastructure capability":
        st.markdown("### Container infrastructure capability")
        st.write(
            "Modern FaaS platforms support container-based functions (AWS Lambda Container Images, Azure Container Apps, "
            "GCP Cloud Run), enabling custom runtimes, larger packages, and specialized tools. OpenFaaS and Knative "
            "run containerized functions on Kubernetes, combining FaaS developer experience with container portability."
        )
        st.write(
            "This is especially useful for scientific workloads, ML inference with GPU libraries, or legacy binaries "
            "that cannot be easily packaged as small zip deployments."
        )
        st.write(
            "**Real-world example (Genomics England):** Genomics England runs bioinformatics workflows as containerized "
            "Lambda functions with custom C++ alignment tools. New genomic datasets arriving in S3 automatically trigger "
            "analysis pipelines without provisioning HPC clusters."
        )
        show_ascii("""
  CONTAINER-BASED FaaS STACK

  Developer writes Dockerfile
          │
          ▼
  Container Registry (ECR / ACR / GCR)
          │
          ▼
  FaaS Platform pulls image on cold start
          │
          ▼
  ┌────────────────────────────┐
  │  Micro-VM (Firecracker)    │
  │  └─ Container Runtime      │
  │     └─ Function Code       │
  │        └─ Custom Tools     │
  │           (e.g., GATK,     │
  │            ffmpeg, etc.)   │
  └────────────────────────────┘
""", "Container FaaS runtime stack")
        show_mermaid("""
flowchart LR
    DataBucket[(Genomics Data Bucket)] --> S3Trigger
    S3Trigger --> ContainerFn1[Alignment Container Fn]
    ContainerFn1 --> ContainerFn2[Variant Calling Fn]
    ContainerFn2 --> ContainerFn3[Annotation Fn]
    ContainerFn3 --> ResultsBucket[(Analysis Results Bucket)]
    ResultsBucket --> ReportFn[Report Generator Fn]
""", "Containerized genomics pipeline")
        rec_box([
            "Use container images when your dependency package exceeds the 250 MB zip limit.",
            "Pin container image tags to digests (not 'latest') for reproducible deployments.",
            "Use multi-stage Docker builds to keep images lean and reduce cold start time.",
            "Consider Firecracker-based runtimes (Lambda) for better security isolation vs. shared containers.",
        ])

    elif sub_choice == "Networking capability":
        st.markdown("### Networking capability")
        st.write(
            "Functions can communicate via HTTP/REST, gRPC, event buses (EventBridge, Pub/Sub), message queues "
            "(SQS, Service Bus, Cloud Tasks), and streaming platforms (Kinesis, Kafka). VPC integration allows "
            "functions to access private resources. Service meshes and API gateways enforce policies at the edge."
        )
        st.write(
            "Choosing the right pattern — synchronous vs. asynchronous, queue vs. stream — is central to building "
            "resilient, observable serverless systems."
        )
        st.write(
            "**Real-world example (Delivery Hero):** Delivery Hero uses AWS EventBridge to fan out order events to "
            "dozens of downstream Lambda functions (notifications, inventory, analytics, loyalty) without any "
            "point-to-point coupling between services."
        )
        show_ascii("""
  SERVERLESS NETWORKING PATTERNS

  Pattern 1: Synchronous (HTTP)
  Client ──► API Gateway ──► Function ──► Response

  Pattern 2: Async Fan-out (Event Bus)
  Producer Fn ──► EventBridge ──► Fn A (notify)
                              ──► Fn B (inventory)
                              ──► Fn C (analytics)

  Pattern 3: Queue-based Decoupling
  Fn A ──► SQS Queue ──► Fn B (batch consumer)

  Pattern 4: Streaming
  Kinesis Stream ──► Function (shard consumer)
""", "Serverless networking patterns")
        show_mermaid("""
flowchart LR
    UploadVideo --> SegmenterFn
    SegmenterFn --> SegmentQueue((SQS Segment Queue))
    SegmentQueue --> TranscodeFn1
    SegmentQueue --> TranscodeFn2
    SegmentQueue --> TranscodeFn3
    TranscodeFn1 --> MergeFn
    TranscodeFn2 --> MergeFn
    TranscodeFn3 --> MergeFn
    MergeFn --> FinalVideo[(Final Output Bucket)]
    MergeFn --> NotifyFn((Notify User))
""", "Parallel video transcoding via queue fan-out")
        rec_box([
            "Prefer async (queue/event) communication over synchronous chains to avoid cascading timeouts.",
            "Use VPC endpoints for functions accessing private RDS/Elasticache to avoid public internet exposure.",
            "Implement circuit breakers in functions that call external HTTP APIs.",
            "Use EventBridge schema registry to enforce event contracts between producer and consumer functions.",
        ])

    elif sub_choice == "Fault tolerance capabilities":
        st.markdown("### Fault tolerance capabilities")
        st.write(
            "FaaS platforms provide built-in retry logic, automatic failover to healthy instances, and dead-letter "
            "queues (DLQs) for unprocessable messages. Idempotency is critical: because retries are automatic, "
            "functions must produce the same result when called multiple times with the same input."
        )
        st.write(
            "Fault tolerance in serverless is less about hardware redundancy and more about designing robust "
            "business logic that can handle partial failures, duplicate events, and backpressure."
        )
        st.write(
            "**Real-world example (Airbnb):** Airbnb's booking confirmation functions are designed to be idempotent "
            "using a transaction ID. If a Lambda retries due to a transient DynamoDB timeout, the second attempt "
            "detects the existing record and skips re-processing — preventing double bookings."
        )
        show_ascii("""
  FAULT TOLERANCE FLOW

  Event arrives
      │
      ▼
  [Function Invoked]
      │
  ┌───▼────────────┐      Success      ┌─────────────┐
  │  Try Execute   │ ────────────────► │  Write to   │
  └───┬────────────┘                   │  DB / Queue │
      │ Failure                        └─────────────┘
      ▼
  [Retry 1] ──► [Retry 2] ──► [Retry N]
      │ All retries exhausted
      ▼
  [Dead-Letter Queue]
      │
  [Alert + Manual Review]
""", "Fault tolerance and retry flow")
        show_mermaid("""
flowchart LR
    IoTDevice --> IngestFn
    IngestFn -->|Success| MetricsDB[(Metrics DB)]
    IngestFn -->|Transient Error| RetryLogic{Retry?}
    RetryLogic -->|Yes - Attempt 2-3| IngestFn
    RetryLogic -->|Exhausted| DLQ((Dead-Letter Queue))
    DLQ --> AlertFn
    AlertFn --> OpsTeam((Ops Alert))
""", "IoT ingestion with retry and DLQ")
        rec_box([
            "Design all functions to be idempotent — use a unique request/transaction ID as a deduplication key.",
            "Configure DLQs on all async event sources (SQS, SNS, EventBridge) and alert on DLQ depth.",
            "Implement exponential backoff with jitter in functions calling external APIs.",
            "Use AWS Step Functions or Azure Durable Functions for complex workflows needing state-aware retries.",
        ])

    elif sub_choice == "Scaling of functions":
        st.markdown("### Scaling of functions")
        st.write(
            "Serverless platforms scale function concurrency automatically from zero to thousands of instances within "
            "seconds. Each invocation runs in an isolated context; the platform manages instance pooling transparently. "
            "Throttling (concurrency limits) prevents downstream services from being overwhelmed."
        )
        st.write(
            "Understanding concurrency, burst limits, and downstream capacity (databases, caches, third-party APIs) "
            "is essential to avoid self-inflicted outages."
        )
        st.write(
            "**Real-world example (Coca-Cola):** Coca-Cola's vending machine platform uses AWS Lambda to handle "
            "purchase events. During peak promotions, invocations spike 50x in minutes — Lambda scales horizontally "
            "with zero operator intervention, then scales back to near-zero overnight."
        )
        show_ascii("""
  SERVERLESS AUTO-SCALING

  Traffic:    ░░░░▓▓▓▓▓▓▓▓████████▓▓▓▓░░░░
              Low  Ramp   Peak      Drain  Low

  Instances:  1    10     1000+     50     1
                          │
                    Concurrency Limit
                    (configurable, prevents
                     DB connection exhaustion)
""", "Serverless auto-scaling over time")
        show_mermaid("""
flowchart LR
    Users --> API_Gateway
    API_Gateway --> FaaS_Instances
    subgraph FaaS_Instances [Auto-scaled Function Pool]
        Fn1
        Fn2
        Fn3
        FnN[Fn ... N]
    end
    FaaS_Instances --> RDSProxy[(RDS Proxy - Connection Pool)]
    RDSProxy --> Database[(Database)]
""", "Auto-scaling with RDS Proxy to protect DB connections")
        rec_box([
            "Set reserved concurrency on critical functions to guarantee capacity during traffic spikes.",
            "Use RDS Proxy (AWS) or connection poolers to prevent Lambda from exhausting DB connections at scale.",
            "Implement account-level concurrency quotas awareness — request limit increases proactively.",
            "Monitor concurrency utilization; set CloudWatch alarms when approaching throttle thresholds.",
        ])

    elif sub_choice == "Data analytics at the edge":
        st.markdown("### Data analytics at the network edge")
        st.write(
            "Edge serverless platforms (Cloudflare Workers, AWS Lambda@Edge, Fastly Compute) run functions in "
            "100+ PoPs globally, reducing round-trip latency from hundreds of milliseconds to single digits. "
            "Common workloads: real-time personalization, A/B testing, image optimization, anomaly detection, "
            "and local data aggregation before cloud upload."
        )
        st.write(
            "Edge analytics often aggregate or filter data locally, sending only meaningful events to the cloud — "
            "reducing bandwidth and central processing costs."
        )
        st.write(
            "**Real-world example (Shopify):** Shopify uses Cloudflare Workers at the edge to apply real-time "
            "storefront personalization (currency, language, pricing) without a round-trip to origin servers — "
            "cutting page load time by ~40% for international users."
        )
        show_ascii("""
  EDGE SERVERLESS TOPOLOGY

  User (Tokyo) ──► [PoP: Tokyo] ──► Edge Function
                                         │
                                   Local processing
                                   (auth, personalize,
                                    A/B test, cache)
                                         │
                              ┌──────────▼──────────┐
                              │ Forward only what's  │
                              │ needed to Origin     │
                              └──────────────────────┘
                              Origin (US-East): heavy logic

  Latency: 5 ms at edge vs 180 ms round-trip to origin
""", "Edge vs. origin latency topology")
        show_mermaid("""
flowchart LR
    Camera --> EdgeFn[Edge Detection Fn - PoP]
    EdgeFn -->|No incident| Discard[Discard Frame]
    EdgeFn -->|Incident detected| CloudAnalytics[(Cloud Analytics)]
    EdgeFn --> LocalAlert[(Local Alert Dashboard)]
    CloudAnalytics --> MLModel[Incident Classification]
""", "Smart city edge analytics flow")
        rec_box([
            "Use edge functions for auth, routing, and personalization — not heavy compute.",
            "Store frequently accessed data in edge KV stores (Cloudflare KV, Lambda@Edge + CloudFront) to avoid origin calls.",
            "Design edge functions to be stateless; propagate state to origin or edge KV only.",
            "Monitor edge function CPU time limits (Cloudflare: 50 ms, Lambda@Edge: 5 s) and profile regularly.",
        ])

    elif sub_choice == "Scientific computing":
        st.markdown("### Scientific computing")
        st.write(
            "Scientific workflows often consist of discrete event-driven steps: data ingestion, quality control, "
            "simulation, post-processing, and visualization. Serverless functions orchestrate these steps — triggered "
            "by new data arrivals or completed upstream tasks — without maintaining long-running cluster infrastructure."
        )
        st.write(
            "Serverless is ideal for bursty workloads where simulations run in waves, and orchestration overhead "
            "needs to be minimal."
        )
        st.write(
            "**Real-world example (NASA JPL):** NASA's Jet Propulsion Laboratory uses AWS Lambda and Step Functions "
            "to process Mars rover telemetry. New data packets arriving in S3 automatically trigger preprocessing, "
            "calibration, and image assembly pipelines — processing years of rover data in weeks of compute."
        )
        show_ascii("""
  SCIENTIFIC WORKFLOW ORCHESTRATION

  [New Data Arrives in Object Store]
            │
            ▼
  [QC / Validation Function]
            │
      ┌─────▼──────┐
      │  Pass QC?  │
      └──┬──────┬──┘
        YES     NO
         │       └──► [Quarantine Bucket]
         ▼
  [Preprocessing Function]
         │
         ▼
  [Simulation / Analysis Function]
  (may fan out to 100s of parallel instances)
         │
         ▼
  [Post-processing + Visualization]
         │
         ▼
  [Results Store + Notification]
""", "Scientific pipeline stages")
        show_mermaid("""
flowchart LR
    SatelliteData[(Satellite Data Bucket)] --> QCFn[QC Validation Fn]
    QCFn -->|Pass| PreprocessFn
    QCFn -->|Fail| Quarantine[(Quarantine Bucket)]
    PreprocessFn --> SimulationFn
    SimulationFn --> ParallelFns
    subgraph ParallelFns [Parallel Simulation Workers]
        Sim1
        Sim2
        SimN[Sim ... N]
    end
    ParallelFns --> PostProcessFn
    PostProcessFn --> Reports[(Reports & Visuals)]
    PostProcessFn --> NotifyFn((Notify Research Team))
""", "Parallel scientific simulation pipeline")
        rec_box([
            "Use AWS Step Functions Map state or Azure Durable Functions fan-out for massively parallel simulations.",
            "Store intermediate results in S3/GCS with meaningful prefixes for easy resumption after failures.",
            "Set function memory based on working set size — more memory = more vCPU in Lambda.",
            "Use spot/preemptible VMs for heavy simulations; FaaS for orchestration and lightweight transforms.",
        ])

    elif sub_choice == "Mobile computing":
        st.markdown("### Mobile computing")
        st.write(
            "Serverless backends are a natural fit for mobile apps: they scale with user growth, require no server "
            "fleet management, and support independent deployment of features. Common patterns include REST/GraphQL "
            "APIs, push notifications, user authentication, in-app purchase processing, and content delivery."
        )
        st.write(
            "Mobile teams benefit from serverless by decoupling backend feature releases from app store updates — "
            "functions can be updated independently while the client remains stable."
        )
        st.write(
            "**Real-world example (McDonald's):** McDonald's mobile order platform (serving tens of millions of "
            "orders per month) uses AWS Lambda for order processing, menu customization, loyalty point calculation, "
            "and push notifications — scaling seamlessly during lunch and dinner rushes."
        )
        show_ascii("""
  MOBILE BACKEND SERVERLESS FLOW

  Mobile App ──► API Gateway ──► AuthFn ──► UserProfileDB
                                   │
                                   └──► OrderFn ──► OrdersDB
                                              │
                                              └──► NotifFn ──► Push Service
""", "Mobile backend serverless flow")
        show_mermaid("""
flowchart LR
    MobileClient --> AuthAPI
    AuthAPI --> AuthFn
    AuthFn --> UserDB[(User DB)]
    MobileClient --> OrderAPI
    OrderAPI --> OrderFn
    OrderFn --> OrdersDB[(Orders DB)]
    OrderFn --> LoyaltyFn
    LoyaltyFn --> PointsDB[(Points DB)]
    OrderFn --> PushFn((Push Notification Fn))
""", "Mobile order + loyalty flow")
        rec_box([
            "Use Cognito, Firebase Auth, or Azure AD B2C for identity; keep auth logic thin in functions.",
            "Implement rate limiting per user/device at the API gateway level.",
            "Use feature flags and configuration services to roll out backend changes gradually.",
            "Log mobile-specific metadata (device type, app version) for better observability and debugging.",
        ])

# ---------------------------------------------------------
# Main app
# ---------------------------------------------------------
st.set_page_config(page_title="Serverless Computing Framework Dashboard", layout="wide")

st.markdown(
    "<h1 style='color:#1a5fb4; font-weight:bold;'>Serverless Computing Framework Dashboard</h1>",
    unsafe_allow_html=True,
)
st.caption("Developed by Randy Singh, Kalsnet (KNet) Consulting Group")

st.sidebar.header("Navigation")
category = st.sidebar.selectbox(
    "Select framework category",
    ["1. Function-as-a-Service (FaaS)"],
)

if category.startswith("1."):
    category_1_faas()

st.markdown("---")
st.caption("Serverless Computing Framework – enhanced synthetic data, exports, and conceptual diagrams.")
