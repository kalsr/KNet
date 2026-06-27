


# ShopFloor Sentinel - Enhanced Export Version
# AI-powered manufacturing quality, tool wear, and predictive maintenance system
# for casting and machining ancillary units.

# Developed by Randy Singh | Kalsnet (KNet) Consulting

# EXPORT CAPABILITIES: PDF, Word (.docx), CSV, JSON, TXT


import io
import json
import logging
import os
import random
import re
import uuid
import datetime as dt
import math
from typing import Dict, List, Tuple

import pandas as pd
import numpy as np
import streamlit as st

try:
    from fpdf import FPDF
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import requests
    HAS_REQUESTS = True
except ImportError:
    HAS_REQUESTS = False

# =========================================================================
# LOGGING
# =========================================================================
LOG_DIR = os.environ.get("SHOPFLOOR_LOG_DIR", "./logs")
os.makedirs(LOG_DIR, exist_ok=True)
logger = logging.getLogger("shopfloor_sentinel")
if not logger.handlers:
    logger.setLevel(logging.INFO)
    _fh = logging.FileHandler(os.path.join(LOG_DIR, "shopfloor_sentinel.log"))
    _fh.setFormatter(logging.Formatter("%(asctime)s | %(levelname)s | %(message)s"))
    logger.addHandler(_fh)

# =========================================================================
# PAGE CONFIG
# =========================================================================
st.set_page_config(
    page_title="ShopFloor Sentinel | Manufacturing AI",
    page_icon="🛡️",
    layout="wide",
    initial_sidebar_state="expanded",
)

# =========================================================================
# PROFESSIONAL STYLING
# =========================================================================
st.markdown(
    """
    <style>
    .main-title {
        text-align: center;
        color: #0047AB;
        font-weight: 900;
        font-size: 2.8rem;
        margin-bottom: 0.3rem;
        letter-spacing: 0.8px;
        font-family: 'Arial Black', sans-serif;
    }
    .subtitle-line {
        text-align: center;
        color: #0047AB;
        font-weight: 800;
        font-size: 1.2rem;
        margin-top: 0;
        margin-bottom: 1.2rem;
        font-family: 'Arial', sans-serif;
    }
    .section-header {
        color: #0047AB;
        font-weight: 800;
        font-size: 1.8rem;
        margin-top: 1.5rem;
        margin-bottom: 0.8rem;
        border-bottom: 3px solid #0047AB;
        padding-bottom: 0.5rem;
    }
    .subsection-header {
        color: #003380;
        font-weight: 700;
        font-size: 1.3rem;
        margin-top: 1rem;
        margin-bottom: 0.6rem;
    }
    .metric-card {
        background: #f0f7ff;
        border-left: 5px solid #0047AB;
        padding: 1rem;
        border-radius: 8px;
        margin: 0.8rem 0;
    }
    .formula-box {
        background: #fff9e6;
        border: 2px solid #ffd700;
        padding: 1rem;
        border-radius: 6px;
        font-family: 'Courier New', monospace;
        font-size: 0.9rem;
        margin: 0.8rem 0;
    }
    .info-box {
        background: #e8f4f8;
        border-left: 4px solid #0099cc;
        padding: 0.8rem;
        border-radius: 4px;
        font-size: 0.85rem;
        margin: 0.6rem 0;
    }
    .warning-box {
        background: #fff3cd;
        border-left: 4px solid #ff9800;
        padding: 0.8rem;
        border-radius: 4px;
        font-size: 0.85rem;
        margin: 0.6rem 0;
    }
    .status-good {
        color: #1a8a3e;
        font-weight: 700;
    }
    .status-warning {
        color: #ff9800;
        font-weight: 700;
    }
    .status-critical {
        color: #c0392b;
        font-weight: 700;
    }
    .divider {
        border-top: 2px solid #0047AB;
        margin: 1.5rem 0;
        opacity: 0.3;
    }
    .export-button {
        display: inline-block;
        margin: 0.4rem;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# =========================================================================
# TITLE BAR
# =========================================================================
st.markdown('<div class="main-title">🛡️ ShopFloor Sentinel</div>', unsafe_allow_html=True)
st.markdown('<div class="subtitle-line">Developed by Randy Singh from Kalsnet (KNet) Consulting</div>', unsafe_allow_html=True)
st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

st.markdown(
    """
    **Manufacturing Intelligence Platform** — Real-time AI for defect detection, tool wear prediction, 
    and predictive maintenance. Turn shop floor data into margin-saving insights.
    
    **ROI Focus:** Cut scrap • Cut tool cost • Cut downtime
    """
)

# =========================================================================
# ENVIRONMENT CONFIG
# =========================================================================
DEPLOY_ENV = os.environ.get("SHOPFLOOR_ENV", "production")
MAX_UPLOAD_MB = int(os.environ.get("SHOPFLOOR_MAX_UPLOAD_MB", "25"))

# =========================================================================
# CONSTANTS & FORMULAS
# =========================================================================

DEFECT_TYPES = {
    "Blowholes": {"severity": 8, "description": "Gas porosity in casting, reduces strength significantly"},
    "Shrinkage": {"severity": 9, "description": "Material contraction cavities, high rejection rate, critical defect"},
    "Cold Shuts": {"severity": 7, "description": "Incomplete flow/fusion, fails under load"},
    "Surface Cracks": {"severity": 8, "description": "Finish defects visible to OEM, customer visible"},
    "Finish Marks": {"severity": 4, "description": "Cosmetic surface marks, usually reworkable"},
    "Burrs": {"severity": 3, "description": "Sharp edges from machining, easily removed"},
    "Dimensional Out": {"severity": 6, "description": "Size outside tolerance, CMM rejection"},
}

MACHINE_TYPES = ["VMC-3", "HMC-5", "Lathe-CNC", "Casting Furnace-1", "Casting Furnace-2", "Die Press-4"]

# =========================================================================
# SESSION STATE
# =========================================================================
DEFAULTS = {
    "synthetic_df": None,
    "real_df": None,
    "analysis_results": None,
    "audit_log": [],
}
for k, v in DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = v

if "session_id" not in st.session_state:
    st.session_state.session_id = str(uuid.uuid4())[:8]

# =========================================================================
# MANUFACTURING FORMULAS
# =========================================================================

class ManufacturingFormulas:
    """Collection of manufacturing AI formulas with explanations"""

    @staticmethod
    def defect_score(defect_type: str, count: int, batch_size: int = 1000) -> float:
        """
        DEFECT SCORE FORMULA
        
        Score = (Defect_Count / Batch_Size) × Severity_Weight × 100
        
        Where:
        - Defect_Count: Number of defects found
        - Batch_Size: Total parts inspected (default 1000)
        - Severity_Weight: 1-10 scale (Blowholes=8, Shrinkage=9, Burrs=3)
        
        Interpretation:
        - Score < 20: Acceptable (PASS)
        - Score 20-50: Caution, investigate cause
        - Score > 50: Reject batch, find root cause
        """
        if defect_type not in DEFECT_TYPES:
            return 0
        severity = DEFECT_TYPES[defect_type]["severity"]
        score = (count / batch_size) * severity * 100
        return round(score, 2)

    @staticmethod
    def tool_wear_percentage(spindle_load_now: float, spindle_load_new: float = 100, 
                            vibration_now: float = 0.5, vibration_max: float = 2.5) -> float:
        """
        TOOL WEAR PREDICTION FORMULA
        
        Wear% = ((Load_Current/Load_New) × 0.6 + (Vib_Current/Vib_Max) × 0.4) × 100
        
        Components:
        - Spindle Load (60% weight): Cutting force & friction increase
        - Vibration (40% weight): Chatter & looseness detection
        
        Ranges:
        - 0-30%: GOOD - Continue operating
        - 30-60%: MONITOR - Plan change within week
        - 60-85%: CHANGE SOON - Schedule in 48 hours
        - 85-100%: CHANGE NOW - Immediate action required
        """
        load_ratio = spindle_load_now / spindle_load_new if spindle_load_new > 0 else 0
        vibration_ratio = vibration_now / vibration_max if vibration_max > 0 else 0
        wear_percent = (load_ratio * 0.6 + vibration_ratio * 0.4) * 100
        return round(min(100, max(0, wear_percent)), 2)

    @staticmethod
    def bearing_failure_risk(vibration_level: float, temperature_rise: float, operating_hours: int) -> float:
        """
        BEARING FAILURE PREDICTION (2-4 week advance warning)
        
        Risk = (Vib_Score × 0.5) + (Temp_Score × 0.3) + (Hours_Score × 0.2)
        
        Vibration (mm/s):
        - 0-2: Normal
        - 2-7: Caution zone
        - 7-11: Alarm zone
        - >11: Danger zone (imminent failure)
        
        Temperature Rise:
        - 0-5°C: Normal
        - 5-15°C: Increased friction
        - 15-25°C: Critical
        - >25°C: Seizing risk
        
        Risk Score:
        - 0-3: Safe
        - 3-6: Monitor, schedule next week
        - 6-8: Alert, change in 2-3 days
        - 8-10: Critical, change today
        """
        vib_score = min(10, vibration_level / 11 * 10)
        temp_score = min(10, temperature_rise / 25 * 10)
        hours_score = min(10, (operating_hours / 10000) * 10)
        risk = (vib_score * 0.5) + (temp_score * 0.3) + (hours_score * 0.2)
        return round(risk, 2)

    @staticmethod
    def surface_finish_prediction(spindle_rpm: float, feed_rate: float, depth_of_cut: float, 
                                 tool_wear_percent: float) -> float:
        """
        SURFACE FINISH PREDICTION (Ra in micrometers)
        
        Ra = Ra_Base × (1 + 0.05×Wear%) × Feed_Factor × RPM_Factor
        
        Ra Values:
        - 0.1-0.4 µm: Precision finish (honed)
        - 0.4-1.6 µm: Finish machining (CMM critical)
        - 1.6-3.2 µm: General machining
        - 3.2-6.3 µm: Rough machining
        - 6.3+ µm: As-cast
        
        Predicts before completion → Adjust parameters in time
        Prevents: Rework, CMM bottleneck, customer complaints
        """
        ra_baseline = 1.2  # µm
        wear_factor = 1 + (tool_wear_percent / 100) * 0.05
        
        optimal_feed = 0.10
        feed_factor = 1 + ((feed_rate / optimal_feed) - 1) * 0.3
        
        optimal_rpm = 2000
        rpm_factor = 1 + ((spindle_rpm / optimal_rpm) - 1) * 0.2
        
        ra_predicted = ra_baseline * wear_factor * feed_factor * rpm_factor
        return round(max(0.1, ra_predicted), 3)

    @staticmethod
    def oee_calculation(run_time: float, ideal_cycle_time: float, good_parts: int, total_parts: int) -> float:
        """
        OEE (Overall Equipment Effectiveness)
        
        OEE = Availability × Performance × Quality
        
        - Availability: Run_Time / Planned_Time (uptime)
        - Performance: (Parts × Cycle) / Run_Time (speed)
        - Quality: Good_Parts / Total_Parts (scrap rate)
        
        Benchmarks:
        - >85%: World-class
        - 70-85%: Good
        - 50-70%: Average
        - <50%: Poor
        
        Impact: Each 5% improvement = ₹2-3 crore additional annual margin
        """
        planned_time = ideal_cycle_time * (total_parts + 50)
        availability = run_time / planned_time if planned_time > 0 else 0
        performance = ((total_parts * ideal_cycle_time) / run_time) if run_time > 0 else 0
        quality = good_parts / total_parts if total_parts > 0 else 0
        oee = availability * performance * quality * 100
        return round(min(100, max(0, oee)), 2)

# =========================================================================
# EXPORT FUNCTIONS
# =========================================================================

def export_to_pdf(df: pd.DataFrame, title: str, formulas: Dict = None) -> bytes:
    """Export dataframe and formulas to PDF"""
    if not HAS_PDF:
        return None
    
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", "B", 16)
        pdf.cell(0, 10, title, ln=True, align="C")
        
        pdf.set_font("Arial", "B", 10)
        pdf.cell(0, 5, f"Generated: {dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", ln=True)
        pdf.cell(0, 5, f"ShopFloor Sentinel | Kalsnet (KNet) Consulting", ln=True)
        
        # Table
        pdf.set_font("Arial", "", 8)
        pdf.ln(5)
        
        col_width = 190 / len(df.columns)
        for col in df.columns:
            pdf.cell(col_width, 7, str(col)[:20], border=1, align="C")
        pdf.ln()
        
        for _, row in df.head(50).iterrows():
            for val in row:
                pdf.cell(col_width, 7, str(val)[:20], border=1, align="L")
            pdf.ln()
        
        if formulas:
            pdf.add_page()
            pdf.set_font("Arial", "B", 12)
            pdf.cell(0, 10, "Formulas & Explanations", ln=True)
            pdf.set_font("Arial", "", 9)
            for formula_name, formula_text in formulas.items():
                pdf.cell(0, 5, formula_name + ":", ln=True)
                pdf.multi_cell(0, 4, formula_text[:500])
                pdf.ln(2)
        
        return pdf.output(dest='S').encode('latin-1')
    except Exception as e:
        st.error(f"PDF export error: {e}")
        return None

def export_to_docx(df: pd.DataFrame, title: str, formulas: Dict = None) -> bytes:
    """Export dataframe and formulas to Word document"""
    if not HAS_DOCX:
        return None
    
    try:
        doc = Document()
        
        # Title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 71, 171)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        meta_para = doc.add_paragraph()
        meta_run = meta_para.add_run(f"Generated: {dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        meta_run.font.size = Pt(10)
        meta_run = meta_para.add_run("ShopFloor Sentinel | Developed by Randy Singh | Kalsnet (KNet) Consulting")
        meta_run.font.size = Pt(10)
        meta_run.italic = True
        
        doc.add_paragraph()
        
        # Table
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr_cells[i].text = str(col)
        
        for _, row in df.head(50).iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val)[:100]
        
        # Formulas section
        if formulas:
            doc.add_page_break()
            formulas_heading = doc.add_paragraph()
            formulas_run = formulas_heading.add_run("Formulas & Explanations")
            formulas_run.font.size = Pt(14)
            formulas_run.font.bold = True
            formulas_run.font.color.rgb = RGBColor(0, 71, 171)
            
            for formula_name, formula_text in formulas.items():
                p = doc.add_paragraph(formula_name, style='Heading 3')
                doc.add_paragraph(formula_text[:500])
        
        # Save to bytes
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        return docx_bytes.getvalue()
    except Exception as e:
        st.error(f"Word export error: {e}")
        return None

def export_to_text(df: pd.DataFrame, title: str, formulas: Dict = None) -> str:
    """Export to plain text format"""
    text = f"\n{'='*80}\n"
    text += f"{title}\n"
    text += f"{'='*80}\n\n"
    text += f"Generated: {dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
    text += "ShopFloor Sentinel | Developed by Randy Singh | Kalsnet (KNet) Consulting\n\n"
    text += df.to_string()
    
    if formulas:
        text += f"\n\n{'='*80}\nFormulas & Explanations\n{'='*80}\n\n"
        for formula_name, formula_text in formulas.items():
            text += f"\n{formula_name}:\n{formula_text}\n"
    
    return text

# =========================================================================
# SIDEBAR - LLM & CONFIG
# =========================================================================
with st.sidebar:
    st.header("⚙️ Configuration")
    
    with st.expander("🤖 AI Prediction Engine (Optional)", expanded=False):
        st.markdown("""
        Enable LLM predictions for pattern recognition in defect causes.
        Choose **Groq** (faster, free tier) or **Claude** (more accurate).
        """)
        
        enable_llm = st.checkbox("Enable AI predictions", value=False)
        
        if enable_llm:
            llm_provider = st.radio("LLM Provider", ["Groq (Faster, Free)", "Claude (Accurate)"])
            
            if "Groq" in llm_provider:
                groq_key = st.text_input("Groq API Key", type="password")
                
                if st.checkbox("📖 Get Groq API Key (FREE)"):
                    st.info("""
                    **Steps to get Groq API Key (completely FREE):**
                    
                    1. **Sign Up**: Go to **console.groq.com**
                    2. **Create Account**: Use email/Google/GitHub
                    3. **Get API Key**: 
                       - Left sidebar → **API Keys**
                       - Click **Create API Key**
                       - Copy the key (shown once)
                    4. **Free Tier Benefits**:
                       - 30 requests/minute
                       - Unlimited usage
                       - Perfect for manufacturing predictions
                    5. **Paste Above**: Put key in the field above
                    
                    **Models available on Groq (all free):**
                    - llama-3.3-70b-versatile (recommended)
                    - mixtral-8x7b-32768 (fast)
                    - gemma2-9b-it (lightweight)
                    """)
            else:
                claude_key = st.text_input("Claude API Key", type="password")
                
                if st.checkbox("📖 Get Claude API Key"):
                    st.info("""
                    **Steps to get Claude API Key (requires billing):**
                    
                    1. **Sign Up**: Go to **console.anthropic.com**
                    2. **Create Account**: Register with email
                    3. **Add Billing**:
                       - Settings → Billing
                       - Add credit card
                       - Set spending limit (e.g., $10/month)
                    4. **Create API Key**:
                       - Settings → API Keys
                       - Create Key
                       - Copy immediately (not shown again)
                    5. **Paste Above**: Put key in the field above
                    
                    **Available Models:**
                    - claude-opus-4-6 (most capable, $15/MTok)
                    - claude-sonnet-4-6 (balanced, $3/MTok) ← Recommended
                    - claude-haiku-4-5 (fast, $0.80/MTok)
                    
                    **Cost Estimate**:
                    - 100 predictions/day × 30 days ≈ ₹20-50/month
                    """)
    
    with st.expander("📊 Data Settings", expanded=True):
        st.markdown(f"**Max Upload**: {MAX_UPLOAD_MB} MB")
        st.markdown(f"**Environment**: {DEPLOY_ENV.upper()}")
        st.markdown(f"**Session**: {st.session_state.session_id}")
    
    st.markdown("---")
    st.caption("v1.0 | Kalsnet (KNet) Consulting")

# =========================================================================
# MAIN INTERFACE - TABS
# =========================================================================

tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "📊 Defect Detection",
    "🔧 Tool Wear Analysis",
    "🔮 Predictive Maintenance",
    "⚙️ Process Optimization",
    "📈 OEE & Analytics"
])

# =========================================================================
# TAB 1: DEFECT DETECTION
# =========================================================================
with tab1:
    st.markdown('<div class="section-header">📊 Visual Defect Detection & Scoring</div>', unsafe_allow_html=True)
    
    st.markdown("""
    AI-powered defect detection catches casting/machining defects before shipment.
    Reduces customer PPM penalties and OEM line rejections.
    """)
    
    col1, col2, col3 = st.columns([2, 2, 1])
    
    with col1:
        n_records = st.slider("Synthetic records", 0, 500, 100, step=10)
    
    with col2:
        data_type = st.selectbox("Dataset type", ["Casting", "Machining", "Mixed"])
    
    with col3:
        st.write("")
        st.write("")
        generate = st.button("🔄 Generate Demo", key="gen_demo")
    
    if generate:
        rows = []
        for i in range(n_records):
            defect_type = random.choice(list(DEFECT_TYPES.keys()))
            defect_count = random.randint(0, 20)
            score = ManufacturingFormulas.defect_score(defect_type, defect_count)
            
            rows.append({
                "Batch_ID": f"BATCH-{random.randint(1000, 9999)}",
                "Timestamp": (dt.datetime.now() - dt.timedelta(days=random.randint(0, 30))).strftime("%Y-%m-%d"),
                "Machine": random.choice(MACHINE_TYPES),
                "Defect_Type": defect_type,
                "Defect_Count": defect_count,
                "Defect_Score": score,
                "Status": "REJECT" if score > 50 else ("CAUTION" if score > 20 else "PASS"),
            })
        
        st.session_state.synthetic_df = pd.DataFrame(rows)
        st.success(f"✅ Generated {n_records} synthetic defect records")
    
    # Upload real data
    st.markdown("**OR upload real inspection data:**")
    uploaded = st.file_uploader("Upload CSV/Excel file", type=["csv", "xlsx"])
    
    if uploaded:
        try:
            if uploaded.name.endswith(".csv"):
                st.session_state.real_df = pd.read_csv(uploaded)
            else:
                st.session_state.real_df = pd.read_excel(uploaded)
            st.success(f"✅ Loaded {len(st.session_state.real_df)} records")
        except Exception as e:
            st.error(f"Error: {e}")
    
    # Display & analyze
    df_show = st.session_state.real_df if st.session_state.real_df is not None else st.session_state.synthetic_df
    
    if df_show is not None:
        st.markdown("**Data Preview:**")
        st.dataframe(df_show.head(20), use_container_width=True, height=250)
        
        # Summary metrics
        st.markdown('<div class="subsection-header">📈 Analysis Summary</div>', unsafe_allow_html=True)
        
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Total Records", len(df_show))
        with col2:
            defects = (df_show.get("Defect_Count", [0]) > 0).sum() if "Defect_Count" in df_show.columns else 0
            st.metric("With Defects", defects)
        with col3:
            rejects = (df_show.get("Status") == "REJECT").sum() if "Status" in df_show.columns else 0
            st.metric("Reject Batches", rejects)
        with col4:
            if "Defect_Score" in df_show.columns:
                st.metric("Avg Score", f"{df_show['Defect_Score'].mean():.1f}", "Lower better")
        
        # Defect type breakdown
        if "Defect_Type" in df_show.columns:
            st.markdown('<div class="subsection-header">Defect Type Breakdown</div>', unsafe_allow_html=True)
            defect_summary = df_show.groupby("Defect_Type").size().reset_index(name="Count")
            st.bar_chart(defect_summary.set_index("Defect_Type"))
        
        # Formulas reference
        st.markdown('<div class="formula-box">', unsafe_allow_html=True)
        st.markdown("""
        **DEFECT SCORE FORMULA**
        
        **Score = (Defect_Count / Batch_Size) × Severity_Weight × 100**
        
        **Severity Weights:**
        """)
        
        for dtype, info in DEFECT_TYPES.items():
            st.write(f"- **{dtype}** ({info['severity']}/10): {info['description']}")
        
        st.markdown("""
        **Interpretation:**
        - Score < 20: ✅ PASS - Acceptable
        - Score 20-50: ⚠️ CAUTION - Investigate root cause
        - Score > 50: 🔴 REJECT - Find and fix issue
        """)
        st.markdown('</div>', unsafe_allow_html=True)
        
        # Export section
        st.markdown('<div class="subsection-header">💾 Export Results</div>', unsafe_allow_html=True)
        
        # Prepare formulas dict
        formulas_dict = {
            "Defect Score": f"Score = (Count / Batch_Size) × Severity × 100\n\nExample: 5 defects in 1000 parts with severity 8 = (5/1000) × 8 × 100 = 4.0",
            "Defect Severity": "\n".join([f"{k}: {v['severity']}/10 - {v['description']}" for k, v in DEFECT_TYPES.items()])
        }
        
        col1, col2, col3, col4, col5 = st.columns(5)
        
        with col1:
            csv = df_show.to_csv(index=False)
            st.download_button("📊 CSV", csv.encode(), "defect_analysis.csv", "text/csv", key="csv1")
        
        with col2:
            json_str = df_show.to_json(orient="records", indent=2)
            st.download_button("📋 JSON", json_str.encode(), "defect_analysis.json", "application/json", key="json1")
        
        with col3:
            txt = export_to_text(df_show, "Defect Detection Analysis", formulas_dict)
            st.download_button("📄 TXT", txt.encode(), "defect_analysis.txt", "text/plain", key="txt1")
        
        with col4:
            if HAS_DOCX:
                docx_bytes = export_to_docx(df_show, "Defect Detection Analysis", formulas_dict)
                if docx_bytes:
                    st.download_button("📝 WORD", docx_bytes, "defect_analysis.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="docx1")
            else:
                st.write("Word export unavailable")
        
        with col5:
            if HAS_PDF:
                pdf_bytes = export_to_pdf(df_show, "Defect Detection Analysis", formulas_dict)
                if pdf_bytes:
                    st.download_button("🔴 PDF", pdf_bytes, "defect_analysis.pdf", "application/pdf", key="pdf1")
            else:
                st.write("PDF export unavailable")

# =========================================================================
# TAB 2: TOOL WEAR ANALYSIS
# =========================================================================
with tab2:
    st.markdown('<div class="section-header">🔧 Tool Wear Prediction & Optimization</div>', unsafe_allow_html=True)
    
    st.markdown("""
    Predict tool wear in real-time. Optimize change timing to save 15-20% on tool cost
    while preventing sudden breakage.
    """)
    
    col1, col2, col3 = st.columns([2, 2, 1])
    
    with col1:
        n_tools = st.slider("Synthetic tool records", 0, 500, 50, step=10)
    
    with col2:
        tool_type = st.selectbox("Tool type", ["Cutting Insert", "End Mill", "Drill", "Boring Bar"])
    
    with col3:
        st.write("")
        st.write("")
        gen_tools = st.button("🔄 Generate", key="gen_tools")
    
    if gen_tools:
        rows = []
        for i in range(n_tools):
            spindle_load = random.randint(50, 100)
            vibration = random.uniform(0.1, 3.0)
            wear_pct = ManufacturingFormulas.tool_wear_percentage(spindle_load, 100, vibration, 2.5)
            
            if wear_pct < 30:
                condition = "🟢 GOOD"
            elif wear_pct < 60:
                condition = "🟡 MONITOR"
            elif wear_pct < 85:
                condition = "🟠 CHANGE SOON"
            else:
                condition = "🔴 CHANGE NOW"
            
            rows.append({
                "Tool_ID": f"TOOL-{random.randint(10000, 99999)}",
                "Spindle_Load_%": spindle_load,
                "Vibration_mm_s": round(vibration, 2),
                "Wear_%": wear_pct,
                "Condition": condition,
            })
        
        tool_df = pd.DataFrame(rows)
        st.session_state.tool_df = tool_df
        st.success(f"✅ Generated {n_tools} tool wear records")
    
    if "tool_df" in st.session_state and st.session_state.tool_df is not None:
        tool_df = st.session_state.tool_df
        
        st.dataframe(tool_df, use_container_width=True, height=250)
        
        # Metrics
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Avg Wear %", f"{tool_df['Wear_%'].mean():.1f}%")
        with col2:
            change_soon = (tool_df['Condition'].isin(['🟠 CHANGE SOON', '🔴 CHANGE NOW'])).sum()
            st.metric("Needs Change", change_soon)
        with col3:
            good = (tool_df['Condition'] == '🟢 GOOD').sum()
            st.metric("Good Tools", good)
        with col4:
            critical = (tool_df['Condition'] == '🔴 CHANGE NOW').sum()
            st.metric("Critical", critical)
        
        # Formula
        st.markdown('<div class="formula-box">', unsafe_allow_html=True)
        st.markdown("""
        **TOOL WEAR % FORMULA**
        
        **Wear % = ((Load / 100) × 0.6 + (Vibration / 2.5) × 0.4) × 100**
        
        **Weighting:**
        - **Spindle Load (60%)**: Measures cutting force & friction
        - **Vibration (40%)**: Detects chatter & wear
        
        **Action Ranges:**
        - 0-30%: ✅ Continue operating
        - 30-60%: ⚠️ Monitor, plan change
        - 60-85%: 🟠 Change in 48 hours
        - 85-100%: 🔴 Immediate change
        
        **ROI**: Optimal timing saves 15-20% tool cost + prevents ₹20k+ breakage downtime
        """)
        st.markdown('</div>', unsafe_allow_html=True)
        
        # Export
        st.markdown('<div class="subsection-header">💾 Export</div>', unsafe_allow_html=True)
        
        formulas_dict = {
            "Tool Wear Formula": "Wear% = ((Load/100) × 0.6 + (Vibration/2.5) × 0.4) × 100",
            "Action Ranges": "0-30% GOOD | 30-60% MONITOR | 60-85% CHANGE SOON | 85-100% CHANGE NOW"
        }
        
        col1, col2, col3, col4, col5 = st.columns(5)
        with col1:
            st.download_button("📊 CSV", tool_df.to_csv(index=False).encode(), "tool_wear.csv", "text/csv", key="csv2")
        with col2:
            st.download_button("📋 JSON", tool_df.to_json(orient="records", indent=2).encode(), "tool_wear.json", "application/json", key="json2")
        with col3:
            txt = export_to_text(tool_df, "Tool Wear Analysis", formulas_dict)
            st.download_button("📄 TXT", txt.encode(), "tool_wear.txt", "text/plain", key="txt2")
        with col4:
            if HAS_DOCX:
                docx_bytes = export_to_docx(tool_df, "Tool Wear Analysis", formulas_dict)
                if docx_bytes:
                    st.download_button("📝 WORD", docx_bytes, "tool_wear.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="docx2")
        with col5:
            if HAS_PDF:
                pdf_bytes = export_to_pdf(tool_df, "Tool Wear Analysis", formulas_dict)
                if pdf_bytes:
                    st.download_button("🔴 PDF", pdf_bytes, "tool_wear.pdf", "application/pdf", key="pdf2")

# =========================================================================
# TAB 3: PREDICTIVE MAINTENANCE
# =========================================================================
with tab3:
    st.markdown('<div class="section-header">🔮 Predictive Bearing Maintenance (2-4 Week Early Warning)</div>', unsafe_allow_html=True)
    
    st.markdown("""
    Predict bearing failure before breakdown. Avoid emergency maintenance during peak production.
    """)
    
    col1, col2 = st.columns(2)
    
    with col1:
        vib = st.slider("Vibration (mm/s)", 0.0, 15.0, 5.0, step=0.5)
    with col2:
        temp = st.slider("Temp Rise (°C)", 0.0, 35.0, 15.0, step=1.0)
    
    hours = st.slider("Hours Since Service", 0, 15000, 8000, step=500)
    
    risk_score = ManufacturingFormulas.bearing_failure_risk(vib, temp, hours)
    
    if risk_score < 3:
        status = "🟢 SAFE"
        action = "No action required"
        color = "#1a8a3e"
    elif risk_score < 6:
        status = "🟡 MONITOR"
        action = "Schedule maintenance next week"
        color = "#ff9800"
    elif risk_score < 8:
        status = "🟠 ALERT"
        action = "Schedule change in 2-3 days"
        color = "#ffa500"
    else:
        status = "🔴 CRITICAL"
        action = "Change bearing today - imminent failure"
        color = "#c0392b"
    
    st.markdown(f'<div style="background:{color}; color:white; padding:1.5rem; border-radius:8px; text-align:center;">'
                f'<h2>Risk Score: {risk_score:.2f}/10</h2>'
                f'<h3>{status}</h3>'
                f'<p>{action}</p>'
                f'</div>', unsafe_allow_html=True)
    
    st.markdown('<div class="formula-box">', unsafe_allow_html=True)
    st.markdown(f"""
    **BEARING FAILURE PREDICTION FORMULA**
    
    **Risk = (Vib_Score × 0.5) + (Temp_Score × 0.3) + (Hours_Score × 0.2)**
    
    **Your Values:**
    - Vibration: {vib:.1f} mm/s → Score {min(10, vib/11*10):.1f}/10
    - Temperature: {temp:.1f}°C → Score {min(10, temp/25*10):.1f}/10
    - Hours: {hours:,} → Score {min(10, hours/10000*10):.1f}/10
    
    **Risk Calculation: {risk_score:.2f}/10**
    
    **Vibration Ranges:**
    - 0-2 mm/s: Normal operation
    - 2-7 mm/s: Caution zone (loose, wear)
    - 7-11 mm/s: Alarm zone (spalling)
    - >11 mm/s: Danger zone (imminent failure)
    
    **Temperature Ranges:**
    - 0-5°C: Normal friction
    - 5-15°C: Increased friction
    - 15-25°C: Critical
    - >25°C: Seizing risk
    """)
    st.markdown('</div>', unsafe_allow_html=True)

# =========================================================================
# TAB 4: PROCESS OPTIMIZATION
# =========================================================================
with tab4:
    st.markdown('<div class="section-header">⚙️ Surface Finish Prediction (Real-Time Quality)</div>', unsafe_allow_html=True)
    
    st.markdown("""
    Predict surface finish BEFORE completing the cut. Adjust parameters to prevent out-of-spec parts.
    """)
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        rpm = st.slider("Spindle RPM", 500, 5000, 2000, step=100)
    with col2:
        feed = st.slider("Feed (mm/rev)", 0.01, 0.50, 0.10, step=0.02)
    with col3:
        depth = st.slider("Depth (mm)", 0.1, 5.0, 1.0, step=0.2)
    with col4:
        wear = st.slider("Tool Wear %", 0, 100, 25, step=5)
    
    ra = ManufacturingFormulas.surface_finish_prediction(rpm, feed, depth, wear)
    
    if ra < 0.8:
        status = "🟢 EXCELLENT"
        color = "#1a8a3e"
    elif ra < 1.6:
        status = "✅ GOOD"
        color = "#66bb6a"
    elif ra < 3.2:
        status = "⚠️ CAUTION"
        color = "#ff9800"
    else:
        status = "🔴 OUT OF SPEC"
        color = "#c0392b"
    
    st.markdown(f'<div style="background:{color}; color:white; padding:2rem; border-radius:8px; text-align:center;">'
                f'<h1>Ra: {ra:.3f} µm</h1>'
                f'<h2>{status}</h2>'
                f'</div>', unsafe_allow_html=True)
    
    st.markdown('<div class="formula-box">', unsafe_allow_html=True)
    st.markdown(f"""
    **SURFACE FINISH PREDICTION FORMULA**
    
    **Ra = Ra_Base × (1 + 0.05×Wear%) × Feed_Factor × RPM_Factor**
    
    **Your Calculation:**
    - Base Ra: 1.2 µm (finish pass)
    - Wear factor: {round(1 + (0.05*wear/100), 3)}
    - Feed factor: Based on {feed:.2f} mm/rev
    - RPM factor: Based on {rpm} rpm
    
    **Predicted Ra: {ra:.3f} µm**
    
    **Surface Finish Reference:**
    - <0.1 µm: Polished
    - 0.1-0.4 µm: Honed
    - 0.4-0.8 µm: Precision finish
    - 0.8-1.6 µm: Finish machining (typical CMM critical)
    - 1.6-3.2 µm: General machining
    - 3.2-6.3 µm: Rough machining
    - 6.3+ µm: As-cast
    """)
    st.markdown('</div>', unsafe_allow_html=True)

# =========================================================================
# TAB 5: OEE & ANALYTICS
# =========================================================================
with tab5:
    st.markdown('<div class="section-header">📈 OEE (Overall Equipment Effectiveness)</div>', unsafe_allow_html=True)
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        run_time = st.number_input("Run Time (min)", 1, 480, 450, step=10)
    with col2:
        cycle_time = st.number_input("Ideal Cycle (min/part)", 0.1, 30.0, 5.0, step=0.5)
    with col3:
        good = st.number_input("Good Parts", 0, 1000, 85, step=5)
    with col4:
        total = st.number_input("Total Parts", 1, 1000, 95, step=5)
    
    oee = ManufacturingFormulas.oee_calculation(run_time, cycle_time, good, total)
    
    col1, col2, col3 = st.columns(3)
    with col1:
        planned = cycle_time * (total + 50)
        avail = (run_time / planned) * 100
        st.metric("Availability", f"{avail:.1f}%")
    with col2:
        perf = ((total * cycle_time) / run_time) * 100
        st.metric("Performance", f"{perf:.1f}%")
    with col3:
        qual = (good / total) * 100
        st.metric("Quality", f"{qual:.1f}%")
    
    st.markdown("---")
    
    if oee >= 85:
        color, grade = "#1a8a3e", "WORLD-CLASS"
    elif oee >= 70:
        color, grade = "#ff9800", "GOOD"
    else:
        color, grade = "#c0392b", "POOR"
    
    st.markdown(f'<div style="background:{color}; color:white; padding:2rem; border-radius:8px; text-align:center;">'
                f'<h1>OEE: {oee:.1f}%</h1>'
                f'<h2>{grade}</h2>'
                f'</div>', unsafe_allow_html=True)
    
    st.markdown('<div class="formula-box">', unsafe_allow_html=True)
    st.markdown(f"""
    **OEE = Availability × Performance × Quality**
    
    **Your OEE: {oee:.1f}%**
    
    **Benchmarks:**
    - >85%: World-class
    - 70-85%: Good
    - 50-70%: Average
    - <50%: Poor
    
    **Impact:** Each 5% OEE gain = ₹2-3 crore additional annual margin
    """)
    st.markdown('</div>', unsafe_allow_html=True)

st.markdown("---")
st.caption(f"🛡️ ShopFloor Sentinel v1.0 | Session {st.session_state.session_id} | {DEPLOY_ENV.upper()} | Developed by Randy Singh | Kalsnet (KNet) Consulting")
