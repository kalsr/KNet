

# CyberSentinel Intelligence Platform
# AI-powered cyber threat detection, threat intelligence, and predictive risk scoring
# Developed by Randy Singh | Kalsnet (KNet) Consulting
# EXPORT CAPABILITIES: PDF, Word (.docx), CSV, JSON, TXT
import io
import json
import logging
import os
import random
import uuid
import datetime as dt
from typing import Dict, List
import pandas as pd
import numpy as np
import streamlit as st

try:
    from fpdf import FPDF
    HAS_PDF = True
except ImportError:
    HAS_PDF = False
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# =========================================================================
# LOGGING
# =========================================================================
LOG_DIR = os.environ.get("CYBERSENTINEL_LOG_DIR", "./logs")
os.makedirs(LOG_DIR, exist_ok=True)
logger = logging.getLogger("cybersentinel")
if not logger.handlers:
    logger.setLevel(logging.INFO)
    _fh = logging.FileHandler(os.path.join(LOG_DIR, "cybersentinel.log"))
    _fh.setFormatter(logging.Formatter("%(asctime)s | %(levelname)s | %(message)s"))
    logger.addHandler(_fh)

# =========================================================================
# PAGE CONFIG
# =========================================================================
st.set_page_config(
    page_title="CyberSentinel Intelligence Platform",
    page_icon="",
    layout="wide",
    initial_sidebar_state="expanded",
)

# =========================================================================
# PROFESSIONAL STYLING
# =========================================================================
st.markdown(
    """
    <style>
    .main-title {
        text-align: center;
        color: #0047AB;
        font-weight: 900;
        font-size: 2.8rem;
        margin-bottom: 0.3rem;
        letter-spacing: 0.8px;
        font-family: 'Arial Black', sans-serif;
    }
    .subtitle-line {
        text-align: center;
        color: #0047AB;
        font-weight: 800;
        font-size: 1.2rem;
        margin-top: 0;
        margin-bottom: 1.2rem;
        font-family: 'Arial', sans-serif;
    }
    .section-header {
        color: #0047AB;
        font-weight: 800;
        font-size: 1.8rem;
        margin-top: 1.5rem;
        margin-bottom: 0.8rem;
        border-bottom: 3px solid #0047AB;
        padding-bottom: 0.5rem;
    }
    .subsection-header {
        color: #003380;
        font-weight: 700;
        font-size: 1.3rem;
        margin-top: 1rem;
        margin-bottom: 0.6rem;
    }
    .formula-box {
        background: #fff9e6;
        border: 2px solid #ffd700;
        padding: 1rem;
        border-radius: 6px;
        font-family: 'Courier New', monospace;
        font-size: 0.9rem;
        margin: 0.8rem 0;
    }
    .divider {
        border-top: 2px solid #0047AB;
        margin: 1.5rem 0;
        opacity: 0.3;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# =========================================================================
# TITLE BAR
# =========================================================================
st.markdown('<div class="main-title"> CyberSentinel Intelligence Platform</div>', unsafe_allow_html=True)
st.markdown('<div class="subtitle-line">Developed by Randy Singh from Kalsnet (KNet) Consulting</div>', unsafe_allow_html=True)
st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
st.markdown(
    """
    **Unified Cyber + Physical/Operational Threat Intelligence Platform** — Real-time AI for
    threat detection, IOC confidence scoring, predictive risk forecasting, and vulnerability/asset
    risk analysis across both digital and physical security domains.

    **Focus:** Detect faster • Prioritize smarter • Predict escalation before it happens — across every domain
    """
)

# =========================================================================
# ENVIRONMENT CONFIG
# =========================================================================
DEPLOY_ENV = os.environ.get("CYBERSENTINEL_ENV", "production")
MAX_UPLOAD_MB = int(os.environ.get("CYBERSENTINEL_MAX_UPLOAD_MB", "25"))

# =========================================================================
# CONSTANTS
# =========================================================================
CYBER_THREAT_TYPES = {
    "Malware Infection": {"severity": 9, "category": "Endpoint", "domain": "Cyber", "description": "Malicious software detected on host, may spread laterally"},
    "Ransomware": {"severity": 10, "category": "Endpoint", "domain": "Cyber", "description": "File encryption / extortion attempt, business-critical"},
    "Phishing Attempt": {"severity": 6, "category": "Email", "domain": "Cyber", "description": "Credential harvesting or malicious attachment via email"},
    "Brute Force Attack": {"severity": 7, "category": "Network", "domain": "Cyber", "description": "Repeated authentication attempts against a service"},
    "DDoS Attack": {"severity": 8, "category": "Network", "domain": "Cyber", "description": "Volumetric or application-layer availability attack"},
    "SQL Injection": {"severity": 9, "category": "Application", "domain": "Cyber", "description": "Malicious query injection against a data-backed application"},
    "Unauthorized Access (Digital)": {"severity": 8, "category": "Identity", "domain": "Cyber", "description": "Access outside of normal permission or behavior baseline"},
    "Data Exfiltration": {"severity": 10, "category": "Data Loss", "domain": "Cyber", "description": "Abnormal outbound transfer of sensitive data"},
    "Insider Threat (Digital)": {"severity": 8, "category": "Identity", "domain": "Cyber", "description": "Anomalous activity by an authenticated internal user"},
    "Port Scanning": {"severity": 3, "category": "Network", "domain": "Cyber", "description": "Reconnaissance activity probing open services"},
    "Privilege Escalation": {"severity": 9, "category": "Identity", "domain": "Cyber", "description": "Attempt to gain elevated system or account rights"},
    "Command & Control (C2) Beacon": {"severity": 10, "category": "Network", "domain": "Cyber", "description": "Outbound traffic consistent with known C2 infrastructure"},
}
PHYSICAL_THREAT_TYPES = {
    "Unauthorized Facility Access": {"severity": 8, "category": "Perimeter", "domain": "Physical/Operational", "description": "Badge tailgating, forced entry, or unrecognized credential at a controlled point"},
    "Perimeter Breach": {"severity": 8, "category": "Perimeter", "domain": "Physical/Operational", "description": "Fence, gate, or barrier compromise detected by sensor or patrol"},
    "Suspicious Package / Object": {"severity": 9, "category": "Safety", "domain": "Physical/Operational", "description": "Unattended or unidentified item reported near a facility or asset"},
    "Insider Threat (Physical)": {"severity": 8, "category": "Personnel", "domain": "Physical/Operational", "description": "Employee or contractor behavior inconsistent with role or access needs"},
    "Equipment Tampering / Sabotage": {"severity": 9, "category": "Operational", "domain": "Physical/Operational", "description": "Evidence of deliberate interference with machinery, utilities, or controls"},
    "Fire / Hazmat Incident": {"severity": 10, "category": "Safety", "domain": "Physical/Operational", "description": "Fire, chemical spill, or hazardous material exposure event"},
    "Workplace Violence Indicator": {"severity": 9, "category": "Personnel", "domain": "Physical/Operational", "description": "Threatening behavior, altercation, or escalation reported on-site"},
    "Supply Chain / Vendor Disruption": {"severity": 6, "category": "Operational", "domain": "Physical/Operational", "description": "Compromise or failure at a third-party vendor affecting operations"},
    "Surveillance / Reconnaissance Activity": {"severity": 5, "category": "Perimeter", "domain": "Physical/Operational", "description": "Repeated observation of a facility consistent with pre-attack planning"},
    "Utility / Environmental Failure": {"severity": 7, "category": "Operational", "domain": "Physical/Operational", "description": "Power, HVAC, or environmental control failure affecting critical operations"},
    "Lost/Stolen Access Credential": {"severity": 7, "category": "Personnel", "domain": "Physical/Operational", "description": "Badge, key, or physical token reported lost or stolen"},
    "Vehicle / Perimeter Ramming Concern": {"severity": 9, "category": "Perimeter", "domain": "Physical/Operational", "description": "Vehicle behavior or bollard/barrier alert near a facility perimeter"},
}
THREAT_TYPES = {**CYBER_THREAT_TYPES, **PHYSICAL_THREAT_TYPES}
DOMAIN_OPTIONS = ["Mixed (Cyber + Physical/Operational)", "Cyber Only", "Physical/Operational Only"]
ASSET_TYPES_CYBER = [
    "Web Server", "Database Server", "Domain Controller", "Endpoint-Workstation",
    "Firewall/Edge Device", "Email Server", "Cloud Storage Bucket", "IoT Device", "VPN Gateway"
]
ASSET_TYPES_PHYSICAL = [
    "Facility - Main Entrance", "Perimeter Fence/Gate", "Data Center (Physical)", "Warehouse/Loading Dock",
    "Executive Office", "Production Floor", "Utility/Mechanical Room", "Parking Structure", "Server Closet"
]
ASSET_TYPES = ASSET_TYPES_CYBER + ASSET_TYPES_PHYSICAL
STATUS_OPTIONS = ["Detected", "Investigating", "Contained", "Resolved", "False Positive"]
IOC_TYPES_CYBER = ["IP Address", "Domain", "File Hash (SHA-256)", "URL", "Email Address"]
IOC_TYPES_PHYSICAL = ["Badge/Credential ID", "Vehicle License Plate", "CCTV Detection Tag", "Access Control Alert", "Tip Line Report"]
IOC_TYPES = IOC_TYPES_CYBER + IOC_TYPES_PHYSICAL
SOURCE_FEEDS_CYBER = ["Internal SIEM", "Commercial Feed A", "Commercial Feed B", "OSINT Feed", "ISAC Sharing", "Government Advisory"]
SOURCE_FEEDS_PHYSICAL = ["Security Guard Report", "CCTV Analytics", "Access Control System", "Local Law Enforcement Tip", "Employee Report"]
SOURCE_FEEDS = SOURCE_FEEDS_CYBER + SOURCE_FEEDS_PHYSICAL

# =========================================================================
# SESSION STATE
# =========================================================================
DEFAULTS = {
    "synthetic_df": None,
    "real_df": None,
    "ioc_df": None,
}
for k, v in DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = v
if "session_id" not in st.session_state:
    st.session_state.session_id = str(uuid.uuid4())[:8]

# =========================================================================
# THREAT INTELLIGENCE FORMULAS
# =========================================================================
class ThreatFormulas:
    """Collection of threat intelligence / prediction formulas with explanations"""

    @staticmethod
    def threat_risk_score(threat_type: str, exploitability: float, asset_criticality: float) -> float:
        """
        THREAT RISK SCORE (TRS)

        TRS = (Severity/10 × 0.4 + Exploitability/10 × 0.3 + Asset_Criticality/10 × 0.3) × 100

        Where:
        - Severity: Inherent severity of the threat type (1-10)
        - Exploitability: How easily this threat can be exploited right now (1-10)
        - Asset_Criticality: Business importance of the targeted asset (1-10)

        Interpretation:
        - Score < 30: LOW - Log and monitor
        - Score 30-60: MEDIUM - Investigate within SLA
        - Score 60-85: HIGH - Escalate to analyst immediately
        - Score > 85: CRITICAL - Incident response activation
        """
        if threat_type not in THREAT_TYPES:
            return 0
        severity = THREAT_TYPES[threat_type]["severity"]
        score = ((severity / 10) * 0.4 + (exploitability / 10) * 0.3 + (asset_criticality / 10) * 0.3) * 100
        return round(min(100, max(0, score)), 2)

    @staticmethod
    def ioc_confidence_score(source_reliability: float, corroboration_count: int, recency_days: int) -> float:
        """
        IOC CONFIDENCE SCORE

        Confidence = (Source_Reliability × 0.40 + Corroboration_Factor × 0.35 + Recency_Factor × 0.25) × 100

        Where:
        - Source_Reliability: 0-1 rating of the intel feed's historical accuracy
        - Corroboration_Factor: min(1, corroborating_sources / 5)
        - Recency_Factor: decays as the indicator ages (fresher = more confident)

        Interpretation:
        - < 40%: LOW - Informational only
        - 40-70%: MEDIUM - Monitor, do not auto-block
        - > 70%: HIGH - Actionable, safe to auto-block
        """
        corroboration_factor = min(1.0, corroboration_count / 5)
        recency_factor = max(0.0, 1 - (recency_days / 90))
        confidence = (source_reliability * 0.40 + corroboration_factor * 0.35 + recency_factor * 0.25) * 100
        return round(min(100, max(0, confidence)), 2)

    @staticmethod
    def escalation_probability(event_frequency_trend: float, avg_confidence: float, avg_severity: float) -> float:
        """
        PREDICTIVE THREAT ESCALATION PROBABILITY

        Escalation% = (Frequency_Trend × 0.5 + Confidence_Norm × 0.3 + Severity_Norm × 0.2) × 100

        Where:
        - Frequency_Trend: normalized rate-of-change in event volume over the observed window (0-1)
        - Confidence_Norm: average IOC confidence, normalized (0-1)
        - Severity_Norm: average threat severity, normalized (0-1)

        This estimates the likelihood a cluster of related events escalates into
        a confirmed incident in the next observation window, based on momentum
        rather than a single point-in-time score.

        Interpretation:
        - < 25%: Unlikely to escalate
        - 25-50%: Possible - keep watching
        - 50-75%: Likely - proactively prepare response
        - > 75%: Highly likely - escalate now
        """
        freq_norm = min(1.0, max(0.0, event_frequency_trend))
        conf_norm = min(1.0, max(0.0, avg_confidence / 100))
        sev_norm = min(1.0, max(0.0, avg_severity / 10))
        escalation = (freq_norm * 0.5 + conf_norm * 0.3 + sev_norm * 0.2) * 100
        return round(min(100, max(0, escalation)), 2)

    @staticmethod
    def attack_surface_risk_index(open_ports: int, known_vulns: int, days_since_patch: int, internet_facing: bool) -> float:
        """
        ATTACK SURFACE RISK INDEX (ASRI)

        ASRI = (Port_Score × 0.25 + Vuln_Score × 0.40 + Patch_Score × 0.20 + Exposure_Score × 0.15)

        Where:
        - Port_Score: min(10, open_ports / 2)
        - Vuln_Score: min(10, known_vulns × 1.5)
        - Patch_Score: min(10, days_since_patch / 30)
        - Exposure_Score: 10 if internet-facing else 3

        Interpretation (0-10 scale):
        - 0-3: Low exposure
        - 3-6: Moderate - schedule hardening
        - 6-8: High - prioritize remediation
        - 8-10: Critical - immediate remediation required
        """
        port_score = min(10, open_ports / 2)
        vuln_score = min(10, known_vulns * 1.5)
        patch_score = min(10, days_since_patch / 30)
        exposure_score = 10 if internet_facing else 3
        asri = port_score * 0.25 + vuln_score * 0.40 + patch_score * 0.20 + exposure_score * 0.15
        return round(min(10, max(0, asri)), 2)


# =========================================================================
# EXPORT FUNCTIONS
# =========================================================================
_PDF_CHAR_MAP = {
    "🟢": "[LOW] ", "🟡": "[MEDIUM] ", "🟠": "[HIGH] ", "🔴": "[CRITICAL] ",
    "✅": "[OK] ", "⚠️": "[WARNING] ", "⚠": "[WARNING] ", "❌": "[X] ",
    "🛡️": "", "🔄": "", "📊": "", "📋": "", "📄": "", "📝": "", "🔧": "", "🔮": "", "⚙️": "", "📈": "", "🌐": "",
    "₹": "Rs. ", "–": "-", "—": "-",
    "’": "'", "‘": "'", "“": '"', "”": '"', "…": "...",
}

def _pdf_safe(value) -> str:
    """Make text safe for FPDF's Latin-1-only core fonts."""
    text = str(value)
    for original, replacement in _PDF_CHAR_MAP.items():
        text = text.replace(original, replacement)
    return text.encode("latin-1", "ignore").decode("latin-1").strip()

def export_to_pdf(df: pd.DataFrame, title: str, formulas: Dict = None) -> bytes:
    if not HAS_PDF:
        return None
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", "B", 16)
        pdf.cell(0, 10, _pdf_safe(title), ln=True, align="C")
        pdf.set_font("Arial", "B", 10)
        pdf.cell(0, 5, f"Generated: {dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", ln=True)
        pdf.cell(0, 5, "CyberSentinel Intelligence Platform | Kalsnet (KNet) Consulting", ln=True)
        pdf.set_font("Arial", "", 8)
        pdf.ln(5)

        col_width = 190 / len(df.columns)
        row_height = 7
        page_bottom = pdf.h - pdf.b_margin

        def draw_header():
            pdf.set_font("Arial", "B", 8)
            for col in df.columns:
                pdf.cell(col_width, row_height, _pdf_safe(col)[:20], border=1, align="C")
            pdf.ln()
            pdf.set_font("Arial", "", 8)

        draw_header()
        for _, row in df.head(50).iterrows():
            if pdf.get_y() + row_height > page_bottom:
                pdf.add_page()
                draw_header()
            for val in row:
                pdf.cell(col_width, row_height, _pdf_safe(val)[:20], border=1, align="L")
            pdf.ln()

        if formulas:
            pdf.add_page()
            pdf.set_font("Arial", "B", 12)
            pdf.cell(0, 10, "Formulas & Explanations", ln=True)
            pdf.set_font("Arial", "", 9)
            for formula_name, formula_text in formulas.items():
                pdf.cell(0, 5, _pdf_safe(formula_name) + ":", ln=True)
                pdf.multi_cell(0, 4, _pdf_safe(formula_text)[:500])
                pdf.ln(2)

        raw_output = pdf.output(dest="S")
        if isinstance(raw_output, str):
            return raw_output.encode("latin-1")
        return bytes(raw_output)
    except Exception as e:
        st.error(f"PDF export error: {e}")
        logger.exception("PDF export failed")
        return None

def export_to_docx(df: pd.DataFrame, title: str, formulas: Dict = None) -> bytes:
    if not HAS_DOCX:
        return None
    try:
        doc = Document()
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 71, 171)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        meta_para = doc.add_paragraph()
        meta_run = meta_para.add_run(f"Generated: {dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        meta_run.font.size = Pt(10)
        meta_run = meta_para.add_run("CyberSentinel Intelligence Platform | Developed by Randy Singh | Kalsnet (KNet) Consulting")
        meta_run.font.size = Pt(10)
        meta_run.italic = True

        doc.add_paragraph()
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr_cells[i].text = str(col)
        for _, row in df.head(50).iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val)[:100]

        if formulas:
            doc.add_page_break()
            formulas_heading = doc.add_paragraph()
            formulas_run = formulas_heading.add_run("Formulas & Explanations")
            formulas_run.font.size = Pt(14)
            formulas_run.font.bold = True
            formulas_run.font.color.rgb = RGBColor(0, 71, 171)
            for formula_name, formula_text in formulas.items():
                doc.add_paragraph(formula_name, style='Heading 3')
                doc.add_paragraph(formula_text[:500])

        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        return docx_bytes.getvalue()
    except Exception as e:
        st.error(f"Word export error: {e}")
        return None

def export_to_text(df: pd.DataFrame, title: str, formulas: Dict = None) -> str:
    text = f"\n{'='*80}\n{title}\n{'='*80}\n\n"
    text += f"Generated: {dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
    text += "CyberSentinel Intelligence Platform | Developed by Randy Singh | Kalsnet (KNet) Consulting\n\n"
    text += df.to_string()
    if formulas:
        text += f"\n\n{'='*80}\nFormulas & Explanations\n{'='*80}\n\n"
        for formula_name, formula_text in formulas.items():
            text += f"\n{formula_name}:\n{formula_text}\n"
    return text

def render_export_row(df: pd.DataFrame, title: str, formulas_dict: Dict, key_prefix: str):
    col1, col2, col3, col4, col5 = st.columns(5)
    with col1:
        csv = df.to_csv(index=False)
        st.download_button("📊 CSV", csv.encode(), f"{key_prefix}.csv", "text/csv", key=f"{key_prefix}_csv")
    with col2:
        json_str = df.to_json(orient="records", indent=2)
        st.download_button("📋 JSON", json_str.encode(), f"{key_prefix}.json", "application/json", key=f"{key_prefix}_json")
    with col3:
        txt = export_to_text(df, title, formulas_dict)
        st.download_button("📄 TXT", txt.encode(), f"{key_prefix}.txt", "text/plain", key=f"{key_prefix}_txt")
    with col4:
        if HAS_DOCX:
            docx_bytes = export_to_docx(df, title, formulas_dict)
            if docx_bytes:
                st.download_button("📝 WORD", docx_bytes, f"{key_prefix}.docx",
                                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key=f"{key_prefix}_docx")
        else:
            st.write("Word export unavailable")
    with col5:
        if HAS_PDF:
            pdf_bytes = export_to_pdf(df, title, formulas_dict)
            if pdf_bytes:
                st.download_button("🔴 PDF", pdf_bytes, f"{key_prefix}.pdf", "application/pdf", key=f"{key_prefix}_pdf")
        else:
            st.write("PDF export unavailable")

# =========================================================================
# SIDEBAR
# =========================================================================
with st.sidebar:
    st.header("⚙️ Configuration")
    with st.expander("📊 Data Settings", expanded=True):
        st.markdown(f"**Max Upload**: {MAX_UPLOAD_MB} MB")
        st.markdown(f"**Environment**: {DEPLOY_ENV.upper()}")
        st.markdown(f"**Session**: {st.session_state.session_id}")
    with st.expander("ℹ️ About", expanded=False):
        st.markdown("""
        CyberSentinel combines rule-based threat scoring, IOC confidence
        rating, and trend-based escalation prediction into a single
        analyst workflow. All scoring logic is transparent and documented
        under each tab's formula box.
        """)
    st.markdown("---")
    st.caption("v1.0 | Kalsnet (KNet) Consulting")

# =========================================================================
# MAIN INTERFACE - TABS
# =========================================================================
tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "🚨 Threat Detection",
    "🌐 Threat Intelligence (IOC)",
    "🔮 Predictive Forecasting",
    "🖥️ Asset & Vulnerability Risk",
    "📈 Dashboard & Analytics",
])

# =========================================================================
# TAB 1: THREAT DETECTION & SCORING
# =========================================================================
with tab1:
    st.markdown('<div class="section-header">🚨 Threat Detection & Risk Scoring</div>', unsafe_allow_html=True)
    st.markdown("""
    Score inbound security events by combining threat severity, exploitability, and the
    business criticality of the targeted asset into a single prioritized risk score.
    """)

    col1, col2, col3, col4 = st.columns([2, 2, 1, 1])
    with col1:
        n_records = st.slider("Synthetic records", 0, 500, 150, step=10)
    with col2:
        domain_choice = st.selectbox("Threat domain", DOMAIN_OPTIONS)
    with col3:
        st.write(""); st.write("")
        generate = st.button("🔄 Generate Demo", key="gen_demo")
    with col4:
        st.write(""); st.write("")
        reset_demo = st.button("🗑️ Reset", key="reset_demo")

    if reset_demo:
        st.session_state.synthetic_df = None
        st.success("✅ Synthetic threat data reset")

    if generate:
        if domain_choice == "Cyber Only":
            pool, asset_pool = CYBER_THREAT_TYPES, ASSET_TYPES_CYBER
        elif domain_choice == "Physical/Operational Only":
            pool, asset_pool = PHYSICAL_THREAT_TYPES, ASSET_TYPES_PHYSICAL
        else:
            pool, asset_pool = THREAT_TYPES, ASSET_TYPES

        rows = []
        for i in range(n_records):
            threat_type = random.choice(list(pool.keys()))
            domain = THREAT_TYPES[threat_type]["domain"]
            exploitability = round(random.uniform(1, 10), 1)
            asset_criticality = round(random.uniform(1, 10), 1)
            score = ThreatFormulas.threat_risk_score(threat_type, exploitability, asset_criticality)

            if score < 30:
                risk_level = "🟢 LOW"
            elif score < 60:
                risk_level = "🟡 MEDIUM"
            elif score < 85:
                risk_level = "🟠 HIGH"
            else:
                risk_level = "🔴 CRITICAL"

            # Physical events use a facility zone reference instead of an IP address
            if domain == "Cyber":
                origin = f"{random.randint(1,223)}.{random.randint(0,255)}.{random.randint(0,255)}.{random.randint(1,254)}"
            else:
                origin = f"ZONE-{random.choice(['A','B','C','D'])}{random.randint(1,12)}"

            rows.append({
                "Event_ID": f"EVT-{random.randint(100000, 999999)}",
                "Timestamp": (dt.datetime.now() - dt.timedelta(hours=random.randint(0, 720))).strftime("%Y-%m-%d %H:%M"),
                "Domain": domain,
                "Source_IP_or_Zone": origin,
                "Threat_Type": threat_type,
                "Category": THREAT_TYPES[threat_type]["category"],
                "Asset_Type": random.choice(asset_pool if asset_pool else ASSET_TYPES),
                "Exploitability": exploitability,
                "Asset_Criticality": asset_criticality,
                "Risk_Score": score,
                "Risk_Level": risk_level,
                "Status": random.choice(STATUS_OPTIONS),
            })
        st.session_state.synthetic_df = pd.DataFrame(rows)
        st.success(f"✅ Generated {n_records} synthetic threat events ({domain_choice})")

    st.markdown("**OR upload real security event data:**")
    uploaded = st.file_uploader("Upload CSV/Excel file", type=["csv", "xlsx"], key="upload_events")
    if uploaded:
        try:
            if uploaded.name.endswith(".csv"):
                st.session_state.real_df = pd.read_csv(uploaded)
            else:
                st.session_state.real_df = pd.read_excel(uploaded)
            st.success(f"✅ Loaded {len(st.session_state.real_df)} records")
        except Exception as e:
            st.error(f"Error: {e}")

    df_show = st.session_state.real_df if st.session_state.real_df is not None else st.session_state.synthetic_df

    if df_show is not None:
        st.markdown("**Data Preview:**")
        st.dataframe(df_show.head(20), use_container_width=True, height=250)

        st.markdown('<div class="subsection-header">📈 Analysis Summary</div>', unsafe_allow_html=True)
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Total Events", len(df_show))
        with col2:
            critical = (df_show.get("Risk_Level", pd.Series(dtype=str)) == "🔴 CRITICAL").sum() if "Risk_Level" in df_show.columns else 0
            st.metric("Critical", critical)
        with col3:
            high = (df_show.get("Risk_Level", pd.Series(dtype=str)) == "🟠 HIGH").sum() if "Risk_Level" in df_show.columns else 0
            st.metric("High", high)
        with col4:
            if "Risk_Score" in df_show.columns:
                st.metric("Avg Risk Score", f"{df_show['Risk_Score'].mean():.1f}")

        if "Domain" in df_show.columns:
            st.markdown('<div class="subsection-header">Cyber vs Physical/Operational Split</div>', unsafe_allow_html=True)
            domain_summary = df_show.groupby("Domain").size().reset_index(name="Count")
            st.bar_chart(domain_summary.set_index("Domain"))

        if "Threat_Type" in df_show.columns:
            st.markdown('<div class="subsection-header">Threat Type Breakdown</div>', unsafe_allow_html=True)
            threat_summary = df_show.groupby("Threat_Type").size().reset_index(name="Count")
            st.bar_chart(threat_summary.set_index("Threat_Type"))

        if "Status" in df_show.columns:
            st.markdown('<div class="subsection-header">Status Breakdown</div>', unsafe_allow_html=True)
            status_summary = df_show.groupby("Status").size().reset_index(name="Count")
            st.bar_chart(status_summary.set_index("Status"))

        st.markdown('<div class="formula-box">', unsafe_allow_html=True)
        st.markdown("""
        **THREAT RISK SCORE (TRS) FORMULA**

        **TRS = (Severity/10 × 0.4 + Exploitability/10 × 0.3 + Asset_Criticality/10 × 0.3) × 100**

        **Severity Weights:**
        """)
        st.markdown("**🌐 Cyber Threats:**")
        for ttype, info in CYBER_THREAT_TYPES.items():
            st.write(f"- **{ttype}** ({info['severity']}/10, {info['category']}): {info['description']}")
        st.markdown("**🏭 Physical/Operational Threats:**")
        for ttype, info in PHYSICAL_THREAT_TYPES.items():
            st.write(f"- **{ttype}** ({info['severity']}/10, {info['category']}): {info['description']}")
        st.markdown("""
        **Interpretation:**
        - Score < 30: 🟢 LOW - Log and monitor
        - Score 30-60: 🟡 MEDIUM - Investigate within SLA
        - Score 60-85: 🟠 HIGH - Escalate to analyst immediately
        - Score > 85: 🔴 CRITICAL - Incident response activation
        """)
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown('<div class="subsection-header">💾 Export Results</div>', unsafe_allow_html=True)
        formulas_dict = {
            "Threat Risk Score": "TRS = (Severity/10 x 0.4 + Exploitability/10 x 0.3 + Asset_Criticality/10 x 0.3) x 100",
            "Threat Severity Reference": "\n".join([f"{k}: {v['severity']}/10 - {v['description']}" for k, v in THREAT_TYPES.items()]),
        }
        render_export_row(df_show, "Threat Detection Analysis", formulas_dict, "threat_detect")

# =========================================================================
# TAB 2: THREAT INTELLIGENCE (IOC)
# =========================================================================
with tab2:
    st.markdown('<div class="section-header">🌐 Threat Intelligence & IOC Confidence</div>', unsafe_allow_html=True)
    st.markdown("""
    Rate incoming Indicators of Compromise (IOCs) from intelligence feeds by combining
    source reliability, cross-source corroboration, and recency into a single confidence score.
    """)

    col1, col2, col3, col4 = st.columns([2, 2, 1, 1])
    with col1:
        n_iocs = st.slider("Synthetic IOC records", 0, 500, 120, step=10)
    with col2:
        ioc_domain_choice = st.selectbox("Intel domain", DOMAIN_OPTIONS, key="ioc_domain")
    with col3:
        st.write(""); st.write("")
        gen_ioc = st.button("🔄 Generate", key="gen_ioc")
    with col4:
        st.write(""); st.write("")
        reset_ioc = st.button("🗑️ Reset", key="reset_ioc")

    if reset_ioc:
        st.session_state.ioc_df = None
        st.success("✅ IOC data reset")

    if gen_ioc:
        if ioc_domain_choice == "Cyber Only":
            type_pool, feed_pool, dom = IOC_TYPES_CYBER, SOURCE_FEEDS_CYBER, "Cyber"
        elif ioc_domain_choice == "Physical/Operational Only":
            type_pool, feed_pool, dom = IOC_TYPES_PHYSICAL, SOURCE_FEEDS_PHYSICAL, "Physical/Operational"
        else:
            type_pool, feed_pool, dom = IOC_TYPES, SOURCE_FEEDS, None

        rows = []
        for i in range(n_iocs):
            ioc_type = random.choice(type_pool)
            source_feed = random.choice(feed_pool)
            domain = dom if dom else ("Cyber" if ioc_type in IOC_TYPES_CYBER else "Physical/Operational")
            source_reliability = round(random.uniform(0.3, 1.0), 2)
            corroboration = random.randint(0, 8)
            recency = random.randint(0, 120)
            confidence = ThreatFormulas.ioc_confidence_score(source_reliability, corroboration, recency)

            if confidence < 40:
                action = "🟢 Informational"
            elif confidence < 70:
                action = "🟡 Monitor"
            else:
                action = "🔴 Actionable - Block"

            rows.append({
                "IOC_ID": f"IOC-{random.randint(10000, 99999)}",
                "Domain": domain,
                "IOC_Type": ioc_type,
                "Source_Feed": source_feed,
                "Source_Reliability": source_reliability,
                "Corroborating_Sources": corroboration,
                "Age_Days": recency,
                "Confidence_%": confidence,
                "Recommended_Action": action,
            })
        st.session_state.ioc_df = pd.DataFrame(rows)
        st.success(f"✅ Generated {n_iocs} synthetic IOC records ({ioc_domain_choice})")

    if st.session_state.ioc_df is not None:
        ioc_df = st.session_state.ioc_df
        st.dataframe(ioc_df, use_container_width=True, height=250)

        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Total IOCs", len(ioc_df))
        with col2:
            st.metric("Avg Confidence", f"{ioc_df['Confidence_%'].mean():.1f}%")
        with col3:
            actionable = (ioc_df["Recommended_Action"] == "🔴 Actionable - Block").sum()
            st.metric("Actionable", actionable)
        with col4:
            informational = (ioc_df["Recommended_Action"] == "🟢 Informational").sum()
            st.metric("Informational", informational)

        st.markdown('<div class="subsection-header">Source Feed Breakdown</div>', unsafe_allow_html=True)
        feed_summary = ioc_df.groupby("Source_Feed").size().reset_index(name="Count")
        st.bar_chart(feed_summary.set_index("Source_Feed"))

        st.markdown('<div class="formula-box">', unsafe_allow_html=True)
        st.markdown("""
        **IOC CONFIDENCE SCORE FORMULA**

        **Confidence = (Source_Reliability × 0.40 + Corroboration_Factor × 0.35 + Recency_Factor × 0.25) × 100**

        - Corroboration_Factor = min(1, corroborating_sources / 5)
        - Recency_Factor = max(0, 1 - (age_days / 90))

        **Interpretation:**
        - < 40%: 🟢 Informational only
        - 40-70%: 🟡 Monitor, do not auto-block
        - > 70%: 🔴 Actionable, safe to auto-block
        """)
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown('<div class="subsection-header">💾 Export Results</div>', unsafe_allow_html=True)
        formulas_dict = {
            "IOC Confidence Score": "Confidence = (Source_Reliability x 0.40 + Corroboration_Factor x 0.35 + Recency_Factor x 0.25) x 100",
        }
        render_export_row(ioc_df, "Threat Intelligence IOC Analysis", formulas_dict, "ioc_intel")

# =========================================================================
# TAB 3: PREDICTIVE FORECASTING
# =========================================================================
with tab3:
    st.markdown('<div class="section-header">🔮 Predictive Threat Escalation Forecasting</div>', unsafe_allow_html=True)
    st.markdown("""
    Estimate the probability that a currently-observed cluster of threat activity escalates
    into a confirmed incident, based on event momentum rather than a single point-in-time score.
    """)

    col1, col2, col3 = st.columns(3)
    with col1:
        freq_trend = st.slider("Event Frequency Trend (0=flat, 1=rapidly rising)", 0.0, 1.0, 0.4, step=0.05)
    with col2:
        avg_conf = st.slider("Average IOC Confidence (%)", 0, 100, 55, step=5)
    with col3:
        avg_sev = st.slider("Average Threat Severity (1-10)", 1.0, 10.0, 6.0, step=0.5)

    escalation = ThreatFormulas.escalation_probability(freq_trend, avg_conf, avg_sev)

    if escalation < 25:
        status, action, color = "🟢 UNLIKELY", "Continue routine monitoring", "#1a8a3e"
    elif escalation < 50:
        status, action, color = "🟡 POSSIBLE", "Keep watching, review in next shift", "#ff9800"
    elif escalation < 75:
        status, action, color = "🟠 LIKELY", "Proactively prepare incident response", "#ffa500"
    else:
        status, action, color = "🔴 HIGHLY LIKELY", "Escalate to incident response now", "#c0392b"

    st.markdown(f'<div style="background:{color}; color:white; padding:1.5rem; border-radius:8px; text-align:center;">'
                f'<h2>Escalation Probability: {escalation:.1f}%</h2>'
                f'<h3>{status}</h3><p>{action}</p></div>', unsafe_allow_html=True)

    st.markdown('<div class="formula-box">', unsafe_allow_html=True)
    st.markdown(f"""
    **PREDICTIVE THREAT ESCALATION FORMULA**

    **Escalation% = (Frequency_Trend × 0.5 + Confidence_Norm × 0.3 + Severity_Norm × 0.2) × 100**

    **Your Values:**
    - Frequency Trend: {freq_trend:.2f} → weighted contribution {freq_trend*0.5*100:.1f}
    - Confidence: {avg_conf}% → weighted contribution {(avg_conf/100)*0.3*100:.1f}
    - Severity: {avg_sev:.1f}/10 → weighted contribution {(avg_sev/10)*0.2*100:.1f}

    **Predicted Escalation Probability: {escalation:.1f}%**

    **Interpretation:**
    - < 25%: Unlikely to escalate
    - 25-50%: Possible - keep watching
    - 50-75%: Likely - proactively prepare response
    - > 75%: Highly likely - escalate now
    """)
    st.markdown('</div>', unsafe_allow_html=True)

    st.markdown('<div class="subsection-header">📈 30-Day Simulated Threat Volume Trend</div>', unsafe_allow_html=True)
    days = pd.date_range(end=dt.datetime.now(), periods=30)
    base = 20 + freq_trend * 40
    trend_vals = [max(0, base + i * freq_trend * 2 + random.uniform(-5, 5)) for i in range(30)]
    trend_df = pd.DataFrame({"Date": days, "Simulated_Event_Volume": trend_vals}).set_index("Date")
    st.line_chart(trend_df)

# =========================================================================
# TAB 4: ASSET & VULNERABILITY RISK
# =========================================================================
with tab4:
    st.markdown('<div class="section-header">🖥️ Attack Surface & Asset Vulnerability Risk</div>', unsafe_allow_html=True)
    st.markdown("""
    Estimate how exposed a given asset is based on open ports, known unpatched vulnerabilities,
    patch age, and internet exposure. Applies equally to digital assets (servers, endpoints)
    and physical assets (facilities, control systems) — for a physical asset, treat "open ports"
    as unmonitored entry points and "internet-facing" as public-facing/street-accessible.
    """)

    col1, col2, col3, col4 = st.columns(4)
    with col1:
        open_ports = st.slider("Open Ports", 0, 30, 6, step=1)
    with col2:
        known_vulns = st.slider("Known Unpatched Vulnerabilities", 0, 15, 3, step=1)
    with col3:
        patch_age = st.slider("Days Since Last Patch", 0, 365, 45, step=5)
    with col4:
        internet_facing = st.toggle("Internet-Facing Asset", value=True)

    asri = ThreatFormulas.attack_surface_risk_index(open_ports, known_vulns, patch_age, internet_facing)

    if asri < 3:
        status, color = "🟢 LOW EXPOSURE", "#1a8a3e"
    elif asri < 6:
        status, color = "🟡 MODERATE - SCHEDULE HARDENING", "#ff9800"
    elif asri < 8:
        status, color = "🟠 HIGH - PRIORITIZE REMEDIATION", "#ffa500"
    else:
        status, color = "🔴 CRITICAL - IMMEDIATE REMEDIATION", "#c0392b"

    st.markdown(f'<div style="background:{color}; color:white; padding:2rem; border-radius:8px; text-align:center;">'
                f'<h1>ASRI: {asri:.2f}/10</h1><h2>{status}</h2></div>', unsafe_allow_html=True)

    st.markdown('<div class="formula-box">', unsafe_allow_html=True)
    st.markdown(f"""
    **ATTACK SURFACE RISK INDEX (ASRI) FORMULA**

    **ASRI = Port_Score×0.25 + Vuln_Score×0.40 + Patch_Score×0.20 + Exposure_Score×0.15**

    - Port_Score = min(10, open_ports / 2) → {min(10, open_ports/2):.1f}
    - Vuln_Score = min(10, known_vulns × 1.5) → {min(10, known_vulns*1.5):.1f}
    - Patch_Score = min(10, days_since_patch / 30) → {min(10, patch_age/30):.1f}
    - Exposure_Score = 10 if internet-facing else 3 → {10 if internet_facing else 3}

    **Predicted ASRI: {asri:.2f}/10**

    **Reference:**
    - 0-3: Low exposure
    - 3-6: Moderate - schedule hardening
    - 6-8: High - prioritize remediation
    - 8-10: Critical - immediate remediation required
    """)
    st.markdown('</div>', unsafe_allow_html=True)

# =========================================================================
# TAB 5: DASHBOARD & ANALYTICS
# =========================================================================
with tab5:
    st.markdown('<div class="section-header">📈 Dashboard & Analytics Overview</div>', unsafe_allow_html=True)

    df_dash = st.session_state.real_df if st.session_state.real_df is not None else st.session_state.synthetic_df

    if df_dash is None:
        st.info("Generate or upload threat detection data in the 'Threat Detection' tab to populate this dashboard.")
    else:
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Total Events", len(df_dash))
        with col2:
            if "Risk_Score" in df_dash.columns:
                st.metric("Avg Risk Score", f"{df_dash['Risk_Score'].mean():.1f}")
        with col3:
            if "Status" in df_dash.columns:
                resolved = (df_dash["Status"] == "Resolved").sum()
                st.metric("Resolved", resolved)
        with col4:
            if "Status" in df_dash.columns:
                open_events = (df_dash["Status"].isin(["Detected", "Investigating"])).sum()
                st.metric("Open", open_events)

        if "Domain" in df_dash.columns:
            st.markdown('<div class="subsection-header">Cyber vs Physical/Operational Split</div>', unsafe_allow_html=True)
            domain_summary = df_dash.groupby("Domain").size().reset_index(name="Count")
            st.bar_chart(domain_summary.set_index("Domain"))

        if "Category" in df_dash.columns:
            st.markdown('<div class="subsection-header">Threats by Category</div>', unsafe_allow_html=True)
            cat_summary = df_dash.groupby("Category").size().reset_index(name="Count")
            st.bar_chart(cat_summary.set_index("Category"))

        if "Asset_Type" in df_dash.columns:
            st.markdown('<div class="subsection-header">Threats by Targeted Asset Type</div>', unsafe_allow_html=True)
            asset_summary = df_dash.groupby("Asset_Type").size().reset_index(name="Count")
            st.bar_chart(asset_summary.set_index("Asset_Type"))

        if "Risk_Score" in df_dash.columns and "Timestamp" in df_dash.columns:
            st.markdown('<div class="subsection-header">Risk Score Over Time</div>', unsafe_allow_html=True)
            try:
                temp = df_dash.copy()
                temp["Timestamp"] = pd.to_datetime(temp["Timestamp"], errors="coerce")
                temp = temp.dropna(subset=["Timestamp"]).sort_values("Timestamp")
                st.line_chart(temp.set_index("Timestamp")["Risk_Score"])
            except Exception:
                st.caption("Timestamp column could not be parsed for the time-series chart.")

st.markdown("---")
st.caption(f"🛡️ CyberSentinel Intelligence Platform v1.0 | Session {st.session_state.session_id} | {DEPLOY_ENV.upper()} | Developed by Randy Singh | Kalsnet (KNet) Consulting")

