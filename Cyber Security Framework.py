




#=============================================================================
  #CYBERSECURITY FRAMEWORK APPLICATION FOR DoD CRITICAL INFRASTRUCTURE
  #@Developed by Randy Singh | KalSnet (KNet) Consulting Group
  #Original Java code: September 2017 — Python/Streamlit conversion: 2024
#=============================================================================
#Run with:
   # pip install streamlit plotly pandas fpdf2 python-docx graphviz
    #streamlit run cybersecurity_framework.py
#=============================================================================


import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import random
import json
import io
from datetime import datetime, timedelta

# ─── Optional export dependencies ─────────────────────────────────────────────
try:
    from fpdf import FPDF
    FPDF_AVAILABLE = True
except ImportError:
    FPDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    from docx.shared import RGBColor, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ══════════════════════════════════════════════════════════════════════════════
#  PAGE CONFIG
# ══════════════════════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="CyberSecurity Framework | KNet Consulting",
    page_icon="",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ══════════════════════════════════════════════════════════════════════════════
#  GLOBAL STYLES
# ══════════════════════════════════════════════════════════════════════════════
st.markdown("""
<style>
  .cs-title {
      text-align:center; font-size:2.3rem; font-weight:900;
      color:#0047AB; letter-spacing:1px; margin-bottom:0;
  }
  .cs-subtitle {
      text-align:center; font-size:1.05rem; font-weight:700;
      color:#0047AB; margin-top:2px; margin-bottom:2px;
  }
  .cs-tagline {
      text-align:center; color:#555; font-size:0.85rem; margin-bottom:14px;
  }
  .cs-hr { border:2px solid #0047AB; margin:6px 0 16px 0; }
  .section-header {
      background:linear-gradient(90deg,#0047AB 0%,#1a6fe8 100%);
      color:white; padding:8px 16px; border-radius:6px;
      font-weight:700; font-size:1rem; margin:10px 0 8px 0;
  }
  .req-card {
      background:#f0f6ff; border-left:4px solid #0047AB;
      border-radius:4px; padding:9px 13px; margin:5px 0; font-size:0.91rem;
  }
  .rec-card {
      background:#fff8e1; border-left:4px solid #f0a500;
      border-radius:4px; padding:9px 13px; margin:5px 0; font-size:0.91rem;
  }
  .sol-card {
      background:#e8f5e9; border-left:4px solid #2e7d32;
      border-radius:4px; padding:9px 13px; margin:5px 0; font-size:0.91rem;
  }
  .ref-card {
      background:#fce4ec; border-left:4px solid #c62828;
      border-radius:4px; padding:9px 13px; margin:5px 0; font-size:0.85rem; font-family:monospace;
  }
  div[data-testid="metric-container"] {
      background:#f0f6ff; border:1px solid #c2d9ff; border-radius:8px; padding:10px;
  }
  section[data-testid="stSidebar"] { background:#001a6b; }
  section[data-testid="stSidebar"] * { color:#e8f0ff !important; }
</style>
""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
#  HEADER
# ══════════════════════════════════════════════════════════════════════════════
st.markdown('<p class="cs-title">CYBERSECURITY FRAMEWORK APPLICATION</p>', unsafe_allow_html=True)
st.markdown('<p class="cs-subtitle">FOR US CRITICAL INFRASTRUCTURE</p>', unsafe_allow_html=True)
st.markdown('<p class="cs-subtitle">Developed by Randy Singh | KalSnet (KNet) Consulting Group</p>', unsafe_allow_html=True)
st.markdown('<p class="cs-tagline">Helping Security Analysts, Planners & Administrators address IoT Security Vulnerabilities in DISA, DoD & USA Critical Infrastructure &nbsp;|&nbsp; Originally authored June 2026</p>', unsafe_allow_html=True)
st.markdown('<hr class="cs-hr">', unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
#  FRAMEWORK DATA  (21 checklist items)
# ══════════════════════════════════════════════════════════════════════════════
FRAMEWORK_ITEMS = {
    1: {
        "title": "Identify - Asset Management",
        "function": "IDENTIFY",
        "summary": "Data, personnel, devices, systems, and facilities that enable the organization to achieve business purposes are identified and managed consistent with their relative importance to business objectives and the organization's risk strategy.",
        "subcategories": [
            "Physical devices and systems within the organization are inventoried.",
            "Software platforms and applications within the organization are inventoried.",
            "Organizational communication and data flows are mapped.",
            "External information systems are catalogued.",
            "Resources (hardware, devices, data, software) are prioritized based on classification, criticality, and business value.",
            "Cybersecurity roles and responsibilities for the entire workforce and third-party stakeholders are established.",
        ],
        "example": "Use a CMDB (Configuration Management Database) like ServiceNow to auto-discover and track all hardware/software assets. Tag assets by classification (CUI, FOUO, UNCLASS).",
        "solution": "Deploy an automated asset discovery tool (Tenable.io, Axonius, or Rumble) integrated with your SIEM. Establish a quarterly asset reconciliation process.",
        "recommendation": "Implement NIST SP 800-171 control 3.4.1. Require all new devices to be registered in the CMDB before network access is granted. Review asset inventory monthly.",
        "references": [
            "CCS CSC 1, COBIT 5 BAI09.01/02, ISA 62443-2-1:2009 4.2.3.4, NIST SP 800-53 Rev.4 CM-8",
            "CCS CSC 2, COBIT 5 BAI09.01/02/05, ISO/IEC 27001:2013 A.8.1.1/2, NIST SP 800-53 CM-8",
            "CCS CSC 1, COBIT 5 DSS05.02, ISO/IEC 27001:2013 A.13.2.1, NIST SP 800-53 AC-4/CA-3",
            "COBIT 5 APO02.02, ISO/IEC 27001:2013 A.11.2.6, NIST SP 800-53 AC-20/SA-9",
            "COBIT 5 APO03.03/04/BAI09.02, ISA 62443-2-1:2009 4.2.3.6, NIST SP 800-53 CP-2/RA-2",
            "COBIT 5 APO01.02/DSS06.03, ISA 62443-2-1:2009 4.3.2.3.3, NIST SP 800-53 CP-2/PS-7",
        ],
    },
    2: {
        "title": "Identify - Business Environment",
        "function": "IDENTIFY",
        "summary": "The organization's mission, objectives, stakeholders, and activities are understood and prioritized. The organization's role in the supply chain is identified and communicated. This information informs cybersecurity roles, responsibilities, and risk management decisions.",
        "subcategories": [
            "The organization's role in the supply chain is identified and communicated.",
            "The organization's place in critical infrastructure and its industry sector is identified and communicated.",
            "Priorities for organizational mission, objectives, and activities are established and communicated.",
            "Dependencies and critical functions for delivery of critical services are established.",
            "Resilience requirements to support delivery of critical services are established.",
        ],
        "example": "Map your agency's position in the DoD supply chain using a BIA (Business Impact Analysis). Identify which systems are mission-critical (Tier 1) vs. supportive (Tier 3).",
        "solution": "Conduct an annual Business Environment Review using CISA's Critical Infrastructure Framework. Maintain a stakeholder register updated quarterly.",
        "recommendation": "Align mission priorities to CISA's 16 Critical Infrastructure Sectors. Document all third-party dependencies. Test resilience via tabletop exercises biannually.",
        "references": [
            "COBIT 5 APO08.04/05/APO10.03-05, ISO/IEC 27001:2013 A.15.1.3/A.15.2.1-2, NIST SP 800-53 CP-2/SA-12",
            "COBIT 5 APO02.06/APO03.01, NIST SP 800-53 PM-8",
            "COBIT 5 APO02.01/06/APO03.01, ISA 62443-2-1:2009 4.2.2.1/4.2.3.6, NIST SP 800-53 PM-11/SA-14",
            "ISO/IEC 27001:2013 A.11.2.2/3/A.12.1.3, NIST SP 800-53 CP-8/PE-9/PE-11/PM-8/SA-14",
            "COBIT 5 DSS04.02, ISO/IEC 27001:2013 A.11.1.4/A.17.1.1-2/A.17.2.1, NIST SP 800-53 CP-2/CP-11/SA-14",
        ],
    },
    3: {
        "title": "Identify - Governance Policies & Procedures",
        "function": "IDENTIFY",
        "summary": "The policies, procedures, and processes to manage and monitor the organization's regulatory, legal, risk, environmental, and operational requirements are understood and inform the management of cybersecurity risk.",
        "subcategories": [
            "Organizational information security policy is established.",
            "Information security roles and responsibilities are coordinated and aligned with internal roles and external partners.",
            "Legal and regulatory requirements regarding cybersecurity, including privacy and civil liberties obligations, are understood and managed.",
            "Governance and risk management processes address cybersecurity risks.",
        ],
        "example": "Publish a DoD-aligned Information Security Policy covering AUP, data classification, incident reporting, and third-party access. Review annually and after major incidents.",
        "solution": "Establish a Cybersecurity Governance Board (CGB) with CISO, Legal, Operations, and HR representatives. Use GRC tools (RSA Archer, ServiceNow GRC) to track policy compliance.",
        "recommendation": "Align policies to FISMA, DFARS 252.204-7012, and CMMC Level 2+. Conduct policy gap assessments annually. Ensure all policies have version control and sign-off records.",
        "references": [
            "ISA 62443-2-1:2009 4.3.2.6, ISO/IEC 27001:2013 A.5.1.1, NIST SP 800-53 -1 controls (all families)",
            "COBIT 5 APO13.12, ISA 62443-2-1:2009 4.3.2.3.3, ISO/IEC 27001:2013 A.6.1.1/A.7.2.1, NIST SP 800-53 PM-1/PS-7",
            "ISA 62443-2-1:2009 4.4.3.7, ISO/IEC 27001:2013 A.18.1, NIST SP 800-53 -1 controls (except PM-1)",
            "COBIT 5 DSS04.02, ISA 62443-2-1:2009 4.2.3.1/3/8/9/11/4.3.2.4.3/6.3, NIST SP 800-53 PM-9/PM-11",
        ],
    },
    4: {
        "title": "Identify - Risk Assessment",
        "function": "IDENTIFY",
        "summary": "The organization understands cybersecurity risk to organizational operations (mission, functions, image, reputation), organizational assets, and individuals.",
        "subcategories": [
            "Asset vulnerabilities are identified and documented.",
            "Threat and vulnerability information is received from information sharing forums and sources.",
            "Threats, both internal and external, are identified and documented.",
            "Potential business impacts and likelihoods are identified.",
            "Threats, vulnerabilities, likelihoods, and impacts are used to determine risk.",
            "Risk responses are identified and prioritized.",
        ],
        "example": "Run quarterly vulnerability scans using Tenable Nessus or Qualys. Participate in DHS AIS (Automated Indicator Sharing) for threat intel feeds. Produce risk-ranked remediation plans.",
        "solution": "Implement a Risk Register using NIST SP 800-30 methodology. Integrate threat intel from ISACs (Information Sharing and Analysis Centers) relevant to DoD sector.",
        "recommendation": "Perform annual Risk Assessments per OMB A-130. Use CVSS 3.x scoring. Prioritize Critical/High vulnerabilities for 30-day remediation. Document all accepted risks with ISSO sign-off.",
        "references": [
            "CCS CSC 4, COBIT 5 APO12.01-04, ISA 62443-2-1:2009 4.2.3/4.2.3.7/9/12, NIST SP 800-53 CA-2/CA-7/RA-3/RA-5/SI-2/SI-4",
            "ISA 62443-2-1:2009 4.2.3/9/12, ISO/IEC 27001:2013 A.6.1.4, NIST SP 800-53 PM-15/PM-16/SI-5",
            "COBIT 5 APO12.01-04, ISA 62443-2-1:2009 4.2.3/9/12, NIST SP 800-53 RA-3/SI-5/PM-12/PM-16",
            "COBIT 5 DSS04.02, NIST SP 800-53 RA-2/RA-3/PM-9/PM-11/SA-14",
            "COBIT 5 APO12.02, ISO/IEC 27001:2013 A.12.6.1, NIST SP 800-53 RA-2/RA-3/PM-16",
            "COBIT 5 APO12.05/APO13.02, NIST SP 800-53 PM-4/PM-9",
        ],
    },
    5: {
        "title": "Identify - Risk Management Strategy",
        "function": "IDENTIFY",
        "summary": "The organization's priorities, constraints, risk tolerances, and assumptions are established and used to support operational risk decisions.",
        "subcategories": [
            "Risk management processes are established, managed, and agreed to by organizational stakeholders.",
            "Organizational risk tolerance is determined and clearly expressed.",
            "The organization's determination of risk tolerance is informed by its role in critical infrastructure and sector-specific risk analysis.",
            "Resilience requirements to support delivery of critical services are established.",
        ],
        "example": "Define a Risk Tolerance Statement (e.g., 'No Critical CVSS >= 9.0 vulnerabilities unpatched for > 15 days'). Tie it to ATO (Authority to Operate) renewal criteria.",
        "solution": "Develop an Enterprise Risk Management (ERM) framework aligned to OMB Circular A-11. Use a risk heat map to visualize and communicate tolerance levels to leadership.",
        "recommendation": "Review risk tolerance annually with senior leadership. Publish risk appetite statements per NIST SP 800-39. Ensure mission-critical systems have a documented BCP (Business Continuity Plan).",
        "references": [
            "COBIT 5 APO12.04/05/APO13.02/BAI02.03/BAI04.02, ISA 62443-2-1:2009 4.3.4.2, NIST SP 800-53 PM-9",
            "COBIT 5 APO12.06, ISA 62443-2-1:2009 4.3.2.6.5, NIST SP 800-53 PM-9",
            "NIST SP 800-53 PM-8/PM-9/PM-11/SA-14",
        ],
    },
    6: {
        "title": "Protect - Access Control",
        "function": "PROTECT",
        "summary": "Access to assets and associated facilities is limited to authorized users, processes, or devices, and to authorized activities and transactions.",
        "subcategories": [
            "Identities and credentials are managed for authorized devices and users.",
            "Physical access to assets is managed and protected.",
            "Remote access is managed.",
            "Access permissions are managed, incorporating least privilege and separation of duties.",
            "Network integrity is protected, incorporating network segregation where appropriate.",
        ],
        "example": "Implement CAC/PIV card authentication for all DoD network access. Deploy CyberArk PAM for privileged accounts. Use Cisco ISE for network access control with 802.1X.",
        "solution": "Deploy Zero Trust Network Access (ZTNA) using SASE architecture. Integrate MFA for all remote access via Microsoft Entra ID or Okta. Enforce RBAC via Active Directory groups.",
        "recommendation": "Enforce least-privilege principle. Review access rights quarterly. Disable inactive accounts after 30 days. Segment networks into security zones. Log all privileged access to SIEM.",
        "references": [
            "CCS CSC 16, COBIT 5 DSS05.04/DSS06.03, ISA 62443-3-3:2013 SR 1.1-1.9, ISO/IEC 27001:2013 A.9.2.1-4, NIST SP 800-53 AC-2/IA Family",
            "COBIT 5 DSS01.04/DSS05.05, ISA 62443-2-1:2009 4.3.3.3.2/8, ISO/IEC 27001:2013 A.11.1.1-6, NIST SP 800-53 PE-2/PE-9",
            "COBIT 5 APO13.01/DSS01.04/DSS05.03, ISA 62443-3-3:2013 SR 1.13/2.6, NIST SP 800-53 AC-17/19/20",
            "CCS CSC 12/15, ISA 62443-2-1:2009 4.3.3.7.3, NIST SP 800-53 AC-2/3/5/6/16",
            "ISA 62443-2-1:2009 4.3.3.4, ISA 62443-3-3:2013 SR 3.1/3.8, NIST SP 800-53 AC-4/SC-7",
        ],
    },
    7: {
        "title": "Protect - Awareness & Training",
        "function": "PROTECT",
        "summary": "The organization's personnel and partners are provided cybersecurity awareness education and are adequately trained to perform their information security-related duties consistent with related policies, procedures, and agreements.",
        "subcategories": [
            "All users are informed and trained.",
            "Privileged users understand roles and responsibilities.",
            "Third-party stakeholders understand roles and responsibilities.",
            "Senior executives understand roles and responsibilities.",
            "Physical and information security personnel understand roles and responsibilities.",
        ],
        "example": "Mandate annual DoD Cyber Awareness Challenge (CAC) training. Conduct phishing simulations monthly via KnowBe4. Provide role-based training for system admins (e.g., SANS GIAC).",
        "solution": "Implement a Security Awareness Training platform (KnowBe4, Proofpoint Security Awareness). Track completion rates in LMS. Report non-compliance to supervisors monthly.",
        "recommendation": "Require 100% annual training completion before account renewal. Include social engineering, insider threat, and data handling modules. Brief senior leadership quarterly on threat landscape.",
        "references": [
            "CCS CSC 9, COBIT 5 APO07.03/BAI05.07, ISA 62443-2-1:2009 4.3.2.4.2, ISO/IEC 27001:2013 A.7.2.2",
            "CCS CSC 9, COBIT 5 APO07.02/DSS06.03, ISA 62443-2-1:2009 4.3.2.4.2-3, NIST SP 800-53 AT-3/PM-13",
            "CCS CSC 9, COBIT 5 APO07.03/APO10.04-05, NIST SP 800-53 PS-7/SA-9",
            "CCS CSC 9, COBIT 5 APO07.03, ISO/IEC 27001:2013 A.6.1.1/A.7.2.2, NIST SP 800-53 AT-3/PM-13",
            "CCS CSC 9, COBIT 5 APO07.03, NIST SP 800-53 AT-3/PM-13",
        ],
    },
    8: {
        "title": "Protect - Data Security",
        "function": "PROTECT",
        "summary": "Information and records (data) are managed consistent with the organization's risk strategy to protect the confidentiality, integrity, and availability of information.",
        "subcategories": [
            "Data-at-rest is protected.",
            "Data-in-transit is protected.",
            "Assets are formally managed throughout removal, transfers, and disposition.",
            "Adequate capacity to ensure availability is maintained.",
            "Protections against data leaks are implemented.",
            "Integrity checking mechanisms are used to verify software, firmware, and information integrity.",
            "Development and testing environments are separate from the production environment.",
        ],
        "example": "Encrypt all data-at-rest with AES-256 using DoD-approved solutions (e.g., BitLocker, NetApp Storage Encryption). Use TLS 1.3 for all data-in-transit. Implement DLP via Microsoft Purview.",
        "solution": "Deploy a Data Classification engine (Microsoft Information Protection or Trellix DLP). Enforce encryption policies via Group Policy. Establish a Media Sanitization Program per NIST SP 800-88.",
        "recommendation": "Classify all data per DoD 5200.01. Encrypt CUI at rest and in transit. Implement immutable backup solutions. Separate dev/test/prod environments with distinct access controls. Conduct DLP audits quarterly.",
        "references": [
            "CCS CSC 17, COBIT 5 APO01.06/BAI02.01/BAI06.01/DSS06.06, ISA 62443-3-3:2013 SR 3.4/4.1, NIST SP 800-53 SC-28",
            "CCS CSC 17, COBIT 5 APO01.06/DSS06.06, ISA 62443-3-3:2013 SR 3.1/3.8/4.1/4.2, NIST SP 800-53 SC-8",
            "COBIT 5 BAI09.03, ISA 62443-2-1:2009 4.3.3.3.9/4.3.4.4.1, ISO/IEC 27001:2013 A.8.2.3/A.8.3.1-3, NIST SP 800-53 CM-8/MP-6/PE-16",
            "COBIT 5 APO13.01, ISA 62443-3-3:2013 SR 7.1/7.2, ISO/IEC 27001:2013 A.12.3.1, NIST SP 800-53 AU-4/CP-2/SC-5",
            "CCS CSC 17, COBIT 5 APO01.06, ISA 62443-3-3:2013 SR 5.2, NIST SP 800-53 AC-4/5/6/SC-7/8/13/31/SI-4",
            "ISA 62443-3-3:2013 SR 3.1/3.3/3.4/3.8, ISO/IEC 27001:2013 A.12.2.1/12.5.1/14.1.2-3, NIST SP 800-53 SI-7",
            "COBIT 5 BAI07.04, ISO/IEC 27001:2013 A.12.1.4, NIST SP 800-53 CM-2",
        ],
    },
    9: {
        "title": "Protect - Information Protection Processes & Procedures",
        "function": "PROTECT",
        "summary": "Security policies (that address purpose, scope, roles, responsibilities, management commitment, and coordination among organizational entities), processes and procedures are maintained and used to manage protection of information systems and assets.",
        "subcategories": [
            "A baseline configuration of IT/ICS is created and maintained.",
            "A System Development Life Cycle (SDLC) to manage systems is implemented.",
            "Configuration change control processes are in place.",
            "Backups of information are conducted, maintained, and tested periodically.",
            "Policy and regulations regarding the physical operating environment for assets are met.",
            "Data is destroyed according to policy.",
            "Protection processes are continuously improved.",
            "Effectiveness of protection technologies is shared with appropriate parties.",
            "Response and recovery plans (IR, BC, DR) are in place and managed.",
            "Response and recovery plans are tested.",
            "Cybersecurity is included in human resources practices.",
            "A vulnerability management plan is developed and implemented.",
        ],
        "example": "Maintain STIG-hardened baselines for all OS/applications. Use Ansible for configuration management. Test backup restoration quarterly. Run IR tabletop exercises biannually.",
        "solution": "Implement a Configuration Management tool (Ansible, Puppet, Chef). Use Veeam or Commvault for backup with 3-2-1 strategy. Maintain an updated IR Plan per NIST SP 800-61.",
        "recommendation": "Apply all STIGs before system deployment. Enforce change management via ITSM (ServiceNow). Test backups monthly. Conduct full DR exercises annually. Include cyber clauses in all HR onboarding.",
        "references": [
            "CCS CSC 3/10, COBIT 5 BAI10.01-03/05, ISA 62443-2-1:2009 4.3.4.3.2-3, NIST SP 800-53 CM-2/3/4/5/6/7/9/SA-10",
            "COBIT 5 APO13.01, ISA 62443-2-1:2009 4.3.4.3.3, NIST SP 800-53 SA-3/4/8/10/11/12/15/17/PL-8",
            "COBIT 5 BAI06.01, ISA 62443-2-1:2009 4.3.4.3.2-3, NIST SP 800-53 CM-3/4/SA-10",
            "COBIT 5 APO13.0, ISA 62443-2-1:2009 4.3.4.3.9, NIST SP 800-53 CP-4/CP-6/CP-9",
            "COBIT 5 DSS01.04/DSS05.05, NIST SP 800-53 PE-10/PE-12/PE-13/PE-14/PE-15/PE-18",
            "COBIT 5 BAI09.03, ISA 62443-2-1:2009 4.3.4.4.4, NIST SP 800-53 MP-6",
            "COBIT 5 APO11.06/DSS04.05, NIST SP 800-53 CA-2/CA-7/CP-2/IR-8/PL-2/PM-6",
            "ISO/IEC 27001:2013 A.16.1.6, NIST SP 800-53 AC-21/CA-7/SI-4",
            "COBIT 5 DSS04.03, ISA 62443-2-1:2009 4.3.2.5.3/4.3.4.5.1, NIST SP 800-53 CP-2/IR-8",
            "ISA 62443-2-1:2009 4.3.2.5.7/4.3.4.5.11, ISO/IEC 27001:2013 A.17.1.3, NIST SP 800-53 CP-4/IR-3/PM-14",
            "COBIT 5 APO07.01-05, ISA 62443-2-1:2009 4.3.3.2.1-3, NIST SP 800-53 PS Family",
            "ISO/IEC 27001:2013 A.12.6.1/A.18.2.2, NIST SP 800-53 RA-3/RA-5/SI-2",
        ],
    },
    10: {
        "title": "Protect - Maintenance of ICS & Information Systems",
        "function": "PROTECT",
        "summary": "Maintenance and repairs of industrial control and information system components are performed consistent with policies and procedures.",
        "subcategories": [
            "Maintenance and repair of organizational assets is performed and logged in a timely manner, with approved and controlled tools.",
            "Remote maintenance of organizational assets is approved, logged, and performed in a manner that prevents unauthorized access.",
        ],
        "example": "Establish an authorized maintenance tool list. Require maintenance windows approved via ITSM change management. Log all remote maintenance sessions in a centralized audit trail.",
        "solution": "Use jump servers / bastion hosts for all remote maintenance. Record sessions with CyberArk Session Manager or BeyondTrust. Require dual authorization for critical system maintenance.",
        "recommendation": "Approve all maintenance tools via the software whitelist. Disable remote maintenance capabilities when not in use. Review maintenance logs monthly for anomalies. Apply DoD maintenance policies per DoDI 8500.01.",
        "references": [
            "COBIT 5 BAI09.03, ISA 62443-2-1:2009 4.3.3.3.7, ISO/IEC 27001:2013 A.11.1.2/A.11.2.4-5, NIST SP 800-53 MA-2/3/5",
            "COBIT 5 DSS05.04, ISA 62443-2-1:2009 4.3.3.6.5-8, ISO/IEC 27001:2013 A.11.2.4/A.15.1.1/A.15.2.1, NIST SP 800-53 MA-4",
        ],
    },
    11: {
        "title": "Protect - Protective Technology",
        "function": "PROTECT",
        "summary": "Technical security solutions are managed to ensure the security and resilience of systems and assets, consistent with related policies, procedures, and agreements.",
        "subcategories": [
            "Audit/log records are determined, documented, implemented, and reviewed in accordance with policy.",
            "Removable media is protected and its use restricted according to policy.",
            "Access to systems and assets is controlled, incorporating the principle of least functionality.",
            "Communications and control networks are protected.",
        ],
        "example": "Deploy a SIEM (Splunk or Microsoft Sentinel) to collect, centralize, and correlate logs. Disable USB ports via Group Policy. Use application whitelisting (Carbon Black or AppLocker).",
        "solution": "Implement DISA-approved SIEM with log retention of 12+ months. Deploy NGFW (Palo Alto or Fortinet) with IPS enabled. Enforce network segmentation using VLANs and micro-segmentation.",
        "recommendation": "Audit logs daily for anomalies. Restrict removable media to approved devices only. Enforce application allowlisting. Protect ICS/OT networks with unidirectional security gateways. Review firewall rules quarterly.",
        "references": [
            "CCS CSC 1, COBIT 5 APO11.04, ISA 62443-3-3:2013 SR 2.8-2.12, ISO/IEC 27001:2013 A.12.4.1-4/A.12.7.1, NIST SP 800-53 AU Family",
            "COBIT5 DSS05.02/APO13.01, ISA 62443-3-3:2013 SR 2.3, ISO/IEC 27001:2013 A.8.2.2-3/A.8.3.1/3/A.11.2.9, NIST SP 800-53 MP-2/4/5/7",
            "COBIT5 DSS05.02, ISA 62443-2-1:2009 4.3.3.5.1-8/4.3.3.6.1-9/4.3.3.7.1-4, NIST SP 800-53 AC-3/CM-7",
            "CCS CSC 7, COBIT 5 DSS05.02/APO13.01, ISA 62443-3-3:2013 SR 3.1/3.5/3.8/4.1/4.3/5.1-3/7.1/7.6, NIST SP 800-53 AC-4/17/18/CP-8/SC-7",
        ],
    },
    12: {
        "title": "Detect - Anomalies & Events",
        "function": "DETECT",
        "summary": "Anomalous activity is detected in a timely manner and the potential impact of events is understood.",
        "subcategories": [
            "A baseline of network operations and expected data flows for users and systems is established and managed.",
            "Detected events are analyzed to understand attack targets and methods.",
            "Event data are aggregated and correlated from multiple sources and sensors.",
            "Impact of events is determined.",
            "Incident alert thresholds are established.",
        ],
        "example": "Use network flow analysis tools (NetFlow, Darktrace) to establish behavioral baselines. Alert on deviations exceeding 3 standard deviations from normal. Correlate events across endpoint, network, and cloud logs.",
        "solution": "Deploy a UEBA (User and Entity Behavior Analytics) solution like Microsoft Sentinel or Splunk UBA. Integrate with threat intel feeds (MISP, OpenCTI). Set alert thresholds based on MITRE ATT&CK TTPs.",
        "recommendation": "Define normal baselines for all critical systems. Set alert tuning cycles monthly to reduce false positives. Map detection rules to MITRE ATT&CK. Conduct threat hunting exercises quarterly.",
        "references": [
            "COBIT 5 DSS03.01, ISA 62443-2-1:2009 4.4.3.3, NIST SP 800-53 AC-4/CA-3/CM-2/SI-4",
            "ISA 62443-2-1:2009 4.3.4.5.6-8, ISA 62443-3-3:2013 SR 2.8-2.12/SR 3.9/6.1-2, NIST SP 800-53 AU-6/CA-7/IR-4/SI-4",
            "ISA 62443-3-3:2013 SR 6.1, NIST SP 800-53 AU-6/CA-7/IR-4/IR-5/IR-8/SI-4",
            "COBIT 5 APO12.06, NIST SP 800-53 CP-2/IR-4/RA-3/SI-4",
            "COBIT 5 APO12.06, ISA 62443-2-1:2009 4.2.3.10, NIST SP 800-53 IR-4/IR-5/IR-8",
        ],
    },
    13: {
        "title": "Detect - Continuous Security Monitoring",
        "function": "DETECT",
        "summary": "The information system and assets are monitored at discrete intervals to identify cybersecurity events and verify the effectiveness of protective measures.",
        "subcategories": [
            "The network is monitored to detect potential cybersecurity events.",
            "The physical environment is monitored to detect potential cybersecurity events.",
            "Personnel activity is monitored to detect potential cybersecurity events.",
            "Malicious code is detected.",
            "Unauthorized mobile code is detected.",
            "External service providers are monitored to detect cybersecurity events.",
            "Monitoring for unauthorized personnel, connections, devices, and software is performed.",
            "Vulnerability scans are performed.",
        ],
        "example": "Deploy Continuous Diagnostics and Mitigation (CDM) per CISA mandate. Use CrowdStrike EDR for endpoint monitoring. Run automated vulnerability scans weekly with Tenable.sc.",
        "solution": "Implement a 24x7 Security Operations Center (SOC) with tiered analyst structure. Use SOAR (Splunk SOAR, Palo Alto XSOAR) to automate tier-1 response. Feed all logs to centralized SIEM.",
        "recommendation": "Achieve CISA CDM Phase 1-4 compliance. Monitor all privileged accounts with enhanced logging. Run vulnerability scans weekly; patch Critical findings within 15 days. Report SOC metrics to leadership monthly.",
        "references": [
            "CCS CSC 14/16, COBIT 5 DSS05.07, ISA 62443-3-3:2013 SR 6.2, NIST SP 800-53 AC-2/AU-12/CA-7/CM-3/SC-5/7/SI-4",
            "ISA 62443-2-1:2009 4.3.3.3.8, NIST SP 800-53 CA-7/PE-3/PE-6/PE-20",
            "ISA 62443-3-3:2013 SR 6.2, ISO/IEC 27001:2013 A.12.4.1, NIST SP 800-53 AC-2/AU-12-13/CA-7/CM-10-11",
            "CCS CSC 5, COBIT 5 DSS05.01, ISA 62443-2-1:2009 4.3.4.3.8, NIST SP 800-53 SI-3",
            "ISA 62443-3-3:2013 SR 2.4, ISO/IEC 27001:2013 A.12.5.1, NIST SP 800-53 SC-18/SI-4/SC-44",
            "COBIT 5 APO07.06, NIST SP 800-53 CA-7/PS-7/SA-4/9/SI-4",
            "NIST SP 800-53 AU-12/CA-7/CM-3/8/PE-3/6/20/SI-4",
            "COBIT 5 BAI03.10, ISA 62443-2-1:2009 4.2.3.1/7, NIST SP 800-53 RA-5",
        ],
    },
    14: {
        "title": "Detect - Detection Processes",
        "function": "DETECT",
        "summary": "Detection processes and procedures are maintained and tested to ensure timely and adequate awareness of anomalous events.",
        "subcategories": [
            "Roles and responsibilities for detection are well defined to ensure accountability.",
            "Detection activities comply with all applicable requirements.",
            "Detection processes are tested.",
            "Event detection information is communicated to appropriate parties.",
            "Detection processes are continuously improved.",
        ],
        "example": "Define RACI matrix for detection roles. Run purple team exercises to test detection coverage against MITRE ATT&CK. Document and track detection gaps via a continuous improvement log.",
        "solution": "Use MITRE ATT&CK Navigator to map detection coverage. Conduct red/blue/purple team exercises biannually. Report detection effectiveness metrics to CISO monthly.",
        "recommendation": "Test all detection use cases quarterly. Measure Mean Time to Detect (MTTD) and set targets (<60 min for critical). Update detection rules after every significant incident. Share detection insights via ISAC.",
        "references": [
            "CCS CSC 5, COBIT 5 DSS05.01, ISA 62443-2-1:2009 4.4.3.1, ISO/IEC 27001:2013 A.6.1.1",
            "ISA 62443-2-1:2009 4.4.3.2, ISO/IEC 27001:2013 A.18.1.4, NIST SP 800-53 CA-2/CA-7/PM-14/SI-4",
            "COBIT 5 APO13.02, ISA 62443-2-1:2009 4.4.3.2, NIST SP 800-53 CA-2/CA-7/PE-3/PM-14/SI-3/4",
            "COBIT 5 APO12.06, ISA 62443-2-1:2009 4.3.4.5.9, NIST SP 800-53 AU-6/CA-2/CA-7/RA-5/SI-4",
            "COBIT 5 APO11.06/DSS04.05, ISA 62443-2-1:2009 4.4.3.4, NIST SP 800-53 CA-2/CA-7/PL-2/RA-5/SI-4/PM-14",
        ],
    },
    15: {
        "title": "Respond - Response Planning",
        "function": "RESPOND",
        "summary": "Response processes and procedures are executed and maintained to ensure timely response to detected cybersecurity events.",
        "subcategories": [
            "Response plan is executed during or after an event.",
        ],
        "example": "Activate the Incident Response Plan (IRP) when a Tier 1 incident is confirmed. Notify CISO within 1 hour. Report to US-CERT within 1 hour per FISMA requirements. Isolate affected systems.",
        "solution": "Maintain an up-to-date IRP per NIST SP 800-61 Rev 2. Integrate IRP with SOAR for automated playbook execution. Pre-position response team on-call rosters and contact lists.",
        "recommendation": "Review and update the IRP annually and after every major incident. Conduct tabletop exercises biannually. Ensure all stakeholders know their response roles. Test communication trees quarterly.",
        "references": [
            "COBIT 5 BAI01.10, CCS CSC 18, ISA 62443-2-1:2009 4.3.4.5.1, ISO/IEC 27001:2013 A.16.1.5, NIST SP 800-53 CP-2/CP-10/IR-4/IR-8",
        ],
    },
    16: {
        "title": "Respond - Response Analysis",
        "function": "RESPOND",
        "summary": "Analysis is conducted to ensure adequate response and support recovery activities.",
        "subcategories": [
            "Notifications from detection systems are investigated.",
            "The impact of the incident is understood.",
            "Forensics are performed.",
            "Incidents are categorized consistent with response plans.",
        ],
        "example": "Use Velociraptor or CrowdStrike Falcon for digital forensics. Classify incidents using DoD's five-tier severity scale. Brief leadership within 4 hours of Tier 1 incident confirmation.",
        "solution": "Establish a DFIR (Digital Forensics and Incident Response) capability with trained analysts. Maintain forensic tools on an air-gapped forensic workstation. Use a dedicated case management system (TheHive).",
        "recommendation": "Preserve forensic evidence per chain-of-custody procedures. Complete impact assessment within 2 hours. Categorize all incidents per CISA severity schema. Conduct after-action reviews within 5 business days.",
        "references": [
            "COBIT 5 DSS02.07, ISA 62443-2-1:2009 4.3.4.5.6-8, ISA 62443-3-3:2013 SR 6.1, NIST SP 800-53 AU-6/CA-7/IR-4/5/PE-6/SI-4",
            "ISA 62443-2-1:2009 4.3.4.5.6-8, ISO/IEC 27001:2013 A.16.1.6, NIST SP 800-53 CP-2/IR-4",
            "ISA 62443-3-3:2013 SR 2.8-2.12/3.9/6.1, ISO/IEC 27001:2013 A.16.1.7, NIST SP 800-53 AU-7/IR-4",
            "ISA 62443-2-1:2009 4.3.4.5.6, ISO/IEC 27001:2013 A.16.1.4, NIST SP 800-53 CP-2/IR-4/5/8",
        ],
    },
    17: {
        "title": "Respond - Mitigation Activities",
        "function": "RESPOND",
        "summary": "Mitigation activities are performed to prevent expansion of an event, mitigate its effects, and eradicate the incident.",
        "subcategories": [
            "Incidents are contained.",
            "Incidents are mitigated.",
            "Newly identified vulnerabilities are mitigated or documented as accepted risks.",
        ],
        "example": "Use automated SOAR playbooks to isolate infected endpoints within minutes. Block malicious IPs at the perimeter firewall. Remove malware using EDR tools. Patch the exploited vulnerability within 72 hours.",
        "solution": "Pre-build containment playbooks in SOAR (network isolation, account lockout, firewall block). Maintain a vulnerability exception process for risks that cannot be immediately patched.",
        "recommendation": "Define containment strategies for all incident categories. Achieve containment of critical incidents within 4 hours. Track all accepted risk exceptions with expiration dates. Review open vulnerabilities weekly.",
        "references": [
            "ISA 62443-2-1:2009 4.3.4.5.6, ISA 62443-3-3:2013 SR 5.1/5.2/5.4, ISO/IEC 27001:2013 A.16.1.5, NIST SP 800-53 IR-4",
            "ISA 62443-2-1:2009 4.3.4.5.6/10, ISO/IEC 27001:2013 A.12.2.1/A.16.1.5, NIST SP 800-53 IR-4",
            "ISO/IEC 27001:2013 A.12.6.1, NIST SP 800-53 CA-7/RA-3/RA-5",
        ],
    },
    18: {
        "title": "Respond - Improving Response Activities",
        "function": "RESPOND",
        "summary": "Organizational response activities are improved by incorporating lessons learned from current and previous detection/response activities.",
        "subcategories": [
            "Response plans incorporate lessons learned.",
            "Response strategies are updated.",
        ],
        "example": "Hold an After-Action Review (AAR) within 5 business days of every Tier 1-2 incident. Update the IRP based on findings. Track improvement actions in a lessons-learned registry.",
        "solution": "Establish a formal lessons-learned process with defined action owners and due dates. Integrate findings into training programs. Share sanitized lessons with the relevant ISAC community.",
        "recommendation": "Mandate AARs for all significant incidents. Track lessons-learned action items to closure in ITSM. Update detection rules and response playbooks within 30 days of AAR completion. Review response KPIs quarterly.",
        "references": [
            "COBIT 5 BAI01.13, ISA 62443-2-1:2009 4.3.4.5.10/4.4.3.4, ISO/IEC 27001:2013 A.16.1.6, NIST SP 800-53 CP-2/IR-4/IR-8",
            "NIST SP 800-53 CP-2/IR-4/IR-8",
        ],
    },
    19: {
        "title": "Recover - Recovery Planning",
        "function": "RECOVER",
        "summary": "Recovery processes and procedures are executed and maintained to ensure timely restoration of systems or assets affected by cybersecurity events.",
        "subcategories": [
            "Recovery plan is executed during or after an event.",
        ],
        "example": "Activate the Disaster Recovery Plan (DRP) when production systems are compromised. Failover to DR site within the defined RTO (Recovery Time Objective). Restore from last known clean backup.",
        "solution": "Maintain a tested DRP per NIST SP 800-34. Use cloud-based DR (AWS GovCloud or Azure Government) for rapid failover. Define and test RTO/RPO for all Tier 1 systems annually.",
        "recommendation": "Define RTO/RPO for all critical systems and document in the DRP. Test full recovery annually; partial tests quarterly. Ensure backups are stored offline and air-gapped. Brief leadership on DR readiness biannually.",
        "references": [
            "CCS CSC 8, COBIT 5 DSS02.05/DSS03.04, ISO/IEC 27001:2013 A.16.1.5, NIST SP 800-53 CP-10/IR-4/IR-8",
        ],
    },
    20: {
        "title": "Recover - Recovery Planning Improvements",
        "function": "RECOVER",
        "summary": "Recovery planning and processes are improved by incorporating lessons learned into future activities.",
        "subcategories": [
            "Recovery plans incorporate lessons learned.",
            "Recovery strategies are updated.",
        ],
        "example": "Update the DRP after each recovery exercise or actual DR event. Benchmark RTO/RPO performance and set improvement targets. Track improvement actions via the ITSM system.",
        "solution": "Conduct formal DRP reviews biannually. Update recovery strategies when technology changes or new risks emerge. Share recovery improvement insights within the DoD community.",
        "recommendation": "Update DRP within 30 days of every DR test or real event. Set aggressive RTO/RPO improvement targets annually. Align recovery improvements to evolving threat landscape. Report recovery maturity to CISO quarterly.",
        "references": [
            "COBIT 5 BAI05.07, ISA 62443-2-1 4.4.3.4, NIST SP 800-53 CP-2/IR-4/IR-8",
            "COBIT 5 BAI07.08, NIST SP 800-53 CP-2/IR-4/IR-8",
        ],
    },
    21: {
        "title": "Recover - Restoration Activities Communications",
        "function": "RECOVER",
        "summary": "Restoration activities are coordinated with internal and external parties such as coordinating centers, service providers, owners of attacking systems, victims, other CSIRTs, and vendors.",
        "subcategories": [
            "Public relations are managed.",
            "Reputation after an event is repaired.",
            "Recovery activities are communicated to internal stakeholders and executive and management teams.",
        ],
        "example": "Coordinate with DISA CSSP and US-CERT during recovery. Designate a single Public Affairs Officer (PAO) for all external communications. Brief senior leadership daily during major recovery operations.",
        "solution": "Pre-draft communication templates for common incident scenarios. Establish a Crisis Communications Plan with legal and PR teams. Use a secure communications channel for recovery coordination.",
        "recommendation": "Activate crisis communications plan within 2 hours of incident declaration. Brief leadership every 4 hours during active recovery. Publish external communications only after legal review. Conduct reputation assessment post-recovery.",
        "references": [
            "COBIT 5 EDM03.02",
            "COBIT 5 MEA03.02",
            "NIST SP 800-53 CP-2/IR-4",
        ],
    },
}

FUNCTION_COLORS = {
    "IDENTIFY": "#0047AB",
    "PROTECT":  "#2e7d32",
    "DETECT":   "#f0a500",
    "RESPOND":  "#c62828",
    "RECOVER":  "#6a1b9a",
}

MENU_ITEMS = [f"{k}. {v['title']}" for k, v in FRAMEWORK_ITEMS.items()]

# ══════════════════════════════════════════════════════════════════════════════
#  SYNTHETIC DATA
# ══════════════════════════════════════════════════════════════════════════════
@st.cache_data
def generate_synthetic_data(n: int) -> pd.DataFrame:
    random.seed(42)
    functions = ["IDENTIFY", "PROTECT", "DETECT", "RESPOND", "RECOVER"]
    item_titles = [v["title"] for v in FRAMEWORK_ITEMS.values()]
    departments = ["CYBERCOM", "DISA/BDE5", "NSA", "TRANSCOM", "DIA", "SOCOM", "CENTCOM", "EUCOM", "NORTHCOM", "PACOM"]
    severities = ["Critical", "High", "Medium", "Low"]
    sev_weights = [0.10, 0.25, 0.40, 0.25]
    statuses = ["Compliant", "Non-Compliant", "Partial", "In Remediation"]
    stat_weights = [0.55, 0.18, 0.17, 0.10]
    infra_types = ["Network", "Cloud", "ICS/OT", "Endpoint", "Application", "Physical"]
    base_date = datetime(2024, 1, 1)
    records = []
    for i in range(1, n + 1):
        func = random.choice(functions)
        status = random.choices(statuses, weights=stat_weights)[0]
        score = random.randint(70, 100) if status == "Compliant" else (
            random.randint(50, 69) if status == "Partial" else random.randint(20, 49))
        records.append({
            "Record_ID": f"CS-{i:04d}",
            "Timestamp": base_date + timedelta(days=random.randint(0, 364), hours=random.randint(0, 23)),
            "Framework_Function": func,
            "Checklist_Item": random.choice([t for t in item_titles if True]),
            "Department": random.choice(departments),
            "Infrastructure_Type": random.choice(infra_types),
            "Severity": random.choices(severities, weights=sev_weights)[0],
            "Compliance_Status": status,
            "Compliance_Score": score,
            "Vulnerabilities_Found": random.randint(0, 50),
            "Mitigated": random.randint(0, 50),
            "MTTD_Hours": round(random.uniform(0.5, 48.0), 1),
            "MTTR_Hours": round(random.uniform(1.0, 120.0), 1),
            "Incident_Flagged": status == "Non-Compliant" and random.random() < 0.5,
        })
    return pd.DataFrame(records)

# ══════════════════════════════════════════════════════════════════════════════
#  EXPORT HELPERS
# ══════════════════════════════════════════════════════════════════════════════
def _pdf_safe(text: str) -> str:
    replacements = [
        ("\u2013", "-"), ("\u2014", "--"), ("\u2018", "'"), ("\u2019", "'"),
        ("\u201c", '"'), ("\u201d", '"'), ("\u2026", "..."), ("\u2022", "*"),
        ("\u00a0", " "), ("\u2265", ">="), ("\u2264", "<="), ("\u00ae", "(R)"),
        ("\u2122", "(TM)"), ("\u00e9", "e"), ("\u00e8", "e"), ("\u00e0", "a"),
        ("\u00fc", "u"), ("\u00f6", "o"), ("\u00e4", "a"), ("&", "and"),
    ]
    for char, rep in replacements:
        text = text.replace(char, rep)
    return text.encode("latin-1", errors="ignore").decode("latin-1")


def build_text_report(item_num: int, data: dict) -> str:
    lines = [
        "=" * 80,
        "  CYBERSECURITY FRAMEWORK APPLICATION FOR DoD CRITICAL INFRASTRUCTURE",
        "  Developed by Randy Singh | KalSnet (KNet) Consulting Group",
        f"  Item {item_num}: {data['title']}",
        f"  Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "=" * 80, "",
        f"FUNCTION: {data['function']}",
        f"SUMMARY: {data['summary']}", "",
        "SUBCATEGORIES:",
    ]
    for i, sub in enumerate(data["subcategories"], 1):
        lines.append(f"  {i}. {sub}")
    lines += ["", f"EXAMPLE: {data['example']}", "",
              f"RECOMMENDED SOLUTION: {data['solution']}", "",
              f"RECOMMENDATION: {data['recommendation']}", "",
              "REFERENCE DOCUMENTS:"]
    for ref in data["references"]:
        lines.append(f"  * {ref}")
    return "\n".join(lines)


def build_json_report(item_num: int, data: dict) -> str:
    payload = {
        "framework": "CyberSecurity Framework for DoD Critical Infrastructure",
        "author": "Randy Singh | KalSnet (KNet) Consulting Group",
        "item_number": item_num,
        "generated": datetime.now().isoformat(),
        "data": data,
    }
    return json.dumps(payload, indent=2, default=str)


def build_pdf_bytes(item_num: int, data: dict) -> bytes:
    if not FPDF_AVAILABLE:
        return b""
    try:
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 15)
        pdf.set_text_color(0, 71, 171)
        pdf.cell(0, 9, _pdf_safe("CYBERSECURITY FRAMEWORK APPLICATION"), new_x="LMARGIN", new_y="NEXT", align="C")
        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 7, _pdf_safe("FOR DoD CRITICAL INFRASTRUCTURE"), new_x="LMARGIN", new_y="NEXT", align="C")
        pdf.cell(0, 7, _pdf_safe("Developed by Randy Singh | KalSnet (KNet) Consulting Group"), new_x="LMARGIN", new_y="NEXT", align="C")
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(80, 80, 80)
        pdf.cell(0, 6, _pdf_safe(f"Item {item_num}: {data['title']}   |   {datetime.now().strftime('%Y-%m-%d %H:%M')}"),
                 new_x="LMARGIN", new_y="NEXT", align="C")
        pdf.set_draw_color(0, 71, 171)
        pdf.set_line_width(0.8)
        y = pdf.get_y() + 2
        pdf.line(10, y, 200, y)
        pdf.ln(5)

        sections = [
            ("FUNCTION", data["function"]),
            ("SUMMARY", data["summary"]),
            ("EXAMPLE", data["example"]),
            ("RECOMMENDED SOLUTION", data["solution"]),
            ("RECOMMENDATION", data["recommendation"]),
        ]
        for heading, body in sections:
            pdf.set_font("Helvetica", "B", 10)
            pdf.set_text_color(0, 71, 171)
            pdf.cell(0, 7, _pdf_safe(heading), new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", "", 9)
            pdf.set_text_color(30, 30, 30)
            pdf.multi_cell(0, 6, _pdf_safe(body))
            pdf.ln(2)

        pdf.set_font("Helvetica", "B", 10)
        pdf.set_text_color(0, 71, 171)
        pdf.cell(0, 7, "SUBCATEGORIES", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(30, 30, 30)
        for i, sub in enumerate(data["subcategories"], 1):
            pdf.multi_cell(0, 6, _pdf_safe(f"  {i}. {sub}"))

        pdf.ln(2)
        pdf.set_font("Helvetica", "B", 10)
        pdf.set_text_color(0, 71, 171)
        pdf.cell(0, 7, "REFERENCE DOCUMENTS", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 8)
        pdf.set_text_color(30, 30, 30)
        for ref in data["references"]:
            pdf.multi_cell(0, 5, _pdf_safe(f"  * {ref}"))

        raw = pdf.output()
        if isinstance(raw, (bytes, bytearray)):
            return bytes(raw)
        return raw.encode("latin-1")
    except Exception as exc:
        try:
            fb = FPDF(); fb.add_page(); fb.set_font("Helvetica", "", 10)
            fb.multi_cell(0, 8, _pdf_safe(f"PDF error: {exc}"))
            raw = fb.output()
            return bytes(raw) if isinstance(raw, (bytes, bytearray)) else raw.encode("latin-1")
        except Exception:
            return b""


def build_docx_bytes(item_num: int, data: dict) -> bytes:
    if not DOCX_AVAILABLE:
        return b""
    try:
        doc = DocxDocument()
        t = doc.add_heading("CYBERSECURITY FRAMEWORK APPLICATION", 0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in t.runs:
            run.font.color.rgb = RGBColor(0, 71, 171)
        for line in ["FOR DoD CRITICAL INFRASTRUCTURE",
                     "Developed by Randy Singh | KalSnet (KNet) Consulting Group"]:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if p.runs:
                p.runs[0].bold = True
                p.runs[0].font.color.rgb = RGBColor(0, 71, 171)
        ts = doc.add_paragraph(f"Item {item_num}: {data['title']}   |   {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        ts.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        for heading, body in [("Function", data["function"]), ("Summary", data["summary"]),
                               ("Example", data["example"]), ("Recommended Solution", data["solution"]),
                               ("Recommendation", data["recommendation"])]:
            h = doc.add_heading(heading, level=2)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0, 71, 171)
            doc.add_paragraph(str(body))

        h = doc.add_heading("Subcategories", level=2)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0, 71, 171)
        for sub in data["subcategories"]:
            doc.add_paragraph(sub, style="List Bullet")

        h = doc.add_heading("Reference Documents", level=2)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0, 71, 171)
        for ref in data["references"]:
            doc.add_paragraph(ref, style="List Bullet")

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except Exception as exc:
        try:
            doc = DocxDocument(); doc.add_paragraph(f"Word error: {exc}")
            buf = io.BytesIO(); doc.save(buf); return buf.getvalue()
        except Exception:
            return b""

# ══════════════════════════════════════════════════════════════════════════════
#  SIDEBAR
# ══════════════════════════════════════════════════════════════════════════════
with st.sidebar:
    st.markdown("## Cyber Framework")
    st.markdown("---")
    nav_mode = st.radio("Navigation Mode", ["Browse Checklist Items", "Synthetic Data & Analytics", "Framework Overview"])
    st.markdown("---")

    selected_item_num = 1
    if nav_mode == "Browse Checklist Items":
        selected_label = st.selectbox("Select Checklist Item", MENU_ITEMS)
        selected_item_num = int(selected_label.split(".")[0])

    st.markdown("---")
    st.markdown("**About**")
    st.markdown("Randy Singh  \nComputer Scientist \nKalSnet (KNet) Consulting")
    st.markdown(" (301) 503-2918 ")

# ── Export sidebar ────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("---")
    st.markdown("### Export Results")

    if nav_mode == "Browse Checklist Items":
        exp_data = FRAMEWORK_ITEMS[selected_item_num]
        exp_num  = selected_item_num
    else:
        exp_data = FRAMEWORK_ITEMS[1]
        exp_num  = 1

    txt_data = build_text_report(exp_num, exp_data)
    st.download_button("Export as Text (.txt)", data=txt_data,
                       file_name=f"cs_framework_item{exp_num}.txt", mime="text/plain", use_container_width=True)

    json_data = build_json_report(exp_num, exp_data)
    st.download_button("Export as JSON (.json)", data=json_data,
                       file_name=f"cs_framework_item{exp_num}.json", mime="application/json", use_container_width=True)

    if FPDF_AVAILABLE:
        pdf_bytes = build_pdf_bytes(exp_num, exp_data)
        st.download_button("Export as PDF (.pdf)", data=pdf_bytes,
                           file_name=f"cs_framework_item{exp_num}.pdf", mime="application/pdf", use_container_width=True)
    else:
        st.warning("PDF: run pip install fpdf2")

    if DOCX_AVAILABLE:
        docx_bytes = build_docx_bytes(exp_num, exp_data)
        st.download_button("Export as Word (.docx)", data=docx_bytes,
                           file_name=f"cs_framework_item{exp_num}.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                           use_container_width=True)
    else:
        st.warning("Word: run pip install python-docx")

# ══════════════════════════════════════════════════════════════════════════════
#  MAIN CONTENT
# ══════════════════════════════════════════════════════════════════════════════

# ── FRAMEWORK OVERVIEW ────────────────────────────────────────────────────────
if nav_mode == "Framework Overview":
    st.markdown('<div class="section-header">NIST Cybersecurity Framework - DoD Critical Infrastructure Overview</div>', unsafe_allow_html=True)

    c1, c2, c3, c4, c5 = st.columns(5)
    for col, (func, color, count) in zip([c1,c2,c3,c4,c5], [
        ("IDENTIFY", "#0047AB", 5), ("PROTECT", "#2e7d32", 6),
        ("DETECT", "#f0a500", 3), ("RESPOND", "#c62828", 4), ("RECOVER", "#6a1b9a", 3)
    ]):
        col.metric(func, f"{count} items")

    st.markdown("---")
    st.markdown("#### NIST CSF 5-Function Flow")
    try:
        import graphviz
        dot = graphviz.Digraph(graph_attr={"rankdir": "LR", "bgcolor": "#f8f9ff", "fontname": "Helvetica"})
        dot.attr("node", shape="box", style="filled", fontname="Helvetica", fontsize="11")
        funcs = [("IDENTIFY","#c6d9ff"),("PROTECT","#c8e6c9"),("DETECT","#fff9c4"),("RESPOND","#ffcdd2"),("RECOVER","#e1bee7")]
        for name, color in funcs:
            dot.node(name, name, fillcolor=color, color="#333")
        for a, b in zip([f[0] for f in funcs], [f[0] for f in funcs[1:]]):
            dot.edge(a, b)
        dot.edge("RECOVER", "IDENTIFY", label="Continuous Improvement", style="dashed")
        st.graphviz_chart(dot, use_container_width=True)
    except Exception:
        st.code("IDENTIFY --> PROTECT --> DETECT --> RESPOND --> RECOVER --> (back to IDENTIFY)")

    st.markdown("---")
    st.markdown("#### Framework Item Summary by Function")
    func_data = {}
    for item in FRAMEWORK_ITEMS.values():
        f = item["function"]
        func_data[f] = func_data.get(f, 0) + 1
    fig_pie = px.pie(values=list(func_data.values()), names=list(func_data.keys()),
                     color=list(func_data.keys()),
                     color_discrete_map=FUNCTION_COLORS,
                     title="21 Checklist Items by CSF Function")
    st.plotly_chart(fig_pie, use_container_width=True)

    st.markdown("#### All 21 Checklist Items at a Glance")
    for num, item in FRAMEWORK_ITEMS.items():
        color = FUNCTION_COLORS[item["function"]]
        st.markdown(
            f'<div class="req-card"><span style="background:{color};color:white;padding:2px 8px;border-radius:3px;font-size:0.78rem;font-weight:700">{item["function"]}</span>'
            f' &nbsp; <b>{num}. {item["title"]}</b><br><span style="color:#555;font-size:0.85rem">{item["summary"][:140]}...</span></div>',
            unsafe_allow_html=True
        )

# ── CHECKLIST ITEM DETAIL ─────────────────────────────────────────────────────
elif nav_mode == "Browse Checklist Items":
    data = FRAMEWORK_ITEMS[selected_item_num]
    func_color = FUNCTION_COLORS[data["function"]]

    st.markdown(
        f'<div class="section-header" style="background:linear-gradient(90deg,{func_color} 0%,{func_color}cc 100%);">'
        f'Item {selected_item_num} of 21 &nbsp;|&nbsp; {data["title"]}</div>',
        unsafe_allow_html=True
    )
    st.markdown(
        f'<span style="background:{func_color};color:white;padding:4px 14px;border-radius:4px;'
        f'font-weight:700;font-size:0.9rem">CSF FUNCTION: {data["function"]}</span>',
        unsafe_allow_html=True
    )
    st.markdown("<br>", unsafe_allow_html=True)

    # Summary
    st.markdown(f'<div class="req-card"><b>Summary:</b><br>{data["summary"]}</div>', unsafe_allow_html=True)

    col_l, col_r = st.columns(2)
    with col_l:
        st.markdown('<div class="sol-card"><b>Example:</b></div>', unsafe_allow_html=True)
        st.markdown(f'<div class="sol-card">{data["example"]}</div>', unsafe_allow_html=True)
        st.markdown('<div class="sol-card"><b>Recommended Solution:</b></div>', unsafe_allow_html=True)
        st.markdown(f'<div class="sol-card">{data["solution"]}</div>', unsafe_allow_html=True)
    with col_r:
        st.markdown('<div class="rec-card"><b>Recommendation:</b></div>', unsafe_allow_html=True)
        st.markdown(f'<div class="rec-card">{data["recommendation"]}</div>', unsafe_allow_html=True)

    # Subcategories
    st.markdown("---")
    st.markdown(f"#### Subcategories ({len(data['subcategories'])} items)")
    for i, sub in enumerate(data["subcategories"], 1):
        st.markdown(f'<div class="req-card"><b>{i}.</b> {sub}</div>', unsafe_allow_html=True)

    # References
    with st.expander("Reference Documents (click to expand)"):
        for ref in data["references"]:
            st.markdown(f'<div class="ref-card">{ref}</div>', unsafe_allow_html=True)

    # Flow diagram for this item
    st.markdown("---")
    st.markdown("#### Process Flow Diagram")
    try:
        import graphviz
        g = graphviz.Digraph(graph_attr={"rankdir": "TD", "bgcolor": "#f8f9ff"})
        g.attr("node", shape="box", style="filled", fontname="Helvetica", fontsize="10")
        g.node("Policy", "Policy / Governance", fillcolor="#dce8ff", color="#0047AB")
        g.node("Implementation", f"Implement:\n{data['title']}", fillcolor=func_color, fontcolor="white", color=func_color)
        g.node("Monitor", "Monitor & Audit", fillcolor="#fff0cc", color="#f0a500")
        g.node("Review", "Review & Improve", fillcolor="#e8f5e9", color="#2e7d32")
        g.node("Report", "Report to Leadership", fillcolor="#fce4ec", color="#c62828")
        g.edge("Policy", "Implementation", label="Drives")
        g.edge("Implementation", "Monitor", label="Feeds")
        g.edge("Monitor", "Review", label="Triggers")
        g.edge("Review", "Policy", label="Updates", style="dashed")
        g.edge("Review", "Report", label="Informs")
        st.graphviz_chart(g, use_container_width=True)
    except Exception:
        st.info("Install graphviz system package for flow diagrams.")

    # Navigation
    st.markdown("---")
    nav_c1, nav_c2, nav_c3 = st.columns([1, 2, 1])
    with nav_c1:
        if selected_item_num > 1:
            prev_label = MENU_ITEMS[selected_item_num - 2]
            st.info(f"Previous: {prev_label[:40]}...")
    with nav_c3:
        if selected_item_num < 21:
            next_label = MENU_ITEMS[selected_item_num]
            st.success(f"Next: {next_label[:40]}...")

# ── SYNTHETIC DATA ────────────────────────────────────────────────────────────
elif nav_mode == "Synthetic Data & Analytics":
    st.markdown('<div class="section-header">Synthetic DoD Cybersecurity Compliance Data & Analytics</div>', unsafe_allow_html=True)
    num_records = st.slider("Number of Synthetic Records", min_value=0, max_value=300, value=150, step=10)

    if num_records == 0:
        st.info("Move the slider above 0 to generate synthetic compliance data.")
    else:
        df = generate_synthetic_data(num_records)

        STATUS_COLORS = {"Compliant": "#2e7d32", "Non-Compliant": "#c62828",
                         "Partial": "#f0a500", "In Remediation": "#0047AB"}

        # KPIs
        k1, k2, k3, k4, k5, k6 = st.columns(6)
        k1.metric("Total Records", num_records)
        k2.metric("Compliant", int((df["Compliance_Status"] == "Compliant").sum()),
                  f'{(df["Compliance_Status"]=="Compliant").mean()*100:.0f}%')
        k3.metric("Non-Compliant", int((df["Compliance_Status"] == "Non-Compliant").sum()))
        k4.metric("Avg Score", f"{df['Compliance_Score'].mean():.1f}")
        k5.metric("Avg MTTD (hrs)", f"{df['MTTD_Hours'].mean():.1f}")
        k6.metric("Incidents", int(df["Incident_Flagged"].sum()))

        st.markdown("---")
        tab_raw, tab_func, tab_dept, tab_sev, tab_trend = st.tabs(
            ["Raw Data", "By Function", "By Department", "By Severity", "Trends"])

        with tab_raw:
            st.dataframe(df, use_container_width=True, height=380)
            st.download_button("Download CSV", df.to_csv(index=False).encode(),
                               "cs_framework_data.csv", "text/csv")

        with tab_func:
            c1, c2 = st.columns(2)
            func_status = df.groupby(["Framework_Function", "Compliance_Status"]).size().reset_index(name="Count")
            fig_fs = px.bar(func_status, x="Framework_Function", y="Count", color="Compliance_Status",
                            barmode="group", color_discrete_map=STATUS_COLORS,
                            title="Compliance Status by CSF Function")
            c1.plotly_chart(fig_fs, use_container_width=True)

            func_score = df.groupby("Framework_Function")["Compliance_Score"].mean().reset_index()
            fig_fsc = px.bar(func_score, x="Framework_Function", y="Compliance_Score",
                             color="Compliance_Score",
                             color_continuous_scale=["#ff4444", "#f0a500", "#2e7d32"],
                             title="Average Compliance Score by CSF Function", range_y=[0, 100])
            fig_fsc.update_layout(coloraxis_showscale=False)
            c2.plotly_chart(fig_fsc, use_container_width=True)

            # Radar
            funcs_r = ["IDENTIFY", "PROTECT", "DETECT", "RESPOND", "RECOVER"]
            scores_r = [df[df["Framework_Function"] == f]["Compliance_Score"].mean() for f in funcs_r]
            fig_radar = go.Figure()
            fig_radar.add_trace(go.Scatterpolar(
                r=scores_r + [scores_r[0]], theta=funcs_r + [funcs_r[0]],
                fill="toself", name="Avg Score", line_color="#0047AB", fillcolor="rgba(0,71,171,0.2)"))
            fig_radar.add_trace(go.Scatterpolar(
                r=[100]*6, theta=funcs_r + [funcs_r[0]],
                fill="toself", name="Target (100)", line_color="#f0a500", fillcolor="rgba(240,165,0,0.05)"))
            fig_radar.update_layout(polar=dict(radialaxis=dict(visible=True, range=[0, 100])),
                                    title="CSF Function Compliance Radar", height=420)
            st.plotly_chart(fig_radar, use_container_width=True)

        with tab_dept:
            dept_status = df.groupby(["Department", "Compliance_Status"]).size().reset_index(name="Count")
            fig_ds = px.bar(dept_status, x="Department", y="Count", color="Compliance_Status",
                            barmode="stack", color_discrete_map=STATUS_COLORS,
                            title="Compliance by Department")
            st.plotly_chart(fig_ds, use_container_width=True)

            infra_counts = df["Infrastructure_Type"].value_counts().reset_index()
            infra_counts.columns = ["Type", "Count"]
            fig_pie = px.pie(infra_counts, values="Count", names="Type",
                             title="Compliance Records by Infrastructure Type",
                             color_discrete_sequence=px.colors.qualitative.Bold)
            st.plotly_chart(fig_pie, use_container_width=True)

        with tab_sev:
            sev_status = df.groupby(["Severity", "Compliance_Status"]).size().reset_index(name="Count")
            sev_order = ["Critical", "High", "Medium", "Low"]
            fig_sev = px.bar(sev_status, x="Severity", y="Count", color="Compliance_Status",
                             barmode="group", color_discrete_map=STATUS_COLORS,
                             category_orders={"Severity": sev_order},
                             title="Compliance Status by Severity")
            st.plotly_chart(fig_sev, use_container_width=True)

            fig_box = px.box(df, x="Framework_Function", y="MTTD_Hours", color="Framework_Function",
                             color_discrete_map=FUNCTION_COLORS,
                             title="Mean Time to Detect (MTTD) Distribution by CSF Function")
            st.plotly_chart(fig_box, use_container_width=True)

        with tab_trend:
            df["Month"] = pd.to_datetime(df["Timestamp"]).dt.to_period("M").astype(str)
            trend = df.groupby(["Month", "Compliance_Status"]).size().reset_index(name="Count")
            fig_trend = px.line(trend, x="Month", y="Count", color="Compliance_Status",
                                markers=True, color_discrete_map=STATUS_COLORS,
                                title="Compliance Trend Over Time")
            st.plotly_chart(fig_trend, use_container_width=True)

            df["Month_dt"] = pd.to_datetime(df["Timestamp"]).dt.to_period("M").astype(str)
            score_trend = df.groupby("Month_dt")["Compliance_Score"].mean().reset_index()
            fig_st = px.line(score_trend, x="Month_dt", y="Compliance_Score", markers=True,
                             title="Average Compliance Score Trend", line_shape="spline")
            fig_st.add_hline(y=80, line_dash="dash", line_color="green", annotation_text="Target 80%")
            st.plotly_chart(fig_st, use_container_width=True)

# ══════════════════════════════════════════════════════════════════════════════
#  FOOTER
# ══════════════════════════════════════════════════════════════════════════════
st.markdown("---")
st.markdown(
    '<div style="text-align:center;color:#888;font-size:0.8rem;">'
    'CyberSecurity Framework Application for DoD Critical Infrastructure &nbsp;|&nbsp; '
    'Developed by Randy Singh &nbsp;|&nbsp; '
    'KalSnet (KNet) Consulting Group &nbsp;|&nbsp; '
    f'KNet/Technology Innovation Team &nbsp;|&nbsp; {datetime.now().year}'
    '</div>',
    unsafe_allow_html=True,
)
