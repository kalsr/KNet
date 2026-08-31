# GraphSentinel: Defense and Cyber Intelligence Platform
# Single File Streamlit Edition
# Developed by Randy Singh from Kalsnet (KNet) Consulting Group
# A production level Streamlit application implementing thirteen Knowledge
# Graph plus AI use cases for cyber defense and defense department analysis.
# Each use case tab supports synthetic data generation with a live record
# count bar chart, or real data upload, a schema explanation with a
# Graphviz flow diagram, interactive graph visualization, AI analysis
# powered by Groq or Gemini, and export of results to PDF, Word, text and
# CSV formats.
# This single file contains all configuration, data generation, graph
# analytics, LLM integration, export utilities, and the Streamlit interface.
import io
import json
import random
import string
from datetime import datetime
import requests
import pandas as pd
import networkx as nx
import plotly.graph_objects as go
import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.units import inch
# =============================================================================
# SECTION 1: USE CASE CONFIGURATION - THIRTEEN DEFENSE AND CYBER USE CASES
# =============================================================================
USE_CASES = {
    "attack_path": {
        "order": 1,
        "title": "Attack Path and Lateral Movement Analysis",
        "graph_provides": "Endpoint to Credential to Privilege to System to Critical Asset relationships",
        "ai_does": "Maps lateral movement paths and explains attacker reachability",
        "description": (
            "This use case builds a knowledge graph connecting endpoints, the "
            "credentials used to authenticate on them, the privilege levels "
            "those credentials grant, the systems reachable with that privilege, "
            "and the critical assets an attacker ultimately wants to reach. "
            "Security teams use this connected view to see exactly how a single "
            "compromised laptop can lead to full domain compromise. The AI layer "
            "reviews the reachable path and explains attacker options in plain "
            "language for incident responders."
        ),
        "node_types": ["Endpoint", "Credential", "Privilege Level", "System", "Critical Asset"],
        "edge_types": [
            ("Endpoint", "Credential", "authenticates_with"),
            ("Credential", "Privilege Level", "grants"),
            ("Privilege Level", "System", "accesses"),
            ("System", "Critical Asset", "can_reach"),
        ],
        "fields": {
            "Endpoint": ["endpoint_id", "hostname", "os_version", "patch_level"],
            "Credential": ["credential_id", "credential_type", "last_rotated_days"],
            "Privilege Level": ["privilege_id", "privilege_name", "privilege_tier"],
            "System": ["system_id", "system_name", "criticality"],
            "Critical Asset": ["asset_id", "asset_name", "business_impact"],
        },
        "formulas": [
            ("Attack Path Length", "Number of edges on the shortest path from a compromised endpoint node to a critical asset node. Shorter paths represent higher immediate risk."),
            ("Reachability Score", "Number of critical asset nodes reachable from a given endpoint divided by the total number of critical asset nodes in the graph."),
            ("Privilege Escalation Risk", "Count of distinct privilege tier increases along the shortest path from endpoint to asset, weighted by the sensitivity of each tier."),
            ("Blast Radius", "Number of distinct systems and assets reachable if a given credential is compromised, used to prioritize credential rotation and segmentation."),
        ],
        "benefits": [
            "Shows the exact attacker path from a compromised endpoint to the crown jewel asset, not just isolated alerts",
            "Prioritizes credential rotation and network segmentation based on measured blast radius",
            "Gives incident responders an explainable attack narrative instead of a raw log dump",
            "Reduces mean time to containment by identifying the highest value choke points to isolate first",
        ],
        "ai_prompt_focus": "explain the attacker reachable path from the compromised endpoint to the critical asset and recommend the highest priority containment action",
    },
    "mitre_attack": {
        "order": 2,
        "title": "MITRE ATT&CK Threat Correlation",
        "graph_provides": "Tactic to Technique to Sub Technique to Threat Group to Indicator relationships",
        "ai_does": "Correlates observed telemetry to known adversary techniques and campaigns",
        "description": (
            "This use case represents the MITRE ATT&CK framework itself as a "
            "knowledge graph, connecting tactics to the techniques that "
            "achieve them, techniques to their sub techniques, techniques to "
            "the threat groups known to use them, and threat groups to "
            "observed indicators of compromise. When new telemetry arrives, "
            "the graph is used to instantly match it against known adversary "
            "campaigns. The AI layer explains which campaign the activity most "
            "closely resembles and why."
        ),
        "node_types": ["Tactic", "Technique", "Sub Technique", "Threat Group", "Indicator"],
        "edge_types": [
            ("Tactic", "Technique", "includes"),
            ("Technique", "Sub Technique", "has_variant"),
            ("Technique", "Threat Group", "used_by"),
            ("Threat Group", "Indicator", "associated_with"),
        ],
        "fields": {
            "Tactic": ["tactic_id", "tactic_name", "kill_chain_phase"],
            "Technique": ["technique_id", "technique_name", "platform"],
            "Sub Technique": ["sub_technique_id", "sub_technique_name"],
            "Threat Group": ["group_id", "group_name", "suspected_origin"],
            "Indicator": ["indicator_id", "indicator_type", "first_seen_date"],
        },
        "formulas": [
            ("Technique Coverage Score", "Number of distinct MITRE ATT&CK techniques observed in current telemetry divided by the total number of techniques known for a suspected threat group."),
            ("Campaign Match Confidence", "Weighted overlap between the observed technique sequence and a known threat group technique profile, normalized to zero to one."),
            ("Indicator Corroboration Count", "Number of independent indicators pointing to the same threat group node, used to raise or lower attribution confidence."),
        ],
        "benefits": [
            "Speeds up threat attribution by matching observed behavior against known adversary technique profiles",
            "Reduces analyst workload with automatic technique to campaign correlation instead of manual lookup",
            "Improves detection engineering by highlighting technique coverage gaps in current monitoring",
            "Provides a standardized, explainable common language for reporting to leadership",
        ],
        "ai_prompt_focus": "correlate the observed techniques to the most likely known threat group and explain the confidence level and reasoning",
    },
    "identity_access": {
        "order": 3,
        "title": "Identity and Access Graph Analysis",
        "graph_provides": "User to Role to Group to Permission to System relationships",
        "ai_does": "Surfaces privilege escalation paths and excessive access risk",
        "description": (
            "This use case connects users to the roles they hold, the groups "
            "those roles belong to, the permissions those groups grant, and "
            "the systems those permissions apply to. Privilege escalation "
            "paths often hide behind layers of nested group membership that "
            "are invisible in a flat access control list. The AI layer walks "
            "the graph to find and explain unexpected escalation paths."
        ),
        "node_types": ["User", "Role", "Group", "Permission", "System"],
        "edge_types": [
            ("User", "Role", "assigned"),
            ("Role", "Group", "member_of"),
            ("Group", "Permission", "grants"),
            ("Permission", "System", "applies_to"),
        ],
        "fields": {
            "User": ["user_id", "user_name", "department", "employment_status"],
            "Role": ["role_id", "role_name", "role_tier"],
            "Group": ["group_id", "group_name", "nested_depth"],
            "Permission": ["permission_id", "permission_name", "sensitivity"],
            "System": ["system_id", "system_name", "criticality"],
        },
        "formulas": [
            ("Escalation Path Depth", "Number of nested group hops a user must traverse to reach a high sensitivity permission, shorter depth combined with high sensitivity indicates higher risk."),
            ("Excess Privilege Score", "Number of permissions a user holds that are unused in the last ninety days divided by the total number of permissions the user holds."),
            ("Access Concentration Ratio", "Number of critical systems a single user or role can reach divided by the total number of critical systems in the graph."),
        ],
        "benefits": [
            "Reveals hidden privilege escalation paths created by nested group membership over time",
            "Identifies excess and unused access for least privilege cleanup campaigns",
            "Supports access certification reviews with a visual, explainable access map",
            "Flags concentration risk where too much critical access sits with too few identities",
        ],
        "ai_prompt_focus": "identify the highest risk privilege escalation paths and recommend specific access reductions",
    },
    "threat_intel_fusion": {
        "order": 4,
        "title": "Threat Intelligence Fusion",
        "graph_provides": "Indicator to Campaign to Malware Family to Infrastructure to Threat Actor relationships",
        "ai_does": "Fuses multi source intelligence feeds and reveals shared infrastructure between campaigns",
        "description": (
            "This use case connects indicators of compromise to the campaigns "
            "they were observed in, the malware families used in those "
            "campaigns, the command and control infrastructure that malware "
            "communicates with, and the threat actors believed to operate "
            "that infrastructure. Campaigns that look unrelated in a single "
            "feed often share infrastructure once fused into one graph. The "
            "AI layer explains the fused picture and attribution confidence."
        ),
        "node_types": ["Indicator", "Campaign", "Malware Family", "Infrastructure", "Threat Actor"],
        "edge_types": [
            ("Indicator", "Campaign", "observed_in"),
            ("Campaign", "Malware Family", "uses"),
            ("Malware Family", "Infrastructure", "communicates_with"),
            ("Infrastructure", "Threat Actor", "operated_by"),
        ],
        "fields": {
            "Indicator": ["indicator_id", "indicator_value", "indicator_type", "source_feed"],
            "Campaign": ["campaign_id", "campaign_name", "first_observed_date"],
            "Malware Family": ["malware_id", "malware_name", "malware_category"],
            "Infrastructure": ["infra_id", "infra_type", "hosting_provider"],
            "Threat Actor": ["actor_id", "actor_name", "suspected_motivation"],
        },
        "formulas": [
            ("Infrastructure Overlap Score", "Number of shared infrastructure nodes between two campaign nodes divided by the total distinct infrastructure used by either campaign."),
            ("Campaign Similarity Index", "Weighted combination of shared malware family, shared infrastructure and shared indicator nodes between two campaigns, normalized to zero to one."),
            ("Actor Attribution Confidence", "Number of independent infrastructure nodes linking a campaign to a specific threat actor, divided by the total infrastructure nodes for that campaign."),
        ],
        "benefits": [
            "Reveals shared infrastructure that links campaigns previously treated as unrelated",
            "Improves attribution confidence by fusing evidence across multiple intelligence feeds",
            "Reduces duplicate investigation effort across teams tracking the same underlying actor",
            "Produces an explainable intelligence picture suitable for executive and partner briefings",
        ],
        "ai_prompt_focus": "explain the fused intelligence picture across campaigns and provide an attribution assessment with confidence level",
    },
    "insider_threat": {
        "order": 5,
        "title": "Insider Threat and Anomalous Behavior Detection",
        "graph_provides": "Employee to Access Event to Resource to Communication to Peer Group relationships",
        "ai_does": "Flags deviations from established access and communication baseline patterns",
        "description": (
            "This use case connects employees to the access events they "
            "perform, the resources those events target, the communications "
            "they send, and the peer group they normally behave like. Insider "
            "risk is rarely a single alarming action, it is a pattern that "
            "deviates from an established baseline. The AI layer reviews the "
            "connected behavior graph and explains which deviations are "
            "significant enough to warrant review, always framed as decision "
            "support for a human investigator."
        ),
        "node_types": ["Employee", "Access Event", "Resource", "Communication", "Peer Group"],
        "edge_types": [
            ("Employee", "Access Event", "performs"),
            ("Access Event", "Resource", "targets"),
            ("Employee", "Communication", "sends"),
            ("Employee", "Peer Group", "belongs_to"),
        ],
        "fields": {
            "Employee": ["employee_id", "department", "role_title", "tenure_months"],
            "Access Event": ["event_id", "event_type", "timestamp", "off_hours_flag"],
            "Resource": ["resource_id", "resource_name", "sensitivity"],
            "Communication": ["comm_id", "channel", "recipient_domain"],
            "Peer Group": ["peer_group_id", "peer_group_name", "baseline_profile"],
        },
        "formulas": [
            ("Behavioral Deviation Score", "Absolute difference between an employee current access pattern and their peer group baseline pattern, normalized by the peer group standard deviation."),
            ("Anomalous Access Frequency", "Count of off hours or first time access events to sensitive resources in a rolling thirty day window."),
            ("Peer Group Divergence", "Number of distinct resources accessed by an employee that no other member of their peer group has accessed in the same window."),
        ],
        "benefits": [
            "Surfaces meaningful behavioral deviation instead of flooding analysts with every unusual event",
            "Uses peer group context so normal role based variation is not mistaken for risk",
            "Supports fair, explainable review that focuses on pattern, not a single incident",
            "Reduces false positive investigation load compared to static threshold based rules",
        ],
        "ai_prompt_focus": "summarize the behavioral deviations for human review and explain why they differ from the peer group baseline, this is decision support not a determination of guilt",
    },
    "software_supply_chain": {
        "order": 6,
        "title": "Supply Chain and Software Bill of Materials Risk",
        "graph_provides": "Application to API to Library to Vulnerability to Vendor relationships",
        "ai_does": "Determines vulnerability blast radius across the software supply chain",
        "description": (
            "This use case maps which applications call which APIs, which "
            "are built on which open source or commercial libraries, which "
            "libraries carry known vulnerabilities, and which vendor supplies "
            "each library. When a new vulnerability is disclosed, the graph "
            "instantly shows every affected application and the vendor "
            "responsible. The AI layer explains blast radius and a "
            "remediation order weighted by mission criticality."
        ),
        "node_types": ["Application", "API", "Library", "Vulnerability", "Vendor"],
        "edge_types": [
            ("Application", "API", "uses"),
            ("API", "Library", "built_on"),
            ("Library", "Vulnerability", "affected_by"),
            ("Library", "Vendor", "supplied_by"),
        ],
        "fields": {
            "Application": ["app_id", "app_name", "mission_criticality"],
            "API": ["api_id", "api_name", "version"],
            "Library": ["library_id", "library_name", "library_version"],
            "Vulnerability": ["cve_id", "severity", "cvss_score"],
            "Vendor": ["vendor_id", "vendor_name", "country_of_origin"],
        },
        "formulas": [
            ("Blast Radius", "Number of distinct application nodes reachable upstream from a vulnerable library node."),
            ("Vendor Concentration Risk", "Number of distinct applications depending on libraries from a single vendor divided by the total number of applications in the graph."),
            ("Patch Debt Score", "Number of days since the vulnerability was published multiplied by the CVSS score, summed across every affected application."),
        ],
        "benefits": [
            "Answers which mission systems are affected by a newly disclosed vulnerability in seconds",
            "Reveals vendor concentration risk that a single supplier compromise could exploit widely",
            "Prioritizes patching using both severity and mission criticality, not severity alone",
            "Supports supply chain risk management reporting with a defensible, explainable trace",
        ],
        "ai_prompt_focus": "explain the vulnerability blast radius across mission applications and recommend a remediation order",
    },
    "intel_investigation": {
        "order": 7,
        "title": "Intelligence Analysis and Connected Data Investigation",
        "graph_provides": "Person to Organization to Location to Event to Communication relationships",
        "ai_does": "Fuses fragmented multi source intelligence into one connected investigative picture",
        "description": (
            "This use case connects persons of interest to the organizations "
            "they are affiliated with, the locations where they have been "
            "observed, the events they participated in, and the "
            "communications they exchanged. Intelligence analysts use this "
            "connected view to investigate relationships across fragmented, "
            "multi source reporting. The AI layer summarizes the network and "
            "highlights the highest priority connections for investigation."
        ),
        "node_types": ["Person", "Organization", "Location", "Event", "Communication"],
        "edge_types": [
            ("Person", "Organization", "affiliated_with"),
            ("Person", "Location", "observed_at"),
            ("Person", "Event", "participated_in"),
            ("Person", "Communication", "exchanged"),
        ],
        "fields": {
            "Person": ["person_id", "known_alias", "nationality", "risk_tier"],
            "Organization": ["org_id", "org_name", "org_type"],
            "Location": ["location_id", "location_name", "region"],
            "Event": ["event_id", "event_type", "event_date"],
            "Communication": ["comm_id", "channel", "date"],
        },
        "formulas": [
            ("Network Centrality Score", "Combination of degree, betweenness and pagerank centrality for a person node, identifying key facilitators in the network."),
            ("Link Confidence Score", "Number of independent source reports supporting a given relationship edge, divided by the total number of source reports referencing either node."),
            ("Investigation Priority Score", "Weighted combination of network centrality, risk tier, and recency of activity, used to rank persons of interest for investigative follow up."),
        ],
        "benefits": [
            "Fuses fragmented, multi source reporting into a single connected investigative picture",
            "Identifies key facilitators and hubs in a network rather than isolated individuals",
            "Speeds up analyst workflow with a visual, explainable relationship map",
            "Supports collaborative investigation across teams working different pieces of the same network",
        ],
        "ai_prompt_focus": "summarize the connected network, identify the highest priority person for follow up investigation, and explain the reasoning",
    },
    "mission_readiness": {
        "order": 8,
        "title": "Mission and Force Readiness Planning",
        "graph_provides": "Mission to Unit to Asset to Location to Capability relationships",
        "ai_does": "Identifies capability gaps and readiness risk for mission execution",
        "description": (
            "This use case connects missions to the units assigned, the "
            "assets those units operate, the locations involved, and the "
            "capabilities each asset provides. Mission planners need to "
            "understand quickly whether the required capability mix is "
            "available at the right location. The AI layer reviews the "
            "graph and explains capability gaps and readiness risk for a "
            "given mission."
        ),
        "node_types": ["Mission", "Unit", "Asset", "Location", "Capability"],
        "edge_types": [
            ("Mission", "Unit", "assigned_to"),
            ("Unit", "Asset", "operates"),
            ("Asset", "Location", "positioned_at"),
            ("Asset", "Capability", "provides"),
        ],
        "fields": {
            "Mission": ["mission_id", "mission_name", "objective", "start_date"],
            "Unit": ["unit_id", "unit_name", "readiness_level"],
            "Asset": ["asset_id", "asset_type", "status"],
            "Location": ["location_id", "location_name", "region"],
            "Capability": ["capability_id", "capability_name", "required_level"],
        },
        "formulas": [
            ("Capability Coverage", "Number of required mission capabilities present among assigned assets divided by the total number of required capabilities."),
            ("Readiness Score", "Average readiness level across all units assigned to a mission, weighted by each unit share of assigned assets."),
            ("Positional Risk", "Count of required assets not yet positioned at the mission location divided by the total number of required assets."),
        ],
        "benefits": [
            "Speeds up mission planning by showing capability gaps at a glance",
            "Improves readiness assessment with a connected unit and asset view",
            "Highlights positional risk before it becomes an operational problem",
            "Supports after action review with a clear relationship based record",
        ],
        "ai_prompt_focus": "identify capability gaps and readiness risk for the mission and explain the reasoning",
    },
    "logistics_resilience": {
        "order": 9,
        "title": "Logistics and Defense Supply Chain Resilience",
        "graph_provides": "Supplier to Component to Depot to Base to Mission relationships",
        "ai_does": "Identifies cascading logistics risk and single points of failure across the defense industrial base",
        "description": (
            "This use case connects suppliers to the components they "
            "provide, the depots that stock those components, the bases "
            "those depots supply, and the missions those bases support. A "
            "disruption at a single supplier can cascade through depots and "
            "bases to threaten mission readiness weeks later. The AI layer "
            "explains the cascading impact and recommends mitigation."
        ),
        "node_types": ["Supplier", "Component", "Depot", "Base", "Mission"],
        "edge_types": [
            ("Supplier", "Component", "provides"),
            ("Component", "Depot", "stocked_at"),
            ("Depot", "Base", "supplies"),
            ("Base", "Mission", "supports"),
        ],
        "fields": {
            "Supplier": ["supplier_id", "supplier_name", "country_of_origin", "reliability_score"],
            "Component": ["component_id", "component_name", "lead_time_days"],
            "Depot": ["depot_id", "depot_name", "current_stock_level"],
            "Base": ["base_id", "base_name", "region"],
            "Mission": ["mission_id", "mission_name", "priority_level"],
        },
        "formulas": [
            ("Single Source Risk", "One divided by the number of distinct suppliers that provide the same component, a value of one means there is no backup supplier."),
            ("Cascading Impact Score", "Number of downstream mission nodes reachable from a disrupted supplier node, divided by the total number of mission nodes in the graph."),
            ("Depot Coverage Ratio", "Current stock level at a depot divided by the average consumption rate of the bases it supplies, expressed in days of coverage remaining."),
        ],
        "benefits": [
            "Reveals single points of failure across multi tier defense supplier networks",
            "Quantifies which missions are most exposed to a given supplier or depot disruption",
            "Shortens response time to supply shocks with pre identified alternative sourcing options",
            "Supports strategic stockpile and depot investment decisions with data instead of guesswork",
        ],
        "ai_prompt_focus": "explain the cascading logistics impact of the disruption and recommend a mitigation priority",
    },
    "counter_terror_financing": {
        "order": 10,
        "title": "Counter Terrorism Financing and Sanctions Network Analysis",
        "graph_provides": "Person to Account to Business to Transaction to Sanctioned Entity relationships",
        "ai_does": "Detects financing networks connected to sanctioned or hostile actors",
        "description": (
            "This use case connects persons to the accounts and businesses "
            "they control, the transactions they make, and any known "
            "sanctioned entities linked to those businesses. Financing "
            "networks that support hostile activity typically move funds "
            "through layers of accounts and shell businesses. The AI layer "
            "reviews the connected transaction network and explains the "
            "path of funds toward a sanctioned entity."
        ),
        "node_types": ["Person", "Account", "Business", "Transaction", "Sanctioned Entity"],
        "edge_types": [
            ("Person", "Account", "controls"),
            ("Account", "Transaction", "sends_or_receives"),
            ("Person", "Business", "owns"),
            ("Business", "Sanctioned Entity", "linked_to"),
        ],
        "fields": {
            "Person": ["person_id", "full_name", "nationality", "watchlist_status"],
            "Account": ["account_id", "account_type", "opened_date"],
            "Business": ["business_id", "business_name", "industry_code"],
            "Transaction": ["transaction_id", "amount", "currency", "date"],
            "Sanctioned Entity": ["entity_id", "entity_name", "sanction_program"],
        },
        "formulas": [
            ("Network Proximity to Sanctioned Entity", "Number of hops in the shortest path between a person node and the nearest sanctioned entity node, shorter distance indicates higher exposure risk."),
            ("Structuring Score", "Count of transactions just below a reporting threshold within a rolling seven day window for a given account."),
            ("Shell Company Density", "Number of businesses sharing the same registered address or beneficial owner divided by the expected average for that jurisdiction."),
        ],
        "benefits": [
            "Detects financing networks that single transaction monitoring alone would miss",
            "Quantifies exposure distance from any person or business to a known sanctioned entity",
            "Identifies shell company clusters used to obscure the true source of funds",
            "Produces an explainable case file suitable for a suspicious activity report or referral",
        ],
        "ai_prompt_focus": "explain the financing network structure, the path of funds toward the sanctioned entity, and the key facilitators involved",
    },
    "personnel_security": {
        "order": 11,
        "title": "Personnel Security and Counter Intelligence",
        "graph_provides": "Person to Foreign Contact to Financial Disclosure to Travel Record to Clearance Level relationships",
        "ai_does": "Flags patterns consistent with insider threat or foreign influence risk",
        "description": (
            "This use case connects cleared personnel to their reported "
            "foreign contacts, financial disclosures, travel records, and "
            "the clearance level they hold. Foreign influence risk rarely "
            "shows up in a single data point, it emerges from a pattern "
            "across contacts, unexplained finances, and travel. The AI "
            "layer reviews the connected record and flags patterns for "
            "human security officer review, always framed as decision "
            "support rather than a determination of wrongdoing."
        ),
        "node_types": ["Person", "Foreign Contact", "Financial Disclosure", "Travel Record", "Clearance Level"],
        "edge_types": [
            ("Person", "Foreign Contact", "has_contact_with"),
            ("Person", "Financial Disclosure", "filed"),
            ("Person", "Travel Record", "traveled_on"),
            ("Person", "Clearance Level", "holds"),
        ],
        "fields": {
            "Person": ["person_id", "role_title", "department", "years_cleared"],
            "Foreign Contact": ["contact_id", "contact_country", "relationship_type"],
            "Financial Disclosure": ["disclosure_id", "disclosure_date", "asset_value_change"],
            "Travel Record": ["travel_id", "destination_country", "travel_date", "purpose"],
            "Clearance Level": ["clearance_id", "clearance_name", "granted_date"],
        },
        "formulas": [
            ("Foreign Influence Risk Score", "Weighted count of foreign contacts in countries of concern, combined with undisclosed or late reported contacts, normalized to zero to one."),
            ("Disclosure Anomaly Score", "Magnitude of unexplained asset value change in a financial disclosure relative to the person reported income."),
            ("Travel Pattern Risk", "Count of undisclosed or unusually frequent travel events to countries of concern within a rolling twelve month window."),
        ],
        "benefits": [
            "Surfaces patterns across contacts, finances and travel that are invisible when reviewed separately",
            "Supports continuous evaluation programs with an explainable, defensible risk indicator",
            "Reduces manual cross referencing effort across multiple personnel security data systems",
            "Frames results as decision support, keeping the human security officer as final adjudicator",
        ],
        "ai_prompt_focus": "summarize the connected pattern across contacts, disclosures and travel for personnel security officer review, this is decision support not an adjudication",
    },
    "digital_twin_infra": {
        "order": 12,
        "title": "Digital Twin of Critical Infrastructure",
        "graph_provides": "Asset to Sensor to Component to Process to Failure Event relationships",
        "ai_does": "Predicts failures and explains root cause across critical infrastructure assets",
        "description": (
            "This use case builds a live connected model of critical "
            "infrastructure, such as a base power grid or a weapons "
            "platform, including its sensors, components, the process steps "
            "it participates in, and its historical failure events. When a "
            "sensor reading drifts, the graph shows exactly which downstream "
            "components and processes are affected. The AI layer explains "
            "the likely root cause and expected time to failure."
        ),
        "node_types": ["Asset", "Sensor", "Component", "Process", "Failure Event"],
        "edge_types": [
            ("Asset", "Sensor", "monitored_by"),
            ("Sensor", "Component", "measures"),
            ("Component", "Process", "participates_in"),
            ("Component", "Failure Event", "experienced"),
        ],
        "fields": {
            "Asset": ["asset_id", "asset_name", "install_date", "mission_criticality"],
            "Sensor": ["sensor_id", "sensor_type", "unit_of_measure"],
            "Component": ["component_id", "component_name", "expected_life_hours"],
            "Process": ["process_id", "process_name", "throughput_rate"],
            "Failure Event": ["failure_id", "failure_type", "failure_date", "downtime_hours"],
        },
        "formulas": [
            ("Remaining Useful Life", "Expected life hours of a component minus cumulative operating hours since install, adjusted downward by the deviation of recent sensor readings from baseline."),
            ("Failure Propagation Score", "Number of downstream process nodes reachable from a failing component divided by total process nodes in the graph."),
            ("Anomaly Deviation", "Absolute difference between current sensor reading and the rolling average, divided by the rolling standard deviation."),
        ],
        "benefits": [
            "Predicts component failure before it causes unplanned downtime on mission critical infrastructure",
            "Explains root cause across the full asset to process chain, not just a single sensor alert",
            "Prioritizes maintenance based on downstream mission process impact",
            "Reduces inspection cost by focusing attention on high deviation components",
        ],
        "ai_prompt_focus": "explain the likely failure cause, downstream mission impact, and recommended maintenance action",
    },
    "data_sovereignty": {
        "order": 13,
        "title": "Data Sovereignty and Deployment Risk Assessment",
        "graph_provides": "Data Asset to Classification Level to Hosting Environment to Jurisdiction to Compliance Requirement relationships",
        "ai_does": "Assesses sovereignty exposure and recommends compliant deployment adjustments",
        "description": (
            "This use case connects data assets to their classification "
            "level, the hosting environment they run on, the jurisdiction "
            "that environment sits in, and the compliance requirements that "
            "jurisdiction imposes. Government agencies increasingly require "
            "sovereign, on premise or air gapped deployment for sensitive "
            "data. The AI layer reviews the graph and flags data assets "
            "whose current hosting arrangement does not meet the compliance "
            "requirement for their classification level."
        ),
        "node_types": ["Data Asset", "Classification Level", "Hosting Environment", "Jurisdiction", "Compliance Requirement"],
        "edge_types": [
            ("Data Asset", "Classification Level", "classified_as"),
            ("Data Asset", "Hosting Environment", "hosted_in"),
            ("Hosting Environment", "Jurisdiction", "located_in"),
            ("Jurisdiction", "Compliance Requirement", "subject_to"),
        ],
        "fields": {
            "Data Asset": ["data_asset_id", "data_asset_name", "owner_department"],
            "Classification Level": ["classification_id", "classification_name", "required_control_tier"],
            "Hosting Environment": ["hosting_id", "hosting_type", "provider"],
            "Jurisdiction": ["jurisdiction_id", "jurisdiction_name", "region"],
            "Compliance Requirement": ["requirement_id", "requirement_name", "mandatory_flag"],
        },
        "formulas": [
            ("Sovereignty Exposure Score", "Count of data asset nodes whose classification level requires a control tier higher than what their current hosting environment provides."),
            ("Jurisdiction Risk Rating", "Weighted count of data assets hosted in jurisdictions flagged as higher risk for data residency, based on the compliance requirement mandatory flag."),
            ("Compliance Gap Count", "Number of mandatory compliance requirements for a jurisdiction that are not currently met by the hosting environment configuration."),
        ],
        "benefits": [
            "Identifies exactly which data assets are hosted in a non compliant or higher risk jurisdiction",
            "Supports sovereign and air gapped deployment planning with a clear, prioritized gap list",
            "Reduces manual cross referencing between data classification and hosting inventories",
            "Provides an explainable, audit ready trail from data asset down to compliance requirement",
        ],
        "ai_prompt_focus": "identify data sovereignty exposure and recommend deployment adjustments to close compliance gaps",
    },
}
USE_CASE_ORDER = [k for k, v in sorted(USE_CASES.items(), key=lambda kv: kv[1]["order"])]
# =============================================================================
# SECTION 2: SYNTHETIC DATA GENERATION AND GRAPH CONSTRUCTION
# =============================================================================
FIRST_NAMES = ["James", "Maria", "Wei", "Amara", "Liam", "Sofia", "Noah", "Priya",
               "Ethan", "Fatima", "Lucas", "Ingrid", "Omar", "Chen", "Ava", "Diego"]
LAST_NAMES = ["Smith", "Garcia", "Chen", "Okafor", "Muller", "Rossi", "Kim", "Singh",
              "Brown", "Nguyen", "Andersson", "Silva", "Khan", "Kowalski", "Dubois"]
CITIES = ["New York", "London", "Singapore", "Toronto", "Mumbai", "Berlin",
          "Sydney", "Sao Paulo", "Dubai", "Tokyo", "Chicago", "Nairobi"]
COUNTRIES = ["United States", "United Kingdom", "Singapore", "Canada", "India",
             "Germany", "Australia", "Brazil", "United Arab Emirates", "Japan"]
def random_name():
    return random.choice(FIRST_NAMES) + " " + random.choice(LAST_NAMES)
def random_id(prefix, n=6):
    suffix = "".join(random.choices(string.digits, k=n))
    return prefix + suffix
def random_field_value(field_name, node_type, index):
    name_lower = field_name.lower()
    if name_lower.endswith("_id"):
        return random_id(node_type[:2].upper(), 5) if index is None else node_type[:2].upper() + str(1000 + index)
    if "name" in name_lower and "user" not in name_lower:
        return random_name()
    if "user_name" in name_lower or "hostname" in name_lower:
        return random.choice(["host", "user", "node"]) + str(random.randint(100, 999))
    if "score" in name_lower or "rating" in name_lower or "level" in name_lower or "pct" in name_lower:
        return round(random.uniform(0, 100), 2)
    if "amount" in name_lower or "value" in name_lower or "balance" in name_lower or "capacity" in name_lower:
        return round(random.uniform(100, 100000), 2)
    if "date" in name_lower or "timestamp" in name_lower:
        y = random.randint(2022, 2026)
        m = random.randint(1, 12)
        d = random.randint(1, 28)
        return "%04d-%02d-%02d" % (y, m, d)
    if "city" in name_lower:
        return random.choice(CITIES)
    if "country" in name_lower or "region" in name_lower or "jurisdiction" in name_lower or "origin" in name_lower:
        return random.choice(COUNTRIES)
    if "status" in name_lower or "flag" in name_lower:
        return random.choice(["Active", "Inactive", "Pending", "Resolved"])
    if "severity" in name_lower:
        return random.choice(["Low", "Medium", "High", "Critical"])
    if "type" in name_lower or "category" in name_lower or "channel" in name_lower or "tier" in name_lower or "segment" in name_lower:
        return random.choice(["Type A", "Type B", "Type C", "Type D"])
    if "days" in name_lower or "hours" in name_lower or "count" in name_lower or "quantity" in name_lower:
        return random.randint(1, 500)
    return "Value " + str(random.randint(1, 999))
def generate_synthetic_dataset(use_case_key, config, nodes_per_type=12, seed=None):
    """
    Generates a synthetic node table and edge table for the given use case
    configuration. Returns a dict with keys nodes (DataFrame) and edges (DataFrame).
    """
    if seed is not None:
        random.seed(seed)
    node_types = config["node_types"]
    fields = config["fields"]
    edge_types = config["edge_types"]
    node_rows = []
    node_ids_by_type = {}
    for ntype in node_types:
        ids_for_type = []
        for i in range(nodes_per_type):
            row = {"node_type": ntype, "node_label": ntype + " " + str(i + 1)}
            for f in fields.get(ntype, []):
                row[f] = random_field_value(f, ntype, i)
            primary_key = fields.get(ntype, ["id"])[0]
            node_id = ntype + "_" + str(i + 1)
            row["node_id"] = node_id
            row[primary_key] = row.get(primary_key, node_id)
            node_rows.append(row)
            ids_for_type.append(node_id)
        node_ids_by_type[ntype] = ids_for_type
    nodes_df = pd.DataFrame(node_rows)
    edge_rows = []
    for source_type, target_type, relation in edge_types:
        source_ids = node_ids_by_type.get(source_type, [])
        target_ids = node_ids_by_type.get(target_type, [])
        if not source_ids or not target_ids:
            continue
        for sid in source_ids:
            num_links = random.randint(1, min(3, len(target_ids)))
            targets = random.sample(target_ids, num_links)
            for tid in targets:
                edge_rows.append({
                    "source_id": sid,
                    "source_type": source_type,
                    "target_id": tid,
                    "target_type": target_type,
                    "relationship": relation,
                    "weight": round(random.uniform(0.1, 1.0), 3),
                })
    edges_df = pd.DataFrame(edge_rows)
    return {"nodes": nodes_df, "edges": edges_df}
def build_networkx_graph(nodes_df, edges_df):
    G = nx.DiGraph()
    if nodes_df is not None and len(nodes_df) > 0:
        for _, row in nodes_df.iterrows():
            node_id = row.get("node_id", row.iloc[0])
            label = row.get("node_label", str(node_id))
            ntype = row.get("node_type", "Entity")
            extra_attrs = row.to_dict()
            extra_attrs.pop("node_id", None)
            extra_attrs.pop("node_label", None)
            extra_attrs.pop("node_type", None)
            G.add_node(node_id, label=label, node_type=ntype, **extra_attrs)
    if edges_df is not None and len(edges_df) > 0:
        for _, row in edges_df.iterrows():
            src = row.get("source_id")
            tgt = row.get("target_id")
            if src is None or tgt is None:
                continue
            if src not in G:
                G.add_node(src, label=str(src), node_type=row.get("source_type", "Entity"))
            if tgt not in G:
                G.add_node(tgt, label=str(tgt), node_type=row.get("target_type", "Entity"))
            rel = row.get("relationship", "related_to")
            weight = row.get("weight", 1.0)
            G.add_edge(src, tgt, relationship=rel, weight=weight)
    return G
def dataframes_from_uploaded(nodes_file, edges_file):
    """
    Reads user uploaded CSV files for nodes and edges into DataFrames.
    Expected node columns: node_id, node_label, node_type, plus any custom fields.
    Expected edge columns: source_id, target_id, relationship, and optional weight.
    """
    nodes_df = pd.read_csv(nodes_file)
    edges_df = pd.read_csv(edges_file)
    if "node_id" not in nodes_df.columns:
        nodes_df["node_id"] = nodes_df.index.astype(str)
    if "node_label" not in nodes_df.columns:
        nodes_df["node_label"] = nodes_df["node_id"]
    if "node_type" not in nodes_df.columns:
        nodes_df["node_type"] = "Entity"
    if "weight" not in edges_df.columns:
        edges_df["weight"] = 1.0
    return {"nodes": nodes_df, "edges": edges_df}
# =============================================================================
# SECTION 3: GRAPH ANALYTICS, VISUALIZATION AND SCHEMA FLOW DIAGRAM
# =============================================================================
def compute_core_metrics(G):
    """
    Computes degree centrality, betweenness centrality and pagerank for
    every node in the graph. Returns a DataFrame indexed by node id.
    """
    if G.number_of_nodes() == 0:
        return pd.DataFrame(columns=["node_id", "label", "node_type",
                                      "degree_centrality", "betweenness_centrality", "pagerank"])
    degree = nx.degree_centrality(G)
    try:
        betweenness = nx.betweenness_centrality(G)
    except Exception:
        betweenness = {n: 0.0 for n in G.nodes()}
    try:
        pagerank = nx.pagerank(G, weight="weight")
    except Exception:
        pagerank = {n: 0.0 for n in G.nodes()}
    rows = []
    for n in G.nodes():
        data = G.nodes[n]
        rows.append({
            "node_id": n,
            "label": data.get("label", str(n)),
            "node_type": data.get("node_type", "Entity"),
            "degree_centrality": round(degree.get(n, 0.0), 4),
            "betweenness_centrality": round(betweenness.get(n, 0.0), 4),
            "pagerank": round(pagerank.get(n, 0.0), 4),
        })
    return pd.DataFrame(rows).sort_values("pagerank", ascending=False).reset_index(drop=True)
def find_top_paths(G, max_paths=5):
    """
    Finds a handful of representative shortest paths between nodes of
    different types that have low out degree and high in degree, useful
    for narrating attack paths, root cause chains or cascading impact.
    """
    paths = []
    nodes = list(G.nodes())
    if len(nodes) < 2:
        return paths
    sources = [n for n in nodes if G.out_degree(n) > 0 and G.in_degree(n) == 0]
    sinks = [n for n in nodes if G.in_degree(n) > 0 and G.out_degree(n) == 0]
    if not sources:
        sources = nodes[: min(5, len(nodes))]
    if not sinks:
        sinks = nodes[: min(5, len(nodes))]
    count = 0
    for s in sources:
        for t in sinks:
            if s == t:
                continue
            try:
                if nx.has_path(G, s, t):
                    p = nx.shortest_path(G, s, t)
                    if len(p) > 1:
                        paths.append(p)
                        count += 1
            except Exception:
                continue
            if count >= max_paths:
                return paths
    return paths
def build_plotly_graph(G, highlight_path=None, max_nodes=250):
    """
    Builds an interactive Plotly figure of the graph using a spring layout.
    Nodes are colored by node_type. If highlight_path is provided, those
    edges are drawn in a distinct color and thicker width.
    """
    if G.number_of_nodes() == 0:
        fig = go.Figure()
        fig.update_layout(title="No graph data available")
        return fig
    if G.number_of_nodes() > max_nodes:
        keep = list(G.nodes())[:max_nodes]
        G = G.subgraph(keep).copy()
    pos = nx.spring_layout(G, seed=42, k=None)
    node_types = sorted(set(nx.get_node_attributes(G, "node_type").values()))
    palette = ["#0b3d91", "#c0392b", "#1e8449", "#b9770e", "#6c3483",
               "#117864", "#a04000", "#2874a6", "#7d3c98", "#154360"]
    color_map = {nt: palette[i % len(palette)] for i, nt in enumerate(node_types)}
    highlight_edges = set()
    if highlight_path:
        for i in range(len(highlight_path) - 1):
            highlight_edges.add((highlight_path[i], highlight_path[i + 1]))
    edge_traces = []
    edge_x_normal, edge_y_normal = [], []
    edge_x_highlight, edge_y_highlight = [], []
    for u, v in G.edges():
        x0, y0 = pos[u]
        x1, y1 = pos[v]
        if (u, v) in highlight_edges:
            edge_x_highlight += [x0, x1, None]
            edge_y_highlight += [y0, y1, None]
        else:
            edge_x_normal += [x0, x1, None]
            edge_y_normal += [y0, y1, None]
    edge_traces.append(go.Scatter(x=edge_x_normal, y=edge_y_normal, mode="lines",
                                   line=dict(width=1, color="#bbbbbb"),
                                   hoverinfo="none", showlegend=False))
    if edge_x_highlight:
        edge_traces.append(go.Scatter(x=edge_x_highlight, y=edge_y_highlight, mode="lines",
                                       line=dict(width=3, color="#e74c3c"),
                                       hoverinfo="none", showlegend=False,
                                       name="Highlighted Path"))
    node_traces = []
    for nt in node_types:
        xs, ys, texts = [], [], []
        for n in G.nodes():
            if G.nodes[n].get("node_type", "Entity") == nt:
                x, y = pos[n]
                xs.append(x)
                ys.append(y)
                label = G.nodes[n].get("label", str(n))
                texts.append(label + " (" + nt + ")")
        node_traces.append(go.Scatter(
            x=xs, y=ys, mode="markers", name=nt,
            marker=dict(size=14, color=color_map[nt], line=dict(width=1, color="#ffffff")),
            text=texts, hoverinfo="text"
        ))
    fig = go.Figure(data=edge_traces + node_traces)
    fig.update_layout(
        showlegend=True,
        margin=dict(l=10, r=10, t=30, b=10),
        xaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
        yaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
        plot_bgcolor="#ffffff",
        legend=dict(orientation="h", yanchor="bottom", y=1.02),
        height=560,
    )
    return fig
def build_schema_flow_diagram(config):
    """
    Builds a Graphviz DOT format flow diagram string describing how the
    entity types in a use case schema connect to one another. Rendered
    with Streamlit st.graphviz_chart in the application.
    """
    lines = ["digraph SchemaFlow {"]
    lines.append('rankdir=LR;')
    lines.append('bgcolor="transparent";')
    lines.append('node [shape=box, style="rounded,filled", fillcolor="#eaf1fb", '
                  'color="#0b3d91", fontname="Helvetica", fontsize=11, fontcolor="#0b3d91", penwidth=2];')
    lines.append('edge [color="#5d6d7e", fontname="Helvetica", fontsize=9, fontcolor="#333333"];')
    for source_type, target_type, relation in config["edge_types"]:
        source_id = source_type.replace(" ", "_")
        target_id = target_type.replace(" ", "_")
        lines.append('"' + source_id + '" [label="' + source_type + '"];')
        lines.append('"' + target_id + '" [label="' + target_type + '"];')
        lines.append('"' + source_id + '" -> "' + target_id + '" [label="' + relation + '"];')
    lines.append("}")
    return "\n".join(lines)
def build_record_count_chart(nodes_df):
    """
    Builds a Plotly bar chart showing the number of synthetic or uploaded
    records generated for each entity type, used as the synthetic data
    summary bar in the application.
    """
    if nodes_df is None or len(nodes_df) == 0 or "node_type" not in nodes_df.columns:
        fig = go.Figure()
        fig.update_layout(title="No data loaded yet")
        return fig
    counts = nodes_df["node_type"].value_counts().reset_index()
    counts.columns = ["node_type", "record_count"]
    fig = go.Figure(data=[
        go.Bar(
            x=counts["node_type"], y=counts["record_count"],
            marker_color="#0b3d91", text=counts["record_count"], textposition="outside",
        )
    ])
    fig.update_layout(
        title="Records Loaded by Entity Type",
        xaxis_title="Entity Type",
        yaxis_title="Record Count",
        margin=dict(l=10, r=10, t=40, b=10),
        height=320,
        plot_bgcolor="#ffffff",
    )
    return fig
# =============================================================================
# SECTION 4: LLM INTEGRATION - GROQ AND GEMINI
# =============================================================================
GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"
GROQ_MODELS_URL = "https://api.groq.com/openai/v1/models"
# Groq periodically retires models. Rather than hard code a single model
# name that eventually returns HTTP 404 model_not_found, keep an ordered
# list of reasonable current candidates and let the caller override with
# whatever the live /models endpoint reports as actually available on
# the caller's account.
GROQ_DEFAULT_MODEL = "llama-3.3-70b-versatile"
GROQ_FALLBACK_MODELS = [
    "llama-3.3-70b-versatile",
    "llama-3.1-8b-instant",
    "openai/gpt-oss-120b",
    "openai/gpt-oss-20b",
    "groq/compound-mini",
]
GEMINI_API_URL_TEMPLATE = "https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={key}"
GEMINI_MODELS_URL_TEMPLATE = "https://generativelanguage.googleapis.com/v1beta/models?key={key}"
# Google periodically retires dated Gemini snapshots (for example
# gemini-2.0-flash) in favor of newer ones. "gemini-flash-latest" is a
# rolling alias Google maintains that always resolves to the current
# recommended fast Gemini model, so it does not go stale the way a
# pinned dated model name does.
GEMINI_DEFAULT_MODEL = "gemini-flash-latest"
GEMINI_FALLBACK_MODELS = [
    "gemini-flash-latest",
    "gemini-pro-latest",
    "gemini-2.5-flash",
    "gemini-2.0-flash",
]
def list_groq_models(api_key, timeout=20):
    """
    Queries Groq's /models endpoint to discover which model ids are
    actually available for this API key right now. Returns a list of
    model id strings, or an empty list if the call fails.
    """
    if not api_key:
        return []
    headers = {"Authorization": "Bearer " + api_key}
    try:
        resp = requests.get(GROQ_MODELS_URL, headers=headers, timeout=timeout)
        if resp.status_code != 200:
            return []
        data = resp.json()
        ids = [m.get("id") for m in data.get("data", []) if m.get("id")]
        # Prefer chat capable text models over whisper or guard models.
        ids = [i for i in ids if "whisper" not in i.lower() and "guard" not in i.lower()
               and "tts" not in i.lower()]
        return sorted(ids)
    except Exception:
        return []
def list_gemini_models(api_key, timeout=20):
    """
    Queries Gemini's ListModels endpoint to discover which model names
    support generateContent for this API key right now. Returns a list
    of model id strings (without the leading "models/" prefix), or an
    empty list if the call fails.
    """
    if not api_key:
        return []
    url = GEMINI_MODELS_URL_TEMPLATE.format(key=api_key)
    try:
        resp = requests.get(url, timeout=timeout)
        if resp.status_code != 200:
            return []
        data = resp.json()
        names = []
        for m in data.get("models", []):
            methods = m.get("supportedGenerationMethods", [])
            if "generateContent" in methods:
                name = m.get("name", "")
                if name.startswith("models/"):
                    name = name[len("models/"):]
                if name:
                    names.append(name)
        return sorted(names)
    except Exception:
        return []
def _is_model_not_found(status_code, response_text):
    if status_code == 404:
        return True
    lowered = (response_text or "").lower()
    return "model_not_found" in lowered or "does not exist" in lowered or "not found" in lowered
def call_groq(api_key, system_prompt, user_prompt, model=None, timeout=60):
    if not api_key:
        return None, "Groq API key is missing. Please enter a key in the sidebar."
    candidates = []
    if model:
        candidates.append(model)
    for m in GROQ_FALLBACK_MODELS:
        if m not in candidates:
            candidates.append(m)
    headers = {
        "Authorization": "Bearer " + api_key,
        "Content-Type": "application/json",
    }
    last_error = None
    for model_name in candidates:
        payload = {
            "model": model_name,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            "temperature": 0.3,
            "max_tokens": 900,
        }
        try:
            resp = requests.post(GROQ_API_URL, headers=headers, data=json.dumps(payload), timeout=timeout)
            if resp.status_code != 200:
                last_error = ("Groq API error for model " + model_name + ", status code "
                              + str(resp.status_code) + ". Details: " + resp.text[:300])
                if _is_model_not_found(resp.status_code, resp.text):
                    # This model was retired or unavailable to this key, try the next candidate.
                    continue
                return None, last_error
            data = resp.json()
            text = data["choices"][0]["message"]["content"]
            return text, None
        except Exception as e:
            last_error = "Groq API request failed for model " + model_name + ": " + str(e)
            continue
    return None, last_error or "Groq API request failed for all candidate models."
def call_gemini(api_key, system_prompt, user_prompt, model=None, timeout=60):
    if not api_key:
        return None, "Gemini API key is missing. Please enter a key in the sidebar."
    candidates = []
    if model:
        candidates.append(model)
    for m in GEMINI_FALLBACK_MODELS:
        if m not in candidates:
            candidates.append(m)
    combined_prompt = system_prompt + "\n\n" + user_prompt
    payload = {
        "contents": [
            {"parts": [{"text": combined_prompt}]}
        ],
        "generationConfig": {"temperature": 0.3, "maxOutputTokens": 900},
    }
    headers = {"Content-Type": "application/json"}
    last_error = None
    for model_name in candidates:
        url = GEMINI_API_URL_TEMPLATE.format(model=model_name, key=api_key)
        try:
            resp = requests.post(url, headers=headers, data=json.dumps(payload), timeout=timeout)
            if resp.status_code != 200:
                last_error = ("Gemini API error for model " + model_name + ", status code "
                              + str(resp.status_code) + ". Details: " + resp.text[:300])
                if _is_model_not_found(resp.status_code, resp.text):
                    continue
                return None, last_error
            data = resp.json()
            candidates_resp = data.get("candidates", [])
            if not candidates_resp:
                last_error = "Gemini API returned no candidates for model " + model_name + "."
                continue
            parts = candidates_resp[0]["content"]["parts"]
            text = "".join([p.get("text", "") for p in parts])
            return text, None
        except Exception as e:
            last_error = "Gemini API request failed for model " + model_name + ": " + str(e)
            continue
    return None, last_error or "Gemini API request failed for all candidate models."
def run_llm_analysis(provider, api_key, system_prompt, user_prompt, model=None):
    """
    Dispatches to the selected provider. provider is either Groq or Gemini.
    Returns a tuple of result text and error message. If the call fails or
    no key is present, a locally generated fallback explanation is returned
    so the application remains usable without live API access. Both
    providers automatically retry against a short list of known-good
    fallback model names if the requested model has been retired or is
    not available on the caller's account, so a single stale model id
    does not silently break AI analysis.
    """
    if provider == "Groq":
        text, error = call_groq(api_key, system_prompt, user_prompt, model=model)
    elif provider == "Gemini":
        text, error = call_gemini(api_key, system_prompt, user_prompt, model=model)
    else:
        text, error = None, "Unknown provider selected."
    if text:
        return text, None
    fallback = build_fallback_summary(user_prompt)
    note = "Live LLM call was not available (" + (error or "no key provided") + "). Showing a rule based summary instead."
    return fallback, note
def build_fallback_summary(user_prompt):
    """
    A simple deterministic fallback used when no API key is configured or a
    live call fails, so the demo experience is not blocked.
    """
    lines = [
        "Rule based summary generated locally, no live LLM call was made.",
        "",
        "Key graph observations were extracted from node and edge statistics",
        "including centrality, pagerank and shortest path analysis.",
        "",
        "To receive a full natural language explanation, add a valid Groq or",
        "Gemini API key in the sidebar, confirm a model is selected under",
        "Model Settings, and rerun the analysis.",
    ]
    return "\n".join(lines)
# =============================================================================
# SECTION 5: EXPORT UTILITIES - PDF, WORD, TEXT, CSV
# =============================================================================
def build_text_report(title, sections):
    """
    sections is a list of tuples (heading, body_text)
    """
    lines = []
    lines.append(title)
    lines.append("=" * len(title))
    lines.append("")
    lines.append("Prepared using the Knowledge Graph AI Use Cases application")
    lines.append("Developed by Randy Singh from Kalsnet KNet Consulting Group")
    lines.append("")
    for heading, body in sections:
        lines.append(heading)
        lines.append("-" * len(heading))
        lines.append(body)
        lines.append("")
    return "\n".join(lines).encode("utf-8")
def build_csv_report(nodes_df, edges_df, metrics_df):
    buffer = io.StringIO()
    buffer.write("NODE TABLE\n")
    if nodes_df is not None and len(nodes_df) > 0:
        nodes_df.to_csv(buffer, index=False)
    buffer.write("\nEDGE TABLE\n")
    if edges_df is not None and len(edges_df) > 0:
        edges_df.to_csv(buffer, index=False)
    buffer.write("\nGRAPH METRICS\n")
    if metrics_df is not None and len(metrics_df) > 0:
        metrics_df.to_csv(buffer, index=False)
    return buffer.getvalue().encode("utf-8")
def build_word_report(title, sections, metrics_df=None):
    doc = Document()
    heading = doc.add_heading(level=0)
    run = heading.add_run(title)
    run.font.color.rgb = RGBColor(0x1f, 0x4e, 0x9c)
    run.font.size = Pt(24)
    run.bold = True
    sub = doc.add_paragraph()
    sub_run = sub.add_run("Knowledge Graph AI Use Cases Application")
    sub_run.bold = True
    sub_run.font.color.rgb = RGBColor(0x1f, 0x4e, 0x9c)
    footer_p = doc.add_paragraph("Developed by Randy Singh from Kalsnet KNet Consulting Group")
    footer_p.runs[0].italic = True
    doc.add_paragraph("")
    for heading_text, body_text in sections:
        doc.add_heading(heading_text, level=1)
        doc.add_paragraph(body_text)
    if metrics_df is not None and len(metrics_df) > 0:
        doc.add_heading("Graph Metrics Summary", level=1)
        top = metrics_df.head(15)
        table = doc.add_table(rows=1, cols=len(top.columns))
        table.style = "Light Grid Accent 1"
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(top.columns):
            hdr_cells[i].text = str(col)
        for _, row in top.iterrows():
            cells = table.add_row().cells
            for i, col in enumerate(top.columns):
                cells[i].text = str(row[col])
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
def build_pdf_report(title, sections, metrics_df=None):
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter,
                             topMargin=0.6 * inch, bottomMargin=0.6 * inch,
                             leftMargin=0.6 * inch, rightMargin=0.6 * inch)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("TitleBlue", parent=styles["Title"],
                                  textColor=colors.HexColor("#1f4e9c"), fontSize=22)
    heading_style = ParagraphStyle("HeadingBlue", parent=styles["Heading2"],
                                    textColor=colors.HexColor("#1f4e9c"))
    body_style = styles["BodyText"]
    footer_style = ParagraphStyle("Footer", parent=styles["Normal"], fontSize=9,
                                   textColor=colors.HexColor("#555555"))
    story = []
    story.append(Paragraph(title, title_style))
    story.append(Paragraph("Knowledge Graph AI Use Cases Application", heading_style))
    story.append(Paragraph("Developed by Randy Singh from Kalsnet KNet Consulting Group", footer_style))
    story.append(Spacer(1, 16))
    for heading_text, body_text in sections:
        story.append(Paragraph(heading_text, heading_style))
        safe_body = body_text.replace("\n", "<br/>")
        story.append(Paragraph(safe_body, body_style))
        story.append(Spacer(1, 10))
    if metrics_df is not None and len(metrics_df) > 0:
        story.append(Paragraph("Graph Metrics Summary", heading_style))
        top = metrics_df.head(15)
        data = [list(top.columns)] + top.astype(str).values.tolist()
        t = Table(data, repeatRows=1)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f4e9c")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTSIZE", (0, 0), (-1, -1), 7.5),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f2f6fc")]),
        ]))
        story.append(t)
    doc.build(story)
    buf.seek(0)
    return buf.read()
# =============================================================================
# SECTION 6: STREAMLIT APPLICATION
# =============================================================================
import streamlit as st
import pandas as pd
from datetime import datetime
st.set_page_config(
    page_title="GraphTheory Defense and Cyber Intelligence Platform",
    page_icon=None,
    layout="wide",
)
# ---------------------------------------------------------------------------
# Global styling, production level look
# ---------------------------------------------------------------------------
st.markdown(
    """
    <style>
    .main-title {
        color: #0b3d91;
        font-weight: 900;
        font-size: 46px;
        margin-bottom: 0px;
        letter-spacing: 0.5px;
    }
    .sub-title {
        color: #0b3d91;
        font-weight: 700;
        font-size: 20px;
        margin-top: 2px;
    }
    .developer-line {
        color: #0b3d91;
        font-weight: 700;
        font-size: 16px;
        margin-top: 4px;
    }
    .usecase-heading {
        color: #0b3d91;
        font-weight: 800;
        font-size: 28px;
        margin-bottom: 4px;
    }
    .kpi-box {
        background-color: #f2f6fc;
        border: 1px solid #d7e0ee;
        border-left: 6px solid #0b3d91;
        border-radius: 6px;
        padding: 14px 18px;
        text-align: center;
    }
    .kpi-value {
        color: #0b3d91;
        font-weight: 800;
        font-size: 26px;
    }
    .kpi-label {
        color: #555555;
        font-size: 12px;
        text-transform: uppercase;
        letter-spacing: 0.5px;
    }
    .benefit-box {
        background-color: #f2f6fc;
        border-left: 5px solid #0b3d91;
        padding: 12px 16px;
        border-radius: 4px;
        margin-bottom: 8px;
    }
    .sidebar-title {
        color: #0b3d91;
        font-weight: 800;
        font-size: 19px;
    }
    section[data-testid="stSidebar"] {
        background-color: #f7f9fc;
        border-right: 1px solid #d7e0ee;
    }
    </style>
    """,
    unsafe_allow_html=True,
)
# ---------------------------------------------------------------------------
# Session state initialization
# ---------------------------------------------------------------------------
if "datasets" not in st.session_state:
    st.session_state["datasets"] = {}
if "analysis_results" not in st.session_state:
    st.session_state["analysis_results"] = {}
if "groq_models" not in st.session_state:
    st.session_state["groq_models"] = []
if "gemini_models" not in st.session_state:
    st.session_state["gemini_models"] = []
# ---------------------------------------------------------------------------
# Sidebar navigation
# ---------------------------------------------------------------------------
st.sidebar.markdown("<div class='sidebar-title'>GraphSentinel</div>", unsafe_allow_html=True)
st.sidebar.caption("Defense and Cyber Intelligence Platform, select a use case below")
use_case_labels = [str(USE_CASES[k]["order"]) + ". " + USE_CASES[k]["title"] for k in USE_CASE_ORDER]
label_to_key = dict(zip(use_case_labels, USE_CASE_ORDER))
selected_label = st.sidebar.radio("Use Cases", use_case_labels, index=0, label_visibility="collapsed")
selected_key = label_to_key[selected_label]
config = USE_CASES[selected_key]
st.sidebar.markdown("---")
st.sidebar.markdown("**AI Provider Settings**")
provider = st.sidebar.selectbox("Select LLM Provider", ["Groq", "Gemini"])
groq_key = st.sidebar.text_input("Groq API Key", type="password", key="groq_key_input")
gemini_key = st.sidebar.text_input("Gemini API Key", type="password", key="gemini_key_input")
active_key = groq_key if provider == "Groq" else gemini_key
# --- Model selection ---------------------------------------------------
# Provider model catalogs change over time and a pinned model id can be
# retired without notice (this is what produced the original 404
# model_not_found errors). Rather than hard code one name, let the user
# fetch the live list of models their key actually has access to, and
# fall back to a short list of known-good candidates when that is not
# possible (no key yet, or the discovery call itself fails).
st.sidebar.markdown("**Model Settings**")
if provider == "Groq":
    fetch_col, _ = st.sidebar.columns([1, 1])
    with fetch_col:
        if st.button("Refresh Groq Models", key="refresh_groq_models"):
            fetched = list_groq_models(groq_key)
            if fetched:
                st.session_state["groq_models"] = fetched
                st.sidebar.success(str(len(fetched)) + " models found")
            else:
                st.sidebar.warning("Could not fetch live model list, using default candidates.")
    model_options = st.session_state["groq_models"] or GROQ_FALLBACK_MODELS
    default_index = model_options.index(GROQ_DEFAULT_MODEL) if GROQ_DEFAULT_MODEL in model_options else 0
    selected_model = st.sidebar.selectbox("Groq Model", model_options, index=default_index, key="groq_model_select")
else:
    fetch_col, _ = st.sidebar.columns([1, 1])
    with fetch_col:
        if st.button("Refresh Gemini Models", key="refresh_gemini_models"):
            fetched = list_gemini_models(gemini_key)
            if fetched:
                st.session_state["gemini_models"] = fetched
                st.sidebar.success(str(len(fetched)) + " models found")
            else:
                st.sidebar.warning("Could not fetch live model list, using default candidates.")
    model_options = st.session_state["gemini_models"] or GEMINI_FALLBACK_MODELS
    default_index = model_options.index(GEMINI_DEFAULT_MODEL) if GEMINI_DEFAULT_MODEL in model_options else 0
    selected_model = st.sidebar.selectbox("Gemini Model", model_options, index=default_index, key="gemini_model_select")
st.sidebar.caption(
    "If a model is retired by the provider, AI Analysis automatically retries "
    "with the next known-good candidate, so a single stale model id will not "
    "block analysis."
)
with st.sidebar.expander("How to get a free Groq API key"):
    st.markdown(
        """
        1. Go to console dot groq dot com in your browser
        2. Sign up for a free account using an email address or a Google account
        3. Once logged in, open the API Keys section from the left menu
        4. Click Create API Key and give it a name
        5. Copy the generated key immediately and paste it into the Groq API Key field in this sidebar
        6. Groq offers a free developer tier with generous request limits, no credit card required to start
        """
    )
with st.sidebar.expander("How to get a free Gemini API key"):
    st.markdown(
        """
        1. Go to aistudio dot google dot com in your browser
        2. Sign in with a Google account
        3. Click Get API Key in the left menu, then click Create API Key
        4. Choose to create the key in a new or existing Google Cloud project
        5. Copy the generated key and paste it into the Gemini API Key field in this sidebar
        6. Google AI Studio provides a free usage tier suitable for development and testing
        """
    )
st.sidebar.markdown("---")
st.sidebar.markdown(
    "<div class='developer-line'>Developed by Randy Singh</div>"
    "<div class='developer-line'>Kalsnet (KNet) Consulting Group</div>",
    unsafe_allow_html=True,
)
# ---------------------------------------------------------------------------
# Header
# ---------------------------------------------------------------------------
st.markdown("<div class='main-title'>GraphTheory</div>", unsafe_allow_html=True)
st.markdown(
    "<div class='sub-title'>Defense and Cyber Intelligence Platform, "
    "Graph Theory Powered Analysis Across Thirteen Mission Critical Use Cases</div>",
    unsafe_allow_html=True,
)
st.markdown(
    "<div class='developer-line'>Developed by Randy Singh from Kalsnet (KNet) Consulting Group</div>",
    unsafe_allow_html=True,
)
st.write("")
# ---------------------------------------------------------------------------
# Use case heading and summary
# ---------------------------------------------------------------------------
st.markdown(
    "<div class='usecase-heading'>" + str(config["order"]) + ". " + config["title"] + "</div>",
    unsafe_allow_html=True,
)
st.write(config["description"])
col_a, col_b = st.columns(2)
with col_a:
    st.info("What the Graph Provides: " + config["graph_provides"])
with col_b:
    st.success("What AI Does: " + config["ai_does"])
current_dataset = st.session_state["datasets"].get(selected_key)
results_store = st.session_state["analysis_results"].get(selected_key, {})
kpi1, kpi2, kpi3, kpi4 = st.columns(4)
total_nodes = len(current_dataset["nodes"]) if current_dataset is not None else 0
total_edges = len(current_dataset["edges"]) if current_dataset is not None else 0
entity_types = len(config["node_types"])
formulas_count = len(config["formulas"])
with kpi1:
    st.markdown("<div class='kpi-box'><div class='kpi-value'>" + str(total_nodes) +
                "</div><div class='kpi-label'>Total Records Loaded</div></div>", unsafe_allow_html=True)
with kpi2:
    st.markdown("<div class='kpi-box'><div class='kpi-value'>" + str(total_edges) +
                "</div><div class='kpi-label'>Total Relationships</div></div>", unsafe_allow_html=True)
with kpi3:
    st.markdown("<div class='kpi-box'><div class='kpi-value'>" + str(entity_types) +
                "</div><div class='kpi-label'>Entity Types in Schema</div></div>", unsafe_allow_html=True)
with kpi4:
    st.markdown("<div class='kpi-box'><div class='kpi-value'>" + str(formulas_count) +
                "</div><div class='kpi-label'>Analytical Formulas Applied</div></div>", unsafe_allow_html=True)
st.markdown("---")
# ---------------------------------------------------------------------------
# Tabs within the selected use case
# ---------------------------------------------------------------------------
tab_data, tab_schema, tab_graph, tab_ai, tab_export, tab_benefits = st.tabs(
    ["Data Source", "Schema, Formulas and Flow Diagram", "Graph Visualization", "AI Analysis", "Export Results", "Benefits"]
)
# ----- Tab 1: Data Source ---------------------------------------------------
with tab_data:
    st.subheader("Choose a Data Source")
    data_mode = st.radio(
        "Data mode",
        ["Use Synthetic Data", "Upload Real Data"],
        horizontal=True,
        key="mode_" + selected_key,
    )
    if data_mode == "Use Synthetic Data":
        num_records = st.slider(
            "Number of records to generate per entity type",
            min_value=5, max_value=100, value=15, key="slider_" + selected_key,
        )
        if st.button("Generate Synthetic Data", key="gen_" + selected_key):
            dataset = generate_synthetic_dataset(selected_key, config, nodes_per_type=num_records, seed=42)
            st.session_state["datasets"][selected_key] = dataset
            st.success("Synthetic data generated for " + config["title"])
    else:
        st.write(
            "Upload two CSV files, one for nodes and one for edges. "
            "Node file should contain columns node_id, node_label, node_type and any custom fields. "
            "Edge file should contain columns source_id, target_id, relationship and optional weight."
        )
        nodes_file = st.file_uploader("Upload Nodes CSV", type=["csv"], key="nodes_upload_" + selected_key)
        edges_file = st.file_uploader("Upload Edges CSV", type=["csv"], key="edges_upload_" + selected_key)
        if nodes_file is not None and edges_file is not None:
            if st.button("Load Uploaded Data", key="load_" + selected_key):
                try:
                    dataset = dataframes_from_uploaded(nodes_file, edges_file)
                    st.session_state["datasets"][selected_key] = dataset
                    st.success("Uploaded data loaded for " + config["title"])
                except Exception as e:
                    st.error("Could not read uploaded files: " + str(e))
    current_dataset = st.session_state["datasets"].get(selected_key)
    if current_dataset is not None:
        st.markdown("### Synthetic Data Summary Bar")
        st.write(
            "This chart shows how many records are currently loaded for each entity type in the schema."
        )
        chart = build_record_count_chart(current_dataset["nodes"])
        st.plotly_chart(chart, use_container_width=True)
        st.markdown("**Node Table Preview**")
        st.dataframe(current_dataset["nodes"].head(100), use_container_width=True)
        st.markdown("**Edge Table Preview**")
        st.dataframe(current_dataset["edges"].head(100), use_container_width=True)
    else:
        st.warning("No data loaded yet for this use case. Generate synthetic data or upload real data above.")
# ----- Tab 2: Schema, Formulas and Flow Diagram ------------------------------
with tab_schema:
    st.subheader("Graph Schema Flow Diagram")
    st.write(
        "The diagram below shows how the entity types in this use case connect to one "
        "another. Each arrow represents a real relationship captured in the knowledge graph."
    )
    dot_diagram = build_schema_flow_diagram(config)
    st.graphviz_chart(dot_diagram, use_container_width=True)
    st.markdown("### Entity Fields")
    st.write("Each entity type below carries the following fields in the graph.")
    field_cols = st.columns(len(config["fields"]))
    for idx, (entity_type, field_list) in enumerate(config["fields"].items()):
        with field_cols[idx % len(field_cols)]:
            st.markdown("**" + entity_type + "**")
            for f in field_list:
                st.write("- " + f)
    st.markdown("### Formulas Used")
    st.write(
        "These are the analytical formulas applied to the graph to produce the "
        "metrics shown in the Graph Visualization and AI Analysis tabs."
    )
    for formula_name, formula_explanation in config["formulas"]:
        with st.expander(formula_name):
            st.write(formula_explanation)
# ----- Tab 3: Graph Visualization -------------------------------------------
with tab_graph:
    st.subheader("Interactive Graph Visualization")
    current_dataset = st.session_state["datasets"].get(selected_key)
    if current_dataset is None:
        st.warning("Load data in the Data Source tab first.")
    else:
        G = build_networkx_graph(current_dataset["nodes"], current_dataset["edges"])
        st.write(
            "Graph contains " + str(G.number_of_nodes()) + " nodes and "
            + str(G.number_of_edges()) + " edges."
        )
        metrics_df = compute_core_metrics(G)
        st.session_state["analysis_results"].setdefault(selected_key, {})["metrics_df"] = metrics_df
        paths = find_top_paths(G, max_paths=5)
        highlight_choice = None
        if paths:
            path_labels = []
            for p in paths:
                labels = [G.nodes[n].get("label", str(n)) for n in p]
                path_labels.append(" then ".join(labels))
            selected_path_label = st.selectbox("Highlight a representative path", ["None"] + path_labels)
            if selected_path_label != "None":
                idx = path_labels.index(selected_path_label)
                highlight_choice = paths[idx]
                st.session_state["analysis_results"][selected_key]["highlighted_path"] = highlight_choice
        fig = build_plotly_graph(G, highlight_path=highlight_choice)
        st.plotly_chart(fig, use_container_width=True)
        st.markdown("### Top Ranked Nodes by Graph Metrics")
        st.write(
            "Degree centrality measures direct connections. Betweenness centrality "
            "measures how often a node sits on the shortest path between other nodes. "
            "Pagerank measures overall structural importance considering edge weight."
        )
        st.dataframe(metrics_df.head(20), use_container_width=True)
# ----- Tab 4: AI Analysis ----------------------------------------------------
with tab_ai:
    st.subheader("AI Powered Analysis")
    current_dataset = st.session_state["datasets"].get(selected_key)
    if current_dataset is None:
        st.warning("Load data in the Data Source tab first.")
    else:
        results_store = st.session_state["analysis_results"].setdefault(selected_key, {})
        metrics_df = results_store.get("metrics_df")
        if metrics_df is None:
            G = build_networkx_graph(current_dataset["nodes"], current_dataset["edges"])
            metrics_df = compute_core_metrics(G)
            results_store["metrics_df"] = metrics_df
        st.write(
            "This step sends a summary of the graph structure and top ranked nodes "
            "to the selected AI provider, which will " + config["ai_prompt_focus"] + "."
        )
        st.caption("Using " + provider + " model: " + selected_model)
        user_question = st.text_area(
            "Optional question to focus the analysis",
            value="Summarize the most important risk or insight visible in this graph.",
            key="question_" + selected_key,
        )
        if st.button("Run AI Analysis", key="run_ai_" + selected_key):
            top_nodes = metrics_df.head(10).to_dict(orient="records")
            system_prompt = (
                "You are an expert analyst for the use case: " + config["title"] + ". "
                "You " + config["ai_does"] + ". "
                "Use the structured graph summary provided to " + config["ai_prompt_focus"] + ". "
                "Be specific, concise and use plain professional language. Do not invent facts "
                "that are not supported by the provided summary."
            )
            user_prompt = (
                "Use case schema relationships: " + str(config["edge_types"]) + "\n"
                "Top ranked nodes by graph importance: " + str(top_nodes) + "\n"
                "Total nodes: " + str(len(current_dataset["nodes"])) + "\n"
                "Total edges: " + str(len(current_dataset["edges"])) + "\n"
                "User question: " + user_question
            )
            with st.spinner("Contacting " + provider + " for analysis..."):
                result_text, note = run_llm_analysis(
                    provider, active_key, system_prompt, user_prompt, model=selected_model
                )
            results_store["ai_result"] = result_text
            results_store["ai_note"] = note
            results_store["ai_question"] = user_question
        if "ai_result" in results_store:
            if results_store.get("ai_note"):
                st.warning(results_store["ai_note"])
            st.markdown("### AI Analysis Result")
            st.write(results_store["ai_result"])
# ----- Tab 5: Export Results --------------------------------------------------
with tab_export:
    st.subheader("Export Results")
    current_dataset = st.session_state["datasets"].get(selected_key)
    results_store = st.session_state["analysis_results"].get(selected_key, {})
    if current_dataset is None:
        st.warning("Load data in the Data Source tab first.")
    else:
        nodes_df = current_dataset["nodes"]
        edges_df = current_dataset["edges"]
        metrics_df = results_store.get("metrics_df")
        ai_result = results_store.get("ai_result", "AI analysis has not been run yet for this use case.")
        report_title = "GraphSentinel Report: " + config["title"]
        generated_on = "Generated on " + datetime.now().strftime("%Y-%m-%d %H:%M")
        sections = [
            ("Use Case Overview", config["description"]),
            ("What the Graph Provides", config["graph_provides"]),
            ("What AI Does", config["ai_does"]),
            ("Data Summary", "Total nodes: " + str(len(nodes_df)) + ". Total edges: " + str(len(edges_df)) + ". " + generated_on),
            ("AI Analysis Result", ai_result),
            ("Benefits", "\n".join(["- " + b for b in config["benefits"]])),
        ]
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            pdf_bytes = build_pdf_report(report_title, sections, metrics_df)
            st.download_button(
                "Download PDF", data=pdf_bytes,
                file_name=selected_key + "_report.pdf", mime="application/pdf",
            )
        with col2:
            word_bytes = build_word_report(report_title, sections, metrics_df)
            st.download_button(
                "Download Word", data=word_bytes,
                file_name=selected_key + "_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        with col3:
            text_bytes = build_text_report(report_title, sections)
            st.download_button(
                "Download Text", data=text_bytes,
                file_name=selected_key + "_report.txt", mime="text/plain",
            )
        with col4:
            csv_bytes = build_csv_report(nodes_df, edges_df, metrics_df)
            st.download_button(
                "Download CSV", data=csv_bytes,
                file_name=selected_key + "_data.csv", mime="text/csv",
            )
        st.markdown("### Report Preview")
        for heading, body in sections:
            st.markdown("**" + heading + "**")
            st.write(body)
# ----- Tab 6: Benefits -------------------------------------------------------
with tab_benefits:
    st.subheader("Benefits of This Use Case")
    for benefit in config["benefits"]:
        st.markdown("<div class='benefit-box'>" + benefit + "</div>", unsafe_allow_html=True)
# ---------------------------------------------------------------------------
# Footer
# ---------------------------------------------------------------------------
st.markdown("---")
st.markdown(
    "<div class='developer-line'>GraphSentinel Defense and Cyber Intelligence Platform</div>"
    "<div class='developer-line'>Developed by Randy Singh from Kalsnet (KNet) Consulting Group</div>",
    unsafe_allow_html=True,
)
