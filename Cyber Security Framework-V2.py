#=============================================================================
# CYBERSECURITY FRAMEWORK APPLICATION FOR DoD CRITICAL INFRASTRUCTURE
# @Developed by Randy Singh | KalSnet (KNet) Consulting Group
# Original Java code: September 2017 — Python/Streamlit conversion: 2024
#=============================================================================
# Run with:
#   pip install streamlit plotly pandas python-docx
#   (optional) pip install fpdf2
#   streamlit run cybersecurity_framework_v2.py
#=============================================================================

import streamlit as st
import pandas as pd
import plotly.graph_objects as go
import json
import io
from datetime import datetime

# ─── Optional export dependencies ─────────────────────────────────────────────
FPDF_AVAILABLE = False
try:
    from fpdf import FPDF
    FPDF_AVAILABLE = True
except ImportError:
    FPDF_AVAILABLE = False

DOCX_AVAILABLE = False
try:
    from docx import Document as DocxDocument
    from docx.shared import RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ══════════════════════════════════════════════════════════════════════════════
#  PAGE CONFIG
# ══════════════════════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="CyberSecurity Framework | KNet Consulting",
    page_icon="🛡️",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ══════════════════════════════════════════════════════════════════════════════
#  GLOBAL STYLES
# ══════════════════════════════════════════════════════════════════════════════
st.markdown(
    """
<style>
  body {
      background-color:#f4f7fc;
      font-family:"Segoe UI",system-ui,-apple-system,BlinkMacSystemFont,sans-serif;
  }
  .cs-title {
      text-align:center; font-size:2.6rem; font-weight:900;
      color:#0047AB; letter-spacing:1px; margin-bottom:0;
  }
  .cs-subtitle {
      text-align:center; font-size:1.1rem; font-weight:700;
      color:#0047AB; margin-top:2px; margin-bottom:2px;
  }
  .cs-tagline {
      text-align:center; color:#0047AB; font-size:0.9rem; margin-bottom:16px;
      font-weight:700;
  }
  .cs-hr { border:2px solid #0047AB; margin:8px 0 18px 0; }
  .section-header {
      background:linear-gradient(90deg,#0047AB 0%,#1a6fe8 100%);
      color:white; padding:10px 18px; border-radius:8px;
      font-weight:800; font-size:1.1rem; margin:12px 0 10px 0;
      box-shadow:0 2px 6px rgba(0,0,0,0.18);
  }
  .req-card {
      background:#f0f6ff; border-left:4px solid #0047AB;
      border-radius:6px; padding:9px 13px; margin:5px 0; font-size:0.92rem;
  }
  .rec-card {
      background:#fff8e1; border-left:4px solid #f0a500;
      border-radius:6px; padding:9px 13px; margin:5px 0; font-size:0.92rem;
  }
  .sol-card {
      background:#e8f5e9; border-left:4px solid #2e7d32;
      border-radius:6px; padding:9px 13px; margin:5px 0; font-size:0.92rem;
  }
  .ref-card {
      background:#fce4ec; border-left:4px solid #c62828;
      border-radius:6px; padding:9px 13px; margin:5px 0; font-size:0.85rem; font-family:monospace;
  }
  section[data-testid="stSidebar"] {
      background:#001a6b;
      overflow-y:auto;
  }
  section[data-testid="stSidebar"] * { color:#e8f0ff !important; }
  .sidebar-box {
      background:rgba(255,255,255,0.06);
      padding:10px 12px;
      border-radius:8px;
      margin-bottom:12px;
      border:1px solid rgba(255,255,255,0.12);
  }
  .sidebar-footer {
      font-size:0.8rem;
      color:#cbd4ff;
      text-align:center;
      margin-top:10px;
  }
</style>
""",
    unsafe_allow_html=True,
)

# ══════════════════════════════════════════════════════════════════════════════
#  HEADER
# ══════════════════════════════════════════════════════════════════════════════
st.markdown(
    '<p class="cs-title">CYBERSECURITY FRAMEWORK APPLICATION</p>',
    unsafe_allow_html=True,
)
st.markdown(
    '<p class="cs-subtitle">FOR US CRITICAL INFRASTRUCTURE</p>',
    unsafe_allow_html=True,
)
st.markdown(
    '<p class="cs-subtitle">Developed by Randy Singh | KalSnet (KNet) Consulting Group</p>',
    unsafe_allow_html=True,
)
st.markdown(
    '<p class="cs-tagline"><b>Helping Security Analysts, Planners & Administrators address IoT Security Vulnerabilities in DISA, DoD & USA Critical Infrastructure &nbsp;|&nbsp; Originally authored June 2026</b></p>',
    unsafe_allow_html=True,
)
st.markdown('<hr class="cs-hr">', unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
#  FRAMEWORK DATA (12 items; extend to 21 if desired)
# ══════════════════════════════════════════════════════════════════════════════
FRAMEWORK_ITEMS = {
    1: {
        "title": "Identify - Asset Management",
        "function": "IDENTIFY",
        "summary": "Data, personnel, devices, systems, and facilities are identified and managed consistent with business objectives and risk strategy.",
        "subcategories": [
            "Physical devices and systems are inventoried.",
            "Software platforms and applications are inventoried.",
            "Organizational communication and data flows are mapped.",
            "External information systems are catalogued.",
            "Resources are prioritized based on classification and criticality.",
            "Cybersecurity roles and responsibilities are established.",
        ],
        "example": "Use a CMDB like ServiceNow to auto-discover and track all hardware/software assets.",
        "solution": "Deploy automated asset discovery integrated with SIEM and reconcile quarterly.",
        "recommendation": "Require all new devices to be registered before network access is granted.",
        "references": [
            "NIST SP 800-53 CM-8, CCS CSC 1.",
        ],
    },
    2: {
        "title": "Identify - Business Environment",
        "function": "IDENTIFY",
        "summary": "Mission, objectives, stakeholders, and supply chain role are understood and prioritized.",
        "subcategories": [
            "Role in the supply chain is identified and communicated.",
            "Place in critical infrastructure sector is identified.",
            "Mission priorities are established and communicated.",
            "Dependencies and critical functions are established.",
            "Resilience requirements are established.",
        ],
        "example": "Map the agency's position in the DoD supply chain using a Business Impact Analysis.",
        "solution": "Conduct annual Business Environment Review using CISA guidance.",
        "recommendation": "Align mission priorities to CISA's 16 Critical Infrastructure Sectors.",
        "references": [
            "NIST SP 800-53 PM-8, COBIT APO02.",
        ],
    },
    3: {
        "title": "Identify - Governance Policies & Procedures",
        "function": "IDENTIFY",
        "summary": "Policies, procedures, and processes for regulatory, legal, and risk requirements inform cybersecurity risk management.",
        "subcategories": [
            "Information security policy is established.",
            "Roles and responsibilities are coordinated.",
            "Legal and regulatory requirements are understood and managed.",
            "Governance and risk management processes address cybersecurity risks.",
        ],
        "example": "Publish a DoD-aligned Information Security Policy and review annually.",
        "solution": "Establish a Cybersecurity Governance Board and use GRC tools.",
        "recommendation": "Align policies to FISMA, DFARS, and CMMC; perform gap assessments.",
        "references": [
            "NIST SP 800-53 PM-1, ISO/IEC 27001 A.5.1.",
        ],
    },
    4: {
        "title": "Identify - Risk Assessment",
        "function": "IDENTIFY",
        "summary": "Cybersecurity risk to operations, assets, and individuals is understood.",
        "subcategories": [
            "Asset vulnerabilities are identified and documented.",
            "Threat and vulnerability information is received from sharing forums.",
            "Threats are identified and documented.",
            "Business impacts and likelihoods are identified.",
            "Risk is determined from threats, vulnerabilities, likelihoods, and impacts.",
            "Risk responses are identified and prioritized.",
        ],
        "example": "Run quarterly vulnerability scans and produce risk-ranked remediation plans.",
        "solution": "Implement a Risk Register using NIST SP 800-30 methodology.",
        "recommendation": "Perform annual Risk Assessments per OMB A-130.",
        "references": [
            "NIST SP 800-53 RA-3/RA-5.",
        ],
    },
    5: {
        "title": "Identify - Risk Management Strategy",
        "function": "IDENTIFY",
        "summary": "Risk tolerances and assumptions are established and used to support operational risk decisions.",
        "subcategories": [
            "Risk management processes are established and agreed.",
            "Risk tolerance is determined and expressed.",
            "Risk tolerance is informed by critical infrastructure role.",
            "Resilience requirements are established.",
        ],
        "example": "Define a Risk Tolerance Statement tied to ATO renewal criteria.",
        "solution": "Develop an Enterprise Risk Management framework.",
        "recommendation": "Review risk tolerance annually with senior leadership.",
        "references": [
            "NIST SP 800-39, COBIT APO12.",
        ],
    },
    6: {
        "title": "Protect - Access Control",
        "function": "PROTECT",
        "summary": "Access to assets and facilities is limited to authorized users, processes, and devices.",
        "subcategories": [
            "Identities and credentials are managed.",
            "Physical access is managed and protected.",
            "Remote access is managed.",
            "Permissions follow least privilege and separation of duties.",
            "Network integrity is protected with segregation.",
        ],
        "example": "Implement CAC/PIV authentication and PAM for privileged accounts.",
        "solution": "Deploy Zero Trust Network Access with MFA.",
        "recommendation": "Review access rights quarterly and log privileged access.",
        "references": [
            "NIST SP 800-53 AC-2/AC-3.",
        ],
    },
    7: {
        "title": "Protect - Awareness & Training",
        "function": "PROTECT",
        "summary": "Personnel and partners receive cybersecurity awareness education and training.",
        "subcategories": [
            "All users are informed and trained.",
            "Privileged users understand responsibilities.",
            "Third-party stakeholders understand responsibilities.",
            "Senior executives understand responsibilities.",
            "Security personnel understand responsibilities.",
        ],
        "example": "Mandate annual Cyber Awareness training and phishing simulations.",
        "solution": "Implement a Security Awareness Training platform.",
        "recommendation": "Require 100% annual training completion before account renewal.",
        "references": [
            "NIST SP 800-53 AT-2/AT-3.",
        ],
    },
    8: {
        "title": "Protect - Data Security",
        "function": "PROTECT",
        "summary": "Information and records are managed to protect confidentiality, integrity, and availability.",
        "subcategories": [
            "Data-at-rest is protected.",
            "Data-in-transit is protected.",
            "Assets are managed through removal and disposition.",
            "Capacity for availability is maintained.",
            "Protections against data leaks are implemented.",
            "Integrity checking mechanisms are used.",
            "Dev/test environments are separate from production.",
        ],
        "example": "Encrypt data-at-rest and use TLS 1.3 for data-in-transit.",
        "solution": "Deploy Data Classification and DLP solutions.",
        "recommendation": "Classify all data per DoD guidance and encrypt CUI.",
        "references": [
            "NIST SP 800-53 SC-28/SC-8.",
        ],
    },
    9: {
        "title": "Protect - Information Protection Processes & Procedures",
        "function": "PROTECT",
        "summary": "Security policies, processes, and procedures manage protection of systems and assets.",
        "subcategories": [
            "Baseline configuration is created and maintained.",
            "SDLC is implemented.",
            "Configuration change control is in place.",
            "Backups are conducted and tested.",
            "Physical environment policies are met.",
            "Data is destroyed according to policy.",
            "Protection processes are continuously improved.",
            "Effectiveness of technologies is shared.",
            "Response and recovery plans are in place.",
            "Response and recovery plans are tested.",
            "Cybersecurity is included in HR practices.",
            "Vulnerability management plan is implemented.",
        ],
        "example": "Maintain STIG baselines and test backup restoration regularly.",
        "solution": "Implement configuration management and backup tools.",
        "recommendation": "Apply STIGs before deployment and enforce change management.",
        "references": [
            "NIST SP 800-53 CM-2/CP-9.",
        ],
    },
    10: {
        "title": "Protect - Maintenance of ICS & Information Systems",
        "function": "PROTECT",
        "summary": "Maintenance and repairs are performed consistent with policies and procedures.",
        "subcategories": [
            "Maintenance is performed and logged with approved tools.",
            "Remote maintenance is approved, logged, and secured.",
        ],
        "example": "Use jump servers for remote maintenance and record sessions.",
        "solution": "Require dual authorization for critical maintenance.",
        "recommendation": "Review maintenance logs monthly for anomalies.",
        "references": [
            "NIST SP 800-53 MA-2/MA-4.",
        ],
    },
    11: {
        "title": "Protect - Protective Technology",
        "function": "PROTECT",
        "summary": "Technical security solutions are managed to ensure security and resilience.",
        "subcategories": [
            "Audit/log records are documented and reviewed.",
            "Removable media is protected and restricted.",
            "Access follows least functionality.",
            "Communications and control networks are protected.",
        ],
        "example": "Deploy SIEM, disable USB ports, and use application whitelisting.",
        "solution": "Implement NGFW with IPS and network segmentation.",
        "recommendation": "Audit logs daily and review firewall rules quarterly.",
        "references": [
            "NIST SP 800-53 AU family/SC-7.",
        ],
    },
    12: {
        "title": "Detect - Anomalies & Events",
        "function": "DETECT",
        "summary": "Anomalous activity is detected and impact is understood.",
        "subcategories": [
            "Baseline of network operations is established.",
            "Events are analyzed to understand targets and methods.",
            "Event data are aggregated and correlated.",
            "Impact of events is determined.",
            "Incident alert thresholds are established.",
        ],
        "example": "Use network flow analysis tools to establish behavioral baselines.",
        "solution": "Deploy UEBA integrated with threat intel feeds.",
        "recommendation": "Map detection rules to MITRE ATT&CK and hunt quarterly.",
        "references": [
            "NIST SP 800-53 SI-4.",
        ],
    },
}

# ══════════════════════════════════════════════════════════════════════════════
#  EXPORT HELPERS
# ══════════════════════════════════════════════════════════════════════════════
def build_text_report(item_id: int, item: dict) -> str:
    lines = [
        "=" * 80,
        "CYBERSECURITY FRAMEWORK APPLICATION",
        "Developed by Randy Singh | KalSnet (KNet) Consulting Group",
        f"Checklist Item: {item_id} - {item['title']}",
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "=" * 80,
        "",
        f"Function: {item['function']}",
        "",
        "Summary:",
        f"  {item['summary']}",
        "",
        "Subcategories:",
    ]
    for sc in item["subcategories"]:
        lines.append(f"  - {sc}")
    lines.append("")
    lines.append("Example:")
    lines.append(f"  {item['example']}")
    lines.append("")
    lines.append("Solution:")
    lines.append(f"  {item['solution']}")
    lines.append("")
    lines.append("Recommendation:")
    lines.append(f"  {item['recommendation']}")
    lines.append("")
    lines.append("References:")
    for ref in item["references"]:
        lines.append(f"  {ref}")
    lines.append("")
    return "\n".join(lines)


def build_json_report(item_id: int, item: dict) -> str:
    payload = {
        "framework": "CyberSecurity Framework Application",
        "author": "Randy Singh | KalSnet (KNet) Consulting Group",
        "checklist_item_id": item_id,
        "title": item["title"],
        "function": item["function"],
        "generated": datetime.now().isoformat(),
        "data": item,
    }
    return json.dumps(payload, indent=2, default=str)


def _pdf_safe(text: str) -> str:
    replacements = [
        ("\u2013", "-"),
        ("\u2014", "--"),
        ("\u2018", "'"),
        ("\u2019", "'"),
        ("\u201c", '"'),
        ("\u201d", '"'),
        ("\u2026", "..."),
        ("\u2022", "-"),
        ("\u00a0", " "),
    ]
    for char, replacement in replacements:
        text = text.replace(char, replacement)
    return text.encode("latin-1", errors="ignore").decode("latin-1")


def build_pdf_bytes(item_id: int, item: dict) -> bytes:
    text_report = build_text_report(item_id, item)

    if FPDF_AVAILABLE:
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 14)
        pdf.set_text_color(0, 71, 171)
        pdf.multi_cell(0, 8, _pdf_safe("CYBERSECURITY FRAMEWORK APPLICATION"), align="C")

        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(30, 30, 30)
        pdf.ln(4)
        pdf.multi_cell(0, 5, _pdf_safe(text_report))

        raw = pdf.output()
        if isinstance(raw, (bytes, bytearray)):
            return bytes(raw)
        return raw.encode("latin-1")

    # Simple text-only PDF fallback
    content = _pdf_safe(text_report)
    header = "%PDF-1.4\n"
    body = f"1 0 obj<<>>endobj\n2 0 obj<< /Length {len(content)} >>\nstream\n{content}\nendstream\nendobj\n"
    trailer = "xref\n0 3\n0000000000 65535 f \ntrailer<< /Size 3 >>\nstartxref\n0\n%%EOF"
    return (header + body + trailer).encode("latin-1", errors="ignore")


def build_docx_bytes(item_id: int, item: dict) -> bytes:
    if not DOCX_AVAILABLE:
        return b""
    doc = DocxDocument()

    t = doc.add_heading("CYBERSECURITY FRAMEWORK APPLICATION", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in t.runs:
        run.font.color.rgb = RGBColor(0, 71, 171)

    sub = doc.add_paragraph(
        "Developed by Randy Singh | KalSnet (KNet) Consulting Group"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if sub.runs:
        sub.runs[0].bold = True
        sub.runs[0].font.color.rgb = RGBColor(0, 71, 171)

    ts_para = doc.add_paragraph(
        f"Checklist Item: {item_id} - {item['title']}   |   {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    ts_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    h = doc.add_heading("Summary", level=2)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 71, 171)
    doc.add_paragraph(item["summary"])

    h = doc.add_heading("Subcategories", level=2)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 71, 171)
    for sc in item["subcategories"]:
        doc.add_paragraph(sc, style="List Bullet")

    h = doc.add_heading("Example", level=2)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 71, 171)
    doc.add_paragraph(item["example"])

    h = doc.add_heading("Solution", level=2)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 71, 171)
    doc.add_paragraph(item["solution"])

    h = doc.add_heading("Recommendation", level=2)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 71, 171)
    doc.add_paragraph(item["recommendation"])

    h = doc.add_heading("References", level=2)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 71, 171)
    for ref in item["references"]:
        doc.add_paragraph(ref, style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# ══════════════════════════════════════════════════════════════════════════════
#  SIDEBAR — NAVIGATION & EXPORT (RADIO BUTTONS)
# ══════════════════════════════════════════════════════════════════════════════
with st.sidebar:
    st.markdown("## Browse Checklist Items")
    st.markdown("---")

    item_labels = [
        f"{i}. {FRAMEWORK_ITEMS[i]['title']}"
        for i in sorted(FRAMEWORK_ITEMS.keys())
    ]
    selected_label = st.radio(
        "Select checklist item",
        item_labels,
        index=0,
        key="item_radio",
    )

    selected_id = int(selected_label.split(".")[0])
    selected_item = FRAMEWORK_ITEMS[selected_id]

    st.markdown('<div class="sidebar-box">', unsafe_allow_html=True)
    st.markdown("**About**")
    st.markdown(
        "Randy Singh  \nComputer Scientist  \nKalSnet (KNet) Consulting Group"
    )
    st.markdown("`(301) 225-9535`")
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown('<div class="sidebar-box">', unsafe_allow_html=True)
    st.markdown("**Export This Checklist Item**")

    export_choice = st.radio(
        "Choose export format",
        ["Text (.txt)", "JSON (.json)", "PDF (.pdf)", "Word (.docx)"],
        index=0,
        key="export_radio",
    )

    text_report = build_text_report(selected_id, selected_item)
    json_report = build_json_report(selected_id, selected_item)
    pdf_bytes = build_pdf_bytes(selected_id, selected_item)
    docx_bytes = build_docx_bytes(selected_id, selected_item)

    if export_choice == "Text (.txt)":
        st.download_button(
            "Download Text",
            data=text_report,
            file_name=f"cyber_framework_item_{selected_id}.txt",
            mime="text/plain",
            key="dl_txt",
        )
    elif export_choice == "JSON (.json)":
        st.download_button(
            "Download JSON",
            data=json_report,
            file_name=f"cyber_framework_item_{selected_id}.json",
            mime="application/json",
            key="dl_json",
        )
    elif export_choice == "PDF (.pdf)":
        st.download_button(
            "Download PDF",
            data=pdf_bytes,
            file_name=f"cyber_framework_item_{selected_id}.pdf",
            mime="application/pdf",
            key="dl_pdf",
        )
    elif export_choice == "Word (.docx)":
        st.download_button(
            "Download Word",
            data=docx_bytes,
            file_name=f"cyber_framework_item_{selected_id}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="dl_docx",
        )

    st.markdown(
        '<div class="sidebar-footer">CyberSecurity Framework | KalSnet (KNet) Consulting</div>',
        unsafe_allow_html=True,
    )

# ══════════════════════════════════════════════════════════════════════════════
#  MAIN CONTENT — SELECTED CHECKLIST ITEM + MERMAID & FLOW DIAGRAMS
# ══════════════════════════════════════════════════════════════════════════════
st.markdown(
    '<div class="section-header">Browse CyberSecurity Framework Checklist Items</div>',
    unsafe_allow_html=True,
)

st.markdown(f"### {selected_item['title']}")
st.markdown(f"**Function:** {selected_item['function']}")

st.markdown(
    f'<div class="req-card"><b>Summary:</b> {selected_item["summary"]}</div>',
    unsafe_allow_html=True,
)

st.markdown("#### Subcategories")
for sc in selected_item["subcategories"]:
    st.markdown(f"- {sc}")

st.markdown(
    f'<div class="sol-card"><b>Solution Pattern:</b> {selected_item["solution"]}</div>',
    unsafe_allow_html=True,
)
st.markdown(
    f'<div class="rec-card"><b>Recommendation:</b> {selected_item["recommendation"]}</div>',
    unsafe_allow_html=True,
)
st.markdown(
    '<div class="ref-card"><b>References:</b><br>'
    + "<br>".join(selected_item["references"])
    + "</div>",
    unsafe_allow_html=True,
)

# ─── Mermaid diagram per function (no f-string) ───────────────────────────────
st.markdown("#### Mermaid Flow Diagram (Function View)")
mermaid_diagram = """
```mermaid
flowchart LR
    A[Start] --> B[FUNCTION_NAME Function]
    B --> C[Policy & Governance]
    C --> D[Controls Implemented]
    D --> E[Monitoring & Detection]
    E --> F[Response & Improvement]
